#!/usr/bin/env python3
"""Local Web API smoke test for the report-generation platform.

The test uses a synthetic CRC 358 + MSI workbook, so it is safe to run before
release or on a staging server without exposing patient data.
"""

from __future__ import annotations

import json
import os
import sys
import time
import urllib.error
import urllib.parse
import urllib.request
from pathlib import Path
from typing import Any


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402
from reportgen.core.golden_case import build_crc_358_msi_golden_excel  # noqa: E402


BASE_URL = os.environ.get("WEB_SMOKE_BASE_URL", "http://127.0.0.1:8000").rstrip("/")
OUTPUT_ROOT = Path(
    os.environ.get("WEB_SMOKE_OUTPUT_ROOT", str(ROOT / "tmp" / "web_smoke"))
)
ADMIN_USERNAME = os.environ.get("WEB_SMOKE_ADMIN_USERNAME", "admin")
ADMIN_PASSWORD = os.environ.get("WEB_SMOKE_ADMIN_PASSWORD", "admin123")
TIMEOUT_SECONDS = int(os.environ.get("WEB_SMOKE_TIMEOUT_SECONDS", "240"))
MIN_DOCX_BYTES = int(os.environ.get("WEB_SMOKE_MIN_DOCX_BYTES", "1000000"))
USER_AGENT = os.environ.get(
    "WEB_SMOKE_USER_AGENT",
    "Mozilla/5.0 ReportGenSmoke/1.0",
)


class SmokeFailure(RuntimeError):
    """Raised when a smoke-test assertion fails."""


def _url(path: str) -> str:
    return f"{BASE_URL}{path}"


def _request(
    method: str,
    path: str,
    *,
    body: bytes | None = None,
    headers: dict[str, str] | None = None,
    timeout: int = 120,
) -> tuple[int, bytes, dict[str, str]]:
    request_headers = {"User-Agent": USER_AGENT, **(headers or {})}
    req = urllib.request.Request(
        _url(path),
        data=body,
        headers=request_headers,
        method=method,
    )
    try:
        with urllib.request.urlopen(req, timeout=timeout) as response:
            return (
                int(response.status),
                response.read(),
                dict(response.headers.items()),
            )
    except urllib.error.HTTPError as exc:
        return int(exc.code), exc.read(), dict(exc.headers.items())


def _json_request(
    method: str,
    path: str,
    *,
    payload: dict[str, Any] | None = None,
    timeout: int = 120,
) -> dict[str, Any]:
    headers = {"Accept": "application/json"}
    body = None
    if payload is not None:
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        headers["Content-Type"] = "application/json"
    status, raw, _headers = _request(
        method,
        path,
        body=body,
        headers=headers,
        timeout=timeout,
    )
    if status < 200 or status >= 300:
        raise SmokeFailure(f"{method} {path} returned HTTP {status}: {raw[:500]!r}")
    try:
        data = json.loads(raw.decode("utf-8"))
    except Exception as exc:
        raise SmokeFailure(f"{method} {path} did not return JSON") from exc
    if data.get("success") is False:
        raise SmokeFailure(f"{method} {path} failed: {data.get('error') or data}")
    return data


def _upload_excel(path: Path) -> dict[str, Any]:
    return _multipart_file_request(
        "/api/v1/excel/upload",
        path,
        fields={},
        timeout=120,
    )


def _multipart_file_request(
    endpoint: str,
    path: Path,
    *,
    fields: dict[str, str],
    timeout: int,
) -> dict[str, Any]:
    boundary = f"----reportgen-smoke-{int(time.time() * 1000)}"
    file_bytes = path.read_bytes()
    parts = [
        f"--{boundary}\r\n".encode("utf-8"),
        (
            'Content-Disposition: form-data; name="file"; '
            f'filename="{path.name}"\r\n'
        ).encode("utf-8"),
        b"Content-Type: application/vnd.openxmlformats-officedocument.spreadsheetml.sheet\r\n\r\n",
        file_bytes,
        b"\r\n",
    ]
    for name, value in fields.items():
        parts.extend(
            [
                f"--{boundary}\r\n".encode("utf-8"),
                f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode("utf-8"),
                value.encode("utf-8"),
                b"\r\n",
            ]
        )
    parts.append(f"--{boundary}--\r\n".encode("utf-8"))
    status, raw, _headers = _request(
        "POST",
        endpoint,
        body=b"".join(parts),
        headers={
            "Accept": "application/json",
            "Content-Type": f"multipart/form-data; boundary={boundary}",
        },
        timeout=timeout,
    )
    if status < 200 or status >= 300:
        raise SmokeFailure(f"{endpoint} returned HTTP {status}: {raw[:500]!r}")
    data = json.loads(raw.decode("utf-8"))
    if not data.get("success"):
        raise SmokeFailure(f"{endpoint} failed: {data}")
    return data


def _generate_file_async(path: Path, clinical_info: dict[str, Any]) -> str:
    payload = _multipart_file_request(
        "/api/v1/reports/generate-file-async",
        path,
        fields={
            "clinical_info": json.dumps(clinical_info, ensure_ascii=False),
            "project_type": "crc_358_msi",
            "project_name": "结直肠癌358基因+MSI",
            "strict_mode": "false",
            "template_contract_mode": "warn",
        },
        timeout=120,
    )
    task_id = (payload.get("data") or {}).get("task_id")
    if not task_id:
        raise SmokeFailure(f"async generation did not return task_id: {payload}")
    return str(task_id)


def _wait_for_task(task_id: str) -> dict[str, Any]:
    deadline = time.monotonic() + TIMEOUT_SECONDS
    last_status = "-"
    while time.monotonic() < deadline:
        payload = _json_request("GET", f"/api/v1/reports/{task_id}", timeout=30)
        data = payload["data"]
        last_status = str(data.get("status") or "-")
        if last_status in {"completed", "failed", "partial_failed", "cancelled"}:
            return data
        time.sleep(3)
    raise SmokeFailure(f"task {task_id} did not finish before timeout; last_status={last_status}")


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise SmokeFailure(message)


def _check_docx_text(path: Path) -> None:
    doc = Document(str(path))
    chunks: list[str] = [p.text for p in doc.paragraphs]
    chunks.extend(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    text = "\n".join(chunks)
    for needle in [
        "本次共检出体细胞变异：2个",
        "与靶向药物用药相关的变异有：1个",
        "微卫星稳定型，MSS",
        "多项临床研究表明，TMB-H的肿瘤",
        "研究表明，MSI-H的实体瘤",
        "ERBB2",
        "c.1979G>A",
        "p.G660D",
    ]:
        _assert(needle in text, f"generated DOCX is missing text: {needle}")


def main() -> int:
    started = time.monotonic()
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    print("Web smoke test")
    print(f"  base_url: {BASE_URL}")
    print(f"  output_root: {OUTPUT_ROOT}")

    status, body, _headers = _request("GET", "/", timeout=30)
    _assert(status == 200, f"frontend index returned HTTP {status}")
    _assert(
        "基因组Panel自动化报告系统" in body.decode("utf-8", errors="ignore"),
        "frontend index did not contain app title",
    )
    print("  ✅ frontend index")

    login = _json_request(
        "POST",
        "/api/v1/auth/login",
        payload={"username": ADMIN_USERNAME, "password": ADMIN_PASSWORD},
        timeout=30,
    )
    _assert(
        bool((login.get("data") or {}).get("access_token")),
        "admin login did not return a token",
    )
    print("  ✅ auth login")

    stats = _json_request("GET", "/api/v1/tasks/stats", timeout=30)
    _assert(isinstance(stats.get("data"), dict), "task stats payload is invalid")
    print("  ✅ task stats")

    ops = _json_request(
        "GET",
        "/api/v1/admin/ops/status?recent_task_limit=1&download_event_limit=1",
        timeout=30,
    )
    ops_raw = json.dumps(ops, ensure_ascii=False)
    _assert(bool((ops.get("data") or {}).get("deployment")), "ops status is missing deployment")
    for forbidden in ["/media/desk16", "/Volumes/KINGSTON", "user_agent", "client_host"]:
        _assert(forbidden not in ops_raw, f"ops status leaked {forbidden!r}")
    print("  ✅ ops status")

    excel_path = build_crc_358_msi_golden_excel(
        OUTPUT_ROOT / "LZ999001_crc_358_msi_golden.xlsx"
    )
    upload = _upload_excel(excel_path)
    upload_data = upload["data"]
    upload_id = upload_data["upload_id"]
    _assert(
        upload_data.get("detected_project_type") == "crc_358_msi",
        f"expected crc_358_msi, got {upload_data.get('detected_project_type')!r}",
    )
    print(f"  ✅ upload + detect ({upload_id})")

    sheets = _json_request("GET", f"/api/v1/excel/{upload_id}/sheets", timeout=30)
    sheet_map = {item["name"]: item for item in sheets["data"]}
    _assert(sheet_map.get("Meta", {}).get("rows") == 1, "Meta sheet row count is wrong")
    _assert(
        sheet_map.get("Meta", {}).get("columns", 0) >= 10,
        "Meta sheet column count is wrong",
    )
    _assert(
        sheet_map.get("Variations", {}).get("rows") == 2,
        "Variations sheet row count is wrong",
    )
    _assert(
        sheet_map.get("Variations", {}).get("columns", 0) >= 10,
        "Variations sheet column count is wrong",
    )
    print("  ✅ sheet list counts")

    meta = _json_request(
        "GET",
        f"/api/v1/excel/{upload_id}/sheets/{urllib.parse.quote('Meta')}?page=1&page_size=5",
        timeout=30,
    )
    _assert(meta["data"]["rows"][0]["样本编号"] == "LZ999001", "Meta preview is wrong")
    variations = _json_request(
        "GET",
        f"/api/v1/excel/{upload_id}/sheets/Variations?page=1&page_size=5",
        timeout=30,
    )
    _assert(
        variations["data"]["rows"][0]["Gene_Symbol"] == "ERBB2",
        "Variations preview is wrong",
    )
    print("  ✅ sheet previews")

    schema = _json_request(
        "GET",
        "/api/v1/clinical-schema?project_type=crc_358_msi",
        timeout=30,
    )
    _assert(len(schema["data"].get("groups") or []) > 0, "clinical schema is empty")
    single_values = _json_request(
        "GET",
        f"/api/v1/excel/{upload_id}/single-values",
        timeout=30,
    )
    clinical_info = single_values["data"]["fields"]
    _assert(clinical_info.get("sample_id") == "LZ999001", "sample_id extraction failed")
    _assert(clinical_info.get("patient_name") == "黄金测试患者", "patient extraction failed")
    print("  ✅ schema + single values")

    task_id = _generate_file_async(excel_path, clinical_info)
    report_data = _wait_for_task(task_id)
    _assert(report_data["status"] == "completed", f"task status is {report_data['status']!r}")
    _assert(report_data["qa_status"] == "PASS", f"QA status is {report_data['qa_status']!r}")
    _assert(len(report_data.get("stage_results") or []) > 0, "stage results are missing")
    print(f"  ✅ report generation ({task_id})")

    qa = _json_request("GET", f"/api/v1/reports/{task_id}/qa", timeout=30)
    _assert(qa["data"].get("status") == "PASS", "QA report endpoint did not return PASS")
    print("  ✅ QA endpoint")

    status, raw_docx, _headers = _request(
        "GET",
        f"/api/v1/reports/{task_id}/download",
        timeout=120,
    )
    _assert(status == 200, f"download returned HTTP {status}")
    _assert(len(raw_docx) >= MIN_DOCX_BYTES, f"DOCX is too small: {len(raw_docx)} bytes")
    downloaded = OUTPUT_ROOT / "downloaded.docx"
    downloaded.write_bytes(raw_docx)
    _check_docx_text(downloaded)
    print(f"  ✅ download + DOCX text ({len(raw_docx)} bytes)")

    summary = {
        "status": "PASS",
        "base_url": BASE_URL,
        "upload_id": upload_id,
        "task_id": task_id,
        "qa_status": report_data["qa_status"],
        "downloaded_docx": str(downloaded),
        "duration_seconds": round(time.monotonic() - started, 3),
    }
    report_path = OUTPUT_ROOT / "web_smoke_report.json"
    report_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"  report: {report_path}")
    print("Web smoke test passed.")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except SmokeFailure as exc:
        print(f"Web smoke test failed: {exc}", file=sys.stderr)
        raise SystemExit(1)
