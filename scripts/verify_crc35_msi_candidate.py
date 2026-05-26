#!/usr/bin/env python3
"""Smoke-verify a local CRC35+MSI candidate DOCX with one or more Excel files.

The script prints only sanitized metrics and SHA prefixes. It writes rendered
DOCX outputs under tmp/ by default and does not expose patient/sample fields.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import zipfile
from pathlib import Path
from typing import Any

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from reportgen.core.enhancer_registry import CRC358Enhancer  # noqa: E402
from reportgen.core.excel_reader import ExcelReader  # noqa: E402
from reportgen.core.field_mapper import FieldMapper  # noqa: E402
from reportgen.core.template_renderer import TemplateRenderer  # noqa: E402


DEFAULT_TEMPLATE = Path(
    "tmp/golden_template_seed/crc_35_msi_golden_template_v0_candidate.docx"
)
DEFAULT_OUTPUT_ROOT = Path("tmp/crc35_candidate_verify")
DEFAULT_PROCESSORS = ("part3_formatted_sections",)


def sha12(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()[:12]


def count_unresolved_markers(path: Path) -> dict[str, int]:
    with zipfile.ZipFile(path) as archive:
        text = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )
    return {
        "jinja_variable": len(re.findall(r"\{\{\s*[^{}]+?\s*\}\}", text)),
        "jinja_block": len(re.findall(r"\{%\s*[^{}]+?\s*%\}", text)),
        "part3_marker": text.count("__PART3_MARKER__"),
    }


def table_row_count(path: Path, table_index: int) -> int:
    doc = Document(str(path))
    if table_index >= len(doc.tables):
        return 0
    return len(doc.tables[table_index].rows)


def verify_one(
    *,
    excel: Path,
    template: Path,
    output_root: Path,
    case_index: int,
) -> dict[str, Any]:
    reader = ExcelReader(config_dir="config", log_level="ERROR")
    mapper = FieldMapper(config_dir="config", log_level="ERROR")
    renderer = TemplateRenderer(log_level="ERROR")

    excel_data = reader.read(str(excel), include_tables=True)
    report_data = mapper.map(excel_data)
    report_data = CRC358Enhancer().enhance(
        report_data,
        excel_data,
        field_mapper=mapper,
        base_path=".",
        project_type="crc_358_msi",
    )

    output_root.mkdir(parents=True, exist_ok=True)
    output = output_root / f"case_{case_index:02d}_{sha12(excel)}.docx"
    renderer.render(
        str(template),
        report_data,
        str(output),
        post_processor_names=DEFAULT_PROCESSORS,
    )
    Document(str(output))

    drug_keys = [
        key
        for key, value in report_data.context.items()
        if key.startswith("drug_") and isinstance(value, list)
    ]
    unresolved = count_unresolved_markers(output)
    table10_rows = table_row_count(output, 10)
    yilitikang_rows = len(report_data.get_table("drug_yilitikang") or [])
    ok = (
        all(count == 0 for count in unresolved.values())
        and table10_rows == yilitikang_rows + 1
        and yilitikang_rows > 0
    )

    return {
        "case": f"case_{case_index:02d}",
        "excel_sha12": sha12(excel),
        "output": str(output),
        "ok": ok,
        "sheets": len(excel_data.sheet_names),
        "context": {
            "variants_2_1_rows": len(report_data.get_table("variants_2_1") or []),
            "chemotherapy_rows": len(report_data.get_table("chemotherapy") or []),
            "drug_table_count": len(drug_keys),
            "nonempty_drug_table_count": sum(
                1 for key in drug_keys if report_data.context[key]
            ),
            "drug_yilitikang_rows": yilitikang_rows,
        },
        "rendered": {
            "table10_rows": table10_rows,
            "unresolved_markers": unresolved,
        },
    }


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--template",
        type=Path,
        default=DEFAULT_TEMPLATE,
        help="Local CRC35 candidate DOCX template",
    )
    parser.add_argument(
        "--output-root",
        type=Path,
        default=DEFAULT_OUTPUT_ROOT,
        help="Directory for rendered smoke outputs",
    )
    parser.add_argument(
        "--excel",
        action="append",
        required=True,
        type=Path,
        help="Excel file to verify; can be repeated",
    )
    return parser


def main() -> int:
    args = build_parser().parse_args()
    results = [
        verify_one(
            excel=excel,
            template=args.template,
            output_root=args.output_root,
            case_index=index,
        )
        for index, excel in enumerate(args.excel, start=1)
    ]
    payload = {
        "template": str(args.template),
        "case_count": len(results),
        "ok": all(item["ok"] for item in results),
        "results": results,
    }
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    return 0 if payload["ok"] else 2


if __name__ == "__main__":
    raise SystemExit(main())
