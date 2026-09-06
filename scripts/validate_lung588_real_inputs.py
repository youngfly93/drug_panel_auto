#!/usr/bin/env python3
# 步骤: 71 肺癌588真实输入脱敏契约验证
# 上游: 外部受控肺癌Excel、panels/lung_588_pdl1/context_contracts/
# 输出: .work/lung588_real_input_audit/validation.json
# 种子: 无（确定性字段映射与契约比对）
"""Validate local lung588 inputs without writing patient identifiers to artifacts."""

from __future__ import annotations

import argparse
import contextlib
import copy
import hashlib
import io
import json
import os
import re
import subprocess
import sys
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from PIL import Image, ImageDraw


ROOT = Path(__file__).resolve().parents[1]
HISTORICAL_CHEMOTHERAPY_GENES = {
    "顺铂": "XPC、MTHFR、GSTP1、ERCC1、XRCC1、GSTM1",
    "长春瑞滨": "ABCB1、CEP72",
    "依托泊苷": "SLCO1B1、DYNC2H1",
    "吉西他滨": "CDA、RRM1",
    "多西他赛": "ERCC1、VEGFA",
    "培美曲塞": "SLC19A1、GGH",
    "紫杉醇": "ERCC1、ABCB1、SOD2、TP53",
    "卡铂": "MTHFR、ERCC1、XRCC1",
}
HISTORICAL_CHEMOTHERAPY_RATINGS = {
    "CASE-LUNG-B": {
        "顺铂": ("可能居中", "可能较高"),
        "长春瑞滨": ("可能较低", "可能较低"),
        "依托泊苷": ("可能较高", "可能较高"),
        "吉西他滨": ("可能较高", "可能较低"),
        "多西他赛": ("可能较低", "可能较高"),
        "培美曲塞": ("可能居中", "可能较高"),
        "紫杉醇": ("可能较低", "可能较高"),
        "卡铂": ("可能较低", "可能居中"),
    },
    "CASE-LUNG-C": {
        "顺铂": ("可能较高", "可能较高"),
        "长春瑞滨": ("可能较高", "可能较低"),
        "依托泊苷": ("可能较高", "可能较高"),
        "吉西他滨": ("可能较高", "可能较低"),
        "多西他赛": ("可能较高", "可能居中"),
        "培美曲塞": ("可能居中", "可能较高"),
        "紫杉醇": ("可能较高", "可能较低"),
        "卡铂": ("可能较高", "可能居中"),
    },
}
HISTORICAL_CHEMOTHERAPY_SUMMARIES = {
    "CASE-LUNG-B": "经分析，可考虑优先选择的化疗方案有吉西他滨单药方案。",
    "CASE-LUNG-C": (
        "经分析，可考虑优先选择的化疗方案有吉西他滨+长春瑞滨、"
        "吉西他滨单药方案、白蛋白结合型紫杉醇单药方案、紫杉醇单药方案。"
    ),
}
HISTORICAL_IRINOTECAN_SAFETY = {
    "CASE-LUNG-B": {
        "result": "UGT1A1基因型为6TA/6TA",
        "dose_evaluation": "正常剂量使用",
    },
    "CASE-LUNG-C": {
        "result": "UGT1A1基因型为6TA/7TA",
        "dose_evaluation": "减少剂量使用",
    },
}
# 2026-09-06 explicit user decision: #14/#15 retain the historical display
# until the report group resolves the medical questions. This verifies display
# compatibility, not clinical validity; previous 2026-09-05 receipts remain intact.
CURRENT_CHEMOTHERAPY_GENES = copy.deepcopy(HISTORICAL_CHEMOTHERAPY_GENES)
CURRENT_CHEMOTHERAPY_RATINGS = copy.deepcopy(HISTORICAL_CHEMOTHERAPY_RATINGS)
CURRENT_CHEMOTHERAPY_SUMMARIES = copy.deepcopy(HISTORICAL_CHEMOTHERAPY_SUMMARIES)
CURRENT_IRINOTECAN_SAFETY = copy.deepcopy(HISTORICAL_IRINOTECAN_SAFETY)
KNOWN_INPUTS = {
    "267a8cbab4d112ea38660dcb1734bb4fb3a7269f50abed6d83a9bf1262ee5646": {
        "alias": "CASE-LUNG-A",
        "contract_id": "case_lung_a",
        "pdl1_tps": 1.0,
        "pdl1_cps": 1.0,
        "pdl1_result": "阳性（低表达）",
        "expected_targeted_drug_count": 2,
        # CASE-LUNG-A has no paired historical final from which a positive
        # exact-event association can be transcribed. Its IFNGR1 event remains
        # covered by the fixed historical IFNGR1/2 gene-group row.
        "expected_immune_positive_count": 0,
        "expected_immune_negative_count": 1,
    },
    "623c96cee1eb7b16cacb62cababba3b790e82007a00a59d0f159efbe025db000": {
        "alias": "CASE-LUNG-B",
        "contract_id": "case_lung_b",
        "pdl1_tps": 50.0,
        "pdl1_cps": 52.0,
        "pdl1_result": "阳性（高表达）",
        "expected_targeted_drug_count": 8,
        "expected_targeted_drug_genes": [
            "TP53",
            "ATM",
            "TSC1",
            "MSH3",
            "BRCA2",
            "MLH1",
            "PMS2",
            "BRAF",
        ],
        "expected_immune_positive_count": 5,
        "expected_immune_negative_count": 0,
    },
    "7b39431044c4a9298f7663c97a47c4df83b5b1e0875d88a64b3e24c05bfa498a": {
        "alias": "CASE-LUNG-C",
        "contract_id": "case_lung_c",
        "pdl1_tps": 5.0,
        "pdl1_cps": 6.0,
        "pdl1_result": "阳性（低表达）",
        "expected_targeted_drug_count": 8,
        "expected_targeted_drug_genes": [
            "BRAF",
            "PTEN",
            "ERBB2",
            "ATM",
            "TP53",
            "TSC2",
            "BRIP1",
            "PIK3CA",
        ],
        "expected_immune_positive_count": 2,
        "expected_immune_negative_count": 1,
    },
}
COMMIT_RE = re.compile(r"^[0-9a-f]{40}$")
UAT_POLICY_PATH = (
    ROOT / "panels" / "lung_588_pdl1" / "uat" / "lung588_risk_based_release_policy.yaml"
)
REPORT_GROUP_UAT_DECISIONS_PATH = (
    ROOT
    / "panels"
    / "lung_588_pdl1"
    / "uat"
    / "lung588_report_group_uat_decisions.yaml"
)
REVIEW_CANDIDATE_CONTRACT_PATH = (
    ROOT
    / "panels"
    / "lung_588_pdl1"
    / "review_baselines"
    / "lung588_historical_review_candidate_v1.yaml"
)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _write_synthetic_pdl1_image(path: Path, alias: str) -> None:
    """Create a deterministic non-clinical image used only for render QA."""

    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (960, 640), "white")
    draw = ImageDraw.Draw(image)
    digest = hashlib.sha256(alias.encode("utf-8")).digest()
    for index in range(80):
        x = (digest[index % len(digest)] * 37 + index * 61) % 920
        y = (digest[(index + 7) % len(digest)] * 29 + index * 43) % 600
        radius = 8 + digest[(index + 13) % len(digest)] % 18
        shade = 75 + digest[(index + 17) % len(digest)] % 120
        draw.ellipse(
            (x, y, x + radius, y + radius),
            fill=(shade, 70, min(220, shade + 35)),
        )
    draw.rectangle((10, 10, 949, 629), outline=(80, 80, 80), width=3)
    image.save(path, format="PNG", optimize=True)


def _source_revision() -> str:
    """Return a quiet immutable source identity in Git and release archives."""

    configured = str(os.environ.get("REPORTGEN_SOURCE_REVISION") or "").strip()
    if COMMIT_RE.fullmatch(configured):
        return configured

    revision_file = ROOT / "REVISION"
    if revision_file.is_file():
        lines = revision_file.read_text(encoding="utf-8").strip().splitlines()
        revision = lines[0] if lines else ""
        if COMMIT_RE.fullmatch(revision):
            return revision

    try:
        completed = subprocess.run(
            ["git", "-C", str(ROOT), "rev-parse", "HEAD"],
            check=False,
            capture_output=True,
            text=True,
            timeout=5,
        )
    except (OSError, subprocess.TimeoutExpired):
        return ""
    revision = completed.stdout.strip()
    return (
        revision if completed.returncode == 0 and COMMIT_RE.fullmatch(revision) else ""
    )


def _load_yaml_mapping(path: Path) -> dict[str, Any]:
    payload = yaml.safe_load(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise RuntimeError(f"expected a mapping in {path}")
    return payload


def _load_uat_release_policy(path: Path = UAT_POLICY_PATH) -> dict[str, Any]:
    policy = _load_yaml_mapping(path)
    case_policy = policy.get("real_case_policy")
    if not isinstance(case_policy, dict):
        raise RuntimeError("lung588 UAT policy is missing real_case_policy")
    if case_policy.get("fixed_minimum_real_case_count") is not None:
        raise RuntimeError(
            "lung588 risk-based UAT policy must not restore a fixed case count"
        )
    case_requirements = policy.get("case_requirements") or {}
    decisions_required = bool(
        case_requirements.get("report_group_decision_required", True)
    )
    expected_fraction = 1.0 if decisions_required else 0.0
    for key in ("required_review_fraction", "required_pass_fraction"):
        value = case_policy.get(key)
        if (
            isinstance(value, bool)
            or not isinstance(value, (int, float))
            or float(value) != expected_fraction
        ):
            raise RuntimeError(f"lung588 UAT policy has invalid {key}: {value!r}")
    if not case_policy.get("require_non_empty_case_set"):
        raise RuntimeError("lung588 UAT policy must require a non-empty real case set")
    return policy


def _load_report_group_uat_decisions(
    path: Path = REPORT_GROUP_UAT_DECISIONS_PATH,
    *,
    expected_policy_id: str | None = None,
) -> dict[str, dict[str, Any]]:
    payload = _load_yaml_mapping(path)
    if expected_policy_id and payload.get("policy_id") != expected_policy_id:
        raise RuntimeError(
            "lung588 report-group UAT register does not match the active policy"
        )
    decisions: dict[str, dict[str, Any]] = {}
    for item in payload.get("cases") or []:
        if not isinstance(item, dict):
            raise RuntimeError("lung588 report-group UAT case entry must be a mapping")
        alias = str(item.get("alias") or "").strip()
        decision = str(item.get("decision") or "").strip().lower()
        if not alias:
            raise RuntimeError("lung588 report-group UAT case alias is required")
        if alias in decisions:
            raise RuntimeError(f"duplicate lung588 report-group UAT alias: {alias}")
        if decision not in {"pass", "fail", "pending"}:
            raise RuntimeError(
                f"invalid lung588 report-group UAT decision for {alias}: {decision!r}"
            )
        p0_count = item.get("p0_count")
        if decision in {"pass", "fail"} and (
            isinstance(p0_count, bool) or not isinstance(p0_count, int) or p0_count < 0
        ):
            raise RuntimeError(
                f"completed lung588 UAT decision for {alias} requires p0_count"
            )
        decisions[alias] = {
            "decision": decision,
            "reviewer": str(item.get("reviewer") or "").strip(),
            "reviewed_at": str(item.get("reviewed_at") or "").strip(),
            "p0_count": p0_count,
        }
    return decisions


def _safe_variant_rows(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        {
            "gene": row.get("gene"),
            "transcript": row.get("transcript"),
            "chromosome": row.get("chromosome"),
            "exon": row.get("exon"),
            "cHGVS": row.get("cHGVS"),
            "pHGVS": row.get("pHGVS"),
            "mutation_type": row.get("mutation_type"),
            "gene_class": row.get("gene_class"),
            "frequency": row.get("frequency"),
        }
        for row in rows
        if isinstance(row, dict)
    ]


def _clinical_info(case: dict[str, Any]) -> dict[str, Any]:
    return {
        "patient_name": case["alias"],
        "sample_id": case["alias"],
        "report_number": case["alias"],
        "project_name": "肺癌588基因+PD-L1",
        "clinical_diagnosis": "肺癌（脱敏验证）",
        "report_date": "2026-07-23",
        "pdl1_tps": case["pdl1_tps"],
        "pdl1_cps": case["pdl1_cps"],
        "pdl1_result": case["pdl1_result"],
        "pdl1_image_path": str(case["pdl1_image_path"]),
        "pdl1_assay_profile_id": ("legacy_unspecified_ihc_transcription_v1"),
        "pdl1_source_record_id": (f"SYNTHETIC-VISUAL-QA-IHC-{case['alias']}"),
        "pdl1_source_record_date": "2026-07-23",
        "pdl1_specimen_id": (f"SYNTHETIC-VISUAL-QA-SPECIMEN-{case['alias']}"),
        "pdl1_image_disposition": "病例专属图像（报告展示）",
        "lung_histology": "非小细胞肺癌",
        "disease_extent": "转移性",
        "prior_systemic_therapy": "已接受",
        "companion_diagnostic_status": "已确认符合",
    }


def _enhance_case(bridge, excel_data, case: dict[str, Any]) -> dict[str, Any]:
    from reportgen.core.context_contract import (
        check_context_contract,
        load_context_contract,
    )
    from reportgen.core.enhancer_registry import get_enhancer, get_panel_registry
    from reportgen.core.report_generator import validate_panel_biomarker_contracts
    from reportgen.core.report_summary import build_report_summary
    from reportgen.models.excel_data import ExcelDataSource
    from reportgen.rules.pdl1 import (
        apply_pdl1_product_display_fields,
        load_pdl1_product_contract,
        validate_pdl1_product_contract,
    )

    working = ExcelDataSource(
        file_path=excel_data.file_path,
        single_values=copy.deepcopy(excel_data.single_values or {}),
        table_data=copy.deepcopy(excel_data.table_data or {}),
        sheet_names=list(excel_data.sheet_names or []),
        metadata=copy.deepcopy(excel_data.metadata or {}),
    )
    clinical_info = _clinical_info(case)
    bridge._inject_clinical_info_into_excel(working, clinical_info)
    registration = get_panel_registry().get("lung_588_pdl1")
    package = registration.package
    report_data = bridge.field_mapper.map(working, panel_package=package)
    report_data = bridge.data_cleaner.validate_and_clean(report_data)
    part3_policy = package.raw.get("part3_knowledge") or {}
    gene_knowledge_provider = (
        bridge._build_gene_knowledge_provider(package)
        if part3_policy.get("enabled", True)
        else None
    )
    report_data = get_enhancer("lung_588_pdl1").enhance(
        report_data,
        working,
        field_mapper=bridge.field_mapper,
        gene_knowledge_provider=gene_knowledge_provider,
        base_path=str(ROOT),
        project_type="lung_588_pdl1",
        panel_package=package,
    )
    pdl1_product_contract = load_pdl1_product_contract(package)
    apply_pdl1_product_display_fields(
        report_data,
        pdl1_product_contract,
    )
    biomarker_failures = validate_panel_biomarker_contracts(
        report_data,
        package.input_contract.get("biomarkers"),
    )
    pdl1_product_failures = validate_pdl1_product_contract(
        report_data,
        pdl1_product_contract,
    )
    context = report_data.get_template_context()
    web_summary = build_report_summary(
        report_data=report_data,
        project_type="lung_588_pdl1",
        project_name="肺癌588基因+PD-L1",
    )
    contract_report = None
    contract_id = case.get("contract_id")
    if contract_id:
        contract_path = package.resolve_context_contract_file(contract_id)
        contract_report = check_context_contract(
            context,
            load_context_contract(contract_path),
            contract_path=contract_path,
        )
    chemotherapy_rows = list(report_data.get_table("chemotherapy_predictions") or [])
    targeted_rows = list(report_data.get_table("targeted_drug_tips") or [])
    guideline_rows = list(report_data.get_table("lung_guideline_drug_results") or [])
    drug_analysis_rows = list(report_data.get_table("drug_analysis_sections") or [])
    chemotherapy_detail_rows = [
        row
        for table_name, table_rows in report_data.context.items()
        if table_name.startswith("drug_") and isinstance(table_rows, list)
        for row in table_rows
        if isinstance(row, dict)
        and any(key in row for key in ("Level", "level", "等级", "Evidence"))
    ]
    return {
        "variant_rows": _safe_variant_rows(
            list(report_data.get_table("all_variants") or [])
        ),
        "targeted_drug_count": len(
            targeted_rows
        ),
        "targeted_drug_genes": [
            str(row.get("gene") or "")
            for row in targeted_rows
            if isinstance(row, dict) and str(row.get("gene") or "")
        ],
        "targeted_drug_intro_count": len(
            list(report_data.get_table("targeted_drug_introductions") or [])
        ),
        "part3_drug_analysis_count": len(drug_analysis_rows),
        "part3_drug_analysis_genes": sorted(
            {
                str(row.get("gene") or "")
                for row in drug_analysis_rows
                if isinstance(row, dict) and str(row.get("gene") or "")
            }
        ),
        "part3_detected_lead_count": sum(
            str(row.get("relation") or "").lstrip().startswith("该样本检出")
            for row in drug_analysis_rows
            if isinstance(row, dict)
        ),
        "reference_count": len(list(report_data.get_table("references") or [])),
        "immune_positive_count": len(
            list(report_data.get_table("immune_positive_variants") or [])
        ),
        "immune_negative_count": len(
            list(report_data.get_table("immune_negative_variants") or [])
        ),
        "lung_guideline_row_count": len(
            guideline_rows
        ),
        "lung_guideline_results": {
            str(row.get("key") or ""): str(row.get("result") or "")
            for row in guideline_rows
            if isinstance(row, dict)
        },
        "immune_fixed_row_counts": {
            category: len(list(report_data.get_table(table_name) or []))
            for category, table_name in (
                ("positive", "immune_positive_results"),
                ("negative", "immune_negative_results"),
                ("hyperprogression", "immune_hyperprogression_results"),
            )
        },
        "chemotherapy_row_counts": {
            category: len(list(report_data.get_table(table_name) or []))
            for category, table_name in (
                ("prediction", "chemotherapy_predictions"),
                ("regimen", "chemotherapy_regimen_predictions"),
                ("dosage", "chemotherapy_dosage_rows"),
            )
        },
        "chemotherapy_base_rows": [
            {
                "drug": str(row.get("drug") or ""),
                "genes": str(row.get("genes") or ""),
                "efficacy": str(row.get("efficacy") or ""),
                "toxicity": str(row.get("toxicity") or ""),
            }
            for row in chemotherapy_rows[:8]
            if isinstance(row, dict)
        ],
        "chemotherapy_summary_text": report_data.get_field(
            "chemotherapy_summary_text"
        ),
        "irinotecan_safety_rows": [
            {
                "result": str(row.get("result") or ""),
                "dose_evaluation": str(row.get("dose_evaluation") or ""),
            }
            for row in report_data.get_table("irinotecan_safety_rows") or []
            if isinstance(row, dict)
        ],
        "chemotherapy_detail_quality": {
            "blank_result_count": sum(
                not str(row.get("Result") or "").strip()
                for row in chemotherapy_detail_rows
            ),
            "level_3_count": sum(
                str(row.get("Level") or row.get("level") or "").strip().upper()
                == "3"
                for row in chemotherapy_detail_rows
            ),
            "english_uncovered_count": sum(
                str(row.get("Result") or "").strip().lower() == "uncovered"
                for row in chemotherapy_detail_rows
            ),
            "english_prefix_drug_count": sum(
                bool(
                    re.match(
                        r"^[A-Za-z]",
                        str(row.get("DrugDisplay") or row.get("药物名称") or "").strip(),
                    )
                )
                for row in chemotherapy_detail_rows
            ),
        },
        "cisplatin_detail_row_count": len(
            list(report_data.get_table("drug_shunbo") or [])
        ),
        "tmb_reference": report_data.get_field("tmb_reference"),
        "undetected_gene_names": [
            str(row.get("gene") or "")
            for row in report_data.get_table("variants_2_1") or []
            if str(row.get("locus") or "") == "未见突变"
        ],
        "biomarkers": {
            "tmb_value": report_data.get_field("tmb_value"),
            "tmb_status": report_data.get_field("tmb_status"),
            "msi_status": report_data.get_field("msi_status"),
            "pdl1_tps": report_data.get_field("pdl1_tps"),
            "pdl1_cps": report_data.get_field("pdl1_cps"),
            "pdl1_result": report_data.get_field("pdl1_result"),
        },
        "biomarker_contract_status": "PASS" if not biomarker_failures else "FAIL",
        "biomarker_failures": biomarker_failures,
        "pdl1_product_contract_status": (
            "PASS" if not pdl1_product_failures else "FAIL"
        ),
        "pdl1_product_failures": pdl1_product_failures,
        "pdl1_input_provenance": "synthetic_visual_qa_only",
        "web_preview": {
            "drug_related_variant_count": web_summary["variants"]["drug_related"],
            "targeted_drug_count": web_summary["drugs"]["targeted_count"],
            "targeted_module_status": web_summary["drugs"]["targeted_status"],
            "immune": web_summary["biomarkers"]["immune"],
        },
        "context_contract": {
            "contract_id": contract_id,
            "status": contract_report["status"]
            if contract_report
            else "NOT_APPLICABLE",
            "summary": contract_report["summary"] if contract_report else {},
        },
    }


def _render_case(
    bridge,
    excel_path: Path,
    case: dict[str, Any],
    output_dir: Path,
    *,
    dpi: int,
) -> dict[str, Any]:
    from reportgen.core.review_candidate_contract import (
        load_review_candidate_contract,
        validate_review_candidate_output,
        validate_review_candidate_template,
    )

    result = bridge.generate_report(
        str(excel_path),
        str(output_dir),
        clinical_info=_clinical_info(case),
        project_type="lung_588_pdl1",
        project_name="肺癌588基因+PD-L1",
        strict_mode=False,
        template_contract_mode="fail",
        qa_visual_render="all",
        qa_visual_render_required=True,
        qa_visual_render_dpi=dpi,
        qa_visual_render_timeout_seconds=180,
    )
    output_file = Path(str(result.get("output_file") or ""))
    qa_path = output_file.with_suffix(".qa.json") if output_file.name else None
    qa_payload: dict[str, Any] = {}
    if qa_path is not None and qa_path.is_file():
        qa_payload = json.loads(qa_path.read_text(encoding="utf-8"))
    visual = (qa_payload.get("checks") or {}).get("visual_render") or {}
    pixel = visual.get("pixel_check") or {}

    contract = load_review_candidate_contract(REVIEW_CANDIDATE_CONTRACT_PATH)
    template_identity = dict(result.get("template_identity") or {})
    selected_template = Path(bridge._resolve_template_path(None, "lung_588_pdl1"))
    template_gate = validate_review_candidate_template(
        contract,
        selected_template,
        template_id=str(template_identity.get("template_id") or ""),
        template_version=str(template_identity.get("version") or ""),
        template_status=str(template_identity.get("status") or ""),
    )
    output_gate = validate_review_candidate_output(
        contract,
        output_file,
        qa_report=qa_payload,
        expected_texts=(str(case["alias"]), str(case["pdl1_result"])),
        require_case_image=True,
    )

    content_failures: list[str] = []
    rendered_semantics: dict[str, Any] = {}
    if output_file.is_file():
        document = Document(output_file)
        visible = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ]
        )
        required_texts = [
            "Gene List for MLseq (n=588)",
            "第三部分：基因变异及相应靶向/免疫药物解析",
            "图1. 免疫组化：PD-L1",
            "报告组评审候选稿（非临床交付）",
        ]
        alias = str(case.get("alias") or "")
        if alias in HISTORICAL_CHEMOTHERAPY_SUMMARIES:
            required_texts.extend(
                (
                    HISTORICAL_CHEMOTHERAPY_SUMMARIES[alias],
                    HISTORICAL_IRINOTECAN_SAFETY[alias]["result"],
                    HISTORICAL_IRINOTECAN_SAFETY[alias]["dose_evaluation"],
                )
            )
        if alias == "CASE-LUNG-C":
            required_texts.extend(
                (
                    "c.1799T>A,p.V600E",
                    "c.1979G>A,p.G660D",
                    "PMID:19935797",
                    "PMID:19966866",
                    "PMID:30449325",
                )
            )
        forbidden_texts = (
            "__PART3_MARKER__",
            "__PDL1_CASE_IMAGE__",
            "n=329",
            "{{",
            "{%",
            "工程草案",
            "报告组二审",
            "报告组复核",
            "脱敏UAT",
            "报告组评审中",
            "依据病例Excel的CtDrug表生成",
            "伊立替康剂量参考（未启用",
        )
        content_failures.extend(
            f"missing:{text}" for text in required_texts if text not in visible
        )

        def matching_table(*tokens: str, column_count: int):
            return next(
                (
                    table
                    for table in document.tables
                    if len(table.columns) == column_count
                    and all(
                        token
                        in "\n".join(cell.text for cell in table.rows[0].cells)
                        for token in tokens
                    )
                ),
                None,
            )

        targeted_table = matching_table(
            "潜在获益靶向药物",
            "可能耐药或慎重药物",
            column_count=4,
        )
        introduction_table = matching_table("药物介绍", column_count=3)
        irinotecan_table = matching_table("剂量安全性评价", column_count=3)
        rendered_semantics = {
            "targeted_summary_row_count": (
                len(targeted_table.rows) - 1 if targeted_table is not None else -1
            ),
            "targeted_intro_row_count": (
                len(introduction_table.rows) - 1
                if introduction_table is not None
                else -1
            ),
            "irinotecan_row_count": (
                len(irinotecan_table.rows) - 1
                if irinotecan_table is not None
                else -1
            ),
        }
        expected_targeted_count = int(case.get("expected_targeted_drug_count") or 0)
        for key in ("targeted_summary_row_count", "targeted_intro_row_count"):
            if rendered_semantics[key] != expected_targeted_count:
                content_failures.append(
                    f"rendered_count:{key}={rendered_semantics[key]}"
                )
        if rendered_semantics["irinotecan_row_count"] != 1:
            content_failures.append(
                "rendered_count:irinotecan_row_count="
                f"{rendered_semantics['irinotecan_row_count']}"
            )
        if (
            int(case.get("expected_targeted_drug_count") or 0) > 0
            and "【待报告组审】" not in visible
        ):
            content_failures.append("missing:【待报告组审】")
        lowered = visible.lower()
        content_failures.extend(
            f"forbidden:{text}" for text in forbidden_texts if text.lower() in lowered
        )
    else:
        content_failures.append("output_missing")

    blank_pages = list(pixel.get("blank_pages") or [])
    low_content_pages = list(pixel.get("unexpected_low_content_pages") or [])
    status = (
        "PASS"
        if result.get("success")
        and result.get("qa_status") in {"PASS", "WARN"}
        and qa_payload.get("status") in {"PASS", "WARN"}
        and not blank_pages
        and not low_content_pages
        and not content_failures
        and template_gate["status"] == "PASS"
        and output_gate["status"] == "PASS"
        else "FAIL"
    )
    failed_codes = sorted(
        {
            str(item.get("code") or "UNKNOWN")
            for gate in (template_gate, output_gate)
            for item in gate.get("errors") or []
        }
    )
    return {
        "status": status,
        "output_alias": output_file.name,
        "_output_file": str(output_file) if output_file.is_file() else "",
        "qa_status": qa_payload.get("status") or result.get("qa_status"),
        "page_count": pixel.get("checked_pages"),
        "blank_page_count": len(blank_pages),
        "unexpected_low_content_page_count": len(low_content_pages),
        "content_failures": content_failures,
        "rendered_semantics": rendered_semantics,
        "template_identity": template_identity,
        "review_candidate_contract": {
            "contract_id": contract["contract_id"],
            "status": (
                "PASS"
                if template_gate["status"] == "PASS" and output_gate["status"] == "PASS"
                else "FAIL"
            ),
            "template_status": template_gate["status"],
            "output_status": output_gate["status"],
            "failed_codes": failed_codes,
        },
        "error_count": len(result.get("errors") or []),
    }


def _build_uat_readiness(
    rows: list[dict[str, Any]],
    *,
    policy: dict[str, Any] | None = None,
    report_group_decisions: dict[str, dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """Evaluate every available real case without a fixed numeric denominator."""

    policy = policy or _load_uat_release_policy()
    case_policy = policy["real_case_policy"]
    case_requirements = policy.get("case_requirements") or {}
    decisions_required = bool(
        case_requirements.get("report_group_decision_required", True)
    )
    reviewer_and_date_required = bool(
        case_requirements.get("report_group_reviewer_and_date_required", True)
    )
    source_required = bool(
        case_requirements.get("verified_case_specific_ihc_source_required", True)
    )
    if report_group_decisions is None:
        report_group_decisions = (
            _load_report_group_uat_decisions(
                expected_policy_id=str(policy["policy_id"])
            )
            if decisions_required
            else {}
        )
    aliases = [str(row.get("alias") or "").strip() for row in rows]
    ngs_structure_pass_count = sum(
        row["auto_detection"]["detected"]
        and row["auto_detection"]["project_type"] == "lung_588_pdl1"
        and row.get("targeted_drug_count", 0)
        == row.get("expected_targeted_drug_count", 0)
        and row.get("immune_positive_count", 0)
        == row.get("expected_immune_positive_count", 0)
        and row.get("immune_negative_count", 0)
        == row.get("expected_immune_negative_count", 0)
        and row["biomarker_contract_status"] == "PASS"
        and row["context_contract"]["status"] in {"PASS", "NOT_APPLICABLE"}
        for row in rows
    )
    pdl1_product_pass_count = sum(
        row["pdl1_product_contract_status"] == "PASS" for row in rows
    )
    verified_case_pdl1_source_count = sum(
        row["pdl1_input_provenance"] == "case_specific_verified_ihc_source"
        for row in rows
    )
    missing_case_alias_count = sum(not alias for alias in aliases)
    missing_decision_aliases = [
        alias for alias in aliases if alias and alias not in report_group_decisions
    ]
    complete_decisions = {
        alias: report_group_decisions[alias]
        for alias in aliases
        if alias in report_group_decisions
        and report_group_decisions[alias]["decision"] in {"pass", "fail"}
        and (
            not reviewer_and_date_required
            or (
                report_group_decisions[alias]["reviewer"]
                and report_group_decisions[alias]["reviewed_at"]
            )
        )
        and isinstance(report_group_decisions[alias].get("p0_count"), int)
        and not isinstance(report_group_decisions[alias].get("p0_count"), bool)
    }
    report_group_reviewed_case_count = len(complete_decisions)
    report_group_passed_case_count = sum(
        item["decision"] == "pass" for item in complete_decisions.values()
    )
    report_group_failed_case_count = sum(
        item["decision"] == "fail" for item in complete_decisions.values()
    )
    p0_count = sum(int(item["p0_count"]) for item in complete_decisions.values())
    required_review_case_count = len(rows) if decisions_required else 0
    p0_allowed = int(case_policy.get("p0_allowed", 0))
    blockers: list[dict[str, str]] = []
    feedback_policy = policy.get("product_owner_feedback_source") or {}
    feedback_entries = [
        entry
        for entry in feedback_policy.get("entries") or []
        if isinstance(entry, dict) and str(entry.get("source") or "").strip()
    ]
    minimum_feedback_entries = int(
        feedback_policy.get("minimum_entry_count", 1)
    )
    if not rows:
        blockers.append(
            {
                "code": "NO_REGISTERED_REAL_CASES",
                "message": "the frozen release has no registered real lung588 cases",
            }
        )
    if missing_case_alias_count:
        blockers.append(
            {
                "code": "REAL_CASE_ALIAS_MISSING",
                "message": f"{missing_case_alias_count} observed cases lack a stable alias",
            }
        )
    if ngs_structure_pass_count != len(rows):
        blockers.append(
            {
                "code": "NGS_STRUCTURE_INCOMPLETE",
                "message": (
                    f"{len(rows) - ngs_structure_pass_count} observed cases "
                    "do not pass the frozen NGS structure contract"
                ),
            }
        )
    if pdl1_product_pass_count != len(rows):
        blockers.append(
            {
                "code": "PDL1_PRODUCT_CONTRACT_BLOCKED",
                "message": (
                    f"{len(rows) - pdl1_product_pass_count} observed cases "
                    "do not pass an enabled PD-L1 product contract"
                ),
            }
        )
    if source_required and verified_case_pdl1_source_count != len(rows):
        blockers.append(
            {
                "code": "PDL1_CASE_SOURCE_NOT_VERIFIED",
                "message": (
                    f"{len(rows) - verified_case_pdl1_source_count} observed "
                    "cases use synthetic machine-QA values rather than a "
                    "verified case-specific IHC source"
                ),
            }
        )
    if decisions_required and missing_decision_aliases:
        blockers.append(
            {
                "code": "REPORT_GROUP_UAT_RECORD_MISSING",
                "message": (
                    f"{len(missing_decision_aliases)} observed cases are absent "
                    "from the report-group UAT register"
                ),
            }
        )
    if decisions_required and report_group_reviewed_case_count != required_review_case_count:
        blockers.append(
            {
                "code": "REPORT_GROUP_UAT_INCOMPLETE",
                "message": (
                    f"{report_group_reviewed_case_count}/"
                    f"{required_review_case_count} observed cases have a complete "
                    "report-group UAT decision, reviewer and date"
                ),
            }
        )
    # A decision is optional for pilot draft access, but an explicitly recorded
    # FAIL remains material evidence for formal release readiness.
    if report_group_failed_case_count:
        blockers.append(
            {
                "code": "REPORT_GROUP_UAT_FAILED",
                "message": (
                    f"{report_group_failed_case_count} observed cases have a "
                    "report-group FAIL decision"
                ),
            }
        )
    if p0_count > p0_allowed:
        blockers.append(
            {
                "code": "P0_DEFECTS_PRESENT",
                "message": f"P0 count {p0_count} exceeds allowed count {p0_allowed}",
            }
        )
    if (
        feedback_policy.get("required")
        and len(feedback_entries) < minimum_feedback_entries
    ):
        blockers.append(
            {
                "code": "PRODUCT_OWNER_FEEDBACK_SOURCE_MISSING",
                "message": (
                    f"{len(feedback_entries)}/{minimum_feedback_entries} "
                    "product-owner feedback source entries are recorded"
                ),
            }
        )
    formal_uat_status = "PASS" if not blockers else "BLOCKED"
    return {
        "scope": "risk_based_all_available_real_cases",
        "policy_id": policy["policy_id"],
        "case_set_policy": case_policy["selection"],
        "fixed_minimum_real_case_count": None,
        "observed_real_input_count": len(rows),
        "required_report_group_review_case_count": required_review_case_count,
        "report_group_decision_required": decisions_required,
        "report_group_reviewer_and_date_required": reviewer_and_date_required,
        "product_owner_feedback_source_count": len(feedback_entries),
        "ngs_structure_pass_count": ngs_structure_pass_count,
        "ngs_structure_status": (
            "PASS" if ngs_structure_pass_count == len(rows) else "FAIL"
        ),
        "pdl1_product_pass_count": pdl1_product_pass_count,
        "pdl1_product_status": (
            "PASS" if pdl1_product_pass_count == len(rows) else "BLOCKED"
        ),
        "verified_case_pdl1_source_count": (verified_case_pdl1_source_count),
        "report_group_reviewed_case_count": (report_group_reviewed_case_count),
        "report_group_passed_case_count": report_group_passed_case_count,
        "report_group_failed_case_count": report_group_failed_case_count,
        "p0_count": p0_count,
        "p0_allowed": p0_allowed,
        "formal_uat_status": formal_uat_status,
        "formal_uat_requirement_met": formal_uat_status == "PASS",
        "blockers": blockers,
    }


def validate_inputs(
    input_dir: Path,
    *,
    render_output_dir: Path | None = None,
    render_dpi: int = 120,
) -> dict[str, Any]:
    os.environ["RG_WEB_UPSTREAM_ROOT"] = str(ROOT)
    # This is an offline promotion gate for a draft Panel, not a production
    # service process. Explicitly open the local product so the same script can
    # be run from a shell that inherited iyun129's disabled-panel boundary.
    os.environ["RG_WEB_DISABLED_PROJECT_TYPES"] = ""
    os.environ["REPORTGEN_DISABLED_PROJECT_TYPES"] = ""
    for import_path in (str(ROOT / "backend"), str(ROOT)):
        if import_path not in sys.path:
            sys.path.insert(0, import_path)

    from app.services.reportgen_bridge import ReportGenBridge

    bridge = ReportGenBridge(
        config_dir=str(ROOT / "config"),
        template_dir=str(ROOT / "templates"),
    )
    located: dict[str, Path] = {}
    for path in sorted(input_dir.glob("*.xlsx")):
        if path.name.startswith(("._", "~$")):
            continue
        digest = _sha256(path)
        if digest in KNOWN_INPUTS:
            located[digest] = path

    missing = sorted(set(KNOWN_INPUTS) - set(located))
    if missing:
        raise RuntimeError(
            f"missing {len(missing)} frozen lung588 inputs; hashes only: {missing}"
        )

    rows: list[dict[str, Any]] = []
    synthetic_image_dir = (
        render_output_dir
        if render_output_dir is not None
        else ROOT / ".work" / "lung588_real_input_audit"
    ) / "synthetic_pdl1_images"
    for digest, case in sorted(
        KNOWN_INPUTS.items(),
        key=lambda item: item[1]["alias"],
    ):
        excel_path = located[digest]
        runtime_case = dict(case)
        image_path = synthetic_image_dir / f"{case['alias']}.png"
        _write_synthetic_pdl1_image(image_path, str(case["alias"]))
        runtime_case["pdl1_image_path"] = image_path
        # Upstream loggers may mention real filenames/sample IDs. Capture both
        # streams so this validation emits only the de-identified payload below.
        with (
            contextlib.redirect_stdout(io.StringIO()),
            contextlib.redirect_stderr(io.StringIO()),
        ):
            excel_data = bridge.read_excel(str(excel_path))
            detected = bridge.detect_project_type(
                str(excel_path), excel_data=excel_data
            )
            result = _enhance_case(bridge, excel_data, runtime_case)
            report_generation = (
                _render_case(
                    bridge,
                    excel_path,
                    runtime_case,
                    render_output_dir,
                    dpi=render_dpi,
                )
                if render_output_dir is not None and case.get("contract_id")
                else {"status": "NOT_RUN"}
            )
        rows.append(
            {
                "alias": case["alias"],
                "source_sha256": digest,
                "expected_targeted_drug_count": case["expected_targeted_drug_count"],
                "expected_immune_positive_count": case[
                    "expected_immune_positive_count"
                ],
                "expected_immune_negative_count": case[
                    "expected_immune_negative_count"
                ],
                "sheet_count": len(excel_data.sheet_names or []),
                "auto_detection": {
                    "detected": bool(detected.get("detected")),
                    "project_type": detected.get("project_type"),
                },
                **result,
                "report_generation": report_generation,
            }
        )

    if render_output_dir is not None:
        from reportgen.core.review_candidate_contract import extract_docx_text

        aliases = [str(row["alias"]) for row in rows]
        for row in rows:
            generation = row["report_generation"]
            output_value = str(generation.pop("_output_file", "") or "")
            output_path = Path(output_value) if output_value else None
            own_alias = str(row["alias"])
            if output_path is None or not output_path.is_file():
                generation["cross_case_leak_status"] = "FAIL"
                generation["cross_case_leaks"] = ["output_missing"]
                generation["status"] = "FAIL"
                continue
            visible = extract_docx_text(output_path)
            leaks = [
                alias for alias in aliases if alias != own_alias and alias in visible
            ]
            own_alias_present = own_alias in visible
            generation["cross_case_leak_status"] = (
                "PASS" if own_alias_present and not leaks else "FAIL"
            )
            generation["cross_case_leaks"] = leaks
            generation["own_alias_present"] = own_alias_present
            if generation["cross_case_leak_status"] != "PASS":
                generation["status"] = "FAIL"

    failures: list[str] = []
    for row in rows:
        if (
            not row["auto_detection"]["detected"]
            or row["auto_detection"]["project_type"] != "lung_588_pdl1"
        ):
            failures.append(
                f"{row['alias']}: lung588 structural identity was not detected"
            )
        if row["targeted_drug_count"] != row["expected_targeted_drug_count"]:
            failures.append(
                f"{row['alias']}: targeted drug rows differ from the exact-event contract"
            )
        expected_targeted_genes = set(row.get("expected_targeted_drug_genes") or [])
        if expected_targeted_genes and set(row["targeted_drug_genes"]) != (
            expected_targeted_genes
        ):
            failures.append(
                f"{row['alias']}: targeted drug genes differ from the historical final"
            )
        if row["targeted_drug_intro_count"] != row["targeted_drug_count"]:
            failures.append(
                f"{row['alias']}: targeted drug introductions are incomplete"
            )
        if row["alias"] == "CASE-LUNG-C" and not expected_targeted_genes <= set(
            row["part3_drug_analysis_genes"]
        ):
            failures.append(
                "CASE-LUNG-C: Part-3 drug analysis does not cover all target genes"
            )
        if row["immune_positive_count"] != row["expected_immune_positive_count"]:
            failures.append(
                f"{row['alias']}: positive immune rows differ from the exact-event contract"
            )
        if row["immune_negative_count"] != row["expected_immune_negative_count"]:
            failures.append(
                f"{row['alias']}: negative immune rows differ from the exact-event contract"
            )
        if row["lung_guideline_row_count"] != 10:
            failures.append(f"{row['alias']}: lung guideline table is not 10 rows")
        if row["lung_guideline_results"].get("ALK") != "未见变异":
            failures.append(f"{row['alias']}: ALK guideline wording differs")
        if row["alias"] == "CASE-LUNG-C":
            if row["lung_guideline_results"].get("BRAF") != "c.1799T>A,p.V600E":
                failures.append("CASE-LUNG-C: BRAF guideline wording differs")
            if row["lung_guideline_results"].get("ERBB2") != "c.1979G>A,p.G660D":
                failures.append("CASE-LUNG-C: ERBB2 guideline wording differs")
        if row["immune_fixed_row_counts"] != {
            "positive": 15,
            "negative": 12,
            "hyperprogression": 8,
        }:
            failures.append(f"{row['alias']}: immune fixed-table rows differ")
        if row["chemotherapy_row_counts"] != {
            "prediction": 27,
            "regimen": 22,
            "dosage": 11,
        }:
            failures.append(f"{row['alias']}: chemotherapy table rows differ")
        expected_ratings = CURRENT_CHEMOTHERAPY_RATINGS.get(row["alias"])
        if expected_ratings is not None:
            expected_base_rows = [
                {
                    "drug": drug,
                    "genes": CURRENT_CHEMOTHERAPY_GENES[drug],
                    "efficacy": ratings[0],
                    "toxicity": ratings[1],
                }
                for drug, ratings in expected_ratings.items()
            ]
            if row["chemotherapy_base_rows"] != expected_base_rows:
                failures.append(
                    f"{row['alias']}: chemotherapy base rows violate corrected source-fidelity contract"
                )
            if row["chemotherapy_summary_text"] != (
                CURRENT_CHEMOTHERAPY_SUMMARIES[row["alias"]]
            ):
                failures.append(
                    f"{row['alias']}: chemotherapy summary violates corrected source-fidelity contract"
                )
            expected_irinotecan = CURRENT_IRINOTECAN_SAFETY[row["alias"]]
            if row["irinotecan_safety_rows"] != [expected_irinotecan]:
                failures.append(
                    f"{row['alias']}: irinotecan safety violates genotype-only review contract"
                )
        if any(row["chemotherapy_detail_quality"].values()):
            failures.append(
                f"{row['alias']}: chemotherapy appendix display normalization failed"
            )
        if row["alias"] in {"CASE-LUNG-B", "CASE-LUNG-C"} and not (
            8 <= row["cisplatin_detail_row_count"] <= 10
        ):
            failures.append(f"{row['alias']}: cisplatin detail row count is unexpected")
        if row["tmb_reference"] != 10:
            failures.append(f"{row['alias']}: tissue TMB reference must be 10")
        forbidden_crc_genes = {"FBXW7", "NF1", "NRAS", "SMAD4", "SMARCA4", "TCF7L2"}
        if forbidden_crc_genes & set(row["undetected_gene_names"]):
            failures.append(f"{row['alias']}: CRC undetected genes leaked into lung report")
        if row["alias"] == "CASE-LUNG-C" and not {"MLH1", "PMS2"} <= set(
            row["undetected_gene_names"]
        ):
            failures.append("CASE-LUNG-C: MLH1/PMS2 are missing from undetected genes")
        preview = row["web_preview"]
        if preview["targeted_drug_count"] != row["expected_targeted_drug_count"]:
            failures.append(
                f"{row['alias']}: web preview targeted-drug count is inconsistent"
            )
        if preview["drug_related_variant_count"] != row["expected_targeted_drug_count"]:
            failures.append(
                f"{row['alias']}: web preview drug-related variant count is inconsistent"
            )
        positive_result = str(preview["immune"].get("positive") or "")
        negative_result = str(preview["immune"].get("negative") or "")
        if row["expected_immune_positive_count"] and not positive_result.startswith(
            "检出（"
        ):
            failures.append(f"{row['alias']}: web preview lost positive immune hits")
        if not row["expected_immune_positive_count"] and positive_result != "未检出":
            failures.append(
                f"{row['alias']}: web preview positive immune zero is wrong"
            )
        if row["expected_immune_negative_count"] and not negative_result.startswith(
            "检出（"
        ):
            failures.append(f"{row['alias']}: web preview lost negative immune hits")
        if not row["expected_immune_negative_count"] and negative_result != "未检出":
            failures.append(
                f"{row['alias']}: web preview negative immune zero is wrong"
            )
        if row["biomarker_contract_status"] != "PASS":
            failures.append(f"{row['alias']}: biomarker contract failed")
        if row["pdl1_product_contract_status"] != "PASS":
            failures.append(f"{row['alias']}: PD-L1 product contract blocked")
        contract_status = row["context_contract"]["status"]
        if contract_status not in {"PASS", "NOT_APPLICABLE"}:
            failures.append(f"{row['alias']}: context contract failed")
        render_status = row["report_generation"]["status"]
        if render_output_dir is not None and row["context_contract"]["contract_id"]:
            if render_status != "PASS":
                failures.append(f"{row['alias']}: rendered report gate failed")
            candidate_contract = (
                row["report_generation"].get("review_candidate_contract") or {}
            )
            if candidate_contract.get("status") != "PASS":
                failures.append(f"{row['alias']}: review-candidate contract failed")
            if row["report_generation"].get("cross_case_leak_status") != "PASS":
                failures.append(f"{row['alias']}: cross-case value leakage detected")
    return {
        "schema_version": "1.0",
        "panel_id": "lung_588_pdl1",
        "status": "FAIL" if failures else "PASS",
        "source_commit": _source_revision(),
        "case_count": len(rows),
        "uat_readiness": _build_uat_readiness(rows),
        "cases": rows,
        "failures": failures,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input-dir", type=Path, required=True)
    parser.add_argument(
        "--output",
        type=Path,
        default=ROOT / ".work" / "lung588_real_input_audit" / "validation.json",
    )
    parser.add_argument(
        "--render-output-dir",
        type=Path,
        help="Optionally render all registered historical cases with full visual QA.",
    )
    parser.add_argument("--render-dpi", type=int, default=120)
    args = parser.parse_args()
    render_output_dir = (
        args.render_output_dir.resolve() if args.render_output_dir is not None else None
    )
    if render_output_dir is not None:
        render_output_dir.mkdir(parents=True, exist_ok=True)
    payload = validate_inputs(
        args.input_dir.resolve(),
        render_output_dir=render_output_dir,
        render_dpi=args.render_dpi,
    )
    output = args.output.resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(
        json.dumps(
            {
                "status": payload["status"],
                "case_count": payload["case_count"],
                "variant_counts": {
                    row["alias"]: len(row["variant_rows"]) for row in payload["cases"]
                },
                "contract_statuses": {
                    row["alias"]: row["context_contract"]["status"]
                    for row in payload["cases"]
                },
                "pdl1_product_statuses": {
                    row["alias"]: row["pdl1_product_contract_status"]
                    for row in payload["cases"]
                },
                "uat_readiness": payload["uat_readiness"],
                "render_statuses": {
                    row["alias"]: row["report_generation"]["status"]
                    for row in payload["cases"]
                },
                "output": str(output),
            },
            ensure_ascii=False,
        )
    )
    return 0 if payload["status"] == "PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
