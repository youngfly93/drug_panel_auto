#!/usr/bin/env python3
"""Apply a structural variable map to a scrubbed golden-template seed DOCX."""

from __future__ import annotations

import argparse
import copy
import json
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Mapping

import yaml
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.build_golden_template_seed import (  # noqa: E402
    DEFAULT_PROTECTED_TOKENS,
    count_tokens_in_zip,
    sha256_file,
)


def load_variable_map(path: str | Path) -> dict[str, Any]:
    spec_path = Path(path)
    with spec_path.open("r", encoding="utf-8") as handle:
        raw = yaml.safe_load(handle) or {}
    if not isinstance(raw, Mapping):
        raise ValueError(f"variable map must be a dict: {spec_path}")
    return dict(raw)


def variable_placeholder(variable: str) -> str:
    variable = str(variable or "").strip()
    if not variable:
        raise ValueError("variable name cannot be empty")
    return "{{ " + variable + " }}"


def expression_placeholder(expression: str) -> str:
    expression = str(expression or "").strip()
    if not expression:
        raise ValueError("placeholder expression cannot be empty")
    return "{{ " + expression + " }}"


def loop_field_expression(alias: str, field: str) -> str:
    alias = str(alias or "row").strip()
    field = str(field or "").strip()
    if not field:
        raise ValueError("loop field cannot be empty")
    if field.isidentifier() and field.isascii():
        return f"{alias}.{field}"
    return f'{alias}[{json.dumps(field, ensure_ascii=False)}]'


def replace_paragraph_text(paragraph, text: str) -> None:
    """Replace paragraph text while preserving the first run's formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def replace_cell_text(cell, text: str) -> None:
    if cell.paragraphs:
        replace_paragraph_text(cell.paragraphs[0], text)
        for paragraph in list(cell.paragraphs[1:]):
            remove_paragraph(paragraph)
    else:
        cell.text = text


def table_row_count(table) -> int:
    return len(table._tbl.tr_lst)


def clone_row_xml(table, row_index: int):
    return copy.deepcopy(table.rows[row_index]._tr)


def insert_row_xml(table, index: int, row_xml) -> None:
    rows = table._tbl.tr_lst
    if index >= len(rows):
        table._tbl.append(row_xml)
    else:
        rows[index].addprevious(row_xml)


def remove_row_at(table, index: int) -> None:
    table._tbl.remove(table._tbl.tr_lst[index])


def clear_row_text(row) -> None:
    for cell in row.cells:
        replace_cell_text(cell, "")


def set_loop_marker_row(row, marker: str) -> None:
    clear_row_text(row)
    replace_cell_text(row.cells[0], marker)


def render_loop_column(alias: str, column: Any) -> str:
    if isinstance(column, str):
        return expression_placeholder(loop_field_expression(alias, column))
    if isinstance(column, Mapping):
        if "expr" in column:
            return expression_placeholder(str(column["expr"]))
        if "field" in column:
            return expression_placeholder(loop_field_expression(alias, str(column["field"])))
        if "text" in column:
            return str(column["text"])
    raise ValueError(f"invalid loop column spec: {column!r}")


def apply_table_loop(doc: Document, item: Mapping[str, Any]) -> dict[str, Any]:
    table_index = int(item["table"])
    table = doc.tables[table_index]
    alias = str(item.get("alias") or "row").strip() or "row"
    collection = str(item["collection"]).strip()
    if not collection:
        raise ValueError("table loop collection cannot be empty")

    template_row_index = int(item["template_row"])
    insert_at = int(item.get("insert_at", template_row_index))
    remove_from = int(item.get("remove_from", template_row_index))
    remove_to_raw = item.get("remove_to", "end")
    original_row_count = table_row_count(table)
    remove_to = (
        original_row_count - 1
        if str(remove_to_raw).lower() == "end"
        else int(remove_to_raw)
    )
    if not (0 <= template_row_index < original_row_count):
        raise IndexError(f"template_row out of range: {template_row_index}")
    if remove_from < 0 or remove_to >= original_row_count or remove_from > remove_to:
        raise IndexError(
            f"invalid remove range {remove_from}:{remove_to} for table {table_index}"
        )

    template_row_xml = clone_row_xml(table, template_row_index)
    start_row_xml = clone_row_xml(table, template_row_index)
    end_row_xml = clone_row_xml(table, template_row_index)

    for _ in range(remove_to - remove_from + 1):
        remove_row_at(table, remove_from)

    insert_row_xml(table, insert_at, end_row_xml)
    insert_row_xml(table, insert_at, template_row_xml)
    insert_row_xml(table, insert_at, start_row_xml)

    start_row = table.rows[insert_at]
    data_row = table.rows[insert_at + 1]
    end_row = table.rows[insert_at + 2]
    set_loop_marker_row(start_row, f"{{%tr for {alias} in {collection} %}}")
    set_loop_marker_row(end_row, "{%tr endfor %}")

    columns = list(item.get("columns") or [])
    if len(columns) > len(data_row.cells):
        raise ValueError(
            f"table loop {item.get('id')!r} declares {len(columns)} columns, "
            f"but table row has {len(data_row.cells)} cells"
        )
    for col_index, column in enumerate(columns):
        replace_cell_text(
            data_row.cells[col_index],
            render_loop_column(alias, column),
        )
    for col_index in range(len(columns), len(data_row.cells)):
        replace_cell_text(data_row.cells[col_index], "")

    return {
        "id": item.get("id"),
        "type": "table_loop",
        "table": table_index,
        "collection": collection,
        "alias": alias,
        "template_row": template_row_index,
        "remove_from": remove_from,
        "remove_to": remove_to,
        "insert_at": insert_at,
        "columns": columns,
        "removed_rows": remove_to - remove_from + 1,
    }


def find_heading_index(paragraphs: list, heading: str) -> int:
    target = str(heading or "").strip()
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == target:
            return index
    raise ValueError(f"heading not found: {target!r}")


def paragraph_after_heading(paragraphs: list, heading: str, nonempty_offset: int):
    start_index = find_heading_index(paragraphs, heading)
    seen = 0
    for paragraph in paragraphs[start_index + 1 :]:
        if not paragraph.text.strip():
            continue
        seen += 1
        if seen == int(nonempty_offset):
            return paragraph
    raise ValueError(
        f"paragraph offset {nonempty_offset!r} after heading {heading!r} not found"
    )


def paragraph_matching_text(paragraphs: list, item: Mapping[str, Any]):
    exact = str(item.get("exact") or "").strip()
    contains = str(item.get("contains") or "").strip()
    if not exact and not contains:
        raise ValueError("paragraph template requires exact or contains")
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if exact and text == exact:
            return paragraph
        if contains and contains in text:
            return paragraph
    needle = exact or contains
    raise ValueError(f"paragraph matching {needle!r} not found")


def variableize_docx(
    source: str | Path,
    output: str | Path,
    variable_map: Mapping[str, Any],
    *,
    protected_tokens: tuple[str, ...] = DEFAULT_PROTECTED_TOKENS,
) -> dict[str, Any]:
    source_path = Path(source)
    output_path = Path(output)
    if source_path.suffix.lower() != ".docx":
        raise ValueError(f"source must be .docx: {source_path}")
    if output_path.suffix.lower() != ".docx":
        raise ValueError(f"output must be .docx: {output_path}")
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(source_path))
    operations: list[dict[str, Any]] = []

    for item in variable_map.get("cell_variables") or []:
        if not isinstance(item, Mapping):
            continue
        table_index = int(item["table"])
        row_index = int(item["row"])
        col_index = int(item["col"])
        variable = str(item["variable"])
        placeholder = variable_placeholder(variable)
        cell = doc.tables[table_index].rows[row_index].cells[col_index]
        before = cell.text
        replace_cell_text(cell, placeholder)
        operations.append(
            {
                "id": item.get("id"),
                "type": "cell",
                "table": table_index,
                "row": row_index,
                "col": col_index,
                "variable": variable,
                "before_preview": before[:80],
            }
        )

    paragraphs = list(doc.paragraphs)
    for item in variable_map.get("paragraph_variables") or []:
        if not isinstance(item, Mapping):
            continue
        heading = str(item["after_heading"])
        nonempty_offset = int(item["nonempty_offset"])
        variable = str(item["variable"])
        placeholder = variable_placeholder(variable)
        paragraph = paragraph_after_heading(paragraphs, heading, nonempty_offset)
        before = paragraph.text
        replace_paragraph_text(paragraph, placeholder)
        operations.append(
            {
                "id": item.get("id"),
                "type": "paragraph",
                "after_heading": heading,
                "nonempty_offset": nonempty_offset,
                "variable": variable,
                "before_preview": before[:120],
            }
        )

    for item in variable_map.get("paragraph_templates") or []:
        if not isinstance(item, Mapping):
            continue
        replacement = str(item.get("text") or "")
        paragraph = paragraph_matching_text(paragraphs, item)
        before = paragraph.text
        replace_paragraph_text(paragraph, replacement)
        operations.append(
            {
                "id": item.get("id"),
                "type": "paragraph_template",
                "match": {
                    "exact": item.get("exact"),
                    "contains": item.get("contains"),
                },
                "before_preview": before[:120],
                "after_preview": replacement[:120],
            }
        )

    for item in variable_map.get("table_loops") or []:
        if not isinstance(item, Mapping):
            continue
        operations.append(apply_table_loop(doc, item))

    doc.save(str(output_path))
    residual_counts = count_tokens_in_zip(output_path, protected_tokens)
    manifest = {
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "source": str(source_path),
        "output": str(output_path),
        "source_sha256": sha256_file(source_path),
        "output_sha256": sha256_file(output_path),
        "template_id": variable_map.get("template_id"),
        "operation_count": len(operations),
        "operations": operations,
        "protected_token_residual_counts": residual_counts,
        "success": not any(count > 0 for count in residual_counts.values()),
    }
    manifest_path = output_path.with_suffix(".variableize.json")
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    manifest["manifest"] = str(manifest_path)
    return manifest


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", type=Path, help="Scrubbed seed DOCX")
    parser.add_argument(
        "--map",
        required=True,
        type=Path,
        help="Variable map YAML",
    )
    parser.add_argument(
        "--output",
        required=True,
        type=Path,
        help="Output variableized DOCX",
    )
    parser.add_argument(
        "--protected-token",
        action="append",
        default=[],
        help="Additional token that must not remain in the output.",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)
    variable_map = load_variable_map(args.map)
    protected_tokens = tuple(
        dict.fromkeys([*DEFAULT_PROTECTED_TOKENS, *args.protected_token])
    )
    manifest = variableize_docx(
        args.source,
        args.output,
        variable_map,
        protected_tokens=protected_tokens,
    )
    print(json.dumps(manifest, ensure_ascii=False, indent=2))
    return 0 if manifest["success"] else 2


if __name__ == "__main__":
    raise SystemExit(main())
