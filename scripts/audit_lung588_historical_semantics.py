#!/usr/bin/env python3
# 步骤: 73 肺癌588历史终版语义审计
# 上游: .work/lung588_historical_refs/repaired/CASE-*.reference.repaired.docx、context_contracts/
# 输出: .work/lung588_historical_semantic_audit/semantic_inventory.json、*.tsv
# 种子: 无（确定性 Word 结构抽取与脱敏合同比对）
"""Extract de-identified clinical semantics from historical lung588 reports.

Historical final reports are evidence of an old display contract, not current
medical truth.  This tool therefore emits *candidates* with
``runtime_eligible=false``.  It never reads the cover table, never emits patient
names/sample identifiers, and rejects its own output if common PII tokens leak.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
from collections import Counter
from pathlib import Path
from typing import Any, Iterable

import yaml
from docx import Document


ALIAS_RE = re.compile(r"^CASE-LUNG-[A-Z0-9-]+$")
GENE_RE = re.compile(r"^[A-Z][A-Z0-9/-]{1,30}$")
C_HGVS_RE = re.compile(r"c\.[0-9_*+?>A-Za-z-]+")
P_HGVS_RE = re.compile(r"p\.[A-Za-z0-9_*?=.-]+")
EVENT_HEADER_RE = re.compile(
    r"^(?P<gene>[A-Z][A-Z0-9/-]+)："
    r"(?P<c_hgvs>c\.[^，,\s]+)"
    r"(?:[，,]\s*(?P<p_hgvs>p\.[^突变\s]+))?"
    r"突变"
)
DRUG_LEVEL_RE = re.compile(r"(?P<drug>[^（）]+?)（(?P<level>[ABCD])）")
PMID_RE = re.compile(r"(?<!\d)([1-9]\d{6,8})(?!\d)")
NCT_RE = re.compile(r"\bNCT\d{8}\b", re.I)
CTR_RE = re.compile(r"\bCTR\d{8,11}\b", re.I)
PII_PATTERNS = {
    "sample_id": re.compile(r"\b(?:LZ|LW)\d{5,}\b", re.I),
    "name_label": re.compile(r"姓名[:：]?"),
    "report_number": re.compile(r"报告编号[:：]?"),
    "sender": re.compile(r"送检者[:：]?"),
}
CANCER_SCOPE_TERMS = (
    "非小细胞肺癌",
    "小细胞肺癌",
    "肺腺癌",
    "肺癌",
    "结直肠癌",
    "结肠癌",
    "直肠癌",
    "黑色素瘤",
    "乳腺癌",
    "卵巢癌",
    "前列腺癌",
    "胰腺癌",
    "胃癌",
    "肝细胞癌",
    "胆道癌",
    "甲状腺癌",
    "胶质瘤",
    "实体瘤",
)
DIRECTION_HEADINGS = {
    "潜在获益靶向/免疫药物解析": "benefit",
    "潜在获益靶向药物解析": "benefit",
    "潜在负相关靶向/免疫药物解析": "caution",
    "潜在负相关靶向药物解析": "caution",
    "研究结论不一致的靶向药物": "conflicting",
}


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _clean(value: Any) -> str:
    return " ".join(str(value or "").split())


def _parse_binding(value: str) -> tuple[str, Path]:
    if "=" not in value:
        raise argparse.ArgumentTypeError("binding must use ALIAS=PATH")
    alias, path = value.split("=", 1)
    alias = alias.strip().upper()
    if not ALIAS_RE.fullmatch(alias):
        raise argparse.ArgumentTypeError(f"invalid de-identified alias: {alias}")
    return alias, Path(path)


def _table_rows(table: Any) -> list[list[str]]:
    return [[_clean(cell.text) for cell in row.cells] for row in table.rows]


def _find_table(document: Document, predicate: Any) -> list[list[str]]:
    # Table 0 is the cover/patient table in both historical references.  It is
    # excluded unconditionally so an unexpected header cannot leak identity.
    for table in document.tables[1:]:
        rows = _table_rows(table)
        if rows and predicate(rows):
            return rows
    raise ValueError("required historical semantic table was not found")


def _variant_parts(text: str) -> tuple[str, str]:
    c_match = C_HGVS_RE.search(text)
    p_match = P_HGVS_RE.search(text)
    return (
        c_match.group(0) if c_match else "",
        p_match.group(0) if p_match else "",
    )


def _parse_drug_cell(text: str) -> tuple[list[dict[str, str]], str]:
    text = _clean(text)
    if text in {"", "-", "--", "—"}:
        return [], ""
    candidates: list[dict[str, str]] = []
    spans: list[tuple[int, int]] = []
    for match in DRUG_LEVEL_RE.finditer(text):
        drug = _clean(match.group("drug")).strip("，,；;、")
        if drug:
            candidates.append(
                {
                    "drug": drug,
                    "historical_evidence_level": match.group("level"),
                    "review_status": "needs_review",
                    "runtime_eligible": False,
                    "secondary_review_status": "pending_report_group_review",
                }
            )
        spans.append(match.span())
    remaining = text
    for start, end in reversed(spans):
        remaining = f"{remaining[:start]} {remaining[end:]}"
    remaining = _clean(remaining).strip("，,；;、")
    return candidates, remaining


def _targeted_summary(document: Document) -> list[dict[str, Any]]:
    rows = _find_table(
        document,
        lambda value: (
            len(value[0]) >= 4
            and value[0][0] == "基因"
            and "突变位点" in value[0][1]
            and "潜在获益" in value[0][2]
        ),
    )
    records: list[dict[str, Any]] = []
    for row in rows[1:]:
        if len(row) < 4:
            continue
        gene = _clean(row[0]).upper()
        if not GENE_RE.fullmatch(gene):
            continue
        c_hgvs, p_hgvs = _variant_parts(row[1])
        benefit, benefit_ungraded = _parse_drug_cell(row[2])
        caution, caution_ungraded = _parse_drug_cell(row[3])
        records.append(
            {
                "gene": gene,
                "c_hgvs": c_hgvs,
                "p_hgvs": p_hgvs,
                "historical_variant_site": _clean(row[1]),
                "benefit_candidates": benefit,
                "caution_candidates": caution,
                "ungraded_benefit_text": benefit_ungraded,
                "ungraded_caution_text": caution_ungraded,
            }
        )
    return records


def _biomarker_and_immune_table(document: Document) -> list[list[str]]:
    return _find_table(
        document,
        lambda rows: (
            any(row and row[0] == "免疫正相关基因" for row in rows)
            and any(row and "肿瘤突变负荷" in row[0] for row in rows)
        ),
    )


def _biomarkers(document: Document) -> dict[str, str]:
    rows = _biomarker_and_immune_table(document)
    mapping = {
        "肿瘤突变负荷（TMB）": "tmb",
        "微卫星不稳定性（MSI）": "msi",
        "PD-L1表达": "pdl1",
        "HLA-I分型": "hla_i",
    }
    result: dict[str, str] = {}
    for row in rows[1:]:
        if len(row) >= 2 and row[0] in mapping:
            result[mapping[row[0]]] = _clean(row[1])
    return result


def _parse_immune_events(text: str) -> list[dict[str, str]]:
    text = re.sub(r"^检出（\d+个）\s*", "", _clean(text))
    markers = list(re.finditer(r"([A-Z][A-Z0-9/-]+)：", text))
    events: list[dict[str, str]] = []
    for index, marker in enumerate(markers):
        gene = marker.group(1)
        start = marker.end()
        end = markers[index + 1].start() if index + 1 < len(markers) else len(text)
        variant_text = _clean(text[start:end])
        c_hgvs, p_hgvs = _variant_parts(variant_text)
        events.append(
            {
                "gene": gene,
                "c_hgvs": c_hgvs,
                "p_hgvs": p_hgvs,
                "historical_direction": "",
                "review_status": "needs_review",
                "runtime_eligible": False,
                "secondary_review_status": "pending_report_group_review",
            }
        )
    return events


def _immune_observations(document: Document) -> list[dict[str, str]]:
    rows = _biomarker_and_immune_table(document)
    direction_by_label = {
        "免疫正相关基因": "positive",
        "免疫负相关基因": "negative",
        "免疫超进展相关基因": "hyperprogression",
    }
    observations: list[dict[str, str]] = []
    for row in rows[1:]:
        if len(row) < 2 or row[0] not in direction_by_label:
            continue
        if not row[1].startswith("检出（"):
            continue
        direction = direction_by_label[row[0]]
        for event in _parse_immune_events(row[1]):
            event["historical_direction"] = direction
            observations.append(event)
    return observations


def _pgx_observations(document: Document) -> list[dict[str, str]]:
    observations: list[dict[str, str]] = []
    required = {"基因", "检测位点", "等级", "检测结果", "用药提示"}
    for table in document.tables[1:]:
        rows = _table_rows(table)
        if not rows or not required.issubset(set(rows[0])):
            continue
        header = rows[0]
        index = {name: header.index(name) for name in required}
        drug_column = next(
            (
                position
                for position, value in enumerate(header)
                if value not in required
            ),
            0,
        )
        default_drug = header[drug_column]
        for row in rows[1:]:
            if len(row) <= max(index.values()):
                continue
            gene = _clean(row[index["基因"]]).upper()
            if not GENE_RE.fullmatch(gene):
                continue
            drug = _clean(row[drug_column]) or default_drug
            observations.append(
                {
                    "drug": drug,
                    "gene": gene,
                    "selector": _clean(row[index["检测位点"]]),
                    "historical_level": _clean(row[index["等级"]]),
                    "observed_genotype": _clean(row[index["检测结果"]]),
                    "historical_interpretation": _clean(row[index["用药提示"]]),
                    "review_status": "needs_review",
                    "runtime_eligible": False,
                    "secondary_review_status": "pending_pgx_review",
                }
            )
    return observations


def _part3_events(document: Document) -> list[dict[str, Any]]:
    paragraphs = [_clean(paragraph.text) for paragraph in document.paragraphs]
    start = next(
        (
            index
            for index, text in enumerate(paragraphs)
            if text == "靶向/免疫药物用药提示解析"
        ),
        None,
    )
    if start is None:
        return []

    direction = ""
    current: dict[str, Any] | None = None
    records: list[dict[str, Any]] = []

    def finish() -> None:
        nonlocal current
        if current is None:
            return
        block = current.pop("_paragraphs")
        joined = "\n".join(block)
        current["source_refs"] = {
            "pmids": sorted(set(PMID_RE.findall(joined))),
            "clinical_trials": sorted(
                {match.upper() for match in NCT_RE.findall(joined)}
            ),
            "china_trials": sorted({match.upper() for match in CTR_RE.findall(joined)}),
        }
        current["cancer_scope_terms"] = [
            term for term in CANCER_SCOPE_TERMS if term in joined
        ]
        current["contains_lung_specific_text"] = any(
            term in joined for term in ("非小细胞肺癌", "肺腺癌", "肺癌")
        )
        current["drug_groups"] = [
            block[index - 1]
            for index, text in enumerate(block)
            if index > 0
            and text == "基因变异与药物关联分析："
            and not block[index - 1].startswith("该样本")
        ]
        current["reference_granularity"] = "event_block_not_drug_specific"
        current["review_status"] = "needs_review"
        current["runtime_eligible"] = False
        current["secondary_review_status"] = "pending_report_group_review"
        records.append(current)
        current = None

    for text in paragraphs[start + 1 :]:
        if not text:
            continue
        if text.startswith("3. 阅读说明"):
            finish()
            break
        if text in DIRECTION_HEADINGS:
            finish()
            direction = DIRECTION_HEADINGS[text]
            continue
        match = EVENT_HEADER_RE.match(text)
        if match:
            finish()
            current = {
                "gene": match.group("gene"),
                "c_hgvs": match.group("c_hgvs"),
                "p_hgvs": match.group("p_hgvs") or "",
                "historical_direction": direction or "unspecified",
                "_paragraphs": [text],
            }
            continue
        if current is not None:
            current["_paragraphs"].append(text)
    finish()
    return records


def _load_contract_variants(path: Path) -> list[dict[str, str]]:
    raw = yaml.safe_load(path.read_text(encoding="utf-8"))
    rows = (((raw or {}).get("tables") or {}).get("all_variants") or {}).get(
        "rows"
    ) or []
    variants: list[dict[str, str]] = []
    for row in rows:
        match = row.get("match") or {}
        expect = row.get("expect") or {}
        variants.append(
            {
                "gene": _clean(((match.get("gene") or {}).get("equals"))).upper(),
                "c_hgvs": _clean(((match.get("cHGVS") or {}).get("equals"))),
                "p_hgvs": _clean(((expect.get("pHGVS") or {}).get("equals"))),
            }
        )
    return variants


def _event_key(row: dict[str, Any]) -> tuple[str, str]:
    return (_clean(row.get("gene")).upper(), _clean(row.get("c_hgvs")))


def _flatten_targeted(
    alias: str,
    targeted: list[dict[str, Any]],
    part3: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    part3_by_event: dict[tuple[str, str], list[dict[str, Any]]] = {}
    for event in part3:
        part3_by_event.setdefault(_event_key(event), []).append(event)

    flattened: list[dict[str, Any]] = []
    for event in targeted:
        event_part3 = part3_by_event.get(_event_key(event), [])
        refs = {
            "pmids": sorted(
                {
                    item
                    for block in event_part3
                    for item in block["source_refs"]["pmids"]
                }
            ),
            "clinical_trials": sorted(
                {
                    item
                    for block in event_part3
                    for item in block["source_refs"]["clinical_trials"]
                }
            ),
            "china_trials": sorted(
                {
                    item
                    for block in event_part3
                    for item in block["source_refs"]["china_trials"]
                }
            ),
        }
        cancer_terms = sorted(
            {term for block in event_part3 for term in block["cancer_scope_terms"]}
        )
        for direction, key in (
            ("benefit", "benefit_candidates"),
            ("caution", "caution_candidates"),
        ):
            for candidate in event[key]:
                flattened.append(
                    {
                        "case_alias": alias,
                        "gene": event["gene"],
                        "c_hgvs": event["c_hgvs"],
                        "p_hgvs": event["p_hgvs"],
                        "direction": direction,
                        "drug": candidate["drug"],
                        "historical_evidence_level": candidate[
                            "historical_evidence_level"
                        ],
                        "historical_source_refs": refs,
                        "historical_cancer_scope_terms": cancer_terms,
                        "reference_granularity": "event_block_not_drug_specific",
                        "review_status": "needs_review",
                        "runtime_eligible": False,
                        "secondary_review_status": "pending_report_group_review",
                    }
                )
    return flattened


def _write_tsv(
    path: Path,
    rows: Iterable[dict[str, Any]],
    fieldnames: list[str],
) -> None:
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, delimiter="\t")
        writer.writeheader()
        for row in rows:
            writer.writerow(
                {
                    field: (
                        json.dumps(row.get(field), ensure_ascii=False, sort_keys=True)
                        if isinstance(row.get(field), (list, dict))
                        else row.get(field, "")
                    )
                    for field in fieldnames
                }
            )


def _assert_no_pii(payloads: Iterable[str]) -> None:
    combined = "\n".join(payloads)
    findings = [
        label for label, pattern in PII_PATTERNS.items() if pattern.search(combined)
    ]
    if findings:
        raise RuntimeError(
            "de-identified audit output failed PII guard: " + ", ".join(findings)
        )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--reference",
        action="append",
        type=_parse_binding,
        required=True,
        help="de-identified reference binding: CASE-LUNG-X=/path/to/repaired.docx",
    )
    parser.add_argument(
        "--contract",
        action="append",
        type=_parse_binding,
        default=[],
        help="de-identified contract binding: CASE-LUNG-X=/path/to/contract.yaml",
    )
    parser.add_argument("--output-dir", type=Path, required=True)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    contracts = dict(args.contract)
    cases: dict[str, dict[str, Any]] = {}
    targeted_rows: list[dict[str, Any]] = []
    immune_rows: list[dict[str, Any]] = []
    pgx_rows: list[dict[str, Any]] = []

    for alias, path in sorted(args.reference):
        if not path.is_file():
            raise FileNotFoundError(f"reference not found for {alias}")
        document = Document(path)
        targeted = _targeted_summary(document)
        part3 = _part3_events(document)
        immune = _immune_observations(document)
        pgx = _pgx_observations(document)
        biomarkers = _biomarkers(document)
        targeted_flat = _flatten_targeted(alias, targeted, part3)
        for row in immune:
            row["case_alias"] = alias
        for row in pgx:
            row["case_alias"] = alias

        contract_check: dict[str, Any] = {"status": "not_supplied"}
        if alias in contracts:
            contract_variants = _load_contract_variants(contracts[alias])
            targeted_keys = {_event_key(row) for row in targeted}
            contract_keys = {_event_key(row) for row in contract_variants}
            contract_check = {
                "status": "PASS" if targeted_keys <= contract_keys else "FAIL",
                "all_variant_count": len(contract_variants),
                "targeted_event_count": len(targeted_keys),
                "non_targeted_events": [
                    row
                    for row in contract_variants
                    if _event_key(row) not in targeted_keys
                ],
                "unexpected_targeted_events": [
                    row for row in targeted if _event_key(row) not in contract_keys
                ],
            }

        cases[alias] = {
            "reference_sha256": _sha256(path),
            "biomarkers": biomarkers,
            "targeted_summary": targeted,
            "immune_observations": immune,
            "pgx_observation_count": len(pgx),
            "pgx_drug_count": len({row["drug"] for row in pgx}),
            "pgx_level_counts": dict(
                sorted(Counter(row["historical_level"] for row in pgx).items())
            ),
            "part3_events": part3,
            "contract_check": contract_check,
        }
        targeted_rows.extend(targeted_flat)
        immune_rows.extend(immune)
        pgx_rows.extend(pgx)

    event_cases: dict[tuple[str, str], list[str]] = {}
    for alias, case in cases.items():
        for row in case["targeted_summary"]:
            event_cases.setdefault(_event_key(row), []).append(alias)

    inventory = {
        "schema_version": 1,
        "status": (
            "PASS"
            if all(
                case["contract_check"]["status"] in {"PASS", "not_supplied"}
                for case in cases.values()
            )
            else "FAIL"
        ),
        "purpose": (
            "historical_display_contract_inventory_only; "
            "no candidate is runtime eligible"
        ),
        "cases": cases,
        "cross_case": {
            "shared_targeted_events": [
                {
                    "gene": gene,
                    "c_hgvs": c_hgvs,
                    "case_count": len(aliases),
                }
                for (gene, c_hgvs), aliases in sorted(event_cases.items())
                if len(aliases) > 1
            ],
            "targeted_candidate_count": len(targeted_rows),
            "immune_observation_count": len(immune_rows),
            "pgx_observation_count": len(pgx_rows),
        },
    }

    args.output_dir.mkdir(parents=True, exist_ok=True)
    inventory_path = args.output_dir / "semantic_inventory.json"
    targeted_path = args.output_dir / "targeted_candidate_review.tsv"
    immune_path = args.output_dir / "immune_candidate_review.tsv"
    pgx_path = args.output_dir / "pgx_observations.tsv"

    inventory_text = (
        json.dumps(
            inventory,
            ensure_ascii=False,
            indent=2,
            sort_keys=True,
        )
        + "\n"
    )
    _write_tsv(
        targeted_path,
        targeted_rows,
        [
            "case_alias",
            "gene",
            "c_hgvs",
            "p_hgvs",
            "direction",
            "drug",
            "historical_evidence_level",
            "historical_source_refs",
            "historical_cancer_scope_terms",
            "reference_granularity",
            "review_status",
            "runtime_eligible",
            "secondary_review_status",
        ],
    )
    _write_tsv(
        immune_path,
        immune_rows,
        [
            "case_alias",
            "gene",
            "c_hgvs",
            "p_hgvs",
            "historical_direction",
            "review_status",
            "runtime_eligible",
            "secondary_review_status",
        ],
    )
    _write_tsv(
        pgx_path,
        pgx_rows,
        [
            "case_alias",
            "drug",
            "gene",
            "selector",
            "historical_level",
            "observed_genotype",
            "historical_interpretation",
            "review_status",
            "runtime_eligible",
            "secondary_review_status",
        ],
    )

    tsv_texts = [
        path.read_text(encoding="utf-8")
        for path in (targeted_path, immune_path, pgx_path)
    ]
    _assert_no_pii([inventory_text, *tsv_texts])
    inventory_path.write_text(inventory_text, encoding="utf-8")

    print(
        json.dumps(
            {
                "status": inventory["status"],
                "cases": len(cases),
                "targeted_candidates": len(targeted_rows),
                "immune_observations": len(immune_rows),
                "pgx_observations": len(pgx_rows),
                "output": inventory_path.name,
            },
            ensure_ascii=False,
        )
    )
    return 0 if inventory["status"] == "PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
