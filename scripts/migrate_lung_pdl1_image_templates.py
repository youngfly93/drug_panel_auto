# 步骤: 肺癌PD-L1模板病例图片占位迁移
# 上游: panels/lung_329_pdl1/templates、panels/lung_588_pdl1/templates
# 输出: 原位更新两份肺癌生产模板（PD-L1表+病例图片+图注）
# 种子: 无
"""Idempotently align lung-panel PD-L1 sections with report-group feedback."""

from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
TEMPLATES = (
    ROOT / "panels/lung_329_pdl1/templates/lung_329_pdl1_golden_template_v2.docx",
    ROOT / "panels/lung_588_pdl1/templates/lung_588_pdl1_golden_template_v0.docx",
)
MARKER = "__PDL1_CASE_IMAGE__"
CAPTION = "图1. 免疫组化：PD-L1"
EXACT_RULE_BOUNDARY = (
    "本项目仅输出经审核的肺癌精确事件规则；须同时满足基因、转录本、"
    "位点及治疗上下文，且交付前须完成病例级复核。泛基因、跨癌种及"
    "未审核规则不展示。"
)
TEXT_REPLACEMENTS = {
    "分析样本基因变异并报告分子事件；患者级用药提示仅在肺癌专属规则完成审核后输出。": (
        "分析样本基因变异并报告分子事件；患者级用药提示仅由经审核的肺癌精确事件规则"
        "在治疗上下文符合时输出，并须病例级复核。"
    ),
    "肺癌专属治疗知识和事件级药物规则尚未完成报告组二审及病例级验收，当前版本不输出患者级用药结论。": EXACT_RULE_BOUNDARY,
    "肺癌专属治疗知识和事件级药物规则当前未启用，本报告不输出患者级用药结论；最终治疗方案须由专业医师结合病理类型、疾病分期及现行诊疗规范综合判断。": EXACT_RULE_BOUNDARY,
}


def _element_text(element) -> str:
    return "".join(node.text or "" for node in element.xpath(".//w:t")).strip()


def _new_paragraph(body, index: int, text: str, *, keep_next: bool) -> Paragraph:
    element = OxmlElement("w:p")
    body.insert(index, element)
    paragraph = Paragraph(element, body)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = keep_next
    run = paragraph.add_run(text)
    run.font.name = "微软雅黑"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    return paragraph


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace paragraph text while retaining its first run formatting."""

    first_rpr = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        first_rpr = deepcopy(paragraph.runs[0]._r.rPr)
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(text)
    if first_rpr is not None:
        run._r.insert(0, first_rpr)


def _align_runtime_boundary_text(doc: Document) -> bool:
    changed = False
    paragraphs = list(doc.paragraphs)
    paragraphs.extend(
        paragraph
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    for paragraph in paragraphs:
        replacement = TEXT_REPLACEMENTS.get((paragraph.text or "").strip())
        if replacement:
            _replace_paragraph_text(paragraph, replacement)
            changed = True

    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        headers = [cell.text.replace("\n", "").strip() for cell in table.rows[0].cells]
        if headers == ["检测基因", "本癌种相关治疗药物", "临床提示", "检测结果"]:
            values = ["见精确事件", "见靶向用药表", EXACT_RULE_BOUNDARY, "需病例级复核"]
        elif headers == ["药物名称", "相关基因", "药物介绍"]:
            values = [
                "未启用",
                "-",
                "本节不从泛基因、跨癌种或未审核证据扩展其它上市药物。",
            ]
        else:
            continue
        for cell, value in zip(table.rows[1].cells, values):
            if cell.text.strip() != value:
                _replace_paragraph_text(cell.paragraphs[0], value)
                changed = True
    return changed


def migrate(path: Path, *, write: bool = True) -> bool:
    doc = Document(path)
    changed = _align_runtime_boundary_text(doc)
    body = doc._element.body
    elements = list(body)

    heading_index = next(
        (
            index
            for index, element in enumerate(elements)
            if element.tag == qn("w:p")
            and _element_text(element).startswith("3.2 PD-L1表达检测结果")
        ),
        None,
    )
    if heading_index is None:
        raise RuntimeError(f"{path.name}: missing PD-L1 section heading")

    table_index = next(
        (
            index
            for index in range(heading_index + 1, len(elements))
            if elements[index].tag == qn("w:tbl")
            and "TPS" in _element_text(elements[index])
            and "CPS" in _element_text(elements[index])
        ),
        None,
    )
    if table_index is None:
        raise RuntimeError(f"{path.name}: missing PD-L1 result table")

    next_heading_index = next(
        (
            index
            for index in range(table_index + 1, len(elements))
            if elements[index].tag == qn("w:p")
            and _element_text(elements[index]).startswith("3.3")
        ),
        None,
    )
    if next_heading_index is None:
        raise RuntimeError(f"{path.name}: missing section following PD-L1")

    existing_block = elements[table_index + 1 : next_heading_index]
    existing_text = [_element_text(element) for element in existing_block]
    if existing_text != [MARKER, CAPTION]:
        # Keep the result table; remove all legacy method/provenance/explanation
        # paragraphs below it. The report-group contract for this section is only
        # TPS, CPS, result, a case image and its caption.
        for element in existing_block:
            body.remove(element)

        current = list(body)
        table_index = current.index(elements[table_index])
        _new_paragraph(body, table_index + 1, MARKER, keep_next=True)
        _new_paragraph(body, table_index + 2, CAPTION, keep_next=False)
        changed = True

    # TPS is a percentage; CPS is a unitless score.
    for node in body.xpath(".//w:t"):
        if (node.text or "").strip() == "{{ pdl1_tps }}":
            node.text = "{{ pdl1_tps }}%"
            changed = True

    if changed and write:
        doc.save(path)
    return changed


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--check",
        action="store_true",
        help="only verify that both templates are already migrated",
    )
    args = parser.parse_args()

    pending: list[Path] = []
    for path in TEMPLATES:
        changed = migrate(path, write=not args.check)
        if changed:
            pending.append(path)
            action = "would update" if args.check else "updated"
        else:
            action = "up to date"
        print(f"{action}: {path.relative_to(ROOT)}")

    if args.check and pending:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
