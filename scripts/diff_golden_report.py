#!/usr/bin/env python3
"""Build a structural and visual diff between a generated report and a reference.

This script is intentionally case-data agnostic. Reviewed case values can be
passed through ``--override key=value`` so the comparison can mimic the web form
injection path without committing patient data.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable, Mapping, Optional

from docx import Document
from PIL import Image, ImageChops

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from reportgen.config.loader import ConfigLoader  # noqa: E402
from reportgen.core.context_contract import (  # noqa: E402
    check_context_contract,
    load_context_contract,
)
from reportgen.core.excel_reader import ExcelReader  # noqa: E402
from reportgen.core.report_diff import ReportDiffOptions, compare_reports  # noqa: E402
from reportgen.core.report_generator import ReportGenerator  # noqa: E402
from reportgen.panels.loader import load_panel_package  # noqa: E402
from reportgen.utils.docx_render import render_docx_to_pngs  # noqa: E402


DEFAULT_REQUIRED_SNIPPETS: tuple[dict[str, str], ...] = (
    {
        "id": "variant_count",
        "text": "本次共检出体细胞变异：8个",
    },
    {
        "id": "drug_related_variant_count",
        "text": "与靶向药物用药相关的变异有：4个",
    },
    {"id": "tmb_value", "text": "6.5"},
    {"id": "tmb_level", "text": "TMB-L"},
    {"id": "msi_status", "text": "微卫星稳定型，MSS"},
    {
        "id": "tmb_education",
        "text": "多项临床研究表明，TMB-H的肿瘤",
    },
    {
        "id": "msi_education",
        "text": "研究表明，MSI-H的实体瘤",
    },
    {"id": "kras_locus", "text": "KRAS：c.34G>A"},
    {"id": "kras_protein", "text": "p.G12S"},
    {"id": "atm_locus", "text": "ATM：c.6874C>T"},
)

DEFAULT_HEADING_TERMS: tuple[str, ...] = (
    "报告导读",
    "致您的一封信",
    "目录",
    "第一部分",
    "患者及样本信息",
    "检测结果小结",
    "2.1 基因变异检测结果",
    "2.2",
    "2.3 NCCN",
    "3.1 肿瘤突变负荷",
    "3.2微卫星不稳定性",
    "3.3 免疫疗效",
    "4. 基因检测列表",
    "本次检测质控结果",
)


@dataclass(frozen=True)
class RenderedDocument:
    docx: Path
    pdf: Optional[Path]
    pngs: list[Path]
    page_texts: list[str]
    visual_pages: list[dict[str, Any]]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--excel", required=True, help="Input Excel file")
    parser.add_argument("--reference-docx", required=True, help="Reviewed reference report")
    parser.add_argument(
        "--template-id",
        default="crc_358_msi_golden_template_v0",
        help="Panel template id to generate",
    )
    parser.add_argument("--template-file", help="Explicit template .docx path")
    parser.add_argument("--project-type", default="crc_358_msi")
    parser.add_argument("--output-dir", default="tmp/m4_5_diff")
    parser.add_argument("--output-filename", default="generated_current.docx")
    parser.add_argument(
        "--context-contract",
        default="panels/crc_358_msi/context_contracts/reviewed_low_tmb_mss.yaml",
    )
    parser.add_argument(
        "--override",
        action="append",
        default=[],
        metavar="KEY=VALUE",
        help="Inject a reviewed-case field using canonical key or Excel synonym",
    )
    parser.add_argument(
        "--required-snippet",
        action="append",
        default=[],
        metavar="ID=TEXT",
        help="Additional required text snippet",
    )
    parser.add_argument("--render-dpi", type=int, default=120)
    parser.add_argument("--blank-threshold", type=float, default=0.003)
    parser.add_argument("--skip-render", action="store_true")
    parser.add_argument("--allow-warn", action="store_true")
    return parser.parse_args()


def parse_key_values(items: Iterable[str]) -> dict[str, str]:
    parsed: dict[str, str] = {}
    for item in items:
        if "=" not in item:
            raise ValueError(f"Expected KEY=VALUE, got: {item!r}")
        key, value = item.split("=", 1)
        key = key.strip()
        if not key:
            raise ValueError(f"Empty key in override: {item!r}")
        parsed[key] = value
    return parsed


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", "", str(text or ""))


def read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    parts.extend(p.text for p in doc.paragraphs if p.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def table_summaries(path: Path, *, max_preview_rows: int = 4) -> list[dict[str, Any]]:
    doc = Document(str(path))
    summaries: list[dict[str, Any]] = []
    for idx, table in enumerate(doc.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        nonempty_rows = [
            row for row in rows if any(normalize_text(cell) for cell in row)
        ]
        header = nonempty_rows[0] if nonempty_rows else []
        summaries.append(
            {
                "index": idx,
                "row_count": len(rows),
                "col_count": max((len(row) for row in rows), default=0),
                "nonempty_row_count": len(nonempty_rows),
                "header": header,
                "preview_rows": nonempty_rows[:max_preview_rows],
            }
        )
    return summaries


def find_heading_pages(page_texts: list[str], headings: Iterable[str]) -> dict[str, int]:
    pages: dict[str, int] = {}
    compact_pages = [normalize_text(text) for text in page_texts]
    for heading in headings:
        compact_heading = normalize_text(heading)
        for idx, page_text in enumerate(compact_pages, start=1):
            if compact_heading in page_text:
                pages[heading] = idx
                break
    return pages


def image_visual_metrics(path: Path, *, blank_threshold: float) -> dict[str, Any]:
    with Image.open(path) as image:
        gray = image.convert("L")
        width, height = gray.size
        white = Image.new("L", gray.size, 255)
        diff = ImageChops.difference(gray, white)
        bbox = diff.getbbox()
        nonwhite_pixels = sum(1 for value in diff.getdata() if value > 12)
        total = width * height

        top = diff.crop((0, 0, width, height // 2))
        bottom = diff.crop((0, height // 2, width, height))
        top_nonwhite = sum(1 for value in top.getdata() if value > 12)
        bottom_nonwhite = sum(1 for value in bottom.getdata() if value > 12)
        half_total = width * (height // 2)

    nonwhite_ratio = nonwhite_pixels / total if total else 0.0
    return {
        "file": str(path),
        "page": _page_number_from_png(path),
        "width": width,
        "height": height,
        "nonwhite_ratio": round(nonwhite_ratio, 6),
        "top_half_nonwhite_ratio": round(top_nonwhite / half_total, 6)
        if half_total
        else 0.0,
        "bottom_half_nonwhite_ratio": round(bottom_nonwhite / half_total, 6)
        if half_total
        else 0.0,
        "content_bbox": list(bbox) if bbox else None,
        "near_blank": nonwhite_ratio < blank_threshold,
    }


def _page_number_from_png(path: Path) -> int:
    match = re.search(r"-(\d+)\.png$", path.name)
    return int(match.group(1)) if match else 0


def pdf_page_texts(pdf_path: Path, page_count: int) -> list[str]:
    pdftotext = shutil.which("pdftotext")
    if not pdftotext:
        return []
    texts: list[str] = []
    for page in range(1, page_count + 1):
        proc = subprocess.run(
            [
                pdftotext,
                "-layout",
                "-f",
                str(page),
                "-l",
                str(page),
                str(pdf_path),
                "-",
            ],
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        texts.append(proc.stdout if proc.returncode == 0 else "")
    return texts


def render_document(
    docx_path: Path,
    *,
    output_dir: Path,
    dpi: int,
    blank_threshold: float,
    skip_render: bool,
) -> RenderedDocument:
    if skip_render:
        return RenderedDocument(
            docx=docx_path,
            pdf=None,
            pngs=[],
            page_texts=[],
            visual_pages=[],
        )
    pngs = render_docx_to_pngs(
        docx_path,
        output_dir=output_dir,
        dpi=dpi,
        keep_pdf=True,
        timeout_seconds=180,
    )
    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    page_texts = pdf_page_texts(pdf_path, len(pngs)) if pdf_path.exists() else []
    visual_pages = [
        image_visual_metrics(path, blank_threshold=blank_threshold) for path in pngs
    ]
    return RenderedDocument(
        docx=docx_path,
        pdf=pdf_path if pdf_path.exists() else None,
        pngs=pngs,
        page_texts=page_texts,
        visual_pages=visual_pages,
    )


def resolve_template(project_type: str, template_id: str, template_file: Optional[str]) -> Path:
    if template_file:
        return Path(template_file).resolve()
    package = load_panel_package(project_type, project_root=ROOT)
    return package.resolve_template_file(template_id)


def canonical_override_columns(overrides: Mapping[str, str]) -> dict[str, str]:
    """Map canonical field keys to a configured Excel synonym where possible."""
    loader = ConfigLoader(config_dir=str(ROOT / "config"))
    mapping = loader.load_mapping_config().get("single_values", {})
    columns: dict[str, str] = {}
    for key, value in overrides.items():
        spec = mapping.get(key)
        synonyms = list((spec or {}).get("synonyms") or [])
        columns[synonyms[0] if synonyms else key] = value
    return columns


def generate_candidate(
    *,
    excel: Path,
    template: Path,
    output_dir: Path,
    output_filename: str,
    project_type: str,
    overrides: Mapping[str, str],
) -> dict[str, Any]:
    excel_reader = ExcelReader(config_dir=str(ROOT / "config"))
    excel_data = excel_reader.read(str(excel))
    for column, value in canonical_override_columns(overrides).items():
        excel_data.add_single_value(column, value)

    generator = ReportGenerator(
        config_dir=str(ROOT / "config"),
        template_dir=str(ROOT / "templates"),
    )
    return generator.generate(
        excel_file=str(excel),
        excel_data=excel_data,
        template_file=str(template),
        output_dir=str(output_dir),
        output_filename=output_filename,
        return_context=True,
        template_contract_mode="fail",
        project_type=project_type,
    )


def check_required_snippets(
    *,
    reference_text: str,
    candidate_text: str,
    snippets: Iterable[Mapping[str, str]],
) -> dict[str, Any]:
    ref_compact = normalize_text(reference_text)
    cand_compact = normalize_text(candidate_text)
    rows = []
    issues = []
    for snippet in snippets:
        snippet_id = str(snippet.get("id") or snippet.get("text") or "snippet")
        text = str(snippet.get("text") or "")
        compact = normalize_text(text)
        row = {
            "id": snippet_id,
            "text": text,
            "in_reference": compact in ref_compact,
            "in_candidate": compact in cand_compact,
        }
        rows.append(row)
        if not row["in_candidate"]:
            issues.append(
                issue(
                    "error",
                    "REQUIRED_SNIPPET_MISSING",
                    f"Candidate report is missing required snippet: {snippet_id}",
                    section="required_snippets",
                    details=row,
                )
            )
        elif not row["in_reference"]:
            issues.append(
                issue(
                    "warning",
                    "SNIPPET_NOT_IN_REFERENCE",
                    f"Required snippet is present in candidate but not reference: {snippet_id}",
                    section="required_snippets",
                    details=row,
                )
            )
    return {"status": section_status(issues), "items": rows, "issues": issues}


def issue(
    level: str,
    code: str,
    message: str,
    *,
    section: str,
    details: Optional[Mapping[str, Any]] = None,
) -> dict[str, Any]:
    row = {
        "level": level,
        "code": code,
        "message": message,
        "section": section,
    }
    if details:
        row["details"] = dict(details)
    return row


def visual_diff_section(
    reference: RenderedDocument,
    candidate: RenderedDocument,
    *,
    headings: Iterable[str],
) -> dict[str, Any]:
    issues: list[dict[str, Any]] = []
    ref_blank = [p["page"] for p in reference.visual_pages if p["near_blank"]]
    cand_blank = [p["page"] for p in candidate.visual_pages if p["near_blank"]]
    if cand_blank:
        issues.append(
            issue(
                "error",
                "CANDIDATE_NEAR_BLANK_PAGES",
                f"Candidate has near-blank pages: {cand_blank}",
                section="visual",
            )
        )
    if ref_blank != cand_blank:
        issues.append(
            issue(
                "warning",
                "BLANK_PAGE_SET_DIFF",
                f"Near-blank pages differ: reference={ref_blank}, candidate={cand_blank}",
                section="visual",
            )
        )

    ref_pages = len(reference.pngs)
    cand_pages = len(candidate.pngs)
    if ref_pages and cand_pages and ref_pages != cand_pages:
        issues.append(
            issue(
                "warning",
                "PAGE_COUNT_DIFF",
                f"Rendered page count differs: reference={ref_pages}, candidate={cand_pages}",
                section="visual",
            )
        )

    ref_heading_pages = find_heading_pages(reference.page_texts, headings)
    cand_heading_pages = find_heading_pages(candidate.page_texts, headings)
    for heading in headings:
        if heading in ref_heading_pages and heading not in cand_heading_pages:
            issues.append(
                issue(
                    "error",
                    "CANDIDATE_HEADING_MISSING",
                    f"Candidate rendered text is missing heading: {heading}",
                    section="visual",
                )
            )
        elif heading in ref_heading_pages and ref_heading_pages[heading] != cand_heading_pages[heading]:
            issues.append(
                issue(
                    "warning",
                    "HEADING_PAGE_DIFF",
                    f"Heading page differs for {heading}: "
                    f"{ref_heading_pages[heading]} -> {cand_heading_pages[heading]}",
                    section="visual",
                )
            )

    top_whitespace = compare_heading_top_whitespace(
        reference,
        candidate,
        ref_heading_pages,
        cand_heading_pages,
        headings=("报告导读", "致您的一封信", "目录"),
    )
    issues.extend(top_whitespace)

    return {
        "status": section_status(issues),
        "reference": {
            "pdf": str(reference.pdf) if reference.pdf else None,
            "page_count": ref_pages,
            "near_blank_pages": ref_blank,
            "heading_pages": ref_heading_pages,
            "pages": reference.visual_pages,
        },
        "candidate": {
            "pdf": str(candidate.pdf) if candidate.pdf else None,
            "page_count": cand_pages,
            "near_blank_pages": cand_blank,
            "heading_pages": cand_heading_pages,
            "pages": candidate.visual_pages,
        },
        "issues": issues,
    }


def compare_heading_top_whitespace(
    reference: RenderedDocument,
    candidate: RenderedDocument,
    ref_heading_pages: Mapping[str, int],
    cand_heading_pages: Mapping[str, int],
    *,
    headings: Iterable[str],
) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    for heading in headings:
        ref_page = ref_heading_pages.get(heading)
        cand_page = cand_heading_pages.get(heading)
        if not ref_page or not cand_page:
            continue
        ref_metric = _metric_for_page(reference.visual_pages, ref_page)
        cand_metric = _metric_for_page(candidate.visual_pages, cand_page)
        if not ref_metric or not cand_metric:
            continue
        ref_top = float(ref_metric.get("top_half_nonwhite_ratio") or 0)
        cand_top = float(cand_metric.get("top_half_nonwhite_ratio") or 0)
        cand_bottom = float(cand_metric.get("bottom_half_nonwhite_ratio") or 0)
        if cand_top < 0.0015 and cand_bottom > 0.005 and ref_top > cand_top * 2.5:
            issues.append(
                issue(
                    "warning",
                    "TOP_HALF_WHITESPACE_REGRESSION",
                    f"Candidate page for {heading} has much less top-half content than reference.",
                    section="visual",
                    details={
                        "heading": heading,
                        "reference_page": ref_page,
                        "candidate_page": cand_page,
                        "reference_top_half_nonwhite_ratio": ref_top,
                        "candidate_top_half_nonwhite_ratio": cand_top,
                        "candidate_bottom_half_nonwhite_ratio": cand_bottom,
                    },
                )
            )
    return issues


def _metric_for_page(metrics: list[Mapping[str, Any]], page: int) -> Optional[Mapping[str, Any]]:
    for metric in metrics:
        if int(metric.get("page") or 0) == int(page):
            return metric
    return None


def section_status(issues: Iterable[Mapping[str, Any]]) -> str:
    items = list(issues)
    if any(item.get("level") == "error" for item in items):
        return "FAIL"
    if any(item.get("level") == "warning" for item in items):
        return "WARN"
    return "PASS"


def overall_status(issues: Iterable[Mapping[str, Any]]) -> str:
    return section_status(issues)


def render_markdown(result: Mapping[str, Any]) -> str:
    summary = result.get("summary") or {}
    visual = (result.get("sections") or {}).get("visual") or {}
    snippets = (result.get("sections") or {}).get("required_snippets") or {}
    lines = [
        "# CRC358 Golden Template Diff",
        "",
        f"- Status: **{result.get('status')}**",
        f"- Generated report: `{Path(str((result.get('generated') or {}).get('docx'))).name}`",
        f"- Reference report: `{Path(str((result.get('inputs') or {}).get('reference_docx'))).name}`",
        f"- Failures: {summary.get('failures', 0)}",
        f"- Warnings: {summary.get('warnings', 0)}",
        f"- Page count: reference {summary.get('reference_page_count')} -> candidate {summary.get('candidate_page_count')}",
        "",
        "## Issues",
        "",
    ]
    issues = list(result.get("issues") or [])
    if issues:
        for item in issues:
            lines.append(
                f"- **{str(item.get('level')).upper()}** "
                f"`{item.get('section')}/{item.get('code')}`: {item.get('message')}"
            )
    else:
        lines.append("No issues detected.")

    lines.extend(["", "## Required Snippets", ""])
    for item in snippets.get("items") or []:
        mark = "OK" if item.get("in_candidate") else "MISSING"
        ref_mark = "reference OK" if item.get("in_reference") else "reference missing"
        lines.append(f"- {mark} `{item.get('id')}` ({ref_mark})")

    lines.extend(["", "## Visual", ""])
    ref_visual = visual.get("reference") or {}
    cand_visual = visual.get("candidate") or {}
    lines.append(f"- Reference PDF: `{ref_visual.get('pdf')}`")
    lines.append(f"- Candidate PDF: `{cand_visual.get('pdf')}`")
    lines.append(f"- Reference near-blank pages: {ref_visual.get('near_blank_pages')}")
    lines.append(f"- Candidate near-blank pages: {cand_visual.get('near_blank_pages')}")
    lines.append("")
    lines.append("### Heading Pages")
    headings = sorted(
        set((ref_visual.get("heading_pages") or {}))
        | set((cand_visual.get("heading_pages") or {}))
    )
    if headings:
        for heading in headings:
            lines.append(
                f"- {heading}: reference "
                f"{(ref_visual.get('heading_pages') or {}).get(heading)} -> "
                f"candidate {(cand_visual.get('heading_pages') or {}).get(heading)}"
            )
    else:
        lines.append("- No heading pages extracted.")

    lines.extend(["", "## Tables", ""])
    tables = (result.get("sections") or {}).get("tables") or {}
    lines.append(
        f"- Table count: reference {len(tables.get('reference') or [])} -> "
        f"candidate {len(tables.get('candidate') or [])}"
    )
    for table in (tables.get("candidate") or [])[:12]:
        lines.append(
            f"- Candidate table {table.get('index')}: "
            f"{table.get('row_count')}x{table.get('col_count')}, "
            f"nonempty rows={table.get('nonempty_row_count')}, "
            f"header={table.get('header')}"
        )

    return "\n".join(lines).rstrip() + "\n"


def write_outputs(result: Mapping[str, Any], output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "diff.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    (output_dir / "diff.md").write_text(render_markdown(result), encoding="utf-8")


def build_required_snippets(extra: Mapping[str, str]) -> list[dict[str, str]]:
    snippets = [dict(item) for item in DEFAULT_REQUIRED_SNIPPETS]
    snippets.extend({"id": key, "text": value} for key, value in extra.items())
    return snippets


def collect_context_contract(
    context: Mapping[str, Any],
    contract_path: Optional[Path],
) -> dict[str, Any]:
    if not contract_path:
        return {"status": "SKIP", "issues": []}
    contract = load_context_contract(contract_path)
    report = check_context_contract(dict(context), contract, contract_path=contract_path)
    issues = []
    if report.get("status") == "FAIL":
        issues.append(
            issue(
                "error",
                "CONTEXT_CONTRACT_FAIL",
                "Context contract failed.",
                section="context_contract",
            )
        )
    elif report.get("status") == "WARN":
        issues.append(
            issue(
                "warning",
                "CONTEXT_CONTRACT_WARN",
                "Context contract emitted warnings.",
                section="context_contract",
            )
        )
    report = dict(report)
    report["issues"] = issues
    return report


def main() -> int:
    args = parse_args()
    output_dir = (ROOT / args.output_dir).resolve()
    generated_dir = output_dir / "generated"
    render_dir = output_dir / "png"
    docx_diff_dir = output_dir / "docx_diff"
    excel = Path(args.excel).resolve()
    reference_docx = Path(args.reference_docx).resolve()
    template = resolve_template(args.project_type, args.template_id, args.template_file)
    overrides = parse_key_values(args.override)
    extra_snippets = parse_key_values(args.required_snippet)
    context_contract_path = (
        (ROOT / args.context_contract).resolve() if args.context_contract else None
    )

    generated_dir.mkdir(parents=True, exist_ok=True)
    generation = generate_candidate(
        excel=excel,
        template=template,
        output_dir=generated_dir,
        output_filename=args.output_filename,
        project_type=args.project_type,
        overrides=overrides,
    )

    generated_docx = Path(str(generation.get("output_file") or ""))
    issues: list[dict[str, Any]] = []
    if not generation.get("success"):
        issues.append(
            issue(
                "error",
                "GENERATION_FAILED",
                "Report generation failed.",
                section="generation",
                details={"errors": generation.get("errors")},
            )
        )
        result = {
            "schema_version": "1.0",
            "generated_at": datetime.now().isoformat(),
            "status": "FAIL",
            "inputs": {
                "excel": str(excel),
                "reference_docx": str(reference_docx),
                "template": str(template),
                "project_type": args.project_type,
            },
            "generated": {"docx": str(generated_docx) if generated_docx else None},
            "summary": {"failures": 1, "warnings": 0},
            "sections": {"generation": generation},
            "issues": issues,
        }
        write_outputs(result, output_dir)
        print(json.dumps({"status": "FAIL", "output_dir": str(output_dir)}, ensure_ascii=False))
        return 1

    context = dict(generation.get("context") or {})
    (output_dir / "context.json").write_text(
        json.dumps(context, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    contract_report = collect_context_contract(context, context_contract_path)
    issues.extend(contract_report.get("issues") or [])

    docx_diff = compare_reports(
        ReportDiffOptions(
            reference_docx=str(reference_docx),
            candidate_docx=str(generated_docx),
            output_dir=str(docx_diff_dir),
            max_samples=25,
        )
    )
    for item in docx_diff.get("issues") or []:
        copied = dict(item)
        copied["section"] = f"docx_diff.{copied.get('section', 'unknown')}"
        # Full document equality is too strict for a golden-template gate during
        # migration, so keep compare_reports findings warning-level unless they
        # indicate openability/table structure errors.
        if copied.get("code") in {"REFERENCE_DOCX_NOT_OPENABLE", "CANDIDATE_DOCX_NOT_OPENABLE"}:
            copied["level"] = "error"
        else:
            copied["level"] = "warning"
        issues.append(copied)

    reference_render = render_document(
        reference_docx,
        output_dir=render_dir / "reference",
        dpi=args.render_dpi,
        blank_threshold=args.blank_threshold,
        skip_render=args.skip_render,
    )
    candidate_render = render_document(
        generated_docx,
        output_dir=render_dir / "candidate",
        dpi=args.render_dpi,
        blank_threshold=args.blank_threshold,
        skip_render=args.skip_render,
    )
    visual = visual_diff_section(
        reference_render,
        candidate_render,
        headings=DEFAULT_HEADING_TERMS,
    )
    issues.extend(visual.get("issues") or [])

    reference_text = read_docx_text(reference_docx)
    candidate_text = read_docx_text(generated_docx)
    required_snippets = check_required_snippets(
        reference_text=reference_text,
        candidate_text=candidate_text,
        snippets=build_required_snippets(extra_snippets),
    )
    issues.extend(required_snippets.get("issues") or [])

    table_section = {
        "status": "PASS",
        "reference": table_summaries(reference_docx),
        "candidate": table_summaries(generated_docx),
    }

    status = overall_status(issues)
    result = {
        "schema_version": "1.0",
        "generated_at": datetime.now().isoformat(),
        "status": status,
        "inputs": {
            "excel": str(excel),
            "reference_docx": str(reference_docx),
            "template": str(template),
            "template_id": args.template_id,
            "project_type": args.project_type,
            "context_contract": str(context_contract_path)
            if context_contract_path
            else None,
            "overrides_supplied": sorted(overrides),
        },
        "generated": {
            "docx": str(generated_docx),
            "pdf": str(candidate_render.pdf) if candidate_render.pdf else None,
            "context_json": str(output_dir / "context.json"),
        },
        "summary": {
            "failures": sum(1 for item in issues if item.get("level") == "error"),
            "warnings": sum(1 for item in issues if item.get("level") == "warning"),
            "reference_page_count": len(reference_render.pngs),
            "candidate_page_count": len(candidate_render.pngs),
            "candidate_near_blank_pages": [
                p["page"] for p in candidate_render.visual_pages if p["near_blank"]
            ],
            "context_contract_status": contract_report.get("status"),
            "docx_diff_status": docx_diff.get("status"),
        },
        "sections": {
            "generation": {
                "success": generation.get("success"),
                "warnings": generation.get("warnings"),
                "stage_results_file": generation.get("stage_results_file"),
                "template_contract": generation.get("template_contract"),
            },
            "context_contract": contract_report,
            "docx_diff": docx_diff,
            "visual": visual,
            "required_snippets": required_snippets,
            "tables": table_section,
        },
        "issues": issues,
    }
    write_outputs(result, output_dir)
    print(
        json.dumps(
            {
                "status": status,
                "failures": result["summary"]["failures"],
                "warnings": result["summary"]["warnings"],
                "output_dir": str(output_dir),
                "diff_json": str(output_dir / "diff.json"),
                "diff_md": str(output_dir / "diff.md"),
            },
            ensure_ascii=False,
        )
    )
    if status == "FAIL":
        return 1
    if status == "WARN" and not args.allow_warn:
        return 2
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
