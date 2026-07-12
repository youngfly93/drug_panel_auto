#!/usr/bin/env python3
"""Generate and validate one synthetic CRC301 report per gap-overlay gene.

The validator deliberately uses one unique fake case per gene.  It verifies
that the CRC301-specific reviewed overlay reaches Part 3, that the synthetic
variant does not acquire a drug claim, and that the complete DOCX renders
without blank pages under the production LibreOffice renderer.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
from pathlib import Path
from typing import Any, Mapping

import yaml
from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from reportgen.core.golden_case import build_crc_301_msi_golden_excel  # noqa: E402
from reportgen.core.report_generator import ReportGenerator  # noqa: E402
from reportgen.panels.loader import load_panel_package  # noqa: E402
from reportgen.utils.docx_render import render_docx_to_pngs  # noqa: E402
from scripts.render_blank_page_check import page_nonwhite_ratio  # noqa: E402


DEFAULT_OVERLAY = ROOT / "panels/crc_301_msi/rules/reviewed_part3_knowledge.yaml"
DEFAULT_OUTPUT = ROOT / ".work/crc301_36_gene_validation"


def compact(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or ""))


def docx_text(path: Path) -> str:
    document = Document(path)
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def load_gene_rows(path: Path) -> list[dict[str, str]]:
    payload = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    rows = [dict(row) for row in payload.get("gene_sections") or []]
    genes = [str(row.get("gene") or "").strip().upper() for row in rows]
    if len(rows) != 36 or len(set(genes)) != 36 or any(not gene for gene in genes):
        raise ValueError(
            f"CRC301 gap overlay must contain 36 unique genes; rows={len(rows)} "
            f"unique={len(set(genes))}"
        )
    return rows


def build_case_workbook(path: Path, *, gene: str, index: int) -> str:
    sample_id = f"LZ9301{index:02d}"
    build_crc_301_msi_golden_excel(path)
    workbook = load_workbook(path)

    meta = workbook["Meta"]
    meta_columns = {str(cell.value): cell.column for cell in meta[1]}
    meta.cell(2, meta_columns["患者姓名"], f"批量验证{index:02d}")
    meta.cell(2, meta_columns["样本编号"], sample_id)
    meta.cell(2, meta_columns["报告编号"], f"VALIDATION-{sample_id}")

    variations = workbook["Variations"]
    columns = {str(cell.value): cell.column for cell in variations[1]}
    if variations.max_row > 1:
        variations.delete_rows(2, variations.max_row - 1)
    values = {
        "Gene_Symbol": gene,
        "Transcript": "NM_TEST.1",
        "Chr": "1",
        "ExIn_ID": "EX1",
        # Deliberately outside realistic transcript coordinates: this exercises
        # gene-level interpretation without accidentally matching a public
        # position-specific drug record.
        "cHGVS": "c.999999A>G",
        "pHGVS_S": "p.M333333V",
        "Freq(%)": round(5.0 + index / 10.0, 1),
        "Function": "Missense",
        "ExistInsmall301": 1,
        "ExistIn552": "Ⅲ类",
        "CLNSIG": "Uncertain_significance",
    }
    for name, value in values.items():
        variations.cell(2, columns[name], value)

    workbook.save(path)
    return sample_id


def rows_for_gene(rows: Any, gene: str) -> list[Mapping[str, Any]]:
    if not isinstance(rows, list):
        return []
    return [
        row
        for row in rows
        if isinstance(row, Mapping)
        and str(row.get("gene") or row.get("Gene") or "").strip().upper() == gene
    ]


def drug_overreach_rows(context: Mapping[str, Any], gene: str) -> list[dict[str, Any]]:
    hits: list[dict[str, Any]] = []
    for key in ("drug_analysis_sections", "drug_benefit_sections", "drug_caution_sections"):
        for row in rows_for_gene(context.get(key), gene):
            hits.append({"context_key": key, "row": dict(row)})

    for key in ("targeted_drug_tips", "variants_2_1"):
        for row in rows_for_gene(context.get(key), gene):
            drug_text = compact(
                " ".join(
                    str(row.get(field) or "")
                    for field in (
                        "drug_name",
                        "benefit_drugs",
                        "caution_drugs",
                        "positive_drugs",
                        "negative_drugs",
                    )
                )
            )
            if drug_text:
                hits.append({"context_key": key, "row": dict(row)})
    return hits


def render_metrics(docx_path: Path, render_dir: Path, dpi: int) -> dict[str, Any]:
    pngs = render_docx_to_pngs(
        docx_path,
        output_dir=render_dir,
        dpi=dpi,
        keep_pdf=False,
        timeout_seconds=180,
    )
    pages = [
        {
            "page": index,
            "nonwhite_ratio": round(page_nonwhite_ratio(path), 6),
        }
        for index, path in enumerate(pngs, start=1)
    ]
    near_blank = [row["page"] for row in pages if row["nonwhite_ratio"] < 0.003]
    return {
        "page_count": len(pages),
        "near_blank_pages": near_blank,
        "min_nonwhite_ratio": min(
            (row["nonwhite_ratio"] for row in pages), default=0.0
        ),
    }


def validate_case(
    *,
    row: Mapping[str, str],
    index: int,
    output_root: Path,
    generator: ReportGenerator,
    template_file: Path,
    dpi: int,
) -> dict[str, Any]:
    gene = str(row["gene"]).strip().upper()
    inputs = output_root / "inputs"
    reports = output_root / "reports"
    render_root = output_root / "render_tmp"
    inputs.mkdir(parents=True, exist_ok=True)
    reports.mkdir(parents=True, exist_ok=True)

    input_path = inputs / f"LZ9301{index:02d}_crc301_{gene}.xlsx"
    sample_id = build_case_workbook(input_path, gene=gene, index=index)
    output_name = f"crc301_gap_{index:02d}_{gene}.docx"

    generation = generator.generate(
        excel_file=str(input_path),
        template_file=str(template_file),
        output_dir=str(reports),
        output_filename=output_name,
        strict_mode=True,
        return_context=True,
        template_contract_mode="fail",
        project_type="crc_301_msi",
        project_name="结直肠癌301基因+MSI",
    )
    result: dict[str, Any] = {
        "index": index,
        "gene": gene,
        "sample_id": sample_id,
        "input_file": str(input_path),
        "output_file": str(generation.get("output_file") or ""),
        "generation_success": bool(generation.get("success")),
        "qa_status": generation.get("qa_status"),
        "errors": list(generation.get("errors") or []),
        "checks": {},
    }
    checks = result["checks"]
    checks["generation_success"] = result["generation_success"]
    checks["qa_pass"] = generation.get("qa_status") == "PASS"
    checks["panel_package_pass"] = (
        (generation.get("panel_package_validation") or {}).get("status") == "PASS"
    )
    checks["rule_provenance_pass"] = (
        (generation.get("rule_provenance") or {}).get("status") == "PASS"
    )

    output_file = Path(result["output_file"]) if result["output_file"] else None
    checks["output_exists"] = bool(output_file and output_file.exists())
    if output_file and output_file.exists():
        text = compact(docx_text(output_file))
        checks["gene_heading_present"] = gene in text
        checks["reviewed_intro_present"] = compact(row.get("intro")) in text
        checks["reviewed_analysis_present"] = compact(row.get("mutation_analysis")) in text
    else:
        checks["gene_heading_present"] = False
        checks["reviewed_intro_present"] = False
        checks["reviewed_analysis_present"] = False

    context = generation.get("context") or {}
    sections = rows_for_gene(context.get("gene_knowledge_sections"), gene)
    checks["single_gene_section"] = len(sections) == 1
    checks["part3_variant_coverage"] = (
        context.get("part3_expected_variant_count") == 1
        and context.get("part3_rendered_variant_count") == 1
        and context.get("part3_expected_variant_keys")
        == context.get("part3_rendered_variant_keys")
    )
    overreach = drug_overreach_rows(context, gene)
    result["drug_overreach_rows"] = overreach
    checks["no_drug_overreach"] = not overreach

    if output_file and output_file.exists():
        render_dir = render_root / f"{index:02d}_{gene}"
        try:
            result["render"] = render_metrics(output_file, render_dir, dpi)
            checks["render_has_pages"] = result["render"]["page_count"] > 0
            checks["no_blank_pages"] = not result["render"]["near_blank_pages"]
        except Exception as exc:  # surfaced as a blocking validation failure
            result["render"] = {"error": f"{type(exc).__name__}: {exc}"}
            checks["render_has_pages"] = False
            checks["no_blank_pages"] = False
        finally:
            shutil.rmtree(render_dir, ignore_errors=True)
    else:
        result["render"] = {"error": "output_not_generated"}
        checks["render_has_pages"] = False
        checks["no_blank_pages"] = False

    result["status"] = "PASS" if all(checks.values()) else "FAIL"
    result["failed_checks"] = [name for name, passed in checks.items() if not passed]
    return result


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--overlay", type=Path, default=DEFAULT_OVERLAY)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--dpi", type=int, default=100)
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    output_root = args.output_dir.resolve()
    output_root.mkdir(parents=True, exist_ok=True)
    rows = load_gene_rows(args.overlay.resolve())
    package = load_panel_package("crc_301_msi", project_root=ROOT)
    template_file = package.resolve_template_file()
    generator = ReportGenerator(config_dir=str(ROOT / "config"), log_level="ERROR")

    results: list[dict[str, Any]] = []
    summary_path = output_root / "validation_summary.json"
    for index, row in enumerate(rows, start=1):
        result = validate_case(
            row=row,
            index=index,
            output_root=output_root,
            generator=generator,
            template_file=template_file,
            dpi=args.dpi,
        )
        results.append(result)
        summary_path.write_text(
            json.dumps({"results": results}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        render = result.get("render") or {}
        print(
            f"[{index:02d}/36] {result['gene']}: {result['status']} "
            f"qa={result['qa_status']} pages={render.get('page_count', 0)} "
            f"blank={render.get('near_blank_pages', [])} "
            f"failed={result['failed_checks']}",
            flush=True,
        )

    failed = [row for row in results if row["status"] != "PASS"]
    summary = {
        "status": "PASS" if not failed else "FAIL",
        "panel_id": "crc_301_msi",
        "overlay": str(args.overlay.resolve()),
        "template": str(template_file),
        "case_count": len(results),
        "passed": len(results) - len(failed),
        "failed": len(failed),
        "failed_genes": [row["gene"] for row in failed],
        "total_rendered_pages": sum(
            int((row.get("render") or {}).get("page_count") or 0) for row in results
        ),
        "results": results,
    }
    summary_path.write_text(
        json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps({k: v for k, v in summary.items() if k != "results"}, ensure_ascii=False, indent=2))
    print(f"summary={summary_path}")
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
