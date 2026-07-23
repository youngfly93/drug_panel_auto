#!/usr/bin/env python3
# 步骤: 70 肺癌588工程模板构建
# 上游: panels/lung_329_pdl1/templates/lung_329_pdl1_golden_template_v1.docx
# 输出: panels/lung_588_pdl1/templates/lung_588_pdl1_golden_template_v0.docx
# 种子: 无（确定性文档变换）
"""Build the independent lung588 draft template without static treatment claims.

The v1 template came from a de-identified historical report.  Its dynamic
variant/PD-L1 scaffolding is useful, but several report-visible tables still
contain fixed historical treatment recommendations or patient genotypes.  This
script preserves v1 as an immutable migration input and produces a lung588 template
whose active report data comes from ReportGen context only.  Unreviewed clinical
sections fail closed with an explicit review notice.
"""

from __future__ import annotations

import argparse
import copy
import hashlib
import json
import os
import re
import tempfile
from pathlib import Path
from typing import Iterable, Sequence
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
import yaml

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = (
    ROOT
    / "panels"
    / "lung_329_pdl1"
    / "templates"
    / "lung_329_pdl1_golden_template_v1.docx"
)
DEFAULT_OUTPUT = (
    ROOT
    / "panels"
    / "lung_588_pdl1"
    / "templates"
    / "lung_588_pdl1_golden_template_v0.docx"
)
GENE_LIST_RULE = (
    ROOT
    / "panels"
    / "lung_588_pdl1"
    / "rules"
    / "knowledge_coverage.yaml"
)
REFERENCE_VARIANT_TEMPLATE = (
    ROOT
    / "panels"
    / "crc_301_msi"
    / "templates"
    / "crc_301_msi_golden_template_v1.docx"
)
EXPECTED_SOURCE_SHA256 = (
    "bbecf470feedab8ba20c4d84f7e4f8f0e2f0714f8f50818dc0419d72d2767e66"
)
EXPECTED_GENE_LIST_SHA256 = (
    "f9e6be05c954a4d3df97f031d453fe1f58ea0689290b11de3c173f4a0edf08f1"
)

REVIEW_NOTICE = (
    "肺癌专属治疗知识和事件级药物规则当前未启用，本报告不输出患者级用药结论；"
    "最终治疗方案须由专业医师结合病理类型、疾病分期及现行诊疗规范综合判断。"
)
CHEMOTHERAPY_NOTICE = (
    "化疗药物多态性与方案模块当前未启用，本报告不输出患者级化疗敏感性、"
    "毒性或剂量结论。"
)
APPENDIX_NOTICE = (
    "本节当前版本暂不提供；相关诊疗信息请咨询专业医师。"
)
IMMUNE_RESEARCH_NOTICE = (
    "本节仅展示研究性免疫相关基因检测事实，不得据此预测个体疗效、"
    "耐药或超进展，也不得自动生成治疗方案。"
)
IMMUNE_RESEARCH_RESULT = "研究性相关标志物，不能单独用于治疗决策。"
PDL1_IMAGE_NOTICE = (
    "附图：本病例未提供可追溯的PD-L1免疫组化图像，故不展示；"
    "TPS、CPS及结果判定须与原始检测记录核对。"
)
PDL1_ASSAY_PROVENANCE_MARKER = "{{ pdl1_assay_provenance }}"
PDL1_SOURCE_PROVENANCE_MARKER = "{{ pdl1_source_provenance }}"
REFERENCE_LINES = (
    "1. PMID:39375078. Chinese guidelines for molecular testing of non-small cell lung cancer "
    "(2024 edition). https://pubmed.ncbi.nlm.nih.gov/39375078/",
    "2. PMID:29355391. Updated Molecular Testing Guideline for the Selection of Lung Cancer "
    "Patients for Treatment With Targeted Tyrosine Kinase Inhibitors. "
    "https://pubmed.ncbi.nlm.nih.gov/29355391/",
    "3. National Cancer Institute. Non-Small Cell Lung Cancer Treatment (PDQ®)—Health "
    "Professional Version. https://www.cancer.gov/types/lung/hp/non-small-cell-lung-treatment-pdq",
)

FORBIDDEN_OUTPUT_TOKENS = (
    "肺癌329",
    "n=329",
    "非小细胞肺癌NCCN指南（2022 V3）",
    "非小细胞肺癌CSCO指南（2022）",
    "多西他赛单药方案",
    "呋喹替尼 Fruquintinib",
    "6TA/6TA",
    "2020年新版NSCLC指南",
    "结直肠癌",
    "5-Fu",
    "帕博利珠单抗",
    "纳武利尤单抗",
    "奥希替尼",
    "PBRM1基因突变的肾透明细胞癌患者免疫治疗获益",
    "TP53 mutations and hepatocellular carcinoma",
    "NCT02576444",
)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _compact(value: object) -> str:
    return re.sub(r"\s+", "", str(value or ""))


def _element_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def _replace_paragraph_text(paragraph, text: str) -> None:
    runs = list(paragraph.runs)
    if runs:
        target_index = next((i for i, run in enumerate(runs) if run.text), 0)
        for index, run in enumerate(runs):
            run.text = text if index == target_index else ""
    else:
        paragraph.add_run(text)


def _remove_paragraph(paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def _replace_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.text = text
        return
    _replace_paragraph_text(cell.paragraphs[0], text)
    for paragraph in list(cell.paragraphs[1:]):
        _remove_paragraph(paragraph)


def _body_paragraph(document: DocumentObject, text: str, *, prefer_last: bool = False):
    expected = _compact(text)
    matches = [p for p in document.paragraphs if _compact(p.text) == expected]
    if not matches:
        raise ValueError(f"paragraph not found: {text!r}")
    return matches[-1] if prefer_last else matches[0]


def _insert_paragraph_after(document: DocumentObject, anchor, text: str):
    paragraph = document.add_paragraph(text)
    anchor._p.addnext(paragraph._p)
    return paragraph


def _remove_explicit_page_break_before(
    document: DocumentObject,
    heading_text: str,
) -> None:
    """Remove a redundant page-break paragraph before a page-break heading."""
    heading = _body_paragraph(document, heading_text)
    previous = heading._p.getprevious()
    if previous is None:
        return
    page_breaks = [
        node
        for node in previous.iter(qn("w:br"))
        if node.get(qn("w:type")) == "page"
    ]
    if page_breaks:
        document.element.body.remove(previous)


def _add_letter_divider_note(document: DocumentObject) -> None:
    heading = _body_paragraph(document, "致您的一封信")
    note = _insert_paragraph_after(
        document,
        heading,
        "感谢您对本次检测的信任。本报告呈现经质量控制的分子检测结果；"
        "治疗相关信息须由临床医师结合病史、病理分型、分期及现行规范综合判断。",
    )
    note.paragraph_format.left_indent = Cm(5.2)
    note.paragraph_format.right_indent = Cm(1.0)
    note.paragraph_format.space_before = Pt(24)
    note.paragraph_format.line_spacing = 1.5
    note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in note.runs:
        run.font.size = Pt(10.5)


def _collapse_between(
    document: DocumentObject,
    start_text: str,
    end_text: str,
    replacement_lines: Sequence[str],
    *,
    replacement_heading: str | None = None,
) -> None:
    """Keep the two boundary headings, replacing all body blocks between them."""
    start = _body_paragraph(document, start_text)
    end = _body_paragraph(document, end_text)
    body_children = list(document.element.body.iterchildren())
    start_index = body_children.index(start._p)
    end_index = body_children.index(end._p)
    if end_index <= start_index:
        raise ValueError(f"invalid block order: {start_text!r} -> {end_text!r}")
    for child in body_children[start_index + 1 : end_index]:
        document.element.body.remove(child)
    if replacement_heading is not None:
        _replace_paragraph_text(start, replacement_heading)
    anchor = start
    for line in replacement_lines:
        anchor = _insert_paragraph_after(document, anchor, line)


def _collapse_after_anchor(
    document: DocumentObject,
    anchor_text: str,
    start_text: str,
    end_text: str,
    replacement_lines: Sequence[str],
    *,
    replacement_heading: str | None = None,
) -> None:
    """Collapse a repeated heading range, selecting the occurrence after anchor."""
    anchor = _body_paragraph(document, anchor_text)
    end = _body_paragraph(document, end_text)
    body_children = list(document.element.body.iterchildren())
    anchor_index = body_children.index(anchor._p)
    end_index = body_children.index(end._p)
    candidates = [
        paragraph
        for paragraph in document.paragraphs
        if _compact(paragraph.text) == _compact(start_text)
        and anchor_index < body_children.index(paragraph._p) < end_index
    ]
    if len(candidates) != 1:
        raise ValueError(
            f"expected one {start_text!r} after {anchor_text!r}, found {len(candidates)}"
        )
    start = candidates[0]
    start_index = body_children.index(start._p)
    for child in body_children[start_index + 1 : end_index]:
        document.element.body.remove(child)
    if replacement_heading is not None:
        _replace_paragraph_text(start, replacement_heading)
    insertion_anchor = start
    for line in replacement_lines:
        insertion_anchor = _insert_paragraph_after(document, insertion_anchor, line)


def _table_text(table, *, rows: int = 3) -> str:
    return _compact(
        " ".join(cell.text for row in table.rows[:rows] for cell in row.cells)
    )


def _find_table(document: DocumentObject, required: Iterable[str]):
    tokens = tuple(_compact(token) for token in required)
    matches = [
        table
        for table in document.tables
        if all(token in _table_text(table) for token in tokens)
    ]
    if len(matches) != 1:
        raise ValueError(f"expected one table containing {required!r}, found {len(matches)}")
    return matches[0]


def _remove_table(table) -> None:
    parent = table._tbl.getparent()
    if parent is not None:
        parent.remove(table._tbl)


def _remove_static_pdl1_image(document: DocumentObject) -> None:
    """Remove the scaffold's patient-specific PD-L1 image and orphaned media part.

    The source scaffold contains a fixed microscopy image between the PD-L1 and
    MSI headings. Reusing it while only replacing TPS/CPS would create a
    cross-case image leak. The builder therefore removes both the visible
    drawing and its unique package relationship, and leaves an explicit
    patient-facing no-image notice.
    """
    start = _body_paragraph(document, "3.2 PD-L1表达检测结果")
    end = _body_paragraph(document, "3.3微卫星不稳定性（MSI）检测结果")
    children = list(document.element.body.iterchildren())
    start_index = children.index(start._p)
    end_index = children.index(end._p)
    if end_index <= start_index:
        raise ValueError("invalid PD-L1 section order")

    image_relationship_ids: set[str] = set()
    image_paragraph_count = 0
    notice_count = 0
    caption_count = 0
    for child in list(children[start_index + 1 : end_index]):
        text = _element_text(child)
        compact_text = _compact(text)
        drawings = list(child.iter(qn("w:drawing")))
        picts = list(child.iter(qn("w:pict")))

        if compact_text == _compact("附图："):
            paragraph = next(
                candidate for candidate in document.paragraphs if candidate._p is child
            )
            _replace_paragraph_text(paragraph, PDL1_IMAGE_NOTICE)
            notice_count += 1
            continue

        if compact_text.startswith(_compact("图1")) and "PD-L1" in text:
            document.element.body.remove(child)
            caption_count += 1
            continue

        if not compact_text and drawings:
            if len(drawings) != 1 or picts:
                raise ValueError("unexpected static PD-L1 image representation")
            relationship_ids = {
                blip.get(qn("r:embed"))
                for blip in child.iter(qn("a:blip"))
                if blip.get(qn("r:embed"))
            }
            if len(relationship_ids) != 1:
                raise ValueError(
                    "static PD-L1 image must use one embedded relationship"
                )
            image_relationship_ids.update(relationship_ids)
            document.element.body.remove(child)
            image_paragraph_count += 1

    if (image_paragraph_count, notice_count, caption_count) != (1, 1, 1):
        raise ValueError(
            "unexpected PD-L1 scaffold assets: "
            f"image={image_paragraph_count}, notice={notice_count}, "
            f"caption={caption_count}"
        )

    remaining_relationship_ids = {
        blip.get(qn("r:embed"))
        for blip in document.element.iter(qn("a:blip"))
        if blip.get(qn("r:embed"))
    }
    for relationship_id in image_relationship_ids:
        if relationship_id in remaining_relationship_ids:
            raise ValueError("static PD-L1 image relationship is still referenced")
        document.part.drop_rel(relationship_id)


def _neutralize_table(table, values: Sequence[str]) -> None:
    if len(table.rows) < 2:
        raise ValueError("table has no body row to preserve")
    template_row = copy.deepcopy(table.rows[1]._tr)
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)
    table._tbl.append(template_row)
    row = table.rows[-1]
    if len(values) != len(row.cells):
        raise ValueError(
            f"neutral row has {len(row.cells)} cells, received {len(values)} values"
        )
    for cell, value in zip(row.cells, values):
        _replace_cell_text(cell, value)


def _set_table_loop(table, collection: str, columns: Sequence[str]) -> None:
    if len(table.rows) < 2:
        raise ValueError("table loop requires a body template row")
    body = copy.deepcopy(table.rows[1]._tr)
    start = copy.deepcopy(table.rows[1]._tr)
    end = copy.deepcopy(table.rows[1]._tr)
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)
    table._tbl.append(start)
    table._tbl.append(body)
    table._tbl.append(end)
    for row in table.rows[1:]:
        for cell in row.cells:
            _replace_cell_text(cell, "")
    _replace_cell_text(table.rows[1].cells[0], f"{{%tr for row in {collection} %}}")
    for cell, field in zip(table.rows[2].cells, columns):
        _replace_cell_text(cell, "{{ row." + field + " }}")
    _replace_cell_text(table.rows[3].cells[0], "{%tr endfor %}")


def _replace_variant_table(document: DocumentObject) -> None:
    stale = _find_table(document, ("突变起始位置", "突变终止位置", "转录本号"))
    _remove_table(stale)

    abbreviated = _find_table(
        document,
        ("基因突变位点信息", "潜在获益靶向/免疫药物"),
    )
    reference = Document(REFERENCE_VARIANT_TEMPLATE)
    full = _find_table(
        reference,
        ("基因突变信息", "转录本号", "潜在获益靶向药物"),
    )
    abbreviated._tbl.addnext(copy.deepcopy(full._tbl))
    _remove_table(abbreviated)


def _load_lung588_gene_list() -> list[str]:
    raw = yaml.safe_load(GENE_LIST_RULE.read_text(encoding="utf-8")) or {}
    genes = [
        str(value).strip().upper()
        for value in raw.get("reportable_genes") or []
        if str(value).strip()
    ]
    if len(genes) != 588 or len(set(genes)) != 588:
        raise ValueError(
            f"lung588 gene contract must contain 588 unique genes, got {len(genes)}"
        )
    digest = hashlib.sha256("\n".join(genes).encode("utf-8")).hexdigest()
    if digest != EXPECTED_GENE_LIST_SHA256:
        raise ValueError(
            "lung588 ordered gene contract changed: "
            f"expected={EXPECTED_GENE_LIST_SHA256}, actual={digest}"
        )
    return genes


def _replace_gene_list(document: DocumentObject) -> None:
    table = _find_table(document, ("Gene List for MLseq (n=329)",))
    if len(table.columns) != 7 or len(table.rows) < 2:
        raise ValueError("unexpected lung329 gene-list table shape")
    body_template = copy.deepcopy(table.rows[1]._tr)
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)

    genes = _load_lung588_gene_list()
    for offset in range(0, len(genes), 7):
        table._tbl.append(copy.deepcopy(body_template))
        row = table.rows[-1]
        # The 329 scaffold uses 0.88 cm rows. With 588 genes that leaves only
        # two rows on a fourth physical page in LibreOffice. Keep seven
        # readable columns but compact each deterministic body row so the tail
        # remains a meaningful page instead of an orphan.
        row.height = Cm(0.72)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        values = genes[offset : offset + 7]
        for index, cell in enumerate(row.cells):
            _replace_cell_text(cell, values[index] if index < len(values) else "")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    _replace_cell_text(table.rows[0].cells[0], "Gene List for MLseq (n=588)")


def _replace_biomarker_cells(document: DocumentObject) -> None:
    table = _find_table(document, ("TMB/MSI/其它生物标志物检测结果", "用药提示"))
    replacements = {
        "肿瘤突变负荷（TMB）": "{{ immuno_tips }}",
        "微卫星不稳定性（MSI）": "{{ msi_tips }}",
        "PD-L1表达": "{{ pdl1_table_interpretation }}",
        "免疫正相关基因": "研究性相关标志物，不能单独用于治疗决策。",
        "免疫负相关基因": "研究性相关标志物，不能单独用于治疗决策。",
        "免疫超进展相关基因": "研究性相关标志物，不能单独用于治疗决策。",
    }
    for row in table.rows[1:]:
        label = _compact(row.cells[0].text)
        for expected, value in replacements.items():
            if _compact(expected) == label:
                _replace_cell_text(row.cells[2], value)
                break


def _insert_summary_count(document: DocumentObject) -> None:
    table = _find_table(document, ("基因", "突变位点", "潜在获益靶向药物"))
    paragraph = document.add_paragraph(
        "*本次共检出体细胞变异：{{ total_variants_count }}个，其中与靶向药物"
        "用药相关的变异有：{{ drug_related_count }}个。"
    )
    table._tbl.addnext(paragraph._p)


def _harden_immune_marker_section(document: DocumentObject) -> None:
    """Retain dynamic result loops while removing unreviewed cross-cancer prose."""
    start = _body_paragraph(
        document,
        "3.4 免疫疗效正相关/负相关/超进展基因检测结果",
    )
    end = _body_paragraph(document, "4.化疗药物相关检测结果")
    children = list(document.element.body.iterchildren())
    start_index = children.index(start._p)
    end_index = children.index(end._p)
    if end_index <= start_index:
        raise ValueError("invalid immune marker section order")

    keep_headings = {
        _compact("免疫治疗正相关基因检测结果"),
        _compact("免疫治疗负相关基因检测结果"),
        _compact("免疫治疗超进展基因检测结果"),
    }
    intro = None
    tables = []
    for child in children[start_index + 1 : end_index]:
        if child.tag == qn("w:tbl"):
            table = next(
                (candidate for candidate in document.tables if candidate._tbl is child),
                None,
            )
            if table is None:
                raise ValueError("unable to resolve immune marker table")
            tables.append(table)
            continue
        if child.tag != qn("w:p"):
            document.element.body.remove(child)
            continue
        text = _compact(_element_text(child))
        if text in keep_headings:
            continue
        if intro is None:
            intro = next(
                paragraph for paragraph in document.paragraphs if paragraph._p is child
            )
            _replace_paragraph_text(intro, IMMUNE_RESEARCH_NOTICE)
            continue
        document.element.body.remove(child)

    if intro is None or len(tables) != 3:
        raise ValueError(
            f"expected immune intro and three dynamic tables, found intro={intro is not None}, "
            f"tables={len(tables)}"
        )
    for table in tables:
        if len(table.rows) != 4 or len(table.columns) != 3:
            raise ValueError("unexpected immune marker table shape")
        _replace_cell_text(table.rows[2].cells[2], IMMUNE_RESEARCH_RESULT)


def _prepare_toc_seed(document: DocumentObject) -> None:
    """Keep the reviewed TOC control but make its first PDF probe complete."""
    sdts = document.element.xpath(".//w:sdt")
    toc = next(
        (
            node
            for node in sdts
            if "TOC"
            in "".join(item.text or "" for item in node.iter(qn("w:instrText")))
        ),
        None,
    )
    if toc is None:
        raise ValueError("Word TOC content control not found")
    content = toc.find(qn("w:sdtContent"))
    if content is None:
        raise ValueError("Word TOC content control has no sdtContent")
    section_break = next(
        (
            copy.deepcopy(child)
            for child in content
            if any(True for _ in child.iter(qn("w:sectPr")))
        ),
        None,
    )
    for child in list(content):
        content.remove(child)
    for text in (
        "第一部分：基本信息",
        "第二部分：检测结果",
        "第三部分：基因变异及相应靶向/免疫药物解析",
        "第四部分：附录",
        "5. 参考文献",
    ):
        paragraph = OxmlElement("w:p")
        run = OxmlElement("w:r")
        node = OxmlElement("w:t")
        node.text = text
        run.append(node)
        paragraph.append(run)
        content.append(paragraph)
    if section_break is not None:
        content.append(section_break)


def _replace_last_reference_heading(document: DocumentObject) -> None:
    candidates = [
        paragraph
        for paragraph in document.paragraphs
        if _compact(paragraph.text) == _compact("参考文献")
    ]
    if not candidates:
        raise ValueError("final reference heading not found")
    _replace_paragraph_text(candidates[-1], "5. 参考文献")


def _visible_text(document: DocumentObject) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join(parts)


def _validate_hardened_document(document: DocumentObject) -> None:
    visible = _visible_text(document)
    found = [token for token in FORBIDDEN_OUTPUT_TOKENS if token in visible]
    if found:
        raise ValueError(f"hardened template retains forbidden static text: {found}")
    required = (
        "__PART3_MARKER__",
        "{{ pdl1_tps }}",
        "{{ pdl1_cps }}",
        "{{ pdl1_result }}",
        PDL1_ASSAY_PROVENANCE_MARKER,
        PDL1_SOURCE_PROVENANCE_MARKER,
        PDL1_IMAGE_NOTICE,
        "肺癌588基因+PD-L1检测项目",
        "Gene List for MLseq (n=588)",
        "{%tr for row in variants_2_1 %}",
        "{%tr for row in nccn_results %}",
        REVIEW_NOTICE,
        CHEMOTHERAPY_NOTICE,
    )
    missing = [token for token in required if token not in visible]
    if missing:
        raise ValueError(f"hardened template is missing required markers: {missing}")
    full_variant_table = [
        table
        for table in document.tables
        if all(
            token in _table_text(table)
            for token in (_compact("转录本号"), _compact("染色体"), _compact("频率"))
        )
    ]
    if len(full_variant_table) != 1:
        raise ValueError(
            f"expected one full dynamic variant table, found {len(full_variant_table)}"
        )
    gene_table = _find_table(document, ("Gene List for MLseq (n=588)",))
    rendered_genes = [
        _compact(cell.text)
        for row in gene_table.rows[1:]
        for cell in row.cells
        if _compact(cell.text)
    ]
    if rendered_genes != _load_lung588_gene_list():
        raise ValueError("rendered lung588 gene list does not match the frozen contract")

    start = _body_paragraph(document, "3.2 PD-L1表达检测结果")
    end = _body_paragraph(document, "3.3微卫星不稳定性（MSI）检测结果")
    children = list(document.element.body.iterchildren())
    pdl1_children = children[
        children.index(start._p) + 1 : children.index(end._p)
    ]
    blank_image_paragraphs = [
        child
        for child in pdl1_children
        if not _compact(_element_text(child))
        and any(True for _ in child.iter(qn("w:drawing")))
    ]
    if blank_image_paragraphs:
        raise ValueError("hardened template retains a static PD-L1 image")


def _normalize_zip_metadata(path: Path) -> None:
    """Make repeated builds byte-stable so the DOCX does not churn in Git."""
    fd, temp_name = tempfile.mkstemp(prefix=f"{path.stem}_", suffix=".docx", dir=path.parent)
    os.close(fd)
    temp_path = Path(temp_name)
    try:
        with ZipFile(path, "r") as source, ZipFile(temp_path, "w", ZIP_DEFLATED) as target:
            for original in sorted(source.infolist(), key=lambda item: item.filename):
                info = copy.copy(original)
                info.date_time = (1980, 1, 1, 0, 0, 0)
                info.extra = b""
                info.comment = b""
                target.writestr(info, source.read(original.filename))
        os.replace(temp_path, path)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def build_template(source: Path, output: Path, *, allow_source_drift: bool = False) -> dict:
    actual_source_hash = _sha256(source)
    if not allow_source_drift and actual_source_hash != EXPECTED_SOURCE_SHA256:
        raise ValueError(
            "lung329 scaffold changed; review it before rebuilding lung588: "
            f"expected={EXPECTED_SOURCE_SHA256}, actual={actual_source_hash}"
        )
    if not REFERENCE_VARIANT_TEMPLATE.is_file():
        raise FileNotFoundError(REFERENCE_VARIANT_TEMPLATE)

    document = Document(source)

    _replace_paragraph_text(
        _body_paragraph(document, "肺癌329基因检测"),
        "肺癌588基因+PD-L1检测项目",
    )
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "使用免疫组化（IHC）方法，对委托人 组织 的PD-L1蛋白表达情况进行检测，给予用药提示。",
        ),
        "使用免疫组化（IHC）方法检测PD-L1蛋白表达，报告TPS、CPS及结果判定。",
    )
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "分析样本的基因变异，寻找与靶向/免疫药物相关的变异，给予用药提示与理论支持。",
        ),
        "分析样本基因变异并报告分子事件；患者级用药提示仅在肺癌专属规则完成审核后输出。",
    )
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "分析化疗药物相关的基因变异，评估化疗的敏感性或毒副作用，为化疗方案的制订提供参考。",
        ),
        "化疗药物多态性模块尚未完成接入和审核，当前版本不输出相关结论。",
    )
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "分析样本微卫星不稳定性，给予用药提示。",
        ),
        "分析样本微卫星不稳定性并报告检测状态；治疗意义须结合临床信息综合判断。",
    )
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "报告第二部分：检测结果（提供本次检测涉及的靶向治疗、免疫治疗、化疗等综合检测结果），是本报告的关键信息。",
        ),
        "报告第二部分：检测结果。提供本次分子检测结果及当前已启用的审核结论，是本报告的关键信息。",
    )
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "报告第三部分：基因变异及相应靶向/免疫药物解析，该部分对第二部分中的基因变异和靶向/免疫药物提示进行详细解析，并且包含本报告的阅读说明。",
        ),
        "报告第三部分：基因变异及相应靶向/免疫药物解析。相关知识未启用时，本节仅显示解释边界和阅读说明。",
    )
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "报告第四部分：附录。提供所检测癌症的诊疗知识、癌症相关重要信号通路、所检测的基因列表信息、以及本报告的参考文献。",
        ),
        "报告第四部分：附录。提供检测基因列表、参考文献、检测方法、质量控制及报告说明。",
    )

    _replace_variant_table(document)
    _replace_gene_list(document)
    _replace_biomarker_cells(document)
    _remove_static_pdl1_image(document)
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "5、因肿瘤存在较大的异质性，送检样本不一定代表肿瘤的全貌。",
        ),
        "4、因肿瘤存在较大的异质性，送检样本不一定代表肿瘤的全貌。",
    )
    _insert_summary_count(document)

    _neutralize_table(
        _find_table(document, ("检测基因", "本癌种相关治疗药物", "临床提示")),
        ("-", "待审核，当前关闭", REVIEW_NOTICE, "-"),
    )
    _neutralize_table(
        _find_table(document, ("化疗药物检测结果", "化疗药物/方案用药提示")),
        ("化疗药物规则待审核", CHEMOTHERAPY_NOTICE),
    )
    _neutralize_table(
        _find_table(document, ("药物名称", "相关基因", "药物介绍", "尼达尼布")),
        ("待审核", "-", REVIEW_NOTICE),
    )
    _neutralize_table(
        _find_table(document, ("药物名称", "相关基因", "药物介绍", "帕博利珠单抗")),
        ("待审核", "PD-L1", REVIEW_NOTICE),
    )

    guideline_table = _find_table(document, ("检测基因", "检测内容", "检测结果", "外显子18"))
    _set_table_loop(guideline_table, "nccn_results", ("gene", "content", "result"))
    _replace_paragraph_text(
        _body_paragraph(document, "2.3 NCCN推荐临床常规靶向药物相关基因检测结果（不限于本癌种）"),
        "2.3 肺癌重点分子事件检测结果",
    )
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "注：1. 以上只列出 NCCN 指南推荐的检测基因，基因检测范围涵盖但不限于上表中列出的检测内容。",
        ),
        "注：本表仅在肺癌重点分子事件及相应指南口径完成审核后展示；当前版本未启用。",
    )
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "2. 检出详情，可查看基因变异检测结果。",
        ),
        "检出详情请查看本报告的基因变异检测结果。",
    )
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "（*以下为本癌种FDA/NMPA批准的抗血管生成类的靶向药物，或国内临床实践中疗效较好的靶向药物，可以作为多线治疗的备选方案。）",
        ),
        REVIEW_NOTICE,
    )

    _collapse_between(
        document,
        "非小细胞肺癌NCCN指南（2022 V3）",
        "3.免疫治疗疗效评估",
        (REVIEW_NOTICE,),
        replacement_heading="肺癌临床指南说明（当前版本未启用）",
    )
    _collapse_between(
        document,
        "备注：1. FoundationOne CDx (324 个基因) TMB 研究表明，对于组织样本，TMB≥10 mut/Mb 为高突变负荷 (仅供参考) 。而对于血液样本，POPLAR 和 OAK 研究表明，bTMB（血液肿瘤突变负荷）≥16 或可成为阿替利珠单抗治疗晚期非小细胞肺癌的生物标记物（仅供参考）。",
        "化疗药物用药提示",
        (REVIEW_NOTICE,),
        replacement_heading="生物标志物结果解释边界",
    )
    _collapse_between(
        document,
        "3.1 肿瘤突变负荷（TMB）水平提示",
        "3.2 PD-L1表达检测结果",
        (
            "{{ tmb_detail_sentence }}",
            "{{ tmb_detail_interpretation }}",
            "{{ tmb_drug_note }}",
        ),
    )
    _collapse_after_anchor(
        document,
        "3.2 PD-L1表达检测结果",
        "用药提示",
        "3.3微卫星不稳定性（MSI）检测结果",
        (
            "{{ pdl1_table_interpretation }}",
            PDL1_ASSAY_PROVENANCE_MARKER,
            PDL1_SOURCE_PROVENANCE_MARKER,
        ),
        replacement_heading="结果解释边界",
    )
    _collapse_between(
        document,
        "3.3微卫星不稳定性（MSI）检测结果",
        "3.4 免疫疗效正相关/负相关/超进展基因检测结果",
        (
            "{{ msi_detail_sentence }}",
            "{{ msi_detail_interpretation }}",
            "{{ msi_tips }}",
        ),
    )
    _harden_immune_marker_section(document)
    _collapse_between(
        document,
        "4.化疗药物相关检测结果",
        "5. 检测结果说明",
        (CHEMOTHERAPY_NOTICE,),
    )
    _collapse_between(
        document,
        "常见问题解答",
        "2. 肺癌诊疗知识",
        (APPENDIX_NOTICE,),
        replacement_heading="1. 常见问题解答",
    )
    _collapse_between(
        document,
        "2. 肺癌诊疗知识",
        "3. 癌症相关信号通路",
        (APPENDIX_NOTICE,),
    )
    _collapse_between(
        document,
        "3.1 文中参考文献及临床试验编号说明",
        "3.2 文中医学及生物学常见名词说明",
        (
            "正文如出现PMID或临床试验登记号，仅用于文献追溯；具体研究适用范围、"
            "入组条件及证据等级须结合原始文献核对，不得把登记号本身作为用药依据。",
        ),
    )
    _replace_last_reference_heading(document)
    _collapse_between(
        document,
        "5. 参考文献",
        "本次检测质控结果",
        REFERENCE_LINES,
    )
    _collapse_between(
        document,
        "3. 癌症相关信号通路",
        "4. 基因检测列表",
        (APPENDIX_NOTICE,),
    )
    _remove_explicit_page_break_before(document, "4.化疗药物相关检测结果")
    _add_letter_divider_note(document)
    _prepare_toc_seed(document)
    _validate_hardened_document(document)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    _normalize_zip_metadata(output)
    return {
        "source": str(source),
        "source_sha256": actual_source_hash,
        "output": str(output),
        "output_sha256": _sha256(output),
        "table_count": len(document.tables),
        "paragraph_count": len(document.paragraphs),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--allow-source-drift", action="store_true")
    args = parser.parse_args()
    result = build_template(
        args.source.resolve(),
        args.output.resolve(),
        allow_source_drift=args.allow_source_drift,
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
