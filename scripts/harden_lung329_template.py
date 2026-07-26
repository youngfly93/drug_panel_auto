#!/usr/bin/env python3
# 步骤: 70 肺癌329受控试运行模板临床内容硬化
# 上游: 受控外部迁移源（不进入Git或生产发布包）
# 输出: panels/lung_329_pdl1/templates/lung_329_pdl1_golden_template_v2.docx
# 种子: 无（确定性文档变换）
"""Build the lung329 v2 controlled-pilot template without static clinical claims.

The migration source is supplied explicitly from controlled external storage
and is never tracked or shipped. Its dynamic variant/PD-L1 scaffolding is
useful, but several report-visible tables contain fixed historical treatment
recommendations or patient genotypes. This script produces a v2 template whose
active report data comes from ReportGen context only. Unreviewed clinical
sections fail closed with an explicit review notice.
"""

from __future__ import annotations

import argparse
import copy
import hashlib
import json
import os
import re
import tempfile
from pathlib import Path
from typing import Iterable, Sequence
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = (
    ROOT
    / "panels"
    / "lung_329_pdl1"
    / "templates"
    / "lung_329_pdl1_golden_template_v2.docx"
)
REFERENCE_VARIANT_TEMPLATE = (
    ROOT
    / "panels"
    / "crc_301_msi"
    / "templates"
    / "crc_301_msi_golden_template_v1.docx"
)
EXPECTED_SOURCE_SHA256 = (
    "bbecf470feedab8ba20c4d84f7e4f8f0e2f0714f8f50818dc0419d72d2767e66"
)

REVIEW_NOTICE = (
    "肺癌专属治疗知识和事件级药物规则尚未完成报告组二审及病例级验收，"
    "当前版本不输出患者级用药结论。"
)
CHEMOTHERAPY_NOTICE = (
    "化疗药物多态性与方案模块尚未完成数据接入、医学审核及脱敏UAT，"
    "当前版本不输出患者级化疗敏感性、毒性或剂量结论。"
)
APPENDIX_NOTICE = (
    "本节历史静态内容已停用；经医学审核并建立可追溯来源后再恢复展示。"
)
IMMUNE_RESEARCH_NOTICE = (
    "免疫正相关、负相关及超进展基因分类尚未完成肺癌专属医学审核，"
    "当前版本不展示患者级基因结果，也不据此预测疗效、耐药或超进展。"
)
PDL1_IMAGE_NOTICE = (
    "附图：本病例未提供可追溯的PD-L1免疫组化图像，故不展示；"
    "TPS、CPS及结果判定须与原始检测记录核对。"
)
PDL1_ASSAY_PROVENANCE_MARKER = "{{ pdl1_assay_provenance }}"
PDL1_SOURCE_PROVENANCE_MARKER = "{{ pdl1_source_provenance }}"
REFERENCE_LINES = (
    "1. PMID:39375078. Chinese guidelines for molecular testing of non-small cell lung cancer "
    "(2024 edition). https://pubmed.ncbi.nlm.nih.gov/39375078/",
    "2. PMID:29355391. Updated Molecular Testing Guideline for the Selection of Lung Cancer "
    "Patients for Treatment With Targeted Tyrosine Kinase Inhibitors. "
    "https://pubmed.ncbi.nlm.nih.gov/29355391/",
    "3. National Cancer Institute. Non-Small Cell Lung Cancer Treatment (PDQ®)—Health "
    "Professional Version. https://www.cancer.gov/types/lung/hp/non-small-cell-lung-treatment-pdq",
)

FORBIDDEN_OUTPUT_TOKENS = (
    "非小细胞肺癌NCCN指南（2022 V3）",
    "非小细胞肺癌CSCO指南（2022）",
    "多西他赛单药方案",
    "呋喹替尼 Fruquintinib",
    "6TA/6TA",
    "2020年新版NSCLC指南",
    "结直肠癌",
    "5-Fu",
    "帕博利珠单抗",
    "纳武利尤单抗",
    "奥希替尼",
    "PBRM1基因突变的肾透明细胞癌患者免疫治疗获益",
    "TP53 mutations and hepatocellular carcinoma",
    "NCT02576444",
)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _compact(value: object) -> str:
    return re.sub(r"\s+", "", str(value or ""))


def _element_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def _replace_paragraph_text(paragraph, text: str) -> None:
    runs = list(paragraph.runs)
    if runs:
        target_index = next((i for i, run in enumerate(runs) if run.text), 0)
        for index, run in enumerate(runs):
            run.text = text if index == target_index else ""
    else:
        paragraph.add_run(text)


def _remove_paragraph(paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def _replace_cell_text(cell, text: str) -> None:
    if not cell.paragraphs:
        cell.text = text
        return
    _replace_paragraph_text(cell.paragraphs[0], text)
    for paragraph in list(cell.paragraphs[1:]):
        _remove_paragraph(paragraph)


def _body_paragraph(document: DocumentObject, text: str, *, prefer_last: bool = False):
    expected = _compact(text)
    matches = [p for p in document.paragraphs if _compact(p.text) == expected]
    if not matches:
        raise ValueError(f"paragraph not found: {text!r}")
    return matches[-1] if prefer_last else matches[0]


def _insert_paragraph_after(document: DocumentObject, anchor, text: str):
    paragraph = document.add_paragraph(text)
    anchor._p.addnext(paragraph._p)
    return paragraph


def _remove_explicit_page_break_before(
    document: DocumentObject,
    heading_text: str,
) -> None:
    """Remove a redundant page-break paragraph before a page-break heading."""
    heading = _body_paragraph(document, heading_text)
    previous = heading._p.getprevious()
    if previous is None:
        return
    page_breaks = [
        node
        for node in previous.iter(qn("w:br"))
        if node.get(qn("w:type")) == "page"
    ]
    if page_breaks:
        document.element.body.remove(previous)


def _add_letter_divider_note(document: DocumentObject) -> None:
    heading = _body_paragraph(document, "致您的一封信")
    note = _insert_paragraph_after(
        document,
        heading,
        "感谢您对本次检测的信任。本报告呈现经质量控制的分子检测结果；"
        "治疗相关信息须由临床医师结合病史、病理分型、分期及现行规范综合判断。",
    )
    note.paragraph_format.left_indent = Cm(5.2)
    note.paragraph_format.right_indent = Cm(1.0)
    note.paragraph_format.space_before = Pt(24)
    note.paragraph_format.line_spacing = 1.5
    note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in note.runs:
        run.font.size = Pt(10.5)


def _collapse_between(
    document: DocumentObject,
    start_text: str,
    end_text: str,
    replacement_lines: Sequence[str],
    *,
    replacement_heading: str | None = None,
) -> None:
    """Keep the two boundary headings, replacing all body blocks between them."""
    start = _body_paragraph(document, start_text)
    end = _body_paragraph(document, end_text)
    body_children = list(document.element.body.iterchildren())
    start_index = body_children.index(start._p)
    end_index = body_children.index(end._p)
    if end_index <= start_index:
        raise ValueError(f"invalid block order: {start_text!r} -> {end_text!r}")
    for child in body_children[start_index + 1 : end_index]:
        document.element.body.remove(child)
    if replacement_heading is not None:
        _replace_paragraph_text(start, replacement_heading)
    anchor = start
    for line in replacement_lines:
        anchor = _insert_paragraph_after(document, anchor, line)


def _collapse_after_anchor(
    document: DocumentObject,
    anchor_text: str,
    start_text: str,
    end_text: str,
    replacement_lines: Sequence[str],
    *,
    replacement_heading: str | None = None,
) -> None:
    """Collapse a repeated heading range, selecting the occurrence after anchor."""
    anchor = _body_paragraph(document, anchor_text)
    end = _body_paragraph(document, end_text)
    body_children = list(document.element.body.iterchildren())
    anchor_index = body_children.index(anchor._p)
    end_index = body_children.index(end._p)
    candidates = [
        paragraph
        for paragraph in document.paragraphs
        if _compact(paragraph.text) == _compact(start_text)
        and anchor_index < body_children.index(paragraph._p) < end_index
    ]
    if len(candidates) != 1:
        raise ValueError(
            f"expected one {start_text!r} after {anchor_text!r}, found {len(candidates)}"
        )
    start = candidates[0]
    start_index = body_children.index(start._p)
    for child in body_children[start_index + 1 : end_index]:
        document.element.body.remove(child)
    if replacement_heading is not None:
        _replace_paragraph_text(start, replacement_heading)
    insertion_anchor = start
    for line in replacement_lines:
        insertion_anchor = _insert_paragraph_after(document, insertion_anchor, line)


def _table_text(table, *, rows: int = 3) -> str:
    return _compact(
        " ".join(cell.text for row in table.rows[:rows] for cell in row.cells)
    )


def _find_table(document: DocumentObject, required: Iterable[str]):
    tokens = tuple(_compact(token) for token in required)
    matches = [
        table
        for table in document.tables
        if all(token in _table_text(table) for token in tokens)
    ]
    if len(matches) != 1:
        raise ValueError(f"expected one table containing {required!r}, found {len(matches)}")
    return matches[0]


def _remove_table(table) -> None:
    parent = table._tbl.getparent()
    if parent is not None:
        parent.remove(table._tbl)


def _remove_static_pdl1_image(document: DocumentObject) -> set[str]:
    """Remove the inherited case-specific PD-L1 image and its package relation."""
    start = _body_paragraph(document, "3.2 PD-L1表达检测结果")
    end = _body_paragraph(document, "3.3微卫星不稳定性（MSI）检测结果")
    children = list(document.element.body.iterchildren())
    start_index = children.index(start._p)
    end_index = children.index(end._p)
    if end_index <= start_index:
        raise ValueError("invalid PD-L1 section order")

    relationship_ids: set[str] = set()
    image_hashes: set[str] = set()
    image_paragraph_count = 0
    label_count = 0
    caption_count = 0
    for child in list(children[start_index + 1 : end_index]):
        text = _element_text(child)
        compact_text = _compact(text)
        drawings = list(child.iter(qn("w:drawing")))

        if compact_text == _compact("附图："):
            document.element.body.remove(child)
            label_count += 1
            continue

        if compact_text.startswith(_compact("图1")) and "PD-L1" in text:
            document.element.body.remove(child)
            caption_count += 1
            continue

        if not compact_text and drawings:
            embedded = {
                blip.get(qn("r:embed"))
                for blip in child.iter(qn("a:blip"))
                if blip.get(qn("r:embed"))
            }
            if len(drawings) != 1 or len(embedded) != 1:
                raise ValueError("unexpected static PD-L1 image representation")
            relationship_id = next(iter(embedded))
            relationship = document.part.rels.get(relationship_id)
            if relationship is None or not relationship.reltype.endswith("/image"):
                raise ValueError("static PD-L1 drawing does not resolve to an image")
            image_hashes.add(hashlib.sha256(relationship.target_part.blob).hexdigest())
            relationship_ids.add(relationship_id)
            document.element.body.remove(child)
            image_paragraph_count += 1

    if (image_paragraph_count, label_count, caption_count) != (1, 1, 1):
        raise ValueError(
            "unexpected PD-L1 scaffold assets: "
            f"image={image_paragraph_count}, label={label_count}, "
            f"caption={caption_count}"
        )

    remaining_relationship_ids = {
        blip.get(qn("r:embed"))
        for blip in document.element.iter(qn("a:blip"))
        if blip.get(qn("r:embed"))
    }
    for relationship_id in relationship_ids:
        if relationship_id in remaining_relationship_ids:
            raise ValueError("static PD-L1 image relationship is still referenced")
        document.part.drop_rel(relationship_id)
    return image_hashes


def _drop_unreferenced_document_media_relationships(
    document: DocumentObject,
) -> int:
    """Remove media relations left behind when static report sections are deleted."""
    used_ids = {
        value
        for element in document.element.iter()
        for attribute in (qn("r:embed"), qn("r:link"))
        if (value := element.get(attribute))
    }
    stale_ids = [
        relationship_id
        for relationship_id, relationship in document.part.rels.items()
        if relationship.reltype.rsplit("/", 1)[-1] in {"image", "hdphoto"}
        and relationship_id not in used_ids
    ]
    for relationship_id in stale_ids:
        document.part.drop_rel(relationship_id)
    return len(stale_ids)


def _validate_output_media(
    path: Path,
    *,
    forbidden_hashes: set[str],
) -> None:
    """Require every stored media part to be referenced and reject source-case images."""
    with ZipFile(path, "r") as archive:
        names = set(archive.namelist())
        media_parts = {name for name in names if name.startswith("word/media/")}
        relationship_targets: set[str] = set()
        for name in names:
            if not name.endswith(".rels"):
                continue
            for target in re.findall(
                rb'Target="(?:\.\./)?media/([^"]+)"',
                archive.read(name),
            ):
                relationship_targets.add(
                    "word/media/" + target.decode("utf-8", errors="strict")
                )
        if media_parts != relationship_targets:
            raise ValueError(
                "hardened template contains missing or orphaned media parts: "
                f"stored_only={sorted(media_parts - relationship_targets)}, "
                f"target_only={sorted(relationship_targets - media_parts)}"
            )
        leaked = [
            name
            for name in sorted(media_parts)
            if hashlib.sha256(archive.read(name)).hexdigest() in forbidden_hashes
        ]
        if leaked:
            raise ValueError(
                "hardened template retains the source-case PD-L1 image: "
                + ", ".join(leaked)
            )


def _neutralize_table(table, values: Sequence[str]) -> None:
    if len(table.rows) < 2:
        raise ValueError("table has no body row to preserve")
    template_row = copy.deepcopy(table.rows[1]._tr)
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)
    table._tbl.append(template_row)
    row = table.rows[-1]
    if len(values) != len(row.cells):
        raise ValueError(
            f"neutral row has {len(row.cells)} cells, received {len(values)} values"
        )
    for cell, value in zip(row.cells, values):
        _replace_cell_text(cell, value)


def _set_table_loop(table, collection: str, columns: Sequence[str]) -> None:
    if len(table.rows) < 2:
        raise ValueError("table loop requires a body template row")
    body = copy.deepcopy(table.rows[1]._tr)
    start = copy.deepcopy(table.rows[1]._tr)
    end = copy.deepcopy(table.rows[1]._tr)
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)
    table._tbl.append(start)
    table._tbl.append(body)
    table._tbl.append(end)
    for row in table.rows[1:]:
        for cell in row.cells:
            _replace_cell_text(cell, "")
    _replace_cell_text(table.rows[1].cells[0], f"{{%tr for row in {collection} %}}")
    for cell, field in zip(table.rows[2].cells, columns):
        _replace_cell_text(cell, "{{ row." + field + " }}")
    _replace_cell_text(table.rows[3].cells[0], "{%tr endfor %}")


def _replace_variant_table(document: DocumentObject) -> None:
    stale = _find_table(document, ("突变起始位置", "突变终止位置", "转录本号"))
    _remove_table(stale)

    abbreviated = _find_table(
        document,
        ("基因突变位点信息", "潜在获益靶向/免疫药物"),
    )
    reference = Document(REFERENCE_VARIANT_TEMPLATE)
    full = _find_table(
        reference,
        ("基因突变信息", "转录本号", "潜在获益靶向药物"),
    )
    abbreviated._tbl.addnext(copy.deepcopy(full._tbl))
    _remove_table(abbreviated)


def _replace_biomarker_cells(document: DocumentObject) -> None:
    table = _find_table(document, ("TMB/MSI/其它生物标志物检测结果", "用药提示"))
    replacements = {
        "肿瘤突变负荷（TMB）": "{{ immuno_tips }}",
        "微卫星不稳定性（MSI）": "{{ msi_tips }}",
        "PD-L1表达": "{{ pdl1_table_interpretation }}",
        "免疫正相关基因": "研究性相关标志物，不能单独用于治疗决策。",
        "免疫负相关基因": "研究性相关标志物，不能单独用于治疗决策。",
        "免疫超进展相关基因": "研究性相关标志物，不能单独用于治疗决策。",
    }
    for row in table.rows[1:]:
        label = _compact(row.cells[0].text)
        for expected, value in replacements.items():
            if _compact(expected) == label:
                _replace_cell_text(row.cells[2], value)
                break


def _insert_summary_count(document: DocumentObject) -> None:
    table = _find_table(document, ("基因", "突变位点", "潜在获益靶向药物"))
    paragraph = document.add_paragraph(
        "*本次共检出体细胞变异：{{ total_variants_count }}个，其中与靶向药物"
        "用药相关的变异有：{{ drug_related_count }}个。"
    )
    table._tbl.addnext(paragraph._p)


def _harden_immune_marker_section(document: DocumentObject) -> None:
    """Replace inherited immune-gene tables with one explicit disabled notice."""
    _collapse_between(
        document,
        "3.4 免疫疗效正相关/负相关/超进展基因检测结果",
        "4.化疗药物相关检测结果",
        (IMMUNE_RESEARCH_NOTICE,),
        replacement_heading="3.4 免疫相关基因模块（未启用）",
    )


def _prepare_toc_seed(document: DocumentObject) -> None:
    """Keep the reviewed TOC control but make its first PDF probe complete."""
    sdts = document.element.xpath(".//w:sdt")
    toc = next(
        (
            node
            for node in sdts
            if "TOC"
            in "".join(item.text or "" for item in node.iter(qn("w:instrText")))
        ),
        None,
    )
    if toc is None:
        raise ValueError("Word TOC content control not found")
    content = toc.find(qn("w:sdtContent"))
    if content is None:
        raise ValueError("Word TOC content control has no sdtContent")
    section_break = next(
        (
            copy.deepcopy(child)
            for child in content
            if any(True for _ in child.iter(qn("w:sectPr")))
        ),
        None,
    )
    for child in list(content):
        content.remove(child)
    for text in (
        "第一部分：基本信息",
        "第二部分：检测结果",
        "第三部分：基因变异及相应靶向/免疫药物解析",
        "第四部分：附录",
        "5. 参考文献",
    ):
        paragraph = OxmlElement("w:p")
        run = OxmlElement("w:r")
        node = OxmlElement("w:t")
        node.text = text
        run.append(node)
        paragraph.append(run)
        content.append(paragraph)
    if section_break is not None:
        content.append(section_break)


def _replace_last_reference_heading(document: DocumentObject) -> None:
    candidates = [
        paragraph
        for paragraph in document.paragraphs
        if _compact(paragraph.text) == _compact("参考文献")
    ]
    if not candidates:
        raise ValueError("final reference heading not found")
    _replace_paragraph_text(candidates[-1], "5. 参考文献")


def _visible_text(document: DocumentObject) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join(parts)


def _validate_hardened_document(document: DocumentObject) -> None:
    visible = _visible_text(document)
    found = [token for token in FORBIDDEN_OUTPUT_TOKENS if token in visible]
    if found:
        raise ValueError(f"hardened template retains forbidden static text: {found}")
    required = (
        "__PART3_MARKER__",
        "{{ pdl1_tps }}",
        "{{ pdl1_cps }}",
        "{{ pdl1_result }}",
        PDL1_ASSAY_PROVENANCE_MARKER,
        PDL1_SOURCE_PROVENANCE_MARKER,
        PDL1_IMAGE_NOTICE,
        "{%tr for row in variants_2_1 %}",
        "{%tr for row in nccn_results %}",
        REVIEW_NOTICE,
        CHEMOTHERAPY_NOTICE,
    )
    missing = [token for token in required if token not in visible]
    if missing:
        raise ValueError(f"hardened template is missing required markers: {missing}")
    full_variant_table = [
        table
        for table in document.tables
        if all(
            token in _table_text(table)
            for token in (_compact("转录本号"), _compact("染色体"), _compact("频率"))
        )
    ]
    if len(full_variant_table) != 1:
        raise ValueError(
            f"expected one full dynamic variant table, found {len(full_variant_table)}"
        )


def _normalize_zip_metadata(path: Path) -> None:
    """Make repeated builds byte-stable so the DOCX does not churn in Git."""
    fd, temp_name = tempfile.mkstemp(prefix=f"{path.stem}_", suffix=".docx", dir=path.parent)
    os.close(fd)
    temp_path = Path(temp_name)
    try:
        with ZipFile(path, "r") as source, ZipFile(temp_path, "w", ZIP_DEFLATED) as target:
            for original in sorted(source.infolist(), key=lambda item: item.filename):
                info = copy.copy(original)
                info.date_time = (1980, 1, 1, 0, 0, 0)
                info.extra = b""
                info.comment = b""
                target.writestr(info, source.read(original.filename))
        os.replace(temp_path, path)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def build_template(source: Path, output: Path, *, allow_source_drift: bool = False) -> dict:
    actual_source_hash = _sha256(source)
    if not allow_source_drift and actual_source_hash != EXPECTED_SOURCE_SHA256:
        raise ValueError(
            "lung329 v1 source changed; review the template before rebuilding v2: "
            f"expected={EXPECTED_SOURCE_SHA256}, actual={actual_source_hash}"
        )
    if not REFERENCE_VARIANT_TEMPLATE.is_file():
        raise FileNotFoundError(REFERENCE_VARIANT_TEMPLATE)

    document = Document(source)

    _replace_paragraph_text(
        _body_paragraph(document, "肺癌329基因检测"),
        "肺癌329基因+PD-L1检测项目",
    )
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "使用免疫组化（IHC）方法，对委托人 组织 的PD-L1蛋白表达情况进行检测，给予用药提示。",
        ),
        "使用免疫组化（IHC）方法检测PD-L1蛋白表达，报告TPS、CPS及结果判定。",
    )
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "分析样本的基因变异，寻找与靶向/免疫药物相关的变异，给予用药提示与理论支持。",
        ),
        "分析样本基因变异并报告分子事件；患者级用药提示仅在肺癌专属规则完成审核后输出。",
    )
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "分析化疗药物相关的基因变异，评估化疗的敏感性或毒副作用，为化疗方案的制订提供参考。",
        ),
        "化疗药物多态性模块尚未完成接入和审核，当前版本不输出相关结论。",
    )

    _replace_variant_table(document)
    _replace_biomarker_cells(document)
    forbidden_media_hashes = _remove_static_pdl1_image(document)
    _insert_summary_count(document)

    _neutralize_table(
        _find_table(document, ("检测基因", "本癌种相关治疗药物", "临床提示")),
        ("-", "待审核，当前关闭", REVIEW_NOTICE, "-"),
    )
    _neutralize_table(
        _find_table(document, ("化疗药物检测结果", "化疗药物/方案用药提示")),
        ("化疗药物规则待审核", CHEMOTHERAPY_NOTICE),
    )
    _neutralize_table(
        _find_table(document, ("药物名称", "相关基因", "药物介绍", "尼达尼布")),
        ("待审核", "-", REVIEW_NOTICE),
    )
    _neutralize_table(
        _find_table(document, ("药物名称", "相关基因", "药物介绍", "帕博利珠单抗")),
        ("待审核", "PD-L1", REVIEW_NOTICE),
    )

    guideline_table = _find_table(document, ("检测基因", "检测内容", "检测结果", "外显子18"))
    _set_table_loop(guideline_table, "nccn_results", ("gene", "content", "result"))
    _replace_paragraph_text(
        _body_paragraph(document, "2.3 NCCN推荐临床常规靶向药物相关基因检测结果（不限于本癌种）"),
        "2.3 肺癌重点分子事件检测结果",
    )
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "（*以下为本癌种FDA/NMPA批准的抗血管生成类的靶向药物，或国内临床实践中疗效较好的靶向药物，可以作为多线治疗的备选方案。）",
        ),
        REVIEW_NOTICE,
    )

    _collapse_between(
        document,
        "非小细胞肺癌NCCN指南（2022 V3）",
        "3.免疫治疗疗效评估",
        (REVIEW_NOTICE,),
        replacement_heading="肺癌临床指南说明（待报告组审核）",
    )
    _collapse_between(
        document,
        "备注：1. FoundationOne CDx (324 个基因) TMB 研究表明，对于组织样本，TMB≥10 mut/Mb 为高突变负荷 (仅供参考) 。而对于血液样本，POPLAR 和 OAK 研究表明，bTMB（血液肿瘤突变负荷）≥16 或可成为阿替利珠单抗治疗晚期非小细胞肺癌的生物标记物（仅供参考）。",
        "化疗药物用药提示",
        (REVIEW_NOTICE,),
        replacement_heading="生物标志物结果解释边界",
    )
    _collapse_between(
        document,
        "3.1 肿瘤突变负荷（TMB）水平提示",
        "3.2 PD-L1表达检测结果",
        (
            "{{ tmb_detail_sentence }}",
            "{{ tmb_detail_interpretation }}",
            "{{ tmb_drug_note }}",
        ),
    )
    _collapse_after_anchor(
        document,
        "3.2 PD-L1表达检测结果",
        "用药提示",
        "3.3微卫星不稳定性（MSI）检测结果",
        (
            "{{ pdl1_table_interpretation }}",
            PDL1_ASSAY_PROVENANCE_MARKER,
            PDL1_SOURCE_PROVENANCE_MARKER,
            PDL1_IMAGE_NOTICE,
        ),
        replacement_heading="结果解释边界",
    )
    _collapse_between(
        document,
        "3.3微卫星不稳定性（MSI）检测结果",
        "3.4 免疫疗效正相关/负相关/超进展基因检测结果",
        (
            "{{ msi_detail_sentence }}",
            "{{ msi_detail_interpretation }}",
            "{{ msi_tips }}",
        ),
    )
    _harden_immune_marker_section(document)
    _collapse_between(
        document,
        "4.化疗药物相关检测结果",
        "5. 检测结果说明",
        (CHEMOTHERAPY_NOTICE,),
    )
    _collapse_between(
        document,
        "常见问题解答",
        "2. 肺癌诊疗知识",
        (APPENDIX_NOTICE,),
        replacement_heading="1. 常见问题解答",
    )
    _collapse_between(
        document,
        "2. 肺癌诊疗知识",
        "3. 癌症相关信号通路",
        (APPENDIX_NOTICE,),
    )
    _collapse_between(
        document,
        "3.1 文中参考文献及临床试验编号说明",
        "3.2 文中医学及生物学常见名词说明",
        (
            "正文如出现PMID或临床试验登记号，仅用于文献追溯；具体研究适用范围、"
            "入组条件及证据等级须由报告组复核，不得把登记号本身作为用药依据。",
        ),
    )
    _replace_last_reference_heading(document)
    _collapse_between(
        document,
        "5. 参考文献",
        "本次检测质控结果",
        REFERENCE_LINES,
    )
    _collapse_between(
        document,
        "3. 癌症相关信号通路",
        "4. 基因检测列表",
        (APPENDIX_NOTICE,),
    )
    _remove_explicit_page_break_before(document, "4.化疗药物相关检测结果")
    _add_letter_divider_note(document)
    _prepare_toc_seed(document)
    _drop_unreferenced_document_media_relationships(document)
    _validate_hardened_document(document)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    _normalize_zip_metadata(output)
    _validate_output_media(output, forbidden_hashes=forbidden_media_hashes)
    return {
        "source": str(source),
        "source_sha256": actual_source_hash,
        "output": str(output),
        "output_sha256": _sha256(output),
        "table_count": len(document.tables),
        "paragraph_count": len(document.paragraphs),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--source",
        type=Path,
        required=True,
        help="受控外部迁移源；该文件不得复制到Git工作树或生产发布包。",
    )
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--allow-source-drift", action="store_true")
    args = parser.parse_args()
    result = build_template(
        args.source.resolve(),
        args.output.resolve(),
        allow_source_drift=args.allow_source_drift,
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
