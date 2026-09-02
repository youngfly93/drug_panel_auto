#!/usr/bin/env python3
# 步骤: 75 肺癌329/588综合版模板内容对齐
# 上游: panels/lung_588_pdl1/templates/lung_588_pdl1_historical_golden_v1.docx、panels/lung_329_pdl1/rules/knowledge_coverage.yaml
# 输出: 同路径329/588评审模板（确定性变量化更新）
# 种子: 无（确定性DOCX XML变换）
"""Align the registered lung comprehensive templates with the final-report contract."""

from __future__ import annotations

import hashlib
import math
from pathlib import Path

import yaml
from docx import Document
from docx.oxml.ns import qn

from build_lung588_historical_golden_template import _set_table_loop_expressions
from build_lung588_template import _normalize_zip_metadata, _replace_cell_text


ROOT = Path(__file__).resolve().parents[1]
LUNG588 = (
    ROOT
    / "panels/lung_588_pdl1/templates/lung_588_pdl1_historical_golden_v1.docx"
)
LUNG329 = ROOT / "panels/lung_329_pdl1/templates/lung_329_pdl1_golden_template_v2.docx"
LUNG329_COVERAGE = ROOT / "panels/lung_329_pdl1/rules/knowledge_coverage.yaml"


def _table_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def _set_or_update_loop(table, collection: str, expressions: tuple[str, ...]) -> None:
    marker = f"{{%tr for row in {collection} %}}"
    if marker not in _table_text(table):
        _set_table_loop_expressions(table, collection, expressions)
        return
    start_index = next(
        index
        for index, row in enumerate(table.rows)
        if marker in " ".join(cell.text for cell in row.cells)
    )
    body = table.rows[start_index + 1]
    if len(body.cells) != len(expressions):
        raise ValueError(f"{collection}: unexpected column count")
    for cell, expression in zip(body.cells, expressions):
        _replace_cell_text(cell, expression)


def _replace_exact_cell_value(document, old: str, new: str) -> None:
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() == old:
                    _replace_cell_text(cell, new)


def _set_paragraph_text_with_cjk_font(paragraph, value: str) -> None:
    """Replace paragraph text without losing an explicit production CJK font."""

    paragraph.text = value
    for run in paragraph.runs:
        run.font.name = "微软雅黑"
        r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        r_fonts.set(qn("w:ascii"), "微软雅黑")
        r_fonts.set(qn("w:hAnsi"), "微软雅黑")
        r_fonts.set(qn("w:eastAsia"), "微软雅黑")


def _align_588(document) -> None:
    if len(document.tables) != 69:
        raise ValueError(f"lung588 template table count changed: {len(document.tables)}")
    tables = document.tables
    _set_or_update_loop(
        tables[1],
        "targeted_drug_tips",
        (
            "{{ row.gene_display }}",
            "{{ row.variant_site }}",
            "{{ row.benefit_drugs }}",
            "{{ row.caution_drugs }}",
        ),
    )
    _set_or_update_loop(
        tables[2],
        "lung_guideline_drug_results",
        (
            "{{ row.gene }}",
            "{{ row.drugs }}",
            "{{ row.clinical_note }}",
            "{{ row.result }}",
        ),
    )
    _set_or_update_loop(
        tables[4],
        "chemotherapy_summary_rows",
        ("化疗药物小结", "{{ row.summary }}"),
    )
    _set_or_update_loop(
        tables[6],
        "targeted_drug_introductions",
        ("{{ row.drug_name }}", "{{ row.gene }}", "{{ row.introduction }}"),
    )
    _set_or_update_loop(
        tables[13],
        "chemotherapy_predictions",
        ("{{ row.drug }}", "{{ row.genes }}", "{{ row.efficacy }}", "{{ row.toxicity }}"),
    )
    _set_or_update_loop(
        tables[14],
        "chemotherapy_regimen_predictions",
        ("{{ row.regimen }}", "{{ row.genes }}", "{{ row.efficacy }}", "{{ row.toxicity }}"),
    )
    _set_or_update_loop(
        tables[15],
        "chemotherapy_dosage_rows",
        ("{{ row.regimen }}", "{{ row.dosage }}"),
    )
    for table in tables[16:49]:
        text = _table_text(table)
        if "{%tr for row in drug_" not in text:
            continue
        start_index = next(
            index
            for index, row in enumerate(table.rows)
            if "{%tr for row in drug_" in " ".join(cell.text for cell in row.cells)
        )
        _replace_cell_text(table.rows[start_index + 1].cells[0], '{{ row.DrugDisplay or "" }}')

    _replace_exact_cell_value(document, "{{ pdl1_result }}", "{{ pdl1_result_display }}")
    _replace_exact_cell_value(document, "{{ pdl1_tps }}%", "{{ pdl1_tps_display }}")
    _replace_exact_cell_value(document, "{{ pdl1_tps }}", "{{ pdl1_tps_display }}")
    _replace_exact_cell_value(document, "{{ pdl1_cps }}", "{{ pdl1_cps_display }}")

    for paragraph in document.paragraphs:
        text = paragraph.text or ""
        if text.startswith("以下药物基因组学明细由本病例Excel的CtDrug表原样映射") or text.startswith(
            "化疗药物小结：{{ chemotherapy_summary_text }}"
        ):
            _set_paragraph_text_with_cjk_font(
                paragraph, "化疗药物小结：{{ chemotherapy_summary_text }}"
            )
        if (
            "映射Excel中的药物基因组学明细供报告组评审；当前不自动形成患者级化疗方案"
            in text
            or text.startswith("依据病例Excel的CtDrug表生成化疗药物小结")
        ):
            _set_paragraph_text_with_cjk_font(
                paragraph,
                "依据病例Excel的CtDrug表生成化疗药物小结、有效性/毒副作用表、"
                "方案用法表及分级位点附录，供报告组评审。",
            )


def _replace_package_text(document, replacements: tuple[tuple[str, str], ...]) -> None:
    for part in document.part.package.parts:
        element = getattr(part, "element", None)
        if element is None:
            element = getattr(part, "_element", None)
        if element is None:
            continue
        for node in element.iter(qn("w:t")):
            value = node.text or ""
            for old, new in replacements:
                value = value.replace(old, new)
            node.text = value


def _replace_package_exact_text(document, replacements: dict[str, str]) -> None:
    """Replace complete text nodes only, avoiding unrelated numeric substrings."""

    for part in document.part.package.parts:
        element = getattr(part, "element", None)
        if element is None:
            element = getattr(part, "_element", None)
        if element is None:
            continue
        for node in element.iter(qn("w:t")):
            if node.text in replacements:
                node.text = replacements[node.text]


def _set_329_gene_table(document) -> None:
    """Build the 329-gene appendix from the governed panel gene universe."""

    coverage = yaml.safe_load(LUNG329_COVERAGE.read_text(encoding="utf-8")) or {}
    genes = [str(value).strip().upper() for value in coverage.get("reportable_genes", [])]
    if len(genes) != 329 or len(set(genes)) != 329:
        raise ValueError("lung329 governed reportable_genes must contain 329 unique genes")

    gene_table = next(
        table
        for table in document.tables
        if "Gene List for MLseq" in table.cell(0, 0).text
    )
    column_count = len(gene_table.columns)
    required_data_rows = math.ceil(len(genes) / column_count)
    for row in list(gene_table.rows[required_data_rows + 1 :]):
        row._element.getparent().remove(row._element)

    _replace_cell_text(gene_table.rows[0].cells[0], "Gene List for MLseq (n=329)")
    cells = [cell for row in gene_table.rows[1:] for cell in row.cells]
    for index, cell in enumerate(cells):
        _replace_cell_text(cell, genes[index] if index < len(genes) else "")


def _save(document, path: Path) -> str:
    document.save(path)
    _normalize_zip_metadata(path)
    return hashlib.sha256(path.read_bytes()).hexdigest()


def main() -> int:
    lung588 = Document(LUNG588)
    _align_588(lung588)
    hash_588 = _save(lung588, LUNG588)

    lung329 = Document(LUNG588)
    _replace_package_text(
        lung329,
        (
            ("肺癌588基因+PD-L1", "肺癌329基因+PD-L1"),
            ("肺癌588", "肺癌329"),
            ("588基因", "329基因"),
            ("n=588", "n=329"),
        ),
    )
    # In the historical source this sentence splits the number and "个基因"
    # across separate Word runs, so substring replacement cannot see the full
    # phrase. Restrict the fallback to an exact text node.
    _replace_package_exact_text(lung329, {"588": "329"})
    _set_329_gene_table(lung329)
    lung329.core_properties.title = "肺癌329基因+PD-L1报告组评审候选模板"
    hash_329 = _save(lung329, LUNG329)
    print(f"lung588_sha256={hash_588}")
    print(f"lung329_sha256={hash_329}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
