#!/usr/bin/env python3
# 步骤: 72 肺癌588受控试运行合成边界报告验证
# 上游: panels/lung_588_pdl1/、config/、肺588模板
# 输出: .work/lung588_pilot_boundary_suite/validation.json 与7份合成报告
# 种子: 20260724（固定合成病例，不含真实患者信息）
"""Generate seven deterministic, de-identified lung588 pilot boundary reports."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import subprocess
import sys
from pathlib import Path
from typing import Any

from docx import Document
from PIL import Image, ImageDraw


ROOT = Path(__file__).resolve().parents[1]
PROFILE_ID = "legacy_unspecified_ihc_transcription_v1"


def _variant(
    gene: str,
    c_hgvs: str,
    p_hgvs: str,
    gene_class: str,
    frequency: float,
    *,
    transcript: str,
    chromosome: str,
    exon: str,
) -> dict[str, Any]:
    return {
        "ExistIn552": gene_class,
        "Gene_Symbol": gene,
        "Transcript": transcript,
        "Chr": chromosome,
        "Exon": exon,
        "cHGVS": c_hgvs,
        "pHGVS_S": p_hgvs,
        "Mutation_Type": "SNV",
        "Freq(%)": frequency,
    }


SCENARIOS: tuple[dict[str, Any], ...] = (
    {
        "id": "SYN-L588-01-PDL1-NEGATIVE",
        "tps": 0,
        "cps": 0,
        "result": "阴性",
        "tmb": 0,
        "msi": "MSS",
        "expected_targeted_drug_count": 0,
        "variants": [
            _variant(
                "TP53",
                "c.734G>A",
                "p.G245D",
                "Ⅱ类",
                30,
                transcript="NM_000546.6",
                chromosome="17",
                exon="7",
            )
        ],
    },
    {
        "id": "SYN-L588-02-PDL1-LOW-LOWER",
        "tps": 1,
        "cps": 1,
        "result": "阳性（低表达）",
        "tmb": 5,
        "msi": "MSS",
        "expected_targeted_drug_count": 1,
        "variants": [
            _variant(
                "BRAF",
                "c.1799T>A",
                "p.V600E",
                "Ⅰ类",
                28.2,
                transcript="NM_004333.6",
                chromosome="7",
                exon="15",
            )
        ],
    },
    {
        "id": "SYN-L588-03-PDL1-LOW-UPPER",
        "tps": 49,
        "cps": 60,
        "result": "阳性（低表达）",
        "tmb": 9.9,
        "msi": "MSS",
        "expected_targeted_drug_count": 0,
        "variants": [
            _variant(
                "BRAF",
                "c.1781A>G",
                "p.D594G",
                "Ⅱ类",
                12.5,
                transcript="NM_004333.6",
                chromosome="7",
                exon="15",
            )
        ],
    },
    {
        "id": "SYN-L588-04-PDL1-HIGH-LOWER",
        "tps": 50,
        "cps": 52,
        "result": "阳性（高表达）",
        "tmb": 10,
        "msi": "MSS",
        "expected_targeted_drug_count": 1,
        "variants": [
            _variant(
                "ERBB2",
                "c.1979G>A",
                "p.G660D",
                "Ⅰ类",
                8.5,
                transcript="NM_004448.4",
                chromosome="17",
                exon="17",
            )
        ],
    },
    {
        "id": "SYN-L588-05-PDL1-HIGH-UPPER",
        "tps": 100,
        "cps": 100,
        "result": "阳性（高表达）",
        "tmb": 20,
        "msi": "MSS",
        "expected_targeted_drug_count": 0,
        "variants": [
            _variant(
                "EGFR",
                "c.2573T>G",
                "p.L858R",
                "Ⅰ类",
                42,
                transcript="NM_005228.5",
                chromosome="7",
                exon="21",
            )
        ],
    },
    {
        "id": "SYN-L588-06-MSIH-TMBH",
        "tps": 5,
        "cps": 6,
        "result": "阳性（低表达）",
        "tmb": 30,
        "msi": "MSI-H",
        "expected_targeted_drug_count": 0,
        "variants": [
            _variant(
                "KRAS",
                "c.34G>T",
                "p.G12C",
                "Ⅰ类",
                25,
                transcript="NM_004985.5",
                chromosome="12",
                exon="2",
            )
        ],
    },
    {
        "id": "SYN-L588-07-MULTI-VARIANT",
        "tps": 5,
        "cps": 6,
        "result": "阳性（低表达）",
        "tmb": 6.3,
        "msi": "MSS",
        "expected_targeted_drug_count": 0,
        "variants": [
            _variant(
                "TP53",
                "c.578A>T",
                "p.H193L",
                "Ⅱ类",
                40 - index,
                transcript="NM_000546.6",
                chromosome="17",
                exon="6",
            )
            for index in range(12)
        ],
    },
)


def _source_revision() -> str:
    configured = str(os.environ.get("REPORTGEN_SOURCE_REVISION") or "").strip()
    if len(configured) == 40 and all(
        character in "0123456789abcdef" for character in configured.lower()
    ):
        return configured.lower()
    revision_file = ROOT / "REVISION"
    if revision_file.is_file():
        revision = revision_file.read_text(encoding="utf-8").strip()
        if len(revision) == 40:
            return revision
    completed = subprocess.run(
        ["git", "-C", str(ROOT), "rev-parse", "HEAD"],
        check=False,
        capture_output=True,
        text=True,
        timeout=5,
    )
    return completed.stdout.strip() if completed.returncode == 0 else ""


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _visible_text(path: Path) -> str:
    document = Document(path)
    return "\n".join(
        [
            *(paragraph.text for paragraph in document.paragraphs),
            *(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ),
        ]
    )


def _write_synthetic_pdl1_image(path: Path, case_id: str) -> None:
    """Write a deterministic, non-clinical image for the image pipeline gate."""

    image = Image.new("RGB", (960, 640), "white")
    draw = ImageDraw.Draw(image)
    digest = hashlib.sha256(case_id.encode("utf-8")).digest()
    for index in range(72):
        x = (digest[index % len(digest)] * 31 + index * 53) % 920
        y = (digest[(index + 9) % len(digest)] * 23 + index * 47) % 600
        diameter = 12 + digest[(index + 15) % len(digest)] % 22
        draw.ellipse(
            (x, y, x + diameter, y + diameter),
            fill=(110, 75, 145 + digest[(index + 3) % len(digest)] % 70),
        )
    draw.rectangle((10, 10, 949, 629), outline=(75, 75, 75), width=3)
    image.save(path, format="PNG", optimize=True)


def _build_excel_data(
    scenario: dict[str, Any],
    xlsx_path: Path,
):
    from reportgen.models.excel_data import ExcelDataSource

    case_id = scenario["id"]
    image_path = xlsx_path.with_suffix(".pdl1.png")
    _write_synthetic_pdl1_image(image_path, case_id)
    form_fields = (
        "pdl1_tps",
        "pdl1_cps",
        "pdl1_result",
        "pdl1_image_path",
        "pdl1_assay_profile_id",
        "pdl1_source_record_id",
        "pdl1_source_record_date",
        "pdl1_specimen_id",
        "pdl1_image_disposition",
        "lung_histology",
        "disease_extent",
        "prior_systemic_therapy",
        "companion_diagnostic_status",
    )
    return ExcelDataSource(
        file_path=str(xlsx_path),
        single_values={
            "患者姓名": case_id,
            "样本编号": case_id,
            "性别": "男",
            "年龄": 60,
            "送检医院": "合成验证机构",
            "癌种": "肺癌",
            "检测项目": "肺癌588基因+PD-L1",
            "报告日期": "2026-07-24",
            "TMB": scenario["tmb"],
            "MSI状态": scenario["msi"],
            "PD-L1 TPS": scenario["tps"],
            "PD-L1 CPS": scenario["cps"],
            "PD-L1结果": scenario["result"],
            "PD-L1病例图片": str(image_path),
            "PD-L1检测方案": PROFILE_ID,
            "PD-L1原始记录编号": f"SYNTHETIC-IHC-{case_id}",
            "PD-L1原始记录日期": "2026-07-24",
            "PD-L1检测标本标识": f"SYNTHETIC-SPECIMEN-{case_id}",
            "PD-L1图像处置": "病例专属图像（报告展示）",
            "肺癌病理类型": "非小细胞肺癌",
            "疾病范围": "转移性",
            "既往系统治疗": "已接受",
            "伴随诊断适配状态": "已确认符合",
        },
        table_data={
            "Variations": scenario["variants"],
            "TMB": [],
            "Msisensor": [],
        },
        sheet_names=["Variations", "TMB", "Msisensor"],
        metadata={
            "field_source_overrides": {
                field: {
                    "source": "form",
                    "source_key": field,
                    "source_detail": "synthetic_controlled_pilot_form",
                }
                for field in form_fields
            }
        },
    )


def run(output_dir: Path, *, require_visual: bool, dpi: int) -> dict[str, Any]:
    os.environ["RG_WEB_UPSTREAM_ROOT"] = str(ROOT)
    os.environ["REPORTGEN_DISABLED_PROJECT_TYPES"] = ""
    for import_path in (str(ROOT), str(ROOT / "backend")):
        if import_path not in sys.path:
            sys.path.insert(0, import_path)

    from reportgen.core.report_generator import ReportGenerator
    from reportgen.panels.loader import load_panel_package

    output_dir.mkdir(parents=True, exist_ok=True)
    template = load_panel_package(
        "lung_588_pdl1",
        project_root=ROOT,
    ).resolve_template_file()
    generator = ReportGenerator(
        config_dir=str(ROOT / "config"),
        log_level="ERROR",
    )
    case_rows: list[dict[str, Any]] = []
    all_ids = {scenario["id"] for scenario in SCENARIOS}
    for scenario in SCENARIOS:
        case_id = scenario["id"]
        xlsx_path = output_dir / f"{case_id}.xlsx"
        xlsx_path.write_bytes(b"REPORTGEN_SYNTHETIC_FIXTURE\n")
        result = generator.generate(
            excel_file=str(xlsx_path),
            excel_data=_build_excel_data(scenario, xlsx_path),
            template_file=str(template),
            output_dir=str(output_dir / case_id),
            output_filename=f"{case_id}.docx",
            strict_mode=True,
            return_context=True,
            template_contract_mode="fail",
            project_type="lung_588_pdl1",
            project_name="肺癌588基因+PD-L1",
            qa_visual_render="all" if require_visual else None,
            qa_visual_render_required=require_visual,
            qa_visual_render_dpi=dpi,
            qa_visual_render_timeout_seconds=180,
        )
        output_path = Path(str(result.get("output_file") or ""))
        failures: list[str] = []
        if not result.get("success"):
            failures.append("generation_failed")
        if not output_path.is_file():
            failures.append("output_missing")
            visible = ""
        else:
            visible = _visible_text(output_path)
            if "__PDL1_CASE_IMAGE__" in visible:
                failures.append("pdl1_image_marker_not_replaced")
            if "图1. 免疫组化：PD-L1" not in visible:
                failures.append("missing_pdl1_image_caption")
            if "22C3" in visible:
                failures.append("invented_22c3_method")
            if "第三部分：基因变异及相应靶向/免疫药物解析" not in visible:
                failures.append("missing_part3_heading")
            if "肺癌专属知识当前未启用" in visible:
                failures.append("stale_part3_disabled_notice")
            if (
                scenario["expected_targeted_drug_count"]
                and "【待报告组审】" not in visible
            ):
                failures.append("missing_targeted_drug_review_notice")
            for other_id in sorted(all_ids - {case_id}):
                if other_id in visible:
                    failures.append(f"cross_case_leak:{other_id}")
        context = result.get("context") or {}
        targeted_count = len(context.get("targeted_drug_tips") or [])
        if targeted_count != scenario["expected_targeted_drug_count"]:
            failures.append("exact_targeted_drug_count_mismatch")
        if not context.get("gene_knowledge_sections"):
            failures.append("part3_sections_missing")

        qa_status = result.get("qa_status")
        qa_payload: dict[str, Any] = {}
        qa_path = Path(str(result.get("qa_report_file") or ""))
        if qa_path.is_file():
            qa_payload = json.loads(qa_path.read_text(encoding="utf-8"))
        if qa_status not in {"PASS", "WARN"}:
            failures.append(f"qa_not_nonblocking:{qa_status or 'missing'}")
        visual_status = (
            ((qa_payload.get("checks") or {}).get("visual_render") or {}).get(
                "status"
            )
        )
        if require_visual and visual_status != "PASS":
            failures.append(f"visual_qa_not_pass:{visual_status or 'missing'}")
        case_rows.append(
            {
                "case_id": case_id,
                "expected_outcome": "PASS",
                "status": "PASS" if not failures else "FAIL",
                "pdl1_tps": scenario["tps"],
                "pdl1_cps": scenario["cps"],
                "pdl1_result": scenario["result"],
                "tmb_value": scenario["tmb"],
                "msi_status": scenario["msi"],
                "input_variant_count": len(scenario["variants"]),
                "expected_targeted_drug_count": scenario[
                    "expected_targeted_drug_count"
                ],
                "runtime_targeted_drug_count": targeted_count,
                "runtime_part3_section_count": len(
                    context.get("gene_knowledge_sections") or []
                ),
                "qa_status": qa_status,
                "visual_qa_status": visual_status,
                "output_sha256": (
                    _sha256(output_path) if output_path.is_file() else ""
                ),
                "failures": failures,
            }
        )

    payload = {
        "schema_version": "1.0",
        "panel_id": "lung_588_pdl1",
        "suite_id": "lung588_controlled_pilot_synthetic_boundary_v1",
        "source_revision": _source_revision(),
        "contains_real_patient_data": False,
        "synthetic_case_count": len(case_rows),
        "passed_case_count": sum(row["status"] == "PASS" for row in case_rows),
        "visual_qa_required": require_visual,
        "status": (
            "PASS"
            if case_rows and all(row["status"] == "PASS" for row in case_rows)
            else "FAIL"
        ),
        "cases": case_rows,
        "release_boundary": {
            "counts_as_real_case_uat": False,
            "counts_as_engineering_boundary_coverage": True,
            "fixed_minimum_real_case_count": None,
            "treatment_inference_allowed": (
                "registered exact events in pilot drafts with visible review/context notices"
            ),
        },
    }
    receipt = output_dir / "validation.json"
    receipt.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return payload


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=(ROOT / ".work" / "lung588_pilot_boundary_suite_20260724"),
    )
    parser.add_argument("--require-visual", action="store_true")
    parser.add_argument("--dpi", type=int, default=120)
    args = parser.parse_args()
    payload = run(
        args.output_dir.resolve(),
        require_visual=args.require_visual,
        dpi=args.dpi,
    )
    print(
        json.dumps(
            {
                "status": payload["status"],
                "synthetic_case_count": payload["synthetic_case_count"],
                "passed_case_count": payload["passed_case_count"],
                "visual_qa_required": payload["visual_qa_required"],
                "output": str(args.output_dir.resolve() / "validation.json"),
            },
            ensure_ascii=False,
        )
    )
    return 0 if payload["status"] == "PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
