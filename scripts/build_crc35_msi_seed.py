#!/usr/bin/env python3
"""Build a local scrubbed CRC35+MSI seed from a selected reviewed DOCX.

This script reads a local ignored candidate manifest or an explicit source
path, extracts known scalar patient/sample tokens, and writes the scrubbed seed
under tmp/. It does not print extracted PHI values.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.build_golden_template_seed import (  # noqa: E402
    build_seed,
    count_tokens_in_zip,
    replace_tokens_in_xml_text,
    sha256_file,
    should_text_process,
)


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
DEFAULT_MANIFEST = Path("tmp/panel_inventory/crc_small_candidate_manifest.local.json")
DEFAULT_OUTPUT = Path("tmp/golden_template_seed/crc_35_msi_seed.docx")


def text(element) -> str:
    return "".join(
        node.text or "" for node in element.xpath(".//w:t", namespaces=NS)
    ).strip()


def load_source_from_manifest(manifest: Path, product: str, selection: str) -> Path:
    data = json.loads(manifest.read_text(encoding="utf-8"))
    try:
        selected = data["products"][product]["selected"]
    except KeyError as exc:
        raise KeyError(f"product not found in local manifest: {product}") from exc
    for item in selected:
        if item.get("selection") == selection:
            return Path(item["source_path"])
    raise KeyError(f"selection not found for {product}: {selection}")


def read_document_root(path: Path):
    with zipfile.ZipFile(path) as archive:
        return etree.fromstring(archive.read("word/document.xml"))


def extract_replacements(root) -> dict[str, str]:
    replacements: dict[str, str] = {}
    scalars = extract_scalar_values(root)
    for key in ("report_number", "sample_id"):
        value = scalars.get(key, "")
        if value:
            replacements[value] = "{{ " + key + " }}"
    return replacements


def extract_scalar_values(root) -> dict[str, str]:
    scalars: dict[str, str] = {}
    tables = root.xpath(".//w:tbl", namespaces=NS)
    if not tables:
        raise ValueError("source DOCX has no tables")

    rows = tables[0].xpath("./w:tr", namespaces=NS)
    cell_texts = [
        [text(cell) for cell in row.xpath("./w:tc", namespaces=NS)]
        for row in rows
    ]
    position_fields = {
        (0, 1): "patient_name",
        (0, 3): "gender",
        (1, 1): "age",
        (1, 3): "sample_type",
        (2, 1): "clinical_diagnosis",
    }
    for (row, col), field in position_fields.items():
        try:
            value = cell_texts[row][col].strip()
        except IndexError:
            continue
        if value:
            scalars[field] = value

    visible = "\n".join(
        text(element)
        for element in root.xpath(".//w:p|.//w:tc", namespaces=NS)
        if text(element)
    )
    report_numbers = sorted(
        set(re.findall(r"(?i)MLJY[-_ ]?[A-Z]{1,3}\d{5,}", visible))
    )
    sample_ids = sorted(
        set(re.findall(r"(?i)(?:LZ|LW|LB|LC|LG|LM)\d{5,}", visible))
    )
    if report_numbers:
        scalars["report_number"] = report_numbers[0]
    if sample_ids:
        scalars["sample_id"] = sample_ids[0]

    for paragraph in [text(p) for p in root.xpath(".//w:p", namespaces=NS)]:
        if "报告日期" not in paragraph and "送检日期" not in paragraph:
            continue
        for date in re.findall(
            r"\d{8}|\d{4}[./-]\d{1,2}[./-]\d{1,2}|\d{4}年\d{1,2}月\d{1,2}日",
            paragraph,
        ):
            scalars.setdefault("report_date", date)

    required = {"patient_name", "gender", "age", "sample_type", "clinical_diagnosis"}
    missing = sorted(required - set(scalars))
    if missing:
        raise ValueError(f"missing scalar fields in CRC35 seed source: {missing}")
    return scalars


def replace_paragraph_text(paragraph, replacement: str) -> None:
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(replacement)
        return
    target_index = next((idx for idx, run in enumerate(runs) if run.text), 0)
    for idx, run in enumerate(runs):
        run.text = replacement if idx == target_index else ""


def patch_targeted_scalars(path: Path, scalars: dict[str, str]) -> None:
    """Patch short/common scalar fields without replacing them globally."""
    doc = Document(str(path))
    if not doc.tables:
        raise ValueError("seed DOCX has no tables")

    table = doc.tables[0]
    cell_map = {
        (0, 1): "patient_name",
        (0, 3): "gender",
        (1, 1): "age",
        (1, 3): "sample_type",
        (2, 1): "clinical_diagnosis",
    }
    for (row, col), field in cell_map.items():
        table.cell(row, col).text = "{{ " + field + " }}"

    patient = scalars.get("patient_name", "")
    gender = scalars.get("gender", "")
    date_pattern = re.compile(
        r"\d{8}|\d{4}[./-]\d{1,2}[./-]\d{1,2}|\d{4}年\d{1,2}月\d{1,2}日"
    )
    for paragraph in doc.paragraphs:
        text_value = paragraph.text
        if patient and patient in text_value and "尊敬" in text_value:
            replace_paragraph_text(
                paragraph,
                "尊敬的      {{ patient_name }}      {{ gender }}士：",
            )
        elif "报告日期" in text_value and date_pattern.search(text_value):
            replace_paragraph_text(
                paragraph,
                date_pattern.sub("{{ report_date }}", text_value),
            )
        elif gender and "尊敬的" in text_value and gender in text_value:
            replace_paragraph_text(
                paragraph,
                "尊敬的      {{ patient_name }}      {{ gender }}士：",
            )

    doc.save(str(path))


def replace_high_risk_residuals(path: Path, scalars: dict[str, str]) -> dict:
    """Scrub high-specific tokens in headers, footers, and text boxes."""
    replacements = {
        value: "{{ " + field + " }}"
        for field in ("patient_name", "report_number", "sample_id", "report_date")
        if (value := scalars.get(field, ""))
    }
    replacement_counts = {
        field: 0
        for field in scalars
        if field in {"patient_name", "report_number", "sample_id", "report_date"}
    }
    if not replacements:
        return {"replacement_counts": replacement_counts, "residual_counts": {}}

    tmp_path = path.with_suffix(".tmp.docx")
    token_to_field = {
        token: replacement.strip("{} ").strip()
        for token, replacement in replacements.items()
    }
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(
        tmp_path, "w", zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item)
            if should_text_process(item.filename):
                text_value = data.decode("utf-8", errors="ignore")
                if item.filename.endswith(".xml"):
                    text_value, counts = replace_tokens_in_xml_text(
                        text_value,
                        replacements,
                    )
                else:
                    counts = {}
                    for token, replacement in replacements.items():
                        count = text_value.count(token)
                        if count:
                            text_value = text_value.replace(token, replacement)
                        counts[token] = count
                for token, count in counts.items():
                    field = token_to_field[token]
                    replacement_counts[field] = replacement_counts.get(field, 0) + count
                data = text_value.encode("utf-8")
            zout.writestr(item, data)
    tmp_path.replace(path)
    residual_counts = count_tokens_in_zip(path, tuple(replacements))
    return {
        "replacement_counts": replacement_counts,
        "residual_counts": {
            token_to_field[token]: count for token, count in residual_counts.items()
        },
    }


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=None)
    parser.add_argument("--manifest", type=Path, default=DEFAULT_MANIFEST)
    parser.add_argument("--product", default="结直肠癌35基因+msi")
    parser.add_argument("--selection", default="median")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return parser


def main() -> int:
    args = build_parser().parse_args()
    source = args.source or load_source_from_manifest(
        args.manifest,
        args.product,
        args.selection,
    )
    root = read_document_root(source)
    scalars = extract_scalar_values(root)
    replacements = extract_replacements(root)
    result = build_seed(
        source=source,
        output=args.output,
        replacements=replacements,
        protected_tokens=tuple(replacements.keys()),
        allow_commit_output=False,
        allow_residual=False,
        project_root=Path.cwd(),
    )
    patch_targeted_scalars(args.output, scalars)
    high_risk = replace_high_risk_residuals(args.output, scalars)
    final_success = not any(high_risk["residual_counts"].values())
    manifest_path = args.output.with_suffix(".manifest.json")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["output_sha256"] = sha256_file(args.output)
    manifest["crc35_targeted_scalar_fields"] = sorted(
        key for key in scalars if key in {
            "age",
            "clinical_diagnosis",
            "gender",
            "patient_name",
            "report_number",
            "report_date",
            "sample_id",
            "sample_type",
        }
    )
    manifest["crc35_high_risk_replacement_counts"] = high_risk[
        "replacement_counts"
    ]
    manifest["crc35_high_risk_residual_counts"] = high_risk["residual_counts"]
    manifest["success"] = bool(result["success"] and final_success)
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(
        json.dumps(
            {
                "output": result["output"],
                "manifest": result["manifest"],
                "success": manifest["success"],
                "global_replacement_count": len(replacements),
                "targeted_scalar_count": len(scalars),
                "high_risk_residual_total": sum(
                    high_risk["residual_counts"].values()
                ),
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0 if manifest["success"] else 2


if __name__ == "__main__":
    raise SystemExit(main())
