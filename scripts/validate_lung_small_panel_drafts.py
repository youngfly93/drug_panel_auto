#!/usr/bin/env python3
# 步骤: 76 小肺癌派生输入逐例工程验收
# 上游: derive_panel_input.py 输出、真实肺癌超集 Excel、可选 PD-L1 转录记录
# 输出: .work/ 下病例隔离 Word、QA、上下文与不含身份的验收回执
# 种子: 无（不生成任何临床字段；病例标识仅用工程别名）
# ruff: noqa: E402
"""Validate derived drafts on the shipping Linux renderer, not clinical parity.

Run serially or launch at most two separate invocations on iyun129. Each output
directory must be new. Raw input files are CASE-LUNG-{A,B,C}.xlsx; derived files
are CASE-LUNG-X-derived-lung_{13,62}.xlsx. PD-L1 never comes from the NGS file.
"""

from __future__ import annotations

import argparse
import copy
import hashlib
import json
import re
import subprocess
import sys
from collections import Counter
from datetime import date
from pathlib import Path

import yaml
from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "backend"))
sys.path.insert(0, str(ROOT / "scripts"))

from app.services.project_identity import resolve_project_identity
from app.services.reportgen_bridge import ReportGenBridge

from reportgen.core.enhancer_registry import get_enhancer
from reportgen.core.template_bridge_358 import load_panel_config
from reportgen.panels.loader import load_panel_package
from scripts.build_lung588_historical_golden_template import DRUG_DETAIL_BINDINGS
from scripts.scan_hardcoded_literals import scan_docx
from scripts.two_case_leak_test import docx_visible_text


def digest(path):
    return hashlib.sha256(Path(path).read_bytes()).hexdigest()


def write_json(path, value):
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2, default=str) + "\n")


def expected_variants(excel_path, package):
    """Primary table oracle: member, coding event and explicit class I/II/III.

    Class III is displayed without drug tips, as in the existing 588 primary
    table. Annotation/CNV notes without c.HGVS are not SNVs.
    """
    contract = package.raw.get("derived_input") or {}
    flag = contract.get("membership_column") or "ExistInsmall588"
    genes = set(contract.get("genes") or [])
    workbook = load_workbook(excel_path, read_only=True, data_only=True)
    try:
        rows = workbook["Variations"].iter_rows(values_only=True)
        header = next(rows)
        selected = []
        for cells in rows:
            row = dict(zip(header, cells))
            category = str(row.get("ExistIn552") or "").strip()
            explicit_class = bool(re.fullmatch(r"(?:Ⅰ|Ⅱ|Ⅲ|I|II|III|1|2|3)类", category))
            gene = str(row.get("Gene_Symbol") or "").strip()
            if (
                row.get(flag) in (1, "1")
                and explicit_class
                and (not genes or gene in genes)
                and str(row.get("cHGVS") or "").strip().startswith("c.")
            ):
                selected.append(
                    {
                        "gene": gene,
                        "c_hgvs": str(row.get("cHGVS") or ""),
                        "p_hgvs": str(row.get("pHGVS_S") or ""),
                    }
                )
        return selected
    finally:
        workbook.close()


def pdl1_fields(source_dir, alias):
    if source_dir is None:
        return {}, {"status": "not_provided"}
    path = source_dir / f"{alias}_pdl1_source_record.yaml"
    if not path.is_file():
        return {}, {"status": "not_provided"}
    record = yaml.safe_load(path.read_text())
    source = record["source_record"]
    values = record["transcribed_values"]
    profile = record["assay_profile"]["assay_profile_id"]
    if profile != "legacy_unspecified_ihc_transcription_v1":
        raise ValueError("Unexpected PD-L1 transcription contract")
    image = source_dir / record["image"]["file"]
    if digest(image) != record["image"]["sha256"]:
        raise ValueError("PD-L1 image source checksum mismatch")
    fields = {
        "pdl1_tps": values["pdl1_tps"],
        "pdl1_cps": values["pdl1_cps"],
        "pdl1_result": values["pdl1_result_contract"],
        "pdl1_image_path": str(image.resolve()),
        "pdl1_assay_profile_id": profile,
        "pdl1_source_record_id": "sha256:" + source["docx_sha256"],
        "pdl1_source_record_date": str(source["report_date"]),
        "pdl1_specimen_id": str(source["sample_id"]),
        "pdl1_image_disposition": "病例专属图像（报告展示）",
    }
    return fields, {
        "status": "historical_transcription_review_only",
        "record_sha256": digest(path),
        "source_docx_sha256": source["docx_sha256"],
        "image_sha256": digest(image),
        "clinical_authorization": "pending_report_group",
    }


def baseline_context(bridge, raw_path, clinical):
    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    excel = bridge.read_excel(str(raw_path))
    bridge._inject_clinical_info_into_excel(excel, clinical)
    report = bridge.field_mapper.map(excel, panel_package=package)
    report = bridge.data_cleaner.validate_and_clean(report)
    report = get_enhancer("lung_588_pdl1").enhance(
        report,
        excel,
        field_mapper=bridge.field_mapper,
        gene_knowledge_provider=bridge._build_gene_knowledge_provider(package),
        base_path=str(ROOT),
        project_type="lung_588_pdl1",
        panel_package=package,
    )
    return report.get_template_context()


def canonical_parity_value(value):
    """Rendering converts source None to empty text, without changing evidence."""
    if value is None:
        return ""
    if isinstance(value, dict):
        return {key: canonical_parity_value(item) for key, item in value.items()}
    if isinstance(value, (list, tuple)):
        return [canonical_parity_value(item) for item in value]
    return value


def inspect_word_scope(document, context, package):
    """Compare final Word guideline, assayed genes and every maintained PGx row."""
    def compact(value):
        return re.sub(r"\s+", "", str("" if value is None else value))
    config = load_panel_config(panel_package=package)
    genes, guidelines, pgx = [], [], []
    guide_table_count = gene_table_count = 0
    merged_guideline = merged_pgx = False
    for table in document.tables:
        kind = None
        pgx_headers = 0
        for index, row in enumerate(table.rows):
            # row.cells repeats the same XML cell for every gridSpan column.
            # Native merging can yield 16 physical grid columns but six fields.
            cells, seen = [], set()
            for cell in row.cells:
                if cell._tc not in seen:
                    cells.append(compact(cell.text))
                    seen.add(cell._tc)
            if cells and re.fullmatch(r"肺癌\d+基因检测列表", cells[0]):
                gene_table_count += 1
                kind = "genes"
            elif cells and cells[0] == "检测基因" and "本癌种相关治疗药物" in "".join(cells):
                guide_table_count += 1
                merged_guideline |= index != 0
                kind = "guideline"
            elif len(cells) == 6 and cells[1:4] == ["基因", "检测位点", "等级"]:
                pgx_headers += 1
                kind = "pgx"
            elif any(cells):
                if kind == "genes":
                    genes.extend(cell for cell in cells if cell)
                elif kind == "guideline":
                    guidelines.append(tuple(cells))
                elif kind == "pgx":
                    pgx.append(tuple(cells))
        merged_pgx |= pgx_headers > 1
    guide_rows = context.get("lung_guideline_drug_results") or []
    columns = len(guidelines[0]) if guidelines else 0
    expected_guidelines = [
        tuple(compact(r.get(key)) for key in ("gene", "drugs", "clinical_note", "result"))
        if columns == 4 else (
            compact(r.get("gene")), compact(str(r.get("drugs") or "") + str(r.get("clinical_note") or "")),
            compact(r.get("result")),
        )
        for r in guide_rows
    ]
    expected_pgx = [
        tuple(compact(row.get(key)) for key in ("DrugDisplay", "Gene", "Locus", "Level", "Genotype", "Result"))
        for _, collection in DRUG_DETAIL_BINDINGS for row in context.get(collection) or []
    ]
    failures = []
    if merged_guideline:
        failures.append("word_guideline_table_merged")
    if merged_pgx:
        failures.append("word_pgx_tables_merged")
    if gene_table_count != 1 or Counter(genes) != Counter(config.crc_important_genes):
        failures.append("word_assayed_gene_list_mismatch")
    configured_guide_genes = [
        compact(row.get("display") or "/".join(row.get("genes") or []))
        for row in config.lung_guideline_drug_rows
    ]
    if (
        guide_table_count != 1 or guidelines != expected_guidelines
        or [compact(row.get("gene")) for row in guide_rows] != configured_guide_genes
    ):
        failures.append("word_guideline_rows_mismatch")
    if Counter(pgx) != Counter(expected_pgx):
        failures.append("word_pgx_detail_rows_mismatch")
    return {
        "assayed_gene_count": len(genes), "guideline_row_count": len(guidelines),
        "pgx_detail_row_count": len(pgx), "expected_pgx_detail_row_count": len(expected_pgx),
        "failures": failures,
    }


def run_case(args, panel, case):
    alias = f"CASE-LUNG-{case}"
    output = args.output_dir / panel / alias
    output.mkdir(parents=True, exist_ok=False)
    package = load_panel_package(panel, project_root=ROOT)
    family = panel.removesuffix("_pdl1")
    source = (
        args.derived_dir / f"{alias}-derived-{family}.xlsx"
        if family in ("lung_13", "lung_62")
        else args.raw_dir / f"{alias}.xlsx"
    )
    bridge = ReportGenBridge(
        config_dir=str(ROOT / "config"), template_dir=str(ROOT / "templates")
    )
    excel = bridge.read_excel(str(source))
    identity = resolve_project_identity(
        bridge,
        excel_path=source,
        excel_data=excel,
        requested_project_type=panel if panel.endswith("_pdl1") else None,
    )
    if identity.project_type != panel:
        raise ValueError("NGS family detection or same-family disambiguation failed")
    clinical = {
        "patient_name": "未提供",
        "sample_id": alias,
        "report_date": date.today().isoformat(),
        "project_name": package.display_name,
    }
    pdl1, pdl1_receipt = (
        pdl1_fields(args.pdl1_source_dir, alias)
        if panel.endswith("_pdl1")
        else ({}, {"status": "not_in_product"})
    )
    clinical.update(pdl1)
    write_json(
        output / "form_source.json", {"clinical_info": clinical, "pdl1": pdl1_receipt}
    )
    bridge._inject_clinical_info_into_excel(excel, clinical)
    result = bridge.generator.generate(
        excel_file=str(source),
        template_file=str(package.resolve_template_file()),
        output_dir=str(output),
        output_filename=f"{alias}-{panel}-draft.docx",
        excel_data=excel,
        return_context=True,
        template_contract_mode="fail",
        project_type=panel,
        project_name=package.display_name,
        qa_visual_render=args.render,
        qa_visual_render_required=args.render == "all",
        qa_visual_render_dpi=args.dpi,
        qa_visual_render_timeout_seconds=300,
        qa_visual_render_output_dir=str(output / "visual"),
    )
    write_json(output / "generation.json", result)
    qa = result.get("qa_report") or {}
    context = result.get("context") or {}
    failures = []
    if not result.get("success"):
        failures.append("generation_failed")
    if qa.get("status") != "PASS":
        failures.append("qa_not_pass")
    expected = expected_variants(source, package)
    rows = context.get("variants_2_1") or []
    if sorted(r["gene"] for r in expected) != sorted(
        str(r.get("gene") or "") for r in rows
    ):
        failures.append("excel_variant_membership_mismatch")
    word_variant_genes, word_targeted_genes = [], []
    word_scope = {"failures": ["word_missing"]}
    report_file = Path(result.get("output_file") or output / "missing.docx")
    if report_file.is_file():
        document = Document(report_file)
        word_scope = inspect_word_scope(document, context, package)
        failures.extend(word_scope["failures"])
        for table in document.tables:
            header = "|".join(c.text for c in table.rows[0].cells)
            if len(table.columns) == 9 and "基因突变信息" in header:
                word_variant_genes = [r.cells[0].text.strip() for r in table.rows[2:]]
                for row in expected:
                    matches = [
                        r
                        for r in table.rows[2:]
                        if r.cells[0].text.strip() == row["gene"]
                    ]
                    if not any(
                        row["p_hgvs"].removeprefix("p.")
                        in "|".join(c.text for c in r.cells)
                        for r in matches
                    ):
                        failures.append("word_variant_hgvs_mismatch:" + row["gene"])
            if len(table.columns) == 4 and "潜在获益靶向药物" in header:
                word_targeted_genes = [r.cells[0].text.strip() for r in table.rows[1:]]
        if word_variant_genes != [str(r.get("gene") or "") for r in rows]:
            failures.append("word_variant_rows_mismatch")
        if word_targeted_genes != [
            str(r.get("gene") or "") for r in context.get("targeted_drug_tips") or []
        ]:
            failures.append("word_targeted_rows_mismatch")
        if "报告组评审草稿（非临床交付）" not in docx_visible_text(report_file):
            failures.append("review_boundary_missing")
    base = baseline_context(
        bridge, args.raw_dir / f"{alias}.xlsx", copy.deepcopy(clinical)
    )
    parity_keys = (
        "tmb_value",
        "tmb_status",
        "msi_status",
        "chemotherapy_predictions",
        "chemotherapy_regimen_predictions",
        "chemotherapy_dosage_rows",
        "chemotherapy_summary_text",
        "irinotecan_safety_rows",
    ) + tuple(collection for _, collection in DRUG_DETAIL_BINDINGS)
    differences = [
        key for key in parity_keys
        if canonical_parity_value(context.get(key)) != canonical_parity_value(base.get(key))
    ]
    if differences:
        failures.append("588_biomarker_pgx_parity_mismatch")
    write_json(
        output / "biomarker_pgx_parity.json",
        {
            "keys": parity_keys,
            "differences": differences,
            "actual": {k: context.get(k) for k in parity_keys},
            "baseline": {k: base.get(k) for k in parity_keys},
        },
    )
    blank = {"status": "NOT_RUN"}
    if args.blank_page_gate and report_file.is_file():
        check = subprocess.run(
            [
                sys.executable,
                str(ROOT / "scripts/render_blank_page_check.py"),
                str(report_file),
                "--require-render",
                "--strict-trailing",
                "--dpi",
                str(args.dpi),
                "--output-dir",
                str(output / "blank_pages"),
                "--json",
            ],
            capture_output=True,
            text=True,
            timeout=360,
        )
        (output / "blank_page_gate.log").write_text(check.stdout + check.stderr)
        try:
            blank = json.loads(check.stdout)
        except ValueError:
            blank = {"status": "FAIL", "exit_code": check.returncode}
        if check.returncode or blank.get("status") != "PASS":
            failures.append("blank_page_gate_failed")
    scan = scan_docx(package.resolve_template_file(), tokens=[])
    if scan.hard:
        failures.append("hardcoded_template_literals")
    if panel == "lung_13" and case == "C":
        if (
            set(word_variant_genes) != {"BRAF", "ERBB2", "TP53", "PIK3CA"}
            or len(word_variant_genes) != 4
        ):
            failures.append("acceptance_C13_four_variants")
        if (
            set(word_targeted_genes) != {"BRAF", "ERBB2", "PIK3CA"}
            or len(word_targeted_genes) != 3
        ):
            failures.append("acceptance_C13_three_targeted_rows")
    summary = {
        "panel_id": panel,
        "case": alias,
        "status": "FAIL" if failures else "PASS",
        "input_sha256": digest(source),
        "template_sha256": digest(package.resolve_template_file()),
        "detection": {
            "ngs_family": identity.detection.get("identity_family"),
            "auto_detected_type": identity.detected_project_type,
            "resolved_type": identity.project_type,
        },
        "generation_success": result.get("success"),
        "qa_status": qa.get("status"),
        "qa_issues": qa.get("issues"),
        "page_count": (qa.get("metrics") or {}).get("visual_render_page_count"),
        "expected_variants": expected,
        "word_variant_genes": word_variant_genes,
        "word_targeted_genes": word_targeted_genes,
        "word_scope": word_scope,
        "parity_differences": differences,
        "hard_literal_count": len(scan.hard),
        "blank_page_gate": blank,
        "pdl1_source": pdl1_receipt,
        "failures": failures,
        "same_case_historical_parity": "not_established_by_derived_inputs",
        "production_eligible": False,
    }
    write_json(output / "receipt.json", summary)
    print(
        json.dumps(
            {
                k: summary[k]
                for k in (
                    "panel_id",
                    "case",
                    "status",
                    "qa_status",
                    "page_count",
                    "failures",
                )
            }
        ),
        flush=True,
    )
    return summary


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--raw-dir", type=Path, required=True)
    parser.add_argument("--derived-dir", type=Path, required=True)
    parser.add_argument("--pdl1-source-dir", type=Path)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument(
        "--panel",
        action="append",
        choices=["lung_13", "lung_62", "lung_62_pdl1", "lung_588", "lung_588_pdl1"],
    )
    parser.add_argument("--case", action="append", choices=["A", "B", "C"])
    parser.add_argument("--render", choices=["all", "none"], default="all")
    parser.add_argument("--dpi", type=int, default=120)
    parser.add_argument("--blank-page-gate", action="store_true")
    args = parser.parse_args()
    if ".work" not in args.output_dir.resolve().parts:
        parser.error("Private output must be under .work")
    summaries = [
        run_case(args, panel, case)
        for panel in args.panel or ["lung_13", "lung_62", "lung_62_pdl1"]
        for case in args.case or ["A", "B", "C"]
    ]
    return int(any(row["status"] != "PASS" for row in summaries))


if __name__ == "__main__":
    raise SystemExit(main())
