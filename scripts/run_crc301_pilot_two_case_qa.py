#!/usr/bin/env python3
"""Run a two-case QA gate for the CRC301 golden-template v1 path.

The inputs are synthetic and contain no patient data. The script explicitly
uses the golden template id so the gate remains stable even if defaults change.
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from reportgen.core.golden_case import build_crc_301_msi_golden_excel  # noqa: E402
from reportgen.core.report_generator import ReportGenerator  # noqa: E402
from reportgen.panels.loader import load_panel_package  # noqa: E402


PANEL_ID = "crc_301_msi"
TEMPLATE_ID = "crc_301_msi_golden_template_v1"


def _write_case_b(path: Path) -> Path:
    """Create an alternate CRC301 workbook with different biomarkers/variants."""
    path.parent.mkdir(parents=True, exist_ok=True)
    panel_name = "结直肠癌301基因+MSI"

    meta = pd.DataFrame(
        [
            {
                "患者姓名": "二例测试患者",
                "样本编号": "LZ999302",
                "报告编号": "MLJY-LZ999302",
                "性别": "女",
                "年龄": 64,
                "临床诊断": "结直肠癌",
                "肿瘤类型": "结直肠癌",
                "样本类型": "组织",
                "取材手段": "穿刺",
                "取材部位": "直肠",
                "项目名称": panel_name,
                "检测项目": panel_name,
                "送检日期": "2026-02-10",
                "报告日期": "2026-02-17",
                "检测方法": "NGS高通量测序",
            }
        ]
    )

    variations = pd.DataFrame(
        [
            {
                "Gene_Symbol": "KRAS",
                "Transcript": "NM_004985.5",
                "Chr": "12",
                "ExIn_ID": "EX2",
                "cHGVS": "c.34G>A",
                "pHGVS_S": "p.G12S",
                "Freq(%)": 46.3,
                "Function": "Missense",
                "ExistInsmall301": 1,
                "ExistIn552": "Ⅰ类",
                "CLNSIG": "Pathogenic",
            },
            {
                "Gene_Symbol": "TP53",
                "Transcript": "NM_000546.6",
                "Chr": "17",
                "ExIn_ID": "EX8",
                "cHGVS": "c.844C>T",
                "pHGVS_S": "p.R282W",
                "Freq(%)": 21.4,
                "Function": "Missense",
                "ExistInsmall301": 1,
                "ExistIn552": "Ⅱ类",
                "CLNSIG": "Pathogenic",
            },
            {
                "Gene_Symbol": "APC",
                "Transcript": "NM_000038.6",
                "Chr": "5",
                "ExIn_ID": "EX16",
                "cHGVS": "c.4348C>T",
                "pHGVS_S": "p.R1450*",
                "Freq(%)": 10.2,
                "Function": "Nonsense",
                "ExistInsmall301": 1,
                "ExistIn552": "Ⅲ类",
                "CLNSIG": "Pathogenic",
            },
        ]
    )

    tmb = pd.DataFrame(
        [
            ["TCGA fit", None, None, None],
            ["SampleTP", "Var_num", "Bed_size", "TMB"],
            ["tissue", 188, 10_000_000, 18.8],
        ]
    )
    msisensor = pd.DataFrame(
        [
            ["control", 1000, 4, 0.4, "MSS"],
            ["tumor", 1000, 460, 46.0, "MSI-H"],
        ],
        columns=["Sample", "Total", "Unstable", "Percent", "Status"],
    )
    qc = pd.DataFrame(
        [
            ["Q30", 96.2],
            ["Coverage", 99.6],
            ["Average sequencing depth", 920],
            ["Insert", 175],
        ]
    )
    empty_cnv = pd.DataFrame(columns=["Gene", "Chr", "Start", "End", "CopyNumber"])
    empty_fusion = pd.DataFrame(
        columns=["Gene1", "Gene2", "Chr1", "Pos1", "Chr2", "Pos2"]
    )

    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        meta.to_excel(writer, sheet_name="Meta", index=False)
        variations.to_excel(writer, sheet_name="Variations", index=False)
        tmb.to_excel(writer, sheet_name="TMB", index=False, header=False)
        msisensor.to_excel(writer, sheet_name="Msisensor", index=False)
        qc.to_excel(writer, sheet_name="QC", index=False, header=False)
        empty_cnv.to_excel(writer, sheet_name="Cnv", index=False)
        empty_fusion.to_excel(writer, sheet_name="Fusion", index=False)

    return path


def _docx_summary(path: Path) -> dict[str, Any]:
    doc = Document(path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    return {
        "exists": path.exists(),
        "tables": len(doc.tables),
        "paragraphs": len(doc.paragraphs),
        "part3_marker_removed": "__PART3_MARKER__" not in text,
        "jinja_removed": "{{" not in text and "{%" not in text,
        "has_reading_section": "3. 阅读说明" in text,
        "has_signature_block": "检测者：" in text and "审核者：" in text,
    }


def _run_case(
    generator: ReportGenerator,
    *,
    case_id: str,
    excel_path: Path,
    template_path: Path,
    output_dir: Path,
) -> dict[str, Any]:
    result = generator.generate(
        excel_file=str(excel_path),
        template_file=str(template_path),
        output_dir=str(output_dir),
        output_filename=f"{case_id}.docx",
        project_type=PANEL_ID,
        template_contract_mode="fail",
        return_context=True,
    )
    output_file = Path(str(result.get("output_file") or output_dir / f"{case_id}.docx"))
    qa_report = result.get("qa_report") or {}
    summary = {
        "case_id": case_id,
        "success": bool(result.get("success")),
        "qa_status": result.get("qa_status") or qa_report.get("status"),
        "error_count": len(result.get("errors") or []),
        "qa_issue_count": len(qa_report.get("issues") or []),
        "output_file": str(output_file),
        "qa_report_file": result.get("qa_report_file"),
        "docx": _docx_summary(output_file) if output_file.exists() else {"exists": False},
    }
    summary["ok"] = (
        summary["success"]
        and summary["qa_status"] == "PASS"
        and summary["error_count"] == 0
        and summary["qa_issue_count"] == 0
        and summary["docx"].get("part3_marker_removed")
        and summary["docx"].get("jinja_removed")
        and summary["docx"].get("has_signature_block")
    )
    return summary


def run(output_root: Path) -> dict[str, Any]:
    output_root.mkdir(parents=True, exist_ok=True)
    inputs_dir = output_root / "inputs"
    reports_dir = output_root / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)

    case_a = build_crc_301_msi_golden_excel(inputs_dir / "crc301_case_a.xlsx")
    case_b = _write_case_b(inputs_dir / "crc301_case_b.xlsx")

    package = load_panel_package(PANEL_ID, project_root=ROOT)
    template_path = package.resolve_template_file(TEMPLATE_ID)
    generator = ReportGenerator(config_dir=str(ROOT / "config"), log_level="ERROR")

    cases = [
        _run_case(
            generator,
            case_id="crc301_case_a",
            excel_path=Path(case_a),
            template_path=template_path,
            output_dir=reports_dir,
        ),
        _run_case(
            generator,
            case_id="crc301_case_b",
            excel_path=case_b,
            template_path=template_path,
            output_dir=reports_dir,
        ),
    ]
    manifest = {
        "panel_id": PANEL_ID,
        "template_id": TEMPLATE_ID,
        "template_status": package.templates[TEMPLATE_ID].status,
        "default_template": package.default_template.template_id,
        "case_count": len(cases),
        "cases": cases,
        "ok": all(case["ok"] for case in cases),
    }
    (output_root / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return manifest


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--output-root",
        type=Path,
        default=Path("tmp/crc301_pilot_two_case_qa"),
    )
    return parser


def main() -> int:
    args = build_parser().parse_args()
    manifest = run(args.output_root)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))
    return 0 if manifest["ok"] else 2


if __name__ == "__main__":
    raise SystemExit(main())
