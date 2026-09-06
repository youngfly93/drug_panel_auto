#!/usr/bin/env python3
# 步骤: 75 从受控小肺癌及无 PD-L1 历史终版蒸馏 draft 包
# 上游: SHA 固定母版、lung_draft_template_maps.yaml、588 肺癌规则
# 输出: .work/ 的私有种子/回执及 panels/<id>/ 可复现脱敏模板/规则
# 种子: 无（确定性 OOXML 变换；不生成病例结论）
# ruff: noqa: E402
"""Build draft packages from their own approved historical layout families.

Patient sources and token manifests stay under .work. This is an engineering
distillation, not medical approval or evidence of historical same-case parity.
"""

from __future__ import annotations

import argparse
import copy
import datetime
import hashlib
import json
import re
import sys
from pathlib import Path

import yaml
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "scripts"))

from scripts.build_golden_template_seed import build_seed, count_tokens_in_zip
from scripts.build_lung588_historical_golden_template import (
    DRUG_DETAIL_BINDINGS,
    _prune_unreferenced_images,
    _replace_basic_information,
    _replace_pdl1_table,
    _scrub_story_content,
    _story_parts,
)
from scripts.build_lung588_template import _normalize_zip_metadata
from scripts.repair_docx_relationships import repair_docx
from scripts.scan_hardcoded_literals import RE_CHGVS, RE_PHGVS, scan_docx
from scripts.variableize_golden_template import (
    replace_cell_text,
    replace_paragraph_text,
    variableize_docx,
)

BASE = ROOT / "panels/lung_588_pdl1"
NOTICE = (
    "报告组评审草稿（非临床交付）：本稿仅用于产品工程验证；"
    "尚未完成真实同案历史对照及报告组医学复核。"
)
APPENDIX_NOTICE = (
    "固定附录沿用历史终版，仅供报告组核对；"
    "指南版本、药物适应证和文献时效性尚待医学复核。"
)
IMMUNE_NOTICE = (
    "本节仅核对来源数据和研究性标志物，不构成患者级免疫治疗获益、耐药或超进展判断。"
)
CHEMO_NOTICE = (
    "本节按现有肺癌588历史展示合同映射 CtDrug；UGT1A1 与长春类解释列入"
    "报告组待决清单，不构成自动用药或剂量建议。"
)
EMPTY_APPROVED_DRUG_NOTICE = (
    "本栏目暂无可列示的结构化药物条目，待报告组复核；"
    "不代表无可用治疗方案。"
)
LABEL_FIELDS = {
    "姓名": "patient_name",
    "性别": "gender",
    "年龄": "age",
    "样本类型": "sample_type",
    "样本编号": "sample_id",
    "报告编号": "report_number",
    "报告编码": "report_number",
    "临床诊断": "clinical_diagnosis",
    "癌症类型": "clinical_diagnosis",
    "检测项目": "project_name",
    "送检单位": "hospital",
    "联系方式": "phone",
    "报告时间": "report_date",
    "报告日期": "report_date",
    "既往用药史": "treatment_history",
    "家族史": "family_history",
    "样本条形码": "sample_id",
    "采样部位": "sample_site",
    "样本采集日期": "collection_date",
    "样本接收日期": "receive_date",
}


def sha(path):
    return hashlib.sha256(Path(path).read_bytes()).hexdigest()


def compact(value):
    return re.sub(r"[\s：:]", "", str(value))


def text(element):
    return "".join(n.text or "" for n in element.iter(qn("w:t")))


def field(name):
    return "{{ " + name + " | default('未提供', true) }}"


def heading(doc, value):
    matches = [p for p in doc.paragraphs if compact(p.text) == compact(value)]
    if len(matches) != 1:
        raise ValueError(f"Expected one body heading {value!r}, found {len(matches)}")
    return matches[0]


def paragraph_after(anchor, value, *, title=False):
    node = OxmlElement("w:p")
    anchor.addnext(node)
    paragraph = Paragraph(node, None)
    run = paragraph.add_run(value)
    run.font.name = "微软雅黑"
    run.font.size = Pt(12 if title else 10.5)
    run.bold = title
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = title
    return node


def ensure_table_separator(table, caption):
    """Protect a logical table boundary from native adjacent-table merging."""
    previous = table._tbl.getprevious()
    cursor = previous
    while cursor is not None and cursor.tag == qn("w:p"):
        value = text(cursor).strip()
        if value and not value.startswith("{%"):
            return  # A surviving caption/paragraph already separates tables.
        cursor = cursor.getprevious()
    if previous is None:
        raise ValueError("A named table boundary requires a preceding body node")
    paragraph_after(previous, caption, title=True)


def replace_between(doc, start, end, values=(), keep_tables=()):
    begin, finish = heading(doc, start)._p, heading(doc, end)._p
    nodes = list(doc.element.body)
    a, b = nodes.index(begin), nodes.index(finish)
    if a >= b:
        raise ValueError("Invalid variable block boundaries")
    for node in nodes[a + 1 : b]:
        doc.element.body.remove(node)
    anchor = begin
    for value in values:
        anchor = paragraph_after(anchor, value)
    for table in keep_tables:
        anchor.addnext(table._tbl)
        anchor = table._tbl
    return anchor


def scalar_map(doc, spec):
    """Infer scalar coordinates only from controlled source table labels."""
    cells, replacements = [], {}
    for ti, table in enumerate(doc.tables):
        if ti > spec["tables"]["basic"] and not (
            spec["gene_count"] == 62 and spec["family"] == "A" and ti < 4
        ):
            continue
        if spec["gene_count"] == 62 and spec["family"] == "A" and ti == 2:
            continue  # horizontal sample headers; values are mapped by column
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                name = LABEL_FIELDS.get(compact(cell.text))
                if name and ci + 1 < len(row.cells):
                    target = row.cells[ci + 1]
                    if target._tc is cell._tc:
                        continue
                    cells.append(
                        {"table": ti, "row": ri, "col": ci + 1, "variable": name}
                    )
                    old = target.text.strip()
                    if (
                        name
                        in {
                            "patient_name",
                            "sample_id",
                            "report_number",
                            "hospital",
                            "phone",
                        }
                        and len(old) >= 2
                        and old not in {"--", "未提供"}
                    ):
                        replacements[old] = field(name)
    # Source report/sample ids can also occur in floating cover boxes.
    joined = "\n".join(text(part.element) for part in _story_parts(doc))
    for token in re.findall(r"(?i)\b(?:LW|NGS|ML|LN)\d{6,}(?:-[A-Z0-9]+)?\b", joined):
        replacements[token] = field("sample_id")
    return cells, replacements


def gene_list(doc, spec):
    genes = spec.get("genes")
    if not genes:
        table = doc.tables[spec["tables"]["genes"]]
        genes = [
            cell.text.strip().upper()
            for row in table.rows[1:]
            for cell in row.cells
            if re.fullmatch(r"[A-Z][A-Z0-9-]*", cell.text.strip().upper())
        ]
    if len(genes) != spec["gene_count"] or len(set(genes)) != len(genes):
        raise ValueError(
            "Historical detection gene list does not match the product count"
        )
    return genes


def loop_spec(table, collection, fields, start=1):
    return {
        "table": table,
        "collection": collection,
        "alias": "row",
        "template_row": start,
        "columns": fields,
    }


def clear_case_regions(doc, panel, spec, tables):
    keep = []
    if panel == "lung_13":
        keep = [tables["introductions"], tables["guideline"]]
    elif panel in {"lung_62", "lung_62_pdl1"}:
        keep = [tables["introductions"]]
    replace_between(
        doc, spec["rich_start"], spec["rich_end"], ("__PART3_MARKER__",), keep
    )
    if panel == "lung_62":
        replace_between(
            doc,
            "2.2 NCCN 推荐临床常规靶向药物相关基因检测结果（不限于本癌种）",
            "2.3 其它潜在获益上市药物提示*",
            ("指南表按本产品基因范围展示，固定说明沿用肺癌588审核规则。",),
            (tables["guideline"],),
        )
    elif panel == "lung_62_pdl1":
        replace_between(
            doc,
            "1.2 NCCN 推荐临床常规靶向药物相关基因检测结果（不限于本癌种）",
            "1.3 其它潜在获益上市药物提示*",
            ("本癌种指南相关检测结果见第一部分；不展示本产品范围外的未检出结论。",),
        )
    elif panel == "lung_588":
        replace_between(
            doc,
            "2.3 NCCN 推荐临床常规靶向药物相关基因检测结果（不限于本癌种）",
            "3. 免疫疗效正相关/负相关/超进展基因检测结果",
            ("本癌种指南相关检测结果见结果小结。",),
        )
    if "obsolete_guideline" in tables:
        table = tables["obsolete_guideline"]
        if table._tbl.getparent() is not None:
            parent = table._tbl.getparent()
            prev = table._tbl.getprevious()
            parent.remove(table._tbl)
            if prev is not None and "NCCN" in text(prev):
                parent.remove(prev)
    # Clear former seed-specific result notes (drug brand/count sentences).
    for p in list(doc.paragraphs):
        value = p.text.strip()
        if "上表涉及的" in value and ("药物" in value or "商品" in value):
            replace_paragraph_text(p, "{{ targeted_drug_brand_summary }}")
        elif "本次共检出" in value or "本次共检测出" in value:
            replace_paragraph_text(
                p,
                "本次检出体细胞变异：{{ total_variants_count }} 个；"
                "靶向药物相关变异：{{ drug_related_count }} 个。",
            )
        elif "化疗药物小结：" in value:
            replace_paragraph_text(p, "化疗药物小结：{{ chemotherapy_summary_text }}")
    if panel == "lung_62_pdl1":
        replace_between(
            doc,
            "2. 用药提示",
            "3. 延伸阅读-免疫治疗",
            ("{{ pdl1_table_interpretation }}",),
        )
        start = heading(doc, "第二部分：PD-L1表达检测结果及解析")._p
        end = heading(doc, "第三部分：肺癌62基因检测结果及解析")._p
        nodes = list(doc.element.body)
        for node in nodes[nodes.index(start) + 1 : nodes.index(end)]:
            # Only the source patient image is in this clinical result block.
            for media in list(node.iter(qn("w:drawing"))) + list(
                node.iter(qn("w:pict"))
            ):
                media.getparent().remove(media)
        paragraph_after(tables["pdl1"]._tbl, "__PDL1_CASE_IMAGE__")
        for p in doc.paragraphs:
            if "定性结果判定标准" in p.text or (
                "TPS" in p.text and "CPS" in p.text and "判定" in p.text
            ):
                replace_paragraph_text(p, "{{ pdl1_classification_notice }}")
            if (
                "PD-L1" in p.text
                and ("阳性" in p.text or "阴性" in p.text)
                and len(p.text) < 180
            ):
                replace_paragraph_text(
                    p,
                    "PD-L1：{{ pdl1_result_display }}；TPS：{{ pdl1_tps_display }}；"
                    "CPS：{{ pdl1_cps_display }}",
                )


def install_shared_modules(doc, spec, tables):
    """Copy only variable table geometry from the existing lung rules template."""
    reference = Document(BASE / "templates/lung_588_pdl1_historical_golden_v1.docx")
    modules = {}
    for table in reference.tables:
        matched = re.search(r"for row in (\w+)", text(table._tbl))
        if matched:
            modules[matched.group(1)] = table
    anchor = tables["variants"]._tbl
    if "targeted" not in tables:
        anchor = paragraph_after(anchor, "靶向药物提示（报告组评审）", title=True)
        node = copy.deepcopy(modules["targeted_drug_tips"]._tbl)
        anchor.addnext(node)
    # Place biomarker/PGx modules before the fixed appendix, not in case prose.
    boundary = heading(doc, spec["rich_end"])._p
    anchor = boundary.getprevious()
    anchor = paragraph_after(
        anchor, "补充检测结果：TMB、MSI 与化疗药物基因组学", title=True
    )
    for value in (
        IMMUNE_NOTICE,
        "{{ tmb_detail_sentence }}",
        "{{ tmb_detail_interpretation }}",
        "{{ msi_detail_sentence }}",
        "{{ msi_detail_interpretation }}",
    ):
        anchor = paragraph_after(anchor, value)
    anchor = paragraph_after(anchor, "{%p if cnv_review_required %}")
    anchor = paragraph_after(anchor, "{{ immune_hyperprogression_result }}")
    anchor = paragraph_after(anchor, "{%p endif %}")
    for collection, title in (
        ("immune_positive_results", "免疫正相关基因（限本产品范围）"),
        ("immune_negative_results", "免疫负相关基因（限本产品范围）"),
        ("immune_hyperprogression_results", "免疫超进展相关基因（限本产品范围）"),
        ("chemotherapy_predictions", "化疗药物来源结果"),
        ("chemotherapy_regimen_predictions", "化疗方案来源结果"),
        ("chemotherapy_dosage_rows", "历史化疗方案剂量展示（待报告组审核）"),
        ("irinotecan_safety_rows", "伊立替康历史剂量展示（待报告组裁决）"),
    ):
        anchor = paragraph_after(anchor, title, title=True)
        if collection == "chemotherapy_predictions":
            anchor = paragraph_after(anchor, CHEMO_NOTICE)
            anchor = paragraph_after(
                anchor, "化疗药物小结：{{ chemotherapy_summary_text }}"
            )
        node = copy.deepcopy(modules[collection]._tbl)
        anchor.addnext(node)
        anchor = node
    anchor = paragraph_after(anchor, "药物基因组学明细（CtDrug 来源核对）", title=True)
    anchor = paragraph_after(anchor, CHEMO_NOTICE)
    for _, collection in DRUG_DETAIL_BINDINGS:
        # Preserve the reviewed 588 table geometry and source-derived rows.
        # Empty drug lists must not create a run of header-only tables.
        anchor = paragraph_after(anchor, "{%p if " + collection + " %}")
        drug_label = modules[collection].cell(0, 0).text.strip()
        if not drug_label or "{" in drug_label:
            raise ValueError("PGx caption must come from the maintained source table header")
        anchor = paragraph_after(anchor, drug_label + "药物基因组学明细", title=True)
        node = copy.deepcopy(modules[collection]._tbl)
        anchor.addnext(node)
        anchor = paragraph_after(node, "{%p endif %}")


def compact_gene_qc_appendix(doc, tables):
    """Remove historical redundant breaks around short gene/QC blocks."""
    gene_table, qc_table = tables["genes"], tables["qc"]
    title = gene_table._tbl.getprevious()
    if title is not None and title.tag == qn("w:p"):
        Paragraph(title, None).paragraph_format.page_break_before = False
        Paragraph(title, None).paragraph_format.keep_with_next = True
        cursor = title.getprevious()
        for _ in range(8):
            if cursor is None or text(cursor).strip():
                break
            previous = cursor.getprevious()
            if cursor.tag == qn("w:p"):
                if any(list(cursor.iter(qn(tag))) for tag in ("w:sectPr", "w:drawing", "w:pict")):
                    break
                cursor.getparent().remove(cursor)
            cursor = previous
    cursor = gene_table._tbl.getnext()
    while cursor is not None and cursor is not qc_table._tbl:
        following = cursor.getnext()
        if cursor.tag == qn("w:p") and not list(cursor.iter(qn("w:sectPr"))):
            if not text(cursor).strip() and not any(
                list(cursor.iter(qn(tag))) for tag in ("w:drawing", "w:pict")
            ):
                cursor.getparent().remove(cursor)
            else:
                Paragraph(cursor, None).paragraph_format.page_break_before = False
                Paragraph(cursor, None).paragraph_format.keep_with_next = True
        cursor = following
    for table in (gene_table, qc_table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.page_break_before = False


def install_refreshable_toc(doc):
    """Replace flat historical TOC caches with a real index, keeping body styles."""
    if any("TOC" in (node.text or "") for node in doc.element.iter(qn("w:instrText"))):
        return
    paragraphs = list(doc.paragraphs)
    entries = [p for p in paragraphs if re.fullmatch(r"toc\s+[1-9]", p.style.name, re.I)]

    def label(value):
        value = str(value).strip().split("\t")[0].strip()
        return compact(re.sub(r"^\d+(?:\.\d+)*[.．、]?\s*", "", value))

    if entries:
        levels = {
            label(p.text): int(re.search(r"[1-9]", p.style.name).group()) - 1
            for p in entries
        }
        toc_heading = None
        last = paragraphs.index(entries[-1])
    else:
        # B-family TOCs can be floating text boxes, not body paragraphs. Their
        # historical page numbers must not remain beside the new native index.
        toc_heading = heading(doc, "目录")
        last = next(i for i, p in enumerate(paragraphs) if p._p is toc_heading._p)
        cache_labels = set()
        toc_media = list(toc_heading._p.iter(qn("w:drawing"))) + list(
            toc_heading._p.iter(qn("w:pict"))
        )
        for media in toc_media:
            labels = []
            for node in media.iter(qn("w:p")):
                value = "".join(
                    "\t" if n.tag == qn("w:tab") else (n.text or "")
                    for n in node.iter() if n.tag in {qn("w:t"), qn("w:tab")}
                )
                labels.append(label(value))
            major = [value for value in labels if re.match(r"第[一二三四五六七八九十]+部分.+", value)]
            if len(major) >= 3:
                # Replace recognizable complete caches and their VML fallback.
                media.getparent().remove(media)
                cache_labels.update(labels)
        if cache_labels:
            # The remaining ornaments belong to the old floating TOC geometry
            # (its vertical rule otherwise crosses the new native text). Only
            # this TOC title's media are removed; other document media stay.
            for media in toc_media:
                parent = media.getparent()
                if parent is not None:
                    parent.remove(media)
        levels = {
            label(p.text): 0 for p in paragraphs[last + 1:]
            if re.fullmatch(r"第[一二三四五六七八九十]+部分\s*[：:].+", p.text.strip())
        }
        # Use surviving body headings from the historical directory, with no
        # inferred labels or inherited case-cell/figure captions.
        for value in cache_labels:
            if value:
                levels.setdefault(value, 1)
    # Both B/C sources contain non-heading paragraphs with direct or inherited
    # outline levels (patient cells, figures, citations and blank paragraphs).
    # Keep their appearance but index only matched historical body headings.
    for node in doc.element.body.iter(qn("w:p")):
        properties = node.get_or_add_pPr()
        outline = properties.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            properties.append(outline)
        outline.set(qn("w:val"), "9")
    # Appendix pathway subsections also have local "参考文献：" captions.
    # Only the final global bibliography belongs to the report directory.
    reference_headings = [p for p in paragraphs[last + 1:] if label(p.text) == label("参考文献")]
    matched = 0
    for p in paragraphs[last + 1:]:
        key = label(p.text)
        if key not in levels:
            continue
        if entries and key == label("参考文献") and p.text.strip().endswith(("：", ":")):
            continue
        if not entries and key == label("参考文献") and p._p is not reference_headings[-1]._p:
            continue
        properties = p._p.get_or_add_pPr()
        node = properties.find(qn("w:outlineLvl"))
        if node is None:
            node = OxmlElement("w:outlineLvl")
            properties.append(node)
        node.set(qn("w:val"), str(levels[key]))
        matched += 1
        if not entries:
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
            cursor = p._p.getprevious()
            page_break = False
            while cursor is not None and cursor.tag == qn("w:p") and not text(cursor).strip():
                if any(list(cursor.iter(qn(tag))) for tag in (
                    "w:sectPr", "w:drawing", "w:pict", "w:fldChar", "w:instrText"
                )):
                    break
                previous = cursor.getprevious()
                page_break |= any(n.get(qn("w:type")) == "page" for n in cursor.iter(qn("w:br")))
                cursor.getparent().remove(cursor)
                cursor = previous
            if page_break:
                p.paragraph_format.page_break_before = True
    if matched < 3:
        raise ValueError("Historical TOC cannot be matched to surviving body headings")
    # Match the working A-family complex-field representation. The empty
    # hand-built content control did not survive the first Linux round-trip.
    # This non-numeric cache is explicitly pending, never a made-up page.
    block = OxmlElement("w:p")
    paragraph = block
    for field_kind in ("begin", "instruction", "separate", "end"):
        run = OxmlElement("w:r")
        paragraph.append(run)
        node = OxmlElement("w:instrText" if field_kind == "instruction" else "w:fldChar")
        if field_kind == "instruction":
            node.text = ' TOC \\o "1-3" \\h \\z \\u '
            node.set(qn("xml:space"), "preserve")
        else:
            node.set(qn("w:fldCharType"), field_kind)
            if field_kind == "begin":
                node.set(qn("w:dirty"), "true")
        run.append(node)
        if field_kind == "separate":
            cache_run = OxmlElement("w:r")
            cache = OxmlElement("w:t")
            cache.text = "目录待排版引擎刷新"
            cache_run.append(cache)
            paragraph.append(cache_run)
    if entries:
        entries[0]._p.addprevious(block)
    else:
        toc_heading._p.addnext(block)
        # The source closes its front-matter section on the TOC title. Move
        # that boundary after the refreshed entries, so the TOC stays with its
        # title and cannot share a body page with patient information.
        sect = toc_heading._p.find(".//" + qn("w:sectPr"))
        if sect is not None:
            boundary = OxmlElement("w:p")
            props = OxmlElement("w:pPr")
            sect.getparent().remove(sect)
            props.append(sect)
            boundary.append(props)
            block.addnext(boundary)
        else:
            first_body_heading = next(
                p for p in paragraphs[last + 1:] if label(p.text) in levels
            )
            first_body_heading.paragraph_format.page_break_before = True
    for p in entries:
        # An old TOC can carry the front-matter section boundary. Preserve it.
        sect = p._p.find(".//" + qn("w:sectPr"))
        if sect is not None:
            boundary = OxmlElement("w:p")
            props = OxmlElement("w:pPr")
            props.append(copy.deepcopy(sect))
            boundary.append(props)
            p._p.addnext(boundary)
        p._p.getparent().remove(p._p)
    update = doc.settings.element.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        doc.settings.element.append(update)
    update.set(qn("w:val"), "true")


def normalize_draft_case_fields(doc, basic_table):
    """Remove stale case summaries and light-background/white-value inheritance."""
    for row in basic_table.rows:
        for cell in row.cells:
            if "{{" not in cell.text:
                continue
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
    for p in doc.paragraphs:
        value = compact(p.text)
        if "检出体细胞变异" in value:
            replace_paragraph_text(
                p,
                "本次检出体细胞变异：{{ total_variants_count }}个；"
                "与靶向药物用药相关的变异有：{{ drug_related_count }}个。详见本例结果表。",
            )
        if value != "肺癌相关重要基因变异及药物提示":
            continue
        p.paragraph_format.keep_with_next = True
        cursor = p._p.getprevious()
        while cursor is not None and cursor.tag == qn("w:p") and not text(cursor).strip():
            if any(list(cursor.iter(qn(tag))) for tag in ("w:sectPr", "w:drawing", "w:pict")):
                break
            previous = cursor.getprevious()
            cursor.getparent().remove(cursor)
            cursor = previous


def normalize_draft_variant_flow(table):
    """Keep one event per row and cancel inherited body-text indentation."""
    for index, row in enumerate(table.rows):
        props = row._tr.get_or_add_trPr()
        for tag in ("w:cantSplit", "w:tblHeader") if index < 2 else ("w:cantSplit",):
            flag = props.find(qn(tag))
            if flag is None:
                flag = OxmlElement(tag)
                props.append(flag)
            flag.set(qn("w:val"), "true")
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(0)
                indent = paragraph._p.get_or_add_pPr().get_or_add_ind()
                indent.set(qn("w:firstLineChars"), "0")
                indent.attrib.pop(qn("w:hangingChars"), None)
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = index < 2


def normalize_optional_drug_block(doc, table):
    """Preserve the source caption, expose empty results, and remove spacers."""
    cursor = table._tbl.getprevious()
    block = []
    while cursor is not None and cursor.tag == qn("w:p"):
        block.append(cursor)
        if "潜在获益上市药物提示" in text(cursor):
            break
        cursor = cursor.getprevious()
    else:
        raise ValueError("Optional drug table has no source section heading")
    title = block[-1]
    for node in block:
        paragraph = Paragraph(node, None)
        paragraph.paragraph_format.page_break_before = False
        paragraph.paragraph_format.keep_with_next = True
    cursor = title.getprevious()
    while cursor is not None and cursor.tag == qn("w:p") and not text(cursor).strip():
        if any(list(cursor.iter(qn(tag))) for tag in ("w:sectPr", "w:drawing", "w:pict")):
            break
        previous = cursor.getprevious()
        cursor.getparent().remove(cursor)
        cursor = previous
    paragraph_after(table._tbl.getprevious(), "{%p if chemotherapy %}")
    anchor = paragraph_after(table._tbl, "{%p else %}")
    anchor = paragraph_after(anchor, EMPTY_APPROVED_DRUG_NOTICE)
    anchor = paragraph_after(anchor, "{%p endif %}")
    cursor = anchor.getnext()
    while cursor is not None:
        following = cursor.getnext()
        if cursor.tag in {qn("w:bookmarkStart"), qn("w:bookmarkEnd")}:
            cursor = following
            continue
        if cursor.tag != qn("w:p") or any(
            list(cursor.iter(qn(tag))) for tag in ("w:sectPr", "w:drawing", "w:pict")
        ):
            break
        if text(cursor).strip():
            # A short result-explanation block can share the same page. Keep
            # the source page boundary before a new major report part.
            if re.fullmatch(r"\d+[.．、]?检测结果说明", compact(text(cursor))):
                Paragraph(cursor, None).paragraph_format.page_break_before = False
            break
        cursor.getparent().remove(cursor)
        cursor = following


def normalize_b_family_faq_flow(doc):
    """Keep bounded question/answer blocks together and let the appendix flow."""
    for table in doc.tables:
        if not re.fullmatch(r"问题\s*\d+", table.cell(0, 0).text.strip()):
            continue
        if len(table.rows) > 4 or len(text(table._tbl)) > 800:
            continue  # Do not bind an unbounded multi-page FAQ into one block.
        paragraphs = list(table._tbl.iter(qn("w:p")))
        for index, node in enumerate(paragraphs):
            paragraph = Paragraph(node, None)
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.keep_with_next = index < len(paragraphs) - 1
        for row in table.rows:
            props = row._tr.get_or_add_trPr()
            if props.find(qn("w:cantSplit")) is None:
                props.append(OxmlElement("w:cantSplit"))
    for paragraph in doc.paragraphs:
        label = re.sub(r"^\d+[.．、]?", "", compact(paragraph.text))
        if label == "常见问题解答":
            # The inherited list shares counters with the preceding glossary.
            # This is the first fixed appendix subsection, not list item 3.
            replace_paragraph_text(paragraph, "1. 常见问题解答")
            paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = 0
        elif label == "肺癌诊疗知识":
            paragraph.paragraph_format.page_break_before = False
            paragraph.paragraph_format.keep_with_next = True


def normalize_fixed_reference_flow(doc, appendix):
    """Keep bounded fixed figure/reference blocks together without rewriting text."""
    nodes = list(doc.element.body)
    for index in range(nodes.index(appendix) + 1, len(nodes)):
        node = nodes[index]
        if node.tag != qn("w:p") or compact(text(node)) != "参考文献":
            continue
        references = []
        for following in nodes[index + 1:]:
            value = text(following).strip()
            if following.tag != qn("w:p") or not re.search(r"\b(?:19|20)\d{2}\b", value):
                break
            references.append(following)
        if not references or len(references) > 8 or sum(len(text(p)) for p in references) > 1200:
            continue
        block = [node, *references]
        caption = nodes[index - 1]
        figure = nodes[index - 2] if index > 1 else None
        if (
            caption.tag == qn("w:p") and 0 < len(text(caption).strip()) <= 120
            and figure is not None and figure.tag == qn("w:p")
            and list(figure.iter(qn("w:drawing")))
        ):
            heights = [int(n.get("cy", "0")) for n in figure.iter(qn("wp:extent"))]
            if heights and 0 < max(heights) <= 5 * 914400:
                block = [figure, caption, *block]
        for position, item in enumerate(block):
            paragraph = Paragraph(item, None)
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.keep_with_next = position < len(block) - 1
        for item in references:
            props = item.get_or_add_pPr()
            grid = props.find(qn("w:snapToGrid"))
            if grid is None:
                grid = OxmlElement("w:snapToGrid")
                props.append(grid)
            grid.set(qn("w:val"), "false")


def normalize_footer_page_totals(doc):
    """Replace historical literal totals with native NUMPAGES fields in place."""
    for part in _story_parts(doc):
        if "/footer" not in str(part.partname):
            continue
        for paragraph in part.element.iter(qn("w:p")):
            texts = paragraph.xpath("./w:r/w:t")
            if any("NUMPAGES" in (item.text or "") for item in paragraph.iter(qn("w:instrText"))):
                continue
            for index, item in enumerate(texts):
                if not (
                    (item.text or "").strip().isdigit() and index > 0 and index + 1 < len(texts)
                    and compact(texts[index - 1].text).endswith("共")
                    and compact(texts[index + 1].text).startswith("页")
                ):
                    continue
                run = item.getparent()
                if len([child for child in run if child.tag != qn("w:rPr")]) != 1:
                    raise ValueError("Unsupported mixed-content footer total run")
                for kind in ("begin", "instruction", "separate", "cache", "end"):
                    new = OxmlElement("w:r")
                    if run.rPr is not None:
                        new.append(copy.deepcopy(run.rPr))
                    if kind == "instruction":
                        element = OxmlElement("w:instrText")
                        element.set(qn("xml:space"), "preserve")
                        element.text = " NUMPAGES "
                    elif kind == "cache":
                        element = OxmlElement("w:t")
                        element.text = "0"  # Native refresh is required before output QA.
                    else:
                        element = OxmlElement("w:fldChar")
                        element.set(qn("w:fldCharType"), kind)
                        if kind == "begin":
                            element.set(qn("w:dirty"), "true")
                    new.append(element)
                    run.addprevious(new)
                run.getparent().remove(run)


def sanitize_final(doc, replacements, panel):
    for table in doc.tables:
        joined = compact(text(table._tbl))
        if "检验者" in joined and "签发者" in joined:
            for row in table.rows:
                for cell in row.cells:
                    for media in list(cell._tc.iter(qn("w:drawing"))) + list(
                        cell._tc.iter(qn("w:pict"))
                    ):
                        media.getparent().remove(media)
                    replace_cell_text(cell, "")
            replace_cell_text(
                table.rows[0].cells[0],
                "检验者：{{ issuer }}    核对者：{{ reviewer }}    签发者：{{ signer }}",
            )
            replace_cell_text(table.rows[-1].cells[-1], "报告日期：{{ report_date }}")
    for p in list(doc.paragraphs):
        if any(
            label in compact(p.text)
            for label in ("检测者", "审核者", "检测人", "审核人")
        ):
            previous = p._p.getprevious()
            if previous is not None and (
                list(previous.iter(qn("w:drawing")))
                or list(previous.iter(qn("w:pict")))
            ):
                previous.getparent().remove(previous)
            # Remove drawings inside the caption too; clearing runs alone can
            # leave VML signature objects and external image relationships.
            for child in list(p._p):
                if child.tag != qn("w:pPr"):
                    p._p.remove(child)
            p.add_run("检测者：        审核者：        报告日期：{{ report_date }}")
        if "对委托人" in p.text and "DNA" in p.text:
            replace_paragraph_text(
                p,
                "对本病例提供的 {{ sample_type }} 样本进行本产品范围内的基因检测，"
                "检测基因列表见附录。",
            )
    for p in doc.paragraphs[:50]:
        value = compact(p.text)
        for label, name in {
            "姓名": "patient_name",
            "报告编号": "report_number",
            "送检日期": "receive_date",
            "报告日期": "report_date",
        }.items():
            if value.startswith(label):
                replace_paragraph_text(p, label + "：" + field(name))
    _scrub_story_content(
        doc,
        replacements=list(replacements.items()),
        sensitive_tokens=list(replacements),
    )
    for part in _story_parts(doc):
        for pnode in part.element.iter(qn("w:p")):
            value = text(pnode).strip()
            date_pattern = r"20\d{2}[-/.年]\d{1,2}[-/.月]\d{1,2}日?"
            if re.fullmatch(date_pattern, value) or (
                "报告" in value and ("日期" in value or "时间" in value)
            ):
                replaced = re.sub(date_pattern, "{{ report_date }}", value)
                if replaced != value:
                    replace_paragraph_text(Paragraph(pnode, None), replaced)
        for rel_id, rel in list(part.rels.items()):
            if rel.is_external and any(
                token in str(rel.target_ref) for token in replacements
            ):
                part.drop_rel(rel_id)
        for node in part.element.iter():
            for attr in ("descr", "title", "name"):
                if attr in node.attrib and any(
                    token in node.attrib[attr] for token in replacements
                ):
                    node.attrib[attr] = "ReportGen"
    _prune_unreferenced_images(doc)
    props = doc.core_properties
    for key in (
        "author",
        "last_modified_by",
        "comments",
        "keywords",
        "category",
        "identifier",
        "language",
        "content_status",
        "version",
    ):
        setattr(props, key, "")
    props.title = panel + " draft template"
    props.subject = "报告组评审草稿；非临床交付"
    props.created = props.modified = datetime.datetime(2000, 1, 1)
    # Exact historical heights prevent long Jinja values from clipping.
    for table in doc.tables:
        for row in table.rows:
            for height in row._tr.xpath("./w:trPr/w:trHeight"):
                height.set(qn("w:hRule"), "atLeast")


def build_template(panel, spec, source, work, output):
    if sha(source) != spec["source_sha256"]:
        raise ValueError("Historical source SHA mismatch")
    work.mkdir(parents=True, exist_ok=True)
    repaired = work / "repaired.docx"
    repair = repair_docx(source, repaired)
    original = Document(repaired)
    genes = gene_list(original, spec)
    cells, replacements = scalar_map(original, spec)
    tokens = tuple(replacements)
    (work / "private_tokens.json").write_text(
        json.dumps(
            {"replacements": replacements, "tokens": tokens},
            ensure_ascii=False,
            indent=2,
        )
    )
    seed = work / "seed.docx"
    seed_result = build_seed(
        repaired,
        seed,
        replacements=replacements,
        protected_tokens=tokens,
        allow_commit_output=False,
        allow_residual=False,
        project_root=ROOT,
    )
    # Macro source regions are fully replaced before literal discovery; no seed
    # event narrative is copied into a scalar knowledge field.
    indices = spec["tables"]
    loops = [
        loop_spec(
            indices["variants"],
            "variants_2_1",
            [
                "gene",
                "transcript",
                "chr",
                "exon",
                "locus",
                "var_type_cn",
                "af_pct",
                "benefit_drugs",
                "caution_drugs",
            ],
            2,
        )
    ]
    if "other_drugs" in indices:
        loops.append(
            loop_spec(indices["other_drugs"], "chemotherapy", ["drug", "gene"])
        )
    loops.append(
        loop_spec(
            indices["introductions"],
            "targeted_drug_introductions",
            ["drug_name", "gene", "introduction"],
        )
    )
    guide_cols = len(original.tables[indices["guideline"]].columns)
    guide_fields = (
        ["gene", "drugs", "clinical_note", "result"]
        if guide_cols == 4
        else ["gene", {"expr": "row.drugs ~ '\\n' ~ row.clinical_note"}, "result"]
    )
    loops.append(
        loop_spec(indices["guideline"], "lung_guideline_drug_results", guide_fields)
    )
    if "targeted" in indices:
        loops.append(
            loop_spec(
                indices["targeted"],
                "targeted_drug_tips",
                ["gene_display", "variant_site", "benefit_drugs", "caution_drugs"],
            )
        )
    variable = work / "variableized.docx"
    variableize_docx(seed, variable, {"cell_variables": cells, "table_loops": loops})
    doc = Document(variable)
    tables = {key: doc.tables[index] for key, index in indices.items()}
    if panel == "lung_588":
        _replace_basic_information(tables["basic"])
    if spec["family"] in {"A", "C"} and len(tables["basic"].columns) == 4:
        # The diagnosis value spans the final three logical cells.
        replace_cell_text(tables["basic"].rows[2].cells[1], field("clinical_diagnosis"))
    if panel == "lung_62":
        for col, name in enumerate(
            ("sample_id", "sample_type", "sample_site", "collection_date", "receive_date")
        ):
            replace_cell_text(doc.tables[2].rows[1].cells[col], field(name))
    if "pdl1" in tables:
        _replace_pdl1_table(tables["pdl1"])
    headers = (
        ["检测基因", "本癌种相关治疗药物", "临床提示", "检测结果"]
        if guide_cols == 4
        else ["检测基因", "本癌种相关治疗药物及临床提示", "检测结果"]
    )
    for cell, value in zip(tables["guideline"].rows[0].cells, headers):
        replace_cell_text(cell, value)
    clear_case_regions(doc, panel, spec, tables)
    if panel == "lung_588":
        # Replace the complete variable immune/chemotherapy block; the fixed
        # appendix, cover, result tables and package identity remain B-family.
        replace_between(
            doc,
            "3. 免疫疗效正相关/负相关/超进展基因检测结果",
            "5. 检测结果说明",
            (IMMUNE_NOTICE, CHEMO_NOTICE),
        )
        for key in ("biomarkers", "chemo_summary"):
            table = tables[key]
            for row in table.rows[1:]:
                for cell in row.cells:
                    replace_cell_text(
                        cell,
                        "{{ chemotherapy_summary_text }}"
                        if key == "chemo_summary"
                        else "{{ tmb_summary }}；{{ msi_summary }}",
                    )
    install_shared_modules(doc, spec, tables)
    ensure_table_separator(tables["guideline"], "肺癌相关重要基因变异及药物提示")
    if not any("total_variants_count" in text(p._p) for p in doc.paragraphs):
        paragraph_after(
            tables["variants"]._tbl.getprevious(),
            "本次检出体细胞变异：{{ total_variants_count }} 个（含Ⅰ/Ⅱ/Ⅲ类）；"
            "靶向药物相关变异：{{ drug_related_count }} 个。",
        )
    # A fixed, explicit gene list is not patient data. No unassayed gene is
    # represented as a negative result in the dynamic undetected set.
    replace_cell_text(
        tables["genes"].rows[0].cells[0], f"肺癌{spec['gene_count']}基因检测列表"
    )
    # These four historical acceptance statuses are not sequencing metrics.
    # Never substitute Q30/depth for a laboratory QC sign-off or invent PASS.
    qc_fields = (
        "qc_extraction_status",
        "qc_library_status",
        "qc_sequencing_status",
        "qc_analysis_status",
    )
    for row, name in zip(tables["qc"].rows[1:], qc_fields):
        replace_cell_text(row.cells[-1], field(name))
    compact_gene_qc_appendix(doc, tables)
    install_refreshable_toc(doc)
    appendix = heading(doc, spec["rich_end"])
    paragraph_after(appendix._p, APPENDIX_NOTICE)
    front = next(p for p in doc.paragraphs if p.text.strip() and "\t" not in p.text)
    paragraph_after(front._p, NOTICE)
    sanitize_final(doc, replacements, panel)
    normalize_draft_case_fields(doc, tables["basic"])
    normalize_draft_variant_flow(tables["variants"])
    if "other_drugs" in tables:
        normalize_optional_drug_block(doc, tables["other_drugs"])
    if panel == "lung_588":
        normalize_b_family_faq_flow(doc)
    normalize_fixed_reference_flow(doc, appendix._p)
    normalize_footer_page_totals(doc)
    fixed = {}
    # Only the fixed appendix may contain a historical literature example.
    after_appendix = False
    for node in doc.element.body:
        if node is appendix._p:
            after_appendix = True
        for pnode in node.iter(qn("w:p")):
            value = text(pnode)
            if not (RE_CHGVS.search(value) or RE_PHGVS.search(value)):
                continue
            if not after_appendix:
                raise ValueError("Unconverted event literal outside the fixed appendix")
            key = f"historical_appendix_{len(fixed) + 1:04}"
            fixed[key] = {
                "purpose": "历史固定附录（非病例解释，待报告组复核）",
                "source": "sha256:" + spec["source_sha256"],
                "text": value,
            }
            replace_paragraph_text(Paragraph(pnode, None), "{{ " + key + " }}")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    _normalize_zip_metadata(output)
    if any(count_tokens_in_zip(output, tokens).values()):
        raise ValueError("Private source token leaked into the template")
    scan = scan_docx(output, tokens=list(tokens))
    if scan.hard:
        raise ValueError("Hardcoded literal gate failed")
    return (
        genes,
        fixed,
        {
            "source_sha256": spec["source_sha256"],
            "template_sha256": sha(output),
            "seed_success": seed_result["success"],
            "repair": repair,
            "hard_literals": len(scan.hard),
            "soft_literals": len(scan.soft),
            "tables": len(doc.tables),
            "paragraphs": len(doc.paragraphs),
        },
    )


def write_yaml(path, payload):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        yaml.safe_dump(payload, allow_unicode=True, sort_keys=False, width=110),
        encoding="utf-8",
    )


def build_package(panel, spec, private_dir, work, packages_dir):
    package_dir = packages_dir / panel
    template_id = panel + "_historical_draft_v1"
    output = package_dir / "templates" / (template_id + ".docx")
    genes, fixed, receipt = build_template(
        panel, spec, private_dir / (panel + ".mother.docx"), work / panel, output
    )
    gene_set = set(genes)
    raw = yaml.safe_load((BASE / "panel.yaml").read_text())
    for key in (
        "release_governance",
        "review_candidate_contract",
        "context_contracts",
        "golden_cases",
    ):
        raw.pop(key, None)
    raw.update(
        panel_id=panel,
        display_name=f"肺癌{spec['gene_count']}基因"
        + ("+PD-L1" if panel.endswith("_pdl1") else ""),
        version="0.1.0",
        status="draft",
        description=NOTICE,
        aliases=[panel],
        default_template=template_id,
    )
    raw["naming"] = {
        "output_pattern": "{patient_name}_{sample_id}_{project_name}_{report_date}_评审草稿.docx",
    }
    if panel == "lung_588":
        raw["aliases"].append("lung588")
    processors = [
        "part3_formatted_sections",
        "rebuild_references",
        "signature_placeholder",
        "immune_table_notes",
        "bullet_lists",
        "blank_page_cleanup",
        "toc_refresh",
        "final_refresh_cleanup",
        "underlines_and_styles",
    ]
    if panel.endswith("_pdl1"):
        processors.insert(3, "pdl1_case_image")
    raw["templates"] = [
        {
            "id": template_id,
            "file": f"templates/{template_id}.docx",
            "version": "0.1.0",
            "status": "draft",
            "description": NOTICE,
            "processors": processors,
        }
    ]
    raw["processors"] = processors
    raw["part3_knowledge"]["cross_cancer_residual_scan"].update(
        start_heading=spec["rich_start"], end_heading=spec["rich_end"]
    )
    raw["input_contract"]["optional_source_fields"] = {
        "phone": ["联系方式", "联系电话"],
        "family_history": ["家族史"],
        "treatment_history": ["既往用药史", "治疗史"],
        "signer": ["签发者", "签发人"],
        "qc_extraction_status": ["核酸提取质控结论"],
        "qc_library_status": ["文库构建质控结论"],
        "qc_sequencing_status": ["测序质控结论"],
        "qc_analysis_status": ["分析质控结论"],
    }
    raw["golden_cases"] = [
        {
            "id": panel + "_derived_draft_contract",
            "runner": "reportgen.core.golden_case:run_golden_case",
            "command": "pytest backend/tests/test_lung_small_panel_contract.py -q",
            "synthetic": True,
            "expected_qa_status": "PASS",
        }
    ]
    raw["template_contract"] = {
        "required_variables": [
            "patient_name",
            "total_variants_count",
            "tmb_detail_sentence",
            "msi_detail_sentence",
        ],
        "required_lists": [
            "variants_2_1",
            "targeted_drug_tips",
            "lung_guideline_drug_results",
            "chemotherapy_predictions",
            "chemotherapy_regimen_predictions",
        ],
        "required_markers": ["__PART3_MARKER__"],
    }
    if panel.endswith("_pdl1"):
        raw["template_contract"]["required_variables"] += [
            "pdl1_tps_display",
            "pdl1_cps_display",
            "pdl1_result_display",
            "pdl1_classification_notice",
        ]
        raw["template_contract"]["required_markers"].append("__PDL1_CASE_IMAGE__")
    else:
        raw["rules"].pop("pdl1_product_contract", None)
        raw["input_contract"]["biomarkers"] = {
            k: v
            for k, v in raw["input_contract"]["biomarkers"].items()
            if not k.startswith("pdl1_")
        }
    flag = spec.get("membership_column") or "ExistInsmall588"
    raw["input_contract"]["missing_source_defaults"] = {"sample_type": "未提供"}
    if spec.get("membership_column"):
        raw["derived_input"] = {
            "membership_column": flag,
            "genes": genes,
            "provenance": "superset_derived_engineering_input_not_a_historical_pair",
        }
        raw["input_contract"]["required_columns"]["Variations"].append(flag)
        raw["input_contract"]["required_columns"]["Hereditary_tumor"] = [
            flag,
            "Gene_Symbol",
        ]
    number = spec["gene_count"]
    pdl1 = panel.endswith("_pdl1")
    pattern = (
        rf"(?:肺癌\s*{number}(?:\s*基因)?|{number}\s*基因)(?:.*PD[\s-]*L1)"
        if pdl1
        else rf"(?:肺癌\s*{number}(?:\s*基因)?|{number}\s*基因)(?!.*PD[\s-]*L1)"
    )
    rules = {
        "keywords": [raw["display_name"]],
        "keyword_groups": [
            {"any": [{"type": "regex", "pattern": pattern}], "weight": 1}
        ],
        "priority": 30 if not pdl1 else 20,
        "structural_fingerprints": [
            {
                "id": f"lung{number}_ngs_family_v1",
                "confidence": 1.0,
                "required_tables": {
                    "Variations": {"required_columns": ["ExistIn552", flag]},
                    "Hereditary_tumor": {"required_columns": [flag, "ExistIn178"]},
                },
            }
        ],
    }
    if number in (62, 588):
        rules["identity_family"] = {
            "id": f"lung_{number}",
            "default_project_type": f"lung_{number}",
            "pdl1": pdl1,
        }
    raw["project_detector_rules"] = rules
    raw["golden_template_source"] = {
        "sha256": spec["source_sha256"],
        "family": spec["family"],
        "builder": "scripts/build_lung_draft_packages.py",
    }
    for source in sorted((BASE / "rules").glob("*.yaml")):
        if source.name.startswith("._"):
            continue
        rule = yaml.safe_load(source.read_text())
        rule["panel_id"] = panel
        if source.name == "panel.yaml":
            rule["important_genes"] = genes
            rule["panel_display_genes"] = [
                {"name": gene, "transcript": "", "chromosome": ""} for gene in genes
            ]
        elif source.name == "guideline_tables.yaml":
            for key in ("nccn_results", "lung_guideline_drug_results"):
                table = rule["guideline_tables"][key]
                table["rows"] = [
                    row
                    for row in table["rows"]
                    if set(row.get("genes") or []) <= gene_set
                ]
        elif source.name == "biomarkers.yaml":
            for table in rule["biomarkers"]["immune_gene_tables"].values():
                if not isinstance(table, dict):
                    continue
                table["genes"] = [g for g in table.get("genes", []) if g in gene_set]
                scoped = []
                for row in table.get("rows", []):
                    group = set(row.get("genes") or [])
                    if row.get("mode") == "co_mutation" and not group <= gene_set:
                        continue
                    if not group & gene_set:
                        continue
                    row["genes"] = [g for g in row["genes"] if g in gene_set]
                    if "selectors" in row:
                        row["selectors"] = [
                            s
                            for s in row["selectors"]
                            if set(s.get("genes") or []) <= gene_set
                        ]
                    scoped.append(row)
                table["rows"] = scoped
        elif source.name == "drugs.yaml":
            policy = rule["targeted_drug_rules"]
            pending = policy["gene_level_review_pending"]
            pending["allowed_genes"] = [
                g
                for g in pending["allowed_genes"]
                if g in gene_set
                and g not in spec.get("excluded_gene_level_drug_tips", [])
            ]
            pending["fallback_overrides"] = {
                g: v
                for g, v in pending["fallback_overrides"].items()
                if g in pending["allowed_genes"]
            }
            selected = []
            for row in policy["reviewed_variant_overrides"]:
                if str(row.get("gene") or "") not in gene_set:
                    continue
                row["panels"] = [panel]
                selected.append(row)
            policy["reviewed_variant_overrides"] = selected
        elif source.name == "report_text.yaml":
            rule.setdefault("texts", {}).update(fixed)
            rule["texts"]["patient_letter_project_label"]["text"] = (
                raw["display_name"] + "检测项目"
            )
        elif source.name == "style.yaml":
            rule["style"].setdefault("toc", {})["mode"] = "native"
        elif source.name == "knowledge_coverage.yaml":
            rule["reportable_genes"] = genes
            rule["contract"]["ordered_gene_list_sha256"] = hashlib.sha256(
                "\n".join(genes).encode()
            ).hexdigest()
        elif source.name.startswith("reviewed_part3_"):
            # Inherit the same lung knowledge and its existing review state;
            # never activate an unreviewed candidate while changing product scope.
            for section in ("gene_sections", "drug_sections"):
                selected = []
                for row in rule.get(section) or []:
                    if str(row.get("gene") or "").upper() not in gene_set:
                        continue
                    allowed = row.get("panels")
                    if allowed is not None and "lung_588_pdl1" not in allowed:
                        continue
                    row["panels"] = [panel]
                    selected.append(row)
                rule[section] = selected
            rule.setdefault("source", {})["consumer_panels"] = [panel]
            rule["source"]["inherited_from_panel"] = "lung_588_pdl1"
            rule["source"]["inherited_rule_sha256"] = sha(source)
        write_yaml(package_dir / "rules" / source.name, rule)
    write_yaml(package_dir / "panel.yaml", raw)
    qa = yaml.safe_load((BASE / "qa.yaml").read_text())
    qa["panel_id"] = panel
    qa["current_output"]["section_aliases"] = {
        "variant_summary": ["基因突变信息"],
        "biomarkers": ["补充检测结果：TMB、MSI 与化疗药物基因组学"],
    }
    if not pdl1:
        qa["current_output"]["required_features"] = {
            k: v
            for k, v in qa["current_output"]["required_features"].items()
            if not k.startswith("pdl1_")
        }
    write_yaml(package_dir / "qa.yaml", qa)
    receipt["gene_count"] = len(genes)
    receipt["membership_column"] = flag
    receipt["source_rule_sha256"] = {
        p.name: sha(p)
        for p in (BASE / "rules").glob("*.yaml")
        if not p.name.startswith("._")
    }
    (work / panel / "build_receipt.json").write_text(
        json.dumps(receipt, indent=2, ensure_ascii=False) + "\n"
    )
    return receipt


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--private-dir", type=Path, required=True)
    parser.add_argument("--work-dir", type=Path, required=True)
    parser.add_argument("--packages-dir", type=Path, required=True)
    parser.add_argument("--panel", action="append")
    args = parser.parse_args()
    if (
        ".work" not in args.work_dir.resolve().parts
        or ".work" not in args.private_dir.resolve().parts
    ):
        raise ValueError("Private source/seed directories must be under .work")
    specs = yaml.safe_load((ROOT / "config/lung_draft_template_maps.yaml").read_text())[
        "panels"
    ]
    for panel in args.panel or specs:
        receipt = build_package(
            panel, specs[panel], args.private_dir, args.work_dir, args.packages_dir
        )
        print(
            json.dumps(
                {
                    "panel": panel,
                    "template_sha256": receipt["template_sha256"],
                    "hard_literals": receipt["hard_literals"],
                    "gene_count": receipt["gene_count"],
                }
            )
        )


if __name__ == "__main__":
    main()
