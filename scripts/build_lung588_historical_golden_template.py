#!/usr/bin/env python3
# 步骤: 74 肺癌588历史终版完整评审模板蒸馏
# 上游: 关系修复后的受控历史终版、私有清洗清单、肺癌588面板规则
# 输出: panels/lung_588_pdl1/templates/lung_588_pdl1_historical_golden_v1.docx
# 种子: 无（确定性文档变换）
"""Distil the historical lung588 final report into a data-driven review template.

The source document is patient-specific and must stay outside Git.  This
builder preserves its report geometry and fixed appendix while replacing all
case-specific result regions with ReportGen/Jinja fields.  The result is a
report-group review candidate, not a medically approved production report.
"""

from __future__ import annotations

import argparse
import copy
import json
import re
import sys
from pathlib import Path
from typing import Any, Iterable, Sequence

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from build_lung588_template import (
    PDL1_ASSAY_PROVENANCE_MARKER,
    PDL1_IMAGE_NOTICE,
    PDL1_SOURCE_PROVENANCE_MARKER,
    _body_paragraph,
    _collapse_between,
    _compact,
    _element_text,
    _find_table,
    _insert_paragraph_after,
    _load_lung588_gene_list,
    _neutralize_table,
    _normalize_zip_metadata,
    _remove_explicit_page_break_before,
    _remove_paragraph,
    _remove_static_pdl1_image,
    _replace_cell_text,
    _replace_paragraph_text,
    _sha256,
)

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from reportgen.core.review_candidate_contract import (  # noqa: E402
    load_review_candidate_contract,
    validate_review_candidate_template,
)


DEFAULT_OUTPUT = (
    ROOT
    / "panels"
    / "lung_588_pdl1"
    / "templates"
    / "lung_588_pdl1_historical_golden_v1.docx"
)
REVIEW_CANDIDATE_CONTRACT = (
    ROOT
    / "panels"
    / "lung_588_pdl1"
    / "review_baselines"
    / "lung588_historical_review_candidate_v1.yaml"
)
TEMPLATE_VERSION = "0.5.3-review.1"
EXPECTED_REPAIRED_SOURCE_SHA256 = (
    "e9cde046db0a93a5b33f13961ff7be25fcce9644a6682b7c132ca9e3604dbb96"
)
EXPECTED_ORIGINAL_SOURCE_SHA256 = (
    "4754ededa67eeeef1b716dd7fb9e907d03c8fd79904a64f48bd271119c9a401b"
)

REVIEW_DRAFT_NOTICE = (
    "报告组评审候选稿（非临床交付）：版式源自历史肺癌588终版；固定医学内容、"
    "规则命中和病例级结论均需报告组逐页复核后方可发布。"
)
PGX_REVIEW_NOTICE = (
    "以下药物基因组学明细由本病例Excel的CtDrug表原样映射，用于报告组检查"
    "字段、分页和解释口径；当前肺癌588面板尚未批准据此自动形成患者级用药结论。"
)
STATIC_APPENDIX_REVIEW_NOTICE = (
    "以下固定附录沿用历史终版版式，仅用于本轮报告组评审；流行病学数据、"
    "指南版本、药物适应证、参考文献和措辞时效性尚未完成当前版本医学复核。"
)
IMMUNE_REVIEW_NOTICE = (
    "本节展示病例数据驱动的研究性分子结果，仅供报告组核对字段和版式；"
    "不得据此单独预测免疫治疗获益、耐药或超进展。"
)
HLA_REVIEW_NOTICE = (
    "HLA-I分型直接来自本病例Excel，仅供报告组核对；等位基因、合子状态与"
    "免疫治疗结局之间的患者级推断当前未启用。"
)
PDL1_CLASSIFICATION_NOTICE_MARKER = "{{ pdl1_classification_notice }}"
LEGACY_FIXED_PDL1_CLASSIFICATION_TEXT = (
    "4、定性结果判定标准：TPS<1%，判定PD-L1蛋白表达为阴性/无表达；"
    "TPS（1-49%），判定PD-L1蛋白表达为低表达；TPS≥50%，判定PD-L1蛋白表达为高表达；"
    "CPS<1，判定PD-L1蛋白表达为阴性；CPS≥1，判定PD-L1蛋白表达为阳性。"
    "目前FDA推荐胃癌/胃食管交界处腺癌、尿路上皮癌、宫颈癌、食管鳞状细胞癌等患者"
    "在使用相关免疫抑制剂时采用CPS评分标准。TPS评分标准主要应用于肺癌。"
)

# Historical tables 17-49 correspond to the named CtDrug mappings below.
# Tables 26, 33 and 34 have no maintained data mapping and are neutralised.
DRUG_DETAIL_BINDINGS: tuple[tuple[int, str], ...] = (
    (17, "drug_shunbo"),
    (18, "drug_kabo"),
    (19, "drug_aoshalibo"),
    (20, "drug_boleihuahewu"),
    (21, "drug_jiaandieling"),
    (22, "drug_zishanchun"),
    (23, "drug_huanlinyanan"),
    (24, "drug_yihuanlinyanan"),
    (25, "drug_yilitikang"),
    (27, "drug_yituopogan"),
    (28, "drug_dakabazuo"),
    (29, "drug_genhuanlei"),
    (30, "drug_bolaimeisin"),
    (31, "drug_kaipeibaibin"),
    (32, "drug_fluorouracil"),
    (35, "drug_jixitabin"),
    (36, "drug_duoxitasai"),
    (37, "drug_peimeiqusai"),
    (38, "drug_changchunjianlei"),
    (39, "drug_mituogenquan"),
    (40, "drug_sananliulin"),
    (41, "drug_mafalan"),
    (42, "drug_tegafur"),
    (43, "drug_atangbaoyin"),
    (44, "drug_letrozole_anastrozole"),
    (45, "drug_disaimisong"),
    (46, "drug_qiangdesong"),
    (47, "drug_tamoxifen"),
    (48, "drug_yiximeitang"),
    (49, "drug_yidabixing"),
)


def _table_text(table, *, rows: int = 3) -> str:
    return _compact(
        " ".join(cell.text for row in table.rows[:rows] for cell in row.cells)
    )


def _set_table_loop_expressions(
    table,
    collection: str,
    expressions: Sequence[str],
    *,
    body_row_index: int = 1,
) -> None:
    """Replace historical body rows with a docxtpl whole-row loop."""
    if len(table.rows) <= body_row_index:
        raise ValueError(
            f"table loop {collection!r} has no row {body_row_index} to clone"
        )
    template_row = copy.deepcopy(table.rows[body_row_index]._tr)
    for row in list(table.rows[body_row_index:]):
        table._tbl.remove(row._tr)

    for _ in range(3):
        table._tbl.append(copy.deepcopy(template_row))
    start, body, end = table.rows[-3:]
    for row in (start, body, end):
        for cell in row.cells:
            _replace_cell_text(cell, "")

    if len(body.cells) != len(expressions):
        raise ValueError(
            f"table loop {collection!r} has {len(body.cells)} columns, "
            f"received {len(expressions)} expressions"
        )
    _replace_cell_text(start.cells[0], f"{{%tr for row in {collection} %}}")
    for cell, expression in zip(body.cells, expressions):
        _replace_cell_text(cell, expression)
    _replace_cell_text(end.cells[0], "{%tr endfor %}")


def _replace_basic_information(table) -> None:
    if len(table.rows) != 4 or len(table.columns) != 5:
        raise ValueError("unexpected historical patient-information table shape")
    replacements = (
        (0, 1, "{{ patient_name }}"),
        (0, 4, "{{ sample_type }}"),
        (1, 1, "{{ gender }}"),
        (1, 4, "{{ sampling_method }}"),
        (2, 1, "{{ age }}"),
        (2, 4, "{{ sample_site }}"),
        (3, 1, "{{ clinical_diagnosis }}"),
        (3, 4, "{{ sample_id }}"),
    )
    _replace_cell_text(table.rows[3].cells[3], "样本编号：")
    for row_index, cell_index, value in replacements:
        _replace_cell_text(table.rows[row_index].cells[cell_index], value)


def _replace_biomarker_summary(table) -> None:
    if len(table.rows) != 8 or len(table.columns) != 3:
        raise ValueError("unexpected historical biomarker-summary table shape")
    values = {
        "肿瘤突变负荷（TMB）": ("{{ tmb_summary }}", "{{ immuno_tips }}"),
        "微卫星不稳定性（MSI）": ("{{ msi_summary }}", "{{ msi_tips }}"),
        "PD-L1表达": ("{{ pdl1_result_display }}", "{{ pdl1_table_interpretation }}"),
        "HLA-I分型": (
            "详见3.4 HLA-I分型检测结果",
            "研究性分型结果，不能单独用于治疗决策。",
        ),
        "免疫正相关基因": (
            "{{ immune_positive_result }}",
            "研究性相关标志物，不能单独用于治疗决策。",
        ),
        "免疫负相关基因": (
            "{{ immune_negative_result }}",
            "研究性相关标志物，不能单独用于治疗决策。",
        ),
        "免疫超进展相关基因": (
            "{{ immune_hyperprogression_result }}",
            "研究性相关标志物，不能单独用于治疗决策。",
        ),
    }
    seen: set[str] = set()
    for row in table.rows[1:]:
        label = _compact(row.cells[0].text)
        for expected, replacement in values.items():
            if label == _compact(expected):
                _replace_cell_text(row.cells[1], replacement[0])
                _replace_cell_text(row.cells[2], replacement[1])
                seen.add(expected)
                break
    missing = sorted(set(values) - seen)
    if missing:
        raise ValueError(f"biomarker summary labels changed: {missing}")


def _replace_pdl1_table(table) -> None:
    if len(table.rows) != 2 or len(table.columns) != 5:
        raise ValueError("unexpected historical PD-L1 table shape")
    values = (
        "PD-L1蛋白表达",
        "免疫组化",
        "{{ pdl1_tps_display }}",
        "{{ pdl1_cps_display }}",
        "{{ pdl1_result_display }}",
    )
    for cell, value in zip(table.rows[1].cells, values):
        _replace_cell_text(cell, value)


def _normalize_gene_list_table(table) -> None:
    """Keep the historical geometry while enforcing the frozen gene symbols."""
    genes = _load_lung588_gene_list()
    cells = [cell for row in table.rows[1:] for cell in row.cells]
    if len(cells) < len(genes):
        raise ValueError(
            f"historical gene table has {len(cells)} body cells for {len(genes)} genes"
        )
    for index, cell in enumerate(cells):
        _replace_cell_text(cell, genes[index] if index < len(genes) else "")
    for row_index, row in enumerate(table.rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Cm(0.88 if row_index == 0 else 0.72)


def _set_hla_loop(table) -> None:
    _set_table_loop_expressions(
        table,
        "hla",
        (
            '{{ row.Locus or row["HLA-A"] or "" }}',
            '{{ row.Type1 or row.HET or "" }} / {{ row.Type2 or row.Allele2 or "" }}',
            '{{ row.HET or row.Zygosity or "" }}',
            "研究性分型结果，不能单独用于治疗决策。",
        ),
    )


def _set_drug_detail_loop(table, collection: str) -> None:
    if len(table.columns) != 6 or len(table.rows) < 2:
        raise ValueError(f"unexpected drug detail table shape for {collection}")
    _set_table_loop_expressions(
        table,
        collection,
        (
            '{{ row.DrugDisplay or "" }}',
            '{{ row.Gene or "" }}',
            '{{ row.Locus or "" }}',
            '{{ row.Level or "" }}',
            '{{ row.Genotype or "" }}',
            '{{ row.Result or "" }}',
        ),
    )


def _remove_signature_artifacts(document: DocumentObject) -> None:
    label = next(
        (
            paragraph
            for paragraph in document.paragraphs
            if "检测者" in (paragraph.text or "")
            and "审核者" in (paragraph.text or "")
            and "报告日期" in (paragraph.text or "")
        ),
        None,
    )
    if label is None:
        raise ValueError("historical signature label not found")
    previous = label._p.getprevious()
    if previous is not None:
        relationship_ids = {
            node.get(qn("r:embed"))
            for node in previous.iter(qn("a:blip"))
            if node.get(qn("r:embed"))
        }
        for relationship_id in relationship_ids:
            if relationship_id in document.part.rels:
                document.part.drop_rel(relationship_id)
        document.element.body.remove(previous)
    _replace_paragraph_text(
        label,
        "检测者：        审核者：        报告日期：{{ report_date_dot }}",
    )


def _replace_joined_token(root, old: str, new: str) -> int:
    changed = 0
    for paragraph in root.iter(qn("w:p")):
        nodes = list(paragraph.iter(qn("w:t")))
        if not nodes:
            continue
        joined = "".join(node.text or "" for node in nodes)
        if old not in joined:
            continue
        replaced = joined.replace(old, new)
        nodes[0].text = replaced
        for node in nodes[1:]:
            node.text = ""
        changed += 1
    return changed


def _story_parts(document: DocumentObject):
    seen: set[str] = set()
    candidates = [document.part]
    for section in document.sections:
        candidates.extend(
            (
                section.header.part,
                section.first_page_header.part,
                section.even_page_header.part,
                section.footer.part,
                section.first_page_footer.part,
                section.even_page_footer.part,
            )
        )
    for part in candidates:
        key = str(part.partname)
        if key in seen or not hasattr(part, "element"):
            continue
        seen.add(key)
        yield part


def _load_private_scrub_manifest(path: Path) -> dict[str, Any]:
    """Load case tokens from an ignored runtime file, never from source code."""
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        raise ValueError(f"cannot read private scrub manifest: {path}") from exc
    if not isinstance(payload, dict):
        raise ValueError("private scrub manifest must be a JSON object")

    raw_replacements = payload.get("replacements")
    if not isinstance(raw_replacements, list) or not raw_replacements:
        raise ValueError("private scrub manifest requires replacements")
    replacements: list[tuple[str, str]] = []
    for item in raw_replacements:
        if not isinstance(item, dict):
            raise ValueError("private scrub replacement must be an object")
        old = str(item.get("from") or "")
        new = str(item.get("to") or "")
        if not old or not new:
            raise ValueError("private scrub replacement requires from/to")
        replacements.append((old, new))

    sensitive_tokens = tuple(
        str(value) for value in payload.get("sensitive_tokens") or () if str(value)
    )
    forbidden_tokens = tuple(
        str(value) for value in payload.get("forbidden_tokens") or () if str(value)
    )
    if not sensitive_tokens or not forbidden_tokens:
        raise ValueError(
            "private scrub manifest requires sensitive_tokens and forbidden_tokens"
        )
    return {
        "source_sha256": str(payload.get("source_sha256") or ""),
        "replacements": tuple(replacements),
        "sensitive_tokens": sensitive_tokens,
        "forbidden_tokens": forbidden_tokens,
    }


def _scrub_story_content(
    document: DocumentObject,
    *,
    replacements: Sequence[tuple[str, str]],
    sensitive_tokens: Sequence[str],
) -> None:
    for part in _story_parts(document):
        root = part.element
        for old, new in replacements:
            _replace_joined_token(root, old, new)

        for instruction in root.iter(qn("w:instrText")):
            value = instruction.text or ""
            if (
                "file://" in value
                or "D:\\" in value
                or any(token in value for token in sensitive_tokens)
            ):
                instruction.text = ' HYPERLINK \\l "重要基因变异及潜在靶向药物信息" '
                continue
            for token in sensitive_tokens:
                value = value.replace(token, "")
            instruction.text = value

        for element in root.iter():
            for attr_name, attr_value in list(element.attrib.items()):
                value = attr_value
                for token in sensitive_tokens:
                    value = value.replace(token, "")
                if value != attr_value:
                    element.set(attr_name, value)


def _stabilize_body_header_layout(document: DocumentObject) -> None:
    """Replace the historical space-padded body header with fixed tab stops.

    The source places the patient label and slogan with a 3,780-twip first-line
    indent plus literal spaces.  That wraps the final slogan glyph when the
    historical CJK font is substituted during LibreOffice rendering.  Preserve
    the reviewed visual positions while expressing them as explicit tab stops.
    """

    seen: set[str] = set()
    matches = []
    for section in document.sections:
        for header in (
            section.header,
            section.first_page_header,
            section.even_page_header,
        ):
            part_name = str(header.part.partname)
            if part_name in seen:
                continue
            seen.add(part_name)
            matches.extend(
                paragraph
                for paragraph in header.paragraphs
                if "姓名：" in (paragraph.text or "")
                and "科技服务人类健康" in (paragraph.text or "")
            )

    if len(matches) != 1:
        raise ValueError(
            "expected one historical body-header patient/slogan paragraph, "
            f"found {len(matches)}"
        )

    paragraph = matches[0]
    _replace_paragraph_text(
        paragraph,
        "\t姓名：{{ patient_name }}\t科技服务人类健康",
    )
    paragraph_properties = paragraph._p.get_or_add_pPr()
    for tag in ("w:tabs", "w:ind"):
        for element in list(paragraph_properties.findall(qn(tag))):
            paragraph_properties.remove(element)

    tabs = OxmlElement("w:tabs")
    patient_tab = OxmlElement("w:tab")
    patient_tab.set(qn("w:val"), "left")
    patient_tab.set(qn("w:pos"), "3400")
    slogan_tab = OxmlElement("w:tab")
    slogan_tab.set(qn("w:val"), "right")
    slogan_tab.set(qn("w:pos"), "8200")
    tabs.extend((patient_tab, slogan_tab))

    paragraph_style = paragraph_properties.find(qn("w:pStyle"))
    insert_at = (
        paragraph_properties.index(paragraph_style) + 1
        if paragraph_style is not None
        else 0
    )
    paragraph_properties.insert(insert_at, tabs)


def _prune_unreferenced_images(document: DocumentObject) -> None:
    for part in _story_parts(document):
        root = part.element
        used = {
            value
            for node in root.iter()
            for value in (
                node.get(qn("r:embed")),
                node.get(qn("r:link")),
                node.get(qn("r:id")),
            )
            if value
        }
        for relationship_id, relationship in list(part.rels.items()):
            if not relationship.reltype.endswith("/image"):
                continue
            if relationship_id not in used:
                part.drop_rel(relationship_id)


def _harden_section_around_table(
    document: DocumentObject,
    start_text: str,
    end_text: str,
    tables: Iterable,
    keep_paragraphs: Iterable[str],
    notice: str,
) -> None:
    start = _body_paragraph(document, start_text)
    end = _body_paragraph(document, end_text)
    children = list(document.element.body.iterchildren())
    start_index = children.index(start._p)
    end_index = children.index(end._p)
    if end_index <= start_index:
        raise ValueError(f"invalid section order: {start_text!r} -> {end_text!r}")
    kept_tables = {table._tbl for table in tables}
    allowed_paragraphs = {_compact(value) for value in keep_paragraphs}
    found_tables: set[object] = set()
    for child in list(children[start_index + 1 : end_index]):
        if child in kept_tables:
            found_tables.add(child)
            continue
        if (
            child.tag == qn("w:p")
            and _compact(_element_text(child)) in allowed_paragraphs
        ):
            continue
        document.element.body.remove(child)
    if found_tables != kept_tables:
        raise ValueError(
            f"section {start_text!r} lost expected tables: "
            f"expected={len(kept_tables)}, found={len(found_tables)}"
        )
    inserted = _insert_paragraph_after(document, start, notice)
    inserted.paragraph_format.space_after = Pt(8)


def _replace_front_matter(document: DocumentObject) -> None:
    first = document.paragraphs[0]
    if not re.fullmatch(r"3{6,}", (first.text or "").strip()):
        raise ValueError("historical cover debug token changed")
    _replace_paragraph_text(first, "")

    cover_fields = {
        "姓    名": "姓    名：                {{ patient_name }}",
        "报告编号": "报告编号：            {{ report_number }}",
        "送检日期": "送检日期：               {{ receive_date_compact }}",
        "报告日期": "报告日期：               {{ report_date_compact }}",
    }
    for label, replacement in cover_fields.items():
        matches = [
            paragraph
            for paragraph in document.paragraphs[:30]
            if label in (paragraph.text or "")
        ]
        if len(matches) != 1:
            raise ValueError(
                f"expected one cover field {label!r}, found {len(matches)}"
            )
        _replace_paragraph_text(matches[0], replacement)

    title = _body_paragraph(document, "检测报告")
    banner = _insert_paragraph_after(document, title, REVIEW_DRAFT_NOTICE)
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in banner.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(192, 0, 0)
        run.font.size = Pt(11)


def _replace_introductory_copy(document: DocumentObject) -> None:
    replacements = {
        "报告第二部分：检测结果（提供本次检测涉及的靶向治疗、免疫治疗、化疗等综合检测结果），是本报告的关键信息。": (
            "报告第二部分：检测结果。展示本次病例数据和当前规则输出，是本轮报告组评审的重点。"
        ),
        "报告第三部分：基因变异及相应靶向/免疫药物解析，该部分对第二部分中的基因变异和靶向/免疫药物提示进行详细解析，并且包含本报告的阅读说明。": (
            "报告第三部分：结构化基因变异解释。未启用或未审核的事件会显示明确说明，不借用历史病例文本。"
        ),
        "报告第四部分：附录。提供所检测癌症的诊疗知识、癌症相关重要信号通路、所检测的基因列表信息、以及本报告的参考文献。": (
            "报告第四部分：历史终版固定附录、588基因列表、参考文献、质控与报告说明，供报告组逐页复核。"
        ),
        "使用免疫组化（IHC）方法，对委托人     组织     的PD-L1蛋白表达情况进行检测，给予用药提示。": (
            "使用免疫组化（IHC）方法检测PD-L1蛋白表达，报告TPS、CPS、结果判定和病例图像状态。"
        ),
        "分析样本微卫星不稳定性，给予用药提示。": (
            "分析样本微卫星不稳定性并报告检测状态；临床意义须结合现行规范复核。"
        ),
        "分析样本存在的基因变异情况，寻找与靶向/免疫药物相关的变异，给予用药提示与理论支持。": (
            "分析样本基因变异并展示当前规则输出；报告组需逐事件复核后方可形成正式结论。"
        ),
        "分析与化疗药物相关的基因变异，评估化疗的敏感性或毒副作用，为化疗方案的制订提供参考。": (
            "分析与化疗药物相关的基因变异，评估化疗的敏感性或毒副作用，"
            "为化疗方案的制订提供参考。"
        ),
    }
    for old, new in replacements.items():
        paragraph = _body_paragraph(document, old)
        if old.startswith("分析与化疗药物相关的基因变异"):
            _set_paragraph_text_with_cjk_font(paragraph, new)
        else:
            _replace_paragraph_text(paragraph, new)


def _set_paragraph_text_with_cjk_font(paragraph, value: str) -> None:
    """Replace one paragraph while preserving an explicit CJK-safe font."""

    paragraph.text = value
    for run in paragraph.runs:
        run.font.name = "微软雅黑"
        r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        r_fonts.set(qn("w:ascii"), "微软雅黑")
        r_fonts.set(qn("w:hAnsi"), "微软雅黑")
        r_fonts.set(qn("w:eastAsia"), "微软雅黑")


def _insert_appendix_review_notices(document: DocumentObject) -> None:
    for heading in (
        "第四部分：附录",
        "肺癌诊疗知识",
        "3. 癌症相关信号通路",
        "5. 参考文献",
    ):
        paragraph = _body_paragraph(document, heading)
        notice = _insert_paragraph_after(
            document, paragraph, STATIC_APPENDIX_REVIEW_NOTICE
        )
        for run in notice.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(192, 0, 0)


def _prepare_historical_toc_seed(document: DocumentObject) -> None:
    """Convert the historical floating TOC into a refreshable block control.

    The historical final report stores its title, decorative line, every TOC row,
    and the following section break in one paragraph made of floating text
    boxes.  ReportGen's final pagination pass deliberately operates on a block
    content control.  Keep the historical title decoration as a normal
    paragraph, then seed a small block control that owns the section break.
    The final processor replaces the seed with complete PAGEREF rows.
    """
    toc_paragraph = _body_paragraph(document, "目    录")
    source = toc_paragraph._p
    source_ppr = source.find(qn("w:pPr"))
    if source_ppr is None:
        raise ValueError("historical TOC paragraph has no paragraph properties")
    section_properties = source_ppr.find(qn("w:sectPr"))
    if section_properties is None:
        raise ValueError("historical TOC paragraph has no following section break")

    title = OxmlElement("w:p")
    title_ppr = copy.deepcopy(source_ppr)
    title_section_properties = title_ppr.find(qn("w:sectPr"))
    if title_section_properties is not None:
        title_ppr.remove(title_section_properties)
    title.append(title_ppr)

    title_run_count = 0
    for child in source:
        if child.tag != qn("w:r"):
            continue
        direct_text = "".join(node.text or "" for node in child.findall(qn("w:t")))
        drawing_names = {node.get("name") or "" for node in child.iter(qn("wp:docPr"))}
        if _compact(direct_text) in {"目", "录"} or any(
            name.startswith(("直接连接符", "椭圆")) for name in drawing_names
        ):
            title.append(copy.deepcopy(child))
            title_run_count += 1
    if title_run_count < 2:
        raise ValueError("historical TOC title decoration could not be isolated")

    content_control = OxmlElement("w:sdt")
    control_properties = OxmlElement("w:sdtPr")
    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), "ReportGen historical lung588 TOC")
    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), "ReportGenHistoricalLung588TOC")
    control_id = OxmlElement("w:id")
    control_id.set(qn("w:val"), "588202608")
    doc_part = OxmlElement("w:docPartObj")
    gallery = OxmlElement("w:docPartGallery")
    gallery.set(qn("w:val"), "Table of Contents")
    doc_part.append(gallery)
    doc_part.append(OxmlElement("w:docPartUnique"))
    control_properties.extend((alias, tag, control_id, doc_part))
    content_control.append(control_properties)
    content_control.append(OxmlElement("w:sdtEndPr"))
    content = OxmlElement("w:sdtContent")
    content_control.append(content)

    for text in (
        "第一部分：基本信息",
        "第二部分：检测结果",
        "第三部分：基因变异及相应靶向/免疫药物解析",
        "第四部分：附录",
        "参考文献",
    ):
        paragraph = OxmlElement("w:p")
        run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        text_node.text = text
        run.append(text_node)
        paragraph.append(run)
        content.append(paragraph)

    boundary = OxmlElement("w:p")
    boundary_properties = OxmlElement("w:pPr")
    boundary_properties.append(copy.deepcopy(section_properties))
    boundary.append(boundary_properties)
    parent = source.getparent()
    if parent is None:
        raise ValueError("historical TOC paragraph is detached")
    insert_at = parent.index(source)
    parent.remove(source)
    parent.insert(insert_at, title)
    parent.insert(insert_at + 1, content_control)
    parent.insert(insert_at + 2, boundary)


def _validate_template(
    document: DocumentObject, forbidden_tokens: Sequence[str]
) -> None:
    visible = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    found = [token for token in forbidden_tokens if token in visible]
    if found:
        raise ValueError(f"historical template retains patient tokens: {found}")

    required = (
        "{{ patient_name }}",
        "{{ sample_id }}",
        "{{ report_number }}",
        "{{ report_date_compact }}",
        "{{ pdl1_tps_display }}",
        "{{ pdl1_cps_display }}",
        "{{ pdl1_result_display }}",
        PDL1_CLASSIFICATION_NOTICE_MARKER,
        "{%tr for row in variants_2_1 %}",
        "{%tr for row in targeted_drug_tips %}",
        "{%tr for row in lung_guideline_drug_results %}",
        "{%tr for row in targeted_drug_introductions %}",
        "{%tr for row in nccn_results %}",
        "{%tr for row in immune_positive_results %}",
        "{%tr for row in chemotherapy_predictions %}",
        "{%tr for row in chemotherapy_regimen_predictions %}",
        "{%tr for row in chemotherapy_dosage_rows %}",
        "{%tr for row in irinotecan_safety_rows %}",
        "{%tr for row in hla %}",
        "__PART3_MARKER__",
        "__PDL1_CASE_IMAGE__",
        "Gene List for MLseq (n=588)",
        REVIEW_DRAFT_NOTICE,
        "{{ chemotherapy_summary_text }}",
    )
    missing = [token for token in required if token not in visible]
    if missing:
        raise ValueError(f"historical template is missing required markers: {missing}")
    for _table_index, collection in DRUG_DETAIL_BINDINGS:
        marker = f"{{%tr for row in {collection} %}}"
        if marker not in visible:
            raise ValueError(f"historical template is missing drug loop {collection}")

    gene_table = _find_table(document, ("Gene List for MLseq (n=588)",))
    rendered_genes = [
        _compact(cell.text)
        for row in gene_table.rows[1:]
        for cell in row.cells
        if _compact(cell.text)
    ]
    if rendered_genes != _load_lung588_gene_list():
        raise ValueError("historical template changed the frozen 588-gene order")
    if len(document.sections) != 5:
        raise ValueError(
            f"historical template must retain five sections, found {len(document.sections)}"
        )


def _scan_package_tokens(path: Path, forbidden_tokens: Sequence[str]) -> None:
    from zipfile import ZipFile

    found: dict[str, list[str]] = {}
    with ZipFile(path) as archive:
        for name in archive.namelist():
            if not name.endswith((".xml", ".rels")):
                continue
            text = archive.read(name).decode("utf-8", "ignore")
            matches = [token for token in forbidden_tokens if token in text]
            if matches:
                found[name] = matches
    if found:
        raise ValueError(f"patient tokens remain in DOCX package: {found}")


def build_template(source: Path, output: Path, scrub_manifest_path: Path) -> dict:
    scrub_manifest = _load_private_scrub_manifest(scrub_manifest_path)
    source_hash = _sha256(source)
    if source_hash == EXPECTED_ORIGINAL_SOURCE_SHA256:
        raise ValueError(
            "the original historical DOCX contains a broken ../NULL relationship; "
            "run scripts/repair_docx_relationships.py first and pass the repaired copy"
        )
    if source_hash != EXPECTED_REPAIRED_SOURCE_SHA256:
        raise ValueError(
            "historical lung588 source changed; review before rebuilding: "
            f"expected={EXPECTED_REPAIRED_SOURCE_SHA256}, actual={source_hash}"
        )
    if scrub_manifest["source_sha256"] != source_hash:
        raise ValueError("private scrub manifest does not match the controlled source")

    document = Document(source)
    if len(document.tables) != 70:
        raise ValueError(f"expected 70 historical tables, found {len(document.tables)}")
    source_tables = list(document.tables)

    _replace_front_matter(document)
    _replace_introductory_copy(document)
    _replace_basic_information(source_tables[0])
    _set_table_loop_expressions(
        source_tables[1],
        "targeted_drug_tips",
        (
            "{{ row.gene_display }}",
            "{{ row.variant_site }}",
            "{{ row.benefit_drugs }}",
            "{{ row.caution_drugs }}",
        ),
    )
    variant_count_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if (paragraph.text or "").startswith("*本次共检出体细胞变异：")
        and "其中与靶向药物用药相关的变异有：" in (paragraph.text or "")
    ]
    if len(variant_count_paragraphs) != 1:
        raise ValueError(
            "expected one historical variant-count sentence, found "
            f"{len(variant_count_paragraphs)}"
        )
    _replace_paragraph_text(
        variant_count_paragraphs[0],
        "*本次共检出体细胞变异：{{ total_variants_count }}个，其中与靶向药物"
        "用药相关的变异有：{{ drug_related_count }}个。",
    )
    _set_table_loop_expressions(
        source_tables[2],
        "lung_guideline_drug_results",
        (
            "{{ row.gene }}",
            "{{ row.drugs }}",
            "{{ row.clinical_note }}",
            "{{ row.result }}",
        ),
    )
    _replace_biomarker_summary(source_tables[3])
    _set_table_loop_expressions(
        source_tables[4],
        "chemotherapy_summary_rows",
        ("化疗药物小结", "{{ row.summary }}"),
    )
    _set_table_loop_expressions(
        source_tables[5],
        "variants_2_1",
        (
            "{{ row.gene }}",
            "{{ row.transcript }}",
            "{{ row.chr }}",
            "{{ row.exon }}",
            "{{ row.locus }}",
            "{{ row.var_type_cn }}",
            "{{ row.af_pct }}",
            "{{ row.benefit_drugs }}",
            "{{ row.caution_drugs }}",
        ),
        body_row_index=2,
    )
    _set_table_loop_expressions(
        source_tables[6],
        "targeted_drug_introductions",
        ("{{ row.drug_name }}", "{{ row.gene }}", "{{ row.introduction }}"),
    )
    _set_table_loop_expressions(
        source_tables[7],
        "nccn_results",
        ("{{ row.gene }}", "{{ row.content }}", "{{ row.result }}"),
    )
    _replace_pdl1_table(source_tables[8])
    _neutralize_table(
        source_tables[9],
        (
            "报告组评审中",
            "PD-L1",
            "历史固定免疫药物介绍已移除；以当前审核规则输出为准。",
        ),
    )
    _set_hla_loop(source_tables[10])
    for table, collection in zip(
        source_tables[11:14],
        (
            "immune_positive_results",
            "immune_negative_results",
            "immune_hyperprogression_results",
        ),
    ):
        _set_table_loop_expressions(
            table,
            collection,
            (
                "{{ row.gene }}",
                "{{ row.result }}",
                "研究性相关标志物，不能单独用于治疗决策。",
            ),
        )

    _set_table_loop_expressions(
        source_tables[14],
        "chemotherapy_predictions",
        ("{{ row.drug }}", "{{ row.genes }}", "{{ row.efficacy }}", "{{ row.toxicity }}"),
    )
    _set_table_loop_expressions(
        source_tables[15],
        "chemotherapy_regimen_predictions",
        ("{{ row.regimen }}", "{{ row.genes }}", "{{ row.efficacy }}", "{{ row.toxicity }}"),
    )
    _set_table_loop_expressions(
        source_tables[16],
        "chemotherapy_dosage_rows",
        ("{{ row.regimen }}", "{{ row.dosage }}"),
    )
    for table_index, collection in DRUG_DETAIL_BINDINGS:
        _set_drug_detail_loop(source_tables[table_index], collection)
    _set_table_loop_expressions(
        source_tables[26],
        "irinotecan_safety_rows",
        ("{{ row.drug }}", "{{ row.result }}", "{{ row.dose_evaluation }}"),
    )
    for table_index in (33, 34):
        title = source_tables[table_index].rows[0].cells[0].text.strip()
        _neutralize_table(
            source_tables[table_index],
            (title, "-", "-", "-", "-", "无维护映射，等待报告组决定是否保留"),
        )
    _normalize_gene_list_table(source_tables[68])

    _remove_static_pdl1_image(document)
    pdl1_image_marker = _body_paragraph(document, PDL1_IMAGE_NOTICE)
    _replace_paragraph_text(pdl1_image_marker, "__PDL1_CASE_IMAGE__")
    pdl1_caption = _insert_paragraph_after(
        document,
        pdl1_image_marker,
        "图1. 免疫组化：PD-L1",
    )
    pdl1_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pdl1_caption.paragraph_format.space_before = Pt(0)
    pdl1_caption.paragraph_format.space_after = Pt(0)
    pdl1_caption.paragraph_format.keep_with_next = True
    for run in pdl1_caption.runs:
        run.font.size = Pt(9)
    _remove_signature_artifacts(document)

    obsolete_guideline_heading = "非小细胞肺癌NCCN指南（2022 V3）"
    _collapse_between(
        document,
        obsolete_guideline_heading,
        "3. 免疫治疗疗效评估",
        (),
    )
    _remove_paragraph(_body_paragraph(document, obsolete_guideline_heading))
    _collapse_between(
        document,
        "3.1 肿瘤突变负荷（TMB）水平提示",
        "3.2 PD-L1表达检测结果",
        (
            "{{ tmb_detail_sentence }}",
            "{{ tmb_detail_interpretation }}",
            "{{ tmb_drug_note }}",
        ),
    )
    _replace_paragraph_text(
        _body_paragraph(document, LEGACY_FIXED_PDL1_CLASSIFICATION_TEXT),
        PDL1_CLASSIFICATION_NOTICE_MARKER,
    )
    _collapse_between(
        document,
        "用药提示",
        "3.3微卫星不稳定性（MSI）检测结果",
        (
            "{{ pdl1_table_interpretation }}",
            PDL1_ASSAY_PROVENANCE_MARKER,
            PDL1_SOURCE_PROVENANCE_MARKER,
        ),
        replacement_heading="结果解释边界",
    )
    _collapse_between(
        document,
        "3.3微卫星不稳定性（MSI）检测结果",
        "3.4 HLA-I分型检测结果",
        (
            "{{ msi_detail_sentence }}",
            "{{ msi_detail_interpretation }}",
            "{{ msi_tips }}",
        ),
    )
    _harden_section_around_table(
        document,
        "3.4 HLA-I分型检测结果",
        "3.5 免疫疗效正相关/负相关/超进展基因检测结果",
        (source_tables[10],),
        (),
        HLA_REVIEW_NOTICE,
    )
    _harden_section_around_table(
        document,
        "3.5 免疫疗效正相关/负相关/超进展基因检测结果",
        "4.化疗药物相关检测结果",
        source_tables[11:14],
        (
            "免疫治疗正相关基因检测结果",
            "免疫治疗负相关基因检测结果",
            "免疫治疗超进展基因检测结果",
        ),
        IMMUNE_REVIEW_NOTICE,
    )
    _set_paragraph_text_with_cjk_font(
        _body_paragraph(
            document,
            "化疗药物小结：经分析，可考虑优先选择的化疗方案有吉西他滨+长春瑞滨、吉西他滨单药方案、白蛋白结合型紫杉醇单药方案、紫杉醇单药方案。",
        ),
        "化疗药物小结：{{ chemotherapy_summary_text }}",
    )
    _replace_paragraph_text(
        _body_paragraph(
            document,
            "*请注意化疗药物毒副作用和有效性除了受遗传因素影响外，还受到临床上的多种因素影响，存在个体差异性。用药时请结合患者临床实际情况，谨遵医嘱。",
        ),
        "*本节为报告组评审数据，不构成化疗敏感性、毒性、剂量或方案建议。",
    )
    _remove_paragraph(
        _body_paragraph(
            document,
            '*详细的化疗基因多态性检测结果和药物信息请参见“4.3临床常用化疗药物评估及解析”。化疗药物的毒副作用和有效性常对应多个基因多态性位点，该报告根据证据支持等级等信息进行综合解读。"/"表示未有研究报道或现有研究结论不一致。',
        )
    )
    _replace_paragraph_text(
        _body_paragraph(document, "v 伊立替康用药剂量参考"),
        "伊立替康剂量安全性评价",
    )
    _remove_explicit_page_break_before(document, "具体的用药方案如下：")
    _remove_explicit_page_break_before(document, "5. 检测结果说明")
    _remove_explicit_page_break_before(document, "2.靶向药物相关检测结果")
    _remove_explicit_page_break_before(document, "肺癌诊疗知识")
    _collapse_between(
        document,
        "第三部分：基因变异及相应靶向/免疫药物解析",
        "3. 阅读说明",
        ("__PART3_MARKER__",),
    )
    _collapse_between(
        document,
        "3.1 文中参考文献及临床试验编号说明",
        "3.2 文中医学及生物学常见名词说明",
        (
            "正文如出现PMID或临床试验登记号，仅用于文献追溯；具体研究适用"
            "范围、入组条件及证据等级须结合原始文献和当前面板审核记录核对。",
        ),
    )
    _insert_appendix_review_notices(document)
    _prepare_historical_toc_seed(document)

    _scrub_story_content(
        document,
        replacements=scrub_manifest["replacements"],
        sensitive_tokens=scrub_manifest["sensitive_tokens"],
    )
    _stabilize_body_header_layout(document)
    _prune_unreferenced_images(document)
    document.core_properties.author = "ReportGen"
    document.core_properties.last_modified_by = "ReportGen"
    document.core_properties.title = "肺癌588基因+PD-L1报告组评审候选模板"
    document.core_properties.subject = "历史金标准版式蒸馏；非临床交付"
    document.core_properties.comments = ""
    _validate_template(document, scrub_manifest["forbidden_tokens"])

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    _normalize_zip_metadata(output)
    # A reload/save pass canonicalises namespace declarations exactly as the
    # family-alignment tool does, so the standalone builder reproduces the
    # registered template byte-for-byte instead of only semantically.
    canonical_document = Document(output)
    canonical_document.save(output)
    _normalize_zip_metadata(output)
    _scan_package_tokens(output, scrub_manifest["forbidden_tokens"])
    contract = load_review_candidate_contract(REVIEW_CANDIDATE_CONTRACT)
    candidate_gate = validate_review_candidate_template(
        contract,
        output,
        template_id="lung_588_pdl1_historical_golden_v1",
        template_version=TEMPLATE_VERSION,
        template_status="pilot",
    )
    if candidate_gate["status"] != "PASS":
        failed_codes = ", ".join(
            str(item.get("code") or "UNKNOWN")
            for item in candidate_gate.get("errors") or []
        )
        raise ValueError(
            "historical template violates the frozen review-candidate contract: "
            f"{failed_codes}"
        )
    rendered = Document(output)
    return {
        "source": str(source),
        "source_sha256": source_hash,
        "scrub_manifest_sha256": _sha256(scrub_manifest_path),
        "output": str(output),
        "output_sha256": _sha256(output),
        "paragraph_count": len(rendered.paragraphs),
        "table_count": len(rendered.tables),
        "section_count": len(rendered.sections),
        "drug_detail_loop_count": len(DRUG_DETAIL_BINDINGS),
        "review_candidate_contract_status": candidate_gate["status"],
        "review_candidate_contract_id": candidate_gate["contract_id"],
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--source",
        type=Path,
        required=True,
        help="关系修复后的受控历史终版；不得提交病例源文件。",
    )
    parser.add_argument(
        "--scrub-manifest",
        type=Path,
        required=True,
        help=".work 下的私有清洗 JSON；不得提交。",
    )
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    result = build_template(
        args.source.resolve(),
        args.output.resolve(),
        args.scrub_manifest.resolve(),
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
