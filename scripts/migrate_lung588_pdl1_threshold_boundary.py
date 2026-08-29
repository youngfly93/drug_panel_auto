#!/usr/bin/env python3
# 步骤: 75 肺癌588未知检测方案PD-L1阈值边界迁移
# 上游: panels/lung_588_pdl1/templates/lung_588_pdl1_historical_golden_v1.docx
# 输出: 原位更新同一模板，用动态来源分级说明取代旧固定TPS/CPS阈值
# 种子: 无（确定性DOCX文本迁移）
"""Remove the legacy universal PD-L1 cutoff paragraph from the lung588 template."""

from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path
import sys
import tempfile

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.build_lung588_historical_golden_template import (  # noqa: E402
    LEGACY_FIXED_PDL1_CLASSIFICATION_TEXT,
    PDL1_CLASSIFICATION_NOTICE_MARKER,
)
from scripts.build_lung588_template import (  # noqa: E402
    _normalize_zip_metadata,
    _replace_paragraph_text,
)

DEFAULT_TEMPLATE = (
    ROOT
    / "panels"
    / "lung_588_pdl1"
    / "templates"
    / "lung_588_pdl1_historical_golden_v1.docx"
)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def migrate_template(path: Path) -> dict[str, object]:
    """Apply the boundary marker once and leave an already migrated file untouched."""

    document = Document(path)
    old_matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == LEGACY_FIXED_PDL1_CLASSIFICATION_TEXT
    ]
    marker_matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == PDL1_CLASSIFICATION_NOTICE_MARKER
    ]
    before = _sha256(path)

    if not old_matches and len(marker_matches) == 1:
        return {
            "status": "PASS",
            "changed": False,
            "template": path.name,
            "sha256_before": before,
            "sha256_after": before,
        }
    if len(old_matches) != 1 or marker_matches:
        raise RuntimeError(
            "unexpected lung588 PD-L1 classification paragraph state: "
            f"legacy={len(old_matches)}, marker={len(marker_matches)}"
        )

    _replace_paragraph_text(old_matches[0], PDL1_CLASSIFICATION_NOTICE_MARKER)
    with tempfile.NamedTemporaryFile(
        dir=path.parent,
        prefix=f".{path.name}.",
        suffix=".tmp",
        delete=False,
    ) as handle:
        temporary = Path(handle.name)
    try:
        document.save(temporary)
        _normalize_zip_metadata(temporary)
        temporary.replace(path)
    finally:
        temporary.unlink(missing_ok=True)

    migrated = Document(path)
    visible = "\n".join(paragraph.text for paragraph in migrated.paragraphs)
    if LEGACY_FIXED_PDL1_CLASSIFICATION_TEXT in visible:
        raise RuntimeError("legacy fixed PD-L1 cutoff paragraph remains after migration")
    if visible.count(PDL1_CLASSIFICATION_NOTICE_MARKER) != 1:
        raise RuntimeError("dynamic PD-L1 classification marker is not unique")
    return {
        "status": "PASS",
        "changed": True,
        "template": path.name,
        "sha256_before": before,
        "sha256_after": _sha256(path),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    args = parser.parse_args()
    print(json.dumps(migrate_template(args.template.resolve()), ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
