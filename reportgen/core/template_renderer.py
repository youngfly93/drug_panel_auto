"""
模板渲染器

负责使用docxtpl渲染Docx模板。
"""

from pathlib import Path
from typing import Any, Optional, Sequence

from docx import Document

from reportgen.core.processors import (
    ProcessorContext,
    build_docx_processors,
    run_processors,
)
from reportgen.core.template_contract import (
    declared_validation_to_dict,
    extract_template_contract,
    validate_contract,
    validate_declared_contract,
)
from reportgen.models.report_data import ReportData
from reportgen.utils.logger import get_logger
from reportgen.utils.validators import validate_docx_file


class TemplateRenderer:
    """
    模板渲染器

    使用docxtpl将报告数据填充到Docx模板。
    """

    WORD_REFRESH_TIMEOUT_SECONDS = 20

    def __init__(self, log_file: Optional[str] = None, log_level: str = "INFO"):
        """
        初始化模板渲染器

        Args:
            log_file: 日志文件路径
            log_level: 日志级别
        """
        self.logger = get_logger(log_file=log_file, level=log_level)
        self.last_processor_report: list[dict[str, Any]] = []

    def render(
        self,
        template_path: str,
        report_data: ReportData,
        output_path: str,
        post_processor_names: Optional[Sequence[str]] = None,
    ) -> str:
        """
        渲染模板并保存

        Args:
            template_path: 模板文件路径
            report_data: 报告数据
            output_path: 输出文件路径

        Returns:
            输出文件路径

        Raises:
            FileNotFoundError: 模板文件不存在
            ValueError: 渲染失败
        """
        # 验证模板文件
        is_valid, error = validate_docx_file(template_path, must_exist=True)
        if not is_valid:
            self.logger.error("模板文件验证失败", template=template_path, error=error)
            raise FileNotFoundError(error)

        # 验证输出路径
        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)

        self.logger.info("开始渲染模板", template=template_path, output=output_path)
        self.last_processor_report = []

        try:
            try:
                from docxtpl import DocxTemplate
            except ModuleNotFoundError as e:
                raise ModuleNotFoundError(
                    "缺少依赖 'docxtpl'，无法渲染docx模板；请先安装 requirements.txt 中的依赖"
                ) from e

            # 加载模板
            doc = DocxTemplate(template_path)

            # 获取模板上下文
            context = self.build_context(report_data)

            self.logger.debug(
                "模板上下文",
                fields=len([k for k, v in context.items() if not isinstance(v, list)]),
                tables=len([k for k, v in context.items() if isinstance(v, list)]),
            )

            # 渲染
            doc.render(context)

            # 保存
            doc.save(output_path)

            self._run_post_render_processors(
                output_path,
                context,
                template_path,
                processor_names=post_processor_names,
            )

            # 验证生成的文件可以被正常打开
            try:
                Document(output_path)
            except Exception as verify_err:
                self.logger.error(
                    "生成的docx文件无法打开，可能已损坏",
                    output=output_path,
                    error=str(verify_err),
                )
                raise ValueError(f"生成的docx文件无法打开: {verify_err}")

            self.logger.info("模板渲染成功", output=output_path)

            return output_path

        except Exception as e:
            self.logger.error(
                "模板渲染失败", template=template_path, output=output_path, error=str(e)
            )
            raise ValueError(f"模板渲染失败: {e}")

    def build_post_render_processors(
        self, processor_names: Optional[Sequence[str]] = None
    ):
        """Build the ordered DOCX post-render processor chain."""
        return build_docx_processors(processor_names)

    def _run_post_render_processors(
        self,
        output_path: str,
        context: dict,
        template_path: str,
        processor_names: Optional[Sequence[str]] = None,
    ) -> None:
        """Run post-render processors and keep an execution report."""
        processor_context = ProcessorContext(
            renderer=self,
            output_path=output_path,
            template_path=template_path,
            template_context=context,
            logger=self.logger,
        )
        results = run_processors(
            self.build_post_render_processors(processor_names), processor_context
        )
        self.last_processor_report = [result.to_dict() for result in results]

        errors = [r for r in self.last_processor_report if r.get("status") == "ERROR"]
        self.logger.info(
            "DOCX后处理完成",
            processors=len(self.last_processor_report),
            errors=len(errors),
        )

    def _normalize_template_context(self, obj):
        """Normalize template context.

        docxtpl/Jinja2 will render None as the string 'None' when used as `{{ var }}`.
        For medical reports this is almost always undesired; normalize missing values
        to empty strings before rendering.
        """
        import math

        if obj is None:
            return ""

        # Handle NaN (float) without importing pandas/numpy
        if isinstance(obj, float):
            try:
                if math.isnan(obj):
                    return ""
            except Exception:
                pass
            return obj

        if isinstance(obj, dict):
            return {k: self._normalize_template_context(v) for k, v in obj.items()}

        if isinstance(obj, list):
            return [self._normalize_template_context(v) for v in obj]

        if isinstance(obj, tuple):
            return tuple(self._normalize_template_context(v) for v in obj)

        return obj

    @staticmethod
    def _truthy(value) -> bool:
        if isinstance(value, bool):
            return value
        return str(value or "").strip().lower() in {"1", "true", "yes", "y", "on"}

    @staticmethod
    def _report_content_config(context: dict | None) -> dict:
        if not isinstance(context, dict):
            return {}
        cfg = context.get("report_content")
        return cfg if isinstance(cfg, dict) else {}

    @staticmethod
    def _panel_style_config(context: dict | None, table_name: str) -> dict:
        if not isinstance(context, dict):
            return {}
        root = context.get("panel_style")
        if not isinstance(root, dict):
            return {}
        table_cfg = root.get(table_name)
        defaults = root.get("defaults")
        merged = dict(defaults) if isinstance(defaults, dict) else {}
        if isinstance(table_cfg, dict):
            merged.update(table_cfg)
        return merged

    @staticmethod
    def _bool_config(value: Any, default: bool) -> bool:
        if value is None:
            return default
        if isinstance(value, bool):
            return value
        text = str(value).strip().lower()
        if text in {"1", "true", "yes", "y", "on"}:
            return True
        if text in {"0", "false", "no", "n", "off"}:
            return False
        return default

    @staticmethod
    def _hex_color_config(value: Any, default: str) -> str:
        text = str(value or default).strip().lstrip("#").upper()
        if len(text) != 6 or not all(ch in "0123456789ABCDEF" for ch in text):
            return default.upper()
        return text

    @staticmethod
    def _float_config(value: Any, default: float) -> float:
        try:
            return float(value)
        except (TypeError, ValueError):
            return default

    @staticmethod
    def _int_text_config(value: Any, default: str) -> str:
        try:
            return str(int(value))
        except (TypeError, ValueError):
            return default

    @staticmethod
    def _set_run_font_name(run, font_name: str) -> None:
        if not font_name:
            return
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        run.font.name = font_name
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:eastAsia"), font_name)

    @staticmethod
    def _format_static_text(text: str, values: dict[str, Any]) -> str:
        try:
            return text.format(**values)
        except Exception:
            return text

    @staticmethod
    def _paragraph_specs(raw: Any) -> list[tuple[str, dict[str, Any]]]:
        specs: list[tuple[str, dict[str, Any]]] = []
        if not isinstance(raw, list):
            return specs

        allowed = {
            "bold",
            "size",
            "color",
            "page_break_before",
            "space_before_pt",
            "space_after_pt",
            "first_line_indent_cm",
            "line_spacing_multiple",
        }
        for item in raw:
            if isinstance(item, str):
                text = item
                options: dict[str, Any] = {}
            elif isinstance(item, dict):
                text = str(item.get("text") or "")
                options = {k: item[k] for k in allowed if k in item}
            else:
                continue
            specs.append((text, options))
        return specs

    def build_context(self, report_data: ReportData) -> dict:
        """Build a normalized template context from ReportData."""
        return self._normalize_template_context(report_data.get_template_context())

    def _remove_template_underlines(self, file_path: str) -> None:
        """Remove explicit run underlines inherited from template placeholders."""
        from docx.oxml.ns import qn

        doc = Document(file_path)
        changed = False

        def remove_underlines(element) -> None:
            nonlocal changed
            for underline in list(element.xpath(".//w:u")):
                if underline.get(qn("w:val")) in {"none", "false", "0"}:
                    continue
                parent = underline.getparent()
                if parent is not None:
                    parent.remove(underline)
                    changed = True

        containers = [doc.element]
        for section in doc.sections:
            containers.extend(
                [
                    section.header._element,
                    section.first_page_header._element,
                    section.even_page_header._element,
                    section.footer._element,
                    section.first_page_footer._element,
                    section.even_page_footer._element,
                ]
            )

        for element in containers:
            remove_underlines(element)

        if changed:
            doc.save(file_path)

    def _normalize_multiline_bullet_paragraphs(self, file_path: str) -> None:
        """Split newline-packed bullet text into separate bullet paragraphs.

        docxtpl renders newline characters inside a placeholder as line breaks in
        the same paragraph. Some reviewed templates reserve blank bullet
        paragraphs after that placeholder, which otherwise become visible empty
        bullet dots after rendering.
        """
        from copy import deepcopy

        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph

        doc = Document(file_path)
        changed = False

        def has_numbering(paragraph) -> bool:
            ppr = paragraph._p.pPr
            return ppr is not None and ppr.numPr is not None

        def has_embedded_object(paragraph) -> bool:
            xml = paragraph._p.xml
            return "<w:drawing" in xml or "<w:pict" in xml

        def non_empty_lines(paragraph) -> list[str]:
            return [
                line.strip()
                for line in (paragraph.text or "").splitlines()
                if line.strip()
            ]

        def is_empty_numbered_placeholder(paragraph) -> bool:
            return (
                has_numbering(paragraph)
                and not (paragraph.text or "").strip()
                and not has_embedded_object(paragraph)
            )

        def set_text(paragraph, text: str) -> None:
            if paragraph.runs:
                paragraph.runs[0].text = text
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(text)

        def remove_paragraph(paragraph) -> None:
            element = paragraph._element
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)

        def insert_numbered_after(paragraph, text: str) -> Paragraph:
            new_p = OxmlElement("w:p")
            ppr = paragraph._p.pPr
            if ppr is not None:
                new_p.append(deepcopy(ppr))
            paragraph._p.addnext(new_p)
            inserted = Paragraph(new_p, paragraph._parent)
            inserted.add_run(text)
            return inserted

        paragraphs = list(doc.paragraphs)
        idx = 0
        while idx < len(paragraphs):
            paragraph = paragraphs[idx]
            lines = non_empty_lines(paragraph)
            if not has_numbering(paragraph) or len(lines) <= 1:
                idx += 1
                continue

            set_text(paragraph, lines[0])
            cursor = paragraph
            next_idx = idx + 1

            for line in lines[1:]:
                if (
                    next_idx < len(paragraphs)
                    and is_empty_numbered_placeholder(paragraphs[next_idx])
                ):
                    cursor = paragraphs[next_idx]
                    set_text(cursor, line)
                    next_idx += 1
                else:
                    cursor = insert_numbered_after(cursor, line)

            while (
                next_idx < len(paragraphs)
                and is_empty_numbered_placeholder(paragraphs[next_idx])
            ):
                remove_paragraph(paragraphs[next_idx])
                next_idx += 1

            changed = True
            paragraphs = list(doc.paragraphs)
            idx += len(lines)

        if changed:
            doc.save(file_path)

    def _remove_empty_numbered_paragraphs(self, file_path: str) -> None:
        """Remove visible empty bullet/numbered paragraphs left by templates."""
        doc = Document(file_path)
        changed = False

        def has_numbering(paragraph) -> bool:
            ppr = paragraph._p.pPr
            return bool(ppr is not None and ppr.numPr is not None)

        def has_embedded_object(paragraph) -> bool:
            xml = paragraph._p.xml
            return "<w:drawing" in xml or "<w:pict" in xml

        def clear_numbering(paragraph) -> bool:
            ppr = paragraph._p.pPr
            if ppr is None or ppr.numPr is None:
                return False
            ppr.remove(ppr.numPr)
            return True

        for paragraph in list(doc.paragraphs):
            if not has_numbering(paragraph):
                continue
            if (paragraph.text or "").strip():
                continue
            if has_embedded_object(paragraph):
                continue
            parent = paragraph._element.getparent()
            if parent is None:
                continue
            parent.remove(paragraph._element)
            changed = True

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if (paragraph.text or "").strip():
                            continue
                        if clear_numbering(paragraph):
                            changed = True

        if changed:
            doc.save(file_path)

    def _restore_detection_content_underlines(self, file_path: str) -> None:
        """Restore intentional fill-in underlines in the detection-content sentence."""
        doc = Document(file_path)
        changed = False

        for paragraph in doc.paragraphs:
            text = paragraph.text or ""
            if "对委托人" not in text or "DNA" not in text or "基因进行检测" not in text:
                continue

            in_fill = False
            for run in paragraph.runs:
                run_text = run.text or ""
                if "对委托人" in run_text:
                    in_fill = True
                    if run.font.underline is not None:
                        run.font.underline = None
                        changed = True
                    continue
                if in_fill and run_text.strip().startswith("的"):
                    in_fill = False
                    if run.font.underline is not None:
                        run.font.underline = None
                        changed = True
                    continue
                if in_fill:
                    if run.font.underline is not True:
                        run.font.underline = True
                        changed = True
                elif run.font.underline is not None:
                    run.font.underline = None
                    changed = True

        if changed:
            doc.save(file_path)

    def _restore_patient_letter_fill_underlines(self, file_path: str) -> None:
        """Restore reviewed fill-in underlines in the patient letter.

        The global template-underline cleanup removes placeholder underlines
        from tables as intended, but the reviewed front letter intentionally
        underlines the patient-name and project-name fill-in spans. Restore
        only those sentence-bounded spans so table/body cleanup stays intact.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        doc = Document(file_path)
        changed = False

        def run_text(run) -> str:
            return "".join(node.text or "" for node in run.xpath(".//w:t"))

        def set_underline(run, enabled: bool) -> None:
            nonlocal changed
            r_pr = run.find(qn("w:rPr"))
            if r_pr is None:
                if not enabled:
                    return
                r_pr = OxmlElement("w:rPr")
                run.insert(0, r_pr)
            underlines = list(r_pr.findall(qn("w:u")))
            if enabled:
                if underlines:
                    return
                underline = OxmlElement("w:u")
                underline.set(qn("w:val"), "single")
                r_pr.append(underline)
                changed = True
                return
            for underline in underlines:
                r_pr.remove(underline)
                changed = True

        for paragraph in doc.element.xpath(".//w:p"):
            paragraph_text = "".join(node.text or "" for node in paragraph.xpath(".//w:t"))
            if not (
                ("尊敬的" in paragraph_text and "先生" in paragraph_text)
                or ("感谢您选择本机构" in paragraph_text and "检测项目" in paragraph_text)
            ):
                continue

            in_name_fill = False
            project_fill_armed = False
            in_project_fill = False

            for run in paragraph.xpath(".//w:r"):
                text = run_text(run)
                if not text:
                    continue

                # Some Word text boxes contain a non-visible fallback run with
                # the whole paragraph duplicated. Do not style that aggregate.
                if len(text) > 120:
                    continue

                if "尊敬的" in text:
                    in_name_fill = True
                    set_underline(run, False)
                    continue
                if "先生" in text and in_name_fill:
                    in_name_fill = False
                    set_underline(run, False)
                    continue
                if in_name_fill:
                    set_underline(run, True)
                    continue

                if "感谢您选择本机构" in text and "精心定" in text:
                    project_fill_armed = True
                    set_underline(run, False)
                    continue
                if project_fill_armed and text.strip() == "的":
                    project_fill_armed = False
                    in_project_fill = True
                    set_underline(run, False)
                    continue
                if "检测项目" in text and in_project_fill:
                    in_project_fill = False
                    set_underline(run, False)
                    continue
                if in_project_fill:
                    set_underline(run, True)

        if changed:
            doc.save(file_path)

    def _restore_msi_result_emphasis(
        self, file_path: str, context: dict | None = None
    ) -> None:
        """Restore reviewed emphasis in the MSI result section.

        The global underline cleanup intentionally removes template placeholder
        underlines from tables, but the reviewed CRC reports emphasize the
        result-specific MSI conclusion and the interpretation bullet matching
        the current MSI status. Apply that emphasis from the generated context
        instead of preserving a source-case hardcoded run layout.
        """
        from docx.oxml.ns import qn

        doc = Document(file_path)
        changed = False
        context = context or {}

        def normalized_status() -> str:
            raw = str(
                context.get("msi_status")
                or context.get("MSI状态")
                or context.get("msi_summary")
                or ""
            ).upper()
            raw = raw.replace(" ", "").replace("_", "-")
            if "MSI-H" in raw or "MSIH" in raw or "高度不稳定" in raw:
                return "MSI-H"
            if "MSI-L" in raw or "MSIL" in raw or "低度不稳定" in raw:
                return "MSI-L"
            if "MSS" in raw or "稳定" in raw:
                return "MSS"
            return ""

        def set_font(run, *, bold: bool = False, underline: bool = False) -> None:
            run.bold = bold
            run.underline = underline
            font_name = "微软雅黑"
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

        def rewrite_conclusion(paragraph) -> bool:
            text = paragraph.text or ""
            prefix = "依据本次检测结果，"
            marker = "该肿瘤样本为"
            if marker not in text:
                return False
            before, after_marker = text.split(marker, 1)
            if prefix in before:
                before = prefix
            if not after_marker.strip():
                return False

            paragraph.clear()
            normal = paragraph.add_run(before)
            set_font(normal)
            label = paragraph.add_run(marker)
            set_font(label, bold=True)
            result = paragraph.add_run(after_marker)
            set_font(result, bold=True, underline=True)
            return True

        def should_underline_interpretation(text: str, status: str) -> bool:
            compact = "".join((text or "").split())
            if status in {"MSS", "MSI-L"}:
                return "肿瘤组织为MSS" in compact and "MSI-L" in compact
            if status == "MSI-H":
                return "肿瘤组织为MSI-H" in compact
            return False

        status = normalized_status()
        for paragraph in doc.paragraphs:
            text = paragraph.text or ""
            if rewrite_conclusion(paragraph):
                changed = True
                continue
            if should_underline_interpretation(text, status):
                for run in paragraph.runs:
                    if run.text:
                        set_font(run, underline=True)
                        changed = True

        if changed:
            doc.save(file_path)

    def _restore_part3_dynamic_styles(
        self, file_path: str, context: dict | None = None
    ) -> None:
        """Restore reviewed Part 3 paragraph/run styles after underline cleanup.

        The golden Part 3 body is rendered dynamically from structured context.
        Later global cleanup removes template-inherited underlines, so the
        intentional reviewed styles for Part 3 need to be restored by semantic
        paragraph role rather than by case-specific text.
        """
        import re
        import xml.etree.ElementTree as ET
        from zipfile import ZipFile

        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        context = context or {}
        doc = Document(file_path)
        paragraphs = list(doc.paragraphs)
        changed = False

        start = None
        end = len(paragraphs)
        for idx, paragraph in enumerate(paragraphs):
            text = (paragraph.text or "").strip()
            if text.startswith("第三部分：基因变异及相应靶向/免疫药物解析"):
                start = idx
                continue
            if start is not None and (
                text.startswith("3. 阅读说明")
                or text.startswith("第四部分：")
                or text.startswith("4. 附录")
            ):
                end = idx
                break
        if start is None:
            return

        label_texts = {
            "基因简介：",
            "基因变异说明：",
            "基因变异解析：",
            "基因变异与药物关联分析：",
            "药物疗效临床解析：",
        }
        main_headings = {
            "基因变异解析",
            "靶向药物/免疫用药提示解析",
        }
        sub_headings = {
            "潜在获益靶向/免疫药物解析",
            "潜在负相关靶向/免疫药物解析",
        }
        variant_header_re = re.compile(r"^u\s+[^：]{1,20}：.+[；;]\s*\d")
        drug_variant_re = re.compile(r"^[A-Za-z0-9_-]+：.+突变相应.*药物$")
        summary_prefix = "在本次检测范围内，检出体细胞变异"
        context_drug_names = {
            line.strip()
            for section_key in ("drug_benefit_sections", "drug_caution_sections")
            for item in (context.get(section_key) or [])
            for line in str((item or {}).get("drug_name") or "").splitlines()
            if line.strip()
        }

        def decimal_outline_num_id() -> int:
            """Find the template's own 1 / 2 / 2.1 numbering id.

            Word ``numId`` values are local to each DOCX. The reviewed source
            uses one id, while the variableized template may assign another
            one after cleanup. Match the numbering definition instead of
            hardcoding a source-document id. Prefer a numbering instance that
            has not already been used in the document, otherwise Word/LibreOffice
            may continue a previous list and render "3. 基因变异解析" instead of
            the reviewed "1. 基因变异解析".
            """
            def paragraph_num_id(paragraph) -> str | None:
                p_pr = paragraph._p.pPr
                if p_pr is None:
                    return None
                num_pr = p_pr.find(qn("w:numPr"))
                if num_pr is None:
                    return None
                num_id_elem = num_pr.find(qn("w:numId"))
                return (
                    num_id_elem.get(qn("w:val"))
                    if num_id_elem is not None
                    else None
                )

            used_all = {
                num_id
                for paragraph in paragraphs
                for num_id in [paragraph_num_id(paragraph)]
                if num_id
            }
            used_before_part3 = {
                num_id
                for paragraph in paragraphs[:start]
                for num_id in [paragraph_num_id(paragraph)]
                if num_id
            }

            try:
                with ZipFile(file_path, "r") as zf:
                    numbering_xml = zf.read("word/numbering.xml")
                nroot = ET.fromstring(numbering_xml)
            except Exception:
                return 9

            ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            w_num = f"{{{ns_w}}}num"
            w_abstract_num = f"{{{ns_w}}}abstractNum"
            w_abstract_num_id = f"{{{ns_w}}}abstractNumId"
            w_num_id = f"{{{ns_w}}}numId"
            w_val = f"{{{ns_w}}}val"
            w_lvl = f"{{{ns_w}}}lvl"
            w_ilvl = f"{{{ns_w}}}ilvl"
            w_num_fmt = f"{{{ns_w}}}numFmt"
            w_lvl_text = f"{{{ns_w}}}lvlText"

            abstract_by_id = {
                elem.get(w_abstract_num_id): elem
                for elem in nroot.findall(w_abstract_num)
            }
            candidates: list[int] = []
            for num in nroot.findall(w_num):
                num_id = num.get(w_num_id)
                abstract_ref = num.find(w_abstract_num_id)
                abstract_id = abstract_ref.get(w_val) if abstract_ref is not None else None
                abstract = abstract_by_id.get(abstract_id)
                if not num_id or abstract is None:
                    continue
                levels = {lvl.get(w_ilvl): lvl for lvl in abstract.findall(w_lvl)}
                level0 = levels.get("0")
                level1 = levels.get("1")
                if level0 is None:
                    continue
                fmt0 = level0.find(w_num_fmt)
                text0 = level0.find(w_lvl_text)
                fmt1 = level1.find(w_num_fmt) if level1 is not None else None
                text1 = level1.find(w_lvl_text) if level1 is not None else None
                if (
                    fmt0 is not None
                    and fmt0.get(w_val) == "decimal"
                    and text0 is not None
                    and "%1" in (text0.get(w_val) or "")
                    and (
                        level1 is None
                        or (
                            fmt1 is not None
                            and fmt1.get(w_val) == "decimal"
                            and text1 is not None
                            and "%1.%2" in (text1.get(w_val) or "")
                        )
                    )
                ):
                    try:
                        candidates.append(int(num_id))
                    except ValueError:
                        continue
            if not candidates:
                return 9
            for candidate in sorted(candidates):
                if str(candidate) not in used_all:
                    return candidate
            for candidate in sorted(candidates):
                if str(candidate) not in used_before_part3:
                    return candidate
            return min(candidates)

        heading_num_id = decimal_outline_num_id()

        def ensure_num_pr(paragraph, num_id: int, ilvl_value: int = 0) -> None:
            p_pr = paragraph._p.get_or_add_pPr()
            existing = p_pr.find(qn("w:numPr"))
            if existing is not None:
                p_pr.remove(existing)
            num_pr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), str(ilvl_value))
            num_id_elem = OxmlElement("w:numId")
            num_id_elem.set(qn("w:val"), str(num_id))
            num_pr.append(ilvl)
            num_pr.append(num_id_elem)
            p_pr.insert(0, num_pr)

        def set_ind_xml(
            paragraph,
            *,
            first_line: int | None = None,
            first_line_chars: int | None = None,
            left: int | None = None,
            hanging: int | None = None,
        ) -> None:
            p_pr = paragraph._p.get_or_add_pPr()
            ind = p_pr.find(qn("w:ind"))
            if ind is None:
                ind = OxmlElement("w:ind")
                p_pr.append(ind)
            for attr in ("firstLine", "firstLineChars", "left", "start", "hanging"):
                q_attr = qn(f"w:{attr}")
                if q_attr in ind.attrib:
                    del ind.attrib[q_attr]
            if first_line is not None:
                ind.set(qn("w:firstLine"), str(first_line))
            if first_line_chars is not None:
                ind.set(qn("w:firstLineChars"), str(first_line_chars))
            if left is not None:
                ind.set(qn("w:left"), str(left))
            if hanging is not None:
                ind.set(qn("w:hanging"), str(hanging))

        def set_outline_level(paragraph, level: int | None) -> None:
            p_pr = paragraph._p.get_or_add_pPr()
            existing = p_pr.find(qn("w:outlineLvl"))
            if existing is not None:
                p_pr.remove(existing)
            if level is None:
                return
            outline = OxmlElement("w:outlineLvl")
            outline.set(qn("w:val"), str(level))
            p_pr.append(outline)

        def set_keep_next(paragraph, enabled: bool) -> None:
            p_pr = paragraph._p.get_or_add_pPr()
            existing = p_pr.find(qn("w:keepNext"))
            if enabled and existing is None:
                p_pr.append(OxmlElement("w:keepNext"))
            elif not enabled and existing is not None:
                p_pr.remove(existing)

        def set_spacing(paragraph, before: int | None = None, after: int | None = 0):
            paragraph.paragraph_format.space_before = (
                Pt(before / 20) if before is not None else None
            )
            paragraph.paragraph_format.space_after = (
                Pt(after / 20) if after is not None else None
            )

        def style_runs(
            paragraph,
            *,
            bold: bool,
            underline: bool,
            size: float,
            color: str | None = None,
            skip_prefix: bool = False,
        ) -> None:
            for run in paragraph.runs:
                if not (run.text or ""):
                    continue
                self._set_run_font_name(run, "微软雅黑")
                if skip_prefix and run.text.strip() == "u":
                    self._set_run_font_name(run, "Wingdings")
                    run.font.bold = False
                    run.font.underline = False
                    run.font.size = Pt(size)
                    continue
                run.font.bold = bold
                run.font.underline = underline
                run.font.size = Pt(size)
                if color:
                    run.font.color.rgb = RGBColor.from_string(color)

        in_drug_section = False
        awaiting_drug_names = False
        for paragraph in paragraphs[start + 1 : end]:
            text = (paragraph.text or "").strip()
            if not text:
                continue

            if text in main_headings:
                in_drug_section = text == "靶向药物/免疫用药提示解析"
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.JUSTIFY
                    if text == "靶向药物/免疫用药提示解析"
                    else None
                )
                ensure_num_pr(paragraph, heading_num_id, ilvl_value=0)
                set_ind_xml(paragraph)
                set_outline_level(paragraph, 0)
                set_spacing(paragraph, before=0, after=0)
                set_keep_next(paragraph, text == "基因变异解析")
                style_runs(paragraph, bold=True, underline=False, size=12)
                changed = True
                continue

            if text in sub_headings:
                in_drug_section = True
                paragraph.alignment = None
                ensure_num_pr(paragraph, heading_num_id, ilvl_value=1)
                set_ind_xml(paragraph)
                set_spacing(paragraph, before=200, after=0)
                set_keep_next(paragraph, True)
                style_runs(paragraph, bold=True, underline=False, size=12)
                changed = True
                continue

            if text.startswith(summary_prefix):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                ensure_num_pr(paragraph, 10, ilvl_value=0)
                set_ind_xml(paragraph)
                set_spacing(paragraph, after=0)
                set_keep_next(paragraph, False)
                style_runs(paragraph, bold=False, underline=False, size=10.5)
                changed = True
                continue

            if variant_header_re.match(text):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                set_spacing(paragraph, before=200, after=0)
                set_keep_next(paragraph, False)
                style_runs(
                    paragraph,
                    bold=True,
                    underline=True,
                    size=12,
                    skip_prefix=True,
                )
                changed = True
                continue

            if in_drug_section and drug_variant_re.match(text):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                ensure_num_pr(paragraph, 11)
                set_ind_xml(paragraph, left=420, hanging=420)
                set_spacing(paragraph, before=200, after=0)
                set_keep_next(paragraph, True)
                style_runs(paragraph, bold=True, underline=False, size=12, color="FF0000")
                awaiting_drug_names = True
                changed = True
                continue

            if in_drug_section and (
                text in context_drug_names
                or (awaiting_drug_names and text not in label_texts)
            ):
                ensure_num_pr(paragraph, 12)
                set_ind_xml(paragraph, left=420, hanging=420)
                set_spacing(paragraph, before=200, after=0)
                set_keep_next(paragraph, True)
                style_runs(paragraph, bold=True, underline=True, size=12, color="0000FF")
                changed = True
                continue

            if text in label_texts:
                awaiting_drug_names = False
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.JUSTIFY
                    if in_drug_section and text in {"基因变异与药物关联分析：", "药物疗效临床解析："}
                    else None
                )
                set_spacing(paragraph, after=0)
                set_keep_next(paragraph, True)
                style_runs(paragraph, bold=True, underline=False, size=10.5)
                changed = True
                continue

            # Narrative text: reviewed report uses two-character first-line
            # indent and both-side alignment for gene/drug analyses.
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_ind_xml(paragraph, first_line=420, first_line_chars=200)
            set_spacing(paragraph, after=0)
            set_keep_next(paragraph, False)
            style_runs(paragraph, bold=False, underline=False, size=10.5)
            changed = True

        if changed:
            doc.save(file_path)

    def _restore_variant_summary_table_style(
        self, file_path: str, context: dict | None = None
    ) -> None:
        """Restore reviewed link-style formatting in the 2.1 variant summary table."""
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor

        doc = Document(file_path)
        changed = False
        style_cfg = self._panel_style_config(context, "variant_summary_table")
        font_name = str(style_cfg.get("font_name") or "微软雅黑").strip()
        header_font_size = self._float_config(style_cfg.get("header_font_size"), 9.0)
        body_font_size = self._float_config(style_cfg.get("body_font_size"), 9.0)
        header_fill = self._hex_color_config(style_cfg.get("header_fill"), "00C4D8")
        header_font_color = RGBColor.from_string(
            self._hex_color_config(style_cfg.get("header_font_color"), "FFFFFF")
        )
        body_font_color = RGBColor.from_string(
            self._hex_color_config(style_cfg.get("body_font_color"), "000000")
        )
        border_color = self._hex_color_config(style_cfg.get("border_color"), "000000")
        border_size = self._int_text_config(style_cfg.get("border_size"), "6")
        link_underline = self._bool_config(style_cfg.get("link_underline"), True)
        link_color = RGBColor.from_string(
            self._hex_color_config(style_cfg.get("link_color"), "0000FF")
        )
        plain_texts = {
            str(value).strip()
            for value in (
                style_cfg.get("plain_texts")
                or ["未见突变", "未检出", "未检出有害变异", "-", "--", "—"]
            )
            if str(value).strip()
        }

        def is_variant_summary_table(table) -> bool:
            if len(table.columns) != 4 or not table.rows:
                return False
            header = " ".join(cell.text.replace("\n", "") for cell in table.rows[0].cells)
            return (
                "基因" in header
                and "突变位点" in header
                and "潜在获益靶向药物" in header
                and "可能耐药或慎重药物" in header
            )

        def set_cell_shading(cell, fill: str) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = tc_pr.find(qn("w:shd"))
            if shd is None:
                shd = OxmlElement("w:shd")
                tc_pr.append(shd)
            shd.set(qn("w:fill"), fill)

        def set_cell_borders(cell) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for side in ("top", "left", "bottom", "right"):
                border = borders.find(qn(f"w:{side}"))
                if border is None:
                    border = OxmlElement(f"w:{side}")
                    borders.append(border)
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), border_size)
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), border_color)

        def is_plain_text(text: str) -> bool:
            normalized = (text or "").strip()
            return normalized in plain_texts or any(
                token in normalized for token in plain_texts if len(token) > 1
            )

        def force_plain_run(run) -> None:
            run.font.bold = False
            run.font.underline = False
            run.font.color.rgb = body_font_color
            try:
                run.style = doc.styles["Default Paragraph Font"]
            except Exception:
                pass

        def style_cell(cell, row_idx: int, col_idx: int) -> None:
            header = row_idx == 0
            normalized_text = (cell.text or "").strip()
            plain_cell = is_plain_text(normalized_text)
            dash_only = normalized_text in {"-", "--", "—"}
            link_cell = (
                row_idx > 0
                and not plain_cell
                and (col_idx == 0 or (col_idx in {2, 3} and not dash_only))
            )

            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            if header:
                set_cell_shading(cell, header_fill)

            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    self._set_run_font_name(run, font_name)
                    run.font.size = Pt(header_font_size if header else body_font_size)
                    if header:
                        run.font.bold = True
                        run.font.underline = False
                        run.font.color.rgb = header_font_color
                    elif link_cell:
                        run.font.bold = False
                        run.font.underline = link_underline
                        run.font.color.rgb = link_color
                    else:
                        force_plain_run(run)
                    if not header and is_plain_text(run.text):
                        force_plain_run(run)

        for table in doc.tables:
            if not is_variant_summary_table(table):
                continue
            for row_idx, row in enumerate(table.rows):
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                if row_idx == 0:
                    row.height = Cm(0.4)
                else:
                    row.height = Cm(0.5)
                for col_idx, cell in enumerate(row.cells):
                    style_cell(cell, row_idx, col_idx)
            changed = True

        if changed:
            doc.save(file_path)

    def _restore_variant_detail_table_style(
        self, file_path: str, context: dict | None = None
    ) -> None:
        """Restore the reviewed 9-column 2.1 variant-detail table styling."""
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        doc = Document(file_path)
        changed = False
        style_cfg = self._panel_style_config(context, "variant_detail_table")
        font_name = str(style_cfg.get("font_name") or "微软雅黑").strip()
        header_font_size = self._float_config(style_cfg.get("header_font_size"), 9.0)
        body_font_size = self._float_config(style_cfg.get("body_font_size"), 9.0)
        header_fill = self._hex_color_config(style_cfg.get("header_fill"), "00C4D8")
        header_font_color = RGBColor.from_string(
            self._hex_color_config(style_cfg.get("header_font_color"), "F9FBFA")
        )
        body_font_color = RGBColor.from_string(
            self._hex_color_config(style_cfg.get("body_font_color"), "000000")
        )
        border_color = self._hex_color_config(style_cfg.get("border_color"), "000000")
        border_size = self._int_text_config(style_cfg.get("border_size"), "6")
        link_underline = self._bool_config(style_cfg.get("link_underline"), True)
        link_color = RGBColor.from_string(
            self._hex_color_config(style_cfg.get("link_color"), "0000FF")
        )
        plain_texts = {
            str(value).strip()
            for value in (
                style_cfg.get("plain_texts")
                or ["未见突变", "未检出", "未检出有害变异", "-", "--", "—"]
            )
            if str(value).strip()
        }

        def is_variant_detail_table(table) -> bool:
            if len(table.columns) != 9 or len(table.rows) < 2:
                return False
            row0 = " ".join(cell.text.replace("\n", "") for cell in table.rows[0].cells)
            row1 = " ".join(cell.text.replace("\n", "") for cell in table.rows[1].cells)
            return (
                "基因名称" in row0
                and "基因突变信息" in row0
                and "靶向药物信息" in row0
                and "转录本号" in row1
                and "潜在获益靶向药物" in row1
            )

        def set_cell_shading(cell, fill: str) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = tc_pr.find(qn("w:shd"))
            if shd is None:
                shd = OxmlElement("w:shd")
                tc_pr.append(shd)
            shd.set(qn("w:fill"), fill)

        def set_cell_borders(cell) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for side in ("top", "left", "bottom", "right"):
                border = borders.find(qn(f"w:{side}"))
                if border is None:
                    border = OxmlElement(f"w:{side}")
                    borders.append(border)
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), border_size)
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), border_color)

        def is_plain_text(text: str) -> bool:
            normalized = (text or "").strip()
            return normalized in plain_texts or any(
                token in normalized for token in plain_texts if len(token) > 1
            )

        def force_plain_run(run) -> None:
            run.font.bold = False
            run.font.underline = False
            run.font.color.rgb = body_font_color
            try:
                run.style = doc.styles["Default Paragraph Font"]
            except Exception:
                pass

        def set_run_font(run, *, header: bool, link: bool) -> None:
            self._set_run_font_name(run, font_name)
            run.font.size = Pt(header_font_size if header else body_font_size)
            run.font.bold = True if header else False
            if header:
                run.font.underline = False
                run.font.color.rgb = header_font_color
            elif link:
                run.font.underline = link_underline
                run.font.color.rgb = link_color
            else:
                force_plain_run(run)
            if not header and is_plain_text(run.text):
                force_plain_run(run)

        for table in doc.tables:
            if not is_variant_detail_table(table):
                continue
            for row_idx, row in enumerate(table.rows):
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                for col_idx, cell in enumerate(row.cells):
                    header = row_idx in {0, 1}
                    text = (cell.text or "").strip()
                    plain_cell = is_plain_text(text)
                    dash_only = text in {"", "-", "--", "—"}
                    link = (
                        row_idx >= 2
                        and not plain_cell
                        and (col_idx == 0 or (col_idx in {7, 8} and not dash_only))
                    )
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_borders(cell)
                    if header:
                        set_cell_shading(cell, header_fill)
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.line_spacing = 1.0
                        for run in paragraph.runs:
                            set_run_font(run, header=header, link=link)
            changed = True

        if changed:
            doc.save(file_path)

    def _restore_biomarker_table_style(
        self, file_path: str, context: dict | None = None
    ) -> None:
        """Restore template typography for the TMB/MSI biomarker result table."""
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        doc = Document(file_path)
        changed = False
        style_cfg = self._panel_style_config(context, "biomarker_table")
        font_name = str(style_cfg.get("font_name") or "微软雅黑").strip()
        header_font_size = self._float_config(style_cfg.get("header_font_size"), 10.0)
        body_font_size = self._float_config(style_cfg.get("body_font_size"), 9.0)
        header_fill = self._hex_color_config(style_cfg.get("header_fill"), "00C4D8")
        header_font_color = RGBColor.from_string(
            self._hex_color_config(style_cfg.get("header_font_color"), "F9FBFA")
        )
        body_font_color = RGBColor.from_string(
            self._hex_color_config(style_cfg.get("body_font_color"), "000000")
        )
        border_color = self._hex_color_config(style_cfg.get("border_color"), "000000")
        border_size = self._int_text_config(style_cfg.get("border_size"), "6")

        def is_biomarker_table(table) -> bool:
            if len(table.columns) != 3 or not table.rows:
                return False
            text = "\n".join(cell.text for row in table.rows for cell in row.cells)
            return "TMB/MSI/其它生物标志物检测结果" in text and "用药提示" in text

        def set_cell_shading(cell, fill: str) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = tc_pr.find(qn("w:shd"))
            if shd is None:
                shd = OxmlElement("w:shd")
                tc_pr.append(shd)
            shd.set(qn("w:fill"), fill)

        def set_cell_borders(cell) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for side in ("top", "left", "bottom", "right"):
                border = borders.find(qn(f"w:{side}"))
                if border is None:
                    border = OxmlElement(f"w:{side}")
                    borders.append(border)
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), border_size)
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), border_color)

        def apply_font(run, *, header: bool) -> None:
            self._set_run_font_name(run, font_name)
            run.font.size = Pt(header_font_size if header else body_font_size)
            run.font.bold = True if header else False
            run.font.underline = False
            run.font.color.rgb = header_font_color if header else body_font_color

        for table in doc.tables:
            if not is_biomarker_table(table):
                continue
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    header = row_idx == 0
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    set_cell_borders(cell)
                    if header:
                        set_cell_shading(cell, header_fill)
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.line_spacing = 1.0
                        for run in paragraph.runs:
                            apply_font(run, header=header)
            changed = True

        if changed:
            doc.save(file_path)

    def _restore_clinical_result_table_style(
        self, file_path: str, context: dict | None = None
    ) -> None:
        """Normalize reviewed borders for 2.2/2.3/3.3 tables and mark detected 3.3 results."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import RGBColor

        doc = Document(file_path)
        changed = False
        style_cfg = self._panel_style_config(context, "clinical_result_tables")
        border_color = self._hex_color_config(style_cfg.get("border_color"), "000000")
        border_size = self._int_text_config(style_cfg.get("border_size"), "6")
        detected_result_color = RGBColor.from_string(
            self._hex_color_config(style_cfg.get("detected_result_color"), "FF0000")
        )
        undetected_values = {
            str(value).strip()
            for value in (
                style_cfg.get("undetected_values")
                or ["", "-", "--", "—", "未检出", "未检出有害变异"]
            )
        }

        def clean_text(value: str) -> str:
            return "".join((value or "").split())

        def header_text(table) -> str:
            if not table.rows:
                return ""
            return clean_text(" ".join(cell.text for cell in table.rows[0].cells))

        def table_kind(table) -> str | None:
            header = header_text(table)
            if (
                len(table.columns) == 3
                and "药物名称" in header
                and "相关基因" in header
                and "药物适应情况" in header
            ):
                return "chemotherapy"
            if (
                len(table.columns) == 3
                and "检测基因" in header
                and "检测内容" in header
                and "检测结果" in header
            ):
                return "nccn"
            if (
                len(table.columns) == 3
                and "基因" in header
                and "检测结果" in header
                and "临床解读" in header
            ):
                return "immune"
            return None

        def set_table_borders(table) -> None:
            tbl_pr = table._tbl.tblPr
            borders = tbl_pr.find(qn("w:tblBorders"))
            if borders is None:
                borders = OxmlElement("w:tblBorders")
                tbl_pr.append(borders)
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = borders.find(qn(f"w:{side}"))
                if border is None:
                    border = OxmlElement(f"w:{side}")
                    borders.append(border)
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), border_size)
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), border_color)

        def set_cell_borders(cell) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for side in ("top", "left", "bottom", "right"):
                border = borders.find(qn(f"w:{side}"))
                if border is None:
                    border = OxmlElement(f"w:{side}")
                    borders.append(border)
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), border_size)
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), border_color)

        def is_detected_result(text: str) -> bool:
            normalized = (text or "").strip()
            compact = clean_text(normalized)
            return bool(
                normalized
                and compact not in {clean_text(value) for value in undetected_values}
                and not compact.startswith("未检出")
            )

        for table in doc.tables:
            kind = table_kind(table)
            if kind is None:
                continue
            set_table_borders(table)
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    set_cell_borders(cell)
                if kind == "immune" and row_idx > 0 and len(row.cells) >= 2:
                    result_cell = row.cells[1]
                    if is_detected_result(result_cell.text):
                        for paragraph in result_cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = detected_result_color
                                run.font.underline = False
            changed = True

        if changed:
            doc.save(file_path)

    def _apply_report_content_fixes(
        self,
        file_path: str,
        context: dict,
        template_path: str | None = None,
    ) -> None:
        """Apply deterministic report-level fixes that are hard to express in Jinja.

        This keeps one shared template usable for both 358 and 301 CRC reports while
        matching the reviewed final-report layout.
        """
        import copy
        import os
        import re
        import tempfile
        from zipfile import ZipFile

        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt

        doc = Document(file_path)
        changed = False
        content_cfg = self._report_content_config(context)

        class ParagraphProxy:
            def __init__(self, element):
                self._p = element

        def iter_all_paragraphs(container):
            for p in container.paragraphs:
                yield p
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from iter_all_paragraphs(cell)

        def replace_in_runs(paragraph, replacer) -> bool:
            updated = False
            for run in paragraph.runs:
                old = run.text
                new = replacer(old)
                if new != old:
                    run.text = new
                    updated = True
            return updated

        def set_paragraph_text(paragraph, text: str) -> bool:
            if paragraph.runs:
                paragraph.runs[0].text = text
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(text)
            return True

        def remove_paragraph(paragraph) -> bool:
            element = paragraph._element
            parent = element.getparent()
            if parent is None:
                return False
            parent.remove(element)
            return True

        def insert_paragraph_after(
            paragraph,
            text: str,
            *,
            bold: bool = False,
            size: float = 10.5,
            color: str | None = None,
            page_break_before: bool = False,
            space_before_pt: float | None = None,
            space_after_pt: float | None = None,
            first_line_indent_cm: float | None = None,
            line_spacing_multiple: float | None = None,
        ):
            new_p = OxmlElement("w:p")
            ppr = None

            def get_ppr():
                nonlocal ppr
                if ppr is None:
                    ppr = OxmlElement("w:pPr")
                return ppr

            if page_break_before:
                pbr = OxmlElement("w:pageBreakBefore")
                get_ppr().append(pbr)
            if first_line_indent_cm is not None:
                ind = OxmlElement("w:ind")
                ind.set(
                    qn("w:firstLine"),
                    str(int(float(first_line_indent_cm) * 567)),
                )
                get_ppr().append(ind)
            if (
                space_before_pt is not None
                or space_after_pt is not None
                or line_spacing_multiple is not None
            ):
                spacing = OxmlElement("w:spacing")
                if space_before_pt is not None:
                    spacing.set(qn("w:before"), str(int(float(space_before_pt) * 20)))
                if space_after_pt is not None:
                    spacing.set(qn("w:after"), str(int(float(space_after_pt) * 20)))
                if line_spacing_multiple is not None:
                    spacing.set(qn("w:line"), str(int(float(line_spacing_multiple) * 240)))
                    spacing.set(qn("w:lineRule"), "auto")
                get_ppr().append(spacing)
            if ppr is not None:
                new_p.append(ppr)
            run = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            if bold:
                rpr.append(OxmlElement("w:b"))
            if size:
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), str(int(size * 2)))
                rpr.append(sz)
                sz_cs = OxmlElement("w:szCs")
                sz_cs.set(qn("w:val"), str(int(size * 2)))
                rpr.append(sz_cs)
            if color:
                c = OxmlElement("w:color")
                c.set(qn("w:val"), color)
                rpr.append(c)
            run.append(rpr)
            t = OxmlElement("w:t")
            t.text = text
            t.set(qn("xml:space"), "preserve")
            run.append(t)
            new_p.append(run)
            paragraph._p.addnext(new_p)
            return ParagraphProxy(new_p)

        def insert_spacer_after(paragraph, space_after_pt: float):
            new_p = OxmlElement("w:p")
            ppr = OxmlElement("w:pPr")
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:after"), str(int(float(space_after_pt) * 20)))
            ppr.append(spacing)
            new_p.append(ppr)
            paragraph._p.addnext(new_p)
            return ParagraphProxy(new_p)

        def set_table_no_borders(table) -> None:
            tbl_pr = table._tbl.tblPr
            borders = tbl_pr.find(qn("w:tblBorders"))
            if borders is None:
                borders = OxmlElement("w:tblBorders")
                tbl_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                elem = borders.find(qn(f"w:{edge}"))
                if elem is None:
                    elem = OxmlElement(f"w:{edge}")
                    borders.append(elem)
                elem.set(qn("w:val"), "nil")

        def cm_to_twips(value: float) -> int:
            return int(float(value) * 567)

        def set_cell_margins_zero(cell) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for edge in ("top", "left", "bottom", "right"):
                elem = tc_mar.find(qn(f"w:{edge}"))
                if elem is None:
                    elem = OxmlElement(f"w:{edge}")
                    tc_mar.append(elem)
                elem.set(qn("w:w"), "0")
                elem.set(qn("w:type"), "dxa")

        def set_fixed_table_widths(table, widths_cm: list[float]) -> None:
            table.autofit = False
            tbl = table._tbl
            tbl_pr = tbl.tblPr
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), "fixed")

            total = sum(cm_to_twips(width) for width in widths_cm)
            tbl_w = tbl_pr.find(qn("w:tblW"))
            if tbl_w is None:
                tbl_w = OxmlElement("w:tblW")
                tbl_pr.append(tbl_w)
            tbl_w.set(qn("w:type"), "dxa")
            tbl_w.set(qn("w:w"), str(total))

            grid = tbl.find(qn("w:tblGrid"))
            if grid is None:
                grid = OxmlElement("w:tblGrid")
                tbl.insert(1, grid)
            for child in list(grid):
                grid.remove(child)
            for width in widths_cm:
                col = OxmlElement("w:gridCol")
                col.set(qn("w:w"), str(cm_to_twips(width)))
                grid.append(col)

            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    width = cm_to_twips(widths_cm[min(idx, len(widths_cm) - 1)])
                    tc_pr = cell._tc.get_or_add_tcPr()
                    tc_w = tc_pr.find(qn("w:tcW"))
                    if tc_w is None:
                        tc_w = OxmlElement("w:tcW")
                        tc_pr.append(tc_w)
                    tc_w.set(qn("w:type"), "dxa")
                    tc_w.set(qn("w:w"), str(width))

        temp_paths: list[str] = []

        def resolve_contact_qr(contact_cfg: dict) -> str | None:
            qr_path = str(contact_cfg.get("qr_image_path") or "").strip()
            if qr_path:
                path = Path(qr_path).expanduser()
                if not path.is_absolute():
                    path = Path(file_path).resolve().parent / path
                if path.exists():
                    return str(path)

            media_name = str(contact_cfg.get("qr_template_media") or "").strip()
            if not media_name or not template_path:
                return None
            tpl = Path(template_path)
            if not tpl.exists():
                return None
            try:
                with ZipFile(tpl) as zf:
                    if media_name not in zf.namelist():
                        return None
                    suffix = Path(media_name).suffix or ".png"
                    handle = tempfile.NamedTemporaryFile(
                        delete=False,
                        suffix=suffix,
                    )
                    with handle:
                        handle.write(zf.read(media_name))
                    temp_paths.append(handle.name)
                    return handle.name
            except Exception:
                return None

        def insert_contact_block_after(paragraph, contact_cfg: dict):
            if not isinstance(contact_cfg, dict) or not self._truthy(
                contact_cfg.get("enabled", True)
            ):
                return paragraph, False

            lines = contact_cfg.get("lines") or []
            lines = [str(line).strip() for line in lines if str(line).strip()]
            if not lines:
                return paragraph, False

            current = paragraph
            space_before = self._float_config(contact_cfg.get("space_before_pt"), 0.0)
            if space_before > 0:
                current = insert_spacer_after(current, space_before)

            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_no_borders(table)
            left_width = self._float_config(contact_cfg.get("left_width_cm"), 10.5)
            right_width = self._float_config(contact_cfg.get("right_width_cm"), 4.0)
            font_size = self._float_config(contact_cfg.get("font_size"), 9.0)
            qr_width = self._float_config(contact_cfg.get("qr_width_cm"), 3.6)

            for idx, cell in enumerate(table.rows[0].cells):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                cell.width = Cm(left_width if idx == 0 else right_width)
                set_cell_margins_zero(cell)
            set_fixed_table_widths(table, [left_width, right_width])

            left = table.rows[0].cells[0]
            left.text = ""
            for idx, line in enumerate(lines):
                p = left.paragraphs[0] if idx == 0 else left.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.space_before = Pt(0)
                run = p.add_run(line)
                run.font.size = Pt(font_size)

            right = table.rows[0].cells[1]
            right.text = ""
            p = right.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            qr = resolve_contact_qr(contact_cfg)
            if qr:
                run = p.add_run()
                run.add_picture(qr, width=Cm(qr_width))

            current._p.addnext(table._tbl)
            return ParagraphProxy(table._tbl), True

        def ensure_page_break_before(paragraph) -> bool:
            ppr = paragraph._p.get_or_add_pPr()
            if ppr.find(qn("w:pageBreakBefore")) is not None:
                return False
            ppr.append(OxmlElement("w:pageBreakBefore"))
            return True

        def normalized_status(value) -> str:
            return str(value or "").strip().upper()

        def has_tmb_h() -> bool:
            status = normalized_status(context.get("tmb_status"))
            if status in {"H", "HIGH", "TMB-H"}:
                return True
            if status:
                return False
            return "TMB-H" in str(context.get("tmb_summary") or "").upper()

        def has_msi_h() -> bool:
            status = normalized_status(context.get("msi_status"))
            if status in {"MSI-H", "MSIH", "HIGH"}:
                return True
            if status:
                return False
            return "MSI-H" in str(context.get("msi_summary") or "").upper()

        def has_known_tmb() -> bool:
            status = normalized_status(context.get("tmb_status"))
            return bool(status and status not in {"未检测", "NOT DETECTED", "NA", "N/A", "NONE"})

        def has_known_msi() -> bool:
            status = normalized_status(context.get("msi_status"))
            return bool(status and status not in {"未检测", "NOT DETECTED", "NA", "N/A", "NONE"})

        def has_immune_biomarker_conflict() -> bool:
            tmb_positive = has_tmb_h()
            msi_positive = has_msi_h()
            return (
                (tmb_positive and has_known_msi() and not msi_positive)
                or (msi_positive and has_known_tmb() and not tmb_positive)
            )

        def compact_date(text: str) -> str:
            return re.sub(r"\b(\d{4})-(\d{2})-(\d{2})\b", r"\1\2\3", text)

        def dotted_date(text: str) -> str:
            return re.sub(r"\b(\d{4})-(\d{2})-(\d{2})\b", r"\1.\2.\3", text)

        def normalized_project_code(value: str) -> str:
            text = str(value or "").strip().upper()
            if not text:
                return ""
            match = re.search(
                r"\b(?:MLJY[-_ ]?)?([A-Z]{1,5}\d{5,}[A-Z]?)\b",
                text,
                flags=re.IGNORECASE,
            )
            return match.group(1).upper() if match else text

        # Cover dates use compact yyyymmdd; signature/report footer uses yyyy.mm.dd.
        for paragraph in doc.paragraphs:
            text = paragraph.text or ""
            if "报告导读" in text:
                break
            if "送检日期" in text or "报告日期" in text:
                changed = replace_in_runs(paragraph, compact_date) or changed

        for paragraph in doc.paragraphs:
            text = paragraph.text or ""
            if "检测者" in text and "报告日期" in text:
                changed = replace_in_runs(paragraph, dotted_date) or changed

        # Keep the front matter aligned with the reviewed report: cover, report
        # guide, patient letter, and TOC are separate pages.
        for paragraph in doc.paragraphs:
            if (paragraph.text or "").strip() == "报告导读":
                changed = ensure_page_break_before(paragraph) or changed
                break

        # The first information table displays project code, not the source MLS file id.
        sample_id = str(context.get("sample_id") or "").strip()
        report_number = str(context.get("report_number") or "").strip()
        project_code = normalized_project_code(sample_id) or normalized_project_code(
            report_number
        )
        if project_code:
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join(cell.text for cell in row.cells)
                    if "项目编码" not in row_text:
                        continue
                    for cell in row.cells:
                        if sample_id and sample_id in cell.text:
                            for p in cell.paragraphs:
                                changed = replace_in_runs(
                                    p, lambda s: s.replace(sample_id, project_code)
                                ) or changed
                    break

        # 301/358 report body should reflect the selected project gene count.
        project_name = str(context.get("project_name") or "").strip()
        panel_match = re.search(r"(\d+)\s*基因", project_name)
        if panel_match:
            panel_count = panel_match.group(1)
            for paragraph in doc.paragraphs:
                if "与肿瘤密切相关的" in (paragraph.text or "") and "个基因" in paragraph.text:
                    # The numeric token is split into its own run in the template.
                    for run in paragraph.runs:
                        if run.text.strip().isdigit():
                            run.text = re.sub(r"\d+", panel_count, run.text)
                            changed = True
                    changed = replace_in_runs(
                        paragraph,
                        lambda s: re.sub(r"与肿瘤密切相关的\d+个基因", f"与肿瘤密切相关的{panel_count}个基因", s),
                    ) or changed

        brand_summary = str(context.get("targeted_drug_brand_summary") or "").strip()
        brand_prefix = "2.上表涉及的已上市的药物名称及对应的商品名称："
        for paragraph in doc.paragraphs:
            if "上表涉及的已上市的药物名称及对应的商品名称" not in (paragraph.text or ""):
                continue
            if brand_summary:
                changed = set_paragraph_text(
                    paragraph,
                    f"{brand_prefix}{brand_summary}",
                ) or changed
            else:
                changed = set_paragraph_text(
                    paragraph,
                    "2.上表未涉及可汇总商品名的已上市药物。",
                ) or changed
            break

        if not has_tmb_h():
            tmb_h_drug_note = re.compile(
                r"并且，\s*FDA\s*已批准帕博利珠单抗用于治疗\s*TMB-H\s*的不可切除或转移性的成人和儿童实体瘤。"
            )
            for paragraph in doc.paragraphs:
                text = paragraph.text or ""
                if (
                    "FDA" not in text
                    or "TMB-H" not in text
                    or "帕博利珠单抗" not in text
                    or "不可切除或转移性的成人和儿童实体瘤" not in text
                ):
                    continue
                new_text = tmb_h_drug_note.sub("", text).strip()
                new_text = re.sub(r"\s+。", "。", new_text)
                if new_text != text:
                    changed = set_paragraph_text(paragraph, new_text) or changed

        if not has_immune_biomarker_conflict():
            for paragraph in list(doc.paragraphs):
                text = paragraph.text or ""
                if "免疫治疗生物标志物包括" in text and "相对独立" in text:
                    changed = remove_paragraph(paragraph) or changed

        # Patient info table should stop at project code; hospital/pathology/QC rows
        # were requested to be removed from this location.
        remove_markers = ("送检医院", "病理号", "平均深度", "Q30", "覆盖度")
        for table in doc.tables:
            if not any("项目编码" in cell.text for row in table.rows for cell in row.cells):
                continue
            for row in list(table.rows):
                row_text = " ".join(cell.text for cell in row.cells)
                if any(marker in row_text for marker in remove_markers):
                    row._tr.getparent().remove(row._tr)
                    changed = True
            break

        # Add the configured consultation line before the patient letter if missing.
        consultation_line = str(context.get("consultation_line") or "").strip()
        has_phone = bool(consultation_line) and any(
            consultation_line in (p.text or "") for p in doc.paragraphs
        )
        if consultation_line and not has_phone:
            for paragraph in doc.paragraphs:
                if "检测结果仅对本次送检样本负责" not in (paragraph.text or ""):
                    continue
                new_p = OxmlElement("w:p")
                ppr = OxmlElement("w:pPr")
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), "left")
                ppr.append(jc)
                new_p.append(ppr)
                r = OxmlElement("w:r")
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rpr.append(b)
                r.append(rpr)
                t = OxmlElement("w:t")
                t.text = consultation_line
                r.append(t)
                new_p.append(r)
                paragraph._p.addnext(new_p)
                changed = True
                break

        # The reviewed CRC report contains a fuller patient letter than the base
        # template. The inserted paragraphs are configured under
        # settings.report_content.patient_letter.
        letter_text = "\n".join(p.text or "" for p in iter_all_paragraphs(doc))
        letter_cfg = content_cfg.get("patient_letter") if isinstance(content_cfg, dict) else {}
        letter_cfg = letter_cfg if isinstance(letter_cfg, dict) else {}
        greeting_specs = self._paragraph_specs(letter_cfg.get("after_greeting"))
        modern_specs = self._paragraph_specs(letter_cfg.get("after_modern_medicine"))
        configured_letter_text = "\n".join(text for text, _ in greeting_specs + modern_specs)
        has_configured_letter = bool(configured_letter_text) and all(
            text and text in letter_text
            for text, _ in greeting_specs + modern_specs
            if "{" not in text
        )
        if (greeting_specs or modern_specs) and not has_configured_letter:
            project_display = str(
                context.get("project_name")
                or context.get("project_display_name")
                or "本次基因检测"
            ).strip()
            format_values = {
                "project_name": project_display,
                "project_display_name": project_display,
            }
            greeting = None
            modern = None
            for paragraph in iter_all_paragraphs(doc):
                text = (paragraph.text or "").strip()
                if text == "您好！" and greeting is None:
                    greeting = paragraph
                if "现代医学已经证明" in text and modern is None:
                    modern = paragraph
            if greeting is not None:
                current = greeting
                for text, options in greeting_specs:
                    current = insert_paragraph_after(
                        current,
                        self._format_static_text(text, format_values),
                        **({"size": 10.5} | options),
                    )
                changed = True
            if modern is not None:
                current = modern
                for text, options in modern_specs:
                    current = insert_paragraph_after(
                        current,
                        self._format_static_text(text, format_values),
                        **({"size": 10.5} | options),
                    )
                changed = True

        # The patient letter starts on a new page.
        for paragraph in iter_all_paragraphs(doc):
            if (paragraph.text or "").strip() != "致您的一封信":
                continue
            changed = ensure_page_break_before(paragraph) or changed
            break

        # Restore configured company-introduction tail pages after the NGS method
        # limitation section.
        full_text = "\n".join(p.text or "" for p in iter_all_paragraphs(doc))
        tail_cfg = content_cfg.get("tail_content") if isinstance(content_cfg, dict) else {}
        tail_cfg = tail_cfg if isinstance(tail_cfg, dict) else {}
        tail_marker = str(tail_cfg.get("marker_text") or "").strip()
        tail_anchor = str(tail_cfg.get("anchor_text") or "").strip()
        tail_specs = self._paragraph_specs(tail_cfg.get("paragraphs"))
        if tail_marker and tail_anchor and tail_specs and tail_marker not in full_text:
            for paragraph in doc.paragraphs:
                if tail_anchor not in (paragraph.text or ""):
                    continue
                current = paragraph
                for text, kwargs in tail_specs:
                    current = insert_paragraph_after(current, text, **kwargs)
                current, contact_changed = insert_contact_block_after(
                    current,
                    tail_cfg.get("contact_block") or {},
                )
                changed = contact_changed or changed
                changed = True
                break

        if tail_marker and tail_marker in full_text:
            contact_cfg = tail_cfg.get("contact_block") or {}
            contact_lines = [
                str(line).strip()
                for line in (contact_cfg.get("lines") or [])
                if str(line).strip()
            ]
            if contact_lines and not all(line in full_text for line in contact_lines):
                for paragraph in doc.paragraphs:
                    if tail_marker not in (paragraph.text or ""):
                        continue
                    current = paragraph
                    # Attach the contact block after the last configured intro
                    # paragraph if it already exists.
                    for candidate in doc.paragraphs:
                        if (
                            "助力我国肿瘤精准医学的发展" in (candidate.text or "")
                            and candidate._p.getparent() is not None
                        ):
                            current = candidate
                    _, contact_changed = insert_contact_block_after(
                        current,
                        contact_cfg,
                    )
                    changed = contact_changed or changed
                    changed = True
                    break

        if changed:
            doc.save(file_path)
        for path in temp_paths:
            try:
                os.unlink(path)
            except OSError:
                pass

    def validate_template_contract(
        self,
        template_path: str,
        context: dict,
        contract_spec: dict | None = None,
    ) -> dict:
        """Validate template references and optional panel-declared requirements.

        Returns:
            A JSON-serializable dict describing missing variables/fields.
        """
        contract = extract_template_contract(template_path)
        validation = validate_contract(contract, context=context)
        declared_report = None
        declared_ok = True
        if contract_spec:
            declared_validation = validate_declared_contract(
                template_path,
                contract,
                contract_spec,
            )
            declared_report = declared_validation_to_dict(
                declared_validation,
                contract_spec,
            )
            declared_ok = bool(declared_validation.ok)

        return {
            "ok": bool(validation.ok and declared_ok),
            "template_path": contract.template_path,
            "required_paths": list(contract.required_paths),
            "required_lists": list(contract.required_lists),
            "loop_row_fields": {
                k: list(v) for k, v in contract.loop_row_fields.items()
            },
            "missing_paths": list(validation.missing_paths),
            "missing_lists": list(validation.missing_lists),
            "missing_row_fields": {
                k: list(v) for k, v in validation.missing_row_fields.items()
            },
            "missing_row_examples": validation.missing_row_examples,
            "declared_contract": declared_report,
        }

    def validate_template(self, template_path: str) -> tuple[bool, Optional[str]]:
        """
        验证模板文件

        Args:
            template_path: 模板文件路径

        Returns:
            (是否有效, 错误消息)
        """
        # 基本文件验证
        is_valid, error = validate_docx_file(template_path, must_exist=True)
        if not is_valid:
            return False, error

        # 尝试加载模板
        try:
            try:
                from docxtpl import DocxTemplate
            except ModuleNotFoundError:
                return False, (
                    "缺少依赖 'docxtpl'，无法校验模板；请先安装 requirements.txt 中的依赖"
                )

            DocxTemplate(template_path)
            self.logger.debug("模板验证成功", template=template_path)
            return True, None
        except Exception as e:
            error_msg = f"模板文件无效: {e}"
            self.logger.error("模板验证失败", template=template_path, error=str(e))
            return False, error_msg

    def get_template_variables(self, template_path: str) -> list[str]:
        """
        获取模板中的变量列表

        Args:
            template_path: 模板文件路径

        Returns:
            变量名列表（包括单值变量和循环变量）
        """
        import re
        from zipfile import ZipFile

        try:
            variables = set()

            # 读取docx文件（本质是zip包）
            with ZipFile(template_path, "r") as zf:
                # 读取document.xml（主要内容）
                xml_files = [
                    "word/document.xml",
                    "word/header1.xml",
                    "word/header2.xml",
                    "word/footer1.xml",
                    "word/footer2.xml",
                ]

                for xml_file in xml_files:
                    try:
                        content = zf.read(xml_file).decode("utf-8")

                        # 提取 {{ variable }} 格式的变量
                        # 匹配单值变量: {{ var }} 或 {{ obj.attr }}
                        single_var_re = (
                            r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*"
                            r"(?:\.[a-zA-Z_][a-zA-Z0-9_]*)?)\s*\}\}"
                        )
                        single_vars = re.findall(
                            single_var_re,
                            content,
                        )
                        variables.update(single_vars)

                        # 提取 {% for item in list %} 中的list变量
                        for_vars = re.findall(
                            r"\{%\s*for\s+\w+\s+in\s+([a-zA-Z_][a-zA-Z0-9_]*)\s*%\}",
                            content,
                        )
                        variables.update(for_vars)

                        # 提取 row.xxx 格式（循环内的字段引用）
                        row_vars = re.findall(r"row\.([a-zA-Z_][a-zA-Z0-9_]*)", content)
                        variables.update([f"row.{v}" for v in row_vars])

                        # 提取 row['xxx'] 格式
                        row_bracket_vars = re.findall(r"row\['([^']+)'\]", content)
                        variables.update([f"row['{v}']" for v in row_bracket_vars])

                    except KeyError:
                        # 文件不存在，跳过
                        continue

            result = sorted(list(variables))
            self.logger.debug(
                "提取模板变量成功", template=template_path, count=len(result)
            )
            return result

        except Exception as e:
            self.logger.error("提取模板变量失败", template=template_path, error=str(e))
            return []

    def validate_template_variables(
        self,
        template_path: str,
        available_variables: list[str],
        *,
        available_row_keys: Optional[set[str]] = None,
    ) -> tuple[bool, list[str], list[str]]:
        """
        校验模板变量是否都有对应的数据源

        Args:
            template_path: 模板文件路径
            available_variables: 可用的变量名列表（来自mapping配置）

        Returns:
            (是否全部匹配, 缺失变量列表, 未使用变量列表)
        """
        template_vars = self.get_template_variables(template_path)

        # 分离单值变量和循环变量
        single_vars = [
            v
            for v in template_vars
            if not v.startswith("row.") and not v.startswith("row[")
        ]

        # 检查缺失的变量（模板中有，但mapping中没有）
        missing_vars = [v for v in single_vars if v not in available_variables]

        # 校验循环行字段（row.* / row['*']）是否在mapping的列定义/同义词中出现
        if available_row_keys is not None:
            row_vars = [
                v for v in template_vars if v.startswith("row.") or v.startswith("row[")
            ]
            for v in row_vars:
                if v.startswith("row."):
                    key = v[len("row.") :]
                    if key and key not in available_row_keys:
                        missing_vars.append(v)
                elif v.startswith("row['") and v.endswith("']"):
                    key = v[len("row['") : -len("']")]
                    if key and key not in available_row_keys:
                        missing_vars.append(v)

        # 检查未使用的变量（mapping中有，但模板中没有）
        unused_vars = [v for v in available_variables if v not in single_vars]

        is_valid = len(missing_vars) == 0

        if missing_vars:
            self.logger.warning(
                "模板变量校验：发现未定义的变量",
                missing_count=len(missing_vars),
                missing_vars=missing_vars[:10],  # 只显示前10个
            )

        return is_valid, missing_vars, unused_vars

    # -------------------- internal helpers --------------------
    def _cleanup_empty_table_rows(self, file_path: str) -> None:
        """打开生成的docx，删除所有完全空白的表格行。

        空白行定义：该行所有单元格的 .text 去除空白后均为空字符串。
        """
        doc = Document(file_path)
        removed = 0
        for tbl in doc.tables:
            # 收集需要删除的 row 索引（从下往上删更安全）
            to_delete = []
            for idx, row in enumerate(tbl.rows):
                if all((cell.text or "").strip() == "" for cell in row.cells):
                    to_delete.append(idx)
            for idx in reversed(to_delete):
                tr = tbl.rows[idx]._tr
                tbl._tbl.remove(tr)
                removed += 1
        # #15: 删除只有表头的空数据表格（CNV/Fusion）
        tables_removed = 0
        cnv_fusion_markers = [
            (["起始位置", "终止位置", "拷贝数"], "CNV"),
            (["基因1", "基因2", "断点"], "Fusion"),
        ]
        for tbl in list(doc.tables):
            if len(tbl.rows) <= 1:  # 只有表头，无数据行
                header_text = " ".join(c.text.strip() for c in tbl.rows[0].cells) if tbl.rows else ""
                for markers, name in cnv_fusion_markers:
                    if all(m in header_text for m in markers):
                        tbl._tbl.getparent().remove(tbl._tbl)
                        tables_removed += 1
                        self.logger.debug(f"移除空{name}表格", table=name)
                        break
        if tables_removed:
            removed += tables_removed  # ensure save happens

        if removed or tables_removed:
            self.logger.debug("移除空白表格行", removed_rows=removed, removed_tables=tables_removed)
            doc.save(file_path)

    def _cleanup_cover_artifacts(self, file_path: str) -> None:
        """删除封面模板里误留的纯数字调试段落。"""
        import re

        doc = Document(file_path)
        removed = 0
        for paragraph in list(doc.paragraphs[:20]):
            text = (paragraph.text or "").strip()
            if not re.fullmatch(r"3{6,}", text):
                continue
            paragraph._element.getparent().remove(paragraph._element)
            removed += 1

        if removed:
            doc.save(file_path)
            self.logger.debug("已清理封面数字残留", removed=removed)

    def _cleanup_trailing_blank_page(self, file_path: str) -> None:
        """删除最后正文后多余的分页符和空白段落。

        只在最后一个有效正文块之后没有任何可见内容时处理，避免误删正常分页。
        """
        import os
        import shutil
        import tempfile
        import xml.etree.ElementTree as ET
        from zipfile import ZIP_DEFLATED, ZipFile

        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        w_body = f"{{{ns_w}}}body"
        w_p = f"{{{ns_w}}}p"
        w_tbl = f"{{{ns_w}}}tbl"
        w_t = f"{{{ns_w}}}t"
        w_br = f"{{{ns_w}}}br"
        w_drawing = f"{{{ns_w}}}drawing"
        w_pict = f"{{{ns_w}}}pict"
        w_sectpr = f"{{{ns_w}}}sectPr"
        w_type = f"{{{ns_w}}}type"
        ignorable_tail_tags = {
            f"{{{ns_w}}}bookmarkStart",
            f"{{{ns_w}}}bookmarkEnd",
            f"{{{ns_w}}}proofErr",
            f"{{{ns_w}}}permStart",
            f"{{{ns_w}}}permEnd",
        }

        def text_of(elem) -> str:
            return "".join((node.text or "") for node in elem.iter(w_t)).strip()

        def has_visible_content(elem) -> bool:
            if text_of(elem):
                return True
            if any(True for _ in elem.iter(w_drawing)):
                return True
            if any(True for _ in elem.iter(w_pict)):
                return True
            return False

        def is_blank_paragraph(elem) -> bool:
            return elem.tag == w_p and not has_visible_content(elem)

        def is_ignorable_tail_element(elem) -> bool:
            return is_blank_paragraph(elem) or elem.tag in ignorable_tail_tags

        with ZipFile(file_path, "r") as zin:
            document_xml = zin.read("word/document.xml")
            document_info = zin.getinfo("word/document.xml")
            other_entries = [
                (info, zin.read(info.filename))
                for info in zin.infolist()
                if info.filename != "word/document.xml"
            ]

        root = ET.fromstring(document_xml)
        body = root.find(f".//{w_body}")
        if body is None:
            return

        children = list(body)
        content_end = len(children)
        if content_end and children[-1].tag == w_sectpr:
            content_end -= 1

        last_content_idx = None
        for idx in range(content_end - 1, -1, -1):
            child = children[idx]
            if child.tag in (w_p, w_tbl) and has_visible_content(child):
                last_content_idx = idx
                break
        if last_content_idx is None:
            return

        trailing = children[last_content_idx + 1:content_end]
        if any(not is_ignorable_tail_element(elem) for elem in trailing):
            return

        changed = 0

        # 移除最后有效段落里位于文档末尾的分页符。
        last_content = children[last_content_idx]
        if last_content.tag == w_p:
            for br in list(last_content.iter(w_br)):
                if br.get(w_type) != "page":
                    continue
                parent = None
                for candidate in last_content.iter():
                    if br in list(candidate):
                        parent = candidate
                        break
                if parent is not None:
                    parent.remove(br)
                    changed += 1

        for elem in trailing:
            if is_blank_paragraph(elem):
                body.remove(elem)
                changed += 1

        if not changed:
            return

        fd, tmp_name = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            with ZipFile(tmp_name, "w", compression=ZIP_DEFLATED) as zout:
                zout.writestr(
                    document_info,
                    ET.tostring(root, encoding="utf-8", xml_declaration=True),
                )
                for info, data in other_entries:
                    zout.writestr(info, data)
            shutil.move(tmp_name, file_path)
        finally:
            if os.path.exists(tmp_name):
                os.unlink(tmp_name)

        self.logger.debug("已清理文档尾部空白页", changed=changed)

    def _remove_blank_page_breaks_before_headings(
        self, file_path: str, headings: tuple[str, ...] | None = None
    ) -> None:
        """Remove blank page-break paragraphs immediately before known headings.

        Some templates keep a blank paragraph containing only ``w:br type=page``
        before a section heading. When the previous table already ends exactly at
        a page boundary, Word/LibreOffice renders that paragraph as a visual
        blank page. This removes only contiguous blank paragraphs before exact
        heading text when at least one of those blanks carries a page break.
        """
        headings = headings or ("5. 参考文献",)
        target_headings = {str(item).strip() for item in headings if str(item).strip()}
        if not target_headings:
            return

        import os
        import shutil
        import tempfile
        import xml.etree.ElementTree as ET
        from zipfile import ZIP_DEFLATED, ZipFile

        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for prefix, uri in {
            "w": ns_w,
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }.items():
            ET.register_namespace(prefix, uri)

        w_body = f"{{{ns_w}}}body"
        w_p = f"{{{ns_w}}}p"
        w_ppr = f"{{{ns_w}}}pPr"
        w_keep_next = f"{{{ns_w}}}keepNext"
        w_t = f"{{{ns_w}}}t"
        w_br = f"{{{ns_w}}}br"
        w_type = f"{{{ns_w}}}type"
        w_last_rendered_page_break = f"{{{ns_w}}}lastRenderedPageBreak"
        w_drawing = f"{{{ns_w}}}drawing"
        w_pict = f"{{{ns_w}}}pict"

        def paragraph_text(elem) -> str:
            return "".join((node.text or "") for node in elem.iter(w_t)).strip()

        def has_page_break(elem) -> bool:
            return any(
                (br.get(w_type) or br.get("type")) == "page"
                for br in elem.iter(w_br)
            ) or any(elem.iter(w_last_rendered_page_break))

        def is_blank_paragraph(elem) -> bool:
            return (
                elem.tag == w_p
                and not paragraph_text(elem)
                and not any(elem.iter(w_drawing))
                and not any(elem.iter(w_pict))
            )

        def ensure_keep_next(elem) -> bool:
            ppr = elem.find(w_ppr)
            if ppr is None:
                ppr = ET.Element(w_ppr)
                elem.insert(0, ppr)
            if ppr.find(w_keep_next) is not None:
                return False
            ppr.append(ET.Element(w_keep_next))
            return True

        with ZipFile(file_path, "r") as zin:
            document_xml = zin.read("word/document.xml")
            document_info = zin.getinfo("word/document.xml")
            other_entries = [
                (info, zin.read(info.filename))
                for info in zin.infolist()
                if info.filename != "word/document.xml"
            ]

        root = ET.fromstring(document_xml)
        body = root.find(w_body)
        if body is None:
            return

        children = list(body)
        to_remove = []
        changed = False

        for idx, elem in enumerate(children):
            if elem.tag != w_p or paragraph_text(elem) not in target_headings:
                continue
            changed = ensure_keep_next(elem) or changed
            cluster = []
            saw_page_break = False
            prev_idx = idx - 1
            while prev_idx >= 0 and is_blank_paragraph(children[prev_idx]):
                prev = children[prev_idx]
                cluster.append(prev)
                saw_page_break = saw_page_break or has_page_break(prev)
                prev_idx -= 1
            if saw_page_break:
                to_remove.extend(cluster)

        removed = 0
        seen = set()
        for paragraph in to_remove:
            marker = id(paragraph)
            if marker in seen:
                continue
            body.remove(paragraph)
            seen.add(marker)
            removed += 1

        changed = changed or bool(removed)
        if not changed:
            return

        fd, tmp_name = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            with ZipFile(tmp_name, "w", compression=ZIP_DEFLATED) as zout:
                zout.writestr(
                    document_info,
                    ET.tostring(root, encoding="utf-8", xml_declaration=True),
                )
                for info, data in other_entries:
                    zout.writestr(info, data)
            shutil.move(tmp_name, file_path)
        finally:
            if os.path.exists(tmp_name):
                os.unlink(tmp_name)

        self.logger.debug("已移除标题前空白分页段落", removed=removed)

    def _normalize_front_matter_spacing(self, file_path: str) -> None:
        """Keep the report guide on its own page without template spacer blanks.

        The reviewed CRC golden source contains a manual page break followed by
        many empty paragraphs before ``报告导读``. Variableizing that source keeps
        those empty paragraphs, so Word renders the guide in the lower half of
        the page. This processor preserves the intended page separation while
        collapsing the spacer cluster to a single page-break paragraph.
        """
        import os
        import shutil
        import tempfile
        import xml.etree.ElementTree as ET
        from zipfile import ZIP_DEFLATED, ZipFile

        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        w_p = f"{{{ns_w}}}p"
        w_r = f"{{{ns_w}}}r"
        w_t = f"{{{ns_w}}}t"
        w_br = f"{{{ns_w}}}br"
        w_drawing = f"{{{ns_w}}}drawing"
        w_type = f"{{{ns_w}}}type"

        def para_text(elem) -> str:
            return "".join((node.text or "") for node in elem.iter(w_t)).strip()

        def has_drawing(elem) -> bool:
            return any(True for _ in elem.iter(w_drawing))

        def has_page_break(elem) -> bool:
            return any(node.attrib.get(w_type) == "page" for node in elem.iter(w_br))

        def is_blank_paragraph(elem) -> bool:
            return elem.tag == w_p and not para_text(elem) and not has_drawing(elem)

        def make_page_break_paragraph() -> ET.Element:
            paragraph = ET.Element(w_p)
            run = ET.SubElement(paragraph, w_r)
            br = ET.SubElement(run, w_br)
            br.set(w_type, "page")
            return paragraph

        with ZipFile(file_path, "r") as zin:
            document_xml = zin.read("word/document.xml")
            other_entries = [
                (info, zin.read(info.filename))
                for info in zin.infolist()
                if info.filename != "word/document.xml"
            ]
            document_info = zin.getinfo("word/document.xml")

        root = ET.fromstring(document_xml)
        changed = False
        removed = 0

        for parent in root.iter():
            children = list(parent)
            if not children:
                continue

            for idx, child in enumerate(children):
                if child.tag != w_p or para_text(child) != "报告导读":
                    continue

                prev_idx = idx - 1
                blank_cluster: list[ET.Element] = []
                saw_page_break = False
                while prev_idx >= 0 and is_blank_paragraph(children[prev_idx]):
                    prev = children[prev_idx]
                    saw_page_break = saw_page_break or has_page_break(prev)
                    blank_cluster.append(prev)
                    prev_idx -= 1

                if not blank_cluster:
                    continue

                for blank in blank_cluster:
                    parent.remove(blank)
                    removed += 1

                insert_at = list(parent).index(child)
                parent.insert(insert_at, make_page_break_paragraph())
                changed = True
                if not saw_page_break:
                    self.logger.debug("报告导读前缺少分页符，已补齐")
                break

        if not changed:
            return

        fd, tmp_name = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            with ZipFile(tmp_name, "w", compression=ZIP_DEFLATED) as zout:
                zout.writestr(
                    document_info,
                    ET.tostring(root, encoding="utf-8", xml_declaration=True),
                )
                for info, data in other_entries:
                    zout.writestr(info, data)
            shutil.move(tmp_name, file_path)
        finally:
            if os.path.exists(tmp_name):
                os.unlink(tmp_name)

        self.logger.debug("已清理报告导读前空白段落", removed=removed)

    def _cleanup_section_spacing(self, file_path: str) -> None:
        """删除章节标题前的空白段落。

        仅删除与章节标题同一父节点、且紧邻标题前的完全空白段落，
        避免误伤正常留白。
        """
        import copy
        import os
        import shutil
        import tempfile
        import xml.etree.ElementTree as ET
        from zipfile import ZIP_DEFLATED, ZipFile

        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        w_p = f"{{{ns_w}}}p"
        w_t = f"{{{ns_w}}}t"
        w_drawing = f"{{{ns_w}}}drawing"
        w_ppr = f"{{{ns_w}}}pPr"
        w_sectpr = f"{{{ns_w}}}sectPr"

        section_titles = (
            "报告导读",
            "致您的一封信",
            "第一部分：基本信息",
            "第二部分：检测结果",
            "第三部分：基因变异及相应靶向/免疫药物解析",
            "第四部分：附录",
        )

        def para_text(elem) -> str:
            return "".join((node.text or "") for node in elem.iter(w_t)).strip()

        def is_blank_paragraph(elem) -> bool:
            if para_text(elem):
                return False
            if any(True for _ in elem.iter(w_drawing)):
                return False
            return True

        def is_section_heading(elem) -> bool:
            text = para_text(elem)
            if not text or len(text) > 120:
                return False
            return any(title in text for title in section_titles)

        with ZipFile(file_path, "r") as zin:
            document_xml = zin.read("word/document.xml")
            other_entries = [(info, zin.read(info.filename)) for info in zin.infolist() if info.filename != "word/document.xml"]

        root = ET.fromstring(document_xml)
        removed = 0

        for parent in root.iter():
            children = list(parent)
            if not children:
                continue

            to_remove = []
            for idx, child in enumerate(children):
                if child.tag != w_p or not is_section_heading(child):
                    continue
                prev_idx = idx - 1
                blank_cluster = []
                while prev_idx >= 0:
                    prev = children[prev_idx]
                    if prev.tag != w_p or not is_blank_paragraph(prev):
                        break
                    blank_cluster.append(prev)
                    prev_idx -= 1
                if not blank_cluster:
                    continue

                sectpr = None
                for blank in blank_cluster:
                    ppr = blank.find(w_ppr)
                    if ppr is None:
                        continue
                    current_sectpr = ppr.find(w_sectpr)
                    if current_sectpr is not None:
                        sectpr = copy.deepcopy(current_sectpr)
                        break

                if sectpr is not None and prev_idx >= 0 and children[prev_idx].tag == w_p:
                    target_p = children[prev_idx]
                    target_ppr = target_p.find(w_ppr)
                    if target_ppr is None:
                        target_ppr = ET.Element(w_ppr)
                        target_p.insert(0, target_ppr)
                    existing_sectpr = target_ppr.find(w_sectpr)
                    if existing_sectpr is not None:
                        target_ppr.remove(existing_sectpr)
                    target_ppr.append(sectpr)

                to_remove.extend(blank_cluster)

            seen = set()
            for elem in to_remove:
                marker = id(elem)
                if marker in seen:
                    continue
                parent.remove(elem)
                seen.add(marker)
                removed += 1

        if removed == 0:
            return

        fd, tmp_name = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            with ZipFile(tmp_name, "w", compression=ZIP_DEFLATED) as zout:
                with ZipFile(file_path, "r") as zin:
                    document_info = zin.getinfo("word/document.xml")
                zout.writestr(
                    document_info,
                    ET.tostring(root, encoding="utf-8", xml_declaration=True),
                )
                for info, data in other_entries:
                    zout.writestr(info, data)
            shutil.move(tmp_name, file_path)
        finally:
            if os.path.exists(tmp_name):
                os.unlink(tmp_name)

        self.logger.debug("清理章节前空白段落", removed=removed)

    def _render_part3_formatted(self, file_path: str, context: dict) -> None:
        """Replace ``__PART3_MARKER__`` with data-driven Part 3 sections.

        对齐参考终版格式：
        - 变异标题：bold=True, size=12pt, color=FF0000(有药物)/0000FF(无药物), 前缀"u "
        - 变异说明：size=10.5pt
        - 基因简介：size=10.5pt
        - 药物标题：bold=True, color=FF0000
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document(file_path)
        content_cfg = self._report_content_config(context)

        # 找到占位标记段落
        placeholder_para = None
        for p in doc.paragraphs:
            if "__PART3_MARKER__" in p.text:
                placeholder_para = p
                break

        if placeholder_para is None:
            return  # 无占位标记，跳过

        def element_text(element) -> str:
            return "".join(node.text or "" for node in element.iter(qn("w:t")))

        def has_following_paragraph(prefix: str) -> bool:
            current = placeholder_para._element.getnext()
            while current is not None:
                if (
                    current.tag == qn("w:p")
                    and element_text(current).strip().startswith(prefix)
                ):
                    return True
                current = current.getnext()
            return False

        amino_table_element = None
        previous = placeholder_para._element.getprevious()
        if previous is not None and previous.tag == qn("w:tbl") and "氨基酸缩写" in element_text(previous):
            amino_table_element = previous
            amino_table_element.getparent().remove(amino_table_element)

        # 获取数据
        sections = context.get("gene_knowledge_sections", [])
        benefit_sections = context.get("drug_benefit_sections", [])
        caution_sections = context.get("drug_caution_sections", [])
        references = context.get("gene_references", [])
        total_count = context.get("total_variants_count", 0)
        drug_count = context.get("drug_related_count", 0)
        has_static_reading_section = has_following_paragraph("3. 阅读说明")
        has_static_reference_section = has_following_paragraph("5. 参考文献")

        # 辅助函数：在指定元素后插入新段落
        def add_para_after(
            prev_element,
            text,
            bold=False,
            size=10.5,
            color=None,
            prefix="",
            page_break_before=False,
            underline=False,
            font_name="微软雅黑",
            prefix_font_name=None,
            justify=False,
            spacing_before=None,
            spacing_after=None,
            first_line_twips=None,
            first_line_chars=None,
            left_twips=None,
            hanging_twips=None,
            num_id=None,
            ilvl=0,
            keep_next=False,
        ):
            new_p = OxmlElement("w:p")

            ppr = None

            def ensure_ppr():
                nonlocal ppr
                if ppr is None:
                    ppr = OxmlElement("w:pPr")
                    new_p.append(ppr)
                return ppr

            if page_break_before:
                ppr = ensure_ppr()
                page_break = OxmlElement("w:pageBreakBefore")
                ppr.append(page_break)
            if keep_next:
                ppr = ensure_ppr()
                ppr.append(OxmlElement("w:keepNext"))
            if num_id is not None:
                ppr = ensure_ppr()
                num_pr = OxmlElement("w:numPr")
                ilvl_elem = OxmlElement("w:ilvl")
                ilvl_elem.set(qn("w:val"), str(ilvl))
                num_id_elem = OxmlElement("w:numId")
                num_id_elem.set(qn("w:val"), str(num_id))
                num_pr.append(ilvl_elem)
                num_pr.append(num_id_elem)
                ppr.append(num_pr)
            if spacing_before is not None or spacing_after is not None:
                ppr = ensure_ppr()
                spacing = OxmlElement("w:spacing")
                if spacing_before is not None:
                    spacing.set(qn("w:before"), str(spacing_before))
                if spacing_after is not None:
                    spacing.set(qn("w:after"), str(spacing_after))
                ppr.append(spacing)
            if any(
                value is not None
                for value in (
                    first_line_twips,
                    first_line_chars,
                    left_twips,
                    hanging_twips,
                )
            ):
                ppr = ensure_ppr()
                ind = OxmlElement("w:ind")
                if left_twips is not None:
                    ind.set(qn("w:left"), str(left_twips))
                if hanging_twips is not None:
                    ind.set(qn("w:hanging"), str(hanging_twips))
                if first_line_twips is not None:
                    ind.set(qn("w:firstLine"), str(first_line_twips))
                if first_line_chars is not None:
                    ind.set(qn("w:firstLineChars"), str(first_line_chars))
                ppr.append(ind)
            if justify:
                ppr = ensure_ppr()
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), "both")
                ppr.append(jc)
            if prefix:
                # 前缀 run
                pr = OxmlElement("w:r")
                prPr = OxmlElement("w:rPr")
                if prefix_font_name:
                    r_fonts = OxmlElement("w:rFonts")
                    r_fonts.set(qn("w:ascii"), prefix_font_name)
                    r_fonts.set(qn("w:hAnsi"), prefix_font_name)
                    r_fonts.set(qn("w:eastAsia"), prefix_font_name)
                    r_fonts.set(qn("w:cs"), prefix_font_name)
                    prPr.append(r_fonts)
                if size:
                    psz = OxmlElement("w:sz")
                    psz.set(qn("w:val"), str(int(size * 2)))
                    prPr.append(psz)
                    pszCs = OxmlElement("w:szCs")
                    pszCs.set(qn("w:val"), str(int(size * 2)))
                    prPr.append(pszCs)
                if len(prPr):
                    pr.append(prPr)
                pt_elem = OxmlElement("w:t")
                pt_elem.text = prefix
                pt_elem.set(qn("xml:space"), "preserve")
                pr.append(pt_elem)
                new_p.append(pr)

            new_r = OxmlElement("w:r")
            # 格式
            rPr = OxmlElement("w:rPr")
            if font_name:
                r_fonts = OxmlElement("w:rFonts")
                r_fonts.set(qn("w:ascii"), font_name)
                r_fonts.set(qn("w:hAnsi"), font_name)
                r_fonts.set(qn("w:eastAsia"), font_name)
                r_fonts.set(qn("w:cs"), font_name)
                rPr.append(r_fonts)
            if bold:
                b_elem = OxmlElement("w:b")
                rPr.append(b_elem)
            if underline:
                u_elem = OxmlElement("w:u")
                u_elem.set(qn("w:val"), "single")
                rPr.append(u_elem)
            if size:
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), str(int(size * 2)))  # half-points
                rPr.append(sz)
                szCs = OxmlElement("w:szCs")
                szCs.set(qn("w:val"), str(int(size * 2)))
                rPr.append(szCs)
            if color:
                c_elem = OxmlElement("w:color")
                c_elem.set(qn("w:val"), color)
                rPr.append(c_elem)
            # 生产报告不加背景色（#5 fix: 去掉浅蓝阴影）

            new_r.append(rPr)
            new_t = OxmlElement("w:t")
            new_t.text = text
            new_t.set(qn("xml:space"), "preserve")
            new_r.append(new_t)
            new_p.append(new_r)

            prev_element.addnext(new_p)
            return new_p

        def add_text_block(prev_element, text, **options):
            current_element = prev_element
            lines = [line.strip() for line in str(text or "").splitlines()]
            lines = [line for line in lines if line]
            for line in lines:
                current_element = add_para_after(current_element, line, **options)
            return current_element

        # 从占位标记位置开始，链式插入
        current = placeholder_para._element

        body_options = {
            "size": 10.5,
            "justify": True,
            "first_line_twips": 420,
            "first_line_chars": 200,
            "spacing_after": 0,
        }
        label_options = {
            "bold": True,
            "size": 10.5,
            "spacing_after": 0,
        }

        # 总述
        current = add_para_after(
            current,
            f"在本次检测范围内，检出体细胞变异：{total_count}个，"
            f"其中与靶向/免疫药物相关的变异：{drug_count}个。"
            "（下面红色标注的为有对应靶向/免疫药物的基因变异。）",
            size=10.5,
            justify=True,
            spacing_after=0,
            num_id=10,
        )

        # === 基因变异解读 ===
        for section in sections:
            header = section.get("header", "")
            has_drug = section.get("has_drug", False)
            header_color = "FF0000" if has_drug else "0000FF"

            # 变异标题：bold, 12pt, red/blue, 前缀圆点 "● "
            current = add_para_after(
                current, header,
                bold=True,
                size=12,
                color=header_color,
                prefix="u ",
                prefix_font_name="Wingdings",
                underline=True,
                justify=True,
                spacing_before=200,
            )

            # 基因简介（紧跟标题，无多余空行）
            intro = section.get("intro", "")
            if intro:
                current = add_para_after(
                    current, "基因简介：", **label_options
                )
                current = add_text_block(current, intro, **body_options)

            # 基因变异说明
            desc = section.get("mutation_desc", "")
            if desc:
                current = add_para_after(
                    current, "基因变异说明：", **label_options
                )
                current = add_text_block(current, desc, **body_options)

            # 基因变异解析
            analysis = section.get("mutation_analysis", "")
            if analysis:
                current = add_para_after(
                    current, "基因变异解析：", **label_options
                )
                current = add_text_block(current, analysis, **body_options)

        # === 靶向药物解析 ===
        if benefit_sections or caution_sections:
            current = add_para_after(current, "")
            current = add_para_after(
                current, "靶向药物/免疫用药提示解析",
                bold=True,
                size=12,
                justify=True,
                spacing_after=0,
                num_id=9,
            )

        # 获益药物
        if benefit_sections:
            current = add_para_after(
                current, "潜在获益靶向/免疫药物解析",
                bold=True,
                size=12,
                spacing_before=200,
                keep_next=True,
                num_id=9,
                ilvl=1,
            )

            for ds in benefit_sections:
                gene = ds.get("gene", "")
                variant = ds.get("variant", "")
                drug_name = ds.get("drug_name", "")
                clinical = ds.get("clinical", "")
                header = ds.get("header")
                if header is None:
                    header = f"{gene}：{variant}突变相应靶向药物"

                if header:
                    current = add_para_after(
                        current,
                        header,
                        bold=True, size=12, color="FF0000",
                        justify=True,
                        spacing_before=200,
                        left_twips=420,
                        hanging_twips=420,
                        num_id=11,
                        keep_next=True,
                    )
                if drug_name:
                    current = add_text_block(
                        current,
                        drug_name,
                        bold=True,
                        size=12,
                        color="0000FF",
                        underline=True,
                        spacing_before=200,
                        left_twips=420,
                        hanging_twips=420,
                        num_id=12,
                        keep_next=True,
                    )
                relation = ds.get("relation", "")
                if relation:
                    current = add_para_after(
                        current,
                        "基因变异与药物关联分析：",
                        **label_options,
                        justify=True,
                    )
                    current = add_text_block(current, relation, **body_options)
                if clinical:
                    current = add_para_after(
                        current, "药物疗效临床解析：", **label_options, justify=True
                    )
                    current = add_text_block(current, clinical, **body_options)

        # 负相关药物
        if caution_sections:
            current = add_para_after(
                current, "潜在负相关靶向/免疫药物解析",
                bold=True,
                size=12,
                spacing_before=200,
                keep_next=True,
                num_id=9,
                ilvl=1,
            )

            for ds in caution_sections:
                gene = ds.get("gene", "")
                variant = ds.get("variant", "")
                drug_name = ds.get("drug_name", "")
                clinical = ds.get("clinical", "")
                header = ds.get("header")
                if header is None:
                    header = f"{gene}：{variant}突变相应负相关靶向药物"

                if header:
                    current = add_para_after(
                        current,
                        header,
                        bold=True, size=12, color="FF0000",
                        justify=True,
                        spacing_before=200,
                        left_twips=420,
                        hanging_twips=420,
                        num_id=11,
                        keep_next=True,
                    )
                if drug_name:
                    current = add_text_block(
                        current,
                        drug_name,
                        bold=True,
                        size=12,
                        color="0000FF",
                        underline=True,
                        spacing_before=200,
                        left_twips=420,
                        hanging_twips=420,
                        num_id=12,
                        keep_next=True,
                    )
                relation = ds.get("relation", "")
                if relation:
                    current = add_para_after(
                        current,
                        "基因变异与药物关联分析：",
                        **label_options,
                        justify=True,
                    )
                    current = add_text_block(current, relation, **body_options)
                if clinical:
                    current = add_para_after(
                        current, "药物疗效临床解析：", **label_options, justify=True
                    )
                    current = add_text_block(current, clinical, **body_options)

        # === 参考文献 ===
        # The CRC golden template keeps the reviewed appendix-level reference
        # section. Only inline references when a marker-only template has no
        # static final reference section after Part 3.
        if references and not has_static_reference_section:
            current = add_para_after(
                current, "参考文献",
                bold=True, size=12,
            )
            current = add_para_after(current, "")
            for ref in references:
                current = add_para_after(current, ref, size=9)

        if not has_static_reading_section:
            reading_blocks = self._paragraph_specs(
                content_cfg.get("part3_reading_blocks")
            )
            for text, options in reading_blocks:
                para_options = {"size": 10.5}
                para_options.update(options)
                current = add_para_after(current, text, **para_options)

        if amino_table_element is not None:
            current.addnext(amino_table_element)
            current = amino_table_element

        # 删除占位标记段落
        placeholder_para._element.getparent().remove(placeholder_para._element)

        doc.save(file_path)
        self.logger.info(
            "Part 3 格式化渲染完成",
            sections=len(sections),
            benefit=len(benefit_sections),
            caution=len(caution_sections),
            references=len(references),
        )

    def _render_signature_placeholder(self, file_path: str, context: dict) -> None:
        """处理签名占位符。

        当前优先保证最终报告不出现原始 `__SIG_IMG__` 文本；
        若提供了有效签名图片，则在占位处插入图片。
        """
        from pathlib import Path

        from docx.oxml import OxmlElement
        from docx.shared import Mm
        from docx.text.run import Run

        doc = Document(file_path)
        signature_path = str(context.get("signature_image_path") or "").strip()
        signature_exists = bool(signature_path) and Path(signature_path).exists()
        updated = False

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if "__SIG_IMG__" not in run.text:
                    continue

                before, after = run.text.split("__SIG_IMG__", 1)
                run.text = before

                if signature_exists:
                    new_r = OxmlElement("w:r")
                    run._r.addnext(new_r)
                    picture_run = Run(new_r, paragraph)
                    picture_run.add_picture(signature_path, width=Mm(22))
                    anchor_run = new_r
                else:
                    anchor_run = run._r

                if after:
                    tail_r = OxmlElement("w:r")
                    anchor_run.addnext(tail_r)
                    tail_run = Run(tail_r, paragraph)
                    tail_run.text = after

                updated = True

        if updated:
            doc.save(file_path)
            self.logger.debug(
                "签名占位处理完成",
                image_inserted=signature_exists,
            )

    def _replace_signature_anchor_images(self, file_path: str, context: dict) -> None:
        """Replace detector/reviewer signature images from context paths.

        Reviewed templates keep the handwritten signatures as positioned
        drawings before the "检测者/审核者/报告日期" line. Keep those anchors and
        dimensions, but redirect their image relationship to per-report uploads
        when supplied by the operator.
        """
        import os
        import shutil
        import tempfile
        from io import BytesIO
        from zipfile import ZIP_DEFLATED, ZipFile

        try:
            from lxml import etree
        except Exception as exc:
            self.logger.debug("缺少 lxml，跳过签名图片替换", error=str(exc))
            return

        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
            "ct": "http://schemas.openxmlformats.org/package/2006/content-types",
        }

        role_paths = {
            "detector": str(
                context.get("detector_signature_image_path")
                or context.get("issuer_signature_image_path")
                or ""
            ).strip(),
            "reviewer": str(context.get("reviewer_signature_image_path") or "").strip(),
        }
        role_paths = {
            role: path for role, path in role_paths.items() if path and Path(path).exists()
        }

        def paragraph_text(paragraph) -> str:
            return "".join(paragraph.xpath(".//w:t/text()", namespaces=ns))

        def drawing_offset(drawing, fallback: int) -> tuple[int, int]:
            offsets = drawing.xpath(".//wp:positionH/wp:posOffset/text()", namespaces=ns)
            try:
                return int(offsets[0]), fallback
            except Exception:
                return fallback, fallback

        def prepare_image(path: str) -> tuple[bytes, str, str]:
            src = Path(path)
            suffix = src.suffix.lower()
            if suffix == ".png":
                return src.read_bytes(), "png", "image/png"
            if suffix in {".jpg", ".jpeg"}:
                return src.read_bytes(), "jpg", "image/jpeg"
            try:
                from PIL import Image

                with Image.open(src) as image:
                    output = BytesIO()
                    image.convert("RGBA").save(output, format="PNG")
                return output.getvalue(), "png", "image/png"
            except Exception as exc:
                raise ValueError(f"签名图片格式不支持或无法转换: {src}") from exc

        def next_relationship_id(root) -> str:
            used = {
                rel.get("Id", "")
                for rel in root.xpath("./rel:Relationship", namespaces=ns)
            }
            max_num = 0
            for rid in used:
                if rid.startswith("rId") and rid[3:].isdigit():
                    max_num = max(max_num, int(rid[3:]))
            candidate = max_num + 1
            while f"rId{candidate}" in used:
                candidate += 1
            return f"rId{candidate}"

        def ensure_content_type(root, extension: str, content_type: str) -> None:
            existing = {
                elem.get("Extension")
                for elem in root.xpath("./ct:Default", namespaces=ns)
            }
            if extension in existing:
                return
            elem = etree.SubElement(
                root, f"{{{ns['ct']}}}Default", Extension=extension, ContentType=content_type
            )
            root.append(elem)

        with ZipFile(file_path, "r") as zin:
            document_xml = zin.read("word/document.xml")
            rels_xml = zin.read("word/_rels/document.xml.rels")
            content_types_xml = zin.read("[Content_Types].xml")
            document_info = zin.getinfo("word/document.xml")
            rels_info = zin.getinfo("word/_rels/document.xml.rels")
            content_types_info = zin.getinfo("[Content_Types].xml")
            entries = [(info, zin.read(info.filename)) for info in zin.infolist()]

        document_root = etree.fromstring(document_xml)
        rels_root = etree.fromstring(rels_xml)
        content_types_root = etree.fromstring(content_types_xml)

        paragraphs = document_root.xpath(".//w:body/w:p", namespaces=ns)
        label_index = next(
            (
                idx
                for idx, paragraph in enumerate(paragraphs)
                if "检测者" in paragraph_text(paragraph)
                and "审核者" in paragraph_text(paragraph)
            ),
            None,
        )
        if label_index is None:
            return

        candidates = paragraphs[max(0, label_index - 2) : label_index + 1]
        drawings: list[Any] = []
        for paragraph in candidates:
            drawings.extend(
                paragraph.xpath(".//wp:anchor | .//wp:inline", namespaces=ns)
            )
        drawing_blips = []
        for order, drawing in enumerate(drawings):
            blip = drawing.xpath(".//a:blip[@r:embed]", namespaces=ns)
            if blip:
                drawing_blips.append((drawing_offset(drawing, order), drawing, blip[0]))
        drawing_blips.sort(key=lambda item: item[0])
        if len(drawing_blips) < 2:
            return

        rels_by_target = {
            rel.get("Target"): rel
            for rel in rels_root.xpath("./rel:Relationship", namespaces=ns)
        }
        replacements: dict[str, bytes] = {}
        changed = False

        def remove_drawing(drawing) -> None:
            drawing_container = drawing.getparent()
            run = drawing_container.getparent() if drawing_container is not None else None
            if run is not None and run.tag.endswith("}r"):
                parent = run.getparent()
                if parent is not None:
                    parent.remove(run)
                    return
            if drawing_container is not None:
                parent = drawing_container.getparent()
                if parent is not None:
                    parent.remove(drawing_container)
                    return
            parent = drawing.getparent()
            if parent is not None:
                parent.remove(drawing)

        for role, (_, drawing, blip) in zip(("detector", "reviewer"), drawing_blips[:2]):
            image_path = role_paths.get(role)
            if not image_path:
                remove_drawing(drawing)
                changed = True
                continue
            image_bytes, extension, content_type = prepare_image(image_path)
            target = f"media/reportgen_signature_{role}.{extension}"
            rel = rels_by_target.get(target)
            if rel is None:
                rid = next_relationship_id(rels_root)
                rel = etree.SubElement(
                    rels_root,
                    f"{{{ns['rel']}}}Relationship",
                    Id=rid,
                    Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
                    Target=target,
                )
                rels_by_target[target] = rel
            else:
                rid = rel.get("Id")
            blip.set(f"{{{ns['r']}}}embed", rid)
            replacements[f"word/{target}"] = image_bytes
            ensure_content_type(content_types_root, extension, content_type)
            changed = True

        if not changed:
            return

        patched = {
            "word/document.xml": etree.tostring(
                document_root,
                xml_declaration=True,
                encoding="UTF-8",
                standalone="yes",
            ),
            "word/_rels/document.xml.rels": etree.tostring(
                rels_root,
                xml_declaration=True,
                encoding="UTF-8",
                standalone="yes",
            ),
            "[Content_Types].xml": etree.tostring(
                content_types_root,
                xml_declaration=True,
                encoding="UTF-8",
                standalone="yes",
            ),
        }
        patched.update(replacements)

        fd, tmp_name = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        written = set()
        try:
            with ZipFile(tmp_name, "w", compression=ZIP_DEFLATED) as zout:
                for info, data in entries:
                    if info.filename == "word/document.xml":
                        zout.writestr(document_info, patched[info.filename])
                    elif info.filename == "word/_rels/document.xml.rels":
                        zout.writestr(rels_info, patched[info.filename])
                    elif info.filename == "[Content_Types].xml":
                        zout.writestr(content_types_info, patched[info.filename])
                    elif info.filename in patched:
                        zout.writestr(info, patched[info.filename])
                    else:
                        zout.writestr(info, data)
                    written.add(info.filename)
                for name, data in patched.items():
                    if name not in written:
                        zout.writestr(name, data)
            shutil.move(tmp_name, file_path)
        finally:
            if os.path.exists(tmp_name):
                os.unlink(tmp_name)

        self.logger.debug(
            "已替换动态签名图片",
            detector=bool(role_paths.get("detector")),
            reviewer=bool(role_paths.get("reviewer")),
        )

    def _normalize_signature_layout(self, file_path: str, context: dict) -> None:
        """Stabilize the detector/reviewer signature block.

        The legacy template stores two handwritten signature images as floating
        anchors in the blank paragraph immediately before the text line:
        "检测者：...审核者：...报告日期：...". Different renderers position those
        anchors slightly differently, so they can overlap the report date. Keep
        the signatures on their own label line and move the date to a separate
        right-aligned paragraph, which is stable across Word/WPS/LibreOffice.
        """
        import os
        import re
        import shutil
        import tempfile
        import xml.etree.ElementTree as ET
        from zipfile import ZIP_DEFLATED, ZipFile

        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        for prefix, uri in {
            "w": ns_w,
            "wp": ns_wp,
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
        }.items():
            ET.register_namespace(prefix, uri)

        w_body = f"{{{ns_w}}}body"
        w_p = f"{{{ns_w}}}p"
        w_ppr = f"{{{ns_w}}}pPr"
        w_jc = f"{{{ns_w}}}jc"
        w_r = f"{{{ns_w}}}r"
        w_t = f"{{{ns_w}}}t"
        wp_anchor = f"{{{ns_wp}}}anchor"
        wp_position_h = f"{{{ns_wp}}}positionH"
        wp_position_v = f"{{{ns_wp}}}positionV"
        wp_pos_offset = f"{{{ns_wp}}}posOffset"

        def para_text(elem) -> str:
            return "".join((node.text or "") for node in elem.iter(w_t))

        def append_run(paragraph, text: str) -> None:
            run = ET.SubElement(paragraph, w_r)
            t = ET.SubElement(run, w_t)
            t.text = text
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

        def replace_para_text(paragraph, text: str) -> None:
            ppr = paragraph.find(w_ppr)
            for child in list(paragraph):
                if child is not ppr:
                    paragraph.remove(child)
            append_run(paragraph, text)

        def make_date_para(text: str) -> ET.Element:
            paragraph = ET.Element(w_p)
            ppr = ET.SubElement(paragraph, w_ppr)
            jc = ET.SubElement(ppr, w_jc)
            jc.set(f"{{{ns_w}}}val", "right")
            append_run(paragraph, text)
            return paragraph

        def set_anchor_offset(anchor, x: int | None = None, y: int | None = None) -> None:
            position_h = anchor.find(wp_position_h)
            if position_h is not None and x is not None:
                pos = position_h.find(wp_pos_offset)
                if pos is not None:
                    pos.text = str(x)
            position_v = anchor.find(wp_position_v)
            if position_v is not None and y is not None:
                pos = position_v.find(wp_pos_offset)
                if pos is not None:
                    pos.text = str(y)

        with ZipFile(file_path, "r") as zin:
            document_xml = zin.read("word/document.xml")
            document_info = zin.getinfo("word/document.xml")
            other_entries = [
                (info, zin.read(info.filename))
                for info in zin.infolist()
                if info.filename != "word/document.xml"
            ]

        root = ET.fromstring(document_xml)
        body = root.find(w_body)
        if body is None:
            return

        children = list(body)
        changed = False

        for idx, elem in enumerate(children):
            if elem.tag != w_p:
                continue
            text = para_text(elem)
            if "检测者" not in text or "审核者" not in text or "报告日期" not in text:
                continue
            if "__SIG_IMG__" in text:
                text = text.replace("__SIG_IMG__", "")

            date_match = re.search(r"报告日期[:：]\s*([0-9.\-/年月日]+)", text)
            report_date = date_match.group(1).strip() if date_match else ""
            if not report_date:
                report_date = str(context.get("report_date") or "").strip()
            report_date = re.sub(r"\b(\d{4})-(\d{2})-(\d{2})\b", r"\1.\2.\3", report_date)

            replace_para_text(elem, "检测者：                    审核者：")
            changed = True

            # Keep the two legacy floating signatures on the label line, but
            # away from the report date. Values are relative to the text column.
            if idx > 0:
                previous = children[idx - 1]
                anchors = sorted(
                    list(previous.iter(wp_anchor)),
                    key=lambda a: int(
                        (
                            a.find(wp_position_h).find(wp_pos_offset).text
                            if a.find(wp_position_h) is not None
                            and a.find(wp_position_h).find(wp_pos_offset) is not None
                            else "0"
                        )
                    ),
                )
                for anchor, x in zip(anchors[:2], (628650, 2200000)):
                    set_anchor_offset(anchor, x=x, y=120000)
                    anchor.set("allowOverlap", "0")
                    changed = True

            date_text = f"报告日期：{report_date}" if report_date else "报告日期："
            next_elem = children[idx + 1] if idx + 1 < len(children) else None
            if next_elem is not None and next_elem.tag == w_p and para_text(next_elem).startswith("报告日期："):
                replace_para_text(next_elem, date_text)
            else:
                body.insert(idx + 1, make_date_para(date_text))
            changed = True
            break

        if not changed:
            return

        fd, tmp_name = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            with ZipFile(tmp_name, "w", compression=ZIP_DEFLATED) as zout:
                zout.writestr(
                    document_info,
                    ET.tostring(root, encoding="utf-8", xml_declaration=True),
                )
                for info, data in other_entries:
                    zout.writestr(info, data)
            shutil.move(tmp_name, file_path)
        finally:
            if os.path.exists(tmp_name):
                os.unlink(tmp_name)

        self.logger.debug("已稳定签名区布局")

    def _fit_tables_to_page_width(self, file_path: str) -> None:
        """缩放超宽表格，使其不超过正文版心。"""
        from docx.oxml.ns import qn

        doc = Document(file_path)
        content_widths = []
        for section in doc.sections:
            try:
                content_widths.append(
                    int(
                        section.page_width.twips
                        - section.left_margin.twips
                        - section.right_margin.twips
                    )
                )
            except Exception:
                continue

        if not content_widths:
            return

        max_width = max(2000, min(content_widths))
        namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        scaled = 0

        for table in doc.tables:
            tbl = table._tbl
            tbl_pr = tbl.find(qn("w:tblPr"))
            if tbl_pr is None:
                continue

            tbl_w = tbl_pr.find(qn("w:tblW"))
            grid = tbl.find(qn("w:tblGrid"))
            grid_cols = [] if grid is None else grid.findall(qn("w:gridCol"))
            grid_width = sum(int(col.get(qn("w:w")) or 0) for col in grid_cols)
            table_width = int(tbl_w.get(qn("w:w")) or 0) if tbl_w is not None else 0
            current_width = max(grid_width, table_width)

            if current_width <= 0 or current_width <= max_width:
                continue

            ratio = max_width / current_width

            if tbl_w is not None:
                tbl_w.set(qn("w:type"), "dxa")
                tbl_w.set(qn("w:w"), str(max_width))

            for col in grid_cols:
                original = int(col.get(qn("w:w")) or 0)
                if original > 0:
                    col.set(qn("w:w"), str(max(240, int(original * ratio))))

            for tc_w in tbl.findall(".//w:tcPr/w:tcW", namespaces):
                if tc_w.get(qn("w:type")) not in (None, "dxa"):
                    continue
                original = int(tc_w.get(qn("w:w")) or 0)
                if original > 0:
                    tc_w.set(qn("w:type"), "dxa")
                    tc_w.set(qn("w:w"), str(max(240, int(original * ratio))))

            scaled += 1

        if scaled:
            doc.save(file_path)
            self.logger.debug("已压缩超宽表格", scaled=scaled, max_width=max_width)

    def _optimize_variant_table_layout(self, file_path: str) -> None:
        """Reflow the 2.1 variant table into a readable portrait-page layout."""
        # The reviewed report keeps the template's 9-column 2.1 table. Earlier
        # versions compacted it to four columns, which broke customer-approved
        # headings, font styling, and colors. Keep this hook as a no-op for
        # backward compatibility with older callers.
        return

        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        doc = Document(file_path)
        changed = False
        # Sum = 8300 twips, matching the A4 text width used by the report.
        # The original template has nine narrow columns; Word/LibreOffice wraps
        # short labels into nearly unreadable fragments. Keep all source fields,
        # but group the genomic details into one logical column.
        widths = [1100, 2600, 3000, 1600]
        headers = [
            "基因",
            "基因突变信息",
            "潜在获益靶向药物（证据等级）",
            "可能耐药或慎重药物（证据等级）",
        ]

        def clean_text(value: str) -> str:
            text = (value or "").strip()
            lines = [line.strip() for line in text.splitlines()]
            return "\n".join(line for line in lines if line)

        def normalize_dash(value: str) -> str:
            text = clean_text(value)
            return text if text and text not in {"-", "--", "—"} else "--"

        def keep_value(value: str) -> bool:
            text = clean_text(value)
            return bool(text and text not in {"-", "--", "—"})

        def set_cell_shading(cell, fill: str) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = tc_pr.find(qn("w:shd"))
            if shd is None:
                shd = OxmlElement("w:shd")
                tc_pr.append(shd)
            shd.set(qn("w:fill"), fill)

        def set_cell_borders(cell, color: str = "000000", size: str = "6") -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for side in ("top", "left", "bottom", "right"):
                border = borders.find(qn(f"w:{side}"))
                if border is None:
                    border = OxmlElement(f"w:{side}")
                    borders.append(border)
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), size)
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), color)

        def set_table_fixed_layout(table) -> None:
            tbl_pr = table._tbl.tblPr
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), "fixed")

            tbl_w = tbl_pr.find(qn("w:tblW"))
            if tbl_w is None:
                tbl_w = OxmlElement("w:tblW")
                tbl_pr.append(tbl_w)
            tbl_w.set(qn("w:type"), "dxa")
            tbl_w.set(qn("w:w"), str(sum(widths)))

            cell_mar = tbl_pr.find(qn("w:tblCellMar"))
            if cell_mar is None:
                cell_mar = OxmlElement("w:tblCellMar")
                tbl_pr.append(cell_mar)
            for side in ("top", "left", "bottom", "right"):
                margin = cell_mar.find(qn(f"w:{side}"))
                if margin is None:
                    margin = OxmlElement(f"w:{side}")
                    cell_mar.append(margin)
                margin.set(qn("w:w"), "70")
                margin.set(qn("w:type"), "dxa")

        def set_cell_width(cell, width: int) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))

        def set_grid(table) -> None:
            grid = table._tbl.tblGrid
            if grid is None:
                grid = OxmlElement("w:tblGrid")
                table._tbl.insert(0, grid)
            for child in list(grid):
                grid.remove(child)
            for width in widths:
                col = OxmlElement("w:gridCol")
                col.set(qn("w:w"), str(width))
                grid.append(col)

        def style_cell(cell, *, header: bool = False, gene: bool = False) -> None:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_borders(cell, color="000000", size="6")
            if header:
                set_cell_shading(cell, "00C4D8")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if header or gene
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    run.font.size = Pt(8.0 if header or gene else 7.4)
                    run.font.bold = bool(header or gene)
                    if header:
                        run.font.color.rgb = RGBColor(255, 255, 255)

        def mutation_info(cells) -> str:
            transcript = clean_text(cells[1].text)
            chrom = clean_text(cells[2].text)
            exon = clean_text(cells[3].text)
            site = clean_text(cells[4].text)
            mutation_type = clean_text(cells[5].text)
            frequency = clean_text(cells[6].text)

            parts = []
            if keep_value(site):
                parts.append(f"位点：{site}")
            if keep_value(transcript):
                parts.append(f"转录本：{transcript}")
            location = []
            if keep_value(chrom):
                location.append(f"染色体 {chrom}")
            if keep_value(exon):
                location.append(f"外显子 {exon}")
            if location:
                parts.append("；".join(location))
            if keep_value(mutation_type):
                parts.append(f"突变类型：{mutation_type}")
            if keep_value(frequency):
                freq = frequency if frequency.endswith("%") else f"{frequency}%"
                parts.append(f"频率：{freq}")
            return "\n".join(parts) if parts else "--"

        def replace_with_compact_table(table) -> None:
            new_table = doc.add_table(rows=1, cols=4)
            try:
                new_table.style = table.style
            except Exception:
                pass

            set_table_fixed_layout(new_table)
            set_grid(new_table)

            header_row = new_table.rows[0]
            header_row.height = None
            for idx, label in enumerate(headers):
                cell = header_row.cells[idx]
                cell.text = label
                set_cell_width(cell, widths[idx])
                style_cell(cell, header=True)
            tr_pr = header_row._tr.get_or_add_trPr()
            tbl_header = tr_pr.find(qn("w:tblHeader"))
            if tbl_header is None:
                tr_pr.append(OxmlElement("w:tblHeader"))

            for source_row in table.rows[2:]:
                cells = source_row.cells
                if len(cells) < 9:
                    continue
                gene = clean_text(cells[0].text)
                info = mutation_info(cells)
                benefit = normalize_dash(cells[7].text)
                caution = normalize_dash(cells[8].text)
                if not gene and info == "--" and benefit == "--" and caution == "--":
                    continue

                row = new_table.add_row()
                values = [gene or "--", info, benefit, caution]
                for idx, value in enumerate(values):
                    cell = row.cells[idx]
                    cell.text = value
                    set_cell_width(cell, widths[idx])
                    style_cell(cell, gene=(idx == 0))

            old_tbl = table._tbl
            old_tbl.addprevious(new_table._tbl)
            old_tbl.getparent().remove(old_tbl)

        for table in list(doc.tables):
            if len(table.columns) != 9 or not table.rows:
                continue
            header_text = " ".join(cell.text for cell in table.rows[0].cells)
            if "基因名称" not in header_text or "靶向药物信息" not in header_text:
                continue
            replace_with_compact_table(table)
            changed = True

        if changed:
            doc.save(file_path)
            self.logger.debug("已优化2.1变异表布局")

    def _normalize_final_section_layout(self, file_path: str) -> None:
        """统一所有 section 的版心，避免正文后半段页面宽度异常。"""
        doc = Document(file_path)
        sections = list(doc.sections)
        if len(sections) < 2:
            return

        text = "\n".join(p.text for p in doc.paragraphs)
        if "第三部分：基因变异及相应靶向/免疫药物解析" not in text:
            return

        reference = sections[0]

        reference_metrics = (
            reference.page_width.twips,
            reference.page_height.twips,
            reference.left_margin.twips,
            reference.right_margin.twips,
            reference.top_margin.twips,
            reference.bottom_margin.twips,
            reference.header_distance.twips,
            reference.footer_distance.twips,
        )
        changed = 0
        for target in sections[1:]:
            target_metrics = (
                target.page_width.twips,
                target.page_height.twips,
                target.left_margin.twips,
                target.right_margin.twips,
                target.top_margin.twips,
                target.bottom_margin.twips,
                target.header_distance.twips,
                target.footer_distance.twips,
            )
            if target_metrics == reference_metrics:
                continue

            target.page_width = reference.page_width
            target.page_height = reference.page_height
            target.left_margin = reference.left_margin
            target.right_margin = reference.right_margin
            target.top_margin = reference.top_margin
            target.bottom_margin = reference.bottom_margin
            target.header_distance = reference.header_distance
            target.footer_distance = reference.footer_distance
            changed += 1

        if changed:
            doc.save(file_path)
            self.logger.debug("已统一页面布局", sections=changed)

    def _restore_reviewed_body_headers(self, file_path: str) -> None:
        """Keep appendix/result sections on the reviewed report header scheme.

        LibreOffice/docxtpl can materialize later sections with blank headers. The
        reviewed template uses the body header (logo + patient name + slogan) and
        its attached watermark/footer rule on the gene-list and tail pages, so
        later sections should link back to the first section that already has
        that body header.
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        doc = Document(file_path)
        sections = list(doc.sections)
        if len(sections) < 3:
            return

        def header_text(section) -> str:
            parts = [p.text for p in section.header.paragraphs]
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts)

        source_idx = None
        for idx, section in enumerate(sections):
            text = header_text(section)
            if "姓名：" in text and "科技服务人类健康" in text:
                source_idx = idx
                break

        if source_idx is None:
            return

        def ensure_footer_top_rule(footer) -> bool:
            changed_footer = False
            targets = footer.paragraphs or []
            if not targets:
                return False

            # The reviewed body/footer style has a black rule above the page
            # number. Some render paths preserve the page number text box but
            # drop the rule, so restore the rule without depending on page text.
            paragraph = targets[0]
            ppr = paragraph._p.get_or_add_pPr()
            p_bdr = ppr.find(qn("w:pBdr"))
            if p_bdr is None:
                p_bdr = OxmlElement("w:pBdr")
                ppr.append(p_bdr)
                changed_footer = True
            top = p_bdr.find(qn("w:top"))
            if top is None:
                top = OxmlElement("w:top")
                p_bdr.append(top)
                changed_footer = True
            attrs = {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:space"): "0",
                qn("w:color"): "000000",
            }
            for key, value in attrs.items():
                if top.get(key) != value:
                    top.set(key, value)
                    changed_footer = True
            return changed_footer

        changed = 0
        for section in sections[source_idx:]:
            text = header_text(section)
            has_body_header = "科技服务人类健康" in text

            if not has_body_header and section.different_first_page_header_footer:
                section.different_first_page_header_footer = False
                changed += 1
            if not has_body_header:
                for hdr in (
                    section.header,
                    section.first_page_header,
                    section.even_page_header,
                ):
                    if not hdr.is_linked_to_previous:
                        hdr.is_linked_to_previous = True
                        changed += 1
            for footer in (
                section.footer,
                section.first_page_footer,
                section.even_page_footer,
            ):
                if ensure_footer_top_rule(footer):
                    changed += 1

        if changed:
            doc.save(file_path)
            self.logger.debug("已恢复后置章节页眉页脚", sections=changed)

    def _normalize_toc_decoration_layout(self, file_path: str) -> None:
        """Restore the reviewed TOC decorative vertical line coordinates.

        The template keeps the cyan TOC line as a floating shape anchored to the
        "目    录" paragraph. Word/WPS/LibreOffice can lay out refreshed TOC text
        slightly differently, so after static TOC rewriting we restore the
        reviewed report's line and circle offsets. We only touch the line and
        small circle anchored to the TOC title.
        """
        import os
        import re
        import shutil
        import tempfile
        import xml.etree.ElementTree as ET
        from zipfile import ZIP_DEFLATED, ZipFile

        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        ns_v = "urn:schemas-microsoft-com:vml"
        for prefix, uri in {
            "w": ns_w,
            "wp": ns_wp,
            "v": ns_v,
            "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
            "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
            "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
            "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
            "o": "urn:schemas-microsoft-com:office:office",
        }.items():
            ET.register_namespace(prefix, uri)

        w_p = f"{{{ns_w}}}p"
        w_t = f"{{{ns_w}}}t"
        wp_anchor = f"{{{ns_wp}}}anchor"
        wp_position_h = f"{{{ns_wp}}}positionH"
        wp_position_v = f"{{{ns_wp}}}positionV"
        wp_pos_offset = f"{{{ns_wp}}}posOffset"
        wp_extent = f"{{{ns_wp}}}extent"
        wp_doc_pr = f"{{{ns_wp}}}docPr"
        v_line = f"{{{ns_v}}}line"
        v_shape = f"{{{ns_v}}}shape"

        line_offset_emu = 862965
        circle_offset_emu = 828675
        line_offset_top_emu = 1119505
        circle_offset_top_emu = 1043305
        line_margin_left_pt = 67.95
        circle_margin_left_pt = 65.25
        line_margin_top_pt = 88.15
        circle_margin_top_pt = 82.15

        def para_text(elem) -> str:
            return "".join((node.text or "") for node in elem.iter(w_t))

        def is_toc_title(elem) -> bool:
            return "目录" in para_text(elem).replace(" ", "")

        def emu_value(elem, attr: str) -> int:
            try:
                return int(elem.get(attr) or "0")
            except ValueError:
                return 0

        def style_pt(style: str, key: str) -> float | None:
            match = re.search(rf"(?:^|;){re.escape(key)}:([-0-9.]+)pt", style)
            if not match:
                return None
            try:
                return float(match.group(1))
            except ValueError:
                return None

        def set_style_pt(style: str, key: str, value: float) -> str:
            formatted = f"{value:g}pt"
            pattern = rf"((?:^|;){re.escape(key)}:)[^;]+"
            if re.search(pattern, style):
                return re.sub(pattern, rf"\g<1>{formatted}", style, count=1)
            if style and not style.endswith(";"):
                style += ";"
            return f"{style}{key}:{formatted};"

        with ZipFile(file_path, "r") as zin:
            document_xml = zin.read("word/document.xml")
            document_info = zin.getinfo("word/document.xml")
            other_entries = [
                (info, zin.read(info.filename))
                for info in zin.infolist()
                if info.filename != "word/document.xml"
            ]

        root = ET.fromstring(document_xml)
        changed = 0

        for para in root.iter(w_p):
            if not is_toc_title(para):
                continue

            for anchor in para.iter(wp_anchor):
                doc_pr = anchor.find(wp_doc_pr)
                name = doc_pr.get("name", "") if doc_pr is not None else ""
                extent = anchor.find(wp_extent)
                if extent is None:
                    continue
                cx = emu_value(extent, f"{{{ns_wp}}}cx") or emu_value(extent, "cx")
                cy = emu_value(extent, f"{{{ns_wp}}}cy") or emu_value(extent, "cy")
                position_h = anchor.find(wp_position_h)
                pos_offset = position_h.find(wp_pos_offset) if position_h is not None else None
                position_v = anchor.find(wp_position_v)
                v_pos_offset = position_v.find(wp_pos_offset) if position_v is not None else None
                if pos_offset is None:
                    continue

                if "直接连接符" in name and cx <= 100000 and cy >= 2000000:
                    if pos_offset.text != str(line_offset_emu):
                        pos_offset.text = str(line_offset_emu)
                        changed += 1
                    if v_pos_offset is not None and v_pos_offset.text != str(line_offset_top_emu):
                        v_pos_offset.text = str(line_offset_top_emu)
                        changed += 1
                elif "椭圆" in name and cx <= 200000 and cy <= 200000:
                    if pos_offset.text != str(circle_offset_emu):
                        pos_offset.text = str(circle_offset_emu)
                        changed += 1
                    if v_pos_offset is not None and v_pos_offset.text != str(circle_offset_top_emu):
                        v_pos_offset.text = str(circle_offset_top_emu)
                        changed += 1

            for shape in list(para.iter(v_line)) + list(para.iter(v_shape)):
                style = shape.get("style") or ""
                height = style_pt(style, "height")
                width = style_pt(style, "width")
                if shape.tag == v_line and height and width and height >= 200 and width <= 5:
                    new_style = set_style_pt(style, "margin-left", line_margin_left_pt)
                    new_style = set_style_pt(new_style, "margin-top", line_margin_top_pt)
                elif (
                    shape.tag == v_shape
                    and height
                    and width
                    and 3 <= height <= 12
                    and 3 <= width <= 12
                ):
                    new_style = set_style_pt(style, "margin-left", circle_margin_left_pt)
                    new_style = set_style_pt(new_style, "margin-top", circle_margin_top_pt)
                else:
                    continue
                if new_style != style:
                    shape.set("style", new_style)
                    changed += 1

        if not changed:
            return

        fd, tmp_name = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            with ZipFile(tmp_name, "w", compression=ZIP_DEFLATED) as zout:
                zout.writestr(
                    document_info,
                    ET.tostring(root, encoding="utf-8", xml_declaration=True),
                )
                for info, data in other_entries:
                    zout.writestr(info, data)
            shutil.move(tmp_name, file_path)
        finally:
            if os.path.exists(tmp_name):
                os.unlink(tmp_name)

        self.logger.debug("已调整目录装饰线位置", changed=changed)

    def _populate_static_toc_page_numbers(
        self, file_path: str, context: dict | None = None
    ) -> None:
        """Write visible TOC page numbers from the final rendered PDF layout.

        LibreOffice can refresh the PAGEREF fields but may leave the tab run
        hidden, and in some environments it also preserves stale page numbers.
        Rendering the final docx to PDF gives us the authoritative pagination
        without requiring Microsoft Word UI permissions.
        """
        import shutil

        if not self._document_contains_toc_or_static_toc(file_path):
            self.logger.debug("文档不包含目录域/静态目录，跳过静态目录页码写回", output=file_path)
            return

        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        pdftotext = shutil.which("pdftotext")
        pdfinfo = shutil.which("pdfinfo")
        if not soffice or not pdftotext or not pdfinfo:
            self.logger.debug(
                "缺少 PDF 布局工具，跳过静态目录页码写回",
                soffice=bool(soffice),
                pdftotext=bool(pdftotext),
                pdfinfo=bool(pdfinfo),
            )
            return

        page_numbers = self._detect_toc_page_numbers_from_pdf_layout(
            file_path=file_path,
            soffice=soffice,
            pdftotext=pdftotext,
            pdfinfo=pdfinfo,
        )
        if not page_numbers:
            return

        stable_numbers = page_numbers
        for attempt in range(1, 5):
            if not self._write_static_toc_page_numbers(file_path, stable_numbers, context):
                return

            visible_numbers = self._read_static_toc_page_numbers(file_path)
            detected_numbers = self._detect_toc_page_numbers_from_pdf_layout(
                file_path=file_path,
                soffice=soffice,
                pdftotext=pdftotext,
                pdfinfo=pdfinfo,
            )
            if self._toc_page_numbers_match(visible_numbers, detected_numbers):
                self.logger.info(
                    "已按最终 PDF 版式写回目录页码",
                    output=file_path,
                    page_numbers=visible_numbers,
                    attempts=attempt,
                )
                return
            if not detected_numbers:
                break
            stable_numbers = detected_numbers

        self.logger.info(
            "已按最终 PDF 版式写回目录页码",
            output=file_path,
            page_numbers=stable_numbers,
            attempts="unstable",
        )

    def _read_static_toc_page_numbers(self, file_path: str) -> dict[str, int]:
        """Read visible static TOC numbers from DOCX XML."""
        import re
        from zipfile import ZipFile

        try:
            from lxml import etree
        except Exception:
            return {}

        labels = (
            "患者及样本信息",
            "检测内容",
            "检测结果小结",
            "靶向药物相关检测结果",
            "免疫治疗疗效评估",
            "检测结果说明",
            "基因变异解析",
            "靶向药物/免疫用药提示解析",
            "阅读说明",
            "常见问题解答",
            "结直肠癌诊疗知识",
            "癌症相关信号通路",
            "基因检测列表",
            "参考文献",
        )
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        try:
            with ZipFile(file_path, "r") as zf:
                root = etree.fromstring(zf.read("word/document.xml"))
        except Exception:
            return {}

        result: dict[str, int] = {}
        for paragraph in root.xpath(".//w:p", namespaces=ns):
            text = "".join(paragraph.xpath(".//w:t/text()", namespaces=ns))
            compact = re.sub(r"\s+", "", text or "")
            if not compact:
                continue
            match = re.search(r"(\d{1,3})$", compact)
            if not match:
                continue
            for label in labels:
                if label in compact and label not in result:
                    result[label] = int(match.group(1))
                    break
        return result

    def _toc_page_numbers_match(
        self,
        visible_numbers: dict[str, int],
        detected_numbers: dict[str, int],
    ) -> bool:
        """Return true when written TOC numbers match rendered target pages."""
        common = [
            label
            for label in detected_numbers
            if label in visible_numbers
        ]
        if not common:
            return False
        return all(visible_numbers[label] == detected_numbers[label] for label in common)

    def _write_static_toc_page_numbers(
        self,
        file_path: str,
        page_numbers: dict[str, int],
        context: dict | None = None,
    ) -> bool:
        """Write visible TOC labels/page numbers while keeping Word jump fields.

        Word/WPS/LibreOffice can disagree on PAGEREF pagination when a document
        contains floating text boxes and section-level page-number restarts.
        Keep the row clickable through ``HYPERLINK \\l`` but write the page
        number as static text derived from the final PDF layout. This prevents
        opening-time field refresh from changing the visible TOC number away
        from the page footer users see.
        """
        import copy
        import os
        import re
        import shutil
        import tempfile
        from zipfile import ZIP_DEFLATED, ZipFile

        try:
            from lxml import etree
        except Exception as exc:
            self.logger.debug("缺少 lxml，跳过静态目录页码写回", error=str(exc))
            return False

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        w_ns = ns["w"]

        def qn(tag: str) -> str:
            return f"{{{w_ns}}}{tag}"

        def normalize_label(text: str) -> str:
            text = re.sub(r"\s+", "", text or "")
            text = re.sub(r"\d+$", "", text)
            text = re.sub(r"^\d+[.．、]", "", text)
            return text

        style_cfg = self._panel_style_config(context, "toc")
        font_name = str(style_cfg.get("font_name") or "微软雅黑").strip()
        section_color = self._hex_color_config(
            style_cfg.get("section_font_color"), "00C4D8"
        )
        item_color = self._hex_color_config(style_cfg.get("item_font_color"), "000000")
        section_size = str(
            int(round(self._float_config(style_cfg.get("section_font_size"), 16.0) * 2))
        )
        item_size = str(
            int(round(self._float_config(style_cfg.get("item_font_size"), 11.0) * 2))
        )
        section_bold = self._bool_config(style_cfg.get("section_bold"), True)
        item_bold = self._bool_config(style_cfg.get("item_bold"), False)
        content_top_padding_twips = max(
            0,
            int(
                round(
                    self._float_config(
                        style_cfg.get("content_top_padding_pt"), 0.0
                    )
                    * 20
                )
            ),
        )

        src = Path(file_path)
        with ZipFile(src, "r") as zin:
            document_xml = zin.read("word/document.xml")
            root = etree.fromstring(document_xml)

            toc_sdt = None
            for sdt in root.xpath(".//w:sdt", namespaces=ns):
                instr = "".join(sdt.xpath(".//w:instrText/text()", namespaces=ns))
                text = "".join(sdt.xpath(".//w:t/text()", namespaces=ns))
                if "TOC" in instr or (
                    "第一部分" in text and "第四部分" in text and "参考文献" in text
                ):
                    toc_sdt = sdt
                    break
            if toc_sdt is None:
                return False

            normalized_numbers = {
                normalize_label(label): str(number)
                for label, number in page_numbers.items()
            }

            toc_groups = [
                (
                    "第一部分：基本信息",
                    ["患者及样本信息", "检测内容"],
                ),
                (
                    "第二部分：检测结果",
                    [
                        "检测结果小结",
                        "靶向药物相关检测结果",
                        "免疫治疗疗效评估",
                        "检测结果说明",
                    ],
                ),
                (
                    "第三部分：基因变异及相应靶向/免疫药物解析",
                    ["基因变异解析", "靶向药物/免疫用药提示解析", "阅读说明"],
                ),
                (
                    "第四部分：附录",
                    ["常见问题解答", "结直肠癌诊疗知识", "癌症相关信号通路", "基因检测列表", "参考文献"],
                ),
            ]

            toc_target_aliases = {
                "患者及样本信息": ["患者信息", "样本信息"],
                "靶向药物相关检测结果": ["靶向药物相关检测结果"],
                "靶向药物/免疫用药提示解析": ["靶向药物/免疫用药提示解析"],
                "阅读说明": ["阅读说明"],
            }
            toc_target_rules = {
                "检测结果小结": {"needles": ["1.检测结果小结"]},
                "靶向药物相关检测结果": {"needles": ["2.靶向药物相关检测结果"]},
                "免疫治疗疗效评估": {"needles": ["3.免疫治疗疗效评估"]},
                "检测结果说明": {
                    "needles": ["4.检测结果说明", "5.检测结果说明"],
                    "fallback": ["检测结果说明"],
                },
                "基因变异解析": {
                    "needles": ["1.基因变异解析"],
                    "fallback": ["基因变异解析"],
                },
                "靶向药物/免疫用药提示解析": {
                    "needles": ["2.靶向药物/免疫用药提示解析"],
                    "fallback": ["靶向药物/免疫用药提示解析"],
                },
                "阅读说明": {"needles": ["3.阅读说明"]},
                "常见问题解答": {"needles": ["1.常见问题解答"]},
                "结直肠癌诊疗知识": {"needles": ["2.结直肠癌诊疗知识"]},
                "癌症相关信号通路": {"needles": ["3.癌症相关信号通路"]},
                # Appendix titles also occur in explanatory prose. Keep these
                # anchors tied to the numbered appendix headings so WPS/Word
                # PAGEREF refreshes and TOC clicks cannot jump to earlier text.
                "基因检测列表": {
                    "needles": ["4.基因检测列表"],
                    "fallback": ["GeneListforMLseq", "基因检测列表"],
                    "prefer_last": True,
                },
                "参考文献": {
                    "needles": ["5.参考文献"],
                    "fallback": ["参考文献"],
                    "prefer_last": True,
                },
            }

            def paragraph_text(elem: Any) -> str:
                return "".join(elem.xpath(".//w:t/text()", namespaces=ns))

            def is_inside_toc(elem: Any) -> bool:
                current = elem
                while current is not None:
                    if current is toc_sdt:
                        return True
                    current = current.getparent()
                return False

            def existing_bookmark_names(elem: Any) -> list[str]:
                return [
                    bookmark.get(qn("name")) or ""
                    for bookmark in elem.xpath(".//w:bookmarkStart", namespaces=ns)
                    if bookmark.get(qn("name"))
                ]

            def preferred_bookmark(names: list[str]) -> str | None:
                for prefix in ("_Toc", "__RefHeading"):
                    for name in names:
                        if name.startswith(prefix):
                            return name
                return names[0] if names else None

            def compact_target(text: str) -> str:
                text = re.sub(r"\s+", "", text or "")
                return (
                    text.replace("．", ".")
                    .replace("。", ".")
                    .replace("、", ".")
                    .replace("：", ":")
                )

            existing_bookmark_ids = [
                int(value)
                for value in root.xpath(".//w:bookmarkStart/@w:id", namespaces=ns)
                if str(value).isdigit()
            ]
            next_bookmark_id = max(existing_bookmark_ids, default=0) + 1

            def make_auto_bookmark_name(label: str) -> str:
                safe = re.sub(r"[^A-Za-z0-9_]+", "_", label)
                safe = safe.strip("_")[:24] or "toc"
                return f"_TocAuto_{safe}_{next_bookmark_id}"

            def add_bookmark_to_paragraph(elem: Any, label: str) -> str:
                nonlocal next_bookmark_id
                bookmark_name = make_auto_bookmark_name(label)
                bookmark_id = str(next_bookmark_id)
                next_bookmark_id += 1
                start = etree.Element(qn("bookmarkStart"))
                start.set(qn("id"), bookmark_id)
                start.set(qn("name"), bookmark_name)
                end = etree.Element(qn("bookmarkEnd"))
                end.set(qn("id"), bookmark_id)

                insert_at = 0
                if len(elem) and elem[0].tag == qn("pPr"):
                    insert_at = 1
                elem.insert(insert_at, start)
                elem.append(end)
                return bookmark_name

            def find_or_create_target_bookmark(label: str) -> str | None:
                paragraph_entries: list[tuple[Any, str, str]] = []
                for paragraph in root.xpath(".//w:p", namespaces=ns):
                    if is_inside_toc(paragraph):
                        continue
                    text = paragraph_text(paragraph)
                    text_norm = normalize_label(text)
                    text_compact = compact_target(text)
                    if text_norm or text_compact:
                        paragraph_entries.append((paragraph, text_norm, text_compact))

                rule = toc_target_rules.get(label)
                if rule:
                    needles = [
                        compact_target(str(needle))
                        for needle in rule.get("needles", [])
                        if compact_target(str(needle))
                    ]
                    matches = [
                        paragraph
                        for paragraph, _text_norm, text_compact in paragraph_entries
                        if any(needle in text_compact for needle in needles)
                    ]
                    if not matches and rule.get("fallback"):
                        fallback_norms = [
                            normalize_label(str(needle))
                            for needle in rule.get("fallback", [])
                            if normalize_label(str(needle))
                        ]
                        matches = [
                            paragraph
                            for paragraph, text_norm, _text_compact in paragraph_entries
                            if any(needle in text_norm for needle in fallback_norms)
                        ]
                    if matches:
                        target = matches[-1] if rule.get("prefer_last") else matches[0]
                        names = existing_bookmark_names(target)
                        selected = preferred_bookmark(names)
                        if selected:
                            return selected
                        return add_bookmark_to_paragraph(target, label)

                candidate_labels = toc_target_aliases.get(label, [label])
                candidate_norms = [normalize_label(candidate) for candidate in candidate_labels]
                candidate_norms = [value for value in candidate_norms if value]

                fallback_para = None
                for paragraph, text_norm, _text_compact in paragraph_entries:
                    if not text_norm:
                        continue
                    if not any(candidate in text_norm for candidate in candidate_norms):
                        continue
                    names = existing_bookmark_names(paragraph)
                    selected = preferred_bookmark(names)
                    if selected:
                        return selected
                    if fallback_para is None:
                        fallback_para = paragraph

                if fallback_para is not None:
                    return add_bookmark_to_paragraph(fallback_para, label)
                return None

            def make_run(
                text: str | None = None,
                *,
                tab: bool = False,
                section: bool = False,
            ) -> Any:
                run = etree.Element(qn("r"))
                r_pr = etree.SubElement(run, qn("rPr"))
                fonts = etree.SubElement(r_pr, qn("rFonts"))
                fonts.set(qn("ascii"), font_name)
                fonts.set(qn("hAnsi"), font_name)
                fonts.set(qn("eastAsia"), font_name)
                fonts.set(qn("cs"), "Times New Roman")
                if section and section_bold or (not section and item_bold):
                    etree.SubElement(r_pr, qn("b"))
                    etree.SubElement(r_pr, qn("bCs"))
                color = etree.SubElement(r_pr, qn("color"))
                color.set(qn("val"), section_color if section else item_color)
                size = etree.SubElement(r_pr, qn("sz"))
                size.set(qn("val"), section_size if section else item_size)
                size_cs = etree.SubElement(r_pr, qn("szCs"))
                size_cs.set(qn("val"), section_size if section else item_size)
                if tab:
                    etree.SubElement(run, qn("tab"))
                else:
                    t = etree.SubElement(run, qn("t"))
                    t.text = text or ""
                return run

            def make_field_run(
                *,
                field_char_type: str | None = None,
                instruction: str | None = None,
            ) -> Any:
                run = etree.Element(qn("r"))
                if field_char_type:
                    field_char = etree.SubElement(run, qn("fldChar"))
                    field_char.set(qn("fldCharType"), field_char_type)
                if instruction is not None:
                    instr = etree.SubElement(run, qn("instrText"))
                    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    instr.text = instruction
                return run

            def make_toc_ppr(section: bool, before_twips: int = 0) -> Any:
                p_pr = etree.Element(qn("pPr"))
                tabs = etree.SubElement(p_pr, qn("tabs"))
                tab = etree.SubElement(tabs, qn("tab"))
                tab.set(qn("val"), "right")
                tab.set(qn("pos"), "7700")
                etree.SubElement(p_pr, qn("adjustRightInd")).set(qn("val"), "0")
                etree.SubElement(p_pr, qn("snapToGrid")).set(qn("val"), "0")
                spacing = etree.SubElement(p_pr, qn("spacing"))
                if before_twips:
                    spacing.set(qn("before"), str(before_twips))
                spacing.set(qn("after"), "0")
                spacing.set(qn("line"), "312")
                spacing.set(qn("lineRule"), "auto")
                indent = etree.SubElement(p_pr, qn("ind"))
                indent.set(qn("left"), "1980")
                indent.set(qn("leftChars"), "900")
                r_pr = etree.SubElement(p_pr, qn("rPr"))
                fonts = etree.SubElement(r_pr, qn("rFonts"))
                fonts.set(qn("ascii"), "Tahoma")
                fonts.set(qn("hAnsi"), "Tahoma")
                fonts.set(qn("eastAsia"), font_name)
                fonts.set(qn("cs"), "Times New Roman")
                if section and section_bold:
                    etree.SubElement(r_pr, qn("b"))
                    etree.SubElement(r_pr, qn("bCs"))
                    color = etree.SubElement(r_pr, qn("color"))
                    color.set(qn("val"), section_color)
                elif not section and item_color:
                    color = etree.SubElement(r_pr, qn("color"))
                    color.set(qn("val"), item_color)
                size = etree.SubElement(r_pr, qn("sz"))
                size.set(qn("val"), section_size if section else item_size)
                size_cs = etree.SubElement(r_pr, qn("szCs"))
                size_cs.set(qn("val"), section_size if section else item_size)
                lang = etree.SubElement(r_pr, qn("lang"))
                lang.set(qn("val"), "en-US")
                lang.set(qn("eastAsia"), "zh-CN")
                lang.set(qn("bidi"), "ar-SA")
                return p_pr

            def make_toc_paragraph(
                label: str,
                number: str | None,
                *,
                section: bool,
                before_twips: int = 0,
            ) -> Any:
                para = etree.Element(qn("p"))
                para.append(make_toc_ppr(section, before_twips=before_twips))
                anchor = find_or_create_target_bookmark(label)
                if anchor:
                    para.append(make_field_run(field_char_type="begin"))
                    para.append(make_field_run(instruction=f' HYPERLINK \\l "{anchor}" '))
                    para.append(make_field_run(field_char_type="separate"))
                para.append(make_run(label, section=section))
                if number is not None:
                    para.append(make_run(tab=True))
                    para.append(make_run(number))
                if anchor:
                    para.append(make_field_run(field_char_type="end"))
                return para

            def make_page_break_paragraph() -> Any:
                para = etree.Element(qn("p"))
                run = etree.SubElement(para, qn("r"))
                br = etree.SubElement(run, qn("br"))
                br.set(qn("type"), "page")
                return para

            toc_content = toc_sdt.find(qn("sdtContent"))
            if toc_content is None:
                return False
            reviewed_section_break = None
            for child in toc_content:
                if child.xpath(".//w:sectPr", namespaces=ns):
                    reviewed_section_break = etree.Element(qn("p"))
                    p_pr = child.find(qn("pPr"))
                    if p_pr is not None:
                        reviewed_section_break.append(copy.deepcopy(p_pr))
                    break
            for child in list(toc_content):
                toc_content.remove(child)

            static_count = 0
            first_visible_toc_entry = True
            for section_label, item_labels in toc_groups:
                present_items = [
                    label
                    for label in item_labels
                    if normalize_label(label) in normalized_numbers
                ]
                if not present_items:
                    continue
                toc_content.append(
                    make_toc_paragraph(
                        section_label,
                        None,
                        section=True,
                        before_twips=(
                            content_top_padding_twips
                            if first_visible_toc_entry
                            else 0
                        ),
                    )
                )
                first_visible_toc_entry = False
                static_count += 1
                for label in present_items:
                    number = normalized_numbers[normalize_label(label)]
                    toc_content.append(make_toc_paragraph(label, number, section=False))
                    static_count += 1

            if not static_count:
                return False

            # The reviewed Word TOC keeps its next-section boundary inside the
            # SDT content. Rebuilding the SDT without restoring it lets
            # "第一部分：基本信息" flow onto the TOC page, crosses the vertical
            # decoration line, and makes footer page numbers count the TOC
            # itself. Preserve that boundary instead of replacing it with a
            # plain page break; only fall back to a page break for synthetic
            # test templates that do not carry section properties.
            if reviewed_section_break is not None:
                toc_content.append(reviewed_section_break)
            else:
                toc_content.append(make_page_break_paragraph())

            patched_xml = etree.tostring(
                root,
                xml_declaration=True,
                encoding="UTF-8",
                standalone="yes",
            )

            fd, tmp_name = tempfile.mkstemp(
                suffix=".docx",
                prefix=f"{src.stem}_toc_",
                dir=str(src.parent),
            )
            os.close(fd)
            tmp_path = Path(tmp_name)
            try:
                with ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
                    for item in zin.infolist():
                        if item.filename == "word/document.xml":
                            data = patched_xml
                        else:
                            data = zin.read(item.filename)
                        zout.writestr(item, data)
                shutil.move(str(tmp_path), str(src))
                self._set_update_fields(str(src))
            finally:
                if tmp_path.exists():
                    tmp_path.unlink()

        return True

    def _remove_update_fields_setting(self, settings_xml: bytes) -> bytes:
        """Disable automatic field refresh after static TOC page numbers are written."""
        try:
            from lxml import etree
        except Exception:
            return settings_xml

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        root = etree.fromstring(settings_xml)
        changed = False
        for elem in root.xpath(".//w:updateFields", namespaces=ns):
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
                changed = True
        if not changed:
            return settings_xml
        return etree.tostring(
            root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone="yes",
        )

    def _detect_toc_page_numbers_from_pdf_layout(
        self,
        *,
        file_path: str,
        soffice: str,
        pdftotext: str,
        pdfinfo: str,
    ) -> dict[str, int]:
        """Render DOCX to PDF and infer report-page numbers for TOC entries."""
        import re
        import shutil
        import subprocess
        import tempfile

        toc_entries: list[tuple[str, tuple[str, ...]]] = [
            ("患者及样本信息", ("第一部分：基本信息患者信息", "患者信息样本信息")),
            ("检测内容", ("第一部分：基本信息患者信息", "检测内容")),
            ("检测结果小结", ("1.检测结果小结",)),
            ("靶向药物相关检测结果", ("2.靶向药物相关检测结果",)),
            ("免疫治疗疗效评估", ("3.免疫治疗疗效评估",)),
            ("临床常用化疗药物评估及解析", ("4.临床常用化疗药物评估及解析",)),
            ("检测结果说明", ("4.检测结果说明", "5.检测结果说明", "检测结果说明")),
            ("基因变异解析", ("基因变异说明", "5.检测结果说明")),
            ("靶向药物/免疫用药提示解析", ("靶向药物/免疫用药提示解析",)),
            ("阅读说明", ("3.阅读说明",)),
            ("常见问题解答", ("1.常见问题解答",)),
            ("结直肠癌诊疗知识", ("2.结直肠癌诊疗知识",)),
            ("癌症相关信号通路", ("3.癌症相关信号通路",)),
            ("基因检测列表", ("4.基因检测列表", "GeneListforMLseq")),
            ("参考文献", ("5.参考文献",)),
        ]

        def normalize(text: str) -> str:
            return re.sub(r"\s+", "", text or "")

        with tempfile.TemporaryDirectory(prefix="reportgen_pdf_") as tmp_dir:
            tmp_dir_path = Path(tmp_dir)
            input_dir = tmp_dir_path / "input"
            output_dir = tmp_dir_path / "output"
            input_dir.mkdir(parents=True, exist_ok=True)
            output_dir.mkdir(parents=True, exist_ok=True)
            tmp_input = input_dir / "input.docx"
            shutil.copy2(file_path, tmp_input)

            cmd = [
                soffice,
                "--headless",
                "--nologo",
                "--nolockcheck",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(tmp_input),
            ]
            try:
                subprocess.run(
                    cmd,
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    timeout=180,
                )
            except Exception as exc:
                self.logger.debug("目录页码 PDF 渲染失败，跳过静态写回", error=str(exc))
                return {}

            pdf_path = output_dir / "input.pdf"
            if not pdf_path.exists():
                return {}

            try:
                info = subprocess.run(
                    [pdfinfo, str(pdf_path)],
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    timeout=30,
                ).stdout
                match = re.search(r"^Pages:\s+(\d+)", info, flags=re.MULTILINE)
                page_count = int(match.group(1)) if match else 0
            except Exception:
                page_count = 0
            if page_count <= 0:
                return {}

            page_texts: dict[int, str] = {}
            normalized_pages: dict[int, str] = {}
            for page in range(1, page_count + 1):
                try:
                    text = subprocess.run(
                        [
                            pdftotext,
                            "-layout",
                            "-f",
                            str(page),
                            "-l",
                            str(page),
                            str(pdf_path),
                            "-",
                        ],
                        check=True,
                        stdout=subprocess.PIPE,
                        stderr=subprocess.PIPE,
                        text=True,
                        timeout=30,
                    ).stdout
                except Exception:
                    text = ""
                page_texts[page] = text
                normalized_pages[page] = normalize(text)

        report_page_numbers = {
            page: number
            for page, text in page_texts.items()
            if (number := self._extract_pdf_footer_page_number(text)) is not None
        }

        toc_page = next(
            (
                page for page, text in normalized_pages.items()
                if "目录" in text and "第一部分" in text and "参考文献" in text
            ),
            None,
        )
        if toc_page is None:
            return {}

        content_start = next(
            (
                page for page in range(toc_page + 1, page_count + 1)
                if "患者信息" in normalized_pages.get(page, "")
                and "检测内容" in normalized_pages.get(page, "")
            ),
            None,
        )
        if content_start is None:
            content_start = toc_page + 1

        page_numbers: dict[str, int] = {}
        for label, needles in toc_entries:
            found_page = None
            for page in range(content_start, page_count + 1):
                page_text = normalized_pages.get(page, "")
                if any(normalize(needle) in page_text for needle in needles):
                    found_page = page
                    break
            if found_page is not None:
                page_numbers[label] = report_page_numbers.get(
                    found_page, found_page - content_start + 1
                )

        return page_numbers

    @staticmethod
    def _extract_pdf_footer_page_number(page_text: str) -> int | None:
        """Return the visible report-page footer number from one PDF text page."""
        import re

        lines = [line.strip() for line in (page_text or "").splitlines() if line.strip()]
        for text in reversed(lines[-12:]):
            if re.fullmatch(r"\d{1,3}", text):
                return int(text)
        return None

    def _compact_gene_list_tables(self, file_path: str, context: dict | None = None) -> None:
        """Align static gene-list tables with the reviewed report layout."""
        from docx.enum.table import WD_ROW_HEIGHT_RULE
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt

        doc = Document(file_path)
        changed = 0
        content_cfg = self._report_content_config(context)
        style_cfg = content_cfg.get("gene_list_table_style")
        style_cfg = style_cfg if isinstance(style_cfg, dict) else {}
        row_height_cm = self._float_config(style_cfg.get("row_height_cm"), 0.88)
        header_font_size = self._float_config(style_cfg.get("header_font_size"), 14.0)
        body_font_size = self._float_config(style_cfg.get("body_font_size"), 10.5)

        for table in doc.tables:
            if not table.rows:
                continue
            header_text = " ".join(cell.text for cell in table.rows[0].cells)
            if "Gene List for MLseq" not in header_text and "基因检测列表" not in header_text:
                continue

            for row_idx, row in enumerate(table.rows):
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                row.height = Cm(row_height_cm)
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.line_spacing = 1.0
                        for run in paragraph.runs:
                            run.font.size = Pt(body_font_size if row_idx else header_font_size)
                            run.font.underline = False
                changed += 1

        if changed:
            doc.save(file_path)
            self.logger.debug("已对齐基因检测列表样式", rows=changed)

    def _normalize_quality_control_tables(self, file_path: str) -> None:
        """Prevent template-inherited underlines in the quality-control table."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        doc = Document(file_path)
        changed = 0
        for table in doc.tables:
            if not table.rows:
                continue
            header_text = " ".join(cell.text.strip() for cell in table.rows[0].cells)
            if "质控项" not in header_text or "质控结果" not in header_text:
                continue
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        for run in paragraph.runs:
                            run.font.underline = False
                            changed += 1

        if changed:
            doc.save(file_path)
            self.logger.debug("已修复质控表下划线继承", runs=changed)

    def _remove_trailing_hla_table(self, file_path: str) -> None:
        """移除文档尾部孤立的 HLA 表格。"""
        doc = Document(file_path)
        if not doc.tables:
            return

        last_table = doc.tables[-1]
        first_row_text = " ".join(cell.text.strip() for cell in last_table.rows[0].cells) if last_table.rows else ""
        if "HLA位点" not in first_row_text:
            return

        last_table._tbl.getparent().remove(last_table._tbl)
        doc.save(file_path)
        self.logger.debug("已移除文档尾部HLA表格")

    def _set_update_fields(self, file_path: str) -> None:
        """设置 docx 的 updateFields 属性，让 Word 打开时自动刷新目录/页码域。"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document(file_path)
        settings = doc.settings.element

        # 添加 <w:updateFields w:val="true"/>
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")

        doc.save(file_path)
        self.logger.debug("已设置 updateFields=true，Word 打开时将刷新目录")

    def _refresh_fields_with_native_engine(self, file_path: str) -> None:
        """优先使用可用的原生排版引擎刷新目录/页码域。"""
        import os
        import sys

        if not self._document_contains_toc(file_path):
            self.logger.debug("文档不包含 TOC 域，跳过目录刷新", output=file_path)
            return

        if sys.platform.startswith("linux"):
            self._refresh_fields_with_libreoffice(file_path)
            return

        if sys.platform == "darwin":
            use_word = str(os.environ.get("REPORTGEN_REFRESH_WITH_WORD") or "").strip().lower()
            if use_word in {"1", "true", "yes", "y", "on"}:
                try:
                    self._refresh_fields_with_word_macos(file_path)
                    return
                except Exception as exc:
                    self.logger.debug(
                        "Word 刷新目录失败，尝试回退",
                        output=file_path,
                        error=str(exc),
                    )
            # Word opened via AppleScript often triggers macOS sandbox prompts
            # for files on external volumes. LibreOffice is the non-interactive
            # default; Word refresh remains opt-in through REPORTGEN_REFRESH_WITH_WORD.
            self._refresh_fields_with_libreoffice(file_path)
            return

        self._refresh_fields_with_libreoffice(file_path)

    def _document_contains_toc(self, file_path: str) -> bool:
        """检查 docx 是否包含 TOC 域。"""
        from zipfile import ZipFile

        try:
            with ZipFile(file_path, "r") as zf:
                document_xml = zf.read("word/document.xml").decode("utf-8", "ignore")
        except Exception as exc:
            raise RuntimeError(f"读取文档目录域失败: {exc}") from exc

        return "TOC" in document_xml

    def _document_contains_toc_or_static_toc(self, file_path: str) -> bool:
        """Check whether a docx has a Word TOC field or our static TOC rows."""
        from zipfile import ZipFile

        try:
            with ZipFile(file_path, "r") as zf:
                document_xml = zf.read("word/document.xml").decode("utf-8", "ignore")
        except Exception as exc:
            raise RuntimeError(f"读取文档目录域失败: {exc}") from exc

        if "TOC" in document_xml:
            return True
        return (
            "HYPERLINK" in document_xml
            and "第一部分" in document_xml
            and "参考文献" in document_xml
        )

    def _refresh_fields_with_word_macos(self, file_path: str) -> None:
        """使用 macOS 上的 Microsoft Word 最佳努力刷新目录和页码域。"""
        import shutil
        import subprocess

        osascript = shutil.which("osascript")
        word_app = Path("/Applications/Microsoft Word.app")
        if not osascript or not word_app.exists():
            raise RuntimeError("未找到 Microsoft Word 或 osascript")

        script_lines = [
            'tell application "Microsoft Word"',
            "set display alerts to alerts none",
            "set wasRunning to running",
            "set preDocCount to count of documents",
            "set visible to false",
            f'open file name "{file_path}"',
            "delay 4",
            "try",
            "repeat with tocRef in (tables of contents of document 1)",
            "update tocRef",
            "update page numbers tocRef",
            "end repeat",
            "end try",
            "try",
            "save document 1",
            "end try",
            "delay 1",
            "set remainingDocs to 0",
            "try",
            "set remainingDocs to count of documents",
            "end try",
            "if (wasRunning is false) and (preDocCount = 0) and (remainingDocs = 0) then",
            "try",
            "quit saving no",
            "end try",
            "end if",
            "return remainingDocs",
            "end tell",
        ]
        command = [osascript, "-l", "AppleScript"]
        for line in script_lines:
            command.extend(["-e", line])

        try:
            result = subprocess.run(
                command,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                timeout=self.WORD_REFRESH_TIMEOUT_SECONDS,
                check=True,
            )
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError(
                "Microsoft Word 刷新目录超时"
                f"（{self.WORD_REFRESH_TIMEOUT_SECONDS}秒），"
                "可能缺少 Apple Events 授权或 Word 首次启动未完成"
            ) from exc
        self.logger.debug(
            "已执行 Microsoft Word 目录刷新流程",
            output=file_path,
            log=result.stdout.strip(),
        )

    def _refresh_fields_with_libreoffice(self, file_path: str) -> None:
        """使用 LibreOffice 最佳努力刷新目录/页码域。

        仅在系统存在 `soffice`/`libreoffice` 且文档中包含 TOC 域时执行。
        失败时抛异常，由上层降级回 `updateFields=true`。
        """
        import shutil

        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            self.logger.debug("未找到 LibreOffice 二进制", output=file_path)
            raise RuntimeError("未找到 LibreOffice，目录页码仅保留 updateFields 兜底")

        if not self._document_contains_toc(file_path):
            self.logger.debug("文档不包含 TOC 域，跳过 LibreOffice 目录刷新", output=file_path)
            return

        try:
            self._refresh_fields_with_libreoffice_uno(file_path, soffice)
            return
        except Exception as exc:
            self.logger.debug(
                "LibreOffice UNO 刷新失败，回退 convert-to 路径",
                output=file_path,
                error=str(exc),
            )

        self._refresh_fields_with_libreoffice_convert(file_path, soffice)

    def _refresh_fields_with_libreoffice_uno(self, file_path: str, soffice: str) -> None:
        """通过 LibreOffice UNO 显式更新目录索引并重新保存 docx。

        如果环境变量 ``REPORTGEN_LO_LISTENER_PORT`` 已设(由 web 端 lifespan
        启动的常驻 LibreOffice listener 注入),则跳过每次的 ``soffice`` 冷启
        和 socket port 探测,直接连那个常驻端口 —— 每次 refresh 大约省 5–8 秒。
        """
        import contextlib
        import os
        import shutil
        import socket
        import subprocess
        import tempfile
        import textwrap

        python_candidates = [Path("/usr/bin/python3"), Path(shutil.which("python3") or "")]
        uno_python = None
        for candidate in python_candidates:
            if not candidate or not str(candidate):
                continue
            if not candidate.exists():
                continue
            probe = subprocess.run(
                [str(candidate), "-c", "import uno"],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            )
            if probe.returncode == 0:
                uno_python = str(candidate)
                break
        if not uno_python:
            raise RuntimeError("未找到支持 python3-uno 的系统 Python")

        persistent_port_raw = os.environ.get("REPORTGEN_LO_LISTENER_PORT", "")
        persistent_port = int(persistent_port_raw) if persistent_port_raw.isdigit() else None

        with tempfile.TemporaryDirectory(prefix="reportgen_lo_") as tmp_dir, contextlib.ExitStack() as stack:
            input_dir = Path(tmp_dir) / "input"
            output_dir = Path(tmp_dir) / "output"
            input_dir.mkdir(parents=True, exist_ok=True)
            output_dir.mkdir(parents=True, exist_ok=True)

            tmp_input = input_dir / "input.docx"
            refreshed = output_dir / "refreshed.docx"
            shutil.copy2(file_path, tmp_input)

            if persistent_port is not None:
                port = persistent_port
                profile_dir = None
            else:
                profile_dir = stack.enter_context(
                    tempfile.TemporaryDirectory(prefix="reportgen_lo_profile_")
                )
                probe_sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                probe_sock.bind(("127.0.0.1", 0))
                port = probe_sock.getsockname()[1]
                probe_sock.close()

            uno_script = Path(tmp_dir) / "refresh_with_uno.py"
            uno_script.write_text(
                textwrap.dedent(
                    """
                    import sys
                    import time
                    import uno
                    from com.sun.star.beans import PropertyValue

                    def prop(name, value):
                        item = PropertyValue()
                        item.Name = name
                        item.Value = value
                        return item

                    input_path, output_path, port = sys.argv[1], sys.argv[2], sys.argv[3]
                    local_ctx = uno.getComponentContext()
                    resolver = local_ctx.ServiceManager.createInstanceWithContext(
                        "com.sun.star.bridge.UnoUrlResolver",
                        local_ctx,
                    )

                    ctx = None
                    last_error = None
                    deadline = time.time() + 30
                    while time.time() < deadline:
                        try:
                            ctx = resolver.resolve(
                                f"uno:socket,host=127.0.0.1,port={port};urp;StarOffice.ComponentContext"
                            )
                            break
                        except Exception as exc:
                            last_error = exc
                            time.sleep(0.5)

                    if ctx is None:
                        raise RuntimeError(f"UNO 连接失败: {last_error}")

                    desktop = ctx.ServiceManager.createInstanceWithContext(
                        "com.sun.star.frame.Desktop",
                        ctx,
                    )
                    load_props = (
                        prop("Hidden", True),
                        prop("ReadOnly", False),
                        prop("UpdateDocMode", 1),
                    )
                    doc = desktop.loadComponentFromURL(
                        uno.systemPathToFileUrl(input_path),
                        "_blank",
                        0,
                        load_props,
                    )
                    try:
                        doc.refresh()
                    except Exception:
                        pass
                    try:
                        doc.TextFields.refresh()
                    except Exception:
                        pass
                    try:
                        indexes = doc.getDocumentIndexes()
                        for idx in range(indexes.getCount()):
                            indexes.getByIndex(idx).update()
                    except Exception:
                        pass
                    time.sleep(1)
                    doc.storeAsURL(
                        uno.systemPathToFileUrl(output_path),
                        (prop("FilterName", "Office Open XML Text"),),
                    )
                    doc.close(True)
                    """
                ).strip(),
                encoding="utf-8",
            )

            listener = None
            if persistent_port is None:
                listener_cmd = [
                    soffice,
                    f"-env:UserInstallation=file://{Path(profile_dir).as_posix()}",
                    "--headless",
                    "--nologo",
                    "--nolockcheck",
                    "--nodefault",
                    f"--accept=socket,host=127.0.0.1,port={port};urp;StarOffice.ComponentContext",
                ]
                listener = subprocess.Popen(
                    listener_cmd,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    text=True,
                )
            try:
                result = subprocess.run(
                    [uno_python, str(uno_script), str(tmp_input), str(refreshed), str(port)],
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    timeout=180,
                )
            except subprocess.TimeoutExpired as exc:
                raise RuntimeError("LibreOffice 刷新目录超时") from exc
            except subprocess.CalledProcessError as exc:
                log = "\n".join(part for part in [(exc.stdout or "").strip(), (exc.stderr or "").strip()] if part)
                raise RuntimeError(
                    f"LibreOffice UNO 刷新失败: {log or exc}"
                ) from exc
            finally:
                if listener is not None:
                    listener.terminate()
                    try:
                        listener.wait(timeout=10)
                    except subprocess.TimeoutExpired:
                        listener.kill()
                        listener.wait(timeout=5)

            if not refreshed.exists():
                log = "\n".join(
                    part for part in [(result.stdout or "").strip(), (result.stderr or "").strip()] if part
                )
                raise RuntimeError(
                    f"LibreOffice UNO 未生成刷新后的文档: {log}"
                )

            try:
                Document(refreshed)
            except Exception as exc:
                raise RuntimeError(f"LibreOffice UNO 生成的文档不可读: {exc}") from exc

            shutil.copy2(refreshed, file_path)
            log = "\n".join(part for part in [(result.stdout or "").strip(), (result.stderr or "").strip()] if part)
            self.logger.info(
                "已使用 LibreOffice 刷新目录/页码",
                output=file_path,
                engine=Path(soffice).name,
                log=log,
            )

    def _refresh_fields_with_libreoffice_convert(self, file_path: str, soffice: str) -> None:
        """回退路径：通过 LibreOffice convert-to 重写 docx。"""
        import shutil
        import subprocess
        import tempfile

        with tempfile.TemporaryDirectory(prefix="reportgen_lo_") as tmp_dir, tempfile.TemporaryDirectory(
            prefix="reportgen_lo_profile_"
        ) as profile_dir:
            input_dir = Path(tmp_dir) / "input"
            output_dir = Path(tmp_dir) / "output"
            input_dir.mkdir(parents=True, exist_ok=True)
            output_dir.mkdir(parents=True, exist_ok=True)

            tmp_input = input_dir / "input.docx"
            shutil.copy2(file_path, tmp_input)

            cmd = [
                soffice,
                f"-env:UserInstallation=file://{Path(profile_dir).as_posix()}",
                "--headless",
                "--nologo",
                "--nolockcheck",
                "--nodefault",
                "--convert-to",
                'docx:Office Open XML Text',
                "--outdir",
                str(output_dir),
                str(tmp_input),
            ]
            try:
                result = subprocess.run(
                    cmd,
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    timeout=180,
                )
            except subprocess.TimeoutExpired as exc:
                raise RuntimeError("LibreOffice convert-to 刷新目录超时") from exc
            except subprocess.CalledProcessError as exc:
                log = "\n".join(part for part in [(exc.stdout or "").strip(), (exc.stderr or "").strip()] if part)
                raise RuntimeError(
                    f"LibreOffice convert-to 刷新目录失败: {log or exc}"
                ) from exc

            refreshed = output_dir / "input.docx"
            if not refreshed.exists():
                log = "\n".join(
                    part for part in [(result.stdout or "").strip(), (result.stderr or "").strip()] if part
                )
                raise RuntimeError(
                    f"LibreOffice convert-to 未生成刷新后的文档: {log}"
                )

            try:
                Document(refreshed)
            except Exception as exc:
                raise RuntimeError(f"LibreOffice convert-to 生成的文档不可读: {exc}") from exc

            shutil.copy2(refreshed, file_path)
            log = "\n".join(part for part in [(result.stdout or "").strip(), (result.stderr or "").strip()] if part)
            self.logger.info(
                "已使用 LibreOffice 刷新目录/页码",
                output=file_path,
                engine=Path(soffice).name,
                log=log,
            )
