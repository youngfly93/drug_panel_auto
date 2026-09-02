"""Build sanitized reference snapshots from legacy DOCX reports.

The snapshots are intended for local regression review. They deliberately avoid
persisting patient names, sample identifiers, dates, absolute paths, or original
filenames unless the caller explicitly opts in.
"""

from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass
from datetime import datetime
from hashlib import sha256
from pathlib import Path
from typing import Any, Iterable, Mapping, Optional

from docx import Document

from reportgen.utils.artifacts import write_json


SAMPLE_ID_RE = re.compile(r"\b(?:LZ|LW|MLF|MLB)\d{6,}[A-Z]?\b", re.IGNORECASE)
REPORT_ID_RE = re.compile(r"\bMLJY[-_ ]?(?:LZ|LW|MLF|MLB)\d{6,}[A-Z]?\b", re.IGNORECASE)
DATE_RE = re.compile(
    r"\b(?:20\d{2}[./-]\d{1,2}[./-]\d{1,2}|20\d{6})\b"
)
PANEL_RE = re.compile(r"结直肠癌\s*(301|358)\s*基因\s*\+?\s*MSI", re.IGNORECASE)
TOTAL_VARIANTS_RE = re.compile(r"本次共检出体细胞变异[:：]?\s*(\d+)\s*个")
DRUG_VARIANTS_RE = re.compile(r"与靶向药物用药相关的变异有[:：]?\s*(\d+)\s*个")
TMB_RE = re.compile(r"(\d+(?:\.\d+)?)\s*(?:mutations/Mb|Muts/Mb|个/Mb).*?(TMB-[HL])")
TMB_RESULT_RE = re.compile(
    r"肿瘤突变负荷为\s*(\d+(?:\.\d+)?)\s*(?:mutations/Mb|Muts/Mb|mut/Mb|个/Mb)"
    r".{0,30}?(TMB-[HL]|水平较高|水平较低)"
)
MSI_RE = re.compile(r"(?:微卫星稳定型[，, ]*)?(MSS|MSI-H|MSI-L)")
MSI_RESULT_RE = re.compile(
    r"该肿瘤样本为\s*(?:微卫星)?(?:"
    r"(?P<mss>稳定[（(]MSS[）)]|稳定型[，,]?MSS)|"
    r"(?P<msih>高度不稳定[（(]MSI-H[）)]|高度不稳定型[，,]?MSI-H)|"
    r"(?P<msil>低度不稳定[（(]MSI-L[）)]|低度不稳定型[，,]?MSI-L)"
    r")"
)
PDL1_RESULT_RE = re.compile(
    r"本次\s*PD\s*[-‐‑‒–—]?\s*L1\s*免疫组化检测结果[:：]\s*"
    r"TPS\s*(?P<tps>\d+(?:\.\d+)?)\s*%?\s*[，,]\s*"
    r"CPS\s*(?P<cps>\d+(?:\.\d+)?)\s*[，,]\s*"
    r"(?P<result>[^。；;\n]+)",
    re.IGNORECASE,
)
GENE_RE = re.compile(r"\b[A-Z][A-Z0-9]{1,9}\b")
PATIENT_NAME_FIELD_RE = re.compile(
    r"(?:患者姓名|受检者姓名|姓名|委托人)\s*[:：]?\s*([\u4e00-\u9fff·]{2,8})"
)
PATIENT_NAME_SALUTATION_RE = re.compile(
    r"尊敬的\s*([\u4e00-\u9fff·]{2,8})\s*(?:先生|女士|患者|家属)"
)
PATIENT_NAME_PREFIX_RE = re.compile(
    r"^([\u4e00-\u9fff·]{2,8})(?=[\-_—\s]|直肠癌|结肠癌|乙状结肠癌|肿瘤|癌)"
)
PATIENT_NAME_EXCLUDES = {
    "姓名",
    "患者",
    "报告",
    "检测",
    "样本",
    "组织",
    "肿瘤",
    "结直肠癌",
    "直肠癌",
    "结肠癌",
    "乙状结肠癌",
    "脉络医学",
    "医学",
    "先生",
    "女士",
    "未知",
}


KEY_SECTIONS = {
    "patient_letter": "致您的一封信",
    "variant_summary": "基因变异检测结果",
    "targeted_drugs": "靶向药物",
    "biomarkers": "TMB/MSI",
    "nccn": "NCCN",
    "variant_interpretation": "基因变异解析",
    "gene_list": "基因检测列表",
    "quality_control": "质控",
    "company_profile": "脉络医学检验简介",
}


@dataclass(frozen=True)
class LegacySnapshotOptions:
    panel: str
    source_dir: str
    output_dir: str
    sample_count: int = 5
    include_source_paths: bool = False
    max_text_chars: int = 12000


def build_legacy_reference_snapshots(
    options: LegacySnapshotOptions,
) -> dict[str, Any]:
    """Build sanitized snapshots for representative legacy DOCX reports."""
    source_dir = Path(options.source_dir).expanduser().resolve()
    output_dir = Path(options.output_dir).expanduser().resolve()
    sample_count = max(1, int(options.sample_count))
    if not source_dir.exists():
        raise FileNotFoundError(f"Legacy source directory not found: {source_dir}")

    docx_files = sorted(
        p
        for p in source_dir.iterdir()
        if p.is_file() and p.suffix.lower() == ".docx" and not p.name.startswith("._")
    )
    snapshots = [
        _snapshot_docx(
            path,
            panel=options.panel,
            include_source_paths=options.include_source_paths,
            max_text_chars=int(options.max_text_chars),
        )
        for path in docx_files
    ]
    readable = [row for row in snapshots if not row.get("read_error")]
    read_errors = [row for row in snapshots if row.get("read_error")]
    selected = _select_representative_snapshots(readable, sample_count=sample_count)

    output_dir.mkdir(parents=True, exist_ok=True)
    samples_dir = output_dir / "samples"
    samples_dir.mkdir(parents=True, exist_ok=True)
    for idx, row in enumerate(selected, start=1):
        sample_id = f"{options.panel}_legacy_ref_{idx:03d}"
        row["reference_id"] = sample_id
        row["source"] = _source_ref(row, include_paths=options.include_source_paths)
        row.pop("_source_path", None)
        row.pop("_source_name", None)
        write_json(samples_dir / f"{sample_id}.json", row)

    manifest = {
        "schema_version": "1.0",
        "generated_at": datetime.now().isoformat(),
        "panel": options.panel,
        "source_dir": str(source_dir) if options.include_source_paths else "<redacted>",
        "include_source_paths": bool(options.include_source_paths),
        "source_docx_count": len(docx_files),
        "readable_docx_count": len(readable),
        "matching_docx_count": len(readable),
        "read_error_count": len(read_errors),
        "selected_count": len(selected),
        "selection_strategy": (
            "one per distinct TMB/MSI/count/table signature, then most complete"
        ),
        "privacy": {
            "patient_names": "redacted",
            "sample_ids": "redacted",
            "report_ids": "redacted",
            "dates": "redacted",
            "source_paths": "redacted"
            if not options.include_source_paths
            else "included_by_user_request",
        },
        "panel_counts": dict(Counter(row.get("detected_panel") for row in readable)),
        "read_errors": [
            {
                "source": _source_ref(row, include_paths=options.include_source_paths),
                "error": row.get("read_error"),
            }
            for row in read_errors
        ],
        "selected_samples": [
            {
                "reference_id": row["reference_id"],
                "source": row["source"],
                "features": row["features"],
                "section_presence": row["section_presence"],
            }
            for row in selected
        ],
    }
    write_json(output_dir / "manifest.json", manifest)
    (output_dir / "README.md").write_text(
        _render_markdown(manifest),
        encoding="utf-8",
    )
    return {
        "status": "PASS",
        "output_dir": str(output_dir),
        "manifest_file": str(output_dir / "manifest.json"),
        "markdown_file": str(output_dir / "README.md"),
        "selected_count": len(selected),
        "source_docx_count": len(docx_files),
        "readable_docx_count": len(readable),
        "matching_docx_count": len(readable),
        "read_error_count": len(read_errors),
    }


def snapshot_docx_report(
    path: str | Path,
    *,
    panel: str,
    include_source_paths: bool = False,
    max_text_chars: int = 12000,
) -> dict[str, Any]:
    """Build one sanitized DOCX snapshot for QA contract checks."""
    snapshot = _snapshot_docx(
        Path(path).expanduser().resolve(),
        panel=panel,
        include_source_paths=include_source_paths,
        max_text_chars=max_text_chars,
    )
    if not include_source_paths:
        snapshot.pop("_source_path", None)
        snapshot.pop("_source_name", None)
    return snapshot


def _snapshot_docx(
    path: Path,
    *,
    panel: str,
    include_source_paths: bool,
    max_text_chars: int,
) -> dict[str, Any]:
    source_name = path.name
    try:
        doc = Document(str(path))
    except Exception as exc:
        return {
            "_source_path": str(path),
            "_source_name": source_name,
            "source_sha256": _sha256_file(path),
            "detected_panel": "unknown",
            "read_error": str(exc),
        }

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = [_table_snapshot(table) for table in doc.tables]
    text = "\n".join(paragraphs + [table["text"] for table in tables])
    patient_names = _patient_names_from_document(source_name, text=text, tables=tables)
    sanitized_text = _sanitize_text(text, patient_names=patient_names)
    detected_panel = _detect_panel(text) or panel
    features = _extract_features(sanitized_text, tables=tables)
    return {
        "_source_path": str(path),
        "_source_name": source_name,
        "source_sha256": _sha256_file(path),
        "source_label": source_name if include_source_paths else "<redacted>",
        "detected_panel": detected_panel,
        "read_error": None,
        "features": features,
        "section_presence": {
            key: needle in text for key, needle in KEY_SECTIONS.items()
        },
        "text_excerpt": sanitized_text[: max(0, int(max_text_chars))],
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "table_shapes": [
            {
                "index": idx,
                "rows": table["rows"],
                "columns": table["columns"],
                "header": _sanitize_table_row(table["header"], patient_names=patient_names),
            }
            for idx, table in enumerate(tables)
        ],
        "selection_key": _selection_key(features, table_count=len(tables)),
    }


def _table_snapshot(table: Any) -> dict[str, Any]:
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    return {
        "rows": len(rows),
        "columns": max((len(row) for row in rows), default=0),
        "header": rows[0] if rows else [],
        "_raw_rows": rows,
        "text": "\n".join("\t".join(row) for row in rows),
    }


def _extract_features(text: str, *, tables: Iterable[Mapping[str, Any]]) -> dict[str, Any]:
    total_variants = _first_int(TOTAL_VARIANTS_RE, text)
    drug_variants = _first_int(DRUG_VARIANTS_RE, text)
    tmb_value, tmb_status = _extract_tmb(text)
    msi_status = _extract_msi_status(text)
    pdl1_tps, pdl1_cps, pdl1_result = _extract_pdl1(text)
    genes = [
        gene
        for gene, _count in Counter(GENE_RE.findall(text)).most_common(20)
        if gene not in {"DNA", "MSI", "TMB", "MSS", "NCCN", "CSCO", "FDA"}
    ]
    table_shapes = Counter(
        f"{table.get('rows', 0)}x{table.get('columns', 0)}" for table in tables
    )
    return {
        "total_variants_count": total_variants,
        "drug_related_count": drug_variants,
        "tmb_value": tmb_value,
        "tmb_status": tmb_status,
        "msi_status": msi_status,
        "pdl1_tps": pdl1_tps,
        "pdl1_cps": pdl1_cps,
        "pdl1_result": pdl1_result,
        "top_gene_tokens": genes[:12],
        "table_shape_counts": dict(sorted(table_shapes.items())),
    }


def _extract_pdl1(
    text: str,
) -> tuple[Optional[float], Optional[float], Optional[str]]:
    """Extract the neutral PD-L1 transcription sentence from a report."""
    match = PDL1_RESULT_RE.search(text)
    if not match:
        return None, None, None
    return (
        float(match.group("tps")),
        float(match.group("cps")),
        match.group("result").strip(),
    )


def _select_representative_snapshots(
    snapshots: list[dict[str, Any]],
    *,
    sample_count: int,
) -> list[dict[str, Any]]:
    selected: list[dict[str, Any]] = []
    seen_keys: set[str] = set()
    ranked = sorted(
        snapshots,
        key=lambda row: (
            row["features"].get("msi_status") or "",
            row["features"].get("tmb_status") or "",
            -(row["features"].get("total_variants_count") or 0),
            -int(row.get("table_count") or 0),
            row.get("source_sha256") or "",
        ),
    )
    for row in ranked:
        key = str(row.get("selection_key"))
        if key in seen_keys:
            continue
        selected.append(row)
        seen_keys.add(key)
        if len(selected) >= sample_count:
            return selected

    for row in ranked:
        if row in selected:
            continue
        selected.append(row)
        if len(selected) >= sample_count:
            break
    return selected


def _selection_key(features: Mapping[str, Any], *, table_count: int) -> str:
    total = features.get("total_variants_count")
    drug = features.get("drug_related_count")
    total_bucket = "unknown" if total is None else ("0" if total == 0 else "1-5" if total <= 5 else "6+")
    drug_bucket = "unknown" if drug is None else ("0" if drug == 0 else "1-3" if drug <= 3 else "4+")
    return "|".join(
        [
            str(features.get("msi_status") or "unknown"),
            str(features.get("tmb_status") or "unknown"),
            total_bucket,
            drug_bucket,
            "many_tables" if table_count >= 15 else "few_tables",
        ]
    )


def _extract_tmb(text: str) -> tuple[Optional[float], Optional[str]]:
    result_match = TMB_RESULT_RE.search(text)
    if result_match:
        raw_status = result_match.group(2)
        if raw_status == "水平较高":
            status = "TMB-H"
        elif raw_status == "水平较低":
            status = "TMB-L"
        else:
            status = raw_status
        return float(result_match.group(1)), status
    fallback = TMB_RE.search(text)
    if not fallback:
        return None, None
    return float(fallback.group(1)), fallback.group(2)


def _extract_msi_status(text: str) -> Optional[str]:
    result_match = MSI_RESULT_RE.search(text)
    if result_match:
        if result_match.group("mss"):
            return "MSS"
        if result_match.group("msih"):
            return "MSI-H"
        if result_match.group("msil"):
            return "MSI-L"
    fallback = MSI_RE.search(text)
    return fallback.group(1) if fallback else None


def _detect_panel(text: str) -> Optional[str]:
    match = PANEL_RE.search(text)
    if not match:
        return None
    return f"crc_{match.group(1)}_msi"


def _first_int(pattern: re.Pattern[str], text: str) -> Optional[int]:
    match = pattern.search(text)
    return int(match.group(1)) if match else None


def _patient_name_from_filename(filename: str) -> Optional[str]:
    stem = Path(filename).stem
    match = PATIENT_NAME_PREFIX_RE.search(stem)
    if not match:
        return None
    value = match.group(1).strip()
    return value if _is_patient_name_candidate(value) else None


def _patient_names_from_document(
    filename: str,
    *,
    text: str,
    tables: Iterable[Mapping[str, Any]],
) -> list[str]:
    candidates: list[str] = []
    filename_name = _patient_name_from_filename(filename)
    if filename_name:
        candidates.append(filename_name)
    candidates.extend(PATIENT_NAME_FIELD_RE.findall(text))
    candidates.extend(PATIENT_NAME_SALUTATION_RE.findall(text))
    for table in tables:
        for row in table.get("_raw_rows") or []:
            cells = [str(cell or "").strip() for cell in row]
            for idx, cell in enumerate(cells):
                if cell in {"姓名", "患者姓名", "受检者姓名", "委托人"} and idx + 1 < len(cells):
                    candidates.append(cells[idx + 1])
                candidates.extend(PATIENT_NAME_FIELD_RE.findall(cell))
    return _unique_patient_names(candidates)


def _unique_patient_names(values: Iterable[str]) -> list[str]:
    names: list[str] = []
    seen: set[str] = set()
    for value in values:
        cleaned = str(value or "").strip()
        if not _is_patient_name_candidate(cleaned) or cleaned in seen:
            continue
        names.append(cleaned)
        seen.add(cleaned)
    return sorted(names, key=len, reverse=True)


def _is_patient_name_candidate(value: str) -> bool:
    if not re.fullmatch(r"[\u4e00-\u9fff·]{2,8}", value):
        return False
    if value in PATIENT_NAME_EXCLUDES:
        return False
    return not any(token in value for token in PATIENT_NAME_EXCLUDES)


def _sanitize_text(
    text: str,
    *,
    patient_names: Iterable[str] | None = None,
) -> str:
    sanitized = str(text or "")
    for patient_name in patient_names or []:
        sanitized = sanitized.replace(patient_name, "<PATIENT_NAME>")
    sanitized = REPORT_ID_RE.sub("<REPORT_ID>", sanitized)
    sanitized = SAMPLE_ID_RE.sub("<SAMPLE_ID>", sanitized)
    sanitized = DATE_RE.sub("<DATE>", sanitized)
    return sanitized


def _sanitize_table_row(
    row: Iterable[Any],
    *,
    patient_names: Iterable[str] | None,
) -> list[str]:
    return [
        _sanitize_text(str(cell or ""), patient_names=patient_names)
        for cell in row
    ]


def _source_ref(row: Mapping[str, Any], *, include_paths: bool) -> str:
    if include_paths:
        return str(row.get("_source_path") or row.get("_source_name") or "")
    digest = str(row.get("source_sha256") or "")
    return f"sha256:{digest[:12]}.docx" if digest else "<redacted>.docx"


def _sha256_file(path: Path) -> str:
    h = sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _render_markdown(manifest: Mapping[str, Any]) -> str:
    lines = [
        "# Legacy Reference Snapshots",
        "",
        f"- Panel: `{manifest.get('panel')}`",
        f"- Source DOCX count: {manifest.get('source_docx_count')}",
        f"- Readable DOCX count: {manifest.get('readable_docx_count')}",
        f"- Read errors: {manifest.get('read_error_count')}",
        f"- Selected samples: {manifest.get('selected_count')}",
        f"- Source paths: {manifest.get('privacy', {}).get('source_paths')}",
        "",
        "## Selected Samples",
        "",
    ]
    for sample in manifest.get("selected_samples") or []:
        features = sample.get("features") or {}
        lines.append(f"### {sample.get('reference_id')}")
        lines.append("")
        lines.append(f"- Source: `{sample.get('source')}`")
        lines.append(f"- Total variants: {features.get('total_variants_count')}")
        lines.append(f"- Drug-related variants: {features.get('drug_related_count')}")
        lines.append(f"- TMB: {features.get('tmb_value')} / {features.get('tmb_status')}")
        lines.append(f"- MSI: {features.get('msi_status')}")
        present = [
            key
            for key, value in (sample.get("section_presence") or {}).items()
            if value
        ]
        lines.append(f"- Sections present: {', '.join(present) if present else 'none'}")
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"
