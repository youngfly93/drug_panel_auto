"""Golden case runner for end-to-end report regression checks.

All fixtures are synthetic.  They provide reproducible release checks without
committing real patient source files.
"""

from __future__ import annotations

import re
from hashlib import sha1
from dataclasses import dataclass
from datetime import datetime
from functools import partial
from pathlib import Path
from typing import Any, Dict, List, Mapping, Optional

import pandas as pd
from docx import Document
from PIL import Image, ImageDraw

from reportgen.core.enhancer_registry import get_panel_registry
from reportgen.core.report_generator import ReportGenerator
from reportgen.models.excel_data import ExcelDataSource
from reportgen.utils.artifacts import write_json
from reportgen.utils.docx_render import render_docx_to_pngs


SUPPORTED_PANELS = {
    "crc_301": "crc_301_msi",
    "crc301": "crc_301_msi",
    "crc_301_msi": "crc_301_msi",
    "crc_358_msi": "crc_358_msi",
    "crc_358": "crc_358_msi",
    "crc358": "crc_358_msi",
    "lung_methylation": "lung_methylation",
    "lung_329": "lung_329_pdl1",
    "lung329": "lung_329_pdl1",
    "lung_329_pdl1": "lung_329_pdl1",
    "lung_13": "lung_13",
    "lung_62": "lung_62",
    "lung_62_pdl1": "lung_62_pdl1",
    "lung_588": "lung_588",
    "lung588": "lung_588",
    "lung_588_pdl1": "lung_588_pdl1",
}


@dataclass(frozen=True)
class GoldenCaseOptions:
    """Options for a single golden case run."""

    panel: str = "crc_358_msi"
    config_dir: str = "config"
    template_dir: str = "templates"
    template_file: Optional[str] = None
    output_root: Optional[str] = None
    log_level: str = "ERROR"
    template_contract_mode: str = "fail"
    render: str = "none"
    render_dpi: int = 120
    render_timeout_seconds: int = 120
    render_required: bool = False
    render_tmp_dir: Optional[str] = None


@dataclass(frozen=True)
class GoldenCaseInput:
    """A fixture path plus optional pre-resolved form-enriched input data."""

    excel_file: Path
    excel_data: Optional[ExcelDataSource] = None


CRC_358_MSI_EXPECTATIONS: Dict[str, Any] = {
    "project_type": "crc_358_msi",
    "project_name": "结直肠癌358基因+MSI",
    "expected_context": {
        "total_variants_count": 2,
        "drug_related_count": 1,
        "targeted_or_immune_related_count": 1,
        "tmb_status": "L",
        "msi_status": "MSS",
    },
    "required_qa_checks": [
        "unrendered_placeholders",
        "empty_numbered_paragraphs",
        "post_processors",
        "variant_detail_table_shape",
        "variant_summary_table_present",
        "targeted_drug_tip_table_present",
        "biomarker_table_present",
        "total_variant_count_text",
        "drug_related_count_text",
        "targeted_or_immune_related_count_text",
    ],
    "required_text": [
        "本次共检出体细胞变异：2个",
        "与靶向药物用药相关的变异有：1个",
        "与靶向/免疫药物相关的变异：1个",
        "6.5mutations/Mb，TMB-L",
        "微卫星稳定型，MSS",
        "多项临床研究表明，TMB-H的肿瘤",
        "研究表明，MSI-H的实体瘤",
        "ERBB2",
        "c.1979G>A",
        "p.G660D",
    ],
}


CRC_301_MSI_EXPECTATIONS: Dict[str, Any] = {
    "project_type": "crc_301_msi",
    "project_name": "结直肠癌301基因+MSI",
    "expected_context": {
        "total_variants_count": 2,
        "drug_related_count": 1,
        "tmb_status": "L",
        "msi_status": "MSS",
    },
    "required_qa_checks": list(CRC_358_MSI_EXPECTATIONS["required_qa_checks"]),
    "required_text": [
        "与肿瘤密切相关的301个基因进行检测",
        "本次共检出体细胞变异：2个",
        "与靶向药物用药相关的变异有：1个",
        "6.5mutations/Mb，TMB-L",
        "微卫星稳定型，MSS",
        "多项临床研究表明，TMB-H的肿瘤",
        "研究表明，MSI-H的实体瘤",
        "ERBB2",
        "c.1979G>A",
        "p.G660D",
    ],
}


LUNG_METHYLATION_EXPECTATIONS: Dict[str, Any] = {
    "project_type": "lung_methylation",
    "project_name": "肺癌甲基化",
    "expected_context": {
        "patient_name": "黄金甲基化患者",
        "sample_id": "LUNG999001",
        "project_name": "肺癌甲基化",
        "methylation_result": "阳性",
    },
    "required_qa_checks": [
        "template_contract",
        "unrendered_placeholders",
        "empty_numbered_paragraphs",
        "post_processors",
    ],
    "required_text": [
        "肺癌甲基化检测报告",
        "黄金甲基化患者",
        "LUNG999001",
        "肺癌甲基化",
        "阳性",
        "SHOX2",
        "RASSF1A",
        "78.5",
    ],
}


_LUNG_PDL1_REQUIRED_QA_CHECKS = [
    "template_contract",
    "rules",
    "docx_openable",
    "unrendered_placeholders",
    "empty_numbered_paragraphs",
    "toc_page_numbers",
    "field_provenance",
    "post_processors",
    "part3_cross_cancer_residuals",
    "pipeline",
]

_LUNG_PDL1_REQUIRED_TEXT = [
    "第三部分：基因变异及相应靶向/免疫药物解析",
    "ERBB2",
    "c.1979G>A",
    "p.G660D",
    "德曲妥珠单抗",
    "【待报告组审】",
    "TPS 50%，CPS 52",
    "阳性（高表达）",
]

LUNG_329_PDL1_EXPECTATIONS: Dict[str, Any] = {
    "project_type": "lung_329_pdl1",
    "project_name": "肺癌329基因+PD-L1",
    "expected_context": {
        "project_name": "肺癌329基因+PD-L1",
        "total_variants_count": 1,
        "drug_related_count": 1,
        "tmb_status": "H",
        "msi_status": "MSS",
    },
    "required_qa_checks": list(_LUNG_PDL1_REQUIRED_QA_CHECKS),
    "required_text": ["肺癌329基因检测", *_LUNG_PDL1_REQUIRED_TEXT],
}

LUNG_588_PDL1_EXPECTATIONS: Dict[str, Any] = {
    "project_type": "lung_588_pdl1",
    "project_name": "肺癌588基因+PD-L1",
    "expected_context": {
        "project_name": "肺癌588基因+PD-L1",
        "total_variants_count": 1,
        "drug_related_count": 1,
        "tmb_status": "H",
        "msi_status": "MSS",
    },
    "required_qa_checks": list(_LUNG_PDL1_REQUIRED_QA_CHECKS),
    "required_text": ["肺癌588基因检测", *_LUNG_PDL1_REQUIRED_TEXT],
}


def run_golden_case(
    options: Optional[GoldenCaseOptions] = None, **kwargs: Any
) -> Dict[str, Any]:
    """Run the configured golden case and return a machine-readable summary."""
    opts = options or GoldenCaseOptions(**kwargs)
    panel = _normalize_panel(opts.panel)
    case = _golden_case_spec(panel)
    expectations = case["expectations"]

    output_root = _resolve_output_root(opts.output_root, panel)
    input_dir = output_root / "input"
    report_dir = output_root / "output"
    input_dir.mkdir(parents=True, exist_ok=True)
    report_dir.mkdir(parents=True, exist_ok=True)

    built_input = case["builder"](input_dir / case["input_filename"])
    if isinstance(built_input, GoldenCaseInput):
        excel_file = built_input.excel_file
        excel_data = built_input.excel_data
    else:
        excel_file = Path(built_input)
        excel_data = None
    template_file = _resolve_template_file(opts)

    generator = ReportGenerator(
        config_dir=opts.config_dir,
        template_dir=opts.template_dir,
        log_level=opts.log_level,
    )
    generation = generator.generate(
        excel_file=str(excel_file),
        excel_data=excel_data,
        template_file=str(template_file),
        output_dir=str(report_dir),
        output_filename=case["output_filename"],
        strict_mode=True,
        return_context=True,
        template_contract_mode=opts.template_contract_mode,
        project_type=panel,
        project_name=expectations["project_name"],
    )

    assertion = assert_golden_case_output(generation, expectations=expectations)
    visual_render = run_visual_render(
        generation.get("output_file"),
        output_root=output_root,
        mode=opts.render,
        dpi=opts.render_dpi,
        timeout_seconds=opts.render_timeout_seconds,
        required=opts.render_required,
        tmp_dir=opts.render_tmp_dir,
    )
    checks = list(assertion["checks"])
    if visual_render["requested"] != "none":
        checks.append(
            {
                "name": "visual_render",
                "passed": visual_render["status"] == "PASS"
                or not opts.render_required,
                "message": visual_render["message"],
                "details": visual_render,
            }
        )

    summary = {
        "schema_version": "1.0",
        "generated_at": datetime.now().isoformat(),
        "panel": panel,
        "ok": bool(generation.get("success"))
        and assertion["ok"]
        and (visual_render["status"] != "FAIL"),
        "input_excel": str(excel_file),
        "template_file": str(template_file),
        "output_root": str(output_root),
        "output_file": generation.get("output_file"),
        "qa_report_file": generation.get("qa_report_file"),
        "field_provenance_file": generation.get("field_provenance_file"),
        "qa_status": generation.get("qa_status"),
        "duration": generation.get("duration"),
        "generation_errors": generation.get("errors") or [],
        "generation_warnings": generation.get("warnings") or [],
        "visual_render": visual_render,
        "checks": checks,
    }
    summary["errors"] = [
        check["message"] for check in summary["checks"] if not check["passed"]
    ] + list(generation.get("errors") or [])

    report_path = output_root / "golden_case_report.json"
    write_json(report_path, summary)
    summary["golden_report_file"] = str(report_path)
    return summary


def run_visual_render(
    output_file: Any,
    *,
    output_root: Path,
    mode: str = "none",
    dpi: int = 120,
    timeout_seconds: int = 120,
    required: bool = False,
    tmp_dir: Optional[str] = None,
) -> Dict[str, Any]:
    """Render a generated golden DOCX to PNG pages when requested."""
    normalized_mode = str(mode or "none").strip().lower()
    if normalized_mode not in {"none", "first", "all"}:
        raise ValueError(f"Unsupported render mode: {mode!r}")

    result: Dict[str, Any] = {
        "requested": normalized_mode,
        "required": bool(required),
        "status": "SKIPPED",
        "message": "visual rendering was not requested",
        "rendered_pages": [],
        "output_dir": None,
        "error": None,
    }
    if normalized_mode == "none":
        return result

    if not output_file:
        result.update(
            {
                "status": "FAIL" if required else "WARN",
                "message": "visual rendering requested but no DOCX output exists",
                "error": "missing_output_file",
            }
        )
        return result

    output_path = Path(str(output_file))
    render_dir = output_root / "rendered_pages" / output_path.stem
    first_page = 1 if normalized_mode == "first" else None
    last_page = 1 if normalized_mode == "first" else None
    try:
        pngs = render_docx_to_pngs(
            output_path,
            output_dir=render_dir,
            dpi=int(dpi),
            first_page=first_page,
            last_page=last_page,
            keep_pdf=False,
            timeout_seconds=int(timeout_seconds),
            tmp_dir=Path(tmp_dir) if tmp_dir else None,
        )
    except Exception as exc:
        result.update(
            {
                "status": "FAIL" if required else "WARN",
                "message": f"visual rendering failed: {exc}",
                "output_dir": str(render_dir),
                "error": str(exc),
            }
        )
        stage = getattr(exc, "stage", None)
        if stage:
            result["stage"] = stage
        command = getattr(exc, "command", None)
        if command:
            result["command"] = list(command)
        stdout = getattr(exc, "stdout", None)
        stderr = getattr(exc, "stderr", None)
        if stdout:
            result["stdout_tail"] = str(stdout)[-2000:]
        if stderr:
            result["stderr_tail"] = str(stderr)[-2000:]
        return result

    result.update(
        {
            "status": "PASS" if pngs else ("FAIL" if required else "WARN"),
            "message": "visual rendering produced PNG pages"
            if pngs
            else "visual rendering completed but produced no PNG pages",
            "rendered_pages": [str(path) for path in pngs],
            "output_dir": str(render_dir),
        }
    )
    return result


def _golden_case_spec(panel: str) -> Dict[str, Any]:
    if panel in {"lung_13", "lung_62", "lung_62_pdl1", "lung_588"}:
        count = panel.split("_")[1]
        pdl1 = panel.endswith("_pdl1")
        name = f"肺癌{count}基因" + ("+PD-L1" if pdl1 else "")
        return {
            "expectations": {
                "project_type": panel,
                "project_name": name,
                "expected_context": {
                    **LUNG_588_PDL1_EXPECTATIONS["expected_context"],
                    "project_name": name,
                },
                "required_qa_checks": [
                    *_LUNG_PDL1_REQUIRED_QA_CHECKS,
                    "case_total_count_consistency", "case_drug_count_consistency",
                ],
                "required_text": [
                    "报告组评审草稿（非临床交付）",
                    "ERBB2", "c.1979G>A", "p.G660D",
                    "SYNTHETIC-PGX-OBSERVATION",
                    *(["TPS 50%，CPS 52", "阳性（高表达）"] if pdl1 else []),
                ],
            },
            "builder": partial(
                _build_lung_pdl1_golden_input,
                panel_id=panel,
                project_name=name,
                sample_id=f"SYN-{panel.upper()}-GOLDEN",
                membership_column=f"ExistInsmall{count}",
                membership_value=1,
                include_pdl1=pdl1,
                include_synthetic_neutral_cnv=True,
                include_synthetic_pgx=True,
            ),
            "input_filename": f"SYN-{panel.upper()}-GOLDEN.xlsx",
            "output_filename": f"golden_{panel}.docx",
        }
    if panel == "crc_301_msi":
        return {
            "expectations": CRC_301_MSI_EXPECTATIONS,
            "builder": build_crc_301_msi_golden_excel,
            "input_filename": "LZ999301_crc_301_msi_golden.xlsx",
            "output_filename": "golden_crc_301_msi.docx",
        }
    if panel == "lung_methylation":
        return {
            "expectations": LUNG_METHYLATION_EXPECTATIONS,
            "builder": build_lung_methylation_golden_excel,
            "input_filename": "LUNG999001_lung_methylation_golden.xlsx",
            "output_filename": "golden_lung_methylation.docx",
        }
    if panel == "lung_329_pdl1":
        return {
            "expectations": LUNG_329_PDL1_EXPECTATIONS,
            "builder": build_lung_329_pdl1_golden_input,
            "input_filename": "SYN-L329-GOLDEN.xlsx",
            "output_filename": "golden_lung_329_pdl1.docx",
        }
    if panel == "lung_588_pdl1":
        return {
            "expectations": LUNG_588_PDL1_EXPECTATIONS,
            "builder": build_lung_588_pdl1_golden_input,
            "input_filename": "SYN-L588-GOLDEN.xlsx",
            "output_filename": "golden_lung_588_pdl1.docx",
        }
    return {
        "expectations": CRC_358_MSI_EXPECTATIONS,
        "builder": build_crc_358_msi_golden_excel,
        "input_filename": "LZ999001_crc_358_msi_golden.xlsx",
        "output_filename": "golden_crc_358_msi.docx",
    }


def build_crc_358_msi_golden_excel(path: Path | str) -> Path:
    """Create a synthetic CRC 358 + MSI workbook suitable for the golden case."""
    return _build_crc_msi_golden_excel(
        path,
        panel_gene_count=358,
        sample_id="LZ999001",
    )


def build_crc_301_msi_golden_excel(path: Path | str) -> Path:
    """Create a synthetic CRC 301 + MSI workbook suitable for the golden case."""
    return _build_crc_msi_golden_excel(
        path,
        panel_gene_count=301,
        sample_id="LZ999301",
    )


def _build_crc_msi_golden_excel(
    path: Path | str,
    *,
    panel_gene_count: int,
    sample_id: str,
) -> Path:
    """Create a synthetic CRC + MSI workbook for the requested CRC panel size."""
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    panel_name = f"结直肠癌{panel_gene_count}基因+MSI"
    panel_flag_column = f"ExistInsmall{panel_gene_count}"

    meta = pd.DataFrame(
        [
            {
                "患者姓名": "黄金测试患者",
                "样本编号": sample_id,
                "报告编号": f"MLJY-{sample_id}",
                "性别": "男",
                "年龄": 58,
                "临床诊断": "结直肠癌",
                "肿瘤类型": "结直肠癌",
                "样本类型": "组织",
                "取材手段": "手术",
                "取材部位": "结肠",
                "项目名称": panel_name,
                "检测项目": panel_name,
                "送检日期": "2026-01-08",
                "报告日期": "2026-01-15",
                "检测方法": "NGS高通量测序",
            }
        ]
    )

    variations = pd.DataFrame(
        [
            {
                "Gene_Symbol": "ERBB2",
                "Transcript": "NM_004448.4",
                "Chr": "17",
                "ExIn_ID": "EX17",
                "cHGVS": "c.1979G>A",
                "pHGVS_S": "p.G660D",
                "Freq(%)": 22.5,
                "Function": "Missense",
                panel_flag_column: 1,
                "ExistIn552": "Ⅰ类",
                "CLNSIG": "Pathogenic",
            },
            {
                "Gene_Symbol": "FBXW7",
                "Transcript": "NM_033632.3",
                "Chr": "4",
                "ExIn_ID": "EX10",
                "cHGVS": "c.1394G>A",
                "pHGVS_S": "p.R465H",
                "Freq(%)": 12.7,
                "Function": "Missense",
                panel_flag_column: 1,
                "ExistIn552": "Ⅲ类",
                "CLNSIG": "Uncertain_significance",
            },
        ]
    )

    tmb = pd.DataFrame(
        [
            ["TCGA fit", None, None, None],
            ["SampleTP", "Var_num", "Bed_size", "TMB"],
            ["tissue", 65, 10_000_000, 6.5],
        ]
    )
    msisensor = pd.DataFrame(
        [
            ["control", 1000, 3, 0.3, "MSS"],
            ["tumor", 1000, 12, 1.2, "MSS"],
        ],
        columns=["Sample", "Total", "Unstable", "Percent", "Status"],
    )
    qc = pd.DataFrame(
        [
            ["Q30", 95.5],
            ["Coverage", 99.8],
            ["Average sequencing depth", 800],
            ["Insert", 180],
        ]
    )

    empty_cnv = pd.DataFrame(columns=["Gene", "Chr", "Start", "End", "CopyNumber"])
    empty_fusion = pd.DataFrame(
        columns=["Gene1", "Gene2", "Chr1", "Pos1", "Chr2", "Pos2"]
    )

    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        meta.to_excel(writer, sheet_name="Meta", index=False)
        variations.to_excel(writer, sheet_name="Variations", index=False)
        tmb.to_excel(writer, sheet_name="TMB", index=False, header=False)
        msisensor.to_excel(writer, sheet_name="Msisensor", index=False)
        qc.to_excel(writer, sheet_name="QC", index=False, header=False)
        empty_cnv.to_excel(writer, sheet_name="Cnv", index=False)
        empty_fusion.to_excel(writer, sheet_name="Fusion", index=False)

    return out


def build_lung_methylation_golden_excel(path: Path | str) -> Path:
    """Create a synthetic lung methylation workbook for M3 panel testing."""
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)

    meta = pd.DataFrame(
        [
            {
                "患者姓名": "黄金甲基化患者",
                "样本编号": "LUNG999001",
                "报告编号": "MLJY-LUNG999001",
                "性别": "女",
                "年龄": 62,
                "临床诊断": "肺癌",
                "肿瘤类型": "肺癌",
                "样本类型": "血浆",
                "项目名称": "肺癌甲基化",
                "检测项目": "肺癌甲基化",
                "甲基化结果": "阳性",
                "送检日期": "2026-02-01",
                "报告日期": "2026-02-08",
                "检测方法": "甲基化特异性测序",
            }
        ]
    )
    sites = pd.DataFrame(
        [
            {
                "基因": "SHOX2",
                "位点": "cg000001",
                "甲基化水平": 78.5,
                "阈值": 10.0,
                "结果": "阳性",
            },
            {
                "基因": "RASSF1A",
                "位点": "cg000002",
                "甲基化水平": 32.1,
                "阈值": 8.0,
                "结果": "阳性",
            },
        ]
    )

    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        meta.to_excel(writer, sheet_name="Meta", index=False)
        sites.to_excel(writer, sheet_name="甲基化位点", index=False)

    return out


def build_lung_329_pdl1_golden_input(path: Path | str) -> GoldenCaseInput:
    """Create the synthetic report-group golden input for lung 329 + PD-L1."""
    return _build_lung_pdl1_golden_input(
        path,
        panel_id="lung_329_pdl1",
        project_name="肺癌329基因+PD-L1",
        sample_id="SYN-L329-GOLDEN",
    )


def build_lung_588_pdl1_golden_input(path: Path | str) -> GoldenCaseInput:
    """Create the synthetic report-group golden input for lung 588 + PD-L1."""
    return _build_lung_pdl1_golden_input(
        path,
        panel_id="lung_588_pdl1",
        project_name="肺癌588基因+PD-L1",
        sample_id="SYN-L588-GOLDEN",
    )


def _build_lung_pdl1_golden_input(
    path: Path | str,
    *,
    panel_id: str,
    project_name: str,
    sample_id: str,
    membership_column: str = "ExistInsmall588",
    membership_value: Any = "Ⅰ类",
    include_pdl1: bool = True,
    include_synthetic_neutral_cnv: bool = False,
    include_synthetic_pgx: bool = False,
) -> GoldenCaseInput:
    """Build a deterministic NGS workbook plus explicit synthetic form data."""
    out = Path(path)
    out.parent.mkdir(parents=True, exist_ok=True)
    image_path = out.with_suffix(".pdl1.png")
    if include_pdl1:
        _write_synthetic_pdl1_image(image_path)

    meta = pd.DataFrame(
        [
            {
                "患者姓名": sample_id,
                "样本编号": sample_id,
                "报告编号": f"MLJY-{sample_id}",
                "性别": "男",
                "年龄": 60,
                "临床诊断": "肺癌",
                "肿瘤类型": "肺癌",
                "样本类型": "组织",
                "送检医院": "合成验证机构",
                "项目名称": project_name,
                "检测项目": project_name,
                "报告日期": "2026-07-24",
            }
        ]
    )
    variation = {
        "ExistIn552": "Ⅰ类",
        membership_column: membership_value,
        "Gene_Symbol": "ERBB2",
        "Transcript": "NM_004448.4",
        "Chr": "17",
        "Exon": "17",
        "cHGVS": "c.1979G>A",
        "pHGVS_S": "p.G660D",
        "Mutation_Type": "SNV",
        "Freq(%)": 8.5,
    }
    tmb = pd.DataFrame(
        [
            ["TCGA fit", None, None, None],
            ["SampleTP", "Var_num", "Bed_size", "TMB"],
            ["tissue", 100, 10_000_000, 10.0],
        ]
    )
    msisensor = pd.DataFrame(
        [["tumor", 1000, 10, 1.0, "MSS"]],
        columns=["Sample", "Total", "Unstable", "Percent", "Status"],
    )
    hereditary = pd.DataFrame(
        columns=["Gene_Symbol", membership_column, "ExistIn178"]
    )
    # A complete synthetic positive-path contract is separate from real CNV
    # absence/ambiguity tests. Never apply these invented observations to a
    # patient workbook or to the real derived-input validator.
    synthetic_cnv = [
        {"Gene": gene, "Status": "neutral"}
        for gene in ("EGFR", "MDM2", "MDM4", "CCND1", "FGF3", "FGF4", "FGF19")
    ]
    synthetic_pgx = [{
        "药物": "顺铂（cisplatin）", "Gene": "ERCC1", "Locus": "rsTEST",
        "Genotype": "AA", "Level": "1B", "Result": "SYNTHETIC-PGX-OBSERVATION",
    }]
    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        meta.to_excel(writer, sheet_name="Meta", index=False)
        pd.DataFrame([variation]).to_excel(
            writer, sheet_name="Variations", index=False
        )
        tmb.to_excel(writer, sheet_name="TMB", index=False, header=False)
        msisensor.to_excel(writer, sheet_name="Msisensor", index=False)
        hereditary.to_excel(writer, sheet_name="Hereditary_tumor", index=False)
        if include_synthetic_neutral_cnv:
            pd.DataFrame(synthetic_cnv).to_excel(writer, sheet_name="Cnv", index=False)
        if include_synthetic_pgx:
            pd.DataFrame(synthetic_pgx).to_excel(writer, sheet_name="CtDrug", index=False)

    form_fields = (
        "pdl1_tps",
        "pdl1_cps",
        "pdl1_result",
        "pdl1_image_path",
        "pdl1_assay_profile_id",
        "pdl1_source_record_id",
        "pdl1_source_record_date",
        "pdl1_specimen_id",
        "pdl1_image_disposition",
        "lung_histology",
        "disease_extent",
        "prior_systemic_therapy",
        "companion_diagnostic_status",
    )
    excel_data = ExcelDataSource(
        file_path=str(out),
        single_values={
            **meta.iloc[0].to_dict(),
            "癌种": "肺癌",
            "TMB": 10,
            "MSI状态": "MSS",
            "PD-L1 TPS": 50,
            "PD-L1 CPS": 52,
            "PD-L1结果": "阳性（高表达）",
            "PD-L1病例图片": str(image_path),
            "PD-L1检测方案": "legacy_unspecified_ihc_transcription_v1",
            "PD-L1原始记录编号": f"SYNTHETIC-IHC-{sample_id}",
            "PD-L1原始记录日期": "2026-07-24",
            "PD-L1检测标本标识": f"SYNTHETIC-SPECIMEN-{sample_id}",
            "PD-L1图像处置": "病例专属图像（报告展示）",
            "肺癌病理类型": "非小细胞肺癌",
            "疾病范围": "转移性",
            "既往系统治疗": "已接受",
            "伴随诊断适配状态": "已确认符合",
        },
        table_data={
            "Variations": [variation],
            "TMB": [],
            "Msisensor": [],
            "Hereditary_tumor": [],
        },
        sheet_names=["Meta", "Variations", "TMB", "Msisensor", "Hereditary_tumor"],
        metadata={
            "field_source_overrides": {
                field: {
                    "source": "form",
                    "source_key": field,
                    "source_detail": "synthetic_golden_case_form",
                }
                for field in form_fields
            },
            "synthetic_fixture": True,
            "panel_id": panel_id,
        },
    )
    if not include_pdl1:
        excel_data.single_values = {
            key: value for key, value in excel_data.single_values.items()
            if not key.startswith("PD-L1")
        }
        excel_data.metadata["field_source_overrides"] = {
            key: value for key, value in excel_data.metadata["field_source_overrides"].items()
            if not key.startswith("pdl1_")
        }
    if include_synthetic_neutral_cnv:
        excel_data.table_data["Cnv"] = synthetic_cnv
        excel_data.sheet_names.append("Cnv")
    if include_synthetic_pgx:
        excel_data.table_data["CtDrug"] = synthetic_pgx
        excel_data.sheet_names.append("CtDrug")
    excel_data.metadata["table_columns"] = {
        "Variations": list(variation),
        "Hereditary_tumor": list(hereditary.columns),
    }
    return GoldenCaseInput(excel_file=out, excel_data=excel_data)


def _write_synthetic_pdl1_image(path: Path) -> None:
    """Write a visibly synthetic image that exercises the media pipeline."""
    image = Image.new("RGB", (960, 640), "white")
    draw = ImageDraw.Draw(image)
    for index in range(72):
        x = (index * 127 + 41) % 920
        y = (index * 83 + 29) % 600
        diameter = 12 + (index * 7) % 22
        draw.ellipse(
            (x, y, x + diameter, y + diameter),
            fill=(110, 75, 145 + (index * 11) % 70),
        )
    draw.rectangle((10, 10, 949, 629), outline=(75, 75, 75), width=3)
    image.save(path, format="PNG", optimize=True)


def assert_golden_case_output(
    generation: Mapping[str, Any],
    *,
    expectations: Optional[Mapping[str, Any]] = None,
) -> Dict[str, Any]:
    """Assert a generated report against frozen golden expectations."""
    expected = expectations or CRC_358_MSI_EXPECTATIONS
    checks: List[Dict[str, Any]] = []

    def check(name: str, passed: bool, message: str, **details: Any) -> None:
        checks.append(
            {"name": name, "passed": bool(passed), "message": message, **details}
        )

    output_file = generation.get("output_file")
    output_path = Path(str(output_file)) if output_file else None
    check(
        "generation_success",
        bool(generation.get("success")),
        "report generator returned success",
        errors=generation.get("errors") or [],
    )
    check(
        "output_docx_exists",
        bool(output_path and output_path.exists()),
        "generated DOCX exists",
        output_file=str(output_file) if output_file else None,
    )

    context = generation.get("context") or {}
    for field, value in (expected.get("expected_context") or {}).items():
        check(
            f"context_{field}",
            context.get(field) == value,
            f"context field {field} matches expected value",
            expected=value,
            actual=context.get(field),
        )

    qa = generation.get("qa_report") or {}
    check(
        "qa_status_pass",
        qa.get("status") == "PASS",
        "QA report status is PASS",
        actual=qa.get("status"),
        issues=qa.get("issues") or [],
    )
    check(
        "field_provenance_present",
        bool(generation.get("field_provenance_file")),
        "field provenance sidecar was generated",
        field_provenance_file=generation.get("field_provenance_file"),
    )

    qa_checks = qa.get("checks") if isinstance(qa, Mapping) else {}
    for qa_check_name in expected.get("required_qa_checks") or []:
        item = qa_checks.get(qa_check_name, {}) if isinstance(qa_checks, Mapping) else {}
        status = item.get("status")
        check(
            f"qa_{qa_check_name}",
            status == "PASS",
            f"QA check {qa_check_name} is PASS",
            actual=status,
            details=item,
        )

    post_processors = (
        generation.get("post_processors") or qa.get("post_processors") or []
    )
    processor_errors = [
        row
        for row in post_processors
        if isinstance(row, Mapping) and row.get("status") == "ERROR"
    ]
    check(
        "post_processors_no_errors",
        not processor_errors,
        "post-render processors completed without ERROR status",
        errors=processor_errors,
    )

    text = _read_docx_text(output_path) if output_path and output_path.exists() else ""
    compact_text = _compact(text)
    for required in expected.get("required_text") or []:
        check(
            f"text_contains_{_slug(required)}",
            _compact(required) in compact_text,
            "required golden text appears in rendered DOCX",
            required=required,
        )

    return {"ok": all(row["passed"] for row in checks), "checks": checks}


def _normalize_panel(panel: str) -> str:
    normalized = SUPPORTED_PANELS.get(str(panel or "").strip())
    if not normalized:
        supported = ", ".join(sorted(SUPPORTED_PANELS))
        raise ValueError(f"Unsupported golden panel: {panel!r}. Supported: {supported}")
    return normalized


def _resolve_output_root(output_root: Optional[str], panel: str) -> Path:
    if output_root:
        return Path(output_root).resolve()
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return (Path("tmp") / "golden_cases" / f"{panel}_{stamp}").resolve()


def _resolve_template_file(opts: GoldenCaseOptions) -> Path:
    if opts.template_file:
        template = Path(opts.template_file)
    else:
        template = None
        try:
            registration = get_panel_registry().get(_normalize_panel(opts.panel))
            if registration and registration.package:
                template = registration.package.resolve_template_file()
        except Exception:
            template = None
        if template is None:
            template = (
                Path(opts.template_dir)
                / "aligned_template_with_cnv_fusion_hla_FIXED.docx"
            )
    template = template.resolve()
    if not template.exists():
        raise FileNotFoundError(f"Golden case template not found: {template}")
    return template


def _read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: List[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _compact(text: Any) -> str:
    return re.sub(r"\s+", "", str(text or ""))


def _slug(text: str) -> str:
    value = re.sub(r"\W+", "_", _compact(text), flags=re.ASCII).strip("_").lower()
    if value:
        return value[:60]
    digest = sha1(str(text).encode("utf-8")).hexdigest()[:12]
    return f"required_text_{digest}"
