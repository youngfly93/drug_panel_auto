"""Machine-readable contracts for report-group review candidates.

Unlike a clinically approved historical golden contract, a review-candidate
contract freezes engineering invariants only: template identity, document
geometry, required review notices, dynamic marker removal, gene denominator,
and blocking visual-QA evidence.  It must not be used to claim medical UAT or
clinical release approval.
"""

from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence

import yaml
from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn


SHA256_RE = re.compile(r"^[0-9a-f]{64}$")


def load_review_candidate_contract(path: str | Path) -> dict[str, Any]:
    """Load and validate the small, non-clinical contract envelope."""

    contract_path = Path(path)
    payload = yaml.safe_load(contract_path.read_text(encoding="utf-8")) or {}
    if not isinstance(payload, Mapping):
        raise ValueError(
            f"review candidate contract must be a mapping: {contract_path}"
        )
    contract = dict(payload)
    if contract.get("schema_version") != "1.0":
        raise ValueError("review candidate contract schema_version must be '1.0'")
    for field in ("contract_id", "panel_id"):
        if not str(contract.get(field) or "").strip():
            raise ValueError(f"review candidate contract requires {field}")

    lifecycle = contract.get("lifecycle") or {}
    if lifecycle.get("status") != "report_group_review_candidate":
        raise ValueError(
            "review candidate lifecycle.status must be 'report_group_review_candidate'"
        )
    if lifecycle.get("clinical_release_status") != "blocked":
        raise ValueError("review candidate clinical release must remain blocked")

    privacy = contract.get("privacy") or {}
    if privacy.get("contains_phi") is not False:
        raise ValueError("review candidate contract privacy.contains_phi must be false")

    template = contract.get("template") or {}
    template_sha = str(template.get("sha256") or "").strip().lower()
    if not SHA256_RE.fullmatch(template_sha):
        raise ValueError("review candidate template.sha256 must be a SHA256")
    if str(template.get("status") or "").strip().lower() != "pilot":
        raise ValueError("review candidate template must remain pilot before approval")
    return contract


def extract_docx_text(path: str | Path) -> str:
    """Extract joined text from body, text boxes, headers, and footers."""

    document = Document(Path(path))
    parts = [document.part]
    for section in document.sections:
        parts.extend(
            (
                section.header.part,
                section.first_page_header.part,
                section.even_page_header.part,
                section.footer.part,
                section.first_page_footer.part,
                section.even_page_footer.part,
            )
        )

    seen: set[str] = set()
    paragraphs: list[str] = []
    for part in parts:
        part_name = str(part.partname)
        if part_name in seen or not hasattr(part, "element"):
            continue
        seen.add(part_name)
        for paragraph in part.element.iter(qn("w:p")):
            text = "".join(node.text or "" for node in paragraph.iter(qn("w:t")))
            if text:
                paragraphs.append(text)
    return "\n".join(paragraphs)


def validate_review_candidate_template(
    contract: Mapping[str, Any],
    template_path: str | Path,
    *,
    template_id: str | None = None,
    template_version: str | None = None,
    template_status: str | None = None,
) -> dict[str, Any]:
    """Validate the frozen candidate template and its package registration."""

    path = Path(template_path)
    checks: list[dict[str, Any]] = []
    errors: list[dict[str, Any]] = []

    def check(code: str, passed: bool, *, expected: Any, actual: Any) -> None:
        item = {
            "code": code,
            "status": "PASS" if passed else "FAIL",
            "expected": expected,
            "actual": actual,
        }
        checks.append(item)
        if not passed:
            errors.append(item)

    template = contract.get("template") or {}
    check("TEMPLATE_EXISTS", path.is_file(), expected=True, actual=path.is_file())
    if not path.is_file():
        return _result(contract, path, checks, errors)

    actual_sha = _sha256(path)
    check(
        "TEMPLATE_SHA256",
        actual_sha == str(template.get("sha256") or "").lower(),
        expected=template.get("sha256"),
        actual=actual_sha,
    )
    if template_id is not None:
        check(
            "TEMPLATE_ID",
            template_id == str(template.get("id") or ""),
            expected=template.get("id"),
            actual=template_id,
        )
    if template_version is not None:
        check(
            "TEMPLATE_VERSION",
            template_version == str(template.get("version") or ""),
            expected=template.get("version"),
            actual=template_version,
        )
    if template_status is not None:
        check(
            "TEMPLATE_STATUS",
            template_status == str(template.get("status") or ""),
            expected=template.get("status"),
            actual=template_status,
        )

    try:
        document = Document(path)
        visible = extract_docx_text(path)
    except Exception as exc:
        check("TEMPLATE_OPENABLE", False, expected="openable", actual=str(exc))
        return _result(contract, path, checks, errors)
    check("TEMPLATE_OPENABLE", True, expected="openable", actual="openable")
    _check_structure(check, document, template.get("structure") or {}, "TEMPLATE")
    _check_text_contract(
        check,
        visible,
        required=template.get("required_texts") or [],
        forbidden=template.get("forbidden_texts") or [],
        prefix="TEMPLATE",
    )
    _check_text_contract(
        check,
        visible,
        required=template.get("required_markers") or [],
        forbidden=[],
        prefix="TEMPLATE_MARKER",
    )
    _check_dynamic_scope(
        check,
        document,
        template.get("dynamic_scope") or {},
        "TEMPLATE",
    )
    _check_gene_list(check, document, contract.get("gene_list") or {}, "TEMPLATE")
    return _result(contract, path, checks, errors, docx_sha256=actual_sha)


def validate_review_candidate_output(
    contract: Mapping[str, Any],
    docx_path: str | Path,
    *,
    qa_report: Mapping[str, Any] | None = None,
    expected_texts: Sequence[str] = (),
    forbidden_texts: Sequence[str] = (),
    expected_page_count: int | None = None,
    require_case_image: bool = False,
) -> dict[str, Any]:
    """Validate one generated report against the review-candidate baseline."""

    path = Path(docx_path)
    checks: list[dict[str, Any]] = []
    errors: list[dict[str, Any]] = []

    def check(code: str, passed: bool, *, expected: Any, actual: Any) -> None:
        item = {
            "code": code,
            "status": "PASS" if passed else "FAIL",
            "expected": expected,
            "actual": actual,
        }
        checks.append(item)
        if not passed:
            errors.append(item)

    check("OUTPUT_EXISTS", path.is_file(), expected=True, actual=path.is_file())
    if not path.is_file():
        return _result(contract, path, checks, errors)

    actual_sha = _sha256(path)
    try:
        document = Document(path)
        visible = extract_docx_text(path)
    except Exception as exc:
        check("OUTPUT_OPENABLE", False, expected="openable", actual=str(exc))
        return _result(contract, path, checks, errors, docx_sha256=actual_sha)
    check("OUTPUT_OPENABLE", True, expected="openable", actual="openable")

    output = contract.get("rendered_output") or {}
    _check_structure(check, document, output.get("structure") or {}, "OUTPUT")
    _check_text_contract(
        check,
        visible,
        required=[*(output.get("required_texts") or []), *expected_texts],
        forbidden=[*(output.get("forbidden_texts") or []), *forbidden_texts],
        prefix="OUTPUT",
    )
    _check_dynamic_scope(
        check,
        document,
        output.get("dynamic_scope") or {},
        "OUTPUT",
    )
    _check_gene_list(check, document, contract.get("gene_list") or {}, "OUTPUT")

    if require_case_image:
        template_shape_count = int(
            ((contract.get("template") or {}).get("structure") or {}).get(
                "inline_shape_count"
            )
            or 0
        )
        check(
            "OUTPUT_CASE_IMAGE",
            len(document.inline_shapes) >= template_shape_count + 1,
            expected={"min": template_shape_count + 1},
            actual=len(document.inline_shapes),
        )

    _check_qa(
        check,
        qa_report,
        output.get("qa") or {},
        actual_docx_sha=actual_sha,
        expected_page_count=expected_page_count,
    )
    return _result(contract, path, checks, errors, docx_sha256=actual_sha)


def _check_structure(
    check, document: DocumentObject, expected: Mapping[str, Any], prefix: str
) -> None:
    fields = {
        "table_count": len(document.tables),
        "section_count": len(document.sections),
        "inline_shape_count": len(document.inline_shapes),
    }
    for field, actual in fields.items():
        if field not in expected:
            continue
        check(
            f"{prefix}_{field.upper()}",
            actual == int(expected[field]),
            expected=int(expected[field]),
            actual=actual,
        )
    if "inline_shape_count_min" in expected:
        minimum = int(expected["inline_shape_count_min"])
        check(
            f"{prefix}_INLINE_SHAPE_COUNT_MIN",
            fields["inline_shape_count"] >= minimum,
            expected={"min": minimum},
            actual=fields["inline_shape_count"],
        )


def _check_text_contract(
    check,
    visible: str,
    *,
    required: Iterable[Any],
    forbidden: Iterable[Any],
    prefix: str,
) -> None:
    compact_visible = _compact(visible)
    for value in required:
        text = str(value or "")
        if not text:
            continue
        check(
            f"{prefix}_REQUIRED_TEXT",
            _compact(text) in compact_visible,
            expected=text,
            actual="present" if _compact(text) in compact_visible else "missing",
        )
    folded = compact_visible.casefold()
    for value in forbidden:
        text = str(value or "")
        if not text:
            continue
        present = _compact(text).casefold() in folded
        check(
            f"{prefix}_FORBIDDEN_TEXT",
            not present,
            expected="absent",
            actual=text if present else "absent",
        )


def _check_dynamic_scope(
    check,
    document: DocumentObject,
    contract: Mapping[str, Any],
    prefix: str,
) -> None:
    """Check patient-facing sections without treating appendix citations as claims."""

    if not contract:
        return
    end_heading = str(contract.get("end_heading") or "")
    blocks: list[str] = []
    boundary_found = False
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            text = "".join(node.text or "" for node in child.iter(qn("w:t")))
            if _compact(text) == _compact(end_heading):
                boundary_found = True
                break
            if text:
                blocks.append(text)
        elif child.tag == qn("w:tbl"):
            text = "".join(node.text or "" for node in child.iter(qn("w:t")))
            if text:
                blocks.append(text)
    check(
        f"{prefix}_DYNAMIC_SCOPE_BOUNDARY",
        boundary_found,
        expected=end_heading,
        actual="present" if boundary_found else "missing",
    )
    if not boundary_found:
        return
    _check_text_contract(
        check,
        "\n".join(blocks),
        required=contract.get("required_texts") or [],
        forbidden=contract.get("forbidden_texts") or [],
        prefix=f"{prefix}_DYNAMIC",
    )


def _check_gene_list(
    check,
    document: DocumentObject,
    contract: Mapping[str, Any],
    prefix: str,
) -> None:
    if not contract:
        return
    header = str(contract.get("header") or "")
    matches = [
        table
        for table in document.tables
        if table.rows and header and header in table.rows[0].cells[0].text
    ]
    check(
        f"{prefix}_GENE_LIST_TABLE",
        len(matches) == 1,
        expected=1,
        actual=len(matches),
    )
    if len(matches) != 1:
        return
    genes = [
        _compact(cell.text)
        for row in matches[0].rows[1:]
        for cell in row.cells
        if _compact(cell.text)
    ]
    expected_count = int(contract.get("count") or 0)
    expected_sha = str(contract.get("ordered_sha256") or "").lower()
    actual_sha = hashlib.sha256("\n".join(genes).encode("utf-8")).hexdigest()
    check(
        f"{prefix}_GENE_LIST_COUNT",
        len(genes) == expected_count,
        expected=expected_count,
        actual=len(genes),
    )
    check(
        f"{prefix}_GENE_LIST_SHA256",
        actual_sha == expected_sha,
        expected=expected_sha,
        actual=actual_sha,
    )


def _check_qa(
    check,
    qa_report: Mapping[str, Any] | None,
    contract: Mapping[str, Any],
    *,
    actual_docx_sha: str,
    expected_page_count: int | None,
) -> None:
    if not contract:
        return
    if not isinstance(qa_report, Mapping):
        check("OUTPUT_QA_PRESENT", False, expected="mapping", actual=None)
        return
    check("OUTPUT_QA_PRESENT", True, expected="mapping", actual="mapping")

    configured_statuses = contract.get("allowed_statuses")
    if isinstance(configured_statuses, (list, tuple, set)):
        allowed_statuses = {
            str(value).strip().upper()
            for value in configured_statuses
            if str(value).strip()
        }
    else:
        allowed_statuses = {str(contract.get("status") or "PASS").strip().upper()}
    actual_status = str(qa_report.get("status") or "")
    check(
        "OUTPUT_QA_STATUS",
        actual_status.upper() in allowed_statuses,
        expected=sorted(allowed_statuses),
        actual=actual_status,
    )
    metrics = qa_report.get("metrics") or {}
    qa_sha = str(metrics.get("output_sha256") or "").lower()
    check(
        "OUTPUT_QA_DOCX_SHA256",
        qa_sha == actual_docx_sha,
        expected=actual_docx_sha,
        actual=qa_sha,
    )

    visual = (qa_report.get("checks") or {}).get("visual_render") or {}
    check(
        "OUTPUT_VISUAL_QA_STATUS",
        visual.get("status") == "PASS",
        expected="PASS",
        actual=visual.get("status"),
    )
    check(
        "OUTPUT_VISUAL_QA_SCOPE",
        visual.get("requested") == str(contract.get("visual_render") or "all")
        and visual.get("required") is True,
        expected={
            "requested": str(contract.get("visual_render") or "all"),
            "required": True,
        },
        actual={
            "requested": visual.get("requested"),
            "required": visual.get("required"),
        },
    )
    pixel = visual.get("pixel_check") or {}
    pages = pixel.get("checked_pages")
    if pages is None:
        pages = metrics.get("visual_render_page_count")
    if expected_page_count is not None:
        check(
            "OUTPUT_PAGE_COUNT_EXACT",
            pages == expected_page_count,
            expected=expected_page_count,
            actual=pages,
        )
    else:
        minimum = int(contract.get("page_count_min") or 0)
        maximum = int(contract.get("page_count_max") or 0)
        valid = (
            isinstance(pages, int)
            and pages >= minimum
            and (maximum <= 0 or pages <= maximum)
        )
        check(
            "OUTPUT_PAGE_COUNT_RANGE",
            valid,
            expected={"min": minimum, "max": maximum or None},
            actual=pages,
        )

    blank_pages = list(pixel.get("blank_pages") or [])
    low_content_pages = list(pixel.get("unexpected_low_content_pages") or [])
    check(
        "OUTPUT_BLANK_PAGES",
        len(blank_pages) <= int(contract.get("max_blank_pages") or 0),
        expected={"max": int(contract.get("max_blank_pages") or 0)},
        actual=blank_pages,
    )
    check(
        "OUTPUT_LOW_CONTENT_PAGES",
        len(low_content_pages) <= int(contract.get("max_low_content_pages") or 0),
        expected={"max": int(contract.get("max_low_content_pages") or 0)},
        actual=low_content_pages,
    )


def _result(
    contract: Mapping[str, Any],
    path: Path,
    checks: list[dict[str, Any]],
    errors: list[dict[str, Any]],
    *,
    docx_sha256: str | None = None,
) -> dict[str, Any]:
    return {
        "schema_version": "1.0",
        "status": "FAIL" if errors else "PASS",
        "contract_id": contract.get("contract_id"),
        "panel_id": contract.get("panel_id"),
        "docx": str(path),
        "docx_sha256": docx_sha256,
        "lifecycle": dict(contract.get("lifecycle") or {}),
        "summary": {
            "check_count": len(checks),
            "pass_count": sum(item["status"] == "PASS" for item in checks),
            "fail_count": len(errors),
        },
        "checks": checks,
        "errors": errors,
    }


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _compact(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or ""))
