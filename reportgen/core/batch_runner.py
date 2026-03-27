"""
Batch pipeline runner (Excel -> DOCX) with validations and optional post-processing.

This module provides a reusable implementation for:
  - batch generation
  - template auto-detection
  - template contract validation (fail/warn)
  - output validations (openable, placeholder residue, key tables)
  - reproducibility artifacts (masked context + meta)
  - highlighting dynamic regions in generated outputs
  - optional page rendering to PNGs

It is intended to be called by:
  - CLI (`reportgen batch-validate`)
  - thin wrappers in `scripts/`
"""

from __future__ import annotations

import json
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

from reportgen.config.loader import ConfigLoader
from reportgen.core.excel_reader import ExcelReader
from reportgen.core.project_detector import ProjectDetector
from reportgen.core.report_generator import ReportGenerator
from reportgen.utils.artifacts import build_meta, write_json
from reportgen.utils.docx_highlighter import highlight_rendered_docx
from reportgen.utils.docx_render import render_docx_to_pngs
from reportgen.utils.privacy import mask_sensitive_data


def _now_ts() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def iter_excel_inputs(inputs: List[str]) -> List[Path]:
    excel_files: List[Path] = []
    for spec in inputs:
        p = Path(spec)
        if p.exists() and p.is_dir():
            excel_files.extend(sorted(p.glob("*.xlsx")))
            continue
        if p.exists() and p.is_file():
            excel_files.append(p)
            continue
        # glob pattern
        excel_files.extend(sorted(Path(".").glob(spec)))

    # normalize + dedupe + filter temp files
    seen: set[Path] = set()
    result: List[Path] = []
    for p in excel_files:
        if p.name.startswith("~$"):
            continue
        rp = p.resolve()
        if rp in seen:
            continue
        seen.add(rp)
        result.append(rp)
    return sorted(result)


def _read_docx_text(doc) -> str:
    parts: List[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _count_placeholders(text: str) -> Dict[str, int]:
    return {
        "{{": text.count("{{"),
        "}}": text.count("}}"),
        "{%": text.count("{%"),
        "%}": text.count("%}"),
    }


def _detect_tables(doc) -> Dict[str, bool]:
    """Heuristic checks for key tables in the aligned template family."""

    def any_row_contains(table, *tokens: str, max_rows: int = 3) -> bool:
        for i in range(min(len(table.rows), max_rows)):
            row_text = " ".join(
                (cell.text or "").strip() for cell in table.rows[i].cells
            )
            if all(t in row_text for t in tokens):
                return True
        return False

    found = {
        "targeted_drug_tips": False,
        "variants_2_1": False,
        "cnv": False,
        "fusion": False,
        "hla": False,
    }

    for tbl in doc.tables:
        if any_row_contains(tbl, "突变位点", "潜在获益"):
            found["targeted_drug_tips"] = True
        # variants_2_1 table has a 2-row header (merged cells in row 0,
        # real labels in row 1).
        if any_row_contains(tbl, "转录本号", "染色体") and any_row_contains(
            tbl, "频率", max_rows=4
        ):
            found["variants_2_1"] = True
        if any_row_contains(tbl, "起始位置", "终止位置", "拷贝数"):
            found["cnv"] = True
        if any_row_contains(tbl, "基因1", "基因2") and any_row_contains(tbl, "断点"):
            found["fusion"] = True
        if any_row_contains(tbl, "HLA位点") and (
            any_row_contains(tbl, "Type 1") or any_row_contains(tbl, "Type1")
        ):
            found["hla"] = True

    return found


@dataclass(frozen=True)
class BatchValidateOptions:
    inputs: List[str]
    name_contains: Optional[str] = None
    template: Optional[str] = None
    config_dir: str = "config"
    output_root: Optional[str] = None
    max_files: Optional[int] = None
    render: str = "none"  # none|first|all
    render_dpi: int = 120
    highlight: bool = False
    highlight_color: str = "D9EAF7"
    highlight_output_root: Optional[str] = None
    emit_context: bool = True
    emit_meta: bool = True
    artifacts_dir_mode: str = "separate"  # separate|alongside
    template_contract: str = "fail"  # none|warn|fail
    log_level: str = "CRITICAL"
    report_file: Optional[str] = None
    show_paths: bool = False


@dataclass(frozen=True)
class BatchValidateRun:
    report: Dict[str, Any]
    report_path: Path
    output_root: Path


def run_batch_generate_validate(
    opts: BatchValidateOptions,
    *,
    progress: Optional[Callable[[str], None]] = None,
) -> BatchValidateRun:
    excel_files = iter_excel_inputs(list(opts.inputs))
    if opts.name_contains:
        token = str(opts.name_contains).lower()
        excel_files = [p for p in excel_files if token in p.name.lower()]
    if opts.max_files is not None:
        excel_files = excel_files[: max(0, int(opts.max_files))]

    if not excel_files:
        raise FileNotFoundError("No Excel files found for the given inputs.")

    output_root = (
        Path(opts.output_root)
        if opts.output_root
        else Path("data/output") / f"batch_validate_{_now_ts()}"
    )
    output_root.mkdir(parents=True, exist_ok=True)
    output_root_resolved = output_root.resolve()

    report_path = (
        Path(opts.report_file)
        if opts.report_file
        else output_root / "validation_report.json"
    )

    def public_path(path: Path) -> str:
        rp = path.resolve()
        if opts.show_paths:
            return str(rp)
        try:
            return str(rp.relative_to(output_root_resolved))
        except Exception:
            return rp.name

    # Sensitive fields for artifact masking.
    config_loader = ConfigLoader(config_dir=opts.config_dir, log_level=opts.log_level)
    sensitive_fields = set(
        config_loader.get_setting("logging.content.sensitive_fields", []) or []
    )
    sensitive_fields.add("sample_id")

    results: List[Dict[str, Any]] = []
    successes = 0
    failures = 0

    # Lazy imports to keep CLI startup light.
    from docx import Document  # python-docx

    excel_reader = ExcelReader(config_dir=opts.config_dir, log_level=opts.log_level)
    detector = ProjectDetector(config_dir=opts.config_dir, log_level=opts.log_level)
    generator = ReportGenerator(config_dir=opts.config_dir, log_level=opts.log_level)

    for idx, excel_path in enumerate(excel_files, start=1):
        start = time.time()
        output_filename = f"report_{idx:03d}.docx"
        output_docx = str((output_root / output_filename).resolve())

        template_path: Optional[str] = opts.template
        errors: List[str] = []
        warnings: List[str] = []
        validation: Dict[str, Any] = {}
        artifacts: Dict[str, Any] = {}

        try:
            excel_data = excel_reader.read(str(excel_path), include_tables=True)

            # 始终尝试检测项目类型（用于 enhancer 分派），
            # 仅在未指定模板时才从检测结果中取模板路径
            det = detector.detect(str(excel_path), excel_data=excel_data)
            detected_project_type = det.get("project_type")
            detected_project_name = det.get("project_name") if det.get("detected") else None
            if not template_path:
                template_path = det.get("template")

            if not template_path:
                raise RuntimeError(
                    "No template resolved (auto-detect failed and "
                    "--template not provided)."
                )
            if not Path(template_path).exists():
                raise FileNotFoundError(f"Template not found: {template_path}")

            gen = generator.generate(
                excel_file=str(excel_path),
                template_file=str(template_path),
                output_dir=str(output_root),
                output_filename=output_filename,
                strict_mode=False,
                excel_data=excel_data,
                return_context=bool(opts.emit_context or opts.emit_meta),
                template_contract_mode=str(opts.template_contract),
                project_type=detected_project_type,
                project_name=detected_project_name,
            )

            if not gen.get("success"):
                errors.extend([str(e) for e in gen.get("errors") or []])
            warnings.extend([str(w) for w in gen.get("warnings") or []])

            if gen.get("output_file"):
                output_docx = str(Path(gen["output_file"]).resolve())

            # --- artifacts: context/meta ---
            template_contract = gen.get("template_contract")
            if opts.emit_context or opts.emit_meta:
                if opts.artifacts_dir_mode == "alongside":
                    artifacts_dir = Path(output_docx).resolve().parent
                else:
                    artifacts_dir = (output_root / "artifacts" / f"{idx:03d}").resolve()
                artifacts_dir.mkdir(parents=True, exist_ok=True)

                if opts.emit_context:
                    ctx = gen.get("context") or {}
                    masked = mask_sensitive_data(ctx, sensitive_keys=sensitive_fields)
                    context_path = artifacts_dir / "context.json"
                    write_json(context_path, masked)
                    artifacts["context_json"] = public_path(context_path)

                if opts.emit_meta:
                    meta = build_meta(
                        excel_path=Path(str(excel_path)).resolve(),
                        template_path=Path(str(template_path)).resolve(),
                        output_docx=(
                            Path(output_docx).resolve() if output_docx else None
                        ),
                        config_dir=Path(opts.config_dir).resolve(),
                        args={
                            "inputs": list(opts.inputs),
                            "name_contains": opts.name_contains,
                            "template": opts.template,
                            "config_dir": opts.config_dir,
                            "output_root": str(output_root),
                            "max_files": opts.max_files,
                            "render": opts.render,
                            "render_dpi": opts.render_dpi,
                            "highlight": opts.highlight,
                            "highlight_color": opts.highlight_color,
                            "highlight_output_root": opts.highlight_output_root,
                            "emit_context": opts.emit_context,
                            "emit_meta": opts.emit_meta,
                            "artifacts_dir_mode": opts.artifacts_dir_mode,
                            "template_contract": opts.template_contract,
                            "log_level": opts.log_level,
                        },
                        template_contract=template_contract,
                        include_paths=bool(opts.show_paths),
                    )
                    meta_path = artifacts_dir / "meta.json"
                    write_json(meta_path, meta)
                    artifacts["meta_json"] = public_path(meta_path)

                artifacts["artifacts_dir"] = public_path(artifacts_dir)

            # --- validations on output docx ---
            doc_ok = False
            placeholder_counts: Dict[str, int] = {}
            table_presence: Dict[str, bool] = {}
            doc_stats: Dict[str, Any] = {}
            issues: List[str] = []

            try:
                doc = Document(output_docx)
                doc_ok = True
                doc_stats = {
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "sections": len(doc.sections),
                    "file_size_bytes": (
                        Path(output_docx).stat().st_size
                        if Path(output_docx).exists()
                        else None
                    ),
                }
                full_text = _read_docx_text(doc)
                placeholder_counts = _count_placeholders(full_text)
                table_presence = _detect_tables(doc)

                if any(v > 0 for v in placeholder_counts.values()):
                    issues.append("unrendered_placeholders")

            except Exception as e:
                issues.append("docx_open_failed")
                errors.append(f"docx_open_failed: {e}")

            # Expectation checks based on Excel content (only when Excel was parsed).
            expected = {
                "cnv": bool(excel_data.get_table_data("Cnv")),
                "fusion": bool(excel_data.get_table_data("Fusion")),
                "hla": bool(excel_data.get_table_data("HLA")),
            }
            missing_expected = []
            for k, exp in expected.items():
                if exp and doc_ok and not table_presence.get(k, False):
                    missing_expected.append(k)
            if missing_expected:
                issues.append(
                    f"missing_expected_tables:{','.join(sorted(missing_expected))}"
                )

            # Optional rendering
            rendered_pages: List[str] = []
            if doc_ok and opts.render != "none":
                pages_dir = (output_root / "pages" / f"{idx:03d}").resolve()
                pages_dir.mkdir(parents=True, exist_ok=True)

                first_page = 1 if opts.render == "first" else None
                last_page = 1 if opts.render == "first" else None

                try:
                    pngs = render_docx_to_pngs(
                        Path(output_docx),
                        output_dir=pages_dir,
                        dpi=int(opts.render_dpi),
                        keep_pdf=False,
                        first_page=first_page,
                        last_page=last_page,
                    )
                    rendered_pages = [
                        p.name for p in pngs if not p.name.startswith("._")
                    ]
                except Exception as e:
                    issues.append("render_failed")
                    errors.append(f"render_failed: {e}")

            # Optional highlighting (writes copies under output/doc/)
            highlight_info: Dict[str, Any] = {}
            if doc_ok and opts.highlight:
                try:
                    highlight_root = (
                        Path(opts.highlight_output_root).resolve()
                        if opts.highlight_output_root
                        else (
                            Path("output/doc/highlighted") / output_root.name
                        ).resolve()
                    )
                    highlight_root.mkdir(parents=True, exist_ok=True)

                    out_name = f"{Path(output_docx).stem}_highlighted.docx"
                    highlighted_docx = (highlight_root / out_name).resolve()
                    summary = highlight_rendered_docx(
                        template_path=str(template_path),
                        input_docx_path=str(output_docx),
                        output_docx_path=str(highlighted_docx),
                        color=str(opts.highlight_color),
                        skip_empty=True,
                    )
                    summary_obj = dict(summary.__dict__)
                    if not opts.show_paths:
                        for key in ("template", "input_docx", "output_docx"):
                            try:
                                summary_obj[key] = Path(
                                    str(summary_obj.get(key) or "")
                                ).name
                            except Exception:
                                summary_obj[key] = ""
                    highlight_info = {
                        "highlighted_docx": (
                            str(highlighted_docx)
                            if opts.show_paths
                            else Path(str(highlighted_docx)).name
                        ),
                        "summary": summary_obj,
                    }
                except Exception as e:
                    issues.append("highlight_failed")
                    errors.append(f"highlight_failed: {e}")

            validation = {
                "docx_openable": doc_ok,
                "doc_stats": doc_stats,
                "placeholders": placeholder_counts,
                "tables_found": table_presence,
                "tables_expected_from_excel": expected,
                "rendered_pages": rendered_pages,
                "issues": issues,
                "template_contract": template_contract,
                "artifacts": artifacts,
                "highlight": highlight_info,
            }

            ok = (
                (len(errors) == 0)
                and doc_ok
                and (not any(v > 0 for v in placeholder_counts.values()))
            )

        except Exception as e:
            ok = False
            errors.append(str(e))

        duration = time.time() - start
        if ok:
            successes += 1
        else:
            failures += 1

        results.append(
            {
                "index": idx,
                "ok": ok,
                "output_docx": (
                    output_docx
                    if (opts.show_paths and Path(output_docx).exists())
                    else (
                        Path(output_docx).name if Path(output_docx).exists() else None
                    )
                ),
                "template": (
                    str(template_path)
                    if (opts.show_paths and template_path)
                    else (Path(str(template_path)).name if template_path else None)
                ),
                "duration_seconds": duration,
                "errors": errors,
                "warnings": warnings,
                "validation": validation,
            }
        )

        if progress:
            if opts.show_paths:
                progress(
                    f"[{idx}/{len(excel_files)}] ok={ok} excel={excel_path} "
                    f"output={output_docx}"
                )
            else:
                progress(
                    f"[{idx}/{len(excel_files)}] ok={ok} "
                    f"output={Path(output_docx).name}"
                )

    report_obj = {
        "generated_at": datetime.now().isoformat(),
        "inputs_count": len(excel_files),
        "successes": successes,
        "failures": failures,
        "output_root": (
            str(output_root_resolved) if opts.show_paths else str(output_root)
        ),
        "results": results,
    }

    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(
        json.dumps(report_obj, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    return BatchValidateRun(
        report=report_obj,
        report_path=report_path.resolve(),
        output_root=output_root.resolve(),
    )
