"""DOCX report diff utilities for regression review.

The diff is deliberately structural rather than visual: it compares extracted
text, table shape/content, basic run styles, and optional QA JSON reports. This
gives the pipeline a deterministic gate even when local Word/LibreOffice
rendering is unavailable.
"""

from __future__ import annotations

import json
import re
from collections import Counter
from dataclasses import dataclass
from datetime import datetime
from difflib import SequenceMatcher, unified_diff
from pathlib import Path
from typing import Any, Mapping, Optional

from docx import Document
from docx.oxml.ns import qn

from reportgen.utils.artifacts import write_json


@dataclass(frozen=True)
class ReportDiffOptions:
    reference_docx: str
    candidate_docx: str
    output_dir: Optional[str] = None
    reference_qa: Optional[str] = None
    candidate_qa: Optional[str] = None
    max_samples: int = 30
    normalize_whitespace: bool = False
    ignore_reference_artifacts: bool = False
    style_metric_policy: str = "warn"


def compare_reports(options: ReportDiffOptions) -> dict[str, Any]:
    """Compare two generated DOCX reports and optional QA sidecars."""
    reference_path = Path(options.reference_docx).resolve()
    candidate_path = Path(options.candidate_docx).resolve()
    max_samples = max(1, int(options.max_samples))

    ref_snapshot = _snapshot_docx(reference_path)
    cand_snapshot = _snapshot_docx(candidate_path)

    sections = {
        "documents": _compare_document_openability(
            ref_snapshot, cand_snapshot, reference_path, candidate_path
        ),
        "text": _compare_text(
            ref_snapshot,
            cand_snapshot,
            max_samples=max_samples,
            normalize_whitespace=options.normalize_whitespace,
            ignore_reference_artifacts=options.ignore_reference_artifacts,
        ),
        "part3": _compare_part3_sections(
            ref_snapshot,
            cand_snapshot,
            max_samples=max_samples,
        ),
        "tables": _compare_tables(
            ref_snapshot,
            cand_snapshot,
            max_samples=max_samples,
            normalize_whitespace=options.normalize_whitespace,
        ),
        "styles": _compare_styles(
            ref_snapshot,
            cand_snapshot,
            max_samples=max_samples,
            policy=options.style_metric_policy,
        ),
        "qa": _compare_qa(
            _load_qa(options.reference_qa, reference_path),
            _load_qa(options.candidate_qa, candidate_path),
            max_samples=max_samples,
        ),
    }
    issues = _collect_issues(sections)
    status = _overall_status(issues)
    result = {
        "schema_version": "1.0",
        "generated_at": datetime.now().isoformat(),
        "status": status,
        "reference_docx": str(reference_path),
        "candidate_docx": str(candidate_path),
        "summary": {
            "failures": sum(1 for item in issues if item["level"] == "error"),
            "warnings": sum(1 for item in issues if item["level"] == "warning"),
            "text_similarity": sections["text"].get("similarity"),
            "table_count": {
                "reference": ref_snapshot.get("table_count"),
                "candidate": cand_snapshot.get("table_count"),
            },
        },
        "issues": issues,
        "sections": sections,
    }

    if options.output_dir:
        output_dir = Path(options.output_dir).resolve()
        write_report_diff_outputs(result, output_dir)
        result["output_dir"] = str(output_dir)
        result["json_file"] = str(output_dir / "report_diff.json")
        result["markdown_file"] = str(output_dir / "report_diff.md")
    return result


def write_report_diff_outputs(result: Mapping[str, Any], output_dir: Path) -> None:
    """Write report_diff.json and report_diff.md."""
    output_dir.mkdir(parents=True, exist_ok=True)
    write_json(output_dir / "report_diff.json", dict(result))
    (output_dir / "report_diff.md").write_text(
        render_report_diff_markdown(result), encoding="utf-8"
    )


def render_report_diff_markdown(result: Mapping[str, Any]) -> str:
    """Render a compact human-readable Markdown diff summary."""
    summary = result.get("summary") or {}
    sections = result.get("sections") or {}
    lines = [
        "# Report Diff",
        "",
        f"- Status: **{result.get('status')}**",
        f"- Reference: `{Path(str(result.get('reference_docx'))).name}`",
        f"- Candidate: `{Path(str(result.get('candidate_docx'))).name}`",
        f"- Failures: {summary.get('failures', 0)}",
        f"- Warnings: {summary.get('warnings', 0)}",
        f"- Text similarity: {summary.get('text_similarity')}",
        "",
        "## Issues",
        "",
    ]
    issues = list(result.get("issues") or [])
    if not issues:
        lines.append("No blocking or warning issues detected.")
    else:
        for issue in issues:
            lines.append(
                f"- **{issue.get('level', '').upper()}** `{issue.get('code')}`: "
                f"{issue.get('message')}"
            )

    text_section = sections.get("text") or {}
    lines.extend(["", "## Text", ""])
    lines.append(f"- Paragraph similarity: {text_section.get('similarity')}")
    lines.append(f"- Changed blocks: {text_section.get('changed_block_count', 0)}")
    for sample in text_section.get("samples") or []:
        lines.append(
            f"- `{sample.get('tag')}` reference[{sample.get('reference_range')}] "
            f"candidate[{sample.get('candidate_range')}]"
        )

    table_section = sections.get("tables") or {}
    lines.extend(["", "## Tables", ""])
    lines.append(
        f"- Table count: {table_section.get('reference_count')} -> "
        f"{table_section.get('candidate_count')}"
    )
    for sample in table_section.get("samples") or []:
        lines.append(f"- {sample.get('message')}")

    part3_section = sections.get("part3") or {}
    lines.extend(["", "## Part 3", ""])
    lines.append(
        f"- Gene sections: {part3_section.get('reference_gene_count')} -> "
        f"{part3_section.get('candidate_gene_count')}"
    )
    lines.append(
        f"- Drug sections: {part3_section.get('reference_drug_count')} -> "
        f"{part3_section.get('candidate_drug_count')}"
    )
    for sample in part3_section.get("samples") or []:
        lines.append(f"- {sample.get('message')}")

    style_section = sections.get("styles") or {}
    lines.extend(["", "## Styles", ""])
    for sample in style_section.get("samples") or []:
        lines.append(f"- {sample.get('message')}")

    qa_section = sections.get("qa") or {}
    lines.extend(["", "## QA", ""])
    lines.append(
        f"- QA status: {qa_section.get('reference_status')} -> "
        f"{qa_section.get('candidate_status')}"
    )
    for sample in qa_section.get("samples") or []:
        lines.append(f"- {sample.get('message')}")

    return "\n".join(lines).rstrip() + "\n"


def _snapshot_docx(path: Path) -> dict[str, Any]:
    snapshot: dict[str, Any] = {"path": str(path), "exists": path.exists()}
    if not path.exists():
        snapshot["openable"] = False
        snapshot["error"] = "file_not_found"
        return snapshot
    try:
        doc = Document(str(path))
    except Exception as exc:
        snapshot["openable"] = False
        snapshot["error"] = str(exc)
        return snapshot

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = [_table_snapshot(table, idx) for idx, table in enumerate(doc.tables)]
    snapshot.update(
        {
            "openable": True,
            "paragraphs": paragraphs,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "tables": tables,
            "style_summary": _style_summary(doc),
        }
    )
    return snapshot


def _table_snapshot(table: Any, index: int) -> dict[str, Any]:
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    row_count = len(rows)
    col_count = max((len(row) for row in rows), default=0)
    header = rows[0] if rows else []
    return {
        "index": index,
        "rows": rows,
        "row_count": row_count,
        "col_count": col_count,
        "header": header,
        "text": "\n".join("\t".join(row) for row in rows),
        "style_summary": _table_style_summary(table),
    }


def _style_summary(doc: Any) -> dict[str, Any]:
    counter: Counter[str] = Counter()
    for paragraph in doc.paragraphs:
        _count_paragraph_runs(counter, paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _count_paragraph_runs(counter, paragraph)
    return dict(sorted(counter.items()))


def _table_style_summary(table: Any) -> dict[str, Any]:
    counter: Counter[str] = Counter()
    for row in table.rows:
        for cell in row.cells:
            borders = cell._tc.xpath("./w:tcPr/w:tcBorders/*")
            if borders:
                counter["cell_border_elements"] += len(borders)
            for paragraph in cell.paragraphs:
                _count_paragraph_runs(counter, paragraph)
    return dict(sorted(counter.items()))


def _count_paragraph_runs(counter: Counter[str], paragraph: Any) -> None:
    if paragraph.style is not None and paragraph.style.name:
        counter[f"paragraph_style:{paragraph.style.name}"] += 1
    for run in paragraph.runs:
        text = run.text or ""
        if not text.strip():
            continue
        font_name = run.font.name
        if font_name:
            counter[f"font:{font_name}"] += 1
        size = run.font.size.pt if run.font.size else None
        if size is not None:
            counter[f"font_size:{round(size, 1)}"] += 1
        if run.bold:
            counter["bold_runs"] += 1
        if run.italic:
            counter["italic_runs"] += 1
        if _is_underlined(run):
            counter["underlined_runs"] += 1
        color = run.font.color.rgb
        if color is not None:
            counter[f"color:{str(color).upper()}"] += 1
            if str(color).upper() not in {"000000", "FFFFFF"}:
                counter["non_black_non_white_color_runs"] += 1


def _is_underlined(run: Any) -> bool:
    if run.font.underline:
        return True
    for underline in run._r.xpath(".//w:u"):
        value = underline.get(qn("w:val"))
        if value not in {None, "none", "0", "false"}:
            return True
    return False


def _compare_document_openability(
    ref: Mapping[str, Any],
    cand: Mapping[str, Any],
    reference_path: Path,
    candidate_path: Path,
) -> dict[str, Any]:
    status = "PASS" if ref.get("openable") and cand.get("openable") else "FAIL"
    issues = []
    if not ref.get("openable"):
        issues.append(
            {
                "level": "error",
                "code": "REFERENCE_DOCX_NOT_OPENABLE",
                "message": f"Reference DOCX cannot be opened: {reference_path}",
            }
        )
    if not cand.get("openable"):
        issues.append(
            {
                "level": "error",
                "code": "CANDIDATE_DOCX_NOT_OPENABLE",
                "message": f"Candidate DOCX cannot be opened: {candidate_path}",
            }
        )
    return {"status": status, "issues": issues}


def _compare_text(
    ref: Mapping[str, Any],
    cand: Mapping[str, Any],
    *,
    max_samples: int,
    normalize_whitespace: bool = False,
    ignore_reference_artifacts: bool = False,
) -> dict[str, Any]:
    ref_paras = list(ref.get("paragraphs") or [])
    cand_paras = list(cand.get("paragraphs") or [])
    if ignore_reference_artifacts:
        ref_paras = [text for text in ref_paras if not _is_reference_artifact(text)]
        cand_paras = [text for text in cand_paras if not _is_reference_artifact(text)]
    if normalize_whitespace:
        ref_paras = [_normalize_soft_whitespace(text) for text in ref_paras]
        cand_paras = [_normalize_soft_whitespace(text) for text in cand_paras]
    matcher = SequenceMatcher(a=ref_paras, b=cand_paras, autojunk=False)
    samples = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        samples.append(
            {
                "tag": tag,
                "reference_range": [i1, i2],
                "candidate_range": [j1, j2],
                "reference_sample": ref_paras[i1:i2][:3],
                "candidate_sample": cand_paras[j1:j2][:3],
            }
        )
        if len(samples) >= max_samples:
            break

    similarity = round(matcher.ratio(), 6)
    status = "PASS" if similarity == 1.0 else "WARN"
    issues = []
    if status == "WARN":
        issues.append(
            {
                "level": "warning",
                "code": "TEXT_DIFF",
                "message": f"DOCX text differs; similarity={similarity}.",
            }
        )
    return {
        "status": status,
        "similarity": similarity,
        "reference_paragraph_count": len(ref_paras),
        "candidate_paragraph_count": len(cand_paras),
        "changed_block_count": len([op for op in matcher.get_opcodes() if op[0] != "equal"]),
        "samples": samples,
        "unified_diff_sample": list(
            unified_diff(
                ref_paras,
                cand_paras,
                fromfile="reference",
                tofile="candidate",
                lineterm="",
            )
        )[: max_samples * 4],
        "issues": issues,
    }


def _compare_tables(
    ref: Mapping[str, Any],
    cand: Mapping[str, Any],
    *,
    max_samples: int,
    normalize_whitespace: bool = False,
) -> dict[str, Any]:
    ref_tables = list(ref.get("tables") or [])
    cand_tables = list(cand.get("tables") or [])
    issues = []
    samples = []
    if len(ref_tables) != len(cand_tables):
        issues.append(
            {
                "level": "error",
                "code": "TABLE_COUNT_DIFF",
                "message": f"Table count differs: {len(ref_tables)} -> {len(cand_tables)}.",
            }
        )
        samples.append(
            {
                "message": f"Table count differs: {len(ref_tables)} -> {len(cand_tables)}",
            }
        )

    for idx, (ref_table, cand_table) in enumerate(zip(ref_tables, cand_tables)):
        ref_shape = [ref_table["row_count"], ref_table["col_count"]]
        cand_shape = [cand_table["row_count"], cand_table["col_count"]]
        if ref_shape != cand_shape:
            issues.append(
                {
                    "level": "error",
                    "code": "TABLE_SHAPE_DIFF",
                    "message": f"Table {idx} shape differs: {ref_shape} -> {cand_shape}.",
                }
            )
            samples.append(
                {
                    "table_index": idx,
                    "message": f"Table {idx} shape differs: {ref_shape} -> {cand_shape}",
                }
            )
            if len(samples) >= max_samples:
                break
            continue
        ref_header = _normalize_row(ref_table.get("header") or [], normalize_whitespace)
        cand_header = _normalize_row(cand_table.get("header") or [], normalize_whitespace)
        if ref_header != cand_header:
            issues.append(
                {
                    "level": "warning",
                    "code": "TABLE_HEADER_DIFF",
                    "message": f"Table {idx} header differs.",
                }
            )
            samples.append(
                {
                    "table_index": idx,
                    "message": f"Table {idx} header differs",
                    "reference": ref_header,
                    "candidate": cand_header,
                }
            )
        cell_diffs = _table_cell_diffs(
            ref_table["rows"],
            cand_table["rows"],
            max_samples,
            normalize_whitespace=normalize_whitespace,
        )
        if cell_diffs:
            issues.append(
                {
                    "level": "warning",
                    "code": "TABLE_CELL_DIFF",
                    "message": f"Table {idx} has {len(cell_diffs)} sampled cell differences.",
                }
            )
            samples.extend(
                {
                    "table_index": idx,
                    "message": f"Table {idx} cell {item['cell']} differs",
                    **item,
                }
                for item in cell_diffs
            )
        if len(samples) >= max_samples:
            break

    return {
        "status": _section_status(issues),
        "reference_count": len(ref_tables),
        "candidate_count": len(cand_tables),
        "samples": samples[:max_samples],
        "issues": issues,
    }


def _compare_part3_sections(
    ref: Mapping[str, Any],
    cand: Mapping[str, Any],
    *,
    max_samples: int,
) -> dict[str, Any]:
    ref_part3 = _extract_part3_sections(list(ref.get("paragraphs") or []))
    cand_part3 = _extract_part3_sections(list(cand.get("paragraphs") or []))
    issues: list[dict[str, Any]] = []
    samples: list[dict[str, Any]] = []

    if not ref_part3["present"] and not cand_part3["present"]:
        return {
            "status": "PASS",
            "present": False,
            "reference_gene_count": 0,
            "candidate_gene_count": 0,
            "reference_drug_count": 0,
            "candidate_drug_count": 0,
            "samples": [],
            "issues": [],
        }

    if ref_part3["present"] != cand_part3["present"]:
        issues.append(
            {
                "level": "error",
                "code": "PART3_PRESENCE_DIFF",
                "message": "Part 3 presence differs between reference and candidate.",
            }
        )

    for section_name, code in (
        ("gene_sections", "PART3_GENE_SECTION_DIFF"),
        ("drug_sections", "PART3_DRUG_SECTION_DIFF"),
    ):
        ref_sections = ref_part3.get(section_name) or {}
        cand_sections = cand_part3.get(section_name) or {}
        ref_keys = set(ref_sections)
        cand_keys = set(cand_sections)
        missing = sorted(ref_keys - cand_keys)
        extra = sorted(cand_keys - ref_keys)
        if missing or extra:
            issues.append(
                {
                    "level": "error",
                    "code": "PART3_SECTION_KEY_DIFF",
                    "message": (
                        f"{section_name} keys differ: missing={missing[:5]}, "
                        f"extra={extra[:5]}."
                    ),
                }
            )
            samples.append(
                {
                    "message": f"{section_name} keys differ",
                    "missing": missing[:max_samples],
                    "extra": extra[:max_samples],
                }
            )
        for key in sorted(ref_keys & cand_keys):
            ref_text = _normalize_part3_text(ref_sections[key]["text"])
            cand_text = _normalize_part3_text(cand_sections[key]["text"])
            if ref_text == cand_text:
                continue
            issues.append(
                {
                    "level": "error",
                    "code": code,
                    "message": f"{section_name} text differs for {key}.",
                }
            )
            if len(samples) < max_samples:
                samples.append(
                    {
                        "message": f"{section_name} text differs for {key}",
                        "key": key,
                        "reference_sample": ref_sections[key]["text"][:300],
                        "candidate_sample": cand_sections[key]["text"][:300],
                    }
                )

    return {
        "status": _section_status(issues),
        "present": bool(ref_part3["present"] or cand_part3["present"]),
        "reference_gene_count": len(ref_part3.get("gene_sections") or {}),
        "candidate_gene_count": len(cand_part3.get("gene_sections") or {}),
        "reference_drug_count": len(ref_part3.get("drug_sections") or {}),
        "candidate_drug_count": len(cand_part3.get("drug_sections") or {}),
        "samples": samples[:max_samples],
        "issues": issues,
    }


def _extract_part3_sections(paragraphs: list[str]) -> dict[str, Any]:
    start = _find_paragraph_index(
        paragraphs, "第三部分：基因变异及相应靶向/免疫药物解析"
    )
    if start is None:
        return {"present": False, "gene_sections": {}, "drug_sections": {}}
    end = len(paragraphs)
    for idx in range(start + 1, len(paragraphs)):
        text = str(paragraphs[idx]).strip()
        if text.startswith("3. 阅读说明") or text.startswith("第四部分：附录"):
            end = idx
            break

    part3 = [str(text).strip() for text in paragraphs[start:end] if str(text).strip()]
    drug_start = _find_paragraph_index(part3, "靶向药物/免疫用药提示解析")
    gene_lines = part3 if drug_start is None else part3[:drug_start]
    drug_lines = [] if drug_start is None else part3[drug_start:]
    return {
        "present": True,
        "gene_sections": _extract_part3_gene_sections(gene_lines),
        "drug_sections": _extract_part3_drug_sections(drug_lines),
    }


def _extract_part3_gene_sections(lines: list[str]) -> dict[str, dict[str, str]]:
    sections: dict[str, dict[str, str]] = {}
    current_key: Optional[str] = None
    current_lines: list[str] = []
    for line in lines:
        if _is_part3_gene_header(line):
            if current_key and current_lines:
                sections[current_key] = {"text": "\n".join(current_lines)}
            current_key = _part3_variant_key_from_header(line)
            current_lines = [line]
            continue
        if current_key:
            current_lines.append(line)
    if current_key and current_lines:
        sections[current_key] = {"text": "\n".join(current_lines)}
    return sections


def _extract_part3_drug_sections(lines: list[str]) -> dict[str, dict[str, str]]:
    sections: dict[str, dict[str, str]] = {}
    current_variant: Optional[dict[str, str]] = None
    current_type = "benefit"
    pending_header = ""
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        if line == "潜在获益靶向/免疫药物解析":
            current_type = "benefit"
            current_variant = None
            pending_header = ""
            idx += 1
            continue
        if line == "潜在负相关靶向/免疫药物解析":
            current_type = "caution"
            current_variant = None
            pending_header = ""
            idx += 1
            continue
        if _is_part3_drug_header(line):
            current_variant = _drug_variant_from_header(line)
            current_type = current_variant.get("type") or current_type
            pending_header = line
            idx += 1
            continue
        if (
            current_variant is None
            or line in {"靶向药物/免疫用药提示解析", "基因变异与药物关联分析：", "药物疗效临床解析："}
            or line.startswith("潜在")
        ):
            idx += 1
            continue

        drug_name = line
        block = [pending_header] if pending_header else []
        block.append(drug_name)
        pending_header = ""
        idx += 1
        if idx < len(lines) and lines[idx] == "基因变异与药物关联分析：":
            block.append(lines[idx])
            idx += 1
            while idx < len(lines) and lines[idx] != "药物疗效临床解析：":
                if _is_part3_drug_header(lines[idx]) or lines[idx].startswith("潜在"):
                    break
                block.append(lines[idx])
                idx += 1
        if idx < len(lines) and lines[idx] == "药物疗效临床解析：":
            block.append(lines[idx])
            idx += 1
            while idx < len(lines):
                if (
                    _is_part3_drug_header(lines[idx])
                    or lines[idx].startswith("潜在")
                    or _is_next_drug_name(lines, idx)
                ):
                    break
                block.append(lines[idx])
                idx += 1
        key = (
            f"{current_type}:"
            f"{current_variant.get('gene')}|{current_variant.get('c_hgvs')}|"
            f"{current_variant.get('p_hgvs')}|{_normalize_part3_text(drug_name)}"
        )
        sections[key] = {"text": "\n".join(item for item in block if item)}
    return sections


def _find_paragraph_index(lines: list[str], needle: str) -> Optional[int]:
    for idx, line in enumerate(lines):
        if str(line).strip() == needle:
            return idx
    return None


def _is_part3_gene_header(text: str) -> bool:
    value = str(text or "").strip()
    if value.startswith("u "):
        value = value[2:].strip()
    if "突变相应" in value:
        return False
    return bool(re.match(r"^[A-Z0-9]+：.+；\d+(?:\.\d+)?%$", value))


def _part3_variant_key_from_header(text: str) -> str:
    value = str(text or "").strip()
    if value.startswith("u "):
        value = value[2:].strip()
    gene, rest = value.split("：", 1)
    variant, _freq = rest.rsplit("；", 1)
    parts = [part.strip() for part in variant.split("，") if part.strip()]
    c_hgvs = parts[0] if parts else ""
    p_hgvs = parts[1] if len(parts) > 1 else ""
    return f"{gene.upper()}|{_part3_hgvs_key(c_hgvs)}|{_part3_hgvs_key(p_hgvs)}"


def _is_part3_drug_header(text: str) -> bool:
    return bool(re.match(r"^[A-Z0-9]+：.+突变相应.+药物$", str(text or "").strip()))


def _drug_variant_from_header(text: str) -> dict[str, str]:
    value = str(text or "").strip()
    gene, rest = value.split("：", 1)
    variant, _tail = rest.split("突变相应", 1)
    parts = [part.strip() for part in variant.split("，") if part.strip()]
    return {
        "gene": gene.upper(),
        "c_hgvs": _part3_hgvs_key(parts[0] if parts else ""),
        "p_hgvs": _part3_hgvs_key(parts[1] if len(parts) > 1 else ""),
        "type": "caution" if "负相关" in value else "benefit",
    }


def _is_next_drug_name(lines: list[str], idx: int) -> bool:
    return idx + 1 < len(lines) and lines[idx + 1] == "基因变异与药物关联分析："


def _part3_hgvs_key(text: str) -> str:
    return re.sub(r"\s+", "", str(text or "")).upper()


def _normalize_part3_text(text: str) -> str:
    value = str(text or "").replace("\u00a0", " ")
    return re.sub(r"\s+", "", value)


def _table_cell_diffs(
    ref_rows: list[list[str]],
    cand_rows: list[list[str]],
    max_samples: int,
    *,
    normalize_whitespace: bool = False,
) -> list[dict[str, Any]]:
    diffs = []
    for r_idx, (ref_row, cand_row) in enumerate(zip(ref_rows, cand_rows)):
        for c_idx, (ref_cell, cand_cell) in enumerate(zip(ref_row, cand_row)):
            ref_cmp = _normalize_soft_whitespace(ref_cell) if normalize_whitespace else ref_cell
            cand_cmp = _normalize_soft_whitespace(cand_cell) if normalize_whitespace else cand_cell
            if ref_cmp != cand_cmp:
                diffs.append(
                    {
                        "cell": [r_idx, c_idx],
                        "reference": ref_cell,
                        "candidate": cand_cell,
                    }
                )
                if len(diffs) >= max_samples:
                    return diffs
    return diffs


def _compare_styles(
    ref: Mapping[str, Any],
    cand: Mapping[str, Any],
    *,
    max_samples: int,
    policy: str = "warn",
) -> dict[str, Any]:
    issues = []
    samples = []
    ref_summary = dict(ref.get("style_summary") or {})
    cand_summary = dict(cand.get("style_summary") or {})
    for key in sorted(set(ref_summary) | set(cand_summary)):
        if ref_summary.get(key, 0) != cand_summary.get(key, 0):
            samples.append(
                {
                    "key": key,
                    "reference": ref_summary.get(key, 0),
                    "candidate": cand_summary.get(key, 0),
                    "message": (
                        f"Style metric {key} differs: "
                        f"{ref_summary.get(key, 0)} -> {cand_summary.get(key, 0)}"
                    ),
                }
            )
            if len(samples) >= max_samples:
                break
    if samples and policy != "summary":
        issues.append(
            {
                "level": "warning",
                "code": "STYLE_DIFF",
                "message": f"Basic DOCX style metrics differ ({len(samples)} sampled).",
            }
        )
    return {
        "status": _section_status(issues),
        "policy": policy,
        "samples": samples,
        "issues": issues,
    }


def _normalize_row(row: list[str], normalize_whitespace: bool) -> list[str]:
    if not normalize_whitespace:
        return row
    return [_normalize_soft_whitespace(cell) for cell in row]


def _normalize_soft_whitespace(text: str) -> str:
    value = str(text or "").replace("\u00a0", " ")
    value = "\n".join(line.rstrip() for line in value.splitlines())
    return value.strip()


def _is_reference_artifact(text: str) -> bool:
    value = str(text or "").strip()
    return bool(value and len(value) >= 12 and len(set(value)) == 1 and value[0].isdigit())


def _compare_qa(
    ref_qa: Optional[Mapping[str, Any]],
    cand_qa: Optional[Mapping[str, Any]],
    *,
    max_samples: int,
) -> dict[str, Any]:
    issues = []
    samples = []
    ref_status = ref_qa.get("status") if ref_qa else None
    cand_status = cand_qa.get("status") if cand_qa else None
    if cand_status == "FAIL":
        issues.append(
            {
                "level": "error",
                "code": "CANDIDATE_QA_FAIL",
                "message": "Candidate QA status is FAIL.",
            }
        )
    elif ref_status and cand_status and ref_status != cand_status:
        issues.append(
            {
                "level": "warning",
                "code": "QA_STATUS_DIFF",
                "message": f"QA status differs: {ref_status} -> {cand_status}.",
            }
        )
    ref_codes = _qa_issue_codes(ref_qa)
    cand_codes = _qa_issue_codes(cand_qa)
    added = sorted(cand_codes - ref_codes)
    removed = sorted(ref_codes - cand_codes)
    if added or removed:
        issues.append(
            {
                "level": "warning",
                "code": "QA_ISSUE_CODE_DIFF",
                "message": f"QA issue codes differ: added={added}, removed={removed}.",
            }
        )
        samples.append({"message": "QA issue code diff", "added": added, "removed": removed})
    return {
        "status": _section_status(issues),
        "reference_status": ref_status,
        "candidate_status": cand_status,
        "reference_issue_codes": sorted(ref_codes),
        "candidate_issue_codes": sorted(cand_codes),
        "samples": samples[:max_samples],
        "issues": issues,
    }


def _load_qa(path: Optional[str], docx_path: Path) -> Optional[Mapping[str, Any]]:
    qa_path = Path(path).resolve() if path else docx_path.with_suffix(".qa.json")
    if not qa_path.exists():
        return None
    try:
        return json.loads(qa_path.read_text(encoding="utf-8"))
    except Exception:
        return {"status": None, "issues": [{"code": "QA_READ_FAILED"}]}


def _qa_issue_codes(qa: Optional[Mapping[str, Any]]) -> set[str]:
    if not qa:
        return set()
    return {
        str(item.get("code"))
        for item in qa.get("issues") or []
        if isinstance(item, Mapping) and item.get("code")
    }


def _collect_issues(sections: Mapping[str, Any]) -> list[dict[str, Any]]:
    issues = []
    for section_name, section in sections.items():
        for item in section.get("issues") or []:
            row = dict(item)
            row["section"] = section_name
            issues.append(row)
    return issues


def _overall_status(issues: list[Mapping[str, Any]]) -> str:
    if any(item.get("level") == "error" for item in issues):
        return "FAIL"
    if any(item.get("level") == "warning" for item in issues):
        return "WARN"
    return "PASS"


def _section_status(issues: list[Mapping[str, Any]]) -> str:
    if any(item.get("level") == "error" for item in issues):
        return "FAIL"
    if any(item.get("level") == "warning" for item in issues):
        return "WARN"
    return "PASS"
