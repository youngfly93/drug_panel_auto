"""
Template contract extraction + validation.

Goal:
  Fail fast (or warn) when a docxtpl/Jinja2 template references variables that
  are not present in the runtime context, to avoid generating reports that look
  "successful" but contain missing content.

Scope (pragmatic, Word-template friendly):
  - Extract `{{ ... }}` variables (top-level + dotted paths).
  - Extract `{% for <var> in <list> %}` loop lists, and the fields referenced
    as `<var>.field` / `<var>['field']` inside the same table.

Notes:
  - This is a heuristic parser intended for docxtpl-style templates where loops
    are usually placed in table rows/cells.
  - It does not aim to fully parse Jinja2 expressions.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Iterable, List, Mapping, Optional, Set, Tuple

from docx import Document

from reportgen.docx_sections import inspect_structural_marker

_JINJA_VAR_RE = re.compile(r"\{\{\s*(?P<expr>.*?)\s*\}\}")
_JINJA_FOR_RE = re.compile(
    r"\{%\s*(?:tr|tc|p|r)?\s*for\s+(?P<var>[a-zA-Z_][a-zA-Z0-9_]*)\s+in\s+"
    r"(?P<list>[a-zA-Z_][a-zA-Z0-9_]*)\s*%}"
)

# Best-effort: grab the leading "variable-ish" token from a Jinja expression.
# Supports dotted paths: foo.bar.baz
_LEADING_PATH_RE = re.compile(
    r"^\s*(?P<path>[a-zA-Z_][a-zA-Z0-9_]*(?:\.[a-zA-Z_][a-zA-Z0-9_]*)*)"
)

_JINJA_BUILTINS = {"loop", "cycler", "namespace", "range", "dict", "list", "lipsum"}


@dataclass(frozen=True)
class TemplateContract:
    template_path: str
    required_paths: Tuple[str, ...]
    required_lists: Tuple[str, ...]
    loop_row_fields: Dict[str, Tuple[str, ...]]  # list_name -> required row fields


@dataclass(frozen=True)
class TemplateTableExpectation:
    """Declared table structure that a template must contain."""

    name: str
    required_headers: Tuple[str, ...] = ()
    columns: Optional[int] = None
    min_columns: Optional[int] = None
    max_columns: Optional[int] = None
    required: bool = True
    max_preview_rows: int = 4


@dataclass(frozen=True)
class DeclaredTemplateContract:
    """Panel-declared template contract loaded from panel.yaml."""

    required_variables: Tuple[str, ...] = ()
    required_lists: Tuple[str, ...] = ()
    required_markers: Tuple[str, ...] = ()
    table_structures: Tuple[TemplateTableExpectation, ...] = ()


@dataclass(frozen=True)
class ContractValidation:
    ok: bool
    missing_paths: Tuple[str, ...]
    missing_lists: Tuple[str, ...]
    missing_row_fields: Dict[str, Tuple[str, ...]]
    missing_row_examples: Dict[
        str, Dict[str, int]
    ]  # list_name -> field -> first missing row index (0-based)


@dataclass(frozen=True)
class DeclaredContractValidation:
    ok: bool
    missing_required_variables: Tuple[str, ...]
    missing_required_lists: Tuple[str, ...]
    missing_required_markers: Tuple[str, ...]
    duplicate_required_markers: Tuple[str, ...]
    marker_counts: Dict[str, int]
    missing_required_tables: Tuple[str, ...]
    table_errors: Dict[str, Tuple[str, ...]]
    table_matches: Dict[str, Dict[str, Any]]


def _iter_doc_text(doc: Document) -> Iterable[str]:
    for p in doc.paragraphs:
        if p.text:
            yield p.text

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text:
                    yield cell.text

    for sec in doc.sections:
        for p in sec.header.paragraphs:
            if p.text:
                yield p.text
        for p in sec.footer.paragraphs:
            if p.text:
                yield p.text


def _extract_loop_contract_from_table(table) -> List[Tuple[str, str, Set[str]]]:
    """Return (list_name, loop_var, required_fields) tuples from one table."""
    found: List[Tuple[str, str, Set[str]]] = []

    # Scan the whole table text, but keep it bounded to this table to avoid
    # cross-contamination.
    table_text_parts: List[str] = []
    for row in table.rows:
        for cell in row.cells:
            if cell.text:
                table_text_parts.append(cell.text)
    table_text = "\n".join(table_text_parts)

    for m in _JINJA_FOR_RE.finditer(table_text):
        loop_var = m.group("var")
        list_name = m.group("list")

        dot_re = re.compile(rf"\b{re.escape(loop_var)}\.([a-zA-Z_][a-zA-Z0-9_]*)\b")
        bracket_re = re.compile(rf"{re.escape(loop_var)}\[['\"]([^'\"]+)['\"]\]")

        fields: Set[str] = set()
        for fm in dot_re.finditer(table_text):
            fields.add(fm.group(1))
        for fm in bracket_re.finditer(table_text):
            fields.add(fm.group(1))

        found.append((list_name, loop_var, fields))

    return found


def extract_template_contract(template_path: str) -> TemplateContract:
    template_p = Path(template_path)
    if not template_p.exists():
        raise FileNotFoundError(f"Template docx not found: {template_p}")
    if template_p.suffix.lower() != ".docx":
        raise ValueError(f"Template must be a .docx file: {template_p}")

    doc = Document(str(template_p))

    loop_vars: Set[str] = set()
    required_lists: Set[str] = set()
    loop_row_fields: Dict[str, Set[str]] = {}

    for tbl in doc.tables:
        for list_name, loop_var, fields in _extract_loop_contract_from_table(tbl):
            loop_vars.add(loop_var)
            required_lists.add(list_name)
            loop_row_fields.setdefault(list_name, set()).update(fields)

    required_paths: Set[str] = set()
    for text in _iter_doc_text(doc):
        for m in _JINJA_VAR_RE.finditer(text):
            expr = m.group("expr") or ""
            # Strip filters: a|b|c -> a
            expr = expr.split("|", 1)[0]
            m2 = _LEADING_PATH_RE.match(expr)
            if not m2:
                continue
            path = m2.group("path")
            if not path:
                continue
            if path in _JINJA_BUILTINS:
                continue

            # Exclude loop variables (e.g. row.xxx) from top-level required paths.
            if path in loop_vars:
                continue
            if any(path.startswith(f"{lv}.") for lv in loop_vars):
                continue

            required_paths.add(path)

    # Sorting for stable output
    return TemplateContract(
        template_path=str(template_p),
        required_paths=tuple(sorted(required_paths)),
        required_lists=tuple(sorted(required_lists)),
        loop_row_fields={
            k: tuple(sorted(v)) for k, v in sorted(loop_row_fields.items())
        },
    )


def parse_declared_template_contract(spec: Any) -> DeclaredTemplateContract:
    """Parse a panel.yaml ``template_contract`` section.

    The accepted YAML shape is intentionally small and portable:

    template_contract:
      required_variables: [patient_name, sample_id]
      required_lists: [variants_2_1]
      required_table_structures:
        variant_detail_table:
          columns: 9
          required_headers: [基因名称, 转录本号]
    """
    if not isinstance(spec, Mapping):
        return DeclaredTemplateContract()

    table_specs: List[TemplateTableExpectation] = []
    raw_tables = (
        spec.get("required_table_structures")
        or spec.get("table_structures")
        or spec.get("tables")
        or {}
    )
    if isinstance(raw_tables, Mapping):
        items = []
        for name, raw in raw_tables.items():
            raw_map = dict(raw) if isinstance(raw, Mapping) else {}
            raw_map.setdefault("name", str(name))
            items.append(raw_map)
    elif isinstance(raw_tables, list):
        items = [item for item in raw_tables if isinstance(item, Mapping)]
    else:
        items = []

    for item in items:
        name = str(item.get("name") or item.get("id") or "").strip()
        if not name:
            continue
        table_specs.append(
            TemplateTableExpectation(
                name=name,
                required_headers=tuple(_str_list(item.get("required_headers"))),
                columns=_int_or_none(item.get("columns")),
                min_columns=_int_or_none(item.get("min_columns")),
                max_columns=_int_or_none(item.get("max_columns")),
                required=bool(item.get("required", True)),
                max_preview_rows=int(item.get("max_preview_rows") or 4),
            )
        )

    required_variables = (
        spec.get("required_variables")
        or spec.get("required_paths")
        or spec.get("variables")
        or []
    )
    required_lists = (
        spec.get("required_lists")
        or spec.get("required_loop_lists")
        or spec.get("lists")
        or []
    )
    required_markers = spec.get("required_markers") or spec.get("markers") or []

    return DeclaredTemplateContract(
        required_variables=tuple(sorted(set(_str_list(required_variables)))),
        required_lists=tuple(sorted(set(_str_list(required_lists)))),
        required_markers=tuple(sorted(set(_str_list(required_markers)))),
        table_structures=tuple(table_specs),
    )


def validate_declared_contract(
    template_path: str,
    extracted_contract: TemplateContract,
    declared_contract: Any,
) -> DeclaredContractValidation:
    """Validate a template against the panel-declared contract."""
    declared = parse_declared_template_contract(declared_contract)
    actual_variables = set(extracted_contract.required_paths)
    actual_lists = set(extracted_contract.required_lists)

    missing_variables = tuple(
        sorted(
            variable
            for variable in declared.required_variables
            if variable not in actual_variables
        )
    )
    missing_lists = tuple(
        sorted(list_name for list_name in declared.required_lists if list_name not in actual_lists)
    )

    marker_document = Document(str(template_path))
    marker_inspections = {
        marker: inspect_structural_marker(marker_document, marker)
        for marker in declared.required_markers
    }
    marker_counts = {
        marker: total for marker, (_indices, total) in marker_inspections.items()
    }
    missing_markers = tuple(
        sorted(
            marker
            for marker, (indices, _total) in marker_inspections.items()
            if len(indices) == 0
        )
    )
    duplicate_markers = tuple(
        sorted(
            marker
            for marker, (_indices, total) in marker_inspections.items()
            if total > 1
        )
    )

    table_snapshots = _inspect_template_tables(
        template_path,
        max_rows=max(
            [table.max_preview_rows for table in declared.table_structures] or [4]
        ),
    )
    missing_tables: List[str] = []
    table_errors: Dict[str, Tuple[str, ...]] = {}
    table_matches: Dict[str, Dict[str, Any]] = {}

    for expected in declared.table_structures:
        candidates = [
            table
            for table in table_snapshots
            if _table_matches_headers(table, expected.required_headers)
        ]
        if not candidates:
            if expected.required:
                missing_tables.append(expected.name)
            continue

        # Prefer the widest match so a merged header row does not hide a valid
        # detail row further down in the table.
        best = max(candidates, key=lambda item: int(item.get("cols") or 0))
        table_matches[expected.name] = best

        errors: List[str] = []
        cols = int(best.get("cols") or 0)
        if expected.columns is not None and cols != expected.columns:
            errors.append(f"expected {expected.columns} columns, found {cols}")
        if expected.min_columns is not None and cols < expected.min_columns:
            errors.append(f"expected at least {expected.min_columns} columns, found {cols}")
        if expected.max_columns is not None and cols > expected.max_columns:
            errors.append(f"expected at most {expected.max_columns} columns, found {cols}")
        if errors:
            table_errors[expected.name] = tuple(errors)

    ok = not (
        missing_variables
        or missing_lists
        or missing_markers
        or duplicate_markers
        or missing_tables
        or table_errors
    )
    return DeclaredContractValidation(
        ok=ok,
        missing_required_variables=missing_variables,
        missing_required_lists=missing_lists,
        missing_required_markers=missing_markers,
        duplicate_required_markers=duplicate_markers,
        marker_counts=marker_counts,
        missing_required_tables=tuple(sorted(missing_tables)),
        table_errors=table_errors,
        table_matches=table_matches,
    )


def declared_validation_to_dict(
    validation: DeclaredContractValidation,
    declared_contract: Any,
) -> Dict[str, Any]:
    """Return a JSON-serializable declared-contract validation report."""
    declared = parse_declared_template_contract(declared_contract)
    return {
        "ok": bool(validation.ok),
        "required_variables": list(declared.required_variables),
        "required_lists": list(declared.required_lists),
        "required_markers": list(declared.required_markers),
        "required_table_structures": [
            {
                "name": table.name,
                "required_headers": list(table.required_headers),
                "columns": table.columns,
                "min_columns": table.min_columns,
                "max_columns": table.max_columns,
                "required": table.required,
            }
            for table in declared.table_structures
        ],
        "missing_required_variables": list(validation.missing_required_variables),
        "missing_required_lists": list(validation.missing_required_lists),
        "missing_required_markers": list(validation.missing_required_markers),
        "duplicate_required_markers": list(validation.duplicate_required_markers),
        "marker_counts": dict(validation.marker_counts),
        "missing_required_tables": list(validation.missing_required_tables),
        "table_errors": {
            name: list(errors) for name, errors in validation.table_errors.items()
        },
        "table_matches": validation.table_matches,
    }


def _get_by_path(obj: Any, path: str) -> bool:
    """Best-effort dotted-path resolver for dict-like contexts."""
    if not path:
        return False
    cur: Any = obj
    for part in path.split("."):
        if isinstance(cur, Mapping):
            if part not in cur:
                return False
            cur = cur[part]
            continue
        if hasattr(cur, part):
            cur = getattr(cur, part)
            continue
        return False
    return True


def validate_contract(
    contract: TemplateContract, *, context: Mapping[str, Any]
) -> ContractValidation:
    missing_paths: List[str] = []
    for p in contract.required_paths:
        if not _get_by_path(context, p):
            missing_paths.append(p)

    missing_lists: List[str] = []
    missing_row_fields: Dict[str, List[str]] = {}
    missing_row_examples: Dict[str, Dict[str, int]] = {}

    for list_name in contract.required_lists:
        if list_name not in context:
            missing_lists.append(list_name)
            continue

        value = context.get(list_name)
        if not isinstance(value, (list, tuple)):
            missing_lists.append(list_name)
            continue

        required_fields = set(contract.loop_row_fields.get(list_name, ()))
        if not required_fields or not value:
            continue

        # Strict: every row should provide every field referenced in the template.
        for field in sorted(required_fields):
            for idx, row in enumerate(value):
                ok = False
                if isinstance(row, Mapping):
                    ok = field in row
                elif hasattr(row, field):
                    ok = True

                if not ok:
                    missing_row_fields.setdefault(list_name, []).append(field)
                    missing_row_examples.setdefault(list_name, {})[field] = idx
                    break

    # Deduplicate and sort
    missing_row_fields_sorted: Dict[str, Tuple[str, ...]] = {
        k: tuple(sorted(set(v))) for k, v in missing_row_fields.items()
    }

    ok = (not missing_paths) and (not missing_lists) and (not missing_row_fields_sorted)

    return ContractValidation(
        ok=ok,
        missing_paths=tuple(sorted(set(missing_paths))),
        missing_lists=tuple(sorted(set(missing_lists))),
        missing_row_fields=missing_row_fields_sorted,
        missing_row_examples=missing_row_examples,
    )


def _str_list(value: Any) -> List[str]:
    if isinstance(value, str):
        return [value.strip()] if value.strip() else []
    if not isinstance(value, (list, tuple, set)):
        return []
    result: List[str] = []
    for item in value:
        text = str(item or "").strip()
        if text:
            result.append(text)
    return result


def _int_or_none(value: Any) -> Optional[int]:
    if value is None or value == "":
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _inspect_template_tables(template_path: str, *, max_rows: int = 4) -> List[Dict[str, Any]]:
    doc = Document(str(template_path))
    snapshots: List[Dict[str, Any]] = []
    for idx, table in enumerate(doc.tables):
        preview = _table_preview(table, max_rows=max_rows)
        snapshots.append(
            {
                "index": idx,
                "rows": len(table.rows),
                "cols": _table_col_count(table),
                "preview": preview,
                "compact_preview": _compact_text(preview),
            }
        )
    return snapshots


def _table_matches_headers(table: Mapping[str, Any], headers: Tuple[str, ...]) -> bool:
    if not headers:
        return True
    compact_preview = str(table.get("compact_preview") or "")
    return all(_compact_text(header) in compact_preview for header in headers)


def _table_col_count(table: Any) -> int:
    if not table.rows:
        return 0
    return max((len(row.cells) for row in table.rows), default=0)


def _table_preview(table: Any, *, max_rows: int) -> str:
    rows: List[str] = []
    for row in table.rows[:max_rows]:
        rows.append(" ".join((cell.text or "").strip() for cell in row.cells))
    return "\n".join(rows)


def _compact_text(value: str) -> str:
    return re.sub(r"\s+", "", value or "")
