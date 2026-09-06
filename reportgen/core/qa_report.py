"""
Post-generation QA report for rendered DOCX outputs.

The checks in this module are intentionally machine-readable and conservative:
they catch known high-risk report regressions without tying the pipeline to one
patient case. Panel-specific assertions are enabled from project/report context.
"""

from __future__ import annotations

import hashlib
import os
import re
import shutil
import subprocess
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Mapping, Optional
from zipfile import ZipFile

from docx import Document

from reportgen.core.pipeline.summary import summarize_stage_results
from reportgen.core.processors import critical_docx_processor_names
from reportgen.models.report_data import ReportData
from reportgen.utils.artifacts import write_json
from reportgen.utils.docx_render import render_docx_to_pngs, renderer_fingerprint

PLACEHOLDER_RE = re.compile(
    r"(\{\{.*?\}\}|\{%.*?%\}|\{#.*?#\}|__[A-Z][A-Z0-9_]{2,}__)",
    re.DOTALL,
)
CRC_PROJECT_TYPES = {"crc_301", "crc_301_msi", "crc_358", "crc_358_msi"}


def _file_sha256(path: Path) -> str | None:
    if not path.is_file():
        return None
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _build_provenance() -> dict[str, Any]:
    """Identify the immutable release (or local Git tree) that produced QA."""
    explicit = str(os.environ.get("REPORTGEN_SOURCE_REVISION") or "").strip()
    if re.fullmatch(r"[0-9a-fA-F]{7,40}", explicit):
        return {
            "source_revision": explicit.lower(),
            "source_kind": "environment",
            "source_dirty": False,
        }

    source_file = Path(__file__).resolve()
    for parent in source_file.parents:
        revision_file = parent / "REVISION"
        if not revision_file.is_file():
            continue
        revision = revision_file.read_text(encoding="utf-8").strip()
        if re.fullmatch(r"[0-9a-fA-F]{7,40}", revision):
            return {
                "source_revision": revision.lower(),
                "source_kind": "release_revision_file",
                "source_dirty": False,
            }

    try:
        revision = subprocess.run(
            ["git", "rev-parse", "HEAD"],
            cwd=source_file.parent,
            check=True,
            capture_output=True,
            text=True,
            timeout=5,
        ).stdout.strip()
        dirty = bool(
            subprocess.run(
                ["git", "status", "--porcelain", "--untracked-files=normal"],
                cwd=source_file.parent,
                check=True,
                capture_output=True,
                text=True,
                timeout=5,
            ).stdout.strip()
        )
        if re.fullmatch(r"[0-9a-fA-F]{40}", revision):
            return {
                "source_revision": revision.lower(),
                "source_kind": "git_worktree",
                "source_dirty": dirty,
            }
    except Exception:
        pass
    return {
        "source_revision": None,
        "source_kind": "unavailable",
        "source_dirty": None,
    }


def build_docx_qa_report(
    *,
    output_file: str,
    report_data: Optional[ReportData] = None,
    project_type: Optional[str] = None,
    project_name: Optional[str] = None,
    template_file: Optional[str] = None,
    generation_id: Optional[str] = None,
    field_provenance: Optional[Mapping[str, Any]] = None,
    field_provenance_file: Optional[str] = None,
    processor_report: Optional[list[Mapping[str, Any]]] = None,
    template_contract: Optional[Mapping[str, Any]] = None,
    rule_provenance: Optional[Mapping[str, Any]] = None,
    stage_results: Optional[list[Mapping[str, Any]]] = None,
    stage_results_file: Optional[str] = None,
    visual_render: Optional[str] = None,
    visual_render_required: bool = False,
    visual_render_dpi: int = 120,
    visual_render_timeout_seconds: int = 120,
    visual_render_output_dir: Optional[str] = None,
    visual_render_tmp_dir: Optional[str] = None,
) -> Dict[str, Any]:
    """Build a structured QA report for one generated DOCX file."""
    generated_at = datetime.now().isoformat()
    output_path = Path(output_file)
    context = report_data.context if report_data is not None else {}
    issues: List[Dict[str, str]] = []
    checks: Dict[str, Any] = {}
    metrics: Dict[str, Any] = {
        "file_size_bytes": output_path.stat().st_size if output_path.exists() else None,
        "output_sha256": _file_sha256(output_path),
    }

    def issue(level: str, code: str, message: str) -> None:
        issues.append({"level": level, "code": code, "message": message})

    checks["template_contract"] = _template_contract_check(template_contract)
    if checks["template_contract"]["status"] == "FAIL":
        issue(
            "error",
            "TEMPLATE_CONTRACT_FAILED",
            checks["template_contract"]["message"],
        )
    checks["rules"] = _rule_provenance_check(rule_provenance)
    if checks["rules"]["status"] == "FAIL":
        issue("error", "RULE_PROVENANCE_FAILED", checks["rules"]["message"])

    if not output_path.exists():
        issue("error", "DOCX_NOT_FOUND", f"Output DOCX not found: {output_file}")
        return attach_pipeline_summary(
            _finalize_report(
                generated_at=generated_at,
                generation_id=generation_id,
                project_type=project_type,
                project_name=project_name,
                template_file=template_file,
                output_file=output_file,
                checks=checks,
                metrics=metrics,
                issues=issues,
                template_contract=template_contract,
                rule_provenance=rule_provenance,
            ),
            stage_results=stage_results,
            stage_results_file=stage_results_file,
        )

    try:
        doc = Document(str(output_path))
        checks["docx_openable"] = {"status": "PASS", "value": True}
    except Exception as exc:
        issue("error", "DOCX_OPEN_FAILED", f"DOCX cannot be opened: {exc}")
        checks["docx_openable"] = {"status": "FAIL", "value": False}
        return attach_pipeline_summary(
            _finalize_report(
                generated_at=generated_at,
                generation_id=generation_id,
                project_type=project_type,
                project_name=project_name,
                template_file=template_file,
                output_file=output_file,
                checks=checks,
                metrics=metrics,
                issues=issues,
                template_contract=template_contract,
                rule_provenance=rule_provenance,
            ),
            stage_results=stage_results,
            stage_results_file=stage_results_file,
        )

    paragraphs = list(_iter_all_paragraphs(doc))
    text = _read_docx_text(doc)
    compact_text = _compact(text)
    part3_text = _read_part3_text(doc, context.get("part3_cross_cancer_residual_scan"))
    table_metrics = _inspect_tables(doc)
    metrics.update(
        {
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "section_count": len(doc.sections),
            **table_metrics,
        }
    )

    placeholder_matches = PLACEHOLDER_RE.findall(text)
    brace_counts = {
        "{{": text.count("{{"),
        "}}": text.count("}}"),
        "{%": text.count("{%"),
        "%}": text.count("%}"),
    }
    placeholder_count = len(placeholder_matches) + sum(brace_counts.values())
    checks["unrendered_placeholders"] = {
        "status": "FAIL" if placeholder_count else "PASS",
        "count": placeholder_count,
        "samples": [_trim_sample(v) for v in placeholder_matches[:5]],
        "brace_counts": brace_counts,
    }
    if placeholder_count:
        issue(
            "error",
            "UNRENDERED_PLACEHOLDER",
            "Rendered DOCX still contains template placeholders.",
        )

    empty_numbered = [
        _paragraph_location(p) for p in paragraphs if _is_empty_numbered_paragraph(p)
    ]
    checks["empty_numbered_paragraphs"] = {
        "status": "FAIL" if empty_numbered else "PASS",
        "count": len(empty_numbered),
        "locations": empty_numbered[:10],
    }
    if empty_numbered:
        issue(
            "error",
            "EMPTY_NUMBERED_PARAGRAPH",
            "Rendered DOCX contains visible empty bullet/numbered paragraphs.",
        )

    toc = _inspect_toc(paragraphs, output_path=output_path)
    checks["toc_page_numbers"] = toc
    if toc["status"] == "FAIL":
        issue("error", "TOC_FIELD_INVALID", toc["message"])
    elif toc["status"] == "WARN":
        issue("warning", "TOC_PAGE_NUMBERS_MISSING", toc["message"])

    checks["visual_render"] = _build_visual_render_check(
        output_path,
        mode=visual_render,
        required=visual_render_required,
        dpi=visual_render_dpi,
        timeout_seconds=visual_render_timeout_seconds,
        output_dir=visual_render_output_dir,
        tmp_dir=visual_render_tmp_dir,
        expected_sparse_pages=_visual_expected_sparse_page_specs(context),
    )
    for visual_issue in _visual_render_issues(checks["visual_render"]):
        issue(**visual_issue)
    metrics["visual_render_page_count"] = len(
        checks["visual_render"].get("rendered_pages") or []
    )

    checks["blank_page_detection"] = _blank_page_detection_check(
        checks["visual_render"]
    ) or {
        "status": "SKIP",
        "message": "DOCX-level QA cannot reliably identify visual blank pages; use PDF/PNG render QA for this check.",
    }

    if field_provenance:
        checks["field_provenance"] = {
            "status": "PASS",
            "file": field_provenance_file,
            "key_field_sources": _field_source_summary(field_provenance),
        }
    else:
        checks["field_provenance"] = {
            "status": "SKIP",
            "file": field_provenance_file,
            "key_field_sources": {},
            "message": "No field provenance report was provided.",
        }

    processor_rows = list(processor_report or [])
    processor_errors = [
        row
        for row in processor_rows
        if isinstance(row, Mapping) and row.get("status") == "ERROR"
    ]
    critical_names = critical_docx_processor_names(project_type)
    critical_processor_errors = [
        row for row in processor_errors if row.get("name") in critical_names
    ]
    checks["post_processors"] = {
        "status": (
            "FAIL"
            if critical_processor_errors
            else (
                "WARN"
                if processor_errors
                else ("PASS" if processor_report is not None else "SKIP")
            )
        ),
        "count": len(processor_rows),
        "error_count": len(processor_errors),
        "errors": processor_errors,
        "critical_error_count": len(critical_processor_errors),
        "critical_errors": critical_processor_errors,
    }
    if processor_errors:
        issue(
            "error" if critical_processor_errors else "warning",
            "POST_PROCESSOR_ERRORS",
            "One or more DOCX post-render processors reported errors.",
        )
    if critical_processor_errors:
        issue(
            "error",
            "CRITICAL_POST_PROCESSOR_ERROR",
            "Part 3 or reference rebuilding failed; the report is not deliverable.",
        )

    table_checks = _build_table_checks(table_metrics, project_type, context)
    checks.update(table_checks)
    for table_issue in _table_issues(table_checks):
        issue(**table_issue)

    style_checks = _build_style_checks(doc, project_type, context)
    checks.update(style_checks)
    for style_issue in _style_issues(style_checks):
        issue(**style_issue)

    business_checks = _build_business_checks(
        compact_text,
        context,
        project_type,
        part3_compact_text=(
            _compact(part3_text) if part3_text is not None else None
        ),
    )
    checks.update(business_checks)
    for business_issue in _business_issues(business_checks):
        issue(**business_issue)

    return attach_pipeline_summary(
        _finalize_report(
            generated_at=generated_at,
            generation_id=generation_id,
            project_type=project_type,
            project_name=project_name,
            template_file=template_file,
            output_file=output_file,
            field_provenance_file=field_provenance_file,
            processor_report=processor_report,
            template_contract=template_contract,
            rule_provenance=rule_provenance,
            checks=checks,
            metrics=metrics,
            issues=issues,
        ),
        stage_results=stage_results,
        stage_results_file=stage_results_file,
    )


def write_docx_qa_report(
    report: Mapping[str, Any], output_file: str, qa_file: Optional[str] = None
) -> str:
    """Write QA JSON next to the generated DOCX and return the JSON path."""
    path = Path(qa_file) if qa_file else Path(output_file).with_suffix(".qa.json")
    write_json(path, dict(report))
    return str(path)


def attach_pipeline_summary(
    report: Mapping[str, Any],
    *,
    stage_results: Optional[list[Mapping[str, Any]]] = None,
    stage_results_file: Optional[str] = None,
) -> Dict[str, Any]:
    """Attach pipeline summary and fold it into QA status."""
    result = dict(report)
    pipeline = summarize_stage_results(
        stage_results,
        stage_results_file=stage_results_file,
    )
    checks = dict(result.get("checks") or {})
    checks["pipeline"] = {
        "status": pipeline["status"],
        "message": _pipeline_message(pipeline),
        "stage_count": pipeline["stage_count"],
        "failed_stages": pipeline["failed_stages"],
        "warning_stages": pipeline["warning_stages"],
        "stage_results_file": pipeline["stage_results_file"],
    }
    issues = list(result.get("issues") or [])
    if pipeline["status"] == "FAIL":
        issues.append(
            {
                "level": "error",
                "code": "PIPELINE_FAILED",
                "message": _pipeline_message(pipeline),
            }
        )
    elif pipeline["status"] == "WARN":
        issues.append(
            {
                "level": "warning",
                "code": "PIPELINE_WARN",
                "message": _pipeline_message(pipeline),
            }
        )

    has_error = any(i.get("level") == "error" for i in issues)
    has_warning = any(i.get("level") == "warning" for i in issues)
    result["pipeline"] = pipeline
    result["checks"] = checks
    result["issues"] = issues
    result["status"] = "FAIL" if has_error else ("WARN" if has_warning else "PASS")
    return result


def _pipeline_message(pipeline: Mapping[str, Any]) -> str:
    status = pipeline.get("status")
    if status == "FAIL":
        return "Generation pipeline failed: " + ", ".join(
            pipeline.get("failed_stages") or []
        )
    if status == "WARN":
        return "Generation pipeline completed with warnings: " + ", ".join(
            pipeline.get("warning_stages") or []
        )
    return "Generation pipeline passed."


def _finalize_report(
    *,
    generated_at: str,
    generation_id: Optional[str],
    project_type: Optional[str],
    project_name: Optional[str],
    template_file: Optional[str],
    output_file: str,
    checks: Mapping[str, Any],
    metrics: Mapping[str, Any],
    issues: List[Dict[str, str]],
    field_provenance_file: Optional[str] = None,
    processor_report: Optional[list[Mapping[str, Any]]] = None,
    template_contract: Optional[Mapping[str, Any]] = None,
    rule_provenance: Optional[Mapping[str, Any]] = None,
) -> Dict[str, Any]:
    has_error = any(i.get("level") == "error" for i in issues)
    has_warning = any(i.get("level") == "warning" for i in issues)
    status = "FAIL" if has_error else ("WARN" if has_warning else "PASS")
    return {
        "schema_version": "1.0",
        "generated_at": generated_at,
        "generation_id": generation_id,
        "status": status,
        "project_type": project_type,
        "project_name": project_name,
        "template_file": str(template_file) if template_file else None,
        "output_file": str(output_file),
        "build_provenance": _build_provenance(),
        "field_provenance_file": str(field_provenance_file)
        if field_provenance_file
        else None,
        "post_processors": list(processor_report or []),
        "template_contract": dict(template_contract) if template_contract else None,
        "rules": dict(rule_provenance) if rule_provenance else None,
        "checks": dict(checks),
        "metrics": dict(metrics),
        "issues": issues,
    }


def _template_contract_check(
    template_contract: Optional[Mapping[str, Any]],
) -> Dict[str, Any]:
    if not template_contract:
        return {
            "status": "SKIP",
            "message": "No template contract report was provided.",
        }

    declared = template_contract.get("declared_contract") or {}
    missing_paths = list(template_contract.get("missing_paths") or [])
    missing_lists = list(template_contract.get("missing_lists") or [])
    missing_row_fields = dict(template_contract.get("missing_row_fields") or {})
    missing_required_variables = list(declared.get("missing_required_variables") or [])
    missing_required_lists = list(declared.get("missing_required_lists") or [])
    missing_required_markers = list(declared.get("missing_required_markers") or [])
    duplicate_required_markers = list(declared.get("duplicate_required_markers") or [])
    missing_required_tables = list(declared.get("missing_required_tables") or [])
    table_errors = dict(declared.get("table_errors") or {})

    ok = bool(template_contract.get("ok"))
    return {
        "status": "PASS" if ok else "FAIL",
        "ok": ok,
        "missing_paths": missing_paths,
        "missing_lists": missing_lists,
        "missing_row_fields": missing_row_fields,
        "missing_required_variables": missing_required_variables,
        "missing_required_lists": missing_required_lists,
        "missing_required_markers": missing_required_markers,
        "duplicate_required_markers": duplicate_required_markers,
        "marker_counts": dict(declared.get("marker_counts") or {}),
        "missing_required_tables": missing_required_tables,
        "table_errors": table_errors,
        "message": (
            "Template contract passed."
            if ok
            else "Template contract failed: "
            f"missing_paths={missing_paths}, missing_lists={missing_lists}, "
            f"missing_required_variables={missing_required_variables}, "
            f"missing_required_lists={missing_required_lists}, "
            f"missing_required_markers={missing_required_markers}, "
            f"duplicate_required_markers={duplicate_required_markers}, "
            f"missing_required_tables={missing_required_tables}, "
            f"table_errors={table_errors}"
        ),
    }


def _rule_provenance_check(
    rule_provenance: Optional[Mapping[str, Any]],
) -> Dict[str, Any]:
    if not rule_provenance:
        return {
            "status": "SKIP",
            "message": "No panel rule provenance was provided.",
            "file_count": 0,
            "files": [],
        }

    ok = bool(rule_provenance.get("ok", True))
    files = list(rule_provenance.get("files") or [])
    issues = list(rule_provenance.get("issues") or [])
    return {
        "status": "PASS" if ok else "FAIL",
        "panel_id": rule_provenance.get("panel_id"),
        "file_count": len(files),
        "files": files,
        "issues": issues,
        "message": (
            f"Panel rules loaded: {len(files)} file(s)."
            if ok
            else f"Panel rule validation failed with {len(issues)} issue(s)."
        ),
    }


def _build_visual_render_check(
    output_path: Path,
    *,
    mode: Optional[str],
    required: bool,
    dpi: int,
    timeout_seconds: int,
    output_dir: Optional[str],
    tmp_dir: Optional[str],
    expected_sparse_pages: Optional[Iterable[Mapping[str, Any]]] = None,
) -> Dict[str, Any]:
    normalized_mode = str(mode or "none").strip().lower()
    result: Dict[str, Any] = {
        "status": "SKIP",
        "requested": normalized_mode,
        "required": bool(required),
        "message": "Visual render QA was not requested.",
        "rendered_pages": [],
        "output_dir": None,
        "error": None,
    }
    if normalized_mode == "none":
        return result
    result["renderer_fingerprint"] = renderer_fingerprint()
    if normalized_mode not in {"first", "all"}:
        result.update(
            {
                "status": "FAIL" if required else "WARN",
                "message": f"Unsupported visual render mode: {mode!r}.",
                "error": "unsupported_mode",
            }
        )
        return result

    render_dir = (
        Path(output_dir)
        if output_dir
        else output_path.parent / "rendered_pages" / output_path.stem
    )
    first_page = 1 if normalized_mode == "first" else None
    last_page = 1 if normalized_mode == "first" else None
    expected_sparse_specs = [
        dict(spec)
        for spec in (expected_sparse_pages or [])
        if isinstance(spec, Mapping)
    ]

    try:
        pngs = render_docx_to_pngs(
            output_path,
            output_dir=render_dir,
            dpi=int(dpi),
            first_page=first_page,
            last_page=last_page,
            keep_pdf=bool(expected_sparse_specs),
            timeout_seconds=int(timeout_seconds),
            tmp_dir=Path(tmp_dir) if tmp_dir else None,
        )
    except Exception as exc:
        result.update(
            {
                "status": "FAIL" if required else "WARN",
                "message": f"Visual render failed: {exc}",
                "output_dir": str(render_dir),
                "error": str(exc),
            }
        )
        for attr in ("stage", "command", "stdout", "stderr"):
            value = getattr(exc, attr, None)
            if not value:
                continue
            if attr == "command":
                result[attr] = list(value)
            elif attr in {"stdout", "stderr"}:
                result[f"{attr}_tail"] = str(value)[-2000:]
            else:
                result[attr] = value
        return result

    result.update(
        {
            "rendered_pages": [str(path) for path in pngs],
            "output_dir": str(render_dir),
        }
    )
    if not pngs:
        result.update(
            {
                "status": "FAIL" if required else "WARN",
                "message": "Visual render completed but produced no PNG pages.",
                "error": "no_rendered_pages",
            }
        )
        return result

    page_texts: Dict[int, str] = {}
    page_text_error: Optional[str] = None
    if expected_sparse_specs:
        pdf_path = render_dir / f"{output_path.stem}.pdf"
        page_texts, page_text_error = _extract_pdf_page_texts(pdf_path, pngs)
        result["page_text_extraction"] = {
            "status": "PASS" if page_text_error is None else "FAIL",
            "page_count": len(page_texts),
            "error": page_text_error,
        }

    pixel_check = _inspect_rendered_png_pages(
        pngs,
        page_texts=page_texts,
        expected_sparse_pages=expected_sparse_specs,
    )
    result["pixel_check"] = pixel_check
    if pixel_check["status"] in {"WARN", "FAIL"}:
        result.update(
            {
                "status": "FAIL" if required else "WARN",
                "message": pixel_check["message"],
            }
        )
        return result

    result.update(
        {
            "status": "PASS",
            "message": (
                "Visual render produced inspectable PNG pages; reviewed "
                "semantic sparse-page exceptions were recorded."
                if pixel_check.get("expected_sparse_pages")
                else "Visual render produced inspectable PNG pages."
            ),
        }
    )
    return result


def _visual_render_issues(check: Mapping[str, Any]) -> Iterable[Dict[str, str]]:
    status = check.get("status")
    if status not in {"WARN", "FAIL"}:
        return

    level = "error" if status == "FAIL" else "warning"
    error = str(check.get("error") or "")
    pixel_check = check.get("pixel_check")
    if isinstance(pixel_check, Mapping):
        if pixel_check.get("blank_pages"):
            code = "VISUAL_RENDER_BLANK_PAGES"
        elif pixel_check.get("unexpected_low_content_pages") or pixel_check.get(
            "low_content_pages"
        ):
            code = "VISUAL_RENDER_LOW_CONTENT"
        elif pixel_check.get("unreadable_pages"):
            code = "VISUAL_RENDER_UNREADABLE_PAGES"
        else:
            code = "VISUAL_RENDER_WARN"
    elif error == "no_rendered_pages":
        code = "VISUAL_RENDER_NO_PAGES"
    elif error == "unsupported_mode":
        code = "VISUAL_RENDER_UNSUPPORTED_MODE"
    else:
        code = "VISUAL_RENDER_FAILED"

    yield {
        "level": level,
        "code": code,
        "message": str(check.get("message") or "Visual render QA reported an issue."),
    }


def _blank_page_detection_check(
    visual_check: Mapping[str, Any],
) -> Optional[Dict[str, Any]]:
    if visual_check.get("requested") == "none":
        return None

    pixel_check = visual_check.get("pixel_check")
    if not isinstance(pixel_check, Mapping):
        return {
            "status": "SKIP",
            "message": "Visual render did not produce inspectable page pixels.",
        }
    if pixel_check.get("status") == "SKIP":
        return {
            "status": "SKIP",
            "message": str(
                pixel_check.get("message")
                or "Rendered page pixel inspection was skipped."
            ),
        }

    blank_pages = list(pixel_check.get("blank_pages") or [])
    low_content_pages = list(
        pixel_check.get("unexpected_low_content_pages")
        or pixel_check.get("low_content_pages")
        or []
    )
    unreadable_pages = list(pixel_check.get("unreadable_pages") or [])
    expected_sparse_pages = list(pixel_check.get("expected_sparse_pages") or [])
    if blank_pages or low_content_pages or unreadable_pages:
        return {
            "status": visual_check.get("status") or "WARN",
            "message": "Visual page inspection found blank or low-content page images.",
            "blank_pages": blank_pages,
            "low_content_pages": low_content_pages,
            "unexpected_low_content_pages": low_content_pages,
            "expected_sparse_pages": expected_sparse_pages,
            "unreadable_pages": unreadable_pages,
            "thresholds": pixel_check.get("thresholds") or {},
        }
    return {
        "status": "PASS",
        "message": "Visual page inspection did not find blank rendered pages.",
        "checked_pages": pixel_check.get("checked_pages") or 0,
        "expected_sparse_pages": expected_sparse_pages,
        "thresholds": pixel_check.get("thresholds") or {},
    }


def _visual_expected_sparse_page_specs(
    context: Mapping[str, Any],
) -> List[Dict[str, Any]]:
    panel_style = context.get("panel_style")
    if not isinstance(panel_style, Mapping):
        return []
    visual_qa = panel_style.get("visual_qa")
    if not isinstance(visual_qa, Mapping):
        return []
    specs = visual_qa.get("expected_sparse_pages")
    if not isinstance(specs, list):
        return []
    return [dict(spec) for spec in specs if isinstance(spec, Mapping)]


def _rendered_page_number(path: Path) -> int:
    match = re.search(r"-(\d+)\.png$", path.name)
    return int(match.group(1)) if match else 0


def _extract_pdf_page_texts(
    pdf_path: Path,
    pngs: Iterable[Path],
) -> tuple[Dict[int, str], Optional[str]]:
    """Extract page text for semantic sparse-page classification.

    Raw page text is intentionally kept in memory and never written to the QA
    JSON because reports can contain patient information.
    """
    command = shutil.which("pdftotext")
    if not command:
        return {}, "pdftotext is unavailable"
    if not pdf_path.is_file():
        return {}, f"Rendered PDF is unavailable: {pdf_path.name}"

    texts: Dict[int, str] = {}
    for path in pngs:
        page_number = _rendered_page_number(Path(path))
        if page_number <= 0:
            continue
        try:
            process = subprocess.run(
                [
                    command,
                    "-layout",
                    "-f",
                    str(page_number),
                    "-l",
                    str(page_number),
                    str(pdf_path),
                    "-",
                ],
                check=False,
                capture_output=True,
                text=True,
                timeout=30,
            )
        except Exception as exc:
            return texts, f"pdftotext failed: {exc}"
        if process.returncode != 0:
            return texts, "pdftotext returned a non-zero exit code"
        texts[page_number] = process.stdout or ""
    return texts, None


def _inspect_rendered_png_pages(
    pngs: Iterable[Path],
    *,
    page_texts: Optional[Mapping[int, str]] = None,
    expected_sparse_pages: Optional[Iterable[Mapping[str, Any]]] = None,
) -> Dict[str, Any]:
    thresholds = {
        "blank_nonwhite_ratio": 0.001,
        "blank_dark_ratio": 0.0002,
        "low_content_dark_ratio": 0.0015,
        # Header/footer text and the light cyan watermark previously made a
        # body-empty page look non-blank. Inspect the central body separately.
        "body_crop_top_fraction": 0.10,
        "body_crop_bottom_fraction": 0.10,
        "body_low_content_dark_ratio": 0.01,
    }
    try:
        from PIL import Image
    except Exception as exc:
        return {
            "status": "SKIP",
            "message": f"Pillow is unavailable; skipped rendered PNG pixel checks: {exc}",
            "checked_pages": 0,
            "thresholds": thresholds,
        }

    pages: List[Dict[str, Any]] = []
    unreadable: List[str] = []
    for path in pngs:
        try:
            pages.append(_rendered_png_page_metrics(Path(path), Image, thresholds))
        except Exception as exc:
            unreadable.append(f"{path}: {exc}")

    blank_pages = [row["path"] for row in pages if row.get("blank")]
    low_content_pages = [row["path"] for row in pages if row.get("low_content")]
    # Covers/back covers can intentionally be sparse. The new body-content
    # gate is therefore applied only to interior pages.
    interior_body_low_content_pages = [
        row["path"]
        for idx, row in enumerate(pages)
        if 0 < idx < len(pages) - 1 and row.get("body_low_content")
    ]
    low_content_pages = list(
        dict.fromkeys(low_content_pages + interior_body_low_content_pages)
    )
    low_content_pages = [path for path in low_content_pages if path not in blank_pages]

    expected_specs = [
        dict(spec)
        for spec in (expected_sparse_pages or [])
        if isinstance(spec, Mapping)
    ]
    spec_match_counts: Dict[str, int] = {}
    accepted_sparse_pages: List[Dict[str, Any]] = []
    unexpected_low_content_pages: List[str] = []
    normalized_page_texts = {
        int(page): re.sub(r"\s+", "", str(text or ""))
        for page, text in (page_texts or {}).items()
    }
    for path in low_content_pages:
        page_number = _rendered_page_number(Path(path))
        page_text = normalized_page_texts.get(page_number, "")
        accepted_spec: Optional[Dict[str, Any]] = None
        for index, spec in enumerate(expected_specs):
            required_text = [
                re.sub(r"\s+", "", str(value or ""))
                for value in (spec.get("required_text") or [])
                if str(value or "").strip()
            ]
            if not required_text or not all(
                marker in page_text for marker in required_text
            ):
                continue
            spec_id = str(spec.get("id") or f"expected_sparse_{index + 1}")
            try:
                max_matches = max(1, int(spec.get("max_matches", 1)))
            except (TypeError, ValueError):
                max_matches = 1
            if spec_match_counts.get(spec_id, 0) >= max_matches:
                continue
            spec_match_counts[spec_id] = spec_match_counts.get(spec_id, 0) + 1
            accepted_spec = {
                "id": spec_id,
                "path": path,
                "page_number": page_number,
                "required_text_count": len(required_text),
            }
            break
        if accepted_spec is None:
            unexpected_low_content_pages.append(path)
        else:
            accepted_sparse_pages.append(accepted_spec)

    if unreadable or blank_pages or unexpected_low_content_pages:
        parts: List[str] = []
        if unreadable:
            parts.append(f"{len(unreadable)} rendered PNG page(s) could not be read")
        if blank_pages:
            parts.append(f"{len(blank_pages)} rendered page(s) look blank")
        if unexpected_low_content_pages:
            parts.append(
                f"{len(unexpected_low_content_pages)} rendered page(s) look "
                "unexpectedly low-content"
            )
        return {
            "status": "WARN",
            "message": "; ".join(parts) + ".",
            "checked_pages": len(pages),
            "pages": pages,
            "blank_pages": blank_pages,
            "low_content_pages": unexpected_low_content_pages,
            "unexpected_low_content_pages": unexpected_low_content_pages,
            "expected_sparse_pages": accepted_sparse_pages,
            "unreadable_pages": unreadable,
            "thresholds": thresholds,
        }

    return {
        "status": "PASS",
        "message": (
            "Rendered PNG pixel checks passed with reviewed semantic sparse-page "
            "exceptions."
            if accepted_sparse_pages
            else "Rendered PNG pixel checks passed."
        ),
        "checked_pages": len(pages),
        "pages": pages,
        "blank_pages": [],
        "low_content_pages": [],
        "unexpected_low_content_pages": [],
        "expected_sparse_pages": accepted_sparse_pages,
        "unreadable_pages": [],
        "thresholds": thresholds,
    }


def _rendered_png_page_metrics(
    path: Path,
    image_module: Any,
    thresholds: Mapping[str, float],
) -> Dict[str, Any]:
    with image_module.open(path) as image:
        original_size = [int(image.width), int(image.height)]
        sampled = image.convert("RGBA")
        body_top = int(image.height * float(thresholds["body_crop_top_fraction"]))
        body_bottom = int(
            image.height * (1.0 - float(thresholds["body_crop_bottom_fraction"]))
        )
        body_sampled = sampled.crop((0, body_top, image.width, body_bottom))
        sampled.thumbnail((300, 300))
        body_sampled.thumbnail((300, 300))
        sampled_data = getattr(sampled, "get_flattened_data", sampled.getdata)
        body_data = getattr(body_sampled, "get_flattened_data", body_sampled.getdata)
        pixels = list(sampled_data())
        body_pixels = list(body_data())

    visible = [pixel for pixel in pixels if pixel[3] > 10]
    visible_count = len(visible) or 1
    nonwhite_count = sum(
        1 for red, green, blue, _alpha in visible if min(red, green, blue) < 248
    )
    dark_count = sum(
        1 for red, green, blue, _alpha in visible if (red + green + blue) / 3 < 210
    )
    body_visible = [pixel for pixel in body_pixels if pixel[3] > 10]
    body_visible_count = len(body_visible) or 1
    body_dark_count = sum(
        1 for red, green, blue, _alpha in body_visible if (red + green + blue) / 3 < 210
    )
    nonwhite_ratio = nonwhite_count / visible_count
    dark_ratio = dark_count / visible_count
    blank = nonwhite_ratio < float(
        thresholds["blank_nonwhite_ratio"]
    ) and dark_ratio < float(thresholds["blank_dark_ratio"])
    low_content = not blank and dark_ratio < float(thresholds["low_content_dark_ratio"])
    body_dark_ratio = body_dark_count / body_visible_count
    body_low_content = body_dark_ratio < float(
        thresholds["body_low_content_dark_ratio"]
    )
    return {
        "path": str(path),
        "file_size_bytes": path.stat().st_size,
        "image_size": original_size,
        "sampled_pixels": visible_count,
        "nonwhite_ratio": round(nonwhite_ratio, 6),
        "dark_ratio": round(dark_ratio, 6),
        "body_dark_ratio": round(body_dark_ratio, 6),
        "blank": blank,
        "low_content": low_content,
        "body_low_content": body_low_content,
    }


def _iter_all_paragraphs(doc) -> Iterable[Any]:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    paragraph._qa_location = (
                        f"table:{table_idx}:r{row_idx}:c{cell_idx}:p{para_idx}"
                    )
                    yield paragraph


def _field_source_summary(field_provenance: Mapping[str, Any]) -> Dict[str, str]:
    fields = field_provenance.get("fields") or {}
    if not isinstance(fields, Mapping):
        return {}
    return {
        str(field): str(info.get("source") or "unknown")
        for field, info in fields.items()
        if isinstance(info, Mapping)
    }


def _read_docx_text(doc) -> str:
    parts: List[str] = []
    for paragraph in _iter_all_paragraphs(doc):
        value = paragraph.text or ""
        if value:
            parts.append(value)
    return "\n".join(parts)


def _read_part3_text(doc, section_policy: Optional[Mapping[str, Any]] = None) -> Optional[str]:
    """Return the body text inside the dynamic Part-3 section, if present."""

    started = False
    policy = section_policy or {}
    start_heading = _compact(str(policy.get("start_heading") or ""))
    end_heading = _compact(str(policy.get("end_heading") or ""))
    parts: List[str] = []
    for paragraph in doc.paragraphs:
        value = str(paragraph.text or "").strip()
        compact = _compact(value)
        if not started:
            if (start_heading and compact == start_heading) or (
                not start_heading
                and compact.startswith("第三部分：基因变异及相应靶向/免疫药物解析")
            ):
                started = True
                parts.append(value)
            continue
        if (end_heading and compact == end_heading) or (
            not end_heading
            and (compact.startswith("3.阅读说明") or compact.startswith("第四部分：附录"))
        ):
            break
        if value:
            parts.append(value)
    return "\n".join(parts) if started else None


def _compact(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def _trim_sample(value: str, limit: int = 120) -> str:
    value = re.sub(r"\s+", " ", value or "").strip()
    return value if len(value) <= limit else value[: limit - 1] + "..."


def _paragraph_location(paragraph: Any) -> str:
    return str(getattr(paragraph, "_qa_location", "body"))


def _is_empty_numbered_paragraph(paragraph: Any) -> bool:
    if (paragraph.text or "").strip():
        return False
    if _paragraph_has_embedded_object(paragraph):
        return False
    ppr = paragraph._p.pPr
    return bool(ppr is not None and ppr.numPr is not None)


def _paragraph_has_embedded_object(paragraph: Any) -> bool:
    xml = paragraph._p.xml
    return "<w:drawing" in xml or "<w:pict" in xml


def _inspect_reportgen_toc_fields(output_path: Path) -> Optional[Dict[str, Any]]:
    """Validate ReportGen-owned PAGEREF fields directly from DOCX XML.

    ``python-docx`` omits paragraphs nested in Word content controls, which is
    where the reviewed TOC lives. XML inspection keeps the QA gate aware of the
    actual reader-refreshable fields and their target bookmarks.
    """
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def qn(tag: str) -> str:
        return f"{{{w_ns}}}{tag}"

    try:
        with ZipFile(output_path, "r") as zf:
            document_root = ET.fromstring(zf.read("word/document.xml"))
            settings_root = ET.fromstring(zf.read("word/settings.xml"))
    except Exception as exc:
        return {
            "status": "FAIL",
            "mode": "PAGEREF",
            "message": f"Unable to inspect TOC field XML: {exc}",
        }

    anchors: List[str] = []
    for node in document_root.iter(qn("instrText")):
        instruction = node.text or ""
        match = re.search(r"\bPAGEREF\s+(_ReportGenToc_\S+)", instruction)
        if match:
            anchors.append(match.group(1))
    if not anchors:
        return None

    bookmark_ids: Dict[str, str] = {}
    for node in document_root.iter(qn("bookmarkStart")):
        name = node.get(qn("name")) or ""
        if name.startswith("_ReportGenToc_"):
            bookmark_ids[name] = node.get(qn("id")) or ""
    bookmark_end_ids = {
        node.get(qn("id")) or "" for node in document_root.iter(qn("bookmarkEnd"))
    }
    unique_anchors = set(anchors)
    missing_bookmarks = sorted(unique_anchors - set(bookmark_ids))
    unclosed_bookmarks = sorted(
        name
        for name in unique_anchors & set(bookmark_ids)
        if bookmark_ids[name] not in bookmark_end_ids
    )
    duplicate_targets = sorted(
        anchor for anchor in unique_anchors if anchors.count(anchor) > 1
    )
    update_values = [
        str(node.get(qn("val")) or "").strip().lower()
        for node in settings_root.iter(qn("updateFields"))
    ]
    update_fields = any(value in {"1", "true", "on"} for value in update_values)

    problems = []
    if missing_bookmarks:
        problems.append(f"missing bookmarks={missing_bookmarks}")
    if unclosed_bookmarks:
        problems.append(f"unclosed bookmarks={unclosed_bookmarks}")
    if duplicate_targets:
        problems.append(f"duplicate targets={duplicate_targets}")
    if not update_fields:
        problems.append("updateFields is not enabled")

    return {
        "status": "FAIL" if problems else "PASS",
        "mode": "PAGEREF",
        "field_count": len(anchors),
        "unique_target_count": len(unique_anchors),
        "bookmark_count": len(bookmark_ids),
        "missing_bookmarks": missing_bookmarks,
        "unclosed_bookmarks": unclosed_bookmarks,
        "duplicate_targets": duplicate_targets,
        "update_fields": update_fields,
        "message": (
            "ReportGen TOC PAGEREF fields and bookmarks are valid."
            if not problems
            else "Invalid ReportGen TOC fields: " + "; ".join(problems)
        ),
    }


def _inspect_toc(
    paragraphs: List[Any], *, output_path: Optional[Path] = None
) -> Dict[str, Any]:
    if output_path is not None:
        field_check = _inspect_reportgen_toc_fields(Path(output_path))
        if field_check is not None:
            return field_check
        native_check = _inspect_native_toc(Path(output_path))
        if native_check is not None:
            return native_check

    texts = [(p.text or "").strip() for p in paragraphs]
    start_idx = None
    for idx, text in enumerate(texts):
        if _compact(text).startswith("目录"):
            start_idx = idx
            break
    if start_idx is None:
        return {"status": "SKIP", "message": "No TOC heading found."}

    toc_lines: List[str] = []
    for text in texts[start_idx + 1 : start_idx + 100]:
        if not text:
            continue
        if text.startswith(("第一部分", "一、基本信息", "致您的一封信")):
            break
        toc_lines.append(text)

    page_like = [
        line
        for line in toc_lines
        if re.search(
            r"(第[一二三四五六七八九十]+部分|\d+(?:\.\d+)+|[一二三四五六七八九十]+、).*\d+\s*$",
            line,
        )
    ]
    if not toc_lines:
        return {"status": "WARN", "line_count": 0, "message": "TOC area is empty."}
    if not page_like:
        return {
            "status": "WARN",
            "line_count": len(toc_lines),
            "message": "TOC lines were found, but no line appears to contain page numbers.",
            "sample_lines": toc_lines[:8],
        }
    return {
        "status": "PASS",
        "line_count": len(toc_lines),
        "page_numbered_line_count": len(page_like),
        "sample_lines": toc_lines[:8],
    }


def _inspect_native_toc(output_path: Path) -> Optional[Dict[str, Any]]:
    """Read native Word TOCs inside SDTs omitted by python-docx paragraphs."""
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(output_path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    bookmarks = {
        node.get("{" + ns["w"] + "}name"): node.get("{" + ns["w"] + "}id")
        for node in root.findall(".//w:bookmarkStart", ns)
    }
    closed_ids = {
        node.get("{" + ns["w"] + "}id")
        for node in root.findall(".//w:bookmarkEnd", ns)
    }
    for block in root.findall(".//w:sdt", ns):
        instructions = [n.text or "" for n in block.findall(".//w:instrText", ns)]
        if not any(re.match(r"\s*TOC\s", value) for value in instructions):
            continue
        lines, numbered = [], []
        for paragraph in block.findall(".//w:p", ns):
            # A heading such as "检测项目 62" is not a cached page number.
            # Word/LibreOffice separate each TOC label and its page with a tab.
            value = "".join(
                "\t" if node.tag == "{" + ns["w"] + "}tab" else (node.text or "")
                for node in paragraph.iter()
                if node.tag in {"{" + ns["w"] + "}t", "{" + ns["w"] + "}tab"}
            ).strip()
            if not value or _compact(value) == "目录":
                continue
            lines.append(value)
            if re.search(r"\t\s*[1-9]\d*\s*$", value):
                numbered.append(value)
        targets = {
            n.get("{" + ns["w"] + "}anchor")
            for n in block.findall(".//w:hyperlink", ns)
        } - {None, ""}
        missing = sorted(targets - set(bookmarks))
        unclosed = sorted(t for t in targets & set(bookmarks) if bookmarks[t] not in closed_ids)
        complete = bool(lines) and len(numbered) == len(lines)
        return {
            "status": "FAIL" if missing or unclosed else ("PASS" if complete else "WARN"),
            "mode": "native_TOC",
            "line_count": len(lines),
            "page_numbered_line_count": len(numbered),
            "missing_bookmarks": missing,
            "unclosed_bookmarks": unclosed,
            "message": (
                "Native TOC has broken bookmark targets." if missing or unclosed else
                "Native TOC contains cached page numbers for every entry." if complete else
                "Native TOC has missing cached page numbers."
            ),
        }
    return None


def _inspect_tables(doc) -> Dict[str, Any]:
    summary = {
        "targeted_drug_tip_tables": [],
        "variant_summary_tables": [],
        "variant_detail_tables": [],
        "bad_variant_detail_tables": [],
        "biomarker_tables": [],
    }

    for idx, table in enumerate(doc.tables):
        col_count = _table_col_count(table)
        row_count = len(table.rows)
        preview = _table_preview(table, max_rows=4)
        compact_preview = _compact(preview)
        entry = {"index": idx, "rows": row_count, "cols": col_count}

        if all(token in compact_preview for token in ("突变位点", "潜在获益")):
            summary["targeted_drug_tip_tables"].append(entry)

        if all(
            token in compact_preview for token in ("基因突变信息", "潜在获益靶向药物")
        ):
            summary["variant_summary_tables"].append(entry)

        has_detail_header = all(
            token in compact_preview
            for token in ("基因名称", "转录本号", "染色体", "外显子")
        )
        if has_detail_header:
            if col_count >= 8:
                summary["variant_detail_tables"].append(entry)
            else:
                summary["bad_variant_detail_tables"].append(entry)

        if all(token in compact_preview for token in ("TMB", "MSI", "用药提示")):
            summary["biomarker_tables"].append(entry)

    return summary


def _table_col_count(table: Any) -> int:
    if not table.rows:
        return 0
    return max((len(row.cells) for row in table.rows), default=0)


def _table_preview(table: Any, *, max_rows: int) -> str:
    rows: List[str] = []
    for row in table.rows[:max_rows]:
        rows.append(" ".join((cell.text or "").strip() for cell in row.cells))
    return "\n".join(rows)


def _is_crc(project_type: Optional[str]) -> bool:
    return (project_type or "").lower() in CRC_PROJECT_TYPES


def _build_table_checks(
    table_metrics: Mapping[str, Any],
    project_type: Optional[str],
    context: Mapping[str, Any],
) -> Dict[str, Any]:
    has_variants = bool(context.get("variants") or context.get("summary_variants"))
    crc = _is_crc(project_type)

    def status_for(count: int, required: bool) -> str:
        if not required:
            return "SKIP"
        return "PASS" if count else "FAIL"

    return {
        "variant_detail_table_shape": {
            "status": "FAIL"
            if table_metrics.get("bad_variant_detail_tables")
            else status_for(
                len(table_metrics.get("variant_detail_tables") or []),
                crc and has_variants,
            ),
            "tables": table_metrics.get("variant_detail_tables") or [],
            "bad_tables": table_metrics.get("bad_variant_detail_tables") or [],
            "required": crc and has_variants,
        },
        "variant_summary_table_present": {
            "status": status_for(
                len(table_metrics.get("variant_summary_tables") or []), crc
            ),
            "tables": table_metrics.get("variant_summary_tables") or [],
            "required": crc,
        },
        "targeted_drug_tip_table_present": {
            "status": status_for(
                len(table_metrics.get("targeted_drug_tip_tables") or []), crc
            ),
            "tables": table_metrics.get("targeted_drug_tip_tables") or [],
            "required": crc,
        },
        "biomarker_table_present": {
            "status": status_for(len(table_metrics.get("biomarker_tables") or []), crc),
            "tables": table_metrics.get("biomarker_tables") or [],
            "required": crc,
        },
    }


def _table_issues(checks: Mapping[str, Any]) -> Iterable[Dict[str, str]]:
    messages = {
        "variant_detail_table_shape": "CRC variant detail table is missing or has an unexpected column shape.",
        "variant_summary_table_present": "CRC 2.1 variant summary table was not detected.",
        "targeted_drug_tip_table_present": "CRC targeted drug tip table was not detected.",
        "biomarker_table_present": "CRC TMB/MSI biomarker table was not detected.",
    }
    for key, check in checks.items():
        if check.get("status") == "FAIL":
            yield {"level": "error", "code": key.upper(), "message": messages[key]}


def _build_style_checks(
    doc: Any,
    project_type: Optional[str],
    context: Mapping[str, Any],
) -> Dict[str, Any]:
    if not _is_crc(project_type):
        return {}

    style = context.get("panel_style")
    if not isinstance(style, Mapping):
        return {
            "docx_style_rules": {
                "status": "SKIP",
                "message": "No panel_style rules were provided in report context.",
                "checked_table_count": 0,
                "failures": [],
            }
        }

    failures: List[Dict[str, Any]] = []
    checked = 0
    table_counts = {
        "variant_summary_table": 0,
        "variant_detail_table": 0,
        "biomarker_table": 0,
        "approved_drug_table": 0,
    }

    for table_idx, table in enumerate(doc.tables):
        if _is_variant_summary_style_table(table):
            table_counts["variant_summary_table"] += 1
            checked += 1
            failures.extend(
                _check_variant_summary_table_style(
                    table,
                    table_idx=table_idx,
                    style=_style_config(style, "variant_summary_table"),
                )
            )
        if _is_variant_detail_style_table(table):
            table_counts["variant_detail_table"] += 1
            checked += 1
            failures.extend(
                _check_variant_detail_table_style(
                    table,
                    table_idx=table_idx,
                    style=_style_config(style, "variant_detail_table"),
                )
            )
        if _is_biomarker_style_table(table):
            table_counts["biomarker_table"] += 1
            checked += 1
            failures.extend(
                _check_biomarker_table_style(
                    table,
                    table_idx=table_idx,
                    style=_style_config(style, "biomarker_table"),
                )
            )
        if _is_approved_drug_table(table):
            table_counts["approved_drug_table"] += 1
            checked += 1
            failures.extend(
                _check_approved_drug_brand_style(table, table_idx=table_idx)
            )

    if not checked:
        # 走到这里的前提：这是 CRC 项目（_is_crc 已判真）且上下文提供了
        # panel_style 样式规则——也就是说这份报告本应包含受样式管理的 2.1
        # 变异表 / 生物标志物表。正常产物会认出 3 张这类表；一张都认不出，
        # 几乎只可能是模板表头/列结构被改动导致认表逻辑（依赖硬编码列数+表头
        # 串）静默失配。此前这里返回 SKIP（绿灯），会让"模板被改坏、所有表
        # 样式丢失"的报告悄悄通过 QA。改判为 FAIL，让这种静默失效立刻报警。
        return {
            "docx_style_rules": {
                "status": "FAIL",
                "message": (
                    "未识别到任何受样式管理的 CRC 表格（2.1 变异表/生物标志物表）。"
                    "正常报告应含这些表；通常是模板表头或列结构改动导致表格识别失配，"
                    "样式可能整体未生效，请核对模板与 panel_style 规则。"
                ),
                "checked_table_count": 0,
                "table_counts": table_counts,
                "failures": [],
            }
        }

    return {
        "docx_style_rules": {
            "status": "FAIL" if failures else "PASS",
            "message": (
                f"DOCX style QA found {len(failures)} issue(s)."
                if failures
                else "DOCX style QA passed."
            ),
            "checked_table_count": checked,
            "table_counts": table_counts,
            "failures": failures[:30],
            "failure_count": len(failures),
        }
    }


def _style_issues(checks: Mapping[str, Any]) -> Iterable[Dict[str, str]]:
    check = checks.get("docx_style_rules")
    if isinstance(check, Mapping) and check.get("status") == "FAIL":
        yield {
            "level": "error",
            "code": "DOCX_STYLE_RULES",
            "message": str(check.get("message") or "DOCX style rules failed."),
        }


def _is_approved_drug_table(table: Any) -> bool:
    return bool(
        table.rows
        and _table_col_count(table) == 3
        and (table.rows[0].cells[0].text or "").strip() == "药物名称"
    )


def _check_approved_drug_brand_style(
    table: Any,
    *,
    table_idx: int,
) -> List[Dict[str, Any]]:
    bracket_re = re.compile(r"(?:\[[^\[\]\r\n]+\])|(?:［[^［］\r\n]+］)")
    failures: List[Dict[str, Any]] = []
    for row_idx, row in enumerate(table.rows[1:], start=1):
        for paragraph in row.cells[0].paragraphs:
            offset = 0
            run_spans: List[tuple[int, int, Any]] = []
            for run in paragraph.runs:
                text = run.text or ""
                run_spans.append((offset, offset + len(text), run))
                offset += len(text)
            full_text = "".join((run.text or "") for run in paragraph.runs)
            for match in bracket_re.finditer(full_text):
                overlapping = [
                    run
                    for start, end, run in run_spans
                    if start < match.end() and end > match.start()
                ]
                if overlapping and all(run.bold is True for run in overlapping):
                    continue
                failures.append(
                    _style_failure(
                        "approved_drug_table",
                        table_idx,
                        row_idx,
                        0,
                        "brand_bracket_bold",
                        True,
                        match.group(0),
                    )
                )
    return failures


def _style_config(root: Mapping[str, Any], table_name: str) -> Dict[str, Any]:
    defaults = root.get("defaults")
    merged = dict(defaults) if isinstance(defaults, Mapping) else {}
    table = root.get(table_name)
    if isinstance(table, Mapping):
        merged.update(table)
    return merged


def _hex_color(value: Any, default: str) -> str:
    text = str(value or default).strip().lstrip("#").upper()
    if len(text) != 6 or not all(ch in "0123456789ABCDEF" for ch in text):
        return default.upper()
    return text


def _bool_style(value: Any, default: bool) -> bool:
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    text = str(value).strip().lower()
    if text in {"1", "true", "yes", "y", "on"}:
        return True
    if text in {"0", "false", "no", "n", "off"}:
        return False
    return default


def _is_variant_summary_style_table(table: Any) -> bool:
    if _table_col_count(table) != 4 or not table.rows:
        return False
    header = _compact(" ".join(cell.text for cell in table.rows[0].cells))
    return all(
        token in header
        for token in ("基因", "突变位点", "潜在获益靶向药物", "可能耐药")
    )


def _is_variant_detail_style_table(table: Any) -> bool:
    if _table_col_count(table) != 9 or len(table.rows) < 2:
        return False
    row0 = _compact(" ".join(cell.text for cell in table.rows[0].cells))
    row1 = _compact(" ".join(cell.text for cell in table.rows[1].cells))
    return all(
        token in row0 for token in ("基因名称", "基因突变信息", "靶向药物信息")
    ) and all(token in row1 for token in ("转录本号", "潜在获益靶向药物"))


def _is_biomarker_style_table(table: Any) -> bool:
    if _table_col_count(table) != 3 or not table.rows:
        return False
    text = _compact("\n".join(cell.text for row in table.rows for cell in row.cells))
    return all(token in text for token in ("TMB", "MSI", "用药提示"))


def _check_variant_summary_table_style(
    table: Any,
    *,
    table_idx: int,
    style: Mapping[str, Any],
) -> List[Dict[str, Any]]:
    expected = {
        "header_fill": _hex_color(style.get("header_fill"), "00C4D8"),
        "header_font_color": _hex_color(style.get("header_font_color"), "FFFFFF"),
        "body_font_color": _hex_color(style.get("body_font_color"), "000000"),
        "link_color": _hex_color(style.get("link_color"), "0000FF"),
        "link_underline": _bool_style(style.get("link_underline"), True),
    }
    failures = _check_header_style(
        table,
        table_idx=table_idx,
        table_name="variant_summary_table",
        header_rows={0},
        expected_fill=expected["header_fill"],
        expected_font_color=expected["header_font_color"],
    )
    for row_idx, row in enumerate(table.rows[1:], start=1):
        for col_idx, cell in enumerate(row.cells):
            dash_only = (cell.text or "").strip() in {"", "-", "--", "—"}
            link_cell = col_idx == 0 or (col_idx in {2, 3} and not dash_only)
            failures.extend(
                _check_cell_runs(
                    cell,
                    table_idx=table_idx,
                    row_idx=row_idx,
                    col_idx=col_idx,
                    table_name="variant_summary_table",
                    expected_color=expected["link_color"]
                    if link_cell
                    else expected["body_font_color"],
                    expected_underline=expected["link_underline"]
                    if link_cell
                    else False,
                )
            )
    return failures


def _check_variant_detail_table_style(
    table: Any,
    *,
    table_idx: int,
    style: Mapping[str, Any],
) -> List[Dict[str, Any]]:
    expected = {
        "header_fill": _hex_color(style.get("header_fill"), "00C4D8"),
        "header_font_color": _hex_color(style.get("header_font_color"), "F9FBFA"),
        "body_font_color": _hex_color(style.get("body_font_color"), "000000"),
        "link_color": _hex_color(style.get("link_color"), "0000FF"),
        "link_underline": _bool_style(style.get("link_underline"), True),
    }
    failures = _check_header_style(
        table,
        table_idx=table_idx,
        table_name="variant_detail_table",
        header_rows={0, 1},
        expected_fill=expected["header_fill"],
        expected_font_color=expected["header_font_color"],
    )
    # 与 _restore_variant_detail_table_style 渲染逻辑保持一致：基因名(col0)仅在该行
    # 检出变异时才是蓝色链接；"未见突变"等行的基因名为黑色（用户要求，c5f6e9b）。
    # 此前 QA 把 col0 无条件当链接，导致每份含未见突变行的报告都报假阳性。
    plain_texts = {
        str(t).strip()
        for t in (
            style.get("plain_texts")
            or ["未见突变", "未检出", "未检出有害变异", "-", "--", "—"]
        )
    }

    def is_plain_text(text: str) -> bool:
        normalized = (text or "").strip()
        return normalized in plain_texts or any(
            token in normalized for token in plain_texts if len(token) > 1
        )

    for row_idx, row in enumerate(table.rows[2:], start=2):
        locus_text = (row.cells[4].text or "").strip() if len(row.cells) > 4 else ""
        row_has_variant = bool(locus_text) and not is_plain_text(locus_text)
        for col_idx, cell in enumerate(row.cells):
            dash_only = (cell.text or "").strip() in {"", "-", "--", "—"}
            link_cell = (col_idx == 0 and row_has_variant) or (
                col_idx in {7, 8} and not dash_only
            )
            failures.extend(
                _check_cell_runs(
                    cell,
                    table_idx=table_idx,
                    row_idx=row_idx,
                    col_idx=col_idx,
                    table_name="variant_detail_table",
                    expected_color=expected["link_color"]
                    if link_cell
                    else expected["body_font_color"],
                    expected_underline=expected["link_underline"]
                    if link_cell
                    else False,
                )
            )
    return failures


def _check_biomarker_table_style(
    table: Any,
    *,
    table_idx: int,
    style: Mapping[str, Any],
) -> List[Dict[str, Any]]:
    expected = {
        "header_fill": _hex_color(style.get("header_fill"), "00C4D8"),
        "header_font_color": _hex_color(style.get("header_font_color"), "F9FBFA"),
        "body_font_color": _hex_color(style.get("body_font_color"), "000000"),
    }
    failures = _check_header_style(
        table,
        table_idx=table_idx,
        table_name="biomarker_table",
        header_rows={0},
        expected_fill=expected["header_fill"],
        expected_font_color=expected["header_font_color"],
    )
    for row_idx, row in enumerate(table.rows[1:], start=1):
        for col_idx, cell in enumerate(row.cells):
            failures.extend(
                _check_cell_runs(
                    cell,
                    table_idx=table_idx,
                    row_idx=row_idx,
                    col_idx=col_idx,
                    table_name="biomarker_table",
                    expected_color=expected["body_font_color"],
                    expected_underline=False,
                )
            )
    return failures


def _check_header_style(
    table: Any,
    *,
    table_idx: int,
    table_name: str,
    header_rows: set[int],
    expected_fill: str,
    expected_font_color: str,
) -> List[Dict[str, Any]]:
    failures: List[Dict[str, Any]] = []
    for row_idx in header_rows:
        if row_idx >= len(table.rows):
            continue
        for col_idx, cell in enumerate(table.rows[row_idx].cells):
            actual_fill = _cell_fill(cell)
            if actual_fill and actual_fill != expected_fill:
                failures.append(
                    _style_failure(
                        table_name,
                        table_idx,
                        row_idx,
                        col_idx,
                        "header_fill",
                        expected_fill,
                        actual_fill,
                    )
                )
            elif not actual_fill:
                failures.append(
                    _style_failure(
                        table_name,
                        table_idx,
                        row_idx,
                        col_idx,
                        "header_fill",
                        expected_fill,
                        None,
                    )
                )
            failures.extend(
                _check_cell_runs(
                    cell,
                    table_idx=table_idx,
                    row_idx=row_idx,
                    col_idx=col_idx,
                    table_name=table_name,
                    expected_color=expected_font_color,
                    expected_underline=False,
                )
            )
    return failures


def _check_cell_runs(
    cell: Any,
    *,
    table_idx: int,
    row_idx: int,
    col_idx: int,
    table_name: str,
    expected_color: str,
    expected_underline: bool,
) -> List[Dict[str, Any]]:
    failures: List[Dict[str, Any]] = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if not (run.text or "").strip():
                continue
            actual_color = _run_color(run)
            if actual_color and actual_color != expected_color:
                failures.append(
                    _style_failure(
                        table_name,
                        table_idx,
                        row_idx,
                        col_idx,
                        "font_color",
                        expected_color,
                        actual_color,
                    )
                )
            elif not actual_color:
                failures.append(
                    _style_failure(
                        table_name,
                        table_idx,
                        row_idx,
                        col_idx,
                        "font_color",
                        expected_color,
                        None,
                    )
                )
            actual_underline = bool(run.font.underline)
            if actual_underline != expected_underline:
                failures.append(
                    _style_failure(
                        table_name,
                        table_idx,
                        row_idx,
                        col_idx,
                        "underline",
                        expected_underline,
                        actual_underline,
                    )
                )
    return failures


def _style_failure(
    table_name: str,
    table_idx: int,
    row_idx: int,
    col_idx: int,
    property_name: str,
    expected: Any,
    actual: Any,
) -> Dict[str, Any]:
    return {
        "table": table_name,
        "table_index": table_idx,
        "row": row_idx,
        "col": col_idx,
        "property": property_name,
        "expected": expected,
        "actual": actual,
    }


def _cell_fill(cell: Any) -> Optional[str]:
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        return None
    value = shd.get(qn("w:fill"))
    return str(value).upper() if value else None


def _run_color(run: Any) -> Optional[str]:
    value = run.font.color.rgb
    return str(value).upper() if value is not None else None


def _build_business_checks(
    compact_text: str,
    context: Mapping[str, Any],
    project_type: Optional[str],
    part3_compact_text: Optional[str] = None,
) -> Dict[str, Any]:
    checks: Dict[str, Any] = {}
    if context.get("cnv_review_required"):
        visible = "CNV待复核" in compact_text
        checks["cnv_source_review"] = {
            "status": "WARN" if visible else "FAIL",
            "genes": list(context.get("cnv_review_genes") or []),
            "message": (
                "CNV source records require review; gain/numeric caller flags are not "
                "confirmed amplification or immune-response conclusions."
                if visible else "Required CNV source-review notice is missing from the report."
            ),
        }
    residual_scan = context.get("part3_cross_cancer_residual_scan")
    if isinstance(residual_scan, Mapping) and residual_scan.get("enabled"):
        configured_terms = [
            str(term).strip()
            for term in residual_scan.get("terms") or []
            if str(term).strip()
        ]
        scan_scope = str(residual_scan.get("scan_scope") or "document").strip()
        scan_text = (
            part3_compact_text
            if scan_scope == "part3" and part3_compact_text is not None
            else compact_text
        )
        matched_terms = [
            term for term in configured_terms if _compact(term) in scan_text
        ]
        checks["part3_cross_cancer_residuals"] = {
            "status": "WARN" if matched_terms else "PASS",
            "severity": "warning",
            "scan_scope": scan_scope,
            "configured_term_count": len(configured_terms),
            "matched_terms": matched_terms,
            "message": str(
                residual_scan.get("notice")
                or "Part 3 contains cross-cancer historical wording pending review."
                ).strip(),
        }
    suppression = context.get("part3_cross_cancer_suppression")
    if isinstance(suppression, Mapping):
        suppressed_count = _as_int(suppression.get("suppressed_field_count")) or 0
        if suppressed_count:
            checks["part3_cross_cancer_suppression"] = {
                "status": "WARN",
                "severity": "warning",
                "suppressed_field_count": suppressed_count,
                "suppressed_row_count": (
                    _as_int(suppression.get("suppressed_row_count")) or 0
                ),
                "message": (
                    "Part-3 historical fields outside the lung-specific review "
                    "scope were hidden; lung-specific replacement wording remains "
                    "pending report-group review."
                ),
            }
    if not _is_crc(project_type):
        return checks

    total_count = _as_int(context.get("total_variants_count"))
    drug_count = _as_int(context.get("drug_related_count"))
    targeted_or_immune_count = _as_int(context.get("targeted_or_immune_related_count"))

    if total_count is not None:
        expected = f"本次共检出体细胞变异：{total_count}个"
        checks["total_variant_count_text"] = {
            "status": "PASS" if _compact(expected) in compact_text else "FAIL",
            "expected": expected,
        }

    if drug_count is not None:
        expected = f"与靶向药物用药相关的变异有：{drug_count}个"
        checks["drug_related_count_text"] = {
            "status": "PASS" if _compact(expected) in compact_text else "FAIL",
            "expected": expected,
        }

    if targeted_or_immune_count is not None:
        expected = f"与靶向/免疫药物相关的变异：{targeted_or_immune_count}个"
        checks["targeted_or_immune_related_count_text"] = {
            "status": "PASS" if _compact(expected) in compact_text else "FAIL",
            "expected": expected,
        }

    raw_brand_warnings = context.get("targeted_drug_brand_warnings")
    if raw_brand_warnings is not None:
        brand_warnings = (
            [str(item) for item in raw_brand_warnings if str(item)]
            if isinstance(raw_brand_warnings, (list, tuple, set))
            else [str(raw_brand_warnings)]
        )
        checks["targeted_drug_brand_mapping"] = {
            "status": "WARN" if brand_warnings else "PASS",
            "warning_codes": brand_warnings,
        }

    raw_drug_consistency = context.get("drug_analysis_consistency")
    if isinstance(raw_drug_consistency, Mapping):
        checks["drug_analysis_consistency"] = {
            "status": (
                "PASS" if raw_drug_consistency.get("status") == "PASS" else "FAIL"
            ),
            "expected_item_count": raw_drug_consistency.get("expected_item_count", 0),
            "rendered_item_count": raw_drug_consistency.get("rendered_item_count", 0),
            "missing": list(raw_drug_consistency.get("missing") or []),
            "unexpected": list(raw_drug_consistency.get("unexpected") or []),
            "duplicates": list(raw_drug_consistency.get("duplicates") or []),
        }

    raw_drug_contract_coverage = context.get("drug_analysis_contract_coverage")
    if isinstance(raw_drug_contract_coverage, Mapping):
        coverage_status = str(raw_drug_contract_coverage.get("status") or "")
        checks["drug_analysis_contract_coverage"] = {
            "status": "PASS" if coverage_status == "PASS" else "WARN",
            "expected_variant_count": raw_drug_contract_coverage.get(
                "expected_variant_count", 0
            ),
            "governed_variant_count": raw_drug_contract_coverage.get(
                "governed_variant_count", 0
            ),
            "coverage_percent": raw_drug_contract_coverage.get(
                "coverage_percent", 0.0
            ),
            "legacy_uncontracted": list(
                raw_drug_contract_coverage.get("legacy_uncontracted") or []
            ),
        }

    raw_domain_coverage = context.get("gene_domain_coverage")
    if isinstance(raw_domain_coverage, Mapping):
        checks["gene_domain_coverage"] = {
            "status": (
                "PASS" if raw_domain_coverage.get("status") == "PASS" else "FAIL"
            ),
            "expected_gene_count": raw_domain_coverage.get(
                "expected_gene_count", 0
            ),
            "covered_gene_count": raw_domain_coverage.get(
                "covered_gene_count", 0
            ),
            "coverage_percent": raw_domain_coverage.get(
                "coverage_percent", 0.0
            ),
            "missing_genes": list(raw_domain_coverage.get("missing_genes") or []),
            "missing_variant_keys": list(
                raw_domain_coverage.get("missing_variant_keys") or []
            ),
            "duplicate_fixed_domain_genes": list(
                raw_domain_coverage.get("duplicate_fixed_domain_genes") or []
            ),
            "duplicate_fixed_domain_variant_keys": list(
                raw_domain_coverage.get("duplicate_fixed_domain_variant_keys")
                or []
            ),
        }

    tmb_status = str(context.get("tmb_status") or "").strip()
    if tmb_status:
        checks["tmb_status_text"] = {
            "status": "PASS" if _compact(tmb_status) in compact_text else "WARN",
            "expected": tmb_status,
        }

    msi_status = str(context.get("msi_status") or "").strip()
    if msi_status:
        checks["msi_status_text"] = {
            "status": "PASS" if _compact(msi_status) in compact_text else "WARN",
            "expected": msi_status,
        }

    return checks


def _business_issues(checks: Mapping[str, Any]) -> Iterable[Dict[str, str]]:
    messages = {
        "total_variant_count_text": "Total variant count text does not match report context.",
        "drug_related_count_text": "Drug-related variant count text does not match report context.",
        "targeted_or_immune_related_count_text": (
            "Targeted-or-immune variant union count text does not match report context."
        ),
        "targeted_drug_brand_mapping": (
            "Targeted drug brand mapping requires explicit configuration review."
        ),
        "drug_analysis_consistency": (
            "Part-2 drug table and Part-3 drug analysis are inconsistent."
        ),
        "drug_analysis_contract_coverage": (
            "Part-3 drug analysis contains legacy rows pending reviewed-contract migration."
        ),
        "gene_domain_coverage": (
            "Part-3 gene knowledge has missing or duplicate fixed protein/domain content."
        ),
        "tmb_status_text": "TMB status from context was not found in rendered text.",
        "msi_status_text": "MSI status from context was not found in rendered text.",
        "part3_cross_cancer_residuals": (
            "Part 3 contains cross-cancer historical wording pending review."
        ),
        "part3_cross_cancer_suppression": (
            "Part-3 historical fields were hidden pending lung-specific review."
        ),
    }
    for key, check in checks.items():
        status = check.get("status")
        message = str(check.get("message") or messages.get(key) or key)
        if status == "FAIL":
            yield {"level": "error", "code": key.upper(), "message": message}
        elif status == "WARN":
            yield {"level": "warning", "code": key.upper(), "message": message}


def _as_int(value: Any) -> Optional[int]:
    if value is None or value == "":
        return None
    try:
        return int(value)
    except Exception:
        return None
