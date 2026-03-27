"""
Generate a customer-facing DOCX that explains:
  - what we extract from Excel (fields/tables),
  - what knowledge bases (DB) are used,
  - the extraction/mapping logic at a high level,
  - where the data lands in the Word template (variables / loop tables).

This is intentionally value-free by default (no patient/sample values) to avoid
leaking identifiers in client-facing docs.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Mapping, Optional, Sequence, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from reportgen.config.loader import ConfigLoader
from reportgen.core.template_contract import extract_template_contract

JINJA_VAR_RE = re.compile(r"\{\{\s*(?P<expr>.*?)\s*\}\}")
JINJA_FOR_RE = re.compile(
    r"\{%\s*for\s+(?P<var>[a-zA-Z_][a-zA-Z0-9_]*)\s+in\s+"
    r"(?P<list>[a-zA-Z_][a-zA-Z0-9_]*)\s*%}"
)
LEADING_PATH_RE = re.compile(
    r"^\s*(?P<path>[a-zA-Z_][a-zA-Z0-9_]*(?:\.[a-zA-Z_][a-zA-Z0-9_]*)*)"
)


@dataclass(frozen=True)
class Location:
    kind: str  # paragraph|table_cell
    area: str  # body|header|footer
    paragraph_index: Optional[int] = None  # 1-based
    table_index: Optional[int] = None  # 1-based within that area
    row: Optional[int] = None  # 1-based
    col: Optional[int] = None  # 1-based
    snippet: str = ""
    table_signature: str = ""


def _clip(text: str, *, limit: int = 42) -> str:
    s = (text or "").strip().replace("\n", " ")
    s = re.sub(r"\s+", " ", s)
    if len(s) <= limit:
        return s
    return s[: limit - 1] + "…"


def _table_signature(tbl, *, max_rows: int = 2) -> str:
    parts: List[str] = []
    for r in range(min(max_rows, len(tbl.rows))):
        row = tbl.rows[r]
        row_text = " | ".join(_clip(cell.text, limit=18) for cell in row.cells)
        parts.append(row_text)
    return " / ".join([p for p in parts if p]).strip()


def _iter_area_blocks(doc: Document, *, area: str) -> Iterable[Tuple[str, int, Any]]:
    """
    Yield (kind, index, obj) where:
      - kind: paragraph|table
      - index: 1-based index within that kind for the area
    """
    if area == "body":
        paragraphs = doc.paragraphs
        tables = doc.tables
    elif area == "header":
        # Use first section header/footer as representative; templates usually
        # keep them consistent across sections.
        hdr = doc.sections[0].header if doc.sections else None
        paragraphs = hdr.paragraphs if hdr else []
        tables = hdr.tables if hdr else []
    elif area == "footer":
        ftr = doc.sections[0].footer if doc.sections else None
        paragraphs = ftr.paragraphs if ftr else []
        tables = ftr.tables if ftr else []
    else:
        raise ValueError(f"Unknown area: {area}")

    for i, p in enumerate(paragraphs, start=1):
        yield ("paragraph", i, p)
    for i, t in enumerate(tables, start=1):
        yield ("table", i, t)


def scan_template_locations(
    template_path: Path,
) -> Tuple[Dict[str, List[Location]], Dict[str, List[Location]]]:
    """
    Return:
      - var_locations: top-level var_name -> locations
      - list_locations: list_name -> locations (where loop starts appear)
    """
    doc = Document(str(template_path))

    # Build loop-var context per table to classify row fields as list fields,
    # not top-level vars.
    # Map: (area, table_index) -> {loop_var: list_name}
    loop_ctx: Dict[Tuple[str, int], Dict[str, str]] = {}

    for area in ("body", "header", "footer"):
        for kind, idx, obj in _iter_area_blocks(doc, area=area):
            if kind != "table":
                continue
            tbl = obj
            ctx: Dict[str, str] = {}
            for row in tbl.rows:
                for cell in row.cells:
                    for m in JINJA_FOR_RE.finditer(cell.text or ""):
                        ctx[m.group("var")] = m.group("list")
            if ctx:
                loop_ctx[(area, idx)] = ctx

    var_locations: Dict[str, List[Location]] = {}
    list_locations: Dict[str, List[Location]] = {}

    for area in ("body", "header", "footer"):
        for kind, idx, obj in _iter_area_blocks(doc, area=area):
            if kind == "paragraph":
                p = obj
                txt = p.text or ""
                if "{{" in txt:
                    for m in JINJA_VAR_RE.finditer(txt):
                        expr = (m.group("expr") or "").split("|", 1)[0]
                        m2 = LEADING_PATH_RE.match(expr)
                        if not m2:
                            continue
                        path = m2.group("path")
                        if not path:
                            continue
                        top = path.split(".", 1)[0]
                        var_locations.setdefault(top, []).append(
                            Location(
                                kind="paragraph",
                                area=area,
                                paragraph_index=idx,
                                snippet=_clip(txt),
                            )
                        )

                if "{% for" in txt:
                    for m in JINJA_FOR_RE.finditer(txt):
                        list_name = m.group("list")
                        list_locations.setdefault(list_name, []).append(
                            Location(
                                kind="paragraph",
                                area=area,
                                paragraph_index=idx,
                                snippet=_clip(txt),
                            )
                        )

            elif kind == "table":
                tbl = obj
                sig = _table_signature(tbl)
                ctx = loop_ctx.get((area, idx), {})

                for r_idx, row in enumerate(tbl.rows, start=1):
                    for c_idx, cell in enumerate(row.cells, start=1):
                        txt = cell.text or ""

                        if "{% for" in txt:
                            for m in JINJA_FOR_RE.finditer(txt):
                                list_name = m.group("list")
                                list_locations.setdefault(list_name, []).append(
                                    Location(
                                        kind="table_cell",
                                        area=area,
                                        table_index=idx,
                                        row=r_idx,
                                        col=c_idx,
                                        snippet=_clip(txt),
                                        table_signature=sig,
                                    )
                                )

                        if "{{" in txt:
                            for m in JINJA_VAR_RE.finditer(txt):
                                expr = (m.group("expr") or "").split("|", 1)[0]
                                m2 = LEADING_PATH_RE.match(expr)
                                if not m2:
                                    continue
                                path = m2.group("path")
                                if not path:
                                    continue
                                top = path.split(".", 1)[0]
                                # Skip loop var (row.xxx) as top-level variables.
                                if top in ctx:
                                    continue
                                var_locations.setdefault(top, []).append(
                                    Location(
                                        kind="table_cell",
                                        area=area,
                                        table_index=idx,
                                        row=r_idx,
                                        col=c_idx,
                                        snippet=_clip(txt),
                                        table_signature=sig,
                                    )
                                )

    return var_locations, list_locations


def _loc_to_str(loc: Location) -> str:
    area_cn = {"body": "正文", "header": "页眉", "footer": "页脚"}.get(
        loc.area, loc.area
    )
    if loc.kind == "paragraph":
        return f"{area_cn}段落#{loc.paragraph_index}：{loc.snippet}"
    if loc.kind == "table_cell":
        sig = f"（表头：{loc.table_signature}）" if loc.table_signature else ""
        return (
            f"{area_cn}表格#{loc.table_index}{sig}(r{loc.row},c{loc.col})："
            f"{loc.snippet}"
        )
    return f"{area_cn}:{loc.snippet}"


def _join_locations(
    locs: Sequence[Location], *, max_items: Optional[int] = None
) -> str:
    if not locs:
        return "模板未引用"
    show = list(locs)
    if max_items is not None:
        show = show[: max(0, int(max_items))]
    lines = [_loc_to_str(loc) for loc in show]
    if max_items is not None and len(locs) > max_items:
        lines.append(f"（其余 {len(locs) - max_items} 处略，见“模板落点索引表”）")
    return "\n".join(lines)


def _set_style_font(
    style,
    *,
    font_name: str,
    size_pt: Optional[float] = None,
    bold: Optional[bool] = None,
) -> None:
    try:
        style.font.name = font_name
        if size_pt is not None:
            style.font.size = Pt(float(size_pt))
        if bold is not None:
            style.font.bold = bool(bold)
        # Ensure East Asia font is set (critical for 宋体 rendering in Word).
        rpr = style._element.get_or_add_rPr()
        rfonts = (
            rpr.rFonts
            if hasattr(rpr, "rFonts") and rpr.rFonts is not None
            else rpr.get_or_add_rFonts()
        )
        rfonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        # Styling should not crash generation; Word can still open.
        pass


def _apply_run_font(run, *, font_name: str) -> None:
    try:
        run.font.name = font_name
        if run._element.rPr is None:
            run._element.get_or_add_rPr()
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


def _apply_document_font(doc: Document, *, font_name: str = "宋体") -> None:
    # Set common styles
    for style_name, size, bold in (
        ("Normal", 11, None),
        ("Heading 1", 16, True),
        ("Heading 2", 14, True),
        ("Heading 3", 12, True),
    ):
        try:
            _set_style_font(
                doc.styles[style_name], font_name=font_name, size_pt=size, bold=bold
            )
        except Exception:
            continue

    def iter_runs() -> Iterable[Any]:
        for p in doc.paragraphs:
            for r in p.runs:
                yield r
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            yield r
        for sec in doc.sections:
            for part in (sec.header, sec.footer):
                for p in part.paragraphs:
                    for r in p.runs:
                        yield r
                for tbl in part.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for r in p.runs:
                                    yield r

    for run in iter_runs():
        _apply_run_font(run, font_name=font_name)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        _apply_run_font(r, font_name="宋体")


def _add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for r in p.runs:
        _apply_run_font(r, font_name="宋体")


def _add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    *,
    col_widths_cm: Optional[List[float]] = None,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                _apply_run_font(r, font_name="宋体")

    for row in rows:
        tr = table.add_row()
        for i, val in enumerate(row):
            tr.cells[i].text = "" if val is None else str(val)
            for p in tr.cells[i].paragraphs:
                for r in p.runs:
                    _apply_run_font(r, font_name="宋体")

    if col_widths_cm:
        from docx.shared import Cm as _Cm

        for i, w in enumerate(col_widths_cm):
            if i < len(table.columns):
                table.columns[i].width = _Cm(float(w))


def _describe_excel_extraction_logic() -> List[str]:
    return [
        "单值字段：默认读取 Excel 第一个sheet 的第1行（横向字段）。",
        "Meta信息：如果存在 Meta/BasicInfo/基本信息 sheet，会尝试以"
        "“横向1行 / Key-Value两列 / 两列无表头”三种格式补充单值字段。",
        "TMB：如果存在 TMB sheet，会扫描多个计算块，优先选择更贴近终版报告的 “TCGA fit” 区块。",
        "MSI：如果存在 Msisensor sheet，优先读取结论字段；若缺失则尝试根据百分比阈值推断 MSI-H/MSI-L/MSS。",
        "HLA：如果存在 HLA sheet，使用专用解析器按 HLA-A/B/C 分块提取 Type1/Type2 及 exon QC 信息。",
        "表格数据：按 mapping.yaml 的 table_data 配置读取相应 sheet（支持 skip_rows），并将每行转为键值对。",
    ]


def _kb_rows(settings: Mapping[str, Any]) -> List[List[str]]:
    kb = settings.get("knowledge_bases", {}) if isinstance(settings, dict) else {}
    if not isinstance(kb, dict):
        kb = {}

    def row_for(
        key: str, display: str, purpose: str, affects: str
    ) -> Optional[List[str]]:
        cfg = kb.get(key)
        if not isinstance(cfg, dict):
            return None
        enabled = bool(cfg.get("enabled", False))
        path = str(cfg.get("path") or "")
        trigger = "enabled=true 且文件存在" if path else "enabled=true 且 path 已配置"

        logic = ""
        if key == "targeted_drug_db":
            filters = (
                cfg.get("filters", {}) if isinstance(cfg.get("filters"), dict) else {}
            )
            logic = (
                "支持按癌种/证据等级/位点匹配过滤；可用 overrides 覆盖缺失条目。"
                if filters
                else "用于将变异映射为“潜在获益/耐药或慎重”药物提示。"
            )
        elif key == "immune_gene_list":
            logic = "用于生成免疫治疗相关基因汇总（正/负相关/超进展相关等）。"
        elif key == "gene_knowledge_db":
            logic = "用于生成“基因简介/基因变异解析/药物疗效临床解析”等段落内容。"
        elif key == "gene_transcript_db":
            logic = "用于补全转录本/染色体信息，以及未检出基因列表的补全。"

        return [
            display,
            "是" if enabled else "否",
            path,
            purpose,
            trigger,
            logic,
            affects,
        ]

    rows: List[List[str]] = []
    for r in (
        row_for(
            "targeted_drug_db",
            "靶向药物提示数据库",
            "药物提示生成",
            "targeted_drug_tips 表格/相关段落",
        ),
        row_for(
            "immune_gene_list", "免疫治疗基因清单", "免疫相关总结", "免疫相关段落/汇总"
        ),
        row_for(
            "gene_knowledge_db",
            "基因知识库",
            "基因诊疗知识段落生成",
            "基因解析/用药解析段落",
        ),
        row_for(
            "gene_transcript_db",
            "转录本-染色体对照库",
            "信息补全",
            "variants/undetected_genes 等表格补全",
        ),
    ):
        if r is not None:
            rows.append(r)
    return rows


def generate_customer_summary_docx(
    *,
    template_path: str,
    config_dir: str,
    output_path: str,
) -> Path:
    template_p = Path(template_path).resolve()
    if not template_p.exists():
        raise FileNotFoundError(f"Template not found: {template_p}")
    if template_p.suffix.lower() != ".docx":
        raise ValueError("template_path must be a .docx file")

    out_p = Path(output_path).resolve()
    out_p.parent.mkdir(parents=True, exist_ok=True)

    loader = ConfigLoader(config_dir=config_dir)
    mapping = loader.load_mapping_config()
    settings = loader.load_settings_config()

    single_values = (
        mapping.get("single_values", {}) if isinstance(mapping, dict) else {}
    )
    table_data = mapping.get("table_data", {}) if isinstance(mapping, dict) else {}

    # Template contract + locations
    contract = extract_template_contract(str(template_p))
    var_locs, list_locs = scan_template_locations(template_p)

    # Build doc
    doc = Document()
    _apply_document_font(doc, font_name="宋体")

    title = "Excel到Docx自动化报告生成说明书（客户版）"
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    _apply_run_font(run, font_name="宋体")
    p.alignment = 1  # center

    _add_paragraph(doc, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    _add_paragraph(doc, f"对应模板：{template_p.name}")
    _add_paragraph(
        doc, "说明：本文档不包含任何患者级样例值，仅解释提取项、逻辑与落点。"
    )

    doc.add_paragraph("")

    _add_heading(doc, "1. 端到端流程概述", level=1)
    for i, line in enumerate(
        [
            "Detect（可选）：识别项目类型并选择模板",
            "Read：读取Excel（单值字段 + 表格数据）",
            "Map：按mapping.yaml将Excel字段/列映射为模板变量与循环表",
            "Clean/Enrich：清洗数据并进行项目增强（例如358相关增强逻辑）",
            "Render：使用docxtpl渲染生成报告DOCX",
            "Post-process（可选）：高亮动态区域、签章、水印等",
            "Validate（可选）：检查报告可打开、占位符残留、关键表格存在等",
        ],
        start=1,
    ):
        _add_paragraph(doc, f"{i}) {line}")

    doc.add_paragraph("")

    _add_heading(doc, "2. Excel数据抽取逻辑（概述）", level=1)
    for line in _describe_excel_extraction_logic():
        _add_paragraph(doc, f"- {line}")

    doc.add_paragraph("")

    _add_heading(doc, "3. 单值字段映射清单（Excel → 模板变量）", level=1)
    single_rows: List[List[str]] = []
    for var in sorted(single_values.keys()):
        cfg = single_values.get(var) or {}
        if not isinstance(cfg, dict):
            continue
        synonyms = cfg.get("synonyms") or []
        if not isinstance(synonyms, list):
            synonyms = [synonyms]

        locs = var_locs.get(var, [])
        single_rows.append(
            [
                str(var),
                str(cfg.get("description") or ""),
                "、".join([str(s) for s in synonyms if str(s).strip()]),
                str(cfg.get("type") or ""),
                "是" if bool(cfg.get("required", False)) else "否",
                (
                    ""
                    if cfg.get("default_value") is None
                    else str(cfg.get("default_value"))
                ),
                str(cfg.get("format_template") or ""),
                str(len(locs)),
                _join_locations(locs, max_items=3),
            ]
        )

    _add_table(
        doc,
        headers=[
            "模板变量",
            "字段说明",
            "Excel字段同义词（匹配）",
            "类型",
            "必填",
            "默认值",
            "格式化",
            "模板出现次数",
            "主要落点（前3处）",
        ],
        rows=single_rows,
        col_widths_cm=[3.2, 4.0, 4.3, 1.5, 1.0, 1.8, 1.8, 1.6, 7.0],
    )

    doc.add_page_break()

    _add_heading(doc, "4. 表格数据映射清单（Excel Sheet → 模板循环表）", level=1)
    table_rows: List[List[str]] = []
    for table_name in sorted(table_data.keys()):
        cfg = table_data.get(table_name) or {}
        if not isinstance(cfg, dict):
            continue
        sheet_name = str(cfg.get("sheet_name") or "")
        required = "是" if bool(cfg.get("required", False)) else "否"
        empty_behavior = str(cfg.get("empty_behavior") or "")

        cols = cfg.get("columns") or {}
        col_lines: List[str] = []
        if isinstance(cols, dict):
            for col_var, col_cfg in cols.items():
                if not isinstance(col_cfg, dict):
                    continue
                syn = col_cfg.get("synonyms") or []
                if not isinstance(syn, list):
                    syn = [syn]
                syn_str = "、".join([str(s) for s in syn if str(s).strip()])
                ctype = str(col_cfg.get("type") or "")
                fmt = str(col_cfg.get("format_template") or "")
                extra = f"（{ctype} {fmt}）".strip()
                extra = extra if extra != "（）" else ""
                col_lines.append(f"{col_var}: {syn_str} {extra}".strip())

        locs = list_locs.get(str(table_name), [])
        table_rows.append(
            [
                str(table_name),
                sheet_name,
                required,
                empty_behavior,
                "\n".join(col_lines[:30])
                + (
                    f"\n…（其余{len(col_lines)-30}列略）" if len(col_lines) > 30 else ""
                ),
                str(len(locs)),
                _join_locations(locs, max_items=3),
            ]
        )

    _add_table(
        doc,
        headers=[
            "模板表名(list)",
            "Excel Sheet",
            "必需",
            "空表处理",
            "列映射（变量: 同义词/类型/格式）",
            "模板出现次数",
            "主要落点（前3处）",
        ],
        rows=table_rows,
        col_widths_cm=[3.2, 2.6, 1.0, 1.8, 9.0, 1.6, 7.0],
    )

    doc.add_page_break()

    _add_heading(doc, "5. 使用到的数据库/知识库（settings.yaml）", level=1)
    kb_rows = _kb_rows(settings)
    if kb_rows:
        _add_table(
            doc,
            headers=[
                "名称",
                "启用",
                "文件路径",
                "用途",
                "触发条件",
                "关键逻辑摘要",
                "影响输出",
            ],
            rows=kb_rows,
            col_widths_cm=[3.0, 1.0, 6.2, 2.4, 2.6, 5.8, 3.0],
        )
    else:
        _add_paragraph(doc, "未在 settings.yaml 中检测到 knowledge_bases 配置。")

    doc.add_page_break()

    _add_heading(doc, "6. 模板落点索引表（更详细）", level=1)
    _add_paragraph(
        doc,
        "说明：索引表列出模板中所有变量/循环表的出现位置，便于和报告正文快速对照定位。",
    )

    index_rows: List[List[str]] = []

    for var in contract.required_paths:
        top = str(var).split(".", 1)[0]
        locs = var_locs.get(top, [])
        index_rows.append(
            ["变量", top, str(len(locs)), _join_locations(locs, max_items=None)]
        )

    for list_name in contract.required_lists:
        locs = list_locs.get(str(list_name), [])
        index_rows.append(
            [
                "循环表",
                str(list_name),
                str(len(locs)),
                _join_locations(locs, max_items=None),
            ]
        )

    _add_table(
        doc,
        headers=["类型", "名称", "出现次数", "位置列表（全部）"],
        rows=index_rows,
        col_widths_cm=[1.8, 3.2, 1.6, 14.0],
    )

    _apply_document_font(doc, font_name="宋体")
    doc.save(str(out_p))
    return out_p
