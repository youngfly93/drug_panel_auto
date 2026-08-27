"""金标样式基线：从报告或模板 DOCX 提取可回归的结构化样式指纹。

用途：把 crc_358/crc_301 等 panel 的**当前正确输出**冻结为基线 JSON，后处理层
任何改动只要跑一遍基线测试，就能立刻发现是否把别处的样式弄坏了（链接色/下划线/
加粗/❖颜色/签名/参考文献覆盖）——把"改完怕崩、静默复发"变成"改完有答案"。

指纹只取**确定性、与样式相关**的信号，刻意排除页码/时间戳/图片关系 ID 等易变项。
"""

from __future__ import annotations

from collections import Counter
from hashlib import sha256
import json
from pathlib import Path
import re
from typing import Any, Union
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn

from reportgen.docx_sections import find_reference_section_bounds

# 关键表：表头特征子串（全部命中即认定该表）。取首个匹配。
_TABLE_SIGNATURES: dict[str, tuple[str, ...]] = {
    "variant_detail": ("基因名称", "频率"),
    "targeted_drug_tips": ("潜在获益靶向药物",),
    "biomarker_results": ("用药提示",),
    "gene_list": ("Gene List for MLseq",),
}

_TEMPLATE_STYLE_PARTS = {
    "word/document.xml",
    "word/fontTable.xml",
    "word/numbering.xml",
    "word/settings.xml",
    "word/styles.xml",
    "word/theme/theme1.xml",
}
_TEMPLATE_RELATED_PART = re.compile(
    r"word/(?:header|footer)\d+\.xml$|"
    r"word/_rels/(?:document|header\d+|footer\d+)\.xml\.rels$"
)


def _run_style(run) -> str:
    """单个 run 的紧凑样式签名：color|b|u（空=纯黑无修饰）。"""
    parts: list[str] = []
    color = run.font.color
    rgb = None
    try:
        rgb = color.rgb if color is not None else None
    except Exception:
        rgb = None
    if rgb is not None:
        parts.append(f"c={str(rgb).upper()}")
    if run.font.bold:
        parts.append("b")
    if run.font.underline:
        parts.append("u")
    return "|".join(parts)


def _cell_signature(cell) -> list[str]:
    """单元格签名：[文本, 去重排序后的非空 run 样式...]。"""
    text = (cell.text or "").strip()
    styles = sorted(
        {
            _run_style(run)
            for paragraph in cell.paragraphs
            for run in paragraph.runs
            if (run.text or "").strip()
        }
    )
    return [text, *styles]


def _table_fingerprint(table) -> list[list[list[str]]]:
    return [[_cell_signature(cell) for cell in row.cells] for row in table.rows]


def _sha256_bytes(value: bytes) -> str:
    return sha256(value).hexdigest()


def _length(value: Any) -> int | None:
    return int(value) if value is not None else None


def _style_id(value: Any) -> str:
    style = getattr(value, "style", None)
    return str(getattr(style, "style_id", "") or "<none>")


def _template_package_hashes(docx_path: Union[str, Path]) -> dict[str, str]:
    """Hash style/layout-bearing package parts without emitting their text."""
    with ZipFile(docx_path) as zf:
        names = sorted(
            name
            for name in zf.namelist()
            if name in _TEMPLATE_STYLE_PARTS
            or _TEMPLATE_RELATED_PART.fullmatch(name)
            or name.startswith("word/media/")
        )
        return {name: _sha256_bytes(zf.read(name)) for name in names}


def extract_template_style_baseline(
    docx_path: Union[str, Path],
) -> dict[str, Any]:
    """Return a PHI-safe source-template style/layout fingerprint.

    Unlike :func:`extract_style_baseline`, this function is meant for template
    sources that do not yet have a fully accepted rendered golden case. It
    never emits paragraph/cell text. Content-bearing XML and media are retained
    only as SHA-256 digests, so any source-template drift remains detectable
    without copying case content into versioned JSON baselines.
    """
    path = Path(docx_path)
    doc = Document(str(path))
    paragraph_styles = Counter(_style_id(paragraph) for paragraph in doc.paragraphs)
    table_styles = Counter(_style_id(table) for table in doc.tables)
    table_shapes = Counter(
        f"{len(table.rows)}x{len(table.columns)}" for table in doc.tables
    )
    section_layout = []
    for section in doc.sections:
        section_layout.append(
            {
                "orientation": str(section.orientation),
                "start_type": str(section.start_type),
                "page_width": _length(section.page_width),
                "page_height": _length(section.page_height),
                "top_margin": _length(section.top_margin),
                "right_margin": _length(section.right_margin),
                "bottom_margin": _length(section.bottom_margin),
                "left_margin": _length(section.left_margin),
                "header_distance": _length(section.header_distance),
                "footer_distance": _length(section.footer_distance),
                "header_linked": bool(section.header.is_linked_to_previous),
                "footer_linked": bool(section.footer.is_linked_to_previous),
            }
        )

    with ZipFile(path) as zf:
        xml_parts = [
            zf.read(name)
            for name in zf.namelist()
            if name == "word/document.xml"
            or re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
        ]
    xml = b"\n".join(xml_parts)
    package_hashes = _template_package_hashes(path)
    canonical = json.dumps(
        package_hashes,
        ensure_ascii=True,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("ascii")

    return {
        "schema_version": 1,
        "metrics": {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "section_count": len(doc.sections),
            "inline_shape_count": len(doc.inline_shapes),
            "drawing_count": len(re.findall(rb"<(?:w:drawing|wp:inline|wp:anchor)\b", xml)),
            "jinja_variable_count": xml.count(b"{{"),
            "jinja_control_count": xml.count(b"{%"),
            "media_part_count": sum(
                name.startswith("word/media/") for name in package_hashes
            ),
        },
        "paragraph_styles": dict(sorted(paragraph_styles.items())),
        "table_styles": dict(sorted(table_styles.items())),
        "table_shapes": dict(sorted(table_shapes.items())),
        "sections": section_layout,
        "package_parts_sha256": package_hashes,
        "aggregate_sha256": _sha256_bytes(canonical),
        "privacy": {
            "paragraph_text_emitted": False,
            "cell_text_emitted": False,
            "media_bytes_emitted": False,
        },
    }


def _part3_bullet_color(doc) -> str | None:
    """第三部分引导段 ❖ 项目符号颜色（None=无/继承）。"""
    num_id = None
    for paragraph in doc.paragraphs:
        text = paragraph.text or ""
        if "在本次检测范围内" in text and "检出体细胞变异" in text:
            p_pr = paragraph._p.find(qn("w:pPr"))
            num_pr = p_pr.find(qn("w:numPr")) if p_pr is not None else None
            num_el = num_pr.find(qn("w:numId")) if num_pr is not None else None
            num_id = num_el.get(qn("w:val")) if num_el is not None else None
            break
    if num_id is None:
        return None
    try:
        numbering = doc.part.numbering_part.element
    except Exception:
        return None
    abstract_id = None
    for num in numbering.findall(qn("w:num")):
        if num.get(qn("w:numId")) == num_id:
            ref = num.find(qn("w:abstractNumId"))
            abstract_id = ref.get(qn("w:val")) if ref is not None else None
            break
    for abstract in numbering.findall(qn("w:abstractNum")):
        if abstract.get(qn("w:abstractNumId")) != abstract_id:
            continue
        for level in abstract.findall(qn("w:lvl")):
            if level.get(qn("w:ilvl")) != "0":
                continue
            rpr = level.find(qn("w:rPr"))
            color = rpr.find(qn("w:color")) if rpr is not None else None
            return (color.get(qn("w:val")) or "").upper() if color is not None else None
    return None


def _signature_fingerprint(doc) -> dict[str, Any]:
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text or ""
        if "检测者" in text and "审核者" in text:
            inline = paragraph._p.findall(".//" + qn("w:drawing") + "/" + qn("wp:inline"))
            anchor = paragraph._p.findall(".//" + qn("wp:anchor"))
            prev_anchor = 0
            if idx > 0:
                prev_anchor = len(doc.paragraphs[idx - 1]._p.findall(".//" + qn("wp:anchor")))
            return {
                "label": text.strip(),
                "inline_images": len(inline),
                "floating_anchors_here_and_prev": len(anchor) + prev_anchor,
            }
    return {}


def _references_fingerprint(doc) -> dict[str, Any]:
    texts = [(p.text or "").strip() for p in doc.paragraphs]
    start, end = find_reference_section_bounds(texts)
    if start is None:
        return {"count": 0, "titleless": 0}
    refs = [t for t in texts[start + 1 : end] if t]
    titleless = [t for t in refs if re.fullmatch(r"PMID:\d+", t)]
    return {"count": len(refs), "titleless": len(titleless)}


def _report_integrity_fingerprint(docx_path: str) -> dict[str, Any]:
    """报告完整性/自洽指纹（确定性，与页码/时间无关）。

    冻结这些信号，可在后处理改动后立即抓出"计数对不上、占位符没渲染、
    结构性段落丢失"这类静默退化——补齐样式基线只盖"样式"、不盖"数值/结构
    自洽"的盲区。刻意只取确定性强的量：
    - variant_count_text / drug_count_text：正文里写明的计数（不是重新计算，
      而是冻结"正文当时说了几个"，改动若让计数措辞或数值漂移即报警）。
    - unrendered_placeholders：``{{ }}`` / ``{%% %%}`` 残留数，正常应为 0。
    - part_section_markers：``第X部分`` 出现次数（结构性骨架，丢段会变）。
    """
    with ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", "ignore")
    text = re.sub(r"<[^>]+>", "", xml)

    def _first_int(pattern: str) -> int | None:
        m = re.search(pattern, text)
        return int(m.group(1)) if m else None

    return {
        "variant_count_text": _first_int(r"检出体细胞变异[：:]\s*(\d+)\s*个"),
        "drug_count_text": _first_int(r"用药相关的变异有[：:]\s*(\d+)\s*个"),
        "unrendered_placeholders": len(re.findall(r"\{\{|\{%", xml)),
        "part_section_markers": len(re.findall(r"第[一二三四五六]部分", text)),
    }


def extract_style_baseline(docx_path: str) -> dict[str, Any]:
    """提取关键表 + 标志位的确定性样式指纹。"""
    doc = Document(docx_path)
    tables: dict[str, Any] = {}
    for table in doc.tables:
        if not table.rows:
            continue
        header = " ".join(cell.text for cell in table.rows[0].cells)
        for name, needles in _TABLE_SIGNATURES.items():
            if name in tables:
                continue
            if all(n in header for n in needles):
                tables[name] = _table_fingerprint(table)
                break
    return {
        "tables": tables,
        "markers": {
            "part3_bullet_color": _part3_bullet_color(doc),
            "signature": _signature_fingerprint(doc),
            "references": _references_fingerprint(doc),
            "integrity": _report_integrity_fingerprint(docx_path),
        },
    }


def diff_baseline(expected: dict[str, Any], actual: dict[str, Any]) -> list[str]:
    """返回人类可读的差异列表（空=一致）。"""
    diffs: list[str] = []

    def walk(path: str, exp: Any, act: Any) -> None:
        if isinstance(exp, dict) and isinstance(act, dict):
            for key in sorted(set(exp) | set(act)):
                walk(f"{path}.{key}", exp.get(key, "<缺>"), act.get(key, "<缺>"))
        elif isinstance(exp, list) and isinstance(act, list):
            if len(exp) != len(act):
                diffs.append(f"{path}: 长度 {len(exp)} → {len(act)}")
            for i in range(min(len(exp), len(act))):
                walk(f"{path}[{i}]", exp[i], act[i])
        elif exp != act:
            diffs.append(f"{path}: {exp!r} → {act!r}")

    walk("", expected, actual)
    return diffs
