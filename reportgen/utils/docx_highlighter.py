"""
DOCX highlighter utilities.

Purpose:
  Highlight dynamic regions in a *rendered* report docx, based on the original
  Jinja2/docxtpl template (.docx).

This helps template maintenance by visually showing which areas are filled from
Excel/mapped variables ({{ ... }}) or loop-generated table rows ({% for ... %}).

Python 3.9+ compatible.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

JINJA_ANY_RE = re.compile(r"(\{\{.*?\}\}|\{%.+?%\})", re.DOTALL)
JINJA_FOR_RE = re.compile(r"\{%\s*for\b", re.IGNORECASE)


def _normalize_color(color: str) -> str:
    c = (color or "").strip()
    if c.startswith("#"):
        c = c[1:]
    c = c.upper()
    if not re.fullmatch(r"[0-9A-F]{6}", c):
        raise ValueError(
            f"Invalid color hex: {color!r}. Expected like 'D9EAF7' or '#D9EAF7'."
        )
    return c


def _clean_text(text: str) -> str:
    s = (text or "").replace("\u200b", "").replace("\xa0", " ")
    s = re.sub(r"\s+", " ", s)
    s = JINJA_ANY_RE.sub("", s)
    return s.strip()


def _add_run_shading(run, *, fill: str, skip_empty: bool) -> bool:
    """Apply run background shading. Returns True if changed."""
    if skip_empty and (run.text or "").strip() == "":
        return False

    fill = _normalize_color(fill)
    rpr = run._r.get_or_add_rPr()

    # If shading exists, reuse it.
    shd = None
    for child in rpr:
        if child.tag == qn("w:shd"):
            shd = child
            break

    if shd is None:
        shd = OxmlElement("w:shd")
        rpr.append(shd)

    # Use 'clear' shading to create a background fill.
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    prev = shd.get(qn("w:fill"))
    shd.set(qn("w:fill"), fill)
    return prev != fill


def _highlight_paragraph(paragraph, *, fill: str, skip_empty: bool) -> int:
    changed = 0
    for run in paragraph.runs:
        if _add_run_shading(run, fill=fill, skip_empty=skip_empty):
            changed += 1
    return changed


def _highlight_cell(cell, *, fill: str, skip_empty: bool) -> int:
    changed = 0
    for p in cell.paragraphs:
        changed += _highlight_paragraph(p, fill=fill, skip_empty=skip_empty)
    return changed


def _table_header_signature(table, *, header_rows: int) -> str:
    """Compute a signature string from header rows for table matching."""
    rows = min(max(header_rows, 1), len(table.rows))
    parts: List[str] = []
    for r in range(rows):
        row = table.rows[r]
        row_text = " | ".join(_clean_text(cell.text) for cell in row.cells)
        parts.append(row_text)
    return "\n".join(parts).strip()


def _table_signature(table) -> Tuple[int, str]:
    """Fallback signature for non-loop tables."""
    cols = len(table.columns)
    parts: List[str] = []
    for r in range(min(2, len(table.rows))):
        row = table.rows[r]
        parts.append(" | ".join(_clean_text(cell.text) for cell in row.cells))
    return cols, "\n".join(parts).strip()


def _similarity(a: str, b: str) -> float:
    if not a and not b:
        return 1.0
    return SequenceMatcher(None, a, b).ratio()


@dataclass(frozen=True)
class TemplateTableSpec:
    index: int
    cols: int
    is_loop: bool
    header_rows: int
    header_var_cells: Tuple[Tuple[int, int], ...]
    var_cells: Tuple[Tuple[int, int], ...]
    signature: str


@dataclass(frozen=True)
class HighlightSummary:
    template: str
    input_docx: str
    output_docx: str
    color: str
    highlighted_runs: int
    matched_tables: int
    tables_processed: int
    paragraphs_highlighted: int


def _extract_template_table_specs(template: Document) -> List[TemplateTableSpec]:
    specs: List[TemplateTableSpec] = []
    for ti, tbl in enumerate(template.tables):
        cols = len(tbl.columns)

        loop_start: Optional[int] = None
        for r_idx, row in enumerate(tbl.rows):
            if any(JINJA_FOR_RE.search(cell.text or "") for cell in row.cells):
                loop_start = r_idx
                break

        is_loop = loop_start is not None
        header_rows = int(loop_start) if is_loop else 0

        header_var_cells: List[Tuple[int, int]] = []
        var_cells: List[Tuple[int, int]] = []

        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                txt = cell.text or ""
                if "{{" in txt:
                    if is_loop and r_idx < header_rows:
                        header_var_cells.append((r_idx, c_idx))
                    elif not is_loop:
                        var_cells.append((r_idx, c_idx))

        signature = (
            _table_header_signature(tbl, header_rows=header_rows)
            if is_loop
            else _table_signature(tbl)[1]
        )

        specs.append(
            TemplateTableSpec(
                index=ti,
                cols=cols,
                is_loop=is_loop,
                header_rows=header_rows,
                header_var_cells=tuple(header_var_cells),
                var_cells=tuple(var_cells),
                signature=signature,
            )
        )
    return specs


def _match_table_indices(
    template: Document, rendered: Document, specs: List[TemplateTableSpec]
) -> Dict[int, int]:
    """Return mapping: template_table_index -> rendered_table_index."""
    if len(template.tables) == len(rendered.tables):
        return {spec.index: spec.index for spec in specs}

    # Fallback matching by signature.
    rendered_sigs: List[Tuple[int, int, str]] = []
    for ri, tbl in enumerate(rendered.tables):
        cols = len(tbl.columns)
        sig = _table_signature(tbl)[1]
        rendered_sigs.append((ri, cols, sig))

    mapping: Dict[int, int] = {}
    used: set[int] = set()

    for spec in specs:
        best: Tuple[float, Optional[int]] = (0.0, None)
        for ri, cols, sig in rendered_sigs:
            if ri in used:
                continue
            if cols != spec.cols:
                continue
            score = _similarity(spec.signature, sig)
            if score > best[0]:
                best = (score, ri)
        if best[1] is not None:
            mapping[spec.index] = int(best[1])
            used.add(int(best[1]))

    return mapping


def highlight_rendered_docx(
    *,
    template_path: str,
    input_docx_path: str,
    output_docx_path: str,
    color: str = "D9EAF7",
    skip_empty: bool = True,
) -> HighlightSummary:
    """Highlight dynamic regions in a rendered docx based on a template docx."""
    template_p = Path(template_path)
    input_p = Path(input_docx_path)
    output_p = Path(output_docx_path)

    if not template_p.exists():
        raise FileNotFoundError(f"Template docx not found: {template_p}")
    if not input_p.exists():
        raise FileNotFoundError(f"Input docx not found: {input_p}")
    if template_p.suffix.lower() != ".docx" or input_p.suffix.lower() != ".docx":
        raise ValueError("Both template and input must be .docx files.")

    fill = _normalize_color(color)

    template = Document(str(template_p))
    rendered = Document(str(input_p))

    specs = _extract_template_table_specs(template)
    table_map = _match_table_indices(template, rendered, specs)

    highlighted_runs = 0
    matched_tables = 0
    tables_processed = 0

    # Highlight body paragraphs containing Jinja markers.
    paragraphs_highlighted = 0
    for i, p in enumerate(template.paragraphs):
        # Only highlight filled values ({{ ... }}), not control statements ({% ... %}),
        # to reduce paragraph index drift when docxtpl removes standalone
        # control blocks.
        if "{{" not in (p.text or ""):
            continue
        if i >= len(rendered.paragraphs):
            continue
        highlighted_runs += _highlight_paragraph(
            rendered.paragraphs[i], fill=fill, skip_empty=skip_empty
        )
        paragraphs_highlighted += 1

    # Highlight headers/footers (by section/paragraph index).
    for si, sec in enumerate(template.sections):
        if si >= len(rendered.sections):
            break
        out_sec = rendered.sections[si]

        for pi, p in enumerate(sec.header.paragraphs):
            if "{{" in (p.text or ""):
                if pi < len(out_sec.header.paragraphs):
                    highlighted_runs += _highlight_paragraph(
                        out_sec.header.paragraphs[pi], fill=fill, skip_empty=skip_empty
                    )
                    paragraphs_highlighted += 1

        for pi, p in enumerate(sec.footer.paragraphs):
            if "{{" in (p.text or ""):
                if pi < len(out_sec.footer.paragraphs):
                    highlighted_runs += _highlight_paragraph(
                        out_sec.footer.paragraphs[pi], fill=fill, skip_empty=skip_empty
                    )
                    paragraphs_highlighted += 1

    # Highlight tables based on template specs.
    for spec in specs:
        tables_processed += 1
        out_idx = table_map.get(spec.index)
        if out_idx is None or out_idx >= len(rendered.tables):
            continue
        matched_tables += 1
        out_tbl = rendered.tables[out_idx]

        if spec.is_loop:
            # Highlight loop-generated data rows.
            # 使用 row.cells 而非 table.columns 来遍历，避免合并单元格导致越界
            for r in range(spec.header_rows, len(out_tbl.rows)):
                for cell in out_tbl.rows[r].cells:
                    highlighted_runs += _highlight_cell(
                        cell, fill=fill, skip_empty=skip_empty
                    )

            # Also highlight any variable cells in header rows.
            for r, c in spec.header_var_cells:
                if r < len(out_tbl.rows) and c < len(out_tbl.columns):
                    highlighted_runs += _highlight_cell(
                        out_tbl.cell(r, c), fill=fill, skip_empty=skip_empty
                    )
        else:
            # Highlight cells that contained variables in template.
            for r, c in spec.var_cells:
                if r < len(out_tbl.rows) and c < len(out_tbl.columns):
                    highlighted_runs += _highlight_cell(
                        out_tbl.cell(r, c), fill=fill, skip_empty=skip_empty
                    )

    output_p.parent.mkdir(parents=True, exist_ok=True)
    rendered.save(str(output_p))

    return HighlightSummary(
        template=str(template_p),
        input_docx=str(input_p),
        output_docx=str(output_p),
        color=fill,
        highlighted_runs=highlighted_runs,
        matched_tables=matched_tables,
        tables_processed=tables_processed,
        paragraphs_highlighted=paragraphs_highlighted,
    )
