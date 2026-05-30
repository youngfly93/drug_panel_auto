import importlib.util
import sys
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.shared import RGBColor
from docxtpl import DocxTemplate

ROOT = Path(__file__).resolve().parents[2]
BACKEND = ROOT / "backend"
for import_path in (str(ROOT), str(BACKEND)):
    if import_path not in sys.path:
        sys.path.insert(0, import_path)

from app.services.reportgen_bridge import ReportGenBridge  # noqa: E402
from reportgen.core.report_generator import ReportGenerator  # noqa: E402
from reportgen.panels.loader import load_panel_package  # noqa: E402
from reportgen.panels.validation import validate_panel_package  # noqa: E402


def _load_seed_builder():
    path = ROOT / "scripts" / "build_golden_template_seed.py"
    spec = importlib.util.spec_from_file_location("build_golden_template_seed", path)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)
    return module


def _load_variableizer():
    path = ROOT / "scripts" / "variableize_golden_template.py"
    spec = importlib.util.spec_from_file_location("variableize_golden_template", path)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)
    return module


def _load_golden_diff():
    path = ROOT / "scripts" / "diff_golden_report.py"
    spec = importlib.util.spec_from_file_location("diff_golden_report", path)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def test_crc358_golden_template_declares_template_level_processors():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    template = package.templates["crc_358_msi_golden_template_v0"]

    assert template.status == "active"
    assert template.processors == (
        "part3_formatted_sections",
        "signature_placeholder",
        "immune_table_notes",
        "bullet_lists",
        "signature_layout",
        "front_matter_spacing",
        "blank_page_cleanup",
        "toc_refresh",
        "final_refresh_cleanup",
        "underlines_and_styles",
    )
    assert package.resolve_template_file(template.template_id).exists()

    report = validate_panel_package("crc_358_msi", project_root=ROOT)
    assert report.ok, report.to_dict()


def test_crc358_golden_template_part3_uses_dynamic_marker():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    template_path = package.resolve_template_file("crc_358_msi_golden_template_v0")
    doc = Document(template_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    text = "\n".join(paragraphs)

    assert "__PART3_MARKER__" in text
    marker_idx = paragraphs.index("__PART3_MARKER__")
    reading_idx = paragraphs.index("3. 阅读说明")
    assert marker_idx < reading_idx
    dynamic_stub = "\n".join(paragraphs[marker_idx:reading_idx])
    assert "p.G12S" not in dynamic_stub
    assert "c.34G>A" not in dynamic_stub
    assert "46.29" not in dynamic_stub


def test_crc301_golden_template_declares_active_default_processors():
    package = load_panel_package("crc_301_msi", project_root=ROOT)
    template = package.templates["crc_301_msi_golden_template_v1"]

    assert template.status == "active"
    assert package.templates["crc_301_msi_golden_template_v0"].status == "deprecated"
    assert package.default_template.template_id == "crc_301_msi_golden_template_v1"
    assert template.processors == (
        "part3_formatted_sections",
        "signature_placeholder",
        "bullet_lists",
        "signature_layout",
        "front_matter_spacing",
        "blank_page_cleanup",
        "toc_refresh",
        "final_refresh_cleanup",
        "underlines_and_styles",
    )
    assert package.resolve_template_file(template.template_id).exists()

    report = validate_panel_package("crc_301_msi", project_root=ROOT)
    assert report.ok, report.to_dict()


def test_crc301_golden_template_part3_uses_dynamic_marker():
    package = load_panel_package("crc_301_msi", project_root=ROOT)
    template_path = package.resolve_template_file("crc_301_msi_golden_template_v1")
    doc = Document(template_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    text = "\n".join(paragraphs)

    assert "__PART3_MARKER__" in text
    marker_idx = paragraphs.index("__PART3_MARKER__")
    reading_idx = paragraphs.index("3. 阅读说明")
    assert marker_idx < reading_idx
    assert "__PATIENT_NAME__" not in text
    assert "__SAMPLE_ID__" not in text
    assert "__REPORT_NUMBER__" not in text
    assert "检测者：" in text
    assert "审核者：" in text
    assert "报告日期：{{ report_date_dot }}" in text


def test_crc301_golden_template_docxtpl_renders_minimal_context(tmp_path):
    package = load_panel_package("crc_301_msi", project_root=ROOT)
    template_path = package.resolve_template_file("crc_301_msi_golden_template_v1")
    output = tmp_path / "crc301_rendered.docx"

    context = {
        "patient_name": "测试患者",
        "sample_id": "TEST301",
        "report_number": "MLJY-TEST301",
        "gender": "男",
        "age": "58",
        "sample_type": "组织",
        "sampling_method": "手术",
        "sample_site": "结肠",
        "clinical_diagnosis": "结直肠癌",
        "report_date_dot": "2026.01.15",
        "total_variants_count": 1,
        "drug_related_count": 1,
        "tmb_summary": "6.0 mutations/Mb，TMB-L",
        "immuno_tips": "未提示",
        "msi_summary": "微卫星稳定型，MSS",
        "msi_tips": "未提示",
        "immune_positive_result": "未检出",
        "immune_negative_result": "未检出",
        "immune_hyperprogression_result": "未检出",
        "tmb_detail_sentence": "本次检测TMB水平较低。",
        "tmb_detail_interpretation": "测试TMB解读。",
        "tmb_drug_note": "",
        "msi_detail_sentence": "本次检测为MSS。",
        "targeted_drug_tips": [
            {
                "gene": "KRAS",
                "variant_site": "p.G12D",
                "benefit_drugs": "--",
                "caution_drugs": "西妥昔单抗",
            }
        ],
        "variants_2_1": [
            {
                "gene": "KRAS",
                "transcript": "NM_004985",
                "chr": "12",
                "exon": "2",
                "locus": "25398284",
                "var_type_cn": "错义突变",
                "af_pct": "12.3%",
                "benefit_drugs": "--",
                "caution_drugs": "西妥昔单抗",
            }
        ],
        "chemotherapy": [
            {"Drug": "瑞戈非尼", "Gene": "VEGFR", "药物适应情况": "测试适应症"}
        ],
        "nccn_results": [{"gene": "KRAS", "content": "EX2", "result": "检出"}],
        "immune_positive_results": [
            {"gene": "TP53", "result": "未检出", "interpretation": "--"}
        ],
        "immune_negative_results": [
            {"gene": "PTEN", "result": "未检出", "interpretation": "--"}
        ],
        "immune_hyperprogression_results": [
            {"gene": "EGFR", "result": "未检出", "interpretation": "--"}
        ],
    }

    template = DocxTemplate(template_path)
    template.render(context)
    template.save(output)

    text = "\n".join(p.text for p in Document(output).paragraphs)
    assert "{{" not in text
    assert "{%" not in text
    assert "TEST301" in text
    assert "检测者：" in text
    assert "审核者：" in text


def test_web_bridge_resolves_panel_template_id():
    bridge = ReportGenBridge(
        config_dir=str(ROOT / "config"),
        template_dir=str(ROOT / "templates"),
    )

    resolved = Path(
        bridge._resolve_template_path(
            "crc_358_msi_golden_template_v0",
            "crc_358_msi",
        )
    )

    assert resolved.name == "crc_358_msi_golden_template_v0.docx"
    assert resolved.exists()
    assert "panels/crc_358_msi/templates" in resolved.as_posix()

    crc301_default = Path(bridge._resolve_template_path(None, "crc_301_msi"))
    assert crc301_default.name == "crc_301_msi_golden_template_v1.docx"
    assert crc301_default.exists()
    assert "panels/crc_301_msi/templates" in crc301_default.as_posix()


def test_report_generator_uses_template_level_processors():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    template_path = package.resolve_template_file("crc_358_msi_golden_template_v0")

    processors = ReportGenerator._get_template_processor_names(
        package,
        str(template_path),
    )

    assert processors == (
        "part3_formatted_sections",
        "signature_placeholder",
        "immune_table_notes",
        "bullet_lists",
        "signature_layout",
        "front_matter_spacing",
        "blank_page_cleanup",
        "toc_refresh",
        "final_refresh_cleanup",
        "underlines_and_styles",
    )


def test_crc358_golden_template_has_no_report_guide_spacer_cluster():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    template_path = package.resolve_template_file("crc_358_msi_golden_template_v0")
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    w_p = f"{{{ns_w}}}p"
    w_t = f"{{{ns_w}}}t"
    w_br = f"{{{ns_w}}}br"
    w_type = f"{{{ns_w}}}type"
    w_drawing = f"{{{ns_w}}}drawing"

    with ZipFile(template_path) as zin:
        root = ET.fromstring(zin.read("word/document.xml"))
    body_items = list(root.iter())
    paragraphs = [elem for elem in body_items if elem.tag == w_p]

    def text(elem):
        return "".join((node.text or "") for node in elem.iter(w_t)).strip()

    def is_blank(elem):
        return elem.tag == w_p and not text(elem) and not any(elem.iter(w_drawing))

    def has_page_break(elem):
        return any(node.attrib.get(w_type) == "page" for node in elem.iter(w_br))

    guide_idx = next(
        idx for idx, paragraph in enumerate(paragraphs) if text(paragraph) == "报告导读"
    )
    cluster = []
    idx = guide_idx - 1
    while idx >= 0 and is_blank(paragraphs[idx]):
        cluster.append(paragraphs[idx])
        idx -= 1

    assert len(cluster) == 1
    assert has_page_break(cluster[0])


def test_golden_seed_builder_scrubs_patient_tokens(tmp_path):
    builder = _load_seed_builder()
    source = tmp_path / "source.docx"
    output = tmp_path / "seed.docx"
    doc = Document()
    doc.add_paragraph("33333333333333333333333333")
    doc.add_paragraph("姓名：黄金测试患者")
    report_no = doc.add_paragraph("报告编号：")
    report_no.add_run("TEST")
    report_no.add_run("-")
    report_no.add_run("REPORT001")
    sample_no = doc.add_paragraph("样本号：")
    sample_no.add_run("SAMPLE")
    sample_no.add_run("001")
    doc.add_paragraph("临床诊断：测试诊断")
    report_date = doc.add_paragraph("报告日期：")
    report_date.add_run("2099")
    report_date.add_run(".12.")
    report_date.add_run("31")
    doc.save(source)
    replacements = {
        "黄金测试患者": "{{ patient_name }}",
        "TEST-REPORT001": "{{ report_number }}",
        "SAMPLE001": "{{ sample_id }}",
        "测试诊断": "{{ clinical_diagnosis }}",
        "2099.12.31": "{{ report_date }}",
    }

    manifest = builder.build_seed(
        source,
        output,
        replacements=replacements,
        protected_tokens=tuple(replacements),
        allow_commit_output=False,
        allow_residual=False,
        project_root=ROOT,
    )

    assert manifest["success"] is True
    assert output.exists()
    assert all(
        count == 0
        for count in manifest["protected_token_residual_counts"].values()
    )
    text = "\n".join(p.text for p in Document(output).paragraphs)
    assert "{{ patient_name }}" in text
    assert "{{ sample_id }}" in text
    assert "{{ report_number }}" in text
    assert "{{ clinical_diagnosis }}" in text
    assert "{{ report_date }}" in text
    assert "333333333333" not in text
    assert manifest["removed_debug_markers"] == 1


def test_golden_template_variableizer_applies_structural_map(tmp_path):
    variableizer = _load_variableizer()
    source = tmp_path / "seed.docx"
    output = tmp_path / "variableized.docx"

    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "姓名："
    patient_cell = table.rows[0].cells[1]
    patient_cell.text = ""
    patient_run = patient_cell.paragraphs[0].add_run("黄金测试患者")
    patient_run.bold = True
    patient_run.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
    table.rows[1].cells[0].text = "性别："
    table.rows[1].cells[1].text = "男"
    table.rows[2].cells[0].text = "静态行"
    table.rows[2].cells[1].text = "删除我"
    loop_table = doc.add_table(rows=4, cols=3)
    loop_table.rows[0].cells[0].text = "基因"
    loop_table.rows[0].cells[1].text = "位点"
    loop_table.rows[0].cells[2].text = "说明"
    loop_table.rows[1].cells[0].text = "KRAS"
    loop_table.rows[1].cells[1].text = "c.34G>A"
    loop_table.rows[1].cells[2].text = "保留样式"
    loop_table.rows[2].cells[0].text = "TP53"
    loop_table.rows[2].cells[1].text = "c.844C>T"
    loop_table.rows[2].cells[2].text = "删除我"
    loop_table.rows[3].cells[0].text = "ATM"
    loop_table.rows[3].cells[1].text = "c.6874C>T"
    loop_table.rows[3].cells[2].text = "删除我"
    doc.add_paragraph("本次共检出体细胞变异：8个。")
    doc.add_paragraph("3.1 肿瘤突变负荷（TMB）水平提示")
    doc.add_paragraph("在本次检测范围内，该样本肿瘤突变负荷为6.5 mutations/Mb。")
    doc.add_paragraph("TMB科普说明。")
    doc.save(source)

    manifest = variableizer.variableize_docx(
        source,
        output,
        {
            "template_id": "unit_test",
            "cell_variables": [
                {"table": 0, "row": 0, "col": 1, "variable": "patient_name"},
                {"table": 0, "row": 1, "col": 1, "variable": "gender"},
            ],
            "paragraph_variables": [
                {
                    "after_heading": "3.1 肿瘤突变负荷（TMB）水平提示",
                    "nonempty_offset": 1,
                    "variable": "tmb_detail_sentence",
                }
            ],
            "paragraph_templates": [
                {
                    "contains": "本次共检出体细胞变异",
                    "text": "本次共检出体细胞变异：{{ total_variants_count }}个。",
                }
            ],
            "table_loops": [
                {
                    "id": "variants",
                    "table": 1,
                    "collection": "variants",
                    "alias": "row",
                    "template_row": 1,
                    "insert_at": 1,
                    "remove_from": 1,
                    "remove_to": "end",
                    "columns": ["gene", "locus", "中文键"],
                }
            ],
        },
    )

    out_doc = Document(output)
    assert manifest["success"] is True
    assert manifest["operation_count"] == 5
    assert out_doc.tables[0].rows[0].cells[1].text == "{{ patient_name }}"
    out_run = next(
        run
        for run in out_doc.tables[0].rows[0].cells[1].paragraphs[0].runs
        if run.text
    )
    assert out_run.bold is True
    assert out_run.font.color.rgb == RGBColor(0x12, 0x34, 0x56)
    assert out_doc.tables[0].rows[1].cells[1].text == "{{ gender }}"
    assert "total_variants_count" in "\n".join(p.text for p in out_doc.paragraphs)
    loop_rows = out_doc.tables[1].rows
    assert [cell.text for cell in loop_rows[1].cells] == [
        "{%tr for row in variants %}",
        "",
        "",
    ]
    assert [cell.text for cell in loop_rows[2].cells] == [
        "{{ row.gene }}",
        "{{ row.locus }}",
        '{{ row["中文键"] }}',
    ]
    assert [cell.text for cell in loop_rows[3].cells] == [
        "{%tr endfor %}",
        "",
        "",
    ]
    assert "{{ tmb_detail_sentence }}" in [p.text for p in out_doc.paragraphs]

    rendered = tmp_path / "rendered.docx"
    template = DocxTemplate(output)
    template.render(
        {
            "patient_name": "张三",
            "gender": "男",
            "tmb_detail_sentence": "TMB 渲染段落",
            "variants": [
                {"gene": "KRAS", "locus": "c.34G>A", "中文键": "A"},
                {"gene": "TP53", "locus": "c.844C>T", "中文键": "B"},
            ],
        }
    )
    template.save(rendered)
    rendered_doc = Document(rendered)
    assert [cell.text for cell in rendered_doc.tables[1].rows[1].cells] == [
        "KRAS",
        "c.34G>A",
        "A",
    ]
    assert [cell.text for cell in rendered_doc.tables[1].rows[2].cells] == [
        "TP53",
        "c.844C>T",
        "B",
    ]


def test_golden_diff_helpers_find_headings_and_tables(tmp_path):
    diff = _load_golden_diff()
    assert diff.find_heading_pages(
        ["封面", "报告导读\n检测结果", "3.1 肿瘤突变负荷（TMB）水平提示"],
        ["报告导读", "3.1 肿瘤突变负荷"],
    ) == {
        "报告导读": 2,
        "3.1 肿瘤突变负荷": 3,
    }

    docx_path = tmp_path / "tables.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "基因"
    table.rows[0].cells[1].text = "位点"
    table.rows[1].cells[0].text = "KRAS"
    table.rows[1].cells[1].text = "c.34G>A"
    doc.save(docx_path)

    summaries = diff.table_summaries(docx_path)
    assert summaries[0]["row_count"] == 2
    assert summaries[0]["col_count"] == 2
    assert summaries[0]["header"] == ["基因", "位点"]


def test_golden_diff_visual_metrics_detect_blank_page(tmp_path):
    from PIL import Image, ImageDraw

    diff = _load_golden_diff()
    blank = tmp_path / "report-1.png"
    content = tmp_path / "report-2.png"
    Image.new("RGB", (200, 300), "white").save(blank)
    img = Image.new("RGB", (200, 300), "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((20, 20, 180, 80), fill="black")
    img.save(content)

    blank_metrics = diff.image_visual_metrics(blank, blank_threshold=0.003)
    content_metrics = diff.image_visual_metrics(content, blank_threshold=0.003)

    assert blank_metrics["near_blank"] is True
    assert content_metrics["near_blank"] is False
    assert content_metrics["top_half_nonwhite_ratio"] > 0


def test_golden_diff_accepts_reference_blank_page_compaction(tmp_path):
    diff = _load_golden_diff()

    reference = diff.RenderedDocument(
        docx=tmp_path / "reference.docx",
        pdf=None,
        pngs=[tmp_path / f"ref-{i}.png" for i in range(1, 11)],
        page_texts=[""] * 4 + ["4. 基因检测列表"] + [""] * 5,
        visual_pages=[
            {
                "page": i,
                "near_blank": i == 2,
                "top_half_nonwhite_ratio": 0.01,
                "bottom_half_nonwhite_ratio": 0.01,
            }
            for i in range(1, 11)
        ],
    )
    candidate = diff.RenderedDocument(
        docx=tmp_path / "candidate.docx",
        pdf=None,
        pngs=[tmp_path / f"cand-{i}.png" for i in range(1, 10)],
        page_texts=[""] * 3 + ["4. 基因检测列表"] + [""] * 5,
        visual_pages=[
            {
                "page": i,
                "near_blank": False,
                "top_half_nonwhite_ratio": 0.01,
                "bottom_half_nonwhite_ratio": 0.01,
            }
            for i in range(1, 10)
        ],
    )

    section = diff.visual_diff_section(
        reference,
        candidate,
        headings=["4. 基因检测列表"],
    )

    assert section["status"] == "PASS"
    assert section["reference"]["blank_pages_accepted_as_reference_artifacts"] is True
    assert section["candidate"]["compact_layout_accepted"] is True
