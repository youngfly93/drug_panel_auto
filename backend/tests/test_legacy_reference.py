import json
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
BACKEND = ROOT / "backend"
for import_path in (str(ROOT), str(BACKEND)):
    if import_path not in sys.path:
        sys.path.insert(0, import_path)

from reportgen.core.legacy_reference import (  # noqa: E402
    LegacySnapshotOptions,
    build_legacy_reference_snapshots,
)


def test_legacy_reference_snapshot_redacts_patient_identifiers(tmp_path):
    source_dir = tmp_path / "legacy"
    source_dir.mkdir()
    docx_path = source_dir / "张三-直肠癌-结直肠癌301基因+msi-mljy-lz123456-终版.docx"
    doc = Document()
    doc.add_paragraph("结直肠癌301基因+MSI检测报告")
    doc.add_paragraph("姓名：张三")
    doc.add_paragraph("样本编号：LZ123456")
    doc.add_paragraph("报告编号：MLJY-LZ123456")
    doc.add_paragraph("报告日期：2025.12.12")
    doc.add_paragraph("本次共检出体细胞变异：2个")
    doc.add_paragraph("与靶向药物用药相关的变异有：1个")
    doc.add_paragraph("6.5 mutations/Mb，TMB-L")
    doc.add_paragraph("微卫星稳定型，MSS")
    doc.add_paragraph("致您的一封信")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "基因"
    table.rows[0].cells[1].text = "突变位点"
    table.rows[1].cells[0].text = "ERBB2"
    table.rows[1].cells[1].text = "c.1979G>A"
    doc.save(docx_path)

    result = build_legacy_reference_snapshots(
        LegacySnapshotOptions(
            panel="crc_301_msi",
            source_dir=str(source_dir),
            output_dir=str(tmp_path / "snapshots"),
            sample_count=1,
        )
    )

    assert result["status"] == "PASS"
    assert result["selected_count"] == 1
    assert result["readable_docx_count"] == 1
    assert result["read_error_count"] == 0
    manifest = json.loads(Path(result["manifest_file"]).read_text(encoding="utf-8"))
    assert manifest["source_dir"] == "<redacted>"
    assert manifest["selected_samples"][0]["features"]["total_variants_count"] == 2
    sample_path = Path(result["output_dir"]) / "samples" / "crc_301_msi_legacy_ref_001.json"
    payload = sample_path.read_text(encoding="utf-8")
    assert "张三" not in payload
    assert "LZ123456" not in payload
    assert "MLJY-LZ123456" not in payload
    assert "2025.12.12" not in payload
    assert "<PATIENT_NAME>" in payload
    assert "<SAMPLE_ID>" in payload
    assert "<REPORT_ID>" in payload
    assert "<DATE>" in payload


def test_legacy_reference_redacts_patient_name_when_filename_has_no_name(tmp_path):
    source_dir = tmp_path / "legacy"
    source_dir.mkdir()
    docx_path = source_dir / "crc301_case_001.docx"
    doc = Document()
    doc.add_paragraph("结直肠癌301基因+MSI检测报告")
    doc.add_paragraph("尊敬的 李四 女士")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "姓名"
    table.rows[0].cells[1].text = "李四"
    table.rows[1].cells[0].text = "样本编号"
    table.rows[1].cells[1].text = "LZ654321"
    doc.add_paragraph("本次共检出体细胞变异：1个")
    doc.add_paragraph("与靶向药物用药相关的变异有：0个")
    doc.save(docx_path)

    result = build_legacy_reference_snapshots(
        LegacySnapshotOptions(
            panel="crc_301_msi",
            source_dir=str(source_dir),
            output_dir=str(tmp_path / "snapshots"),
            sample_count=1,
        )
    )

    sample_path = Path(result["output_dir"]) / "samples" / "crc_301_msi_legacy_ref_001.json"
    payload = sample_path.read_text(encoding="utf-8")
    assert "李四" not in payload
    assert "LZ654321" not in payload
    assert "<PATIENT_NAME>" in payload


def test_quality_gate_runs_legacy_reference_when_source_root_exists(tmp_path):
    from reportgen.core.qa_gate import QualityGateOptions, run_quality_gate

    source_root = tmp_path / "legacy_reports_by_panel"
    source_dir = source_root / "crc_301_msi"
    source_dir.mkdir(parents=True)
    docx_path = source_dir / "crc301_case_001.docx"
    doc = Document()
    doc.add_paragraph("结直肠癌301基因+MSI检测报告")
    doc.add_paragraph("姓名：王五")
    doc.add_paragraph("本次共检出体细胞变异：3个")
    doc.add_paragraph("与靶向药物用药相关的变异有：1个")
    doc.add_paragraph("6.5 mutations/Mb，TMB-L")
    doc.add_paragraph("该肿瘤样本为微卫星稳定型，MSS")
    doc.add_paragraph("基因变异检测结果")
    doc.add_paragraph("TMB/MSI")
    doc.add_paragraph("基因检测列表")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "基因"
    table.rows[0].cells[1].text = "突变位点"
    table.rows[1].cells[0].text = "KRAS"
    table.rows[1].cells[1].text = "c.34G>A"
    doc.save(docx_path)

    result = run_quality_gate(
        QualityGateOptions(
            project_root=str(ROOT),
            output_root=str(tmp_path / "gate"),
            run_lint=False,
            run_pytest=False,
            run_golden=False,
            run_legacy_reference=True,
            legacy_source_root=str(source_root),
            legacy_panels=("crc_301_msi",),
            legacy_sample_count=1,
            legacy_required=True,
        )
    )

    legacy_step = next(
        step for step in result["steps"] if step["name"] == "legacy_reference_crc_301_msi"
    )
    assert result["status"] == "PASS"
    assert legacy_step["status"] == "PASS"
    assert legacy_step["summary"]["selected_count"] == 1
    assert Path(legacy_step["manifest_file"]).exists()


def test_quality_gate_can_require_legacy_reference_source(tmp_path):
    from reportgen.core.qa_gate import QualityGateOptions, run_quality_gate

    result = run_quality_gate(
        QualityGateOptions(
            project_root=str(ROOT),
            output_root=str(tmp_path / "gate"),
            run_lint=False,
            run_pytest=False,
            run_golden=False,
            run_legacy_reference=True,
            legacy_source_root=str(tmp_path / "missing"),
            legacy_required=True,
        )
    )

    legacy_step = next(
        step for step in result["steps"] if step["name"] == "legacy_reference"
    )
    assert result["status"] == "FAIL"
    assert legacy_step["status"] == "FAIL"
    assert legacy_step["issues"][0]["code"] == "LEGACY_REFERENCE_SOURCE_MISSING"


def test_current_output_contract_uses_panel_qa_profile(tmp_path):
    from reportgen.core import qa_gate

    docx_path = tmp_path / "golden_crc_301_msi.docx"
    doc = Document()
    doc.add_paragraph("结直肠癌301基因+MSI检测报告")
    doc.add_paragraph("本次共检出体细胞变异：3个")
    doc.add_paragraph("与靶向药物用药相关的变异有：1个")
    doc.add_paragraph("6.5 mutations/Mb，TMB-L")
    doc.add_paragraph("该肿瘤样本为微卫星稳定型，MSS")
    doc.add_paragraph("基因变异检测结果")
    doc.add_paragraph("TMB/MSI")
    doc.add_paragraph("基因检测列表")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "基因"
    table.rows[0].cells[1].text = "突变位点"
    table.rows[1].cells[0].text = "KRAS"
    table.rows[1].cells[1].text = "c.34G>A"
    doc.save(docx_path)

    steps = qa_gate._run_current_output_steps(
        qa_gate.QualityGateOptions(
            project_root=str(ROOT),
            panels=("crc_301_msi",),
        ),
        project_root=ROOT,
        output_root=tmp_path / "current_output",
        prior_steps=[
            {
                "name": "golden_crc_301_msi_reference",
                "status": "PASS",
                "panel": "crc_301_msi",
                "output_file": str(docx_path),
            }
        ],
    )

    assert steps[0]["name"] == "current_output_crc_301_msi"
    assert steps[0]["status"] == "PASS"
    assert Path(steps[0]["snapshot_file"]).exists()
    assert steps[0]["qa_profile_file"].endswith("panels/crc_301_msi/qa.yaml")
