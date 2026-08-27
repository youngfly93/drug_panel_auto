# ruff: noqa: E402
"""All declared ReportGen templates must have an explicit drift baseline.

Source-template checks are fast and PHI-safe. They do not replace rendered
golden-case QA; the matrix records rendered baselines separately so a static
template hash can never be mistaken for end-to-end acceptance evidence.
"""

from __future__ import annotations

import json
import os
import sys
from pathlib import Path

import pytest
import yaml
from docx import Document

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from reportgen.core.style_baseline import (
    diff_baseline,
    extract_template_style_baseline,
)
from reportgen.panels.loader import PanelPackageLoader

MATRIX_PATH = Path(__file__).with_name("template_style_baseline_matrix.yaml")
UPDATE_ENV = "REPORTGEN_UPDATE_TEMPLATE_BASELINES"


def _matrix_entries() -> list[dict]:
    payload = yaml.safe_load(MATRIX_PATH.read_text(encoding="utf-8")) or {}
    assert payload.get("schema_version") == "1.0"
    entries = payload.get("templates") or []
    assert isinstance(entries, list)
    return entries


def _declared_templates() -> dict[tuple[str, str], tuple[object, object]]:
    declared = {}
    for package in PanelPackageLoader(project_root=ROOT).load_all():
        for template in package.templates.values():
            declared[(package.panel_id, template.template_id)] = (package, template)
    return declared


def _entry_id(entry: dict) -> str:
    return f"{entry['panel_id']}::{entry['template_id']}"


def test_template_style_matrix_covers_every_declared_template_once():
    entries = _matrix_entries()
    actual_keys = [
        (str(entry.get("panel_id") or ""), str(entry.get("template_id") or ""))
        for entry in entries
    ]
    declared = _declared_templates()

    assert len(actual_keys) == len(set(actual_keys)), "模板样式基线矩阵存在重复项"
    assert set(actual_keys) == set(declared), (
        "模板样式基线矩阵必须覆盖 panel.yaml 声明的每个模板；"
        f" missing={sorted(set(declared) - set(actual_keys))}"
        f" extra={sorted(set(actual_keys) - set(declared))}"
    )

    for entry in entries:
        key = (entry["panel_id"], entry["template_id"])
        package, template = declared[key]
        assert entry["status"] == template.status
        assert bool(entry["default"]) == (template.template_id == package.default_template_id)
        assert str(entry.get("source_baseline") or "").endswith(".json")
        rendered = str(entry.get("rendered_output_baseline") or "")
        if rendered:
            assert (Path(__file__).parent / rendered).exists()


@pytest.mark.parametrize("entry", _matrix_entries(), ids=_entry_id)
def test_declared_template_source_style_baseline(entry: dict):
    package, template = _declared_templates()[(entry["panel_id"], entry["template_id"])]
    template_path = package.resolve_template_file(template.template_id)
    baseline_path = Path(__file__).parent / entry["source_baseline"]

    assert template_path.exists(), template_path
    actual = extract_template_style_baseline(template_path)
    if os.environ.get(UPDATE_ENV) == "1":
        baseline_path.parent.mkdir(parents=True, exist_ok=True)
        baseline_path.write_text(
            json.dumps(actual, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        return

    assert baseline_path.exists(), (
        f"缺模板源基线 {baseline_path}; 仅在确认模板变化有意后设置 {UPDATE_ENV}=1 更新"
    )
    expected = json.loads(baseline_path.read_text(encoding="utf-8"))
    diffs = diff_baseline(expected, actual)
    assert not diffs, (
        f"[{_entry_id(entry)}] 模板源样式/布局发生漂移。"
        f"若变化有意，设置 {UPDATE_ENV}=1 更新并复核差异。\n"
        + "\n".join(diffs[:40])
    )


def test_template_source_fingerprint_never_emits_document_text(tmp_path):
    secret = "PRIVATE_PATIENT_LZ999999"
    template_path = tmp_path / "privacy.docx"
    document = Document()
    document.add_paragraph(secret)
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = secret
    document.save(template_path)

    baseline = extract_template_style_baseline(template_path)

    assert secret not in json.dumps(baseline, ensure_ascii=False)
    assert baseline["privacy"] == {
        "paragraph_text_emitted": False,
        "cell_text_emitted": False,
        "media_bytes_emitted": False,
    }


def test_template_source_fingerprint_detects_content_drift(tmp_path):
    template_path = tmp_path / "drift.docx"
    document = Document()
    document.add_paragraph("before")
    document.save(template_path)
    before = extract_template_style_baseline(template_path)

    document = Document(template_path)
    document.paragraphs[0].text = "after"
    document.save(template_path)
    after = extract_template_style_baseline(template_path)

    assert before["aggregate_sha256"] != after["aggregate_sha256"]
