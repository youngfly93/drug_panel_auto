from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.template_fit_analyzer import mine_section_titles  # noqa: E402


def _write_docx(path: Path, *, case_suffix: int) -> None:
    doc = Document()
    doc.add_paragraph("第一部分：检测报告摘要")
    doc.add_paragraph("报告编号：MLJY-LZ25")
    doc.add_paragraph("送检日期：2025")
    doc.add_paragraph(f"样本编号：LZ25{case_suffix:04d}")
    doc.add_paragraph("咨询电话：000-00000000")
    doc.save(path)


def test_section_mining_excludes_recurring_case_metadata_and_contacts(
    tmp_path: Path,
) -> None:
    reports = []
    for case_suffix in range(3):
        path = tmp_path / f"case-{case_suffix}.docx"
        _write_docx(path, case_suffix=case_suffix)
        reports.append(path)

    result = mine_section_titles(reports, [], min_doc_freq=0.6)
    lines = {row[0] for row in result["lines"]}

    assert "第一部分：检测报告摘要" in lines
    assert not any("报告编号" in row for row in lines)
    assert not any("送检日期" in row for row in lines)
    assert not any("样本编号" in row for row in lines)
    assert not any("电话" in row for row in lines)
