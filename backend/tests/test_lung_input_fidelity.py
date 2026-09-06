"""Real-XLSX regressions for CNV parsing and molecule/genotype fidelity."""

from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook
from reportgen.core.excel_reader import ExcelReader
from reportgen.core.qa_report import _build_business_checks
from reportgen.core.template_bridge_358 import (
    PanelConfig,
    _build_lung_chemotherapy_tables,
    _build_nccn_and_immune_fields,
    build_immune_variants,
    load_panel_config,
)
from reportgen.core.template_renderer import TemplateRenderer
from reportgen.core.validation import build_tmb_fields
from reportgen.models.report_data import ReportData
from reportgen.panels.loader import load_panel_package

ROOT = Path(__file__).resolve().parents[2]
SUMMARY_HEADER = ["Gene", "Cnvkit", "AvgCP", "ExonNum"]
CALL_HEADER = ["#Chr", "Start", "End", "Status", "CopyNum(X)", "Gene", "AvgCP"]


def _xlsx(tmp_path, *, cnv=None, ctdrug=None, ct1000=None):
    workbook = Workbook()
    workbook.active.title = "Meta"
    workbook.active.append(["Key", "Value"])
    workbook.active.append(["sample_id", "SYN-LUNG-FIDELITY"])
    for name, rows in (("Cnv", cnv), ("CtDrug", ctdrug), ("Ct1000", ct1000)):
        if rows is not None:
            sheet = workbook.create_sheet(name)
            for row in rows:
                sheet.append(row)
    path = tmp_path / "SYN-LUNG-FIDELITY.xlsx"
    workbook.save(path)
    return ExcelReader(config_dir=str(ROOT / "config"), log_level="ERROR").read(str(path))


@pytest.fixture(params=["lung_329_pdl1", "lung_588_pdl1"])
def lung_config(request):
    package = load_panel_package(request.param, project_root=ROOT)
    return load_panel_config(base_path=str(ROOT), panel_package=package)


@pytest.mark.parametrize("prefix_rows", [0, 2, 5])
def test_cnv_semantic_headers_preserve_both_blocks_and_source_rows(tmp_path, prefix_rows):
    source = _xlsx(
        tmp_path,
        cnv=[
            *[["CNV export"] for _ in range(prefix_rows)],
            SUMMARY_HEADER,
            ["EGFR", 0, 24.5, 28],
            ["FGF3", 1, 7.25, 3],
            [],
            CALL_HEADER,
            ["chr7", 100, 200, "gain", 12.25, "EGFR", 12.25],
            ["chr11", 300, 400, "gain", 3.625, "FGF3", 3.625],
        ],
    )
    assert [row["Gene"] for row in source.get_table_data("Cnv")] == ["EGFR", "FGF3"]
    assert source.get_table_data("Cnv")[0]["_source_row"] == prefix_rows + 6
    assert float(source.get_table_data("Cnv")[0]["CopyNum(X)"]) == 12.25
    assert float(source.get_table_data("Cnv_summary")[0]["AvgCP"]) == 24.5
    assert str(source.get_table_data("Cnv_summary")[0]["Cnvkit"]) == "0"
    assert source.metadata["cnv_parse"]["status"] == "parsed"


def test_malformed_nonempty_cnv_is_not_silently_treated_as_negative(tmp_path):
    with pytest.raises(ValueError, match="CNV|Cnv"):
        _xlsx(tmp_path, cnv=[["wrong_gene_column", "mystery"], ["EGFR", 24.5]])


def test_valid_empty_cnv_header_remains_distinguishable_from_missing(tmp_path):
    source = _xlsx(tmp_path, cnv=[SUMMARY_HEADER, [], CALL_HEADER])
    assert source.get_table_data("Cnv") == []
    assert source.metadata["cnv_parse"]["status"] == "parsed"
    assert "Gene" in source.metadata["table_columns"]["Cnv"]


def test_legacy_copy_number_header_supports_a_valid_empty_cnv_sheet(tmp_path):
    source = _xlsx(tmp_path, cnv=[["Gene", "Chr", "Start", "End", "CopyNumber"]])
    assert source.get_table_data("Cnv") == []
    assert source.metadata["cnv_parse"]["status"] == "parsed"
    assert "CopyNumber" in source.metadata["table_columns"]["Cnv"]


def test_numeric_copy_number_is_retained_for_review_not_called_amplification(tmp_path, lung_config):
    source = _xlsx(
        tmp_path,
        cnv=[["Gene", "Chr", "Start", "End", "CopyNumber"], ["EGFR", "7", 100, 200, 8]],
    )
    assert float(source.get_table_data("Cnv")[0]["CopyNumber"]) == 8
    data = ReportData()
    _build_nccn_and_immune_fields(data, [], source, panel_config=lung_config)
    text = data.get_field("imm_hyper_EGFR_AMP")
    assert "CopyNumber=8" in text and "待复核" in text
    assert build_immune_variants(source, panel_config=lung_config)["hyperprogression"] == []


@pytest.mark.parametrize("status", ["gain", "loss", "AMP", "not amplified", 0, 1])
def test_lung_cnv_source_is_visible_but_not_a_new_immune_claim(tmp_path, lung_config, status):
    source = _xlsx(
        tmp_path,
        cnv=[
            CALL_HEADER,
            ["chr7", 100, 200, status, 8.0, "EGFR", 8.0],
            ["chr11", 300, 400, status, 6.0, "FGF3", 6.0],
        ],
    )
    data = ReportData()
    _build_nccn_and_immune_fields(data, [], source, panel_config=lung_config)
    for key in ("EGFR_AMP", "FGF3"):
        text = data.get_field(f"imm_hyper_{key}")
        assert "CNV" in text
        assert "复核" in text
        assert "未检出有害变异" not in text
    assert build_immune_variants(source, panel_config=lung_config)["hyperprogression"] == []
    assert data.get_field("cnv_review_required") is True
    assert "待复核" in data.get_field("immuno_hyperprogression_genes")


def test_summary_only_numeric_cnvkit_is_not_an_amplification_call(tmp_path, lung_config):
    source = _xlsx(tmp_path, cnv=[SUMMARY_HEADER, ["EGFR", 1, 30.0, 28]])
    data = ReportData()
    _build_nccn_and_immune_fields(data, [], source, panel_config=lung_config)
    assert "复核" in data.get_field("imm_hyper_EGFR_AMP")
    assert str(source.get_table_data("Cnv")[0]["Cnvkit"]) == "1"
    assert build_immune_variants(source, panel_config=lung_config)["hyperprogression"] == []


@pytest.mark.parametrize("star28", ["6TA/6TA", "6TA/7TA", "7TA/7TA"])
@pytest.mark.parametrize("star6", ["GG", "AG", "AA"])
def test_ugt1a1_both_loci_are_displayed_without_automatic_dose(
    tmp_path, lung_config, star28, star6
):
    source = _xlsx(
        tmp_path,
        ctdrug=[
            ["药物", "检测基因", "检测位点", "基因型"],
            ["伊立替康（irinotecan）", "UGT1A1", "rs8175347", star28],
            ["伊立替康（irinotecan）", "UGT1A1", "rs4148323", star6],
        ],
    )
    data = ReportData()
    _build_lung_chemotherapy_tables(data, source, lung_config)
    row = data.get_table("irinotecan_safety_rows")[0]
    assert f"*28（rs8175347）：{star28}" in row["result"]
    assert f"*6（rs4148323）：{star6}" in row["result"]
    assert "待医学复核" in row["dose_evaluation"]
    assert "正常剂量使用" not in str(row)
    assert "减少剂量使用" not in str(row)


@pytest.mark.parametrize(
    "calls, expected", [([], "未提供"), (["GG", "AA"], "冲突"), (["Uncovered"], "未覆盖")]
)
def test_ugt1a1_missing_conflicting_and_uncovered_calls_are_not_conflated(
    tmp_path, lung_config, calls, expected
):
    source = _xlsx(
        tmp_path,
        ctdrug=[
            ["药物", "检测基因", "检测位点", "基因型"],
            ["伊立替康", "UGT1A1", "rs8175347", "6TA/7TA"],
            *[["伊立替康", "UGT1A1", "rs4148323", value] for value in calls],
        ],
    )
    data = ReportData()
    _build_lung_chemotherapy_tables(data, source, lung_config)
    row = data.get_table("irinotecan_safety_rows")[0]
    assert expected in row["result"]
    assert "待医学复核" in row["dose_evaluation"]


@pytest.mark.parametrize(
    "molecule", ["长春新碱（vincristine）", "长春碱（vinblastine）", "长春瑞滨+顺铂"]
)
def test_wrong_or_combined_molecule_cannot_supply_vinorelbine_prediction(
    tmp_path, lung_config, molecule
):
    source = _xlsx(
        tmp_path,
        ct1000=[
            [None, molecule, None, None, "ABCB1"],
            ["sensitive H! Toxicity L!"],
            [None, "吉西他滨（gemcitabine）", None, None, "CDA"],
            ["sensitive H! Toxicity L!"],
        ],
        ctdrug=[
            ["药物", "检测基因", "用药提示（仅供参考）", "等级"],
            [molecule, "CEP72", "疗效较好，毒性较低", "1B"],
            ["长春瑞滨（vinorelbine）", "CASP7", "疗效较好，毒性较低", "3"],
        ],
    )
    data = ReportData()
    _build_lung_chemotherapy_tables(data, source, lung_config)
    rows = {row["drug"]: row for row in data.get_table("chemotherapy_predictions")}
    assert rows["长春瑞滨"]["efficacy"] == "未提供同药物证据"
    assert rows["长春瑞滨"]["toxicity"] == "未提供同药物证据"
    assert rows["长春瑞滨"]["genes"] == "/"
    assert rows["吉西他滨+长春瑞滨"]["efficacy"] == "未提供同药物证据"
    assert "吉西他滨+长春瑞滨" not in data.get_field("chemotherapy_summary_text")


@pytest.mark.parametrize(
    "molecule", ["长春瑞滨（vinorelbine）", "vinorelbine", "VINORELBINE（长春瑞滨）"]
)
def test_exact_vinorelbine_summary_is_accepted(tmp_path, lung_config, molecule):
    source = _xlsx(
        tmp_path,
        ct1000=[
            [None, molecule, None, None, "SYNTHETIC_GENE"],
            ["sensitive H! Toxicity L!"],
        ],
    )
    data = ReportData()
    _build_lung_chemotherapy_tables(data, source, lung_config)
    row = next(
        row for row in data.get_table("chemotherapy_predictions") if row["drug"] == "长春瑞滨"
    )
    assert row["efficacy"] == "可能较高"
    assert row["toxicity"] == "可能较低"
    assert row["genes"] == "SYNTHETIC_GENE"


def test_mixed_molecules_in_one_ct1000_block_are_not_a_vinorelbine_source(tmp_path, lung_config):
    source = _xlsx(
        tmp_path,
        ct1000=[
            [None, "长春瑞滨（vinorelbine）", None, None, "GENE_A"],
            [None, "长春新碱（vincristine）", None, None, "GENE_B"],
            ["sensitive H! Toxicity L!"],
        ],
    )
    assert len(source.get_table_data("Ct1000_summary")[0]["source_drugs"]) == 2
    data = ReportData()
    _build_lung_chemotherapy_tables(data, source, lung_config)
    row = next(
        row for row in data.get_table("chemotherapy_predictions") if row["drug"] == "长春瑞滨"
    )
    assert row["efficacy"] == "未提供同药物证据"


def test_conflicting_exact_drug_summaries_are_not_resolved_by_row_order(tmp_path, lung_config):
    source = _xlsx(
        tmp_path,
        ct1000=[
            [None, "vinorelbine", None, None, "GENE_A"],
            ["sensitive H! Toxicity L!"],
            [None, "vinorelbine", None, None, "GENE_A"],
            ["sensitive L! Toxicity H!"],
        ],
    )
    data = ReportData()
    _build_lung_chemotherapy_tables(data, source, lung_config)
    row = next(
        row for row in data.get_table("chemotherapy_predictions") if row["drug"] == "长春瑞滨"
    )
    assert row["efficacy"] == "未提供同药物证据"


@pytest.mark.parametrize(
    "status, expected", [("loss", "未检出"), ("gain", "未检出"), ("AMP", "CNV:AMP")]
)
def test_generic_amplification_row_does_not_treat_loss_or_gain_as_amp(tmp_path, status, expected):
    source = _xlsx(tmp_path, cnv=[CALL_HEADER, ["chr17", 100, 200, status, 4.0, "ERBB2", 4.0]])
    config = PanelConfig(
        nccn_result_rows=[{"key": "ERBB2_AMP", "genes": ["ERBB2"], "match": "扩增"}]
    )
    data = ReportData()
    _build_nccn_and_immune_fields(data, [], source, panel_config=config)
    assert data.get_field("nccn_ERBB2_AMP") == expected


def test_lung_unconfirmed_met_cnv_is_not_called_a_joint_met_egfr_event(tmp_path, lung_config):
    source = _xlsx(tmp_path, cnv=[CALL_HEADER, ["chr7", 100, 200, "gain", 4.0, "MET", 4.0]])
    data = ReportData()
    _build_nccn_and_immune_fields(
        data,
        [{"gene": "EGFR", "cHGVS": "c.2573T>G", "pHGVS": "p.L858R", "gene_class": "Ⅰ类"}],
        source,
        panel_config=lung_config,
    )
    row = next(
        row
        for row in data.get_table("lung_guideline_drug_results")
        if row["key"] == "MET_EGFR_JOINT"
    )
    assert "无法确认联合变异" in row["result"]


@pytest.mark.parametrize("notice, expected", [("CNV待复核：EGFR", "WARN"), ("未检出", "FAIL")])
def test_cnv_review_notice_is_required_in_rendered_report(notice, expected):
    checks = _build_business_checks(
        notice, {"cnv_review_required": True, "cnv_review_genes": ["EGFR"]}, "lung_588_pdl1"
    )
    assert checks["cnv_source_review"]["status"] == expected


@pytest.mark.parametrize(
    "value, expected_status, expected_direction",
    [(9.0, "L", "低于"), (10.0, "H", "等于"), (11.0, "H", "高于")],
)
def test_tmb_threshold_wording_matches_the_numeric_comparison(
    value, expected_status, expected_direction
):
    fields = build_tmb_fields(value, sample_type="组织")
    assert fields["tmb_status"] == expected_status
    assert f"本次检测结果{expected_direction}参考值" in fields["tmb_summary"]


@pytest.mark.parametrize(
    "panel_id, grouped", [("lung_329_pdl1", True), ("lung_588_pdl1", True), ("crc_358_msi", False)]
)
def test_short_lung_drug_analysis_keeps_review_tail_with_its_source_block(
    tmp_path, panel_id, grouped
):
    path = tmp_path / "synthetic-layout.docx"
    document = Document()
    texts = [
        "第三部分：基因变异及相应靶向/免疫药物解析",
        "靶向药物/免疫用药提示解析",
        "潜在获益靶向/免疫药物解析",
        "GENE：c.1A>T突变相应靶向药物",
        "SYNTHETIC-DRUG",
        "基因变异与药物关联分析：",
        "用于排版测试的合成来源说明。" * 12,
        "药物疗效临床解析：",
        "该段历史知识尚未完成肺癌专属审核；待报告组复核。",
        "3. 阅读说明",
    ]
    for text in texts:
        document.add_paragraph(text)
    document.save(path)
    renderer = TemplateRenderer(log_level="ERROR")
    renderer._restore_part3_dynamic_styles(
        str(path),
        {
            "project_type": panel_id,
            "drug_benefit_sections": [{"drug_name": "SYNTHETIC-DRUG"}],
        },
    )
    result = Document(path)
    assert [p.text for p in result.paragraphs] == texts
    for index in (5, 6, 7):
        assert bool(result.paragraphs[index].paragraph_format.keep_with_next) is grouped
    assert not result.paragraphs[8].paragraph_format.keep_with_next
