# ruff: noqa: E402, I001

import sys
import subprocess
import json
import shutil
from datetime import date
from pathlib import Path
from types import SimpleNamespace
from zipfile import ZipFile

import pytest
import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Inches
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from reportgen.config.loader import ConfigLoader
from reportgen.core.batch_runner import (
    BatchValidateOptions,
    _expected_tables_from_excel,
)
from reportgen.core.data_cleaner import DataCleaner
from reportgen.core.excel_reader import ExcelReader
from reportgen.core.field_provenance import (
    build_field_provenance_report,
    write_field_provenance_report,
)
from reportgen.core.field_mapper import FieldMapper
from reportgen.core.golden_case import (
    CRC_358_MSI_EXPECTATIONS,
    GoldenCaseOptions,
    LUNG_METHYLATION_EXPECTATIONS,
    assert_golden_case_output,
    build_crc_301_msi_golden_excel,
    build_crc_358_msi_golden_excel,
    build_lung_329_pdl1_golden_input,
    build_lung_588_pdl1_golden_input,
    build_lung_methylation_golden_excel,
    run_golden_case,
    run_visual_render,
)
from reportgen.core.legacy_reference import snapshot_docx_report
from reportgen.core.project_detector import ProjectDetector
from reportgen.core.processors import (
    CRITICAL_DOCX_PROCESSOR_NAMES,
    critical_docx_processor_names,
    ProcessorContext,
    run_processors,
)
from reportgen.core.processors.docx import (
    _run_final_refresh_cleanup,
    _run_front_matter_spacing,
    _run_underlines_and_styles,
)
from reportgen.core.qa_report import build_docx_qa_report, write_docx_qa_report
from reportgen.core.report_summary import build_report_summary, write_report_summary
from reportgen.core.report_generator import ReportGenerator
from reportgen.core.signature_library import resolve_signature_path, signature_options
from reportgen.core.report_diff import ReportDiffOptions, compare_reports
from reportgen.core.template_bridge_358 import (
    PanelConfig,
    _build_targeted_drug_brand_summary_result,
    _build_nccn_and_immune_fields,
    _compact_drug_display_tables,
    _variant_override_matches,
    _patch_reviewed_variant_override_rows,
    build_targeted_drug_brand_summary,
    build_tmb_summary,
    build_variants_for_template,
    enhance_report_data,
    load_panel_config,
)
from reportgen.core.template_renderer import TemplateRenderer
from reportgen.core.validation import validate_excel_data_common
from reportgen.models.excel_data import ExcelDataSource
from reportgen.models.report_data import ReportData
from reportgen.panels.loader import (
    PanelPackageLoader,
    load_panel_package,
    validate_panel_package_config,
)
from reportgen.panels.registry import PanelRegistry
from reportgen.panels.validation import (
    validate_panel_package,
    validate_panel_package_path,
    validate_panel_registry,
)
from reportgen.rules import PanelRuleEngine, load_rule_package
from reportgen.rules.targeted_drugs import load_targeted_drug_rule_context
from app.services import clinical_info_service
from reportgen.rules.evaluators import apply_report_text_rules, collect_report_texts
from reportgen.knowledge.gene_knowledge import GeneKnowledgeProvider
from reportgen.docx_sections import find_reference_section_bounds
from reportgen.utils import docx_render


def _read_docx_part(docx_path: Path, part_name: str) -> str:
    with ZipFile(docx_path) as zf:
        return zf.read(part_name).decode("utf-8")


def _toc_sdt_xml(docx_path: Path) -> str:
    """Return the serialized XML of the TOC content control only.

    The report body legitimately contains external HYPERLINK fields (reference
    URLs), so field-absence checks for the TOC must be scoped to the TOC SDT
    rather than the whole document.
    """
    import re

    from lxml import etree

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(docx_path) as zf:
        root = etree.fromstring(zf.read("word/document.xml"))

    def sdt_text(elem) -> str:
        return "".join(elem.xpath(".//w:t/text()", namespaces=ns))

    for sdt in root.xpath(".//w:sdt", namespaces=ns):
        text = sdt_text(sdt)
        # The TOC SDT carries 参考文献 plus a 第N部分 part marker. (A partial
        # rebuild may include only the appendix part, so match any 第N部分.)
        if "参考文献" in text and re.search(r"第[一二三四]部分", text):
            return etree.tostring(sdt, encoding="unicode")
    raise AssertionError("TOC content control not found")


def _excel(
    tmp_path: Path,
    *,
    single_values=None,
    variations=None,
    tables=None,
) -> ExcelDataSource:
    path = tmp_path / "unknown.xlsx"
    path.write_bytes(b"placeholder")
    table_data = dict(tables or {})
    if variations is not None:
        table_data["Variations"] = variations
    return ExcelDataSource(
        file_path=str(path),
        single_values=single_values or {},
        table_data=table_data,
        sheet_names=list(table_data),
        metadata={},
    )


def test_data_cleaner_exposes_date_display_aliases():
    report_data = ReportData()
    report_data.set_field("receive_date", "2025.11.21")
    report_data.set_field("report_date", "2025-12-04")

    DataCleaner(log_level="ERROR").validate_and_clean(report_data)

    assert report_data.get_field("receive_date") == "2025-11-21"
    assert report_data.get_field("receive_date_compact") == "20251121"
    assert report_data.get_field("receive_date_dot") == "2025.11.21"
    assert report_data.get_field("report_date") == "2025-12-04"
    assert report_data.get_field("report_date_compact") == "20251204"
    assert report_data.get_field("report_date_dot") == "2025.12.04"


def test_data_cleaner_exposes_visible_aliases_for_missing_optional_dates():
    report_data = ReportData()

    DataCleaner(log_level="ERROR").validate_and_clean(report_data)

    assert report_data.get_field("receive_date") is None
    assert report_data.get_field("receive_date_compact") == "未提供"
    assert report_data.get_field("receive_date_dot") == "未提供"


def test_field_mapper_derives_report_number_from_sample_id(tmp_path):
    excel_data = _excel(tmp_path)
    excel_data.metadata["sample_id_from_filename"] = "lz258792"

    report_data = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR").map(excel_data)

    assert report_data.get_field("report_number") == "MLJY-LZ258792"


def test_field_mapper_keeps_explicit_report_number(tmp_path):
    excel_data = _excel(tmp_path, single_values={"报告编号": "CUSTOM-REPORT-001"})
    excel_data.metadata["sample_id_from_filename"] = "lz258792"

    report_data = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR").map(excel_data)

    assert report_data.get_field("report_number") == "CUSTOM-REPORT-001"


def test_field_mapper_receive_date_becomes_cover_compact_alias(tmp_path):
    excel_data = _excel(
        tmp_path,
        single_values={"送检日期": "2025.11.21", "报告日期": "2025-12-04"},
    )
    excel_data.metadata["sample_id_from_filename"] = "lz258792"

    report_data = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR").map(excel_data)
    DataCleaner(log_level="ERROR").validate_and_clean(report_data)

    assert report_data.get_field("receive_date") == "2025-11-21"
    assert report_data.get_field("receive_date_compact") == "20251121"
    assert report_data.get_field("report_date_compact") == "20251204"


class _FakeDrugLookup:
    def _lookup_targeted_drugs_for_variant(
        self,
        gene,
        *,
        c_point,
        p_point,
        variant_level="",
        cancer_type="",
    ):
        if gene == "KRAS":
            return "Avutometinib+Defactinib（C）\n司美替尼（C）", "西妥昔单抗（A）", 100.0
        if gene == "ATM":
            return "奥拉帕利+帕博利珠单抗（C）", "--", 100.0
        return f"{gene}获益药（C）", "--", 100.0


class _FakeGeneKnowledgeProvider:
    def load(self, base_path=None):
        return None

    def build_all_gene_knowledge_sections(self, variants, cancer_type=""):
        return [
            {
                "header": f"{row.get('gene')}：{row.get('cHGVS')}",
                "has_drug": row.get("benefit_drugs") not in ("", "--", None),
                "intro": "基因简介",
                "mutation_desc": "位点说明",
                "mutation_analysis": "解析内容",
            }
            for row in variants
        ]

    def build_drug_analysis_sections(self, variants):
        return []

    def build_all_references_flat(self, variants, max_per_gene=5):
        return [f"{row.get('gene')} reference" for row in variants]

    def build_references(self, variants, max_per_gene=5):
        return {row.get("gene"): [f"{row.get('gene')} reference"] for row in variants}


class _FakeDrugSectionGeneKnowledgeProvider(_FakeGeneKnowledgeProvider):
    def build_drug_analysis_sections(self, variants):
        return [
            {
                "gene": row.get("gene"),
                "drug_name": row.get("benefit_drugs"),
                "drug_type": "benefit",
            }
            for row in variants
        ]


def test_reviewed_override_does_not_create_absent_erbb2_tip(tmp_path):
    excel_data = _excel(tmp_path, variations=[])
    report_data = enhance_report_data(
        ReportData(),
        excel_data,
        base_path=str(ROOT),
    )

    tips = report_data.get_table("targeted_drug_tips")
    assert all(row.get("gene") != "ERBB2" for row in tips)
    assert report_data.get_field("msi_status") == "未检测"
    assert report_data.get_field("tmb_status") == "未检测"


def test_reviewed_override_applies_when_erbb2_variant_is_detected(tmp_path):
    excel_data = _excel(
        tmp_path,
        variations=[
            {
                "Gene_Symbol": "ERBB2",
                "Transcript": "NM_004448.4",
                "Chr": "17",
                "ExIn_ID": "EX17",
                "cHGVS": "c.1979G>A",
                "pHGVS_S": "p.G660D",
                "Freq(%)": 12.3,
                "ExistInsmall358": 1,
                "ExistIn552": 1,
                "CLNSIG": "Pathogenic",
            }
        ],
    )
    report_data = enhance_report_data(
        ReportData(),
        excel_data,
        base_path=str(ROOT),
    )

    tips = report_data.get_table("targeted_drug_tips")
    erbb2 = [row for row in tips if row.get("gene") == "ERBB2"]
    assert erbb2
    assert "曲妥珠单抗" in erbb2[0].get("benefit_drugs", "")


def test_nccn_mutation_rows_do_not_include_cnv_or_fusion(tmp_path):
    excel_data = _excel(
        tmp_path,
        variations=[
            {
                "Gene_Symbol": "ERBB2",
                "Transcript": "NM_004448.4",
                "Chr": "17",
                "ExIn_ID": "EX20",
                "cHGVS": "c.2324_2325ins12",
                "pHGVS_S": "p.Y772_A775dup",
                "Freq(%)": 12.3,
                "ExistInsmall358": 1,
                "ExistIn552": "Ⅱ类",
                "CLNSIG": "Pathogenic",
            }
        ],
        tables={
            "Cnv": [
                {"Gene": "ERBB2", "Status": "扩增"},
                {"Gene": "EGFR", "Status": "扩增"},
            ],
            "Fusion": [
                {"Gene1": "BICC1", "Gene2": "FGFR2"},
            ],
        },
    )
    report_data = enhance_report_data(
        ReportData(),
        excel_data,
        base_path=str(ROOT),
    )

    assert report_data.get_field("nccn_ERBB2_MUT") == ("c.2324_2325ins12，p.Y772_A775dup")
    assert report_data.get_field("nccn_ERBB2_AMP") == "CNV:扩增"
    assert report_data.get_field("nccn_FGFR123_MUT") == "未检出"
    assert report_data.get_field("nccn_FGFR123_FUSION") == ("FGFR2：融合:BICC1-FGFR2")
    assert report_data.get_field("imm_hyper_EGFR_AMP") == "CNV:扩增"


def test_nccn_result_rows_are_driven_by_guideline_rules(tmp_path):
    excel_data = _excel(
        tmp_path,
        variations=[
            {
                "Gene_Symbol": "EGFR",
                "Transcript": "NM_005228.5",
                "Chr": "7",
                "ExIn_ID": "EX21",
                "cHGVS": "c.2573T>G",
                "pHGVS_S": "p.L858R",
                "ExistIn552": "Ⅱ类",
            }
        ],
    )
    report_data = ReportData()
    panel_config = PanelConfig(
        nccn_result_rows=[
            {
                "key": "CUSTOM_EGFR_EX21",
                "genes": ["EGFR"],
                "match": "外显子21",
            }
        ]
    )

    _build_nccn_and_immune_fields(
        report_data,
        [
            {
                "gene": "EGFR",
                "cHGVS": "c.2573T>G",
                "pHGVS": "p.L858R",
                "exon": "EX21",
                "gene_class": "Ⅱ类",
            }
        ],
        excel_data,
        panel_config=panel_config,
    )

    assert report_data.get_field("nccn_CUSTOM_EGFR_EX21") == "c.2573T>G，p.L858R"
    assert report_data.get_field("nccn_EGFR_EX21") is None
    assert report_data.get_table("nccn_results") == [
        {
            "key": "CUSTOM_EGFR_EX21",
            "gene": "EGFR",
            "genes": "EGFR",
            "content": "外显子21",
            "match": "外显子21",
            "result": "c.2573T>G，p.L858R",
            "interpretation": "",
            "检测基因": "EGFR",
            "检测内容": "外显子21",
            "检测结果": "c.2573T>G，p.L858R",
            "基因": "EGFR",
            "临床解读": "",
        }
    ]


def test_grouped_nccn_and_immune_rows_keep_gene_identity_and_vaf_order(tmp_path):
    report_data = ReportData()
    variants = [
        {
            "gene": "TP53",
            "cHGVS": "c.499C>T",
            "pHGVS": "p.Q167*",
            "frequency": "1.16",
            "gene_class": "Ⅱ类",
        },
        {
            "gene": "KRAS",
            "cHGVS": "c.34G>T",
            "pHGVS": "p.G12C",
            "frequency": "13.58",
            "gene_class": "Ⅰ类",
        },
        {
            "gene": "ATR",
            "cHGVS": "c.1291delA",
            "pHGVS": "p.R431Gfs*8",
            "frequency": "0.98",
            "gene_class": "Ⅱ类",
        },
        {
            "gene": "BRCA1",
            "cHGVS": "c.505C>T",
            "pHGVS": "p.Q169*",
            "frequency": "1.10",
            "gene_class": "Ⅱ类",
        },
    ]
    panel_config = PanelConfig(
        nccn_result_rows=[{"key": "BRCA12_MUT", "genes": ["BRCA1", "BRCA2"], "match": "突变"}],
        immune_positive_rows=[
            {"key": "KRAS_TP53", "genes": ["TP53", "KRAS"], "mode": "co_mutation"},
            {"key": "DDR", "genes": ["ATR", "BRCA1"], "mode": "gene_group"},
        ],
    )

    _build_nccn_and_immune_fields(
        report_data,
        variants,
        _excel(tmp_path),
        panel_config=panel_config,
    )

    assert report_data.get_field("nccn_BRCA12_MUT") == "BRCA1：c.505C>T，p.Q169*"
    assert report_data.get_field("imm_pos_KRAS_TP53").splitlines() == [
        "KRAS：c.34G>T，p.G12C",
        "TP53：c.499C>T，p.Q167*",
    ]
    assert report_data.get_field("imm_pos_DDR").splitlines() == [
        "BRCA1：c.505C>T，p.Q169*",
        "ATR：c.1291delA，p.R431Gfs*8",
    ]


def test_crc_guideline_rule_contains_active_nccn_rows():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    rule = PanelRuleEngine.from_panel_package(package).get("guideline_tables")
    rows = rule["guideline_tables"]["nccn_results"]["rows"]

    assert rule["version"] == "0.2.0"
    assert rule["status"] == "active"
    assert len(rows) >= 30
    assert any(row["key"] == "KRAS_EX2" for row in rows)
    assert any(row["key"] == "FGFR123_FUSION" for row in rows)


def test_crc_drug_rule_contains_active_approved_rows():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    engine = PanelRuleEngine.from_panel_package(package)
    rule = engine.get("drugs")
    rows = rule["approved_drug_rows"]

    assert rule["version"] == "0.7.1"
    assert rule["status"] == "active"
    assert len(rows) == 7
    assert "瑞戈非尼" in rows[0]["drug"]
    assert rows[0]["enabled"] is True
    assert rule["drug_rules"]["approved_drug_rows_source"] == ("drugs.yaml:approved_drug_rows")
    assert rule["drug_rules"]["approved_drug_rows_display_mode"] == ("exclude_if_listed_in_part2")
    transition = rule["drug_rules"]["approved_drug_rows_contract_transition"]
    assert transition == {
        "proposed_contract_id": "crc358_dynamic_7_drug_differencing_20260720",
        "prior_contract_id": "crc358_reviewed_historical_table_contract",
        "transition_status": "report_group_approved_supersession",
        "approved_at": "2026-07-22",
        "approval_receipt_id": ("report_group_crc358_domain_presentation_20260724"),
        "source_ref": {
            "type": "report_group_feedback",
            "id": "feedback_20260720_crc_approved_drug_dynamic_display",
        },
    }
    assert rows[0]["secondary_review_status"] == "report_group_approved"
    assert rows[0]["secondary_review_receipt_id"] == (
        "report_group_crc358_domain_presentation_20260724"
    )
    assert "crc_approved_drugs" not in engine.get("panel_rules")


def test_load_panel_config_prefers_drugs_yaml_approved_rows():
    panel_config = load_panel_config(base_path=str(ROOT), panel_id="crc_358_msi")

    assert len(panel_config.approved_drug_rows) == 7
    assert "瑞戈非尼" in panel_config.approved_drug_rows[0]["drug"]
    assert panel_config.approved_drug_rows_display_mode == ("exclude_if_listed_in_part2")


def test_crc_biomarker_rule_contains_active_immune_tables():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    engine = PanelRuleEngine.from_panel_package(package)
    rule = engine.get("biomarkers")
    tables = rule["biomarkers"]["immune_gene_tables"]

    assert rule["version"] == "0.3.0"
    assert rule["status"] == "active"
    assert len(tables["positive"]["genes"]) == 54
    assert len(tables["negative"]["genes"]) == 12
    assert len(tables["hyperprogression"]["genes"]) == 8
    assert {row["key"] for row in tables["positive"]["rows"]} >= {
        "KRAS_TP53",
        "DDR",
    }
    assert {row["key"] for row in tables["negative"]["rows"]} >= {
        "EGFR_L858R",
        "KRAS_STK11",
    }
    positive_modes = {row["key"]: row["mode"] for row in tables["positive"]["rows"]}
    negative_modes = {row["key"]: row["mode"] for row in tables["negative"]["rows"]}
    assert positive_modes["PDCD1LG2"] == "non_sequence_biomarker"
    assert negative_modes["B2M"] == "confirmed_functional_loss"
    assert negative_modes["IFNGR12"] == "confirmed_functional_loss"
    assert {row["key"] for row in tables["hyperprogression"]["rows"]} >= {"EGFR_AMP"}
    assert "immune_positive_genes" not in engine.get("panel_rules")
    assert "immune_negative_genes" not in engine.get("panel_rules")
    assert "immune_hyperprogression_genes" not in engine.get("panel_rules")


def test_load_panel_config_prefers_biomarkers_yaml_immune_tables():
    panel_config = load_panel_config(base_path=str(ROOT), panel_id="crc_358_msi")

    assert "TP53" in panel_config.immune_positive_genes
    assert "PTEN" in panel_config.immune_negative_genes
    assert "EGFR" in panel_config.immune_hyperprogression_genes
    assert {row["key"] for row in panel_config.immune_positive_rows} >= {
        "KRAS_TP53",
        "DDR",
    }
    assert {row["key"] for row in panel_config.immune_negative_rows} >= {
        "EGFR_L858R",
        "KRAS_STK11",
    }
    assert {row["key"] for row in panel_config.immune_hyperprogression_rows} >= {"EGFR_AMP"}


def test_load_panel_config_uses_package_identity_without_redundant_panel_id():
    crc358 = load_panel_package("crc_358_msi", project_root=ROOT)
    crc301 = load_panel_package("crc_301_msi", project_root=ROOT)

    crc358_config = load_panel_config(panel_package=crc358)
    crc301_config = load_panel_config(panel_package=crc301)

    assert "FANCD2" in crc358_config.immune_positive_genes
    assert "FANCD2" not in crc301_config.immune_positive_genes


def test_immune_table_rows_are_driven_by_biomarker_rules(tmp_path):
    report_data = ReportData()
    panel_config = PanelConfig(
        nccn_result_rows=[],
        immune_positive_rows=[
            {"key": "CUSTOM_ATM", "genes": ["ATM"], "mode": "direct"},
        ],
        immune_negative_rows=[],
        immune_hyperprogression_rows=[],
    )

    _build_nccn_and_immune_fields(
        report_data,
        [
            {
                "gene": "ATM",
                "cHGVS": "c.6874C>T",
                "pHGVS": "p.Q2292*",
                "gene_class": "Ⅱ类",
            }
        ],
        _excel(tmp_path),
        panel_config=panel_config,
    )

    assert report_data.get_field("imm_pos_CUSTOM_ATM") == "c.6874C>T，p.Q2292*"
    assert report_data.get_table("immune_positive_results") == [
        {
            "key": "CUSTOM_ATM",
            "gene": "ATM",
            "genes": "ATM",
            "content": "",
            "match": "",
            "result": "c.6874C>T，p.Q2292*",
            "interpretation": "检出有害变异时可能疗效较好。",
            "检测基因": "ATM",
            "检测内容": "",
            "检测结果": "c.6874C>T，p.Q2292*",
            "基因": "ATM",
            "临床解读": "检出有害变异时可能疗效较好。",
        }
    ]


def test_field_mapper_adds_legacy_fusion_aliases(tmp_path):
    excel_data = _excel(
        tmp_path,
        tables={
            "Fusion": [
                {
                    "Gene1": "EML4",
                    "Chr1": "chr2",
                    "Pos1": 42491832,
                    "Gene2": "ALK",
                    "Chr2": "chr2",
                    "Pos2": 29446394,
                    "Sv_type": "fusion",
                }
            ]
        },
    )

    report_data = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR").map(excel_data)

    rows = report_data.get_table("fusion")
    assert rows
    assert "#Est_Type" in rows[0]
    assert rows[0]["#Est_Type"] == "fusion"
    assert "Freq1" in rows[0]


def test_missing_clnsig_is_not_defaulted_to_pathogenic(tmp_path):
    excel_data = _excel(
        tmp_path,
        variations=[
            {
                "Gene_Symbol": "TP53",
                "Transcript": "NM_000546.6",
                "Chr": "17",
                "ExIn_ID": "EX8",
                "cHGVS": "c.817C>T",
                "pHGVS_S": "p.R273C",
                "Freq(%)": 33.0,
                "ExistInsmall358": 1,
                "ExistIn552": 1,
                "CLNSIG": "*",
            }
        ],
    )
    variants = build_variants_for_template(
        excel_data,
        filter_class_i_ii_only=False,
        important_genes_only=False,
        panel_config=load_panel_config(base_path=str(ROOT)),
    )

    assert variants
    assert variants[0]["clinical_significance"] == "临床意义未明"


def test_tmb_missing_is_explicit_not_low(tmp_path):
    summary = build_tmb_summary(_excel(tmp_path))
    assert summary["tmb_value"] == "未检测"
    assert summary["tmb_status"] == "未检测"
    assert summary["tmb_level_cn"] == "未检测"


def test_tmb_invalid_is_explicit_format_error(tmp_path):
    summary = build_tmb_summary(_excel(tmp_path, single_values={"TMB": "abc"}))
    assert summary["tmb_value"] == "未检测（格式错误）"
    assert summary["tmb_status"] == "未检测"
    assert summary["tmb_summary"] == "未检测（格式错误）"


def test_field_mapper_invalid_tmb_uses_same_format_error(tmp_path):
    report_data = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR").map(
        _excel(tmp_path, single_values={"TMB": "abc"})
    )

    assert report_data.get_field("tmb_value") == "未检测（格式错误）"
    assert report_data.get_field("tmb_status") == "未检测"
    assert report_data.get_field("tmb_summary") == "未检测（格式错误）"


def test_field_mapper_valid_tmb_overrides_default_status(tmp_path):
    report_data = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR").map(
        _excel(tmp_path, single_values={"TMB": 7.10499382716049, "MSI状态": "MSI-H"})
    )

    assert report_data.get_field("tmb_value") == "7.1"
    assert report_data.get_field("tmb") == "7.1 mutations/Mb"
    assert report_data.get_field("tmb_status") == "L"
    assert report_data.get_field("tmb_level_cn") == "低"
    assert "TMB-L" in report_data.get_field("tmb_summary")
    assert report_data.get_field("msi_status") == "MSI-H"
    assert "TMB-H的肿瘤" in report_data.get_field("immuno_tips")
    assert "#帕博利珠单抗" in report_data.get_field("immuno_tips")
    assert "TMB水平较低" in report_data.get_field("tmb_detail_sentence")
    assert "MSI-H" in report_data.get_field("msi_detail_sentence")
    assert "MSI-H" in report_data.get_field("msi_tips")


def test_report_summary_extracts_biomarkers_variants_and_drugs():
    report_data = ReportData(
        context={
            "patient_name": "测试患者",
            "sample_id": "CASE001",
            "clinical_diagnosis": "结直肠癌",
            "project_name": "结直肠癌358基因+MSI",
            "tmb_summary": "12.9 mutations/Mb，TMB-H",
            "tmb_value": "12.9",
            "tmb_status": "H",
            "msi_summary": "微卫星稳定型，MSS",
            "msi_status": "MSS",
            "total_variants_count": 2,
            "drug_related_count": 1,
            "variants_2_1": [
                {
                    "gene": "KRAS",
                    "locus": "c.34G>A，p.G12S",
                    "gene_class": "Ⅱ类",
                    "af_pct": "18.2%",
                    "benefit_drugs": "--",
                    "caution_drugs": "西妥昔单抗（A）",
                },
                {
                    "gene": "APC",
                    "locus": "c.3927del",
                    "gene_class": "Ⅲ类",
                    "af_pct": "10.0%",
                    "benefit_drugs": "--",
                    "caution_drugs": "--",
                },
            ],
            "targeted_drug_tips": [
                {
                    "gene": "KRAS",
                    "variant_site": "c.34G>A，p.G12S",
                    "benefit_drugs": "--",
                    "caution_drugs": "西妥昔单抗（A）",
                },
                {
                    "gene": "APC",
                    "variant_site": "c.3927del",
                    "benefit_drugs": "--",
                    "caution_drugs": "--",
                },
            ],
            "chemotherapy": [{"Drug": "瑞戈非尼", "Gene": "VEGFR"}],
        }
    )

    summary = build_report_summary(
        report_data=report_data,
        project_type="crc_358_msi",
        project_name="结直肠癌358基因+MSI",
        generation_id="CASE001",
        output_file="/tmp/fake.docx",
        qa_report={"status": "PASS", "issues": []},
    )

    assert summary["patient"]["sample_id"] == "CASE001"
    assert summary["biomarkers"]["tmb"]["status"] == "H"
    assert summary["biomarkers"]["msi"]["status"] == "MSS"
    assert summary["variants"]["total"] == 2
    assert summary["variants"]["drug_related"] == 1
    assert summary["variants"]["by_class"] == {"Ⅱ类": 1, "Ⅲ类": 1}
    assert summary["variants"]["key_rows"][0]["gene"] == "KRAS"
    assert summary["drugs"]["targeted_count"] == 1
    assert summary["drugs"]["displayed_variant_count"] == 2
    assert summary["drugs"]["targeted_rows"][1]["benefit_drugs"] == "--"
    assert summary["drugs"]["targeted_rows"][1]["caution_drugs"] == "--"
    assert summary["qa"]["status"] == "PASS"


def test_report_summary_flags_draft_panel_status_without_cancer_name():
    summary = build_report_summary(
        report_data=ReportData(),
        project_type="lung_329_pdl1",
        project_name="肺癌329基因+PD-L1",
        panel_status="draft",
        template_status="draft",
        qa_report={"status": "PASS", "issues": []},
    )

    assert summary["panel"] == {"status": "draft", "template_status": "draft"}
    assert any("draft" in item and "勿直接交付" in item for item in summary["manual_review"])


def test_report_summary_flags_pilot_panel_status():
    summary = build_report_summary(
        report_data=ReportData(),
        project_type="lung_methylation",
        project_name="肺癌甲基化",
        panel_status="pilot",
        template_status="pilot",
        qa_report={"status": "PASS", "issues": []},
    )

    assert any("pilot" in item and "试运行" in item for item in summary["manual_review"])


def test_report_summary_active_panel_status_has_no_draft_guard():
    summary = build_report_summary(
        report_data=ReportData(),
        project_type="crc_358_msi",
        project_name="结直肠癌358基因+MSI",
        panel_status="active",
        template_status="active",
        qa_report={"status": "PASS", "issues": []},
    )

    assert summary["panel"] == {"status": "active", "template_status": "active"}
    assert summary["manual_review"] == []


def test_report_generator_reads_status_from_panel_package():
    package = load_panel_package(ROOT / "panels" / "lung_329_pdl1")
    template_file = package.resolve_template_file()

    assert ReportGenerator._panel_status(package) == "pilot"
    assert ReportGenerator._template_status(package, str(template_file)) == "pilot"


def test_write_report_summary_uses_docx_sidecar(tmp_path):
    output_file = tmp_path / "case.docx"
    output_file.write_bytes(b"PK\x03\x04")

    sidecar = write_report_summary({"schema_version": "1.0"}, str(output_file))

    assert Path(sidecar) == tmp_path / "case.summary.json"
    assert json.loads(Path(sidecar).read_text(encoding="utf-8"))["schema_version"] == "1.0"


def test_field_mapper_dynamic_tmb_msi_narratives_match_mss_low_tmb(tmp_path):
    report_data = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR").map(
        _excel(tmp_path, single_values={"TMB": 7.74681481481482, "MSI状态": "MSS"})
    )

    assert report_data.get_field("tmb_value") == "7.7"
    assert report_data.get_field("tmb_status") == "L"
    assert "7.7 mutations/Mb" in report_data.get_field("tmb_detail_sentence")
    assert "TMB水平较低" in report_data.get_field("tmb_detail_sentence")
    assert "2020年6月，FDA批准帕博利珠单抗" in report_data.get_field("tmb_detail_interpretation")
    assert "帕博利珠单抗、纳武利尤单抗" in report_data.get_field("tmb_drug_note")
    assert "微卫星稳定（MSS）型" in report_data.get_field("msi_detail_sentence")
    assert "林奇综合征" in report_data.get_field("msi_detail_interpretation")
    assert "MSI-H的实体瘤通常具有免疫原性" in report_data.get_field("msi_tips")


def test_mixed_reviewed_class_labels_control_counts_and_drug_rows(tmp_path):
    variations = [
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "TP53",
            "Transcript": "NM_000546.6",
            "Chr": "chr17",
            "ExIn_ID": "EX8",
            "cHGVS": "c.844C>T",
            "pHGVS_S": "p.R282W",
            "Function": "Missense",
            "Freq(%)": 67.29,
            "CLNSIG": "Pathogenic/Likely_pathogenic",
        },
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "KRAS",
            "Transcript": "NM_004985.5",
            "Chr": "chr12",
            "ExIn_ID": "EX2",
            "cHGVS": "c.34G>A",
            "pHGVS_S": "p.G12S",
            "Function": "Missense",
            "Freq(%)": 46.29,
            "CLNSIG": "Pathogenic",
        },
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "APC",
            "Transcript": "NM_000038.6",
            "Chr": "chr5",
            "ExIn_ID": "EX16E",
            "cHGVS": "c.4348C>T",
            "pHGVS_S": "p.R1450*",
            "Function": "Nonsense",
            "Freq(%)": 41.12,
            "CLNSIG": "Pathogenic",
        },
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "APC",
            "Transcript": "NM_000038.6",
            "Chr": "chr5",
            "ExIn_ID": "EX16E",
            "cHGVS": "c.2387_2388del",
            "pHGVS_S": "p.Y796Wfs*2",
            "Function": "Frameshift",
            "Freq(%)": 37.52,
            "CLNSIG": "Pathogenic",
        },
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "SETD2",
            "Transcript": "NM_014159.7",
            "Chr": "chr3",
            "ExIn_ID": "EX8",
            "cHGVS": "c.4930G>T",
            "pHGVS_S": "p.G1644*",
            "Function": "Nonsense",
            "Freq(%)": 22.15,
            "CLNSIG": "*",
        },
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "EPHA2",
            "Transcript": "NM_004431.5",
            "Chr": "chr1",
            "ExIn_ID": "EX2",
            "cHGVS": "c.153+2T>C",
            "pHGVS_S": "*",
            "Function": "Splice-5",
            "Freq(%)": 3.33,
            "CLNSIG": "*",
        },
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "FANCI",
            "Transcript": "NM_001113378.2",
            "Chr": "chr15",
            "ExIn_ID": "EX26",
            "cHGVS": "c.2879G>A",
            "pHGVS_S": "p.R960Q",
            "Function": "Missense",
            "Freq(%)": 2.18,
            "CLNSIG": "Uncertain_significance",
        },
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "ATM",
            "Transcript": "NM_000051.4",
            "Chr": "chr11",
            "ExIn_ID": "EX47",
            "cHGVS": "c.6874C>T",
            "pHGVS_S": "p.Q2292*",
            "Function": "Nonsense",
            "Freq(%)": 0.53,
            "CLNSIG": "Pathogenic",
        },
        {
            "ExistIn552": 1,
            "ExistInsmall358": 1,
            "Gene_Symbol": "FBXW7",
            "Transcript": "NM_001349798.2",
            "Chr": "chr4",
            "ExIn_ID": "EX4",
            "cHGVS": "c.349_351del",
            "pHGVS_S": "p.E117del",
            "Function": "CDS-indel",
            "Freq(%)": 1.31,
            "CLNSIG": "Uncertain_significance",
        },
    ]

    report_data = enhance_report_data(
        ReportData(),
        _excel(
            tmp_path,
            single_values={"TMB": 6.463, "MSI状态": "MSS", "样本类型": "组织"},
            variations=variations,
        ),
        field_mapper=_FakeDrugLookup(),
        base_path=str(ROOT),
    )

    assert report_data.get_field("total_variants_count") == 8
    assert report_data.get_field("drug_related_count") == 4
    # 靶向药物子集为 4 个；免疫相关变异均与该子集重合，
    # 因此第三部分“靶向/免疫”去重并集仍应为 4，不能相加为 7。
    assert report_data.get_field("targeted_or_immune_related_count") == 4
    assert [row["gene"] for row in report_data.get_table("variants")] == [
        "TP53",
        "KRAS",
        "SETD2",
        "ATM",
    ]
    assert "FBXW7" not in {row.get("gene") for row in report_data.get_table("targeted_drug_tips")}
    assert report_data.get_field("immune_positive_count") == 3
    assert "KRAS：c.34G>A，p.G12S" in report_data.get_field("immune_positive_result")
    assert report_data.get_field("immune_negative_result") == "未检出"
    assert report_data.get_field("imm_neg_KRAS_STK11") == "未检出有害变异"
    assert "TP53：c.844C>T，p.R282W" in report_data.get_field("imm_pos_KRAS_TP53")
    assert "ATM：c.6874C>T，p.Q2292*" in report_data.get_field("imm_pos_DDR")
    assert "西妥昔单抗[爱必妥]" in report_data.get_field("targeted_drug_brand_summary")
    assert "FBXW7" not in report_data.get_field("targeted_drug_brand_summary")


def test_drug_related_count_matches_variants_2_1_table(tmp_path):
    """The footnote drug count (drug_related_count) must be counted from the
    same variants_2_1 table the reader sees in section 2.1 — not a separate
    variant pipeline. Regression: the footnote showed 4 while the 2.1 table
    listed 6 drug-bearing rows (audit "变异计数一致" failure).
    """
    report_data = ReportData()
    # FieldMapper populates variants_2_1 before the enhancers run. Pre-set it
    # with 2 drug-bearing rows (+ 1 plain row + 1 未见突变 baseline). The
    # variations below would otherwise be counted as 3 drug-bearing by the old
    # `variants` pipeline (_FakeDrugLookup gives every gene a drug), so a result
    # of 2 proves the count comes from variants_2_1.
    report_data.set_table(
        "variants_2_1",
        [
            {
                "gene": "TP53",
                "locus": "c.1A>T",
                "af_pct": "30",
                "benefit_drugs": "DrugX（C）",
                "caution_drugs": "--",
            },
            {
                "gene": "ATR",
                "locus": "c.2A>T",
                "af_pct": "5",
                "benefit_drugs": "--",
                "caution_drugs": "DrugY（A）",
            },
            {
                "gene": "APC",
                "locus": "c.3A>T",
                "af_pct": "20",
                "benefit_drugs": "--",
                "caution_drugs": "--",
            },
            {
                "gene": "BRAF",
                "locus": "未见突变",
                "af_pct": "--",
                "benefit_drugs": "--",
                "caution_drugs": "--",
            },
        ],
    )
    variations = [
        {
            "ExistIn552": "Ⅰ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "KRAS",
            "Transcript": "NM_a",
            "Chr": "chr12",
            "ExIn_ID": "EX2",
            "cHGVS": "c.34G>T",
            "pHGVS_S": "p.G12C",
            "Function": "Missense",
            "Freq(%)": 40.0,
            "CLNSIG": "Pathogenic",
        },
        {
            "ExistIn552": "Ⅰ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "PIK3CA",
            "Transcript": "NM_b",
            "Chr": "chr3",
            "ExIn_ID": "EX10",
            "cHGVS": "c.1624G>A",
            "pHGVS_S": "p.E542K",
            "Function": "Missense",
            "Freq(%)": 15.0,
            "CLNSIG": "Pathogenic",
        },
        {
            "ExistIn552": "Ⅰ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "ATM",
            "Transcript": "NM_c",
            "Chr": "chr11",
            "ExIn_ID": "EX47",
            "cHGVS": "c.6874C>T",
            "pHGVS_S": "p.Q2292*",
            "Function": "Nonsense",
            "Freq(%)": 10.0,
            "CLNSIG": "Pathogenic",
        },
    ]
    result = enhance_report_data(
        report_data,
        _excel(
            tmp_path,
            single_values={"TMB": 6.0, "MSI状态": "MSS", "样本类型": "组织"},
            variations=variations,
        ),
        field_mapper=_FakeDrugLookup(),
        base_path=str(ROOT),
    )
    assert result.get_field("drug_related_count") == 2


def test_targeted_or_immune_count_includes_immune_only_variant_once(tmp_path):
    """Part 3 count is a distinct-variant union, not the targeted subset.

    The targeted table contains KRAS and TP53.  The immune pipeline contains
    the same KRAS variant plus immune-only DNMT3A, so the union must be 3.
    """
    report_data = ReportData()
    report_data.set_table(
        "variants_2_1",
        [
            {
                "gene": "KRAS",
                "locus": "c.34G>T, p.G12C",
                "benefit_drugs": "索托拉西布（A）",
                "caution_drugs": "--",
            },
            {
                "gene": "TP53",
                "locus": "c.499C>T, p.Q167*",
                "benefit_drugs": "AZD1775（C）",
                "caution_drugs": "--",
            },
        ],
    )
    variations = [
        {
            "ExistIn552": "Ⅰ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "KRAS",
            "Transcript": "NM_004985.5",
            "Chr": "chr12",
            "ExIn_ID": "EX2",
            "cHGVS": "c.34G>T",
            "pHGVS_S": "p.G12C",
            "Function": "Missense",
            "Freq(%)": 13.58,
            "CLNSIG": "Pathogenic",
        },
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "DNMT3A",
            "Transcript": "NM_022552.5",
            "Chr": "chr2",
            "ExIn_ID": "EX12",
            "cHGVS": "c.1367del",
            "pHGVS_S": "p.K456Sfs*195",
            "Function": "Frameshift",
            "Freq(%)": 1.18,
            "CLNSIG": "Pathogenic",
        },
    ]

    result = enhance_report_data(
        report_data,
        _excel(
            tmp_path,
            single_values={"TMB": 6.0, "MSI状态": "MSS", "样本类型": "组织"},
            variations=variations,
        ),
        field_mapper=_FakeDrugLookup(),
        base_path=str(ROOT),
    )

    assert result.get_field("drug_related_count") == 2
    assert result.get_field("targeted_or_immune_related_count") == 3


def test_variants_2_1_keeps_all_detected_panel_variants(tmp_path):
    variations = [
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "APC",
            "Transcript": "NM_000038.6",
            "Chr": "chr5",
            "ExIn_ID": "EX16E",
            "cHGVS": "c.2387_2388del",
            "pHGVS_S": "p.Y796Wfs*2",
            "Function": "Frameshift",
            "Freq(%)": 37.52,
        },
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "EPHA2",
            "Transcript": "NM_004431.5",
            "Chr": "chr1",
            "ExIn_ID": "EX2",
            "cHGVS": "c.153+2T>C",
            "pHGVS_S": "*",
            "Function": "Splice-5",
            "Freq(%)": 3.33,
        },
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "FANCI",
            "Transcript": "NM_001113378.2",
            "Chr": "chr15",
            "ExIn_ID": "EX26",
            "cHGVS": "c.2879G>A",
            "pHGVS_S": "p.R960Q",
            "Function": "Missense",
            "Freq(%)": 2.18,
        },
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 0,
            "Gene_Symbol": "NOTPANEL",
            "Transcript": "NM_TEST",
            "Chr": "chr1",
            "ExIn_ID": "EX1",
            "cHGVS": "c.1A>T",
            "pHGVS_S": "p.K1M",
            "Function": "Missense",
            "Freq(%)": 1.0,
        },
    ]
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")

    rows = mapper._build_variants_2_1(
        _excel(tmp_path, variations=variations),
        ReportData(),
    )

    genes = [row["gene"] for row in rows]
    assert "EPHA2" in genes
    assert "FANCI" in genes
    assert "NOTPANEL" not in genes
    apc = next(row for row in rows if row["gene"] == "APC")
    assert apc["var_type_cn"] == "缺失突变"
    epha2 = next(row for row in rows if row["gene"] == "EPHA2")
    assert epha2["locus"] == "c.153+2T>C"
    assert epha2["exon"] == "内含子2"
    assert epha2["benefit_drugs"] == "--"
    fanci = next(row for row in rows if row["gene"] == "FANCI")
    assert fanci["var_type_cn"] == "点突变"


def test_flt3_reviewed_override_brings_out_curated_drugs(tmp_path):
    """FLT3 c.2537G>A (Ⅱ类) has only gene-level KB entries (excluded by
    require_position_match) and the KB drugs are AML-specific. A reviewed
    per-variant override pins the clinically-curated CRC TKIs so the variant is
    drug-related (matches the reviewed final report and keeps the footnote count
    in sync with the 2.1 table).
    """
    variations = [
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "FLT3",
            "Transcript": "NM_004119.3",
            "Chr": "chr13",
            "ExIn_ID": "EX20",
            "cHGVS": "c.2537G>A",
            "pHGVS_S": "p.G846D",
            "Function": "Missense",
            "Freq(%)": 1.02,
        },
    ]
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")
    rows = mapper._build_variants_2_1(_excel(tmp_path, variations=variations), ReportData())
    flt3 = next(row for row in rows if row["gene"] == "FLT3")
    for drug in ("瑞戈非尼", "索拉非尼", "舒尼替尼"):
        assert drug in flt3["benefit_drugs"], flt3["benefit_drugs"]
    assert flt3["caution_drugs"] == "--"


def test_immune_table_conditional_notes_and_numbering():
    """Immune-table footnotes are conditional and renumber correctly:
      * TMB-H            → append the FDA TMB-H sentence to note 2.
      * TMB/MSI 结果不一致 → insert the biomarker-independence note (note 3).
    The trailing 药物名称 note becomes "4." when the conflict note is present,
    "3." otherwise. TMB-L or both-high add neither.
    """
    from docx import Document

    def build_doc():
        doc = Document()
        doc.add_paragraph(
            "备注：1. FoundationOne CDx (324 个基因) TMB 研究表明，"
            "对于组织样本，TMB≥10 mut/Mb 为高突变负荷。"
        )
        doc.add_paragraph("#帕博利珠单抗、#纳武利尤单抗均已获FDA和/或NMPA批准用于治疗结直肠癌。")
        doc.add_paragraph(
            "3. 上表涉及的已上市的药物名称及对应的商品名称："
            "帕博利珠单抗[可瑞达]、纳武利尤单抗[欧狄沃]"
        )
        return doc

    renderer = TemplateRenderer(log_level="ERROR")

    def run(ctx):
        doc = build_doc()
        renderer._apply_immune_table_notes_to_doc(doc, ctx)
        texts = [p.text for p in doc.paragraphs]
        full = "\n".join(texts)
        drug_note_num = next(t[:2] for t in texts if "上表涉及" in t)
        return full, drug_note_num

    # TMB-H + MSS (inconsistent): both notes added, drug note renumbered to 4.
    full, num = run({"tmb_status": "H", "msi_status": "MSS"})
    assert "帕博利珠单抗用于治疗 TMB-H" in full
    assert "免疫治疗生物标志物包括" in full
    assert num == "4."

    # TMB-L + MSS: neither note, drug note stays 3.
    full, num = run({"tmb_status": "L", "msi_status": "MSS"})
    assert "帕博利珠单抗用于治疗 TMB-H" not in full
    assert "免疫治疗生物标志物包括" not in full
    assert num == "3."

    # TMB-H + MSI-H (both high, consistent): TMB-H sentence only, no conflict note.
    full, num = run({"tmb_status": "H", "msi_status": "MSI-H"})
    assert "帕博利珠单抗用于治疗 TMB-H" in full
    assert "免疫治疗生物标志物包括" not in full
    assert num == "3."


def test_variants_2_1_detected_rows_sorted_by_frequency_desc(tmp_path):
    """Detected-variant rows are ordered by 频率 high→low, grouped by gene.

    A gene's variants share merged 基因名/转录本/染色体 cells in the reviewed
    report, so they must stay adjacent — genes are ordered by their highest
    frequency, and variants within a gene by frequency descending.
    """
    variations = [
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "TP53",
            "Transcript": "NM_000546.6",
            "Chr": "chr17",
            "ExIn_ID": "EX8",
            "cHGVS": "c.844C>T",
            "pHGVS_S": "p.R282W",
            "Function": "Missense",
            "Freq(%)": 50.0,
        },
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "APC",
            "Transcript": "NM_000038.6",
            "Chr": "chr5",
            "ExIn_ID": "EX10",
            "cHGVS": "c.994C>T",
            "pHGVS_S": "p.R332*",
            "Function": "Nonsense",
            "Freq(%)": 10.5,
        },
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "KRAS",
            "Transcript": "NM_004985.5",
            "Chr": "chr12",
            "ExIn_ID": "EX2",
            "cHGVS": "c.34G>T",
            "pHGVS_S": "p.G12C",
            "Function": "Missense",
            "Freq(%)": 30.0,
        },
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "APC",
            "Transcript": "NM_000038.6",
            "Chr": "chr5",
            "ExIn_ID": "EX16",
            "cHGVS": "c.4348C>T",
            "pHGVS_S": "p.R1450*",
            "Function": "Nonsense",
            "Freq(%)": 40.0,
        },
    ]
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")
    rows = mapper._build_variants_2_1(_excel(tmp_path, variations=variations), ReportData())

    def af(row):
        try:
            return float(str(row.get("af_pct")).replace("%", "").strip())
        except (TypeError, ValueError):
            return None

    detected = [(row["gene"], af(row)) for row in rows if af(row) is not None]
    # genes by max freq (TP53 50 > APC 40 > KRAS 30); APC's two variants stay
    # together and ordered 40 then 10.5. Non-integer VAF display keeps the
    # reviewed report's two-decimal precision.
    assert detected == [("TP53", 50.0), ("APC", 40.0), ("APC", 10.5), ("KRAS", 30.0)]
    assert (
        next(
            row["af_pct"]
            for row in rows
            if row["gene"] == "APC" and row["locus"].startswith("c.994C>T")
        )
        == "10.50"
    )


def test_targeted_drug_tips_summary_sorted_by_frequency_desc(tmp_path, monkeypatch):
    """The 1.检测结果小结 summary table (targeted_drug_tips) is ordered by
    frequency high→low and grouped by gene — consistent with the 2.1 table.
    """
    variations = [
        {
            "ExistIn552": "Ⅰ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "GENEB",
            "Transcript": "NM_1",
            "Chr": "chr1",
            "ExIn_ID": "EX1",
            "cHGVS": "c.100A>T",
            "pHGVS_S": "p.K34N",
            "Function": "Missense",
            "Freq(%)": 50.0,
        },
        {
            "ExistIn552": "Ⅰ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "GENEA",
            "Transcript": "NM_2",
            "Chr": "chr2",
            "ExIn_ID": "EX1",
            "cHGVS": "c.200G>C",
            "pHGVS_S": "p.G67A",
            "Function": "Missense",
            "Freq(%)": 10.0,
        },
        {
            "ExistIn552": "Ⅰ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "GENEC",
            "Transcript": "NM_3",
            "Chr": "chr3",
            "ExIn_ID": "EX1",
            "cHGVS": "c.300T>A",
            "pHGVS_S": "p.S100T",
            "Function": "Missense",
            "Freq(%)": 30.0,
        },
        {
            "ExistIn552": "Ⅰ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "GENEA",
            "Transcript": "NM_2",
            "Chr": "chr2",
            "ExIn_ID": "EX2",
            "cHGVS": "c.400C>G",
            "pHGVS_S": "p.P133A",
            "Function": "Missense",
            "Freq(%)": 40.0,
        },
    ]
    # This test only verifies row ordering. Simulate an unavailable production KB
    # so the legacy CtDrug fallback can provide simple synthetic drug rows.
    ctdrug = [
        {
            "检测基因": "GENEA",
            "药物": "DrugA",
            "证据等级": "A",
            "用药提示（仅供参考）": "敏感，推荐使用",
        },
        {
            "检测基因": "GENEB",
            "药物": "DrugB",
            "证据等级": "A",
            "用药提示（仅供参考）": "敏感，推荐使用",
        },
        {
            "检测基因": "GENEC",
            "药物": "DrugC",
            "证据等级": "A",
            "用药提示（仅供参考）": "敏感，推荐使用",
        },
    ]
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")
    monkeypatch.setattr(mapper, "_load_targeted_drug_db", lambda: None)
    rows = mapper._build_targeted_drug_tips(
        _excel(tmp_path, variations=variations, tables={"CtDrug": ctdrug}),
        ReportData(),
    )

    genes = [row["gene"] for row in rows]
    # gene max freq: GENEB 50 > GENEA 40 > GENEC 30; GENEA's two sites adjacent
    # and ordered 40 then 10.
    assert genes == ["GENEB", "GENEA", "GENEA", "GENEC"], genes
    # The internal sort-only frequency field must not leak into output rows.
    assert all("af" not in row for row in rows)


def test_crc_targeted_summary_only_displays_drug_matched_variants(tmp_path):
    """The Part 2 drug table excludes detected variants without a qualified
    drug conclusion; full display applies to each matched row's drug list.
    """
    variations = [
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "TP53",
            "cHGVS": "c.646G>A",
            "pHGVS_S": "p.V216M",
            "Freq(%)": 93.79,
        },
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "SOS1",
            "cHGVS": "c.2536G>A",
            "pHGVS_S": "p.E846K",
            "Freq(%)": 57.77,
        },
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "GNAS",
            "cHGVS": "c.1030G>A",
            "pHGVS_S": "p.E344K",
            "Freq(%)": 19.04,
        },
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "KMT2B",
            "cHGVS": "c.1681C>T",
            "pHGVS_S": "p.P561S",
            "Freq(%)": 1.69,
        },
        {
            "ExistIn552": "Ⅲ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "SMARCA4",
            "cHGVS": "c.839C>A",
            "pHGVS_S": "p.P280H",
            "Freq(%)": 1.27,
        },
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "RNF43",
            "cHGVS": "c.350_351delinsA",
            "pHGVS_S": "p.R117Hfs*41",
            "Freq(%)": 0.90,
        },
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "BARD1",
            "cHGVS": "c.590delA",
            "pHGVS_S": "p.K197Rfs*15",
            "Freq(%)": 0.70,
        },
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "ATM",
            "cHGVS": "c.3673C>T",
            "pHGVS_S": "p.Q1225*",
            "Freq(%)": 0.53,
        },
        # A classified row outside the selected panel must not leak into the summary.
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 0,
            "Gene_Symbol": "ERBB2",
            "cHGVS": "c.1979G>A",
            "pHGVS_S": "p.G660D",
            "Freq(%)": 99.0,
        },
    ]
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    context = load_targeted_drug_rule_context(package)
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")
    report_data = ReportData(context={"cancer_type": "结直肠癌"})

    rows = mapper._build_targeted_drug_tips(
        _excel(tmp_path, variations=variations),
        report_data,
        targeted_drug_rules=context,
    )

    assert rows == []
    assert mapper._lookup_reviewed_variant_override_drugs(
        "ATM",
        "c.3673C>T",
        "p.Q1225*",
        variant_level="Ⅱ类",
        targeted_drug_rules=context,
    ) == ("--", "--")


def test_immune_positive_summary_includes_class_i_ii_without_clnsig(tmp_path):
    """The immune-summary positive list must include Ⅰ/Ⅱ class variants even
    when CLNSIG is blank (e.g. frameshift PMS2/ATR), matching the 3.3 detail
    table. Regression: a CLNSIG pathogenicity whitelist used to drop blank-CLNSIG
    variants, so the front summary showed fewer variants than the 3.3 detail.
    """
    from reportgen.core.template_bridge_358 import build_immune_variants

    variations = [
        # PMS2 frameshift, Ⅰ类, BLANK CLNSIG — was wrongly dropped from summary.
        {
            "ExistIn552": "Ⅰ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "PMS2",
            "Transcript": "NM_000535",
            "Chr": "chr7",
            "ExIn_ID": "EX11",
            "cHGVS": "c.1273delT",
            "pHGVS_S": "p.S425Lfs*23",
            "Function": "Frameshift",
            "Freq(%)": 30.0,
            "CLNSIG": "",
        },
        # ATR frameshift, Ⅰ类, BLANK CLNSIG — was wrongly dropped from summary.
        {
            "ExistIn552": "Ⅰ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "ATR",
            "Transcript": "NM_001184",
            "Chr": "chr3",
            "ExIn_ID": "EX9",
            "cHGVS": "c.1291delA",
            "pHGVS_S": "p.R431Gfs*8",
            "Function": "Frameshift",
            "Freq(%)": 5.0,
            "CLNSIG": "",
        },
        # KRAS missense, Ⅰ类, pathogenic — always included.
        {
            "ExistIn552": "Ⅰ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "KRAS",
            "Transcript": "NM_004985",
            "Chr": "chr12",
            "ExIn_ID": "EX2",
            "cHGVS": "c.34G>T",
            "pHGVS_S": "p.G12C",
            "Function": "Missense",
            "Freq(%)": 13.0,
            "CLNSIG": "Pathogenic",
        },
    ]

    result = build_immune_variants(
        _excel(tmp_path, variations=variations), filter_class_i_ii_only=True
    )
    positive_genes = {v["gene"].upper() for v in result["positive"]}
    # Blank-CLNSIG Ⅰ/Ⅱ class variants are included alongside the pathogenic one.
    assert {"PMS2", "ATR", "KRAS"} <= positive_genes


def test_field_mapper_uses_panel_biomarker_gene_sets_without_cross_panel_leakage(
    tmp_path,
):
    """Panel-local biomarkers.yaml is authoritative before the enhancer runs.

    CRC358 intentionally includes FANCD2 after report-group review, while CRC301
    does not.  A shared FieldMapper must therefore produce different summaries
    for the same source rows and must not cache one panel's rule into the next
    request.  ERBB2 is present in the legacy public XLSX but absent from both
    panel-local positive sets, so it also guards against legacy fallback leakage.
    """
    variations = [
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "ExistInsmall301": 1,
            "Gene_Symbol": "FANCD2",
            "cHGVS": "c.1630C>T",
            "pHGVS_S": "p.Q544*",
        },
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "ExistInsmall301": 1,
            "Gene_Symbol": "RAD50",
            "cHGVS": "c.1093C>T",
            "pHGVS_S": "p.R365*",
        },
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "ExistInsmall301": 1,
            "Gene_Symbol": "ERBB2",
            "cHGVS": "c.1979G>A",
            "pHGVS_S": "p.G660D",
        },
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 0,
            "ExistInsmall301": 0,
            "Gene_Symbol": "ATM",
            "cHGVS": "c.1A>T",
            "pHGVS_S": "p.K1*",
        },
    ]
    excel_data = _excel(tmp_path, variations=variations)
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")
    crc358 = load_panel_package("crc_358_msi", project_root=ROOT)
    crc301 = load_panel_package("crc_301_msi", project_root=ROOT)

    crc358_result = mapper.map(excel_data, panel_package=crc358)
    crc301_result = mapper.map(excel_data, panel_package=crc301)

    assert crc358_result.get_field("immuno_positive_genes").splitlines() == [
        "检出（2个）",
        "FANCD2：c.1630C>T，p.Q544*",
        "RAD50：c.1093C>T，p.R365*",
    ]
    assert crc301_result.get_field("immuno_positive_genes").splitlines() == [
        "检出（1个）",
        "RAD50：c.1093C>T，p.R365*",
    ]


def test_declared_panel_biomarker_rule_does_not_silently_fallback(tmp_path):
    rule_path = tmp_path / "biomarkers.yaml"
    rule_path.write_text(
        yaml.safe_dump(
            {
                "schema_version": "1.0",
                "panel_id": "broken_panel",
                "rule_id": "biomarkers",
                "version": "0.0.1",
                "status": "active",
                "biomarkers": {},
            },
            allow_unicode=True,
        ),
        encoding="utf-8",
    )
    package = SimpleNamespace(
        panel_id="broken_panel",
        rules={"biomarkers": "rules/biomarkers.yaml"},
        resolve_rule_file=lambda _rule_name: rule_path,
    )
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")

    with pytest.raises(ValueError, match="no immune_gene_tables"):
        mapper._load_immune_gene_sets(panel_package=package)


def test_fancd2_panel_rule_survives_mapper_enhancer_and_word_render(tmp_path):
    """Regression chain for feedback #12: rule -> mapper -> enhancer -> DOCX."""
    variations = [
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "FANCD2",
            "Transcript": "NM_006567.4",
            "Chr": "chr3",
            "ExIn_ID": "EX18",
            "cHGVS": "c.1630C>T",
            "pHGVS_S": "p.Q544*",
            "Function": "Nonsense",
            "Freq(%)": 5.0,
        },
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "RAD50",
            "Transcript": "NM_005732.4",
            "Chr": "chr5",
            "ExIn_ID": "EX8",
            "cHGVS": "c.1093C>T",
            "pHGVS_S": "p.R365*",
            "Function": "Nonsense",
            "Freq(%)": 4.0,
        },
    ]
    excel_data = _excel(tmp_path, variations=variations)
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")

    report_data = mapper.map(excel_data, panel_package=package)
    assert "FANCD2：c.1630C>T，p.Q544*" in report_data.get_field("immuno_positive_genes")

    enhance_report_data(
        report_data,
        excel_data,
        field_mapper=mapper,
        base_path=str(ROOT),
        panel_id=package.panel_id,
        panel_package=package,
    )
    assert [row["gene"] for row in report_data.get_table("immune_positive_variants")] == [
        "FANCD2",
        "RAD50",
    ]
    assert report_data.get_field("imm_pos_DDR").splitlines() == [
        "FANCD2：c.1630C>T，p.Q544*",
        "RAD50：c.1093C>T，p.R365*",
    ]

    template_path = tmp_path / "immune_summary_template.docx"
    output_path = tmp_path / "immune_summary_output.docx"
    doc = Document()
    doc.add_paragraph("免疫治疗正相关基因")
    doc.add_paragraph("{{ immuno_positive_genes }}")
    doc.add_paragraph("DDR")
    doc.add_paragraph("{{ imm_pos_DDR }}")
    doc.save(template_path)

    TemplateRenderer(log_level="ERROR").render(
        str(template_path),
        report_data,
        str(output_path),
        post_processor_names=[],
    )
    rendered_text = "\n".join(p.text for p in Document(output_path).paragraphs)
    assert "检出（2个）" in rendered_text
    assert "FANCD2：c.1630C>T，p.Q544*" in rendered_text
    assert "RAD50：c.1093C>T，p.R365*" in rendered_text


def test_immune_negative_and_hyperprogression_summary_use_count_header():
    from reportgen.core.template_bridge_358 import format_immune_result

    variants = [
        {
            "gene": "DNMT3A",
            "cHGVS": "c.1367delA",
            "pHGVS": "p.K456Sfs*195",
        }
    ]

    text = format_immune_result(variants, "hyperprogression")

    assert text == "检出（1个）\nDNMT3A：c.1367delA，p.K456Sfs*195"
    assert "检出：" not in text


def test_reviewed_variant_override_replaces_existing_targeted_tip():
    report_data = ReportData()
    report_data.set_table(
        "variants_2_1",
        [
            {
                "gene": "KRAS",
                "locus": "c.34G>A,\np.G12S",
                "benefit_drugs": "old",
                "caution_drugs": "old",
            }
        ],
    )
    report_data.set_table(
        "targeted_drug_tips",
        [
            {
                "gene": "KRAS",
                "variant_site": "c.34G>A,\np.G12S",
                "benefit_drugs": "Avutometinib+Defactinib（C）",
                "caution_drugs": "西妥昔单抗（A）",
            }
        ],
    )
    panel_config = PanelConfig(
        reviewed_variant_overrides=[
            {
                "gene": "KRAS",
                "c_hgvs": "c.34G>A",
                "p_hgvs": "p.G12S",
                "benefit_drugs": ["司美替尼（C）", "Defactinib+Avutometinib（C）"],
                "caution_drugs": ["西妥昔单抗（A）", "帕尼单抗（A）"],
            }
        ]
    )

    _patch_reviewed_variant_override_rows(report_data, panel_config)

    variant_row = report_data.get_table("variants_2_1")[0]
    tip_row = report_data.get_table("targeted_drug_tips")[0]
    assert "司美替尼（C）" in variant_row["benefit_drugs"]
    assert "Avutometinib+Defactinib" not in tip_row["benefit_drugs"]
    assert "Defactinib+Avutometinib（C）" in tip_row["benefit_drugs"]
    assert "帕尼单抗（A）" in tip_row["caution_drugs"]


@pytest.mark.parametrize(
    ("transcript", "expected_benefit"),
    [
        ("NM_004985.5", "转录本限定药物（A）"),
        ("NM_004985.4", "old"),
        ("", "old"),
    ],
)
def test_transcript_bound_override_only_patches_exact_runtime_event(
    transcript,
    expected_benefit,
):
    report_data = ReportData()
    report_data.set_table(
        "variants_2_1",
        [
            {
                "gene": "KRAS",
                "transcript": transcript,
                "locus": "c.34G>A,\np.G12S",
                "gene_class": "Ⅱ类",
                "benefit_drugs": "old",
                "caution_drugs": "old",
            }
        ],
    )
    panel_config = PanelConfig(
        reviewed_variant_overrides=[
            {
                "gene": "KRAS",
                "transcript": "NM_004985.5",
                "c_hgvs": "c.34G>A",
                "p_hgvs": "p.G12S",
                "variant_level": ["Ⅱ类"],
                "benefit_drugs": ["转录本限定药物（A）"],
                "caution_drugs": ["--"],
            }
        ]
    )

    _patch_reviewed_variant_override_rows(report_data, panel_config)

    row = report_data.get_table("variants_2_1")[0]
    assert row["benefit_drugs"] == expected_benefit


def test_long_drug_lists_are_compacted_for_word_tables_but_kept_full_in_summary():
    long_list = "\n".join(f"药物{i}（C）" for i in range(1, 8))
    report_data = ReportData(
        context={
            "variants_2_1": [
                {
                    "gene": "ERBB2",
                    "locus": "c.1979G>A,\np.G660D",
                    "benefit_drugs": long_list,
                    "caution_drugs": "--",
                }
            ],
            "targeted_drug_tips": [
                {
                    "gene": "ERBB2",
                    "variant_site": "c.1979G>A,\np.G660D",
                    "benefit_drugs": long_list,
                    "caution_drugs": "--",
                }
            ],
        }
    )

    _compact_drug_display_tables(report_data)

    variant_row = report_data.get_table("variants_2_1")[0]
    tip_row = report_data.get_table("targeted_drug_tips")[0]
    assert variant_row["benefit_drugs"].splitlines() == [
        "药物1（C）",
        "药物2（C）",
        "药物3（C）",
        "药物4（C）",
        "药物5（C）",
        "另2项详见第三部分",
    ]
    assert variant_row["benefit_drugs_full"] == long_list
    assert tip_row["benefit_drugs_full"] == long_list
    assert variant_row["caution_drugs"] == "--"
    assert tip_row["caution_drugs"] == "--"

    summary = build_report_summary(
        report_data=report_data,
        project_type="crc_358_msi",
        qa_report={"status": "PASS", "issues": []},
    )

    assert summary["variants"]["key_rows"][0]["benefit_drugs"] == long_list
    assert summary["drugs"]["targeted_rows"][0]["benefit_drugs"] == long_list


def test_crc_panel_can_show_all_drugs_in_part2():
    long_list = "\n".join(f"药物{i}（C）" for i in range(1, 8))
    report_data = ReportData(
        context={"variants_2_1": [{"benefit_drugs": long_list, "caution_drugs": "--"}]}
    )

    _compact_drug_display_tables(report_data, max_items=None)

    assert report_data.get_table("variants_2_1")[0]["benefit_drugs"] == long_list
    assert "另" not in report_data.get_table("variants_2_1")[0]["benefit_drugs"]


def test_fanca_reviewed_rule_only_matches_class_ii_loss_of_function():
    rule = {
        "gene": "FANCA",
        "variant_level": ["Ⅱ类"],
        "applicability": "loss_of_function",
    }
    assert _variant_override_matches(rule, "FANCA", "c.100_101del", "p.K34Rfs*5", gene_class="Ⅱ类")
    assert not _variant_override_matches(rule, "FANCA", "c.100A>G", "p.K34R", gene_class="Ⅱ类")
    assert not _variant_override_matches(
        rule, "FANCA", "c.100_101del", "p.K34Rfs*5", gene_class="Ⅲ类"
    )


def test_report_group_feedback_variants_have_exact_panel_scoped_rules():
    rules = yaml.safe_load(
        (ROOT / "panels/crc_358_msi/rules/drugs.yaml").read_text(encoding="utf-8")
    )["targeted_drug_rules"]["reviewed_variant_overrides"]
    expected = {
        "FANCD2": ("c.1630C>T", "p.Q544*", 2),
        "RAD50": ("c.1093C>T", "p.R365*", 4),
        "PALB2": ("c.47del", "p.K16Sfs*2", 2),
        "RAD51D": ("c.685C>T", "p.Q229*", 4),
    }

    for gene, (c_hgvs, p_hgvs, research_count) in expected.items():
        rule = next(row for row in rules if row.get("gene") == gene and row.get("c_hgvs") == c_hgvs)
        assert rule["panels"] == ["crc_358_msi"]
        assert rule["review_status"] == "approved_for_runtime"
        assert rule["secondary_review_status"] == "report_group_approved"
        assert rule["benefit_drugs"] == ["--"]
        assert rule["caution_drugs"] == ["--"]
        assert len(rule["research_drugs"]) == research_count
        assert _variant_override_matches(rule, gene, c_hgvs, p_hgvs, gene_class="Ⅱ类")
        assert not _variant_override_matches(rule, gene, "c.999999A>G", "p.X1Y", gene_class="Ⅱ类")
        assert not _variant_override_matches(rule, gene, c_hgvs, p_hgvs, gene_class="Ⅲ类")


def test_research_only_drugs_follow_the_same_class_boundary_as_part2(tmp_path):
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    panel_config = load_panel_config(
        base_path=str(ROOT),
        panel_id="crc_358_msi",
        panel_package=package,
        config_path=str(package.resolve_rule_file("panel_rules")),
    )

    def build(level):
        return build_variants_for_template(
            _excel(
                tmp_path,
                variations=[
                    {
                        "Gene_Symbol": "FANCD2",
                        "Transcript": "NM_001018115.3",
                        "Chr": "3",
                        "ExIn_ID": "EX19",
                        "cHGVS": "c.1630C>T",
                        "pHGVS_S": "p.Q544*",
                        "Freq(%)": 12.3,
                        "ExistInsmall358": 1,
                        "ExistIn552": level,
                    }
                ],
            ),
            filter_class_i_ii_only=False,
            important_genes_only=False,
            panel_config=panel_config,
        )[0]

    assert "芦卡帕利" in build("Ⅱ类")["research_drugs"]
    assert build("Ⅲ类")["research_drugs"] == "--"


def test_feedback_ddr_genes_are_in_runtime_immune_rows():
    panel_config = load_panel_config(base_path=str(ROOT), panel_id="crc_358_msi")
    ddr_row = next(row for row in panel_config.immune_positive_rows if row.get("key") == "DDR")

    assert {"FANCD2", "RAD50", "PALB2", "RAD51D"} <= set(panel_config.immune_positive_genes)
    assert {"FANCD2", "RAD50", "PALB2", "RAD51D"} <= {
        str(gene).upper() for gene in ddr_row["genes"]
    }


def test_reviewed_part3_drug_sections_bind_to_exact_feedback_variants():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    provider = GeneKnowledgeProvider(
        {
            "enabled": True,
            "gene_knowledge_db": {
                "enabled": True,
                "path": "missing.xlsx",
                "reviewed_part3_overlay_paths": (
                    ReportGenerator._resolve_panel_reviewed_part3_overlays(package)
                ),
            },
        }
    )
    provider.load(base_path=str(ROOT))
    variants = [
        {
            "gene": gene,
            "cHGVS": c_hgvs,
            "pHGVS": p_hgvs,
            "frequency": "1.00",
            "benefit_drugs": "--",
            "caution_drugs": "--",
            "research_drugs": "研究性方案",
        }
        for gene, c_hgvs, p_hgvs in (
            ("FANCD2", "c.1630C>T", "p.Q544*"),
            ("RAD50", "c.1093C>T", "p.R365*"),
            ("PALB2", "c.47del", "p.K16Sfs*2"),
            ("RAD51D", "c.685C>T", "p.Q229*"),
        )
    ]

    sections = provider.build_drug_analysis_sections(variants)

    assert [section["gene"] for section in sections] == [
        "FANCD2",
        "RAD50",
        "PALB2",
        "RAD51D",
    ]
    assert all(section["drug_type"] == "research" for section in sections)
    assert all("NCT" in section["clinical"] for section in sections)
    assert all("不能" in section["relation"] for section in sections)


def test_immune_only_part3_variant_is_marked_as_therapy_associated(tmp_path):
    class NoTargetedDrugLookup:
        def _lookup_targeted_drugs_for_variant(self, *_args, **_kwargs):
            return "--", "--", 0.0

    report_data = ReportData()
    report_data.set_field(
        "report_content",
        {
            "part3_variant_scope": "summary_variants",
            "part3_reference_variant_scope": "summary_variants",
        },
    )
    result = enhance_report_data(
        report_data,
        _excel(
            tmp_path,
            variations=[
                {
                    "ExistIn552": "Ⅱ类",
                    "ExistInsmall358": 1,
                    "Gene_Symbol": "DNMT3A",
                    "cHGVS": "c.1367del",
                    "pHGVS_S": "p.K456Sfs*195",
                    "Function": "Frameshift",
                    "Freq(%)": 1.18,
                    "CLNSIG": "Pathogenic",
                }
            ],
        ),
        field_mapper=NoTargetedDrugLookup(),
        gene_knowledge_provider=GeneKnowledgeProvider({"enabled": False}),
        base_path=str(ROOT),
    )

    section = result.get_table("gene_knowledge_sections")[0]
    assert section["gene"] == "DNMT3A"
    assert section["has_drug"] is True
    assert section["header_color"] == "FF0000"


def test_crc301_additional_overlay_is_loaded_after_shared_crc_overlay():
    package = load_panel_package("crc_301_msi", project_root=ROOT)
    paths = ReportGenerator._resolve_panel_reviewed_part3_overlays(package)
    provider = GeneKnowledgeProvider(
        {
            "enabled": True,
            "gene_knowledge_db": {
                "enabled": True,
                "path": "missing.xlsx",
                "reviewed_part3_overlay_paths": paths,
            },
        }
    )
    provider.load(base_path=str(ROOT))

    sections = provider.build_all_gene_knowledge_sections(
        [
            {
                "gene": "ABCB1",
                "cHGVS": "c.1A>G",
                "pHGVS": "p.M1V",
                "frequency": "1%",
            }
        ]
    )

    relative_paths = [Path(path).relative_to(ROOT).as_posix() for path in paths]
    assert relative_paths[:2] == [
        "panels/crc_358_msi/rules/reviewed_part3_knowledge.yaml",
        "panels/crc_301_msi/rules/reviewed_part3_knowledge.yaml",
    ]
    assert relative_paths[-2:] == [
        "panels/crc_358_msi/rules/reviewed_part3_drug_contracts_20260721.yaml",
        "panels/crc_358_msi/rules/reviewed_part3_medical_corrections_20260722.yaml",
    ]
    assert len(sections) == 1
    assert "P-糖蛋白" in sections[0]["intro"]
    assert "任意体细胞变异不能直接解释为化疗耐药" in sections[0]["mutation_analysis"]


def test_report_group_approved_gene_rows_keep_reviewed_conservative_wording():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    provider = GeneKnowledgeProvider(
        {
            "enabled": True,
            "gene_knowledge_db": {
                "enabled": True,
                "path": "missing.xlsx",
                "reviewed_part3_overlay_paths": (
                    ReportGenerator._resolve_panel_reviewed_part3_overlays(package)
                ),
            },
        }
    )
    provider.load(base_path=str(ROOT))

    overlay = yaml.safe_load(
        (ROOT / "panels/crc_358_msi/rules/reviewed_part3_knowledge.yaml").read_text(
            encoding="utf-8"
        )
    )
    expected = {
        row["gene"]: row
        for row in overlay["gene_sections"]
        if row.get("review_status") == "approved_for_runtime"
    }

    for gene in ("CHD2", "HIST1H3B", "HLA-DPA1", "WDR90", "ZNF703"):
        sections = provider.build_all_gene_knowledge_sections(
            [
                {
                    "gene": gene,
                    "cHGVS": "c.999A>G",
                    "pHGVS": "p.X333Y",
                    "frequency": "1%",
                }
            ]
        )
        assert len(sections) == 1
        assert sections[0]["intro"] == expected[gene]["intro"]
        assert sections[0]["mutation_analysis"].endswith(expected[gene]["mutation_analysis"])
        assert sections[0]["mutation_analysis"].count("编码的蛋白全长") == 1


def test_excel_reader_resolves_msi_sheet_alias_reordered_columns_and_tumor_row(tmp_path):
    import pandas as pd

    path = tmp_path / "msi_alias.xlsx"
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        pd.DataFrame({"A": [1]}).to_excel(writer, index=False, sheet_name="Meta")
        pd.DataFrame(
            {
                "Status": ["MSS", "MSI-H"],
                "Sample": ["normal", "tumor"],
                "Percent": [1.1, 45.2],
            }
        ).to_excel(writer, index=False, sheet_name="MSI sensor")

    result = ExcelReader(config_dir=str(ROOT / "config"), log_level="ERROR").read(str(path))

    assert result.single_values["MSI状态"] == "MSI-H"
    assert float(result.single_values["MSI百分比"]) == 45.2


def test_excel_reader_infers_msi_status_from_percent_when_status_missing(tmp_path):
    import pandas as pd

    path = tmp_path / "msi_percent_only.xlsx"
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        pd.DataFrame({"A": [1]}).to_excel(writer, index=False, sheet_name="Meta")
        pd.DataFrame({"Sample": ["normal", "tumor"], "MSI Percentage": [1.1, "45.2%"]}).to_excel(
            writer, index=False, sheet_name="MSIsensor"
        )

    result = ExcelReader(config_dir=str(ROOT / "config"), log_level="ERROR").read(str(path))

    assert result.single_values["MSI状态"] == "MSI-H"


def test_excel_reader_reads_summary_first_legacy_msisensor_export(tmp_path):
    import pandas as pd

    path = tmp_path / "legacy_msisensor_summary.xlsx"
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        pd.DataFrame({"A": [1]}).to_excel(writer, index=False, sheet_name="Meta")
        pd.DataFrame(
            [
                ["msisensor", 18, 0, 0],
                [None, None, None, None],
                ["Chr", "Pos", "Pvalue", "Name"],
            ],
            columns=[
                "SOFT",
                "Total_Number_of_Sites",
                "Number_of_Somatic_Sites",
                "%",
            ],
        ).to_excel(writer, index=False, sheet_name="Msisensor")

    result = ExcelReader(config_dir=str(ROOT / "config"), log_level="ERROR").read(str(path))

    assert result.single_values["MSI状态"] == "MSS"
    assert float(result.single_values["MSI百分比"]) == 0.0


def test_excel_reader_prefers_msisensor2_over_legacy_summary(tmp_path):
    import pandas as pd

    path = tmp_path / "combined_msisensor_summary.xlsx"
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        pd.DataFrame(
            [
                ["msisensor", 22, 0, 0, None],
                ["msisensor2", 114, 3, 45.5, "MSI-H"],
            ],
            columns=[
                "SOFT",
                "Total_Number_of_Sites",
                "Number_of_Somatic_Sites",
                "%",
                "Status",
            ],
        ).to_excel(writer, index=False, sheet_name="Msisensor")

    result = ExcelReader(config_dir=str(ROOT / "config"), log_level="ERROR").read(str(path))

    assert result.single_values["MSI状态"] == "MSI-H"
    assert float(result.single_values["MSI百分比"]) == pytest.approx(45.5)


def test_part3_variant_scope_can_follow_summary_variants(tmp_path):
    report_data = ReportData()
    report_data.set_field(
        "report_content",
        {
            "part3_variant_scope": "summary_variants",
            "part3_reference_variant_scope": "summary_variants",
        },
    )
    report_data = enhance_report_data(
        report_data,
        _excel(
            tmp_path,
            variations=[
                {
                    "ExistIn552": "Ⅱ类",
                    "ExistInsmall358": 1,
                    "Gene_Symbol": "TP53",
                    "Transcript": "NM_000546.6",
                    "Chr": "chr17",
                    "ExIn_ID": "EX8",
                    "cHGVS": "c.844C>T",
                    "pHGVS_S": "p.R282W",
                    "Function": "Missense",
                    "Freq(%)": 67.29,
                    "CLNSIG": "Pathogenic",
                },
                {
                    "ExistIn552": "Ⅲ类",
                    "ExistInsmall358": 1,
                    "Gene_Symbol": "APC",
                    "Transcript": "NM_000038.6",
                    "Chr": "chr5",
                    "ExIn_ID": "EX16E",
                    "cHGVS": "c.4348C>T",
                    "pHGVS_S": "p.R1450*",
                    "Function": "Nonsense",
                    "Freq(%)": 41.12,
                    "CLNSIG": "Pathogenic",
                },
            ],
        ),
        field_mapper=_FakeDrugLookup(),
        gene_knowledge_provider=_FakeGeneKnowledgeProvider(),
        base_path=str(ROOT),
    )

    headers = [row["header"] for row in report_data.get_table("gene_knowledge_sections")]
    assert headers == ["TP53：c.844C>T", "APC：c.4348C>T"]
    assert report_data.get_table("gene_references") == [
        "TP53 reference",
        "APC reference",
    ]


def test_part3_drug_analysis_scope_follows_summary_variants(tmp_path):
    report_data = ReportData()
    report_data.set_field(
        "report_content",
        {
            "part3_variant_scope": "summary_variants",
            "part3_reference_variant_scope": "summary_variants",
        },
    )
    report_data = enhance_report_data(
        report_data,
        _excel(
            tmp_path,
            variations=[
                {
                    "ExistIn552": "Ⅱ类",
                    "ExistInsmall358": 1,
                    "Gene_Symbol": "TP53",
                    "cHGVS": "c.844C>T",
                    "pHGVS_S": "p.R282W",
                    "Function": "Missense",
                    "Freq(%)": 67.29,
                    "CLNSIG": "Pathogenic",
                },
                {
                    "ExistIn552": "Ⅲ类",
                    "ExistInsmall358": 1,
                    "Gene_Symbol": "APC",
                    "cHGVS": "c.4348C>T",
                    "pHGVS_S": "p.R1450*",
                    "Function": "Nonsense",
                    "Freq(%)": 41.12,
                    "CLNSIG": "Pathogenic",
                },
            ],
        ),
        field_mapper=_FakeDrugLookup(),
        gene_knowledge_provider=_FakeDrugSectionGeneKnowledgeProvider(),
        base_path=str(ROOT),
    )

    genes = [row["gene"] for row in report_data.get_table("drug_analysis_sections")]
    assert genes == ["TP53", "APC"]


def test_part3_and_word_tables_share_grouped_vaf_order(tmp_path):
    """Part 2 and Part 3 keep each gene contiguous under one VAF policy."""
    report_data = ReportData()
    report_data.set_field(
        "report_content",
        {
            "part3_variant_scope": "summary_variants",
            "part3_reference_variant_scope": "summary_variants",
        },
    )
    variations = [
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "APC",
            "cHGVS": "c.994C>T",
            "pHGVS_S": "p.R332*",
            "Freq(%)": 22.03,
        },
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "APC",
            "cHGVS": "c.4666_4667insA",
            "pHGVS_S": "p.T1556Nfs*3",
            "Freq(%)": 17.50,
        },
        {
            "ExistIn552": "Ⅱ类",
            "ExistInsmall358": 1,
            "Gene_Symbol": "SMAD4",
            "cHGVS": "c.1081C>T",
            "pHGVS_S": "p.R361C",
            "Freq(%)": 20.00,
        },
    ]

    report_data = enhance_report_data(
        report_data,
        _excel(tmp_path, variations=variations),
        field_mapper=_FakeDrugLookup(),
        gene_knowledge_provider=_FakeDrugSectionGeneKnowledgeProvider(),
        base_path=str(ROOT),
    )

    # Table views keep APC adjacent so OOXML vertical merging remains valid.
    assert [(row["gene"], row["frequency"]) for row in report_data.get_table("all_variants")] == [
        ("APC", "22.03"),
        ("APC", "17.50"),
        ("SMAD4", "20"),
    ]

    # Part 3 uses the same grouped policy: genes by max VAF, then sites by VAF.
    assert [row["header"] for row in report_data.get_table("gene_knowledge_sections")] == [
        "APC：c.994C>T",
        "APC：c.4666_4667insA",
        "SMAD4：c.1081C>T",
    ]
    assert [row["gene"] for row in report_data.get_table("drug_analysis_sections")] == [
        "APC",
        "APC",
        "SMAD4",
    ]
    assert report_data.get_table("gene_references") == [
        "APC reference",
        "APC reference",
        "SMAD4 reference",
    ]


def test_part3_marker_renders_from_context_without_case_stub(tmp_path):
    docx_path = tmp_path / "part3_marker.docx"
    doc = Document()
    doc.add_paragraph("第三部分：基因变异及相应靶向/免疫药物解析")
    doc.add_paragraph("基因变异解析")
    doc.add_paragraph("__PART3_MARKER__")
    doc.add_paragraph("3. 阅读说明")
    doc.add_paragraph("5. 参考文献")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._render_part3_formatted(
        str(docx_path),
        {
            "total_variants_count": 2,
            "drug_related_count": 1,
            "targeted_or_immune_related_count": 2,
            "gene_knowledge_sections": [
                {
                    "gene": "BRAF",
                    "header": "BRAF：c.1799T>A，p.V600E；12.30%",
                    "has_drug": True,
                    "intro": "BRAF intro",
                    "mutation_desc": "BRAF desc",
                    "mutation_analysis": "BRAF analysis",
                }
            ],
            "drug_benefit_sections": [
                {
                    "gene": "BRAF",
                    "variant": "c.1799T>A，p.V600E",
                    "drug_name": "维莫非尼",
                    "clinical": "BRAF clinical",
                }
            ],
            "drug_caution_sections": [],
            "drug_research_sections": [
                {
                    "gene": "BRAF",
                    "variant": "c.1799T>A，p.V600E",
                    "drug_name": "研究方案甲",
                    "relation": "仅为临床试验筛选线索。",
                    "clinical": "不构成潜在获益结论。",
                }
            ],
            "gene_references": ["BRAF reference"],
        },
    )

    paragraphs = [p.text for p in Document(docx_path).paragraphs]
    text = "\n".join(paragraphs)
    assert "__PART3_MARKER__" not in text
    assert "BRAF：c.1799T>A，p.V600E；12.30%" in text
    assert "其中与靶向/免疫药物相关的变异：2个" in text
    assert "维莫非尼" in text
    assert "临床试验/研究性方案解析" in text
    assert "研究方案甲" in text
    assert "基因变异与研究方案关联：" in text
    assert "临床试验/研究证据：" in text
    assert "BRAF reference" not in text
    assert paragraphs.count("3. 阅读说明") == 1
    assert "p.G12S" not in text
    assert "c.34G>A" not in text
    assert "46.29" not in text


def test_targeted_drug_brand_summary_uses_final_drug_columns_only():
    summary = build_targeted_drug_brand_summary(
        [
            {
                "gene": "KRAS",
                "benefit_drugs": "Avutometinib+Defactinib（C）\n司美替尼（C）",
                "caution_drugs": "西妥昔单抗（A）\n帕尼单抗（A）",
            },
            {
                "gene": "ATM",
                "benefit_drugs": "奥拉帕利+帕博利珠单抗（C）",
                "caution_drugs": "--",
            },
        ],
        base_path=str(ROOT),
    )

    assert summary == (
        "司美替尼[科赛优]、帕尼单抗[维克替比]、Defactinib[FAKZYNJA]、"
        "奥拉帕利[利普卓]、西妥昔单抗[爱必妥]、Avutometinib[AVMAPKI]、"
        "帕博利珠单抗[可瑞达]。"
    )


def test_targeted_brand_summary_has_explicit_order_and_no_nested_false_match(
    tmp_path,
):
    config_dir = tmp_path / "config"
    config_dir.mkdir()
    (config_dir / "drug_brands.yaml").write_text(
        """
schema_version: "1.1"
targeted_summary_order: [DrugA, DrugB, 西罗莫司, 替西罗莫司, ABC, BCDE]
brands:
  DrugB: B牌
  DrugA: A牌
  西罗莫司: 西罗莫司牌
  替西罗莫司: 替西罗莫司牌
  ABC: ABC牌
  BCDE: BCDE牌
""".strip()
        + "\n",
        encoding="utf-8",
    )

    summary = build_targeted_drug_brand_summary(
        [
            {
                "benefit_drugs": "DrugB（C）\n替西罗莫司（C）\nDrugA（C）\nABCDE（C）",
                "caution_drugs": "--",
            }
        ],
        base_path=str(tmp_path),
    )

    assert summary == ("DrugA[A牌]、DrugB[B牌]、替西罗莫司[替西罗莫司牌]、BCDE[BCDE牌]。")
    assert "西罗莫司[西罗莫司牌]" not in summary
    assert "ABC[ABC牌]" not in summary


def test_targeted_brand_summary_warns_when_config_is_unavailable(monkeypatch):
    monkeypatch.setattr(
        "reportgen.core.template_bridge_358._load_drug_brand_config",
        lambda **_kwargs: ({}, []),
    )

    summary, warnings = _build_targeted_drug_brand_summary_result(
        [
            {
                "benefit_drugs": "奥拉帕利（C）",
                "caution_drugs": "--",
            }
        ],
        base_path=str(ROOT),
    )

    assert summary == ""
    assert warnings == ["BRAND_CONFIG_MISSING"]


def test_qa_report_exposes_targeted_brand_mapping_warning(tmp_path):
    docx_path = tmp_path / "brand_mapping_warning.docx"
    doc = Document()
    doc.add_paragraph("本次共检出体细胞变异：0个")
    doc.save(docx_path)
    report_data = ReportData()
    report_data.set_field("targeted_drug_brand_warnings", ["NO_CONFIGURED_BRAND_MATCH"])

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        report_data=report_data,
        project_type="crc_358_msi",
    )

    assert qa["checks"]["targeted_drug_brand_mapping"] == {
        "status": "WARN",
        "warning_codes": ["NO_CONFIGURED_BRAND_MATCH"],
    }
    assert any(
        issue["code"] == "TARGETED_DRUG_BRAND_MAPPING" and issue["level"] == "warning"
        for issue in qa["issues"]
    )


def test_targeted_brand_processor_replaces_golden_template_static_summary(tmp_path):
    docx_path = tmp_path / "targeted_brand_summary.docx"
    doc = Document()
    doc.add_paragraph("2.上表涉及的已上市的药物名称及对应的商品名称：旧商品名。")
    doc.add_paragraph(
        "3. 上表涉及的已上市的药物名称及对应的商品名称："
        "帕博利珠单抗[可瑞达]、纳武利尤单抗[欧狄沃]。"
    )
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._apply_targeted_drug_brand_summary(
        str(docx_path),
        {"targeted_drug_brand_summary": ("奥拉帕利[利普卓]、纳武利尤单抗[欧狄沃]。")},
    )

    paragraphs = [p.text for p in Document(docx_path).paragraphs]
    assert paragraphs[0] == (
        "2.上表涉及的已上市的药物名称及对应的商品名称：奥拉帕利[利普卓]、纳武利尤单抗[欧狄沃]。"
    )
    assert paragraphs[1] == (
        "3. 上表涉及的已上市的药物名称及对应的商品名称："
        "帕博利珠单抗[可瑞达]、纳武利尤单抗[欧狄沃]。"
    )


def test_case2_fixture_tmb_msi_mapping_if_available():
    fixture = ROOT / "output/xlsx_for_win/case2_highTMB_MSIH_MLB0002.result.xlsx"
    if not fixture.exists():
        pytest.skip("case2 fixture is not available in this checkout")

    excel_data = ExcelReader(config_dir=str(ROOT / "config"), log_level="ERROR").read(
        str(fixture), include_tables=True
    )
    report_data = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR").map(excel_data)

    assert excel_data.single_values["TMB"] == pytest.approx(7.10499382716049)
    assert excel_data.single_values["MSI状态"] == "MSI-H"
    assert report_data.get_field("tmb_value") == "7.1"
    assert report_data.get_field("tmb_status") == "L"
    assert report_data.get_field("msi_status") == "MSI-H"


def test_field_mapper_updates_msi_status_cn_from_mss(tmp_path):
    report_data = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR").map(
        _excel(tmp_path, single_values={"MSI状态": "MSS"})
    )

    assert report_data.get_field("msi_status") == "MSS"
    assert report_data.get_field("msi_status_cn") == "微卫星稳定型，MSS"


def test_missing_report_date_is_filled_with_generation_date():
    report_data = ReportData()
    generator = ReportGenerator(config_dir=str(ROOT / "config"), log_level="ERROR")

    generator._mark_missing_report_date(report_data)

    assert report_data.get_field("report_date") == date.today().isoformat()
    assert "缺失必填字段: report_date" not in report_data.validation_errors


def test_common_validation_warns_with_today_backfill(tmp_path):
    warnings = validate_excel_data_common(_excel(tmp_path), today=date(2026, 4, 19))

    report_date_warnings = [w for w in warnings if w.get("field") == "report_date"]
    assert report_date_warnings
    assert "系统将使用生成报告当天日期" in report_date_warnings[0]["message"]


def test_detector_does_not_use_panel_column_numbers_as_crc_signal(tmp_path):
    excel_data = _excel(
        tmp_path,
        single_values={"TMB": 3.58},
        variations=[{"ExistInsmall358": 1, "Gene_Symbol": "TP53", "cHGVS": "c.1A>T"}],
    )
    detector = ProjectDetector(config_dir=str(ROOT / "config"), log_level="ERROR")
    result = detector.detect(str(Path(excel_data.file_path)), excel_data=excel_data)

    assert result["project_type"] != "crc_358_msi"


def test_detector_uses_trusted_filename_project_tokens(tmp_path):
    detector = ProjectDetector(config_dir=str(ROOT / "config"), log_level="ERROR")
    excel_data = _excel(tmp_path)

    path_301 = tmp_path / "_MLS2600000001_结直肠癌301基因+MSI_终版.xlsx"
    path_358 = tmp_path / "_MLS2600000002_结直肠癌358基因+MSI_终版.xlsx"
    path_lung = tmp_path / "_LUNG999001_肺癌甲基化_终版.xlsx"
    path_301.write_bytes(b"placeholder")
    path_358.write_bytes(b"placeholder")
    path_lung.write_bytes(b"placeholder")

    assert detector.detect(str(path_301), excel_data=excel_data)["project_type"] == "crc_301_msi"
    assert detector.detect(str(path_358), excel_data=excel_data)["project_type"] == "crc_358_msi"
    assert (
        detector.detect(str(path_lung), excel_data=excel_data)["project_type"] == "lung_methylation"
    )


def test_crc_panel_enhancer_accepts_legacy_aliases():
    from reportgen.core.enhancer_registry import (
        CRC358Enhancer,
        NoopEnhancer,
        UnknownPanelError,
        get_enhancer,
        get_panel_registry,
        get_registered_project_types,
        is_registered_project_type,
        normalize_project_type,
    )

    assert isinstance(get_enhancer("crc_301"), CRC358Enhancer)
    assert isinstance(get_enhancer("crc_358"), CRC358Enhancer)
    assert normalize_project_type("crc_358") == "crc_358_msi"
    assert normalize_project_type("CRC358") == "crc_358_msi"
    assert normalize_project_type("crc301") == "crc_301_msi"
    assert is_registered_project_type("crc_358")
    assert "crc_358_msi" in get_registered_project_types()
    assert get_panel_registry().get("crc_358").package.panel_id == "crc_358_msi"
    assert get_panel_registry().get("crc_301").package.panel_id == "crc_301_msi"
    lung = get_panel_registry().get("lung_methylation")
    assert isinstance(lung.enhancer, NoopEnhancer)
    assert lung.package.panel_id == "lung_methylation"
    with pytest.raises(UnknownPanelError):
        get_enhancer("unknown_panel")


def test_report_generator_rejects_unknown_project_type(tmp_path):
    generator = ReportGenerator(config_dir=str(ROOT / "config"), log_level="ERROR")

    result = generator.generate(
        excel_file=str(tmp_path / "missing.xlsx"),
        template_file=str(ROOT / "templates" / "aligned_template_with_cnv_fusion_hla_FIXED.docx"),
        output_dir=str(tmp_path),
        project_type="unknown_panel",
    )

    assert result["success"] is False
    assert any("未注册的Panel项目类型" in error for error in result["errors"])
    assert result["stage_results"][0]["name"] == "PanelResolutionStage"
    assert result["stage_results"][0]["status"] == "FAIL"
    assert result["stage_results"][0]["duration_ms"] is not None


def _write_minimal_panel_package(
    tmp_path: Path,
    panel_id: str,
    *,
    directory_name=None,
    aliases=None,
) -> Path:
    panel_dir = tmp_path / "panels" / (directory_name or panel_id)
    (panel_dir / "templates").mkdir(parents=True)
    (panel_dir / "rules").mkdir(parents=True)
    (panel_dir / "templates" / "standard.docx").write_bytes(b"placeholder")
    (panel_dir / "rules" / "panel.yaml").write_text(
        yaml.safe_dump(
            {
                "schema_version": "1.0",
                "panel_id": panel_id,
                "rule_id": "panel_rules",
                "version": "0.1.0",
                "status": "draft",
                "rules": [],
            },
            sort_keys=False,
            allow_unicode=True,
        ),
        encoding="utf-8",
    )
    (panel_dir / "rules" / "report_text.yaml").write_text(
        yaml.safe_dump(
            {
                "schema_version": "1.0",
                "panel_id": panel_id,
                "rule_id": "report_text",
                "version": "0.1.0",
                "status": "draft",
                "texts": {"summary": "synthetic"},
            },
            sort_keys=False,
            allow_unicode=True,
        ),
        encoding="utf-8",
    )
    (panel_dir / "mappings.yaml").write_text("mappings: {}\n", encoding="utf-8")
    (panel_dir / "qa.yaml").write_text(
        yaml.safe_dump(
            {
                "schema_version": "1.0",
                "panel_id": panel_id,
                "current_output": {
                    "enabled": False,
                    "source": "golden_reference",
                    "required_features": {},
                    "required_sections": {},
                    "require_table_shapes": "warn",
                    "privacy_checks": {
                        "report_id": "fail",
                        "sample_id": "fail",
                        "date": "fail",
                    },
                },
                "legacy_reference": {
                    "enabled": False,
                    "sample_count": 1,
                    "required_features": {},
                    "required_sections": {},
                    "require_table_shapes": "warn",
                    "privacy_checks": {
                        "source_dir": "fail",
                        "report_id": "fail",
                        "sample_id": "fail",
                        "date": "fail",
                    },
                },
            },
            sort_keys=False,
            allow_unicode=True,
        ),
        encoding="utf-8",
    )
    panel_yaml = panel_dir / "panel.yaml"
    panel_yaml.write_text(
        yaml.safe_dump(
            {
                "schema_version": "1.0",
                "panel_id": panel_id,
                "display_name": panel_id.replace("_", " ").title(),
                "version": "0.1.0",
                "status": "draft",
                "aliases": aliases or [],
                "default_template": "standard",
                "templates": [
                    {
                        "id": "standard",
                        "file": "templates/standard.docx",
                        "status": "draft",
                    }
                ],
                "mappings": {"default": "mappings.yaml"},
                "rules": {
                    "panel_rules": "rules/panel.yaml",
                    "report_text": "rules/report_text.yaml",
                },
                "input_contract": {"required_tables": ["Variations"]},
                "template_contract": {"required_variables": ["patient_name"]},
                "golden_cases": [
                    {
                        "id": f"{panel_id}_synthetic",
                        "runner": "reportgen.core.golden_case:run_golden_case",
                        "synthetic": True,
                        "expected_qa_status": "PASS",
                    }
                ],
            },
            sort_keys=False,
            allow_unicode=True,
        ),
        encoding="utf-8",
    )
    return panel_yaml


def test_panel_package_loader_reads_crc358_package():
    package = load_panel_package("crc_358_msi", project_root=ROOT)

    assert package.panel_id == "crc_358_msi"
    assert package.display_name == "结直肠癌358基因+MSI"
    assert "crc_358" in package.aliases
    assert package.default_template.template_id == "crc_358_msi_golden_template_v0"
    assert package.resolve_template_file().exists()
    assert package.resolve_mapping_file().name == "mapping.yaml"
    assert package.resolve_rule_file("panel_rules").name == "crc.yaml"
    assert "panels/crc_358_msi/rules" in str(package.resolve_rule_file("panel_rules"))
    assert package.resolve_rule_file("panel_rules").exists()
    assert "variant_tables" in package.processors
    assert package.input_contract["required_tables"] == ["Variations"]
    assert "variant_detail" in package.template_contract["required_table_structures"]
    assert package.qa_profile["panel_id"] == "crc_358_msi"
    assert package.qa_profile["current_output"]["enabled"] is True
    assert package.qa_profile["legacy_reference"]["enabled"] is True
    assert package.qa_profile["legacy_reference"]["source_dir_name"] == "crc_358_msi"
    assert package.golden_cases[0]["id"] == "crc_358_msi_synthetic_low_tmb_mss"


def test_panel_package_loader_loads_all_packages():
    loader = PanelPackageLoader(project_root=ROOT)
    packages = loader.load_all()

    assert {package.panel_id for package in packages} >= {
        "crc_301_msi",
        "crc_358_msi",
        "lung_methylation",
    }


def test_panel_package_loader_reads_crc301_package():
    package = load_panel_package("crc_301_msi", project_root=ROOT)

    assert package.panel_id == "crc_301_msi"
    assert package.display_name == "结直肠癌301基因+MSI"
    assert "crc_301" in package.aliases
    assert package.default_template.template_id == "crc_301_msi_golden_template_v1"
    assert package.resolve_template_file().exists()
    assert package.resolve_rule_file("panel_rules").name == "crc.yaml"
    assert "panels/crc_301_msi/rules" in str(package.resolve_rule_file("panel_rules"))
    assert package.input_contract["required_columns"]["Variations"] == [
        "Gene_Symbol",
        "cHGVS",
        "ExistInsmall301",
        "ExistIn552",
    ]
    assert package.qa_profile["legacy_reference"]["enabled"] is True
    assert package.qa_profile["current_output"]["required_sections"]["gene_list"] == "warn"
    assert package.qa_profile["legacy_reference"]["required_sections"]["gene_list"] == "warn"


def test_panel_package_loader_reads_lung_methylation_package():
    package = load_panel_package("lung_methylation", project_root=ROOT)

    assert package.panel_id == "lung_methylation"
    assert package.display_name == "肺癌甲基化"
    assert package.raw["status"] == "draft"
    assert package.default_template.template_id == "lung_methylation_minimal_v1"
    assert package.default_template.status == "draft"
    assert package.resolve_template_file().exists()
    assert package.resolve_rule_file("panel_rules").name == "lung_methylation.yaml"
    assert package.input_contract["required_tables"] == ["甲基化位点"]
    assert "methylation_sites" in package.template_contract["required_lists"]
    assert package.qa_profile["current_output"]["enabled"] is False
    assert package.qa_profile["legacy_reference"]["enabled"] is False
    assert package.golden_cases[0]["id"] == "lung_methylation_synthetic_positive"
    assert package.golden_cases[0]["counts_as_production_evidence"] is False


def test_panel_package_registry_validator_accepts_builtin_packages():
    report = validate_panel_registry(project_root=ROOT)

    assert report.status == "PASS"
    assert report.panels_checked == [
        "crc_301_msi",
        "crc_358_msi",
        "endometrial_29",
        "lung_329_pdl1",
        "lung_588_pdl1",
        "lung_methylation",
    ]
    assert report.errors == []


def test_panel_package_validator_rejects_invalid_qa_profile(tmp_path):
    panel_yaml = _write_minimal_panel_package(tmp_path, "bad_qa")
    qa_path = panel_yaml.parent / "qa.yaml"
    payload = yaml.safe_load(qa_path.read_text(encoding="utf-8"))
    payload["legacy_reference"]["privacy_checks"]["sample_id"] = "block"
    payload["current_output"]["source"] = "uploaded_file"
    qa_path.write_text(
        yaml.safe_dump(payload, sort_keys=False, allow_unicode=True),
        encoding="utf-8",
    )

    report = validate_panel_package_path(
        panel_yaml,
        project_root=tmp_path,
        panels_dir=tmp_path,
    )

    assert report.status == "FAIL"
    assert any(issue.code == "QA_LEGACY_SEVERITY_INVALID" for issue in report.errors)
    assert any(issue.code == "QA_CURRENT_SOURCE_INVALID" for issue in report.errors)


def test_panel_package_validator_rejects_processor_order_and_dependencies(tmp_path):
    panel_yaml = _write_minimal_panel_package(tmp_path, "bad_processors")
    payload = yaml.safe_load(panel_yaml.read_text(encoding="utf-8"))
    payload["processors"] = [
        "underlines_and_styles",
        "toc_refresh",
        "blank_page_cleanup",
    ]
    panel_yaml.write_text(
        yaml.safe_dump(payload, sort_keys=False, allow_unicode=True),
        encoding="utf-8",
    )

    report = validate_panel_package_path(
        panel_yaml,
        project_root=tmp_path,
        panels_dir=tmp_path / "panels",
    )

    codes = {issue.code for issue in report.errors}
    assert "PROCESSOR_ORDER_INVALID" in codes
    assert "PROCESSOR_DEPENDENCY_MISSING" in codes
    assert "PROCESSOR_DEPENDENCY_ORDER" in codes


def test_panel_package_validator_accepts_alias_lookup():
    report = validate_panel_package_path(
        ROOT / "panels" / "crc_358_msi" / "panel.yaml",
        project_root=ROOT,
    )
    alias_report = validate_panel_package("crc_358", project_root=ROOT)

    assert report.status == "PASS"
    assert alias_report.status == "PASS"


def test_rule_package_loader_records_crc_rule_provenance():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    report = load_rule_package(package)
    provenance = report.to_provenance()

    assert report.ok
    assert provenance["status"] == "PASS"
    rule_names = {item["rule_name"] for item in provenance["files"]}
    assert {
        "panel_rules",
        "report_text",
        "biomarkers",
        "guideline_tables",
        "drugs",
        "style",
    } <= rule_names
    report_text = next(item for item in provenance["files"] if item["rule_name"] == "report_text")
    assert report_text["version"] == "0.2.0"
    assert report_text["sha256"]


def test_report_text_rules_drive_tmb_msi_copy():
    report_data = ReportData(
        context={
            "tmb_value": "6.5",
            "tmb_status": "L",
            "sample_type": "组织",
            "msi_status": "MSS",
        }
    )
    text_rules = {
        "tmb_table_immuno_tips": "自定义TMB表格提示",
        "tmb_detail_interpretation": "自定义TMB章节科普",
        "tmb_drug_note": "自定义TMB药物清单",
        "msi_educational_tips": "自定义MSI表格提示",
        "msi_crc_interpretation": "自定义MSI章节解读",
    }

    applied = apply_report_text_rules(report_data, text_rules)

    assert applied["immuno_tips"] == "tmb_table_immuno_tips"
    assert report_data.get_field("immuno_tips") == "自定义TMB表格提示"
    assert report_data.get_field("tmb_detail_interpretation") == "自定义TMB章节科普"
    assert report_data.get_field("tmb_drug_note") == "自定义TMB药物清单"
    assert report_data.get_field("msi_tips") == "自定义MSI表格提示"
    assert report_data.get_field("msi_detail_interpretation") == "自定义MSI章节解读"


def test_crc_report_text_rule_contains_active_tmb_msi_copy():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    rule = PanelRuleEngine.from_panel_package(package).get("report_text")
    text_rules = collect_report_texts(rule)

    assert rule["version"] == "0.2.0"
    assert rule["status"] == "active"
    assert "TMB-H的肿瘤" in text_rules["tmb_table_immuno_tips"]
    assert "2020年6月，FDA批准帕博利珠单抗" in text_rules["tmb_detail_interpretation"]
    assert "MSI-H的实体瘤通常具有免疫原性" in text_rules["msi_educational_tips"]


def test_crc_style_rule_contains_active_table_tokens():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    engine = PanelRuleEngine.from_panel_package(package)
    rule = engine.get("style")
    style = rule["style"]

    assert rule["version"] == "0.3.0"
    assert rule["status"] == "active"
    assert style["part3_gene_analysis"] == {
        "fixed_domain_label": "基因结构域：",
        "default_label": "基因变异解析：",
        "merge_fixed_domain_with_first_narrative": True,
    }
    assert style["variant_summary_table"]["link_underline"] is True
    assert "未见突变" in style["variant_summary_table"]["plain_texts"]
    assert style["variant_detail_table"]["link_color"] == "0000FF"
    assert "未见突变" in style["variant_detail_table"]["plain_texts"]
    assert style["toc"]["section_font_color"] == "00C4D8"
    assert style["toc"]["content_top_padding_pt"] == 57
    assert style["biomarker_table"]["header_fill"] == "00B7C7"
    assert style["clinical_result_tables"]["border_color"] == "000000"
    assert style["clinical_result_tables"]["detected_result_color"] == "FF0000"
    assert "style" not in engine.get("panel_rules")


def test_report_generator_loads_panel_style_from_style_rule():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    style = ReportGenerator._load_panel_style_config(package)

    assert style["variant_summary_table"]["link_underline"] is True
    assert style["variant_detail_table"]["link_color"] == "0000FF"
    assert style["toc"]["section_font_size"] == 16
    assert style["biomarker_table"]["header_fill"] == "00B7C7"
    assert style["clinical_result_tables"]["border_color"] == "000000"


def test_panel_package_validator_rejects_missing_report_text_rule(tmp_path):
    panel_yaml = _write_minimal_panel_package(tmp_path, "bad_panel")
    (panel_yaml.parent / "rules" / "report_text.yaml").unlink()

    report = validate_panel_package_path(
        panel_yaml,
        project_root=tmp_path,
        panels_dir=tmp_path / "panels",
    )

    assert report.status == "FAIL"
    assert any(issue.code == "DECLARED_FILE_MISSING" for issue in report.errors)


def test_panel_package_validator_rejects_duplicate_rule_key(tmp_path):
    panel_yaml = _write_minimal_panel_package(tmp_path, "bad_panel")
    (panel_yaml.parent / "rules" / "report_text.yaml").write_text(
        "\n".join(
            [
                'schema_version: "1.0"',
                'panel_id: "bad_panel"',
                'rule_id: "report_text"',
                'version: "0.1.0"',
                "texts:",
                '  summary: "first"',
                "texts:",
                '  summary: "second"',
            ]
        ),
        encoding="utf-8",
    )

    report = validate_panel_package_path(
        panel_yaml,
        project_root=tmp_path,
        panels_dir=tmp_path / "panels",
    )

    assert report.status == "FAIL"
    assert any(issue.code == "RULE_DUPLICATE_KEY" for issue in report.errors)


def test_panel_package_validator_rejects_missing_declared_file(tmp_path):
    panel_yaml = _write_minimal_panel_package(tmp_path, "bad_panel")
    (panel_yaml.parent / "templates" / "standard.docx").unlink()

    report = validate_panel_package_path(
        panel_yaml,
        project_root=tmp_path,
        panels_dir=tmp_path / "panels",
    )

    assert report.status == "FAIL"
    assert {issue.code for issue in report.errors} >= {"DECLARED_FILE_MISSING"}


def test_panel_package_validator_rejects_directory_name_mismatch(tmp_path):
    panel_yaml = _write_minimal_panel_package(
        tmp_path,
        "actual_panel",
        directory_name="wrong_panel",
    )

    report = validate_panel_package_path(
        panel_yaml,
        project_root=tmp_path,
        panels_dir=tmp_path / "panels",
    )

    assert report.status == "FAIL"
    assert any(issue.code == "PANEL_DIR_NAME_MISMATCH" for issue in report.errors)


def test_panel_registry_validator_rejects_alias_collisions(tmp_path):
    _write_minimal_panel_package(tmp_path, "alpha_panel", aliases=["shared_alias"])
    _write_minimal_panel_package(tmp_path, "beta_panel", aliases=["shared_alias"])

    report = validate_panel_registry(
        project_root=tmp_path,
        panels_dir=tmp_path / "panels",
    )

    assert report.status == "FAIL"
    assert any(issue.code == "REGISTRY_ALIAS_COLLISION" for issue in report.errors)


def test_panel_registry_rejects_runtime_alias_collision():
    registry = PanelRegistry()
    registry.register("alpha_panel", object(), aliases=("shared_alias",))

    with pytest.raises(ValueError, match="shared_alias"):
        registry.register("beta_panel", object(), aliases=("shared_alias",))

    with pytest.raises(ValueError, match="alias"):
        registry.register("shared_alias", object())


def test_report_generator_blocks_invalid_panel_package_before_excel_read(tmp_path, monkeypatch):
    panel_yaml = _write_minimal_panel_package(tmp_path, "bad_panel")
    (panel_yaml.parent / "templates" / "standard.docx").unlink()
    package = PanelPackageLoader(
        project_root=tmp_path,
        panels_dir=tmp_path / "panels",
    ).load_file(panel_yaml)

    import reportgen.core.report_generator as report_generator_module

    monkeypatch.setattr(
        report_generator_module,
        "normalize_project_type",
        lambda project_type: "bad_panel",
    )
    monkeypatch.setattr(
        ReportGenerator,
        "_get_panel_registration",
        staticmethod(lambda project_type: SimpleNamespace(package=package)),
    )

    generator = ReportGenerator(config_dir=str(ROOT / "config"), log_level="ERROR")

    def fail_read(*_args, **_kwargs):
        raise AssertionError("excel should not be read when panel package is invalid")

    monkeypatch.setattr(generator.excel_reader, "read", fail_read)

    result = generator.generate(
        excel_file=str(tmp_path / "missing.xlsx"),
        template_file=str(tmp_path / "missing_template.docx"),
        output_dir=str(tmp_path / "out"),
        project_type="bad_panel",
    )

    assert result["success"] is False
    assert "Panel Package校验失败" in result["errors"][0]
    assert result["generation_id"] == "generation"
    failure_stage_file = Path(result["stage_results_file"])
    assert failure_stage_file.name == "generation.stage_results.json"
    assert failure_stage_file.exists()
    stage_names = [stage["name"] for stage in result["stage_results"]]
    assert stage_names == ["PanelResolutionStage", "PanelPackageValidationStage"]
    assert result["stage_results"][-1]["status"] == "FAIL"
    validation = result["panel_package_validation"]
    assert validation["status"] == "FAIL"
    assert any(issue["code"] == "DECLARED_FILE_MISSING" for issue in validation["issues"])


def test_panel_validate_cli_reports_invalid_package(tmp_path):
    panel_yaml = _write_minimal_panel_package(tmp_path, "bad_panel")
    (panel_yaml.parent / "templates" / "standard.docx").unlink()

    proc = subprocess.run(
        [
            sys.executable,
            "-m",
            "reportgen.cli",
            "panel",
            "validate",
            "--project-root",
            str(tmp_path),
        ],
        cwd=str(ROOT),
        text=True,
        capture_output=True,
        check=False,
    )

    assert proc.returncode == 1
    assert "DECLARED_FILE_MISSING" in proc.stdout


def test_panel_package_schema_rejects_missing_default_template():
    ok, errors = validate_panel_package_config(
        {
            "schema_version": "1.0",
            "panel_id": "bad_panel",
            "display_name": "Bad Panel",
            "default_template": "missing_template",
            "templates": [{"id": "declared", "file": "templates/a.docx"}],
        }
    )

    assert not ok
    assert any("default_template" in error for error in errors)


def test_panel_package_schema_rejects_cyclic_gene_symbol_aliases():
    ok, errors = validate_panel_package_config(
        {
            "schema_version": "1.0",
            "panel_id": "bad_alias_panel",
            "display_name": "Bad Alias Panel",
            "default_template": "declared",
            "templates": [{"id": "declared", "file": "templates/a.docx"}],
            "gene_symbol_aliases": {"GENE_A": "GENE_B", "GENE_B": "GENE_A"},
        }
    )

    assert not ok
    assert any("contains a cycle" in error for error in errors)


def test_panel_package_validator_rejects_missing_required_marker(tmp_path):
    panel_yaml = _write_minimal_panel_package(tmp_path, "marker_panel")
    template_path = tmp_path / "panels" / "marker_panel" / "templates" / "standard.docx"
    template_doc = Document()
    template_doc.add_paragraph("模板正文，无 Part 3 标记")
    template_doc.save(template_path)
    payload = yaml.safe_load(panel_yaml.read_text(encoding="utf-8"))
    payload["template_contract"]["required_markers"] = ["__PART3_MARKER__"]
    panel_yaml.write_text(
        yaml.safe_dump(payload, allow_unicode=True, sort_keys=False),
        encoding="utf-8",
    )

    report = validate_panel_package_path(
        panel_yaml,
        project_root=tmp_path,
        panels_dir=tmp_path / "panels",
    )

    assert report.ok is False
    assert any(issue.code == "TEMPLATE_MARKER_MISSING" for issue in report.errors)


def test_panel_package_validator_accepts_marker_split_across_runs(tmp_path):
    panel_yaml = _write_minimal_panel_package(tmp_path, "split_marker_panel")
    template_path = tmp_path / "panels" / "split_marker_panel" / "templates" / "standard.docx"
    template_doc = Document()
    paragraph = template_doc.add_paragraph()
    paragraph.add_run("__PART3")
    paragraph.add_run("_MARKER__")
    template_doc.save(template_path)
    payload = yaml.safe_load(panel_yaml.read_text(encoding="utf-8"))
    payload["template_contract"]["required_markers"] = ["__PART3_MARKER__"]
    panel_yaml.write_text(
        yaml.safe_dump(payload, allow_unicode=True, sort_keys=False),
        encoding="utf-8",
    )

    report = validate_panel_package_path(
        panel_yaml,
        project_root=tmp_path,
        panels_dir=tmp_path / "panels",
    )

    marker_errors = {
        "TEMPLATE_MARKER_MISSING",
        "TEMPLATE_MARKER_DUPLICATED",
        "TEMPLATE_MARKER_CHECK_FAILED",
    }
    assert marker_errors.isdisjoint({issue.code for issue in report.errors})


def test_template_contract_fails_when_declared_variable_is_removed(tmp_path):
    template_path = tmp_path / "missing_declared_variable.docx"
    doc = Document()
    doc.add_paragraph("患者：{{ patient_name }}")
    doc.save(template_path)

    report = TemplateRenderer(log_level="ERROR").validate_template_contract(
        str(template_path),
        {"patient_name": "张三", "sample_id": "LZ000001"},
        contract_spec={
            "required_variables": ["patient_name", "sample_id"],
        },
    )

    assert report["ok"] is False
    assert report["missing_paths"] == []
    assert "sample_id" in report["declared_contract"]["missing_required_variables"]


@pytest.mark.parametrize("marker_count", [0, 2])
def test_template_contract_requires_exactly_one_structural_marker(tmp_path, marker_count):
    template_path = tmp_path / f"part3_marker_{marker_count}.docx"
    doc = Document()
    doc.add_paragraph("报告正文")
    for _ in range(marker_count):
        doc.add_paragraph("__PART3_MARKER__")
    doc.save(template_path)

    report = TemplateRenderer(log_level="ERROR").validate_template_contract(
        str(template_path),
        {},
        contract_spec={"required_markers": ["__PART3_MARKER__"]},
    )

    declared = report["declared_contract"]
    assert report["ok"] is False
    assert declared["marker_counts"]["__PART3_MARKER__"] == marker_count
    if marker_count == 0:
        assert declared["missing_required_markers"] == ["__PART3_MARKER__"]
        assert declared["duplicate_required_markers"] == []
    else:
        assert declared["missing_required_markers"] == []
        assert declared["duplicate_required_markers"] == ["__PART3_MARKER__"]


@pytest.mark.parametrize("with_body_marker", [False, True])
def test_template_contract_rejects_part3_marker_in_table(
    tmp_path,
    with_body_marker,
):
    template_path = tmp_path / f"table_marker_{with_body_marker}.docx"
    doc = Document()
    if with_body_marker:
        doc.add_paragraph("__PART3_MARKER__")
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "__PART3_MARKER__"
    doc.save(template_path)

    report = TemplateRenderer(log_level="ERROR").validate_template_contract(
        str(template_path),
        {},
        contract_spec={"required_markers": ["__PART3_MARKER__"]},
    )

    declared = report["declared_contract"]
    assert report["ok"] is False
    if with_body_marker:
        assert declared["duplicate_required_markers"] == ["__PART3_MARKER__"]
        assert declared["marker_counts"]["__PART3_MARKER__"] == 2
    else:
        assert declared["missing_required_markers"] == ["__PART3_MARKER__"]
        assert declared["marker_counts"]["__PART3_MARKER__"] == 1


def test_part3_renderer_rejects_duplicate_marker_in_one_paragraph(tmp_path):
    template_path = tmp_path / "duplicate_part3_marker.docx"
    doc = Document()
    doc.add_paragraph("__PART3_MARKER__ / __PART3_MARKER__")
    doc.save(template_path)

    with pytest.raises(ValueError, match="expected exactly one"):
        TemplateRenderer(log_level="ERROR")._render_part3_formatted(
            str(template_path),
            {},
        )


def test_template_contract_fails_when_declared_table_shape_changes(tmp_path):
    template_path = tmp_path / "bad_variant_table.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=8)
    headers = [
        "基因名称",
        "转录本号",
        "染色体",
        "外显子",
        "位点",
        "突变 类型",
        "频率 (%)",
        "潜在获益靶向药物",
    ]
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    doc.save(template_path)

    report = TemplateRenderer(log_level="ERROR").validate_template_contract(
        str(template_path),
        {},
        contract_spec={
            "required_table_structures": {
                "variant_detail": {
                    "columns": 9,
                    "required_headers": [
                        "基因名称",
                        "转录本号",
                        "染色体",
                        "外显子",
                        "频率",
                        "潜在获益靶向药物",
                    ],
                }
            }
        },
    )

    assert report["ok"] is False
    errors = report["declared_contract"]["table_errors"]["variant_detail"]
    assert any("expected 9 columns" in error for error in errors)


def test_template_contract_extracts_docxtpl_table_row_loops(tmp_path):
    template_path = tmp_path / "loop_template.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "基因"
    table.rows[0].cells[1].text = "检测结果"
    table.rows[1].cells[0].text = "{%tr for row in immune_positive_results %}"
    table.rows[2].cells[0].text = "{{ row.gene }}"
    table.rows[2].cells[1].text = "{{ row.result }}"
    doc.save(template_path)

    report = TemplateRenderer(log_level="ERROR").validate_template_contract(
        str(template_path),
        {"immune_positive_results": [{"gene": "KRAS", "result": "检出"}]},
        contract_spec={
            "required_lists": ["immune_positive_results"],
        },
    )

    assert report["ok"] is True
    assert report["required_lists"] == ["immune_positive_results"]
    assert report["loop_row_fields"]["immune_positive_results"] == ["gene", "result"]
    assert "row.gene" not in report["required_paths"]


def test_report_generator_fails_bad_panel_template_before_rendering(tmp_path):
    template_path = tmp_path / "bad_panel_template.docx"
    doc = Document()
    doc.add_paragraph("患者：{{ patient_name }}")
    doc.save(template_path)
    output_dir = tmp_path / "out"
    excel_data = _excel(
        tmp_path,
        single_values={
            "患者姓名": "张三",
            "样本编号": "LZ000001",
            "报告日期": "2026.05.17",
            "MSI状态": "MSS",
        },
        variations=[],
    )
    excel_data.metadata["table_columns"] = {
        "Variations": [
            "Gene_Symbol",
            "cHGVS",
            "ExistInsmall358",
            "ExistIn552",
        ]
    }

    result = ReportGenerator(
        config_dir=str(ROOT / "config"),
        log_level="ERROR",
    ).generate(
        excel_file=str(tmp_path / "case.xlsx"),
        template_file=str(template_path),
        output_dir=str(output_dir),
        output_filename="should_not_render.docx",
        excel_data=excel_data,
        project_type="crc_358_msi",
        template_contract_mode="fail",
    )

    assert result["success"] is False
    assert result["output_file"] is None
    assert not (output_dir / "should_not_render.docx").exists()
    stage_by_name = {stage["name"]: stage for stage in result["stage_results"]}
    assert stage_by_name["TemplateContractStage"]["status"] == "FAIL"
    assert "TemplateRenderStage" not in stage_by_name
    declared = result["template_contract"]["declared_contract"]
    assert "sample_id" in declared["missing_required_variables"]
    assert "variant_detail" in declared["missing_required_tables"]


def test_report_generator_blocks_missing_part3_marker_even_in_warn_mode(tmp_path):
    template_path = tmp_path / "missing_part3_marker.docx"
    doc = Document()
    doc.add_paragraph("患者：{{ patient_name }}")
    doc.save(template_path)
    output_dir = tmp_path / "out"
    excel_data = _excel(
        tmp_path,
        single_values={
            "患者姓名": "测试患者",
            "样本编号": "TEST-PART3",
            "报告日期": "2026.07.10",
            "MSI状态": "MSS",
        },
        variations=[],
    )
    excel_data.metadata["table_columns"] = {
        "Variations": [
            "Gene_Symbol",
            "cHGVS",
            "ExistInsmall358",
            "ExistIn552",
        ]
    }

    result = ReportGenerator(
        config_dir=str(ROOT / "config"),
        log_level="ERROR",
    ).generate(
        excel_file=str(tmp_path / "case.xlsx"),
        template_file=str(template_path),
        output_dir=str(output_dir),
        output_filename="must_not_render.docx",
        excel_data=excel_data,
        project_type="crc_358_msi",
        template_contract_mode="warn",
    )

    assert result["success"] is False
    assert result["output_file"] is None
    assert not (output_dir / "must_not_render.docx").exists()
    stage_by_name = {stage["name"]: stage for stage in result["stage_results"]}
    assert stage_by_name["TemplateContractStage"]["status"] == "FAIL"
    assert stage_by_name["TemplateContractStage"]["issues"][0]["code"] == ("PART3_CONTRACT_FAILED")
    assert "TemplateRenderStage" not in stage_by_name


def test_report_generator_blocks_partial_part3_variant_coverage(
    tmp_path,
    monkeypatch,
):
    original = GeneKnowledgeProvider.build_all_gene_knowledge_sections

    def drop_last_section(self, variants, cancer_type=""):
        return original(self, variants, cancer_type)[:-1]

    monkeypatch.setattr(
        GeneKnowledgeProvider,
        "build_all_gene_knowledge_sections",
        drop_last_section,
    )
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    output_dir = tmp_path / "out"
    result = ReportGenerator(
        config_dir=str(ROOT / "config"),
        log_level="ERROR",
    ).generate(
        excel_file=str(tmp_path / "case.xlsx"),
        template_file=str(package.resolve_template_file()),
        output_dir=str(output_dir),
        output_filename="partial_part3.docx",
        excel_data=_excel(
            tmp_path,
            single_values={
                "患者姓名": "测试患者",
                "样本编号": "TEST-PART3-COVERAGE",
                "报告日期": "2026.07.10",
                "癌种": "结直肠癌",
                "MSI状态": "MSS",
            },
            variations=[
                {
                    "ExistIn552": "Ⅱ类",
                    "ExistInsmall358": 1,
                    "Gene_Symbol": "TP53",
                    "cHGVS": "c.844C>T",
                    "pHGVS_S": "p.R282W",
                    "Function": "Missense",
                    "Freq(%)": 67.29,
                },
                {
                    "ExistIn552": "Ⅱ类",
                    "ExistInsmall358": 1,
                    "Gene_Symbol": "APC",
                    "cHGVS": "c.4348C>T",
                    "pHGVS_S": "p.R1450*",
                    "Function": "Nonsense",
                    "Freq(%)": 41.12,
                },
            ],
        ),
        project_type="crc_358_msi",
        template_contract_mode="warn",
    )

    assert result["success"] is False
    assert result["output_file"] is None
    assert not (output_dir / "partial_part3.docx").exists()
    stage = next(row for row in result["stage_results"] if row["name"] == "TemplateContractStage")
    assert stage["status"] == "FAIL"
    assert "variant coverage mismatch" in stage["issues"][0]["message"]


def test_msi_percentage_label_conflict_is_warned(tmp_path):
    from app.services.reportgen_bridge import ReportGenBridge

    excel_data = _excel(
        tmp_path,
        single_values={"MSI状态": "MSS", "MSI百分比": 40.0},
    )
    bridge = ReportGenBridge(
        config_dir=str(ROOT / "config"),
        template_dir=str(ROOT / "templates"),
    )

    warnings = bridge.validate_excel_data(excel_data)
    assert any(w.get("field") == "msi_status" and w.get("level") == "warning" for w in warnings)


def test_batch_options_accept_forced_project_type():
    opts = BatchValidateOptions(
        inputs=["dummy.xlsx"],
        project_type="crc_358_msi",
        project_name="结直肠癌358基因+MSI",
    )
    assert opts.project_type == "crc_358_msi"


def test_hla_expected_table_matches_default_hidden_policy(tmp_path):
    excel_data = _excel(tmp_path, tables={"HLA": [{"Locus": "HLA-A", "Type1": "01:01"}]})

    assert _expected_tables_from_excel(excel_data, show_hla_table=False)["hla"] is False
    assert _expected_tables_from_excel(excel_data, show_hla_table=True)["hla"] is True


def test_consultation_phone_is_enabled_in_config():
    settings = yaml.safe_load((ROOT / "config" / "settings.yaml").read_text(encoding="utf-8"))
    report_content = settings.get("report_content") or {}
    assert report_content.get("consultation_phone") == "022-87190699"
    assert report_content.get("patient_letter", {}).get("after_greeting")
    assert report_content.get("part3_reading_blocks")
    assert report_content.get("tail_content", {}).get("paragraphs")


def test_configured_letter_and_tail_content_are_inserted(tmp_path):
    docx_path = tmp_path / "configured_content.docx"
    doc = Document()
    doc.add_paragraph("您好！")
    doc.add_paragraph("现代医学已经证明，肿瘤发生与基因异常有关。")
    doc.add_paragraph("致您的一封信")
    doc.add_paragraph("部分基因与药物对应关系，目前仅限于临床试验科学研究阶段。")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._apply_report_content_fixes(
        str(docx_path),
        {
            "project_name": "结直肠癌358基因+MSI",
            "report_content": {
                "patient_letter": {
                    "after_greeting": ["项目：{project_name}"],
                    "after_modern_medicine": ["固定结尾段"],
                },
                "tail_content": {
                    "anchor_text": "部分基因与药物对应关系，目前仅限于临床试验科学研究阶段",
                    "marker_text": "尾页标题",
                    "paragraphs": [
                        {"text": "尾页标题", "bold": True, "page_break_before": True},
                        {
                            "text": "尾页正文",
                            "first_line_indent_cm": 0.7,
                            "line_spacing_multiple": 1.3,
                            "space_after_pt": 5,
                        },
                    ],
                },
            },
        },
    )

    doc = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "项目：结直肠癌358基因+MSI" in text
    assert "固定结尾段" in text
    assert "尾页标题" in text
    assert "尾页正文" in text
    tail_paragraph = next(p for p in doc.paragraphs if p.text == "尾页正文")
    ppr = tail_paragraph._p.get_or_add_pPr()
    assert ppr.find(qn("w:ind")).get(qn("w:firstLine")) == "396"
    spacing = ppr.find(qn("w:spacing"))
    assert spacing.get(qn("w:line")) == "312"
    assert spacing.get(qn("w:after")) == "100"


def test_configured_letter_content_is_inserted_inside_table_cells(tmp_path):
    docx_path = tmp_path / "configured_table_letter.docx"
    doc = Document()
    doc.add_paragraph("致您的一封信")
    table = doc.add_table(rows=2, cols=1)
    table.rows[0].cells[0].text = "您好！"
    table.rows[1].cells[0].text = "现代医学已经证明，肿瘤发生与基因异常有关。"
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._apply_report_content_fixes(
        str(docx_path),
        {
            "project_name": "结直肠癌358基因+MSI",
            "report_content": {
                "patient_letter": {
                    "after_greeting": ["项目：{project_name}"],
                    "after_modern_medicine": ["固定补充段一", "固定补充段二"],
                },
            },
        },
    )

    doc = Document(docx_path)
    text = "\n".join(
        [p.text for p in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    assert "项目：结直肠癌358基因+MSI" in text
    assert "固定补充段一" in text
    assert "固定补充段二" in text


def test_report_content_fixes_project_code_is_not_lz_specific(tmp_path):
    docx_path = tmp_path / "project_code.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "项目编码"
    table.rows[0].cells[1].text = "MLJY-MLB2509307001"
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._apply_report_content_fixes(
        str(docx_path),
        {
            "sample_id": "MLJY-MLB2509307001",
            "report_number": "MLJY-MLB2509307001",
        },
    )

    rendered = Document(docx_path)
    assert rendered.tables[0].rows[0].cells[1].text == "MLB2509307001"


def test_configured_tail_contact_block_uses_template_qr(tmp_path):
    from zipfile import ZipFile

    docx_path = tmp_path / "tail_contact.docx"
    doc = Document()
    doc.add_paragraph("部分基因与药物对应关系，目前仅限于临床试验科学研究阶段。")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._apply_report_content_fixes(
        str(docx_path),
        {
            "report_content": {
                "tail_content": {
                    "anchor_text": "部分基因与药物对应关系，目前仅限于临床试验科学研究阶段",
                    "marker_text": "脉络医学检验简介",
                    "paragraphs": [
                        {"text": "脉络医学检验简介", "page_break_before": True},
                        {"text": "简介正文"},
                    ],
                    "contact_block": {
                        "enabled": True,
                        "lines": ["地址：测试地址", "电话：022-87190699"],
                        "qr_template_media": "word/media/image37.png",
                        "space_before_pt": 12,
                    },
                },
            },
        },
        str(ROOT / "templates/aligned_template_with_cnv_fusion_hla_FIXED.docx"),
    )

    doc = Document(docx_path)
    text = "\n".join(
        [p.text for p in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    assert "脉络医学检验简介" in text
    assert "地址：测试地址" in text
    assert "电话：022-87190699" in text
    assert len(doc.inline_shapes) == 1
    with ZipFile(docx_path) as zf:
        assert any(name.startswith("word/media/") for name in zf.namelist())


def test_excel_reader_extracts_lowercase_lz_sample_id_from_filename(tmp_path):
    path = tmp_path / "上传使用Excel表：lz000001.xlsx"
    path.write_bytes(b"placeholder")

    sample_id = ExcelReader(
        config_dir=str(ROOT / "config"),
        log_level="ERROR",
    )._extract_sample_id_from_filename(str(path))

    assert sample_id == "LZ000001"


def test_patient_info_can_be_loaded_from_external_runtime_path(monkeypatch, tmp_path):
    patient_info = tmp_path / "patient_info.yaml"
    patient_info.write_text(
        yaml.safe_dump(
            {
                "defaults": {"issuer": "签发人"},
                "project_info": {"project_name": "结直肠癌358基因+MSI"},
                "patients": {
                    "LZ000001": {
                        "patient_name": "测试患者",
                        "report_date": "2026-01-01",
                    }
                },
            },
            allow_unicode=True,
            sort_keys=False,
        ),
        encoding="utf-8",
    )
    monkeypatch.setenv("REPORTGEN_PATIENT_INFO_PATH", str(patient_info))

    info = ConfigLoader(config_dir=str(ROOT / "config"), log_level="ERROR").load_patient_info(
        "LZ000001"
    )

    assert info["patient_name"] == "测试患者"
    assert info["report_date"] == "2026-01-01"
    assert info["issuer"] == "签发人"
    assert info["project_name"] == "结直肠癌358基因+MSI"


def test_reviewed_atm_drug_override_matches_final_report_scope():
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")

    benefit, caution, score = mapper._lookup_targeted_drugs_for_variant(
        "ATM",
        c_point="c.6874C>T",
        p_point="p.Q2292*",
        variant_level="Ⅱ类",
        cancer_type="乙状结肠癌",
    )

    assert score == 100.0
    assert "奥拉帕利+帕博利珠单抗（C）" in benefit
    assert "Tuvusertib" not in benefit
    assert "Peposertib" not in benefit
    assert caution == "--"


def test_drug_analysis_is_limited_to_final_displayed_drugs():
    provider = GeneKnowledgeProvider({"enabled": False})
    provider._loaded = True
    provider._drug_full_cache = {
        "ATM": [
            {
                "type": "benefit",
                "drug": ("奥拉帕利（Olaparib）、芦卡帕利（Rucaparib）、Tuvusertib+Peposertib"),
                "relation": "奥拉帕利相关说明。Tuvusertib联合治疗说明。",
                "clinical": "芦卡帕利相关说明。Peposertib与Tuvusertib研究说明。",
            }
        ]
    }

    sections = provider.build_drug_analysis_sections(
        [
            {
                "gene": "ATM",
                "cHGVS": "c.6874C>T",
                "pHGVS": "p.Q2292*",
                "benefit_drugs": "奥拉帕利（C）\n芦卡帕利（C）",
                "caution_drugs": "--",
            }
        ]
    )

    assert len(sections) == 1
    section = sections[0]
    assert "奥拉帕利" in section["drug_name"]
    assert "芦卡帕利" in section["drug_name"]
    assert "Tuvusertib" not in section["drug_name"]
    assert "Peposertib" not in section["clinical"]


def test_drug_analysis_respects_variant_specific_p_point():
    provider = GeneKnowledgeProvider({"enabled": False})
    provider._loaded = True
    provider._drug_full_cache = {
        "TP53": [
            {
                "type": "benefit",
                "drug": "Eprenetapopt",
                "p_point": "",
                "relation": "general TP53 relation",
                "clinical": "general TP53 clinical",
            },
            {
                "type": "benefit",
                "drug": "Eprenetapopt、PC14586",
                "p_point": "p.Y220C",
                "relation": "p.Y220C specific relation",
                "clinical": "p.Y220C specific clinical",
            },
        ]
    }

    sections = provider.build_drug_analysis_sections(
        [
            {
                "gene": "TP53",
                "cHGVS": "c.844C>T",
                "pHGVS": "p.R282W",
                "benefit_drugs": "Eprenetapopt（C）",
                "caution_drugs": "--",
            }
        ]
    )

    assert len(sections) == 1
    # relation 现在会被自动前置「该样本检出…突变。」变异描述开头，故用 endswith
    # 断言匹配到的是 general（非 p.Y220C 特异）relation。
    assert sections[0]["relation"].endswith("general TP53 relation")
    assert "Y220C" not in sections[0]["clinical"]


def test_drug_analysis_matches_kras_wildcard_p_point():
    provider = GeneKnowledgeProvider({"enabled": False})
    provider._loaded = True
    provider._drug_full_cache = {
        "KRAS": [
            {
                "type": "benefit",
                "drug": "索托拉西布",
                "p_point": "p.G12C",
                "relation": "G12C only",
                "clinical": "G12C clinical",
            },
            {
                "type": "benefit",
                "drug": "司美替尼",
                "p_point": "p.G12X(X为除C、D外的任何氨基酸)",
                "relation": "G12X wildcard",
                "clinical": "G12X clinical",
            },
        ]
    }

    sections = provider.build_drug_analysis_sections(
        [
            {
                "gene": "KRAS",
                "cHGVS": "c.34G>A",
                "pHGVS": "p.G12S",
                "benefit_drugs": "索托拉西布（C）\n司美替尼（C）",
                "caution_drugs": "--",
            }
        ]
    )

    assert len(sections) == 1
    assert sections[0]["drug_name"] == "司美替尼"
    # relation 被自动前置变异描述开头，用 endswith 断言匹配到通配 p.G12X relation。
    assert sections[0]["relation"].endswith("G12X wildcard")


def test_gene_knowledge_uses_reviewed_columns_without_intro_domain_tail():
    provider = GeneKnowledgeProvider({"enabled": False})
    provider._loaded = True
    provider._gene_intro_cache = {
        "KRAS": provider._strip_intro_domain_tail(
            "KRAS",
            "KRAS基因简介。\nKRAS基因编码的蛋白全长为189个氨基酸，"
            "主要包含Hypervariable region（166-185位氨基酸）。",
        )
    }
    provider._gene_analysis_cache = {"KRAS": "generic KRAS analysis"}
    provider._reviewed_gene_analysis_cache = {
        "KRAS": {
            "domain_text": (
                "KRAS基因编码的蛋白全长为189个氨基酸，主要包含RAS结构域（1-166位氨基酸）。"
            ),
            "expert_text": "据OncoKB/JAXCKB数据库记载，该为已知激活突变，对蛋白功能有重要影响。",
            "cancer_text": "KRAS突变会激活下游信号通路，且已被证实是anti-EGFR抗体药物耐药的标志。",
        }
    }

    section = provider.build_gene_knowledge_section(
        gene="KRAS",
        c_hgvs="c.34G>A",
        p_hgvs="p.G12S",
        frequency=46.29,
        mutation_type="Missense",
        has_drug=True,
    )

    assert section["intro"] == "KRAS基因简介。"
    assert "Hypervariable region" not in section["intro"]
    assert "RAS结构域" in section["mutation_analysis"]
    assert "已知激活突变" in section["mutation_analysis"]
    assert "疾病的发生发展及用药相关" in section["mutation_analysis"]


def test_signature_placeholder_is_removed_without_image(tmp_path):
    docx_path = tmp_path / "signature.docx"
    doc = Document()
    doc.add_paragraph("签名：__SIG_IMG__")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._render_signature_placeholder(str(docx_path), {})

    assert "__SIG_IMG__" not in "\n".join(p.text for p in Document(docx_path).paragraphs)


def test_signature_layout_keeps_report_date_on_signature_line(tmp_path):
    from shutil import copyfile

    source = ROOT / "templates/aligned_template_with_cnv_fusion_hla_FIXED.docx"
    docx_path = tmp_path / "signature_layout.docx"
    copyfile(source, docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._render_signature_placeholder(str(docx_path), {})
    renderer._apply_report_content_fixes(str(docx_path), {"report_date": "2026-04-26"})
    renderer._normalize_signature_layout(str(docx_path), {"report_date": "2026-04-26"})

    paragraphs = [p.text.strip() for p in Document(docx_path).paragraphs if p.text.strip()]
    signature_lines = [p for p in paragraphs if p.startswith("检测者：")]
    assert signature_lines
    assert any("报告日期：2026.04.26" in p for p in signature_lines)
    assert "报告日期：2026.04.26" not in paragraphs
    signature_para = next(
        p for p in Document(docx_path).paragraphs if p.text.strip().startswith("检测者：")
    )
    assert "w:keepLines" in signature_para._p.xml


def test_signature_layout_keeps_detection_explanation_with_signature(tmp_path):
    docx_path = tmp_path / "signature_block.docx"
    doc = Document()
    doc.add_paragraph("前文")
    doc.add_paragraph("4. 检测结果说明")
    for text in ("说明一", "说明二", "说明三", "说明四"):
        doc.add_paragraph(text)
    doc.add_paragraph("检测者：  审核者：  报告日期：2026.07.14")
    doc.add_paragraph("第三部分：基因变异及相应靶向/免疫药物解析")
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._normalize_signature_layout(str(docx_path), {})
    renderer._normalize_signature_layout(str(docx_path), {})

    paragraphs = Document(docx_path).paragraphs
    block = paragraphs[1:6]
    signature = paragraphs[6]
    assert all("w:keepNext" in paragraph._p.xml for paragraph in block)
    assert "w:keepNext" not in signature._p.xml
    assert "w:keepLines" in signature._p.xml


def test_detector_and_reviewer_signature_images_are_context_driven(tmp_path):
    from PIL import Image

    def write_png(path: Path, color: tuple[int, int, int]) -> None:
        Image.new("RGB", (12, 6), color).save(path)

    old_detector = tmp_path / "old_detector.png"
    old_reviewer = tmp_path / "old_reviewer.png"
    new_detector = tmp_path / "new_detector.png"
    new_reviewer = tmp_path / "new_reviewer.png"
    write_png(old_detector, (255, 0, 0))
    write_png(old_reviewer, (0, 255, 0))
    write_png(new_detector, (0, 0, 255))
    write_png(new_reviewer, (255, 255, 0))

    docx_path = tmp_path / "signature_dynamic.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_picture(str(old_detector), width=Cm(2))
    paragraph.add_run(" ")
    paragraph.add_run().add_picture(str(old_reviewer), width=Cm(2))
    doc.add_paragraph(
        "检测者：                    审核者：                    报告日期：2026.05.24"
    )
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._replace_signature_anchor_images(
        str(docx_path),
        {
            "detector_signature_image_path": str(new_detector),
            "reviewer_signature_image_path": str(new_reviewer),
        },
    )

    with ZipFile(docx_path) as zf:
        assert zf.read("word/media/reportgen_signature_detector.png") == new_detector.read_bytes()
        assert zf.read("word/media/reportgen_signature_reviewer.png") == new_reviewer.read_bytes()
        document_xml = zf.read("word/document.xml").decode("utf-8")
        rels_xml = zf.read("word/_rels/document.xml.rels").decode("utf-8")
    assert "reportgen_signature_detector.png" in rels_xml
    assert "reportgen_signature_reviewer.png" in rels_xml
    assert "reportgen_signature_detector" not in document_xml


def test_signature_processors_replace_uploaded_images_after_layout(tmp_path):
    from PIL import Image

    def write_png(path: Path, color: tuple[int, int, int]) -> None:
        Image.new("RGB", (12, 6), color).save(path)

    old_detector = tmp_path / "old_detector.png"
    old_reviewer = tmp_path / "old_reviewer.png"
    new_detector = tmp_path / "new_detector.png"
    new_reviewer = tmp_path / "new_reviewer.png"
    write_png(old_detector, (255, 0, 0))
    write_png(old_reviewer, (0, 255, 0))
    write_png(new_detector, (0, 0, 255))
    write_png(new_reviewer, (255, 255, 0))

    docx_path = tmp_path / "signature_processor_chain.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_picture(str(old_detector), width=Cm(2))
    paragraph.add_run(" ")
    paragraph.add_run().add_picture(str(old_reviewer), width=Cm(2))
    doc.add_paragraph(
        "检测者：                    审核者：                    报告日期：2026.05.24"
    )
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    context = {
        "detector_signature_image_path": str(new_detector),
        "reviewer_signature_image_path": str(new_reviewer),
        "report_date": "2026.05.24",
    }
    renderer._run_post_render_processors(
        str(docx_path),
        context,
        str(docx_path),
        processor_names=["signature_placeholder", "signature_layout"],
    )
    # Inline signatures are rendered as the final step (underlines_and_styles);
    # invoke directly to mirror that processor.
    renderer._render_inline_signatures(str(docx_path), context)

    rendered = Document(str(docx_path))
    label_inline = _sig_label_inline_drawings(rendered)
    assert label_inline is not None and len(label_inline) == 2

    with ZipFile(docx_path) as zf:
        media = {zf.read(name) for name in zf.namelist() if name.startswith("word/media/")}
        document_xml = zf.read("word/document.xml").decode("utf-8")

    assert new_detector.read_bytes() in media
    assert new_reviewer.read_bytes() in media
    assert "检测者" in document_xml
    assert "审核者" in document_xml


def test_template_signature_images_are_removed_without_context_paths(tmp_path):
    from PIL import Image

    def write_png(path: Path, color: tuple[int, int, int]) -> None:
        Image.new("RGB", (12, 6), color).save(path)

    old_detector = tmp_path / "old_detector.png"
    old_reviewer = tmp_path / "old_reviewer.png"
    write_png(old_detector, (255, 0, 0))
    write_png(old_reviewer, (0, 255, 0))

    docx_path = tmp_path / "signature_blank.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_picture(str(old_detector), width=Cm(2))
    paragraph.add_run(" ")
    paragraph.add_run().add_picture(str(old_reviewer), width=Cm(2))
    doc.add_paragraph(
        "检测者：                    审核者：                    报告日期：2026.05.24"
    )
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._replace_signature_anchor_images(str(docx_path), {})

    with ZipFile(docx_path) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")

    assert "<w:drawing>" not in document_xml
    assert "检测者" in document_xml
    assert "审核者" in document_xml


def test_signature_library_resolves_names_to_paths(tmp_path):
    detector_png = tmp_path / "detector.png"
    reviewer_png = tmp_path / "reviewer.png"
    detector_png.write_bytes(b"detector")
    reviewer_png.write_bytes(b"reviewer")
    (tmp_path / "signatures.yaml").write_text(
        f"""
detector:
  张三: {detector_png.name}
reviewer:
  - name: 李四
    path: {reviewer_png}
""",
        encoding="utf-8",
    )

    assert signature_options(tmp_path, "detector") == ["张三"]
    assert resolve_signature_path(tmp_path, "detector", " 张三 ") == str(detector_png.resolve())
    assert resolve_signature_path(tmp_path, "reviewer", "李四") == str(reviewer_png)
    assert resolve_signature_path(tmp_path, "reviewer", "王五") == ""


def test_signature_library_prefers_external_runtime_storage(tmp_path, monkeypatch):
    project = tmp_path / "release"
    config_dir = project / "config"
    runtime_root = tmp_path / "runtime-storage"
    runtime_image = runtime_root / "signatures" / "registry" / "detector.png"
    config_dir.mkdir(parents=True)
    runtime_image.parent.mkdir(parents=True)
    runtime_image.write_bytes(b"runtime-signature")
    (config_dir / "signatures.yaml").write_text(
        "detector:\n  TEST: storage/signatures/registry/detector.png\n",
        encoding="utf-8",
    )
    monkeypatch.setenv("RG_WEB_STORAGE_ROOT", str(runtime_root))

    assert resolve_signature_path(config_dir, "detector", "TEST") == str(runtime_image.resolve())


def test_report_generator_fills_signature_paths_from_library(tmp_path):
    detector_png = tmp_path / "detector.png"
    reviewer_png = tmp_path / "reviewer.png"
    detector_png.write_bytes(b"detector")
    reviewer_png.write_bytes(b"reviewer")
    (tmp_path / "signatures.yaml").write_text(
        f"""
detector:
  张三: {detector_png}
reviewer:
  李四: {reviewer_png}
""",
        encoding="utf-8",
    )
    report_data = ReportData(
        context={
            "issuer": "张三",
            "reviewer": "李四",
            "reviewer_signature_image_path": "/explicit/reviewer.png",
        }
    )
    generator = object.__new__(ReportGenerator)
    generator.config_dir = str(tmp_path)

    generator._resolve_signature_image_fields(report_data)

    assert report_data.get_field("detector_signature_image_path") == str(detector_png)
    assert report_data.get_field("reviewer_signature_image_path") == "/explicit/reviewer.png"


def test_clinical_schema_exposes_signature_people_as_editable_selects(monkeypatch):
    monkeypatch.setattr(
        clinical_info_service,
        "_load_mapping_yaml",
        lambda: {
            "single_values": {
                "issuer": {
                    "synonyms": ["检测者"],
                    "type": "string",
                    "default_value": "",
                    "description": "检测者",
                },
                "reviewer": {
                    "synonyms": ["审核者"],
                    "type": "string",
                    "default_value": "",
                    "description": "审核者",
                },
                "detector_signature_image_path": {
                    "synonyms": ["检测者签名图片路径"],
                    "type": "string",
                    "default_value": None,
                    "description": "检测者签名图片",
                },
                "reviewer_signature_image_path": {
                    "synonyms": ["审核者签名图片路径"],
                    "type": "string",
                    "default_value": None,
                    "description": "审核者签名图片",
                },
            }
        },
    )
    monkeypatch.setattr(
        clinical_info_service,
        "signature_options",
        lambda _config_dir, role: ["张三"] if role == "detector" else ["李四"],
    )

    schema = clinical_info_service.get_clinical_form_schema("crc_358_msi")
    fields = {field.key: field for group in schema.groups for field in group.fields}

    assert fields["issuer"].ui.component == "select"
    assert fields["issuer"].ui.options == ["张三"]
    assert fields["issuer"].ui.allow_create is True
    assert fields["reviewer"].ui.options == ["李四"]
    assert fields["detector_signature_image_path"].ui.component == "file-upload"
    assert fields["reviewer_signature_image_path"].ui.component == "file-upload"


def test_crc358_clinical_schema_does_not_require_receive_date(monkeypatch):
    monkeypatch.setattr(
        clinical_info_service,
        "_load_mapping_yaml",
        lambda: {
            "single_values": {
                "receive_date": {
                    "synonyms": ["收样日期"],
                    "type": "date",
                    "required": False,
                    "description": "样本接收/送检日期",
                },
                "report_date": {
                    "synonyms": ["报告日期"],
                    "type": "date",
                    "required": True,
                    "description": "报告生成日期",
                },
            }
        },
    )

    schema = clinical_info_service.get_clinical_form_schema("crc_358_msi")
    fields = {field.key: field for group in schema.groups for field in group.fields}

    assert fields["receive_date"].required is False
    assert fields["report_date"].required is True


def test_template_renderer_removes_explicit_underlines(tmp_path):
    from zipfile import ZipFile

    docx_path = tmp_path / "underlines.docx"
    doc = Document()
    doc.add_paragraph().add_run("姓名").font.underline = True
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("TP53").font.underline = True
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._remove_template_underlines(str(docx_path))

    rendered = Document(docx_path)
    assert rendered.paragraphs[0].runs[0].font.underline is None
    assert rendered.tables[0].cell(0, 0).paragraphs[0].runs[0].font.underline is None
    with ZipFile(docx_path) as zf:
        assert b"<w:u" not in zf.read("word/document.xml")


def test_detection_content_fill_underline_is_restored(tmp_path):
    docx_path = tmp_path / "detection_content.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("对委托人").font.underline = True
    paragraph.add_run("    组织    ").font.underline = True
    paragraph.add_run(
        "的 DNA 样本微卫星不稳定性（MSI）以及与肿瘤密切相关的358个基因进行检测："
    ).font.underline = True
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._remove_template_underlines(str(docx_path))
    renderer._restore_detection_content_underlines(str(docx_path))

    runs = Document(docx_path).paragraphs[0].runs
    assert runs[0].font.underline is None
    assert runs[1].font.underline is True
    assert runs[2].font.underline is None


def test_patient_letter_fill_underlines_are_restored(tmp_path):
    docx_path = tmp_path / "patient_letter.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("尊敬的").font.underline = True
    paragraph.add_run("     ").font.underline = True
    paragraph.add_run("苏雨起").font.underline = True
    paragraph.add_run("     ").font.underline = True
    paragraph.add_run("先生：").font.underline = True
    paragraph.add_run("感谢您选择本机构为您精心定").font.underline = True
    paragraph.add_run("的").font.underline = True
    paragraph.add_run("      ").font.underline = True
    paragraph.add_run("结直肠癌358基因+MSI      ").font.underline = True
    paragraph.add_run("检测项目。").font.underline = True
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._remove_template_underlines(str(docx_path))
    renderer._restore_patient_letter_fill_underlines(str(docx_path))

    runs = Document(docx_path).paragraphs[0].runs
    assert runs[0].font.underline is None
    assert runs[1].font.underline is True
    assert runs[2].font.underline is True
    assert runs[3].font.underline is True
    assert runs[4].font.underline is None
    assert runs[5].font.underline is None
    assert runs[6].font.underline is None
    assert runs[7].font.underline is True
    assert runs[8].font.underline is True
    assert runs[9].font.underline is None


def test_pdf_footer_page_number_is_used_for_static_toc_detection():
    page_text = """
姓名：苏雨起                 科技服务人类健康

3. 免疫治疗疗效评估
3.1 肿瘤突变负荷（TMB）水平提示



                                    12
"""

    assert TemplateRenderer._extract_pdf_footer_page_number(page_text) == 12
    assert TemplateRenderer._extract_pdf_footer_page_number("正文\\n没有页脚") is None


def test_report_generator_normalizes_duplicate_project_name_suffix():
    assert (
        ReportGenerator._normalize_project_name("结直肠癌358基因+MSI检测项目", "crc_358_msi")
        == "结直肠癌358基因+MSI"
    )
    assert (
        ReportGenerator._normalize_project_name(
            "结直肠癌358基因+MSI检测项目检测项目", "crc_358_msi"
        )
        == "结直肠癌358基因+MSI"
    )


def test_variant_summary_table_restores_reviewed_link_style(tmp_path):
    docx_path = tmp_path / "variant_summary.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=4)
    headers = [
        "基因",
        "突变位点",
        "潜在获益靶向药物\n（证据等级）",
        "可能耐药或慎重药物\n（证据等级）",
    ]
    values = ["KRAS", "c.34G>A,\np.G12S", "司美替尼（C）", "西妥昔单抗（A）"]
    for idx, value in enumerate(headers):
        table.rows[0].cells[idx].text = value
    for idx, value in enumerate(values):
        table.rows[1].cells[idx].text = value
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._remove_template_underlines(str(docx_path))
    renderer._restore_variant_summary_table_style(str(docx_path))

    doc = Document(docx_path)
    table = doc.tables[0]
    gene_run = table.rows[1].cells[0].paragraphs[0].runs[0]
    site_run = table.rows[1].cells[1].paragraphs[0].runs[0]
    drug_run = table.rows[1].cells[2].paragraphs[0].runs[0]
    resist_run = table.rows[1].cells[3].paragraphs[0].runs[0]

    assert gene_run.font.underline is True
    assert str(gene_run.font.color.rgb) == "0000FF"
    assert site_run.font.underline is False
    assert str(site_run.font.color.rgb) == "000000"
    assert drug_run.font.underline is True
    assert str(drug_run.font.color.rgb) == "0000FF"
    assert resist_run.font.underline is True
    assert str(resist_run.font.color.rgb) == "0000FF"


def test_variant_summary_table_can_disable_link_underlines_from_panel_style(tmp_path):
    docx_path = tmp_path / "variant_summary_plain.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=4)
    headers = [
        "基因",
        "突变位点",
        "潜在获益靶向药物\n（证据等级）",
        "可能耐药或慎重药物\n（证据等级）",
    ]
    values = ["KRAS", "c.34G>A,\np.G12S", "司美替尼（C）", "西妥昔单抗（A）"]
    for idx, value in enumerate(headers):
        table.rows[0].cells[idx].text = value
    for idx, value in enumerate(values):
        table.rows[1].cells[idx].text = value
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._restore_variant_summary_table_style(
        str(docx_path),
        {
            "panel_style": {
                "variant_summary_table": {
                    "link_underline": False,
                    "link_color": "000000",
                }
            }
        },
    )

    table = Document(docx_path).tables[0]
    gene_run = table.rows[1].cells[0].paragraphs[0].runs[0]
    drug_run = table.rows[1].cells[2].paragraphs[0].runs[0]
    assert gene_run.font.underline is False
    assert str(gene_run.font.color.rgb) == "000000"
    assert drug_run.font.underline is False
    assert str(drug_run.font.color.rgb) == "000000"


def test_variant_summary_table_keeps_undetected_text_plain(tmp_path):
    docx_path = tmp_path / "variant_summary_undetected.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=4)
    headers = [
        "基因",
        "突变位点",
        "潜在获益靶向药物\n（证据等级）",
        "可能耐药或慎重药物\n（证据等级）",
    ]
    values = ["FBXW7", "未见突变", "--", "--"]
    for idx, value in enumerate(headers):
        table.rows[0].cells[idx].text = value
    for idx, value in enumerate(values):
        table.rows[1].cells[idx].text = value
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._restore_variant_summary_table_style(str(docx_path))

    table = Document(docx_path).tables[0]
    undetected_run = table.rows[1].cells[1].paragraphs[0].runs[0]
    benefit_run = table.rows[1].cells[2].paragraphs[0].runs[0]

    assert undetected_run.text == "未见突变"
    assert undetected_run.font.underline is False
    assert str(undetected_run.font.color.rgb) == "000000"
    assert benefit_run.font.underline is False
    assert str(benefit_run.font.color.rgb) == "000000"


def test_variant_detail_table_keeps_undetected_drug_text_plain(tmp_path):
    docx_path = tmp_path / "variant_detail_undetected.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=9)
    row0 = [
        "基因名称",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "靶向药物信息",
        "靶向药物信息",
    ]
    row1 = [
        "基因名称",
        "转录本号",
        "染色体",
        "外显子",
        "位点",
        "突变\n类型",
        "频率\n(%)",
        "潜在获益靶向药物\n（证据等级）",
        "可能耐药或\n慎重药物\n（证据等级）",
    ]
    row2 = [
        "FBXW7",
        "NM_033632.4",
        "4",
        "10",
        "未见突变",
        "--",
        "--",
        "未见突变",
        "--",
    ]
    for row, values in zip(table.rows, [row0, row1, row2]):
        for idx, value in enumerate(values):
            row.cells[idx].text = value
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._restore_variant_detail_table_style(str(docx_path))

    table = Document(docx_path).tables[0]
    site_run = table.rows[2].cells[4].paragraphs[0].runs[0]
    drug_run = table.rows[2].cells[7].paragraphs[0].runs[0]

    assert site_run.font.underline is False
    assert str(site_run.font.color.rgb) == "000000"
    assert drug_run.font.underline is False
    assert str(drug_run.font.color.rgb) == "000000"
    assert drug_run.style.name == "Default Paragraph Font"


def test_variant_detail_table_restores_reviewed_template_style(tmp_path):
    docx_path = tmp_path / "variant_detail.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=9)
    row0 = [
        "基因名称",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "靶向药物信息",
        "靶向药物信息",
    ]
    row1 = [
        "基因名称",
        "转录本号",
        "染色体",
        "外显子",
        "位点",
        "突变\n类型",
        "频率\n(%)",
        "潜在获益靶向药物\n（证据等级）",
        "可能耐药或\n慎重药物\n（证据等级）",
    ]
    row2 = [
        "KRAS",
        "NM_004985.5",
        "12",
        "2",
        "c.34G>A,\np.G12S",
        "点突变",
        "46.29",
        "司美替尼（C）",
        "西妥昔单抗（A）",
    ]
    for row, values in zip(table.rows, [row0, row1, row2]):
        for idx, value in enumerate(values):
            row.cells[idx].text = value
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._restore_variant_detail_table_style(str(docx_path))

    doc = Document(docx_path)
    table = doc.tables[0]
    header_run = table.rows[0].cells[0].paragraphs[0].runs[0]
    gene_run = table.rows[2].cells[0].paragraphs[0].runs[0]
    site_run = table.rows[2].cells[4].paragraphs[0].runs[0]
    drug_run = table.rows[2].cells[7].paragraphs[0].runs[0]

    assert header_run.font.name == "微软雅黑"
    assert header_run.font.size.pt == 9
    assert header_run.font.bold is True
    assert str(header_run.font.color.rgb) == "F9FBFA"
    assert gene_run.font.underline is True
    assert str(gene_run.font.color.rgb) == "0000FF"
    assert site_run.font.underline is False
    assert str(site_run.font.color.rgb) == "000000"
    assert drug_run.font.underline is True
    assert str(drug_run.font.color.rgb) == "0000FF"
    assert table.rows[2].cells[7].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_variant_detail_table_can_disable_link_underlines_from_panel_style(tmp_path):
    docx_path = tmp_path / "variant_detail_plain.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=9)
    row0 = [
        "基因名称",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "靶向药物信息",
        "靶向药物信息",
    ]
    row1 = [
        "基因名称",
        "转录本号",
        "染色体",
        "外显子",
        "位点",
        "突变\n类型",
        "频率\n(%)",
        "潜在获益靶向药物\n（证据等级）",
        "可能耐药或\n慎重药物\n（证据等级）",
    ]
    row2 = [
        "KRAS",
        "NM_004985.5",
        "12",
        "2",
        "c.34G>A,\np.G12S",
        "点突变",
        "46.29",
        "司美替尼（C）",
        "西妥昔单抗（A）",
    ]
    for row, values in zip(table.rows, [row0, row1, row2]):
        for idx, value in enumerate(values):
            row.cells[idx].text = value
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._restore_variant_detail_table_style(
        str(docx_path),
        {
            "panel_style": {
                "variant_detail_table": {
                    "link_underline": False,
                    "link_color": "000000",
                }
            }
        },
    )

    table = Document(docx_path).tables[0]
    gene_run = table.rows[2].cells[0].paragraphs[0].runs[0]
    drug_run = table.rows[2].cells[7].paragraphs[0].runs[0]
    assert gene_run.font.underline is False
    assert str(gene_run.font.color.rgb) == "000000"
    assert drug_run.font.underline is False
    assert str(drug_run.font.color.rgb) == "000000"


def test_static_toc_page_numbers_keep_reviewed_toc_style(tmp_path):
    import re

    package = load_panel_package("crc_358_msi", project_root=ROOT)
    template_path = package.resolve_template_file("crc_358_msi_golden_template_v0")
    docx_path = tmp_path / "toc_style.docx"
    shutil.copy2(template_path, docx_path)

    context = {
        "panel_style": {
            "toc": {
                "font_name": "微软雅黑",
                "section_font_color": "00C4D8",
                "section_font_size": 16,
                "section_bold": True,
                "item_font_color": "000000",
                "item_font_size": 11,
                "item_bold": False,
                "content_top_padding_pt": 57,
            }
        }
    }
    ok = TemplateRenderer(log_level="ERROR")._write_static_toc_page_numbers(
        str(docx_path),
        {
            "患者及样本信息": 1,
            "检测内容": 1,
            "检测结果小结": 2,
            "靶向药物相关检测结果": 4,
            "免疫治疗疗效评估": 10,
            "检测结果说明": 15,
            "基因变异解析": 16,
            "靶向药物/免疫用药提示解析": 20,
            "阅读说明": 25,
            "参考文献": 71,
        },
        context,
    )

    assert ok is True
    xml = _read_docx_part(docx_path, "word/document.xml")
    section_idx = xml.index("第一部分：基本信息")
    item_idx = xml.index("患者及样本信息", section_idx)
    section_xml = xml[section_idx - 800 : section_idx + 200]
    item_xml = xml[item_idx - 800 : item_idx + 300]
    assert 'w:val="00C4D8"' in section_xml
    assert 'w:val="32"' in section_xml
    assert "<w:b" in section_xml
    assert 'w:val="000000"' in item_xml
    assert 'w:val="22"' in item_xml
    assert "<w:u" not in section_xml
    assert "<w:u" not in item_xml
    assert 'w:leader="dot"' not in xml
    # Page numbers are live PAGEREF fields tied to ReportGen-owned bookmarks.
    # Existing malformed template _Toc* anchors must never be reused.
    toc_xml = _toc_sdt_xml(docx_path)
    assert "HYPERLINK" not in toc_xml
    assert "PAGEREF _ReportGenToc_" in toc_xml
    assert 'w:dirty="true"' in toc_xml
    assert "1.检测结果小结" not in xml
    assert "靶向药物/免疫用药提示解析" in xml
    assert '<w:ind w:left="1980" w:leftChars="900"/>' in xml
    assert '<w:spacing w:before="1140" w:after="0" w:line="312" w:lineRule="auto"/>' in xml
    assert "<w:sectPr>" in xml
    settings_xml = _read_docx_part(docx_path, "word/settings.xml")
    assert "<w:updateFields" in settings_xml
    assert 'w:val="true"' in settings_xml

    # The first two entries used to bind to the introductory sentence on report
    # page 3 because it also mentions “患者及样本信息、检测内容”. They must
    # instead target the exact Part 1 heading and the exact 检测内容 heading.
    from lxml import etree

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(docx_path) as zf:
        root = etree.fromstring(zf.read("word/document.xml"))
    toc = next(
        sdt
        for sdt in root.xpath(".//w:sdt", namespaces=ns)
        if "第一部分" in "".join(sdt.xpath(".//w:t/text()", namespaces=ns))
        and "参考文献" in "".join(sdt.xpath(".//w:t/text()", namespaces=ns))
    )

    def row_anchor(label):
        paragraph = next(
            p
            for p in toc.xpath(".//w:p", namespaces=ns)
            if label in "".join(p.xpath(".//w:t/text()", namespaces=ns))
        )
        instruction = "".join(paragraph.xpath(".//w:instrText/text()", namespaces=ns))
        return re.search(r"PAGEREF (_ReportGenToc_[^ ]+)", instruction).group(1)

    def bookmark_text(name):
        start = root.xpath(".//w:bookmarkStart[@w:name=$name]", namespaces=ns, name=name)[0]
        paragraph = start.xpath("ancestor::w:p[1]", namespaces=ns)[0]
        return "".join(paragraph.xpath(".//w:t/text()", namespaces=ns)).strip()

    patient_target = bookmark_text(row_anchor("患者及样本信息"))
    content_target = bookmark_text(row_anchor("检测内容"))
    assert "第一部分：基本信息" in patient_target
    assert content_target == "检测内容"


def test_lung588_disabled_part3_toc_is_explicit_and_drops_active_analysis_rows(
    tmp_path,
):
    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    template_path = package.resolve_template_file()
    docx_path = tmp_path / "lung588_disabled_part3_toc.docx"
    shutil.copy2(template_path, docx_path)

    ok = TemplateRenderer(log_level="ERROR")._write_static_toc_page_numbers(
        str(docx_path),
        {
            "本部分未启用": 47,
            "阅读说明": 47,
            "参考文献": 79,
        },
        {
            "part3_knowledge_status": "disabled",
            "panel_style": {"toc": {}},
        },
    )

    assert ok is True
    toc_xml = _toc_sdt_xml(docx_path)
    assert "本部分未启用" in toc_xml
    assert "基因变异解析" not in toc_xml
    assert "靶向药物/免疫用药提示解析" not in toc_xml
    visible_numbers = TemplateRenderer(log_level="ERROR")._read_static_toc_page_numbers(
        str(docx_path)
    )
    assert visible_numbers["本部分未启用"] == 47
    assert visible_numbers["阅读说明"] == 47


def test_toc_uses_valid_reportgen_bookmarks_and_cached_page_numbers(tmp_path):
    import re

    from lxml import etree

    package = load_panel_package("crc_358_msi", project_root=ROOT)
    template_path = package.resolve_template_file("crc_358_msi_golden_template_v0")
    docx_path = tmp_path / "toc_targets.docx"
    shutil.copy2(template_path, docx_path)

    ok = TemplateRenderer(log_level="ERROR")._write_static_toc_page_numbers(
        str(docx_path),
        {
            "基因检测列表": 73,
            "参考文献": 75,
        },
        {"panel_style": {"toc": {}}},
    )

    assert ok is True
    # TOC page fields point only to dedicated ReportGen bookmarks; stale/malformed
    # template _Toc* bookmarks are not trusted.
    toc_xml = _toc_sdt_xml(docx_path)
    assert "HYPERLINK" not in toc_xml
    anchors = re.findall(r"PAGEREF (_ReportGenToc_[^ ]+) ", toc_xml)
    assert len(anchors) == 2
    assert len(set(anchors)) == 2

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(docx_path) as zf:
        root = etree.fromstring(zf.read("word/document.xml"))

    bookmark_names = set(root.xpath(".//w:bookmarkStart/@w:name", namespaces=ns))
    assert set(anchors) <= bookmark_names

    def paragraph_text(elem):
        return "".join(elem.xpath(".//w:t/text()", namespaces=ns))

    toc = [
        sdt
        for sdt in root.xpath(".//w:sdt", namespaces=ns)
        if "参考文献" in paragraph_text(sdt) and re.search(r"第[一二三四]部分", paragraph_text(sdt))
    ]
    assert toc, "TOC sdt not found"
    toc = toc[0]

    def toc_row_number(label: str) -> str | None:
        needle = label.replace(" ", "")
        for paragraph in toc.xpath(".//w:p", namespaces=ns):
            text = re.sub(r"\s+", "", paragraph_text(paragraph))
            match = re.search(r"(\d{1,3})$", text)
            if match and needle in text:
                return match.group(1)
        return None

    # Static page numbers land in the right rows.
    assert toc_row_number("基因检测列表") == "73"
    assert toc_row_number("参考文献") == "75"

    from reportgen.core.qa_report import _inspect_toc

    toc_check = _inspect_toc([], output_path=docx_path)
    assert toc_check["status"] == "PASS"
    assert toc_check["mode"] == "PAGEREF"
    assert toc_check["field_count"] == 2
    assert toc_check["unique_target_count"] == 2
    assert toc_check["missing_bookmarks"] == []
    assert toc_check["unclosed_bookmarks"] == []
    assert toc_check["duplicate_targets"] == []
    assert toc_check["update_fields"] is True


def test_set_word_compat_pagination_adds_printer_metrics(tmp_path):
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    template_path = package.resolve_template_file("crc_358_msi_golden_template_v0")
    docx_path = tmp_path / "compat.docx"
    shutil.copy2(template_path, docx_path)

    TemplateRenderer(log_level="ERROR")._set_word_compat_pagination(str(docx_path))

    settings_xml = _read_docx_part(docx_path, "word/settings.xml")
    # Word-compat pagination flags keep the cached preview closer to Word; live
    # PAGEREF fields provide the exact reader-side page numbers.
    assert "<w:usePrinterMetrics" in settings_xml
    assert "<w:doNotUseHTMLParagraphAutoSpacing" in settings_xml
    # Legacy compat flags must precede <w:compatSetting> in the schema.
    assert settings_xml.index("usePrinterMetrics") < settings_xml.index("compatSetting")


def test_variant_detail_no_mutation_gene_is_not_link_styled(tmp_path):
    """In the 9-column 2.1 variant-detail table, a 未见突变 (no-mutation) row
    must show its gene name as plain black text — NOT a blue underlined link.
    Only genes with an actual detected variant are link-styled (matches the
    reviewed final report). Regression guard: this had silently come back.
    """
    from docx.oxml.ns import qn

    doc = Document()
    table = doc.add_table(rows=4, cols=9)
    hdr0 = [
        "基因名称",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "靶向药物信息",
        "靶向药物信息",
    ]
    hdr1 = [
        "基因名称",
        "转录本号",
        "染色体",
        "外显子",
        "位点",
        "突变类型",
        "频率",
        "潜在获益靶向药物",
        "可能耐药或慎重药物",
    ]
    variant_row = [
        "TP53",
        "NM_000546",
        "chr17",
        "7",
        "c.844C>T",
        "missense",
        "30",
        "AZD1775（C）",
        "--",
    ]
    no_mut_row = ["BRAF", "", "", "", "未见突变", "", "", "--", "--"]
    for ci, value in enumerate(hdr0):
        table.rows[0].cells[ci].text = value
    for ci, value in enumerate(hdr1):
        table.rows[1].cells[ci].text = value
    for ci, value in enumerate(variant_row):
        table.rows[2].cells[ci].text = value
    for ci, value in enumerate(no_mut_row):
        table.rows[3].cells[ci].text = value
    docx_path = tmp_path / "variant_detail.docx"
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._restore_variant_detail_table_style(
        str(docx_path), {"panel_style": {"variant_detail_table": {}}}
    )

    out = Document(docx_path)
    result_table = out.tables[0]

    def gene_color_underline(row_idx):
        cell = result_table.rows[row_idx].cells[0]
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if not run.text.strip():
                    continue
                rpr = run._element.find(qn("w:rPr"))
                color = rpr.find(qn("w:color")) if rpr is not None else None
                underline = rpr.find(qn("w:u")) if rpr is not None else None
                return (
                    (color.get(qn("w:val")) if color is not None else None),
                    (underline.get(qn("w:val")) if underline is not None else None),
                )
        return (None, None)

    # Real variant row: gene is a blue underlined link.
    color, underline = gene_color_underline(2)
    assert color and color.lower() == "0000ff"
    assert underline and underline != "none"

    # No-mutation row: gene is plain black, no underline.
    color, underline = gene_color_underline(3)
    assert (color or "000000").lower() == "000000"
    assert (underline or "none") == "none"


def test_biomarker_table_restores_template_typography(tmp_path):
    docx_path = tmp_path / "biomarker.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "TMB/MSI/其它生物标志物检测结果"
    table.rows[0].cells[1].text = "TMB/MSI/其它生物标志物检测结果"
    table.rows[0].cells[2].text = "用药提示"
    table.rows[1].cells[0].text = "肿瘤突变负荷（TMB）"
    table.rows[1].cells[1].text = "6.5 mutations/Mb，TMB-L"
    table.rows[1].cells[
        2
    ].text = "多项临床研究表明，TMB-H的肿瘤对免疫检查点抑制剂有更强的免疫应答效果"
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._restore_biomarker_table_style(str(docx_path))

    doc = Document(docx_path)
    table = doc.tables[0]
    header_run = table.rows[0].cells[2].paragraphs[0].runs[0]
    body_run = table.rows[1].cells[2].paragraphs[0].runs[0]

    assert header_run.font.name == "微软雅黑"
    assert header_run.font.size.pt == 10
    assert header_run.font.bold is True
    assert str(header_run.font.color.rgb) == "F9FBFA"
    assert body_run.font.name == "微软雅黑"
    assert body_run.font.size.pt == 9
    assert body_run.font.bold is False
    assert body_run.font.underline is False
    assert str(body_run.font.color.rgb) == "000000"
    assert table.rows[1].cells[2].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_biomarker_table_uses_panel_style_tokens(tmp_path):
    docx_path = tmp_path / "biomarker_panel_style.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "TMB/MSI/其它生物标志物检测结果"
    table.rows[0].cells[1].text = "TMB/MSI/其它生物标志物检测结果"
    table.rows[0].cells[2].text = "用药提示"
    table.rows[1].cells[0].text = "微卫星不稳定性（MSI）"
    table.rows[1].cells[1].text = "MSS"
    table.rows[1].cells[2].text = "研究表明，MSI-H的实体瘤通常具有免疫原性"
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._restore_biomarker_table_style(
        str(docx_path),
        {
            "panel_style": {
                "biomarker_table": {
                    "header_fill": "00B7C7",
                    "header_font_color": "FFFFFF",
                    "body_font_color": "000000",
                    "font_name": "微软雅黑",
                    "header_font_size": 10,
                    "body_font_size": 9,
                    "border_color": "000000",
                    "border_size": "6",
                }
            }
        },
    )

    table = Document(docx_path).tables[0]
    header_cell = table.rows[0].cells[2]
    header_run = header_cell.paragraphs[0].runs[0]
    shd = header_cell._tc.get_or_add_tcPr().find(qn("w:shd"))

    assert shd is not None
    assert shd.get(qn("w:fill")) == "00B7C7"
    assert str(header_run.font.color.rgb) == "FFFFFF"


def test_clinical_result_tables_restore_black_borders_and_red_detected_results(tmp_path):
    docx_path = tmp_path / "clinical_result_tables.docx"
    doc = Document()

    chemo = doc.add_table(rows=2, cols=3)
    chemo.rows[0].cells[0].text = "药物名称"
    chemo.rows[0].cells[1].text = "相关基因"
    chemo.rows[0].cells[2].text = "药物适应情况"
    chemo.rows[1].cells[0].text = "瑞戈非尼"
    chemo.rows[1].cells[1].text = "VEGFR"
    chemo.rows[1].cells[2].text = "FDA批准用于治疗结直肠癌。"

    nccn = doc.add_table(rows=2, cols=3)
    nccn.rows[0].cells[0].text = "检测基因"
    nccn.rows[0].cells[1].text = "检测内容"
    nccn.rows[0].cells[2].text = "检测结果"
    nccn.rows[1].cells[0].text = "KRAS"
    nccn.rows[1].cells[1].text = "外显子2"
    nccn.rows[1].cells[2].text = "c.34G>A，p.G12S"

    immune = doc.add_table(rows=3, cols=3)
    immune.rows[0].cells[0].text = "基因"
    immune.rows[0].cells[1].text = "检测结果"
    immune.rows[0].cells[2].text = "临床解读"
    immune.rows[1].cells[0].text = "KRAS"
    immune.rows[1].cells[1].text = "c.34G>A，p.G12S"
    immune.rows[1].cells[2].text = "检出有害变异时可能疗效较好。"
    immune.rows[2].cells[0].text = "PTEN"
    immune.rows[2].cells[1].text = "未检出有害变异"
    immune.rows[2].cells[2].text = "检出有害变异时可能耐药。"
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._restore_clinical_result_table_style(str(docx_path))

    doc = Document(docx_path)

    def cell_border_colors(table):
        colors = set()
        for row in table.rows:
            for cell in row.cells:
                borders = cell._tc.get_or_add_tcPr().find(qn("w:tcBorders"))
                assert borders is not None
                for child in borders:
                    colors.add(child.get(qn("w:color")))
        return colors

    assert cell_border_colors(doc.tables[0]) == {"000000"}
    assert cell_border_colors(doc.tables[1]) == {"000000"}
    assert cell_border_colors(doc.tables[2]) == {"000000"}

    detected_run = doc.tables[2].rows[1].cells[1].paragraphs[0].runs[0]
    undetected_run = doc.tables[2].rows[2].cells[1].paragraphs[0].runs[0]
    assert str(detected_run.font.color.rgb) == "FF0000"
    assert detected_run.font.underline is False
    assert str(undetected_run.font.color.rgb) != "FF0000"


def test_report_content_fixes_remove_tmb_h_only_notes_when_tmb_low(tmp_path):
    docx_path = tmp_path / "immune_notes.docx"
    doc = Document()
    doc.add_paragraph("2.上表涉及的已上市的药物名称及对应的商品名称：旧商品名。")
    doc.add_paragraph(
        "#帕博利珠单抗均已获FDA和/或NMPA批准用于治疗结直肠癌。"
        "并且，FDA 已批准帕博利珠单抗用于治疗 TMB-H 的不可切除或转移性的成人和儿童实体瘤。"
    )
    doc.add_paragraph(
        "目前已知的免疫治疗生物标志物包括TMB、MSI、PD-L1表达、"
        "免疫疗效正相关/负相关/超进展基因等。值得注意的是，"
        "TMB、MSI、PD-L1表达预测生物标志物是相对独立的预测指标。"
    )
    doc.add_paragraph("4. 上表涉及的已上市的药物名称及对应的商品名称：免疫药品清单。")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._apply_report_content_fixes(
        str(docx_path),
        {
            "tmb_status": "L",
            "msi_status": "MSS",
            "targeted_drug_brand_summary": "西妥昔单抗[爱必妥]。",
        },
    )

    text = "\n".join(p.text for p in Document(docx_path).paragraphs)
    assert "TMB-H 的不可切除或转移性的成人和儿童实体瘤" not in text
    assert "相对独立" not in text
    assert "FDA和/或NMPA批准用于治疗结直肠癌" in text
    assert "西妥昔单抗[爱必妥]" in text


def test_report_content_fixes_keep_conflict_note_for_tmb_high_mss(tmp_path):
    docx_path = tmp_path / "immune_notes_high.docx"
    doc = Document()
    doc.add_paragraph("2.上表涉及的已上市的药物名称及对应的商品名称：旧商品名。")
    doc.add_paragraph(
        "并且，FDA 已批准帕博利珠单抗用于治疗 TMB-H 的不可切除或转移性的成人和儿童实体瘤。"
    )
    doc.add_paragraph(
        "目前已知的免疫治疗生物标志物包括TMB、MSI、PD-L1表达。"
        "值得注意的是，TMB、MSI、PD-L1表达预测生物标志物是相对独立的预测指标。"
    )
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._apply_report_content_fixes(
        str(docx_path),
        {
            "tmb_status": "H",
            "msi_status": "MSS",
            "targeted_drug_brand_summary": "帕博利珠单抗[可瑞达]。",
        },
    )

    text = "\n".join(p.text for p in Document(docx_path).paragraphs)
    assert "TMB-H 的不可切除或转移性的成人和儿童实体瘤" in text
    assert "相对独立" in text


def test_multiline_bullet_placeholder_splits_into_real_bullets(tmp_path):
    docx_path = tmp_path / "multiline_bullets.docx"
    doc = Document()

    def add_bullet(text: str = ""):
        paragraph = doc.add_paragraph(text)
        ppr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "1")
        num_pr.append(ilvl)
        num_pr.append(num_id)
        ppr.append(num_pr)
        return paragraph

    add_bullet("第一条\n第二条\n第三条")
    add_bullet()
    add_bullet()
    add_bullet()
    doc.add_paragraph("注：")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._normalize_multiline_bullet_paragraphs(str(docx_path))

    rendered = Document(docx_path)
    bullets = [
        p.text for p in rendered.paragraphs if p._p.pPr is not None and p._p.pPr.numPr is not None
    ]
    assert bullets == ["第一条", "第二条", "第三条"]


def test_empty_numbered_paragraphs_are_removed(tmp_path):
    docx_path = tmp_path / "empty_numbered.docx"
    doc = Document()

    def add_bullet(text: str = ""):
        paragraph = doc.add_paragraph(text)
        ppr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "1")
        num_pr.append(ilvl)
        num_pr.append(num_id)
        ppr.append(num_pr)
        return paragraph

    add_bullet("保留")
    add_bullet("")
    doc.add_paragraph("")
    add_bullet("也保留")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._remove_empty_numbered_paragraphs(str(docx_path))

    rendered = Document(docx_path)
    bullets = [
        p.text for p in rendered.paragraphs if p._p.pPr is not None and p._p.pPr.numPr is not None
    ]
    assert bullets == ["保留", "也保留"]
    assert any(not p.text for p in rendered.paragraphs)


def test_empty_numbered_paragraph_with_image_is_preserved(tmp_path):
    from PIL import Image

    docx_path = tmp_path / "numbered_image.docx"
    image_path = tmp_path / "figure.png"
    Image.new("RGB", (320, 120), "white").save(image_path)
    doc = Document()
    paragraph = doc.add_paragraph("")
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(num_id)
    ppr.append(num_pr)
    paragraph.add_run().add_picture(str(image_path), width=Inches(2.0))
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._remove_empty_numbered_paragraphs(str(docx_path))

    rendered = Document(docx_path)
    assert len(rendered.inline_shapes) == 1
    assert any("<w:drawing" in paragraph._p.xml for paragraph in rendered.paragraphs)


def test_multiline_bullet_cleanup_preserves_numbered_image_paragraph(tmp_path):
    from PIL import Image

    docx_path = tmp_path / "numbered_image_after_multiline.docx"
    image_path = tmp_path / "figure.png"
    Image.new("RGB", (320, 120), "white").save(image_path)
    doc = Document()

    def add_bullet(text: str = ""):
        paragraph = doc.add_paragraph(text)
        ppr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), "1")
        num_pr.append(num_id)
        ppr.append(num_pr)
        return paragraph

    add_bullet("第一条\n第二条")
    image_paragraph = add_bullet("")
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(2.0))
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._normalize_multiline_bullet_paragraphs(str(docx_path))

    rendered = Document(docx_path)
    assert len(rendered.inline_shapes) == 1
    assert any("<w:drawing" in paragraph._p.xml for paragraph in rendered.paragraphs)


def test_blank_page_break_before_references_heading_is_removed(tmp_path):
    docx_path = tmp_path / "reference_break.docx"
    doc = Document()
    doc.add_paragraph("Gene List for MLseq (n=358)")
    blank = doc.add_paragraph()
    blank.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("5. 参考文献")
    doc.add_paragraph("PMID: 23066310 KRAS mutation testing in metastatic colorectal cancer.")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._remove_blank_page_breaks_before_headings(str(docx_path))

    rendered = Document(docx_path)
    texts = [p.text for p in rendered.paragraphs]
    assert texts == [
        "Gene List for MLseq (n=358)",
        "5. 参考文献",
        "PMID: 23066310 KRAS mutation testing in metastatic colorectal cancer.",
    ]
    assert all('w:type="page"' not in p._p.xml for p in rendered.paragraphs)


def test_blank_page_break_before_quality_control_heading_is_removed(tmp_path):
    docx_path = tmp_path / "quality_control_break.docx"
    doc = Document()
    doc.add_paragraph("72. 2019 JCO Abstract 5005 TOPARP-B.")
    blank = doc.add_paragraph()
    blank.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("本次检测质控结果")
    doc.add_paragraph("质控表正文")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._remove_blank_page_breaks_before_headings(
        str(docx_path),
        ("本次检测质控结果",),
    )

    rendered = Document(docx_path)
    assert [p.text for p in rendered.paragraphs] == [
        "72. 2019 JCO Abstract 5005 TOPARP-B.",
        "本次检测质控结果",
        "质控表正文",
    ]
    assert all('w:type="page"' not in p._p.xml for p in rendered.paragraphs)


def test_part3_section_spacing_removes_empty_cluster_but_preserves_section(tmp_path):
    docx_path = tmp_path / "part3_section_spacing.docx"
    doc = Document()
    signature = doc.add_paragraph("检测者：  审核者：  报告日期：2026.06.04")
    for _ in range(6):
        doc.add_paragraph("")
    doc.add_section(WD_SECTION.NEW_PAGE)
    heading = "第三部分：基因变异及相应靶向/免疫药物解析"
    doc.add_paragraph(heading)
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._cleanup_section_spacing(docx_path, (heading,))
    renderer._cleanup_section_spacing(docx_path, (heading,))

    rendered = Document(docx_path)
    assert [p.text for p in rendered.paragraphs] == [signature.text, heading]
    assert "sectPr" in rendered.paragraphs[0]._p.xml
    assert len(rendered.sections) == 2


def test_standalone_page_break_before_pathway_table_is_removed(tmp_path):
    docx_path = tmp_path / "pathway_break.docx"
    doc = Document()
    doc.add_paragraph("参考文献：")
    doc.add_paragraph("Cancer Cell. 2007;12:104-107.")
    doc.add_paragraph("")
    standalone = doc.add_paragraph("")
    standalone.add_run().add_break(WD_BREAK.PAGE)
    first = doc.add_table(rows=1, cols=2)
    first.cell(0, 0).text = "通路名称"
    first.cell(0, 1).text = "MAPK/ERK Signaling Pathway"

    doc.add_paragraph("MAPK/ERK Signaling Pathway")
    pathway_reference_label = doc.add_paragraph("参考文献：")
    pathway_reference_label.paragraph_format.space_after = Pt(2)
    embedded = doc.add_paragraph("Nat Rev Cancer. 2003;3:650-665.")
    embedded.paragraph_format.space_after = Pt(2)
    embedded.add_run().add_break(WD_BREAK.PAGE)
    second = doc.add_table(rows=1, cols=2)
    second.cell(0, 0).text = "通路名称"
    second.cell(0, 1).text = "Wnt Signaling Pathway"
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._remove_standalone_page_breaks_before_pathway_tables(docx_path)
    renderer._remove_standalone_page_breaks_before_pathway_tables(docx_path)

    rendered = Document(docx_path)
    empty_breaks = [
        p for p in rendered.paragraphs if not p.text.strip() and 'w:type="page"' in p._p.xml
    ]
    assert empty_breaks == []
    assert not any(not p.text.strip() for p in rendered.paragraphs[:2])
    reference_label = next(p for p in rendered.paragraphs if p.text.strip() == "参考文献：")
    assert "keepNext" in reference_label._p.xml
    assert (
        'w:type="page"'
        in next(p for p in rendered.paragraphs if p.text.startswith("Nat Rev Cancer"))._p.xml
    )
    assert pathway_reference_label.text in [p.text for p in rendered.paragraphs]
    compacted_label = [
        p
        for p in rendered.paragraphs
        if p.text == pathway_reference_label.text and "keepNext" in p._p.xml
    ][-1]
    compacted_reference = next(
        p for p in rendered.paragraphs if p.text.startswith("Nat Rev Cancer")
    )
    assert 'w:after="0"' in compacted_label._p.xml
    assert 'w:after="0"' in compacted_reference._p.xml
    assert len(rendered.tables) == 2


def test_template_tmb_msi_patient_narratives_are_dynamic():
    template = ROOT / "templates/aligned_template_with_cnv_fusion_hla_FIXED.docx"
    doc = Document(template)
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    text = "\n".join(texts)

    assert "13.5mutations/Mb" not in text
    assert "该肿瘤样本为 微卫星稳定（MSS）型" not in text
    assert "MSI-H的实体瘤通常具有免疫原性" not in text
    for token in [
        "{{ tmb_detail_sentence }}",
        "{{ tmb_detail_interpretation }}",
        "{{ tmb_drug_note }}",
        "{{ msi_detail_sentence }}",
        "{{ msi_detail_interpretation }}",
        "{{ msi_tips }}",
    ]:
        assert token in text


def test_patient_letter_is_native_docx_text_not_legacy_textbox():
    from zipfile import ZipFile
    import xml.etree.ElementTree as ET

    template = ROOT / "templates/aligned_template_with_cnv_fusion_hla_FIXED.docx"
    doc = Document(template)
    text_parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)
    text = "\n".join(text_parts)

    assert "致您的一封信" in text
    assert "尊敬的 {{ patient_name }} {{ patient_salutation }}：" in text
    assert "现代医学已经证明" in text

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(template) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    textbox_text = "\n".join(
        "".join(t.text or "" for t in textbox.findall(".//w:t", ns))
        for textbox in root.findall(".//w:txbxContent", ns)
    )

    assert "现代医学已经证明" not in textbox_text
    assert "尊敬的" not in textbox_text


def test_patient_letter_body_block_is_not_pushed_to_right():
    from zipfile import ZipFile
    import xml.etree.ElementTree as ET

    template = ROOT / "templates/aligned_template_with_cnv_fusion_hla_FIXED.docx"
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(template) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))

    matching_tables = []
    for table in root.findall(".//w:tbl", ns):
        text = "".join(t.text or "" for t in table.findall(".//w:t", ns))
        if "现代医学已经证明" in text:
            matching_tables.append(table)

    assert len(matching_tables) == 1
    widths = [
        int(tc_w.get(f"{{{ns['w']}}}w") or "0")
        for tc_w in matching_tables[0].findall(".//w:tcPr/w:tcW", ns)[:2]
    ]
    grid_widths = [
        int(grid_col.get(f"{{{ns['w']}}}w") or "0")
        for grid_col in matching_tables[0].findall("./w:tblGrid/w:gridCol", ns)[:2]
    ]
    assert widths[0] <= 800
    assert widths[1] >= 8000
    assert grid_widths[0] <= 800
    assert grid_widths[1] >= 8000


def test_patient_letter_salutation_follows_gender():
    generator = ReportGenerator(config_dir=str(ROOT / "config"), log_level="ERROR")
    report_data = ReportData()

    report_data.set_field("gender", "女")
    generator._set_patient_salutation(report_data)
    assert report_data.get_field("patient_salutation") == "女士"

    report_data.set_field("gender", "男")
    generator._set_patient_salutation(report_data)
    assert report_data.get_field("patient_salutation") == "先生"


def test_clinical_diagnosis_populates_basic_info_display_only():
    generator = ReportGenerator(config_dir=str(ROOT / "config"), log_level="ERROR")
    report_data = ReportData()
    report_data.set_field("cancer_type", "-")
    report_data.set_field("clinical_diagnosis", "乙状结肠癌")

    generator._apply_clinical_diagnosis_for_display(report_data)

    assert report_data.get_field("cancer_type") == "乙状结肠癌"


def test_toc_decoration_line_is_moved_left_and_up(tmp_path):
    from shutil import copyfile
    from zipfile import ZipFile
    import xml.etree.ElementTree as ET

    source = ROOT / "templates/aligned_template_with_cnv_fusion_hla_FIXED.docx"
    docx_path = tmp_path / "toc_layout.docx"
    copyfile(source, docx_path)

    TemplateRenderer(log_level="ERROR")._normalize_toc_decoration_layout(str(docx_path))

    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    w_p = f"{{{ns_w}}}p"
    w_t = f"{{{ns_w}}}t"
    wp_anchor = f"{{{ns_wp}}}anchor"
    wp_doc_pr = f"{{{ns_wp}}}docPr"
    wp_extent = f"{{{ns_wp}}}extent"
    wp_position_h = f"{{{ns_wp}}}positionH"
    wp_position_v = f"{{{ns_wp}}}positionV"
    wp_pos_offset = f"{{{ns_wp}}}posOffset"

    with ZipFile(docx_path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))

    line_offsets = []
    circle_offsets = []
    for para in root.iter(w_p):
        text = "".join(t.text or "" for t in para.iter(w_t)).replace(" ", "")
        if "目录" not in text:
            continue
        for anchor in para.iter(wp_anchor):
            doc_pr = anchor.find(wp_doc_pr)
            extent = anchor.find(wp_extent)
            if doc_pr is None or extent is None:
                continue
            name = doc_pr.get("name", "")
            cx = int(extent.get("cx") or "0")
            cy = int(extent.get("cy") or "0")
            pos_h = anchor.find(wp_position_h)
            pos_v = anchor.find(wp_position_v)
            if pos_h is None or pos_v is None:
                continue
            x = int((pos_h.find(wp_pos_offset).text or "0"))
            y = int((pos_v.find(wp_pos_offset).text or "0"))
            if "直接连接符" in name and cx <= 100000 and cy >= 2000000:
                line_offsets.append((x, y))
            elif "椭圆" in name and cx <= 200000 and cy <= 200000:
                circle_offsets.append((x, y))

    assert line_offsets
    assert circle_offsets
    assert all(x == 862965 and y == 1119505 for x, y in line_offsets)
    assert all(x == 828675 and y == 1043305 for x, y in circle_offsets)


def test_section_layouts_are_normalized(tmp_path):
    docx_path = tmp_path / "sections.docx"
    doc = Document()
    doc.add_paragraph("第三部分：基因变异及相应靶向/免疫药物解析")
    doc.sections[0].right_margin = Cm(2.0)
    second = doc.add_section(WD_SECTION.NEW_PAGE)
    second.right_margin = Cm(3.5)
    doc.add_paragraph("第四部分：附录")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._normalize_final_section_layout(str(docx_path))

    margins = {round(section.right_margin.cm, 2) for section in Document(docx_path).sections}
    assert margins == {2.0}


def test_appendix_sections_reuse_reviewed_body_header(tmp_path):
    docx_path = tmp_path / "headers.docx"
    doc = Document()
    doc.add_paragraph("致您的一封信")
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.header.is_linked_to_previous = False
    body_section.header.paragraphs[0].text = "姓名：测试患者              科技服务人类健康"
    doc.add_paragraph("第三部分：基因变异及相应靶向/免疫药物解析")
    appendix_section = doc.add_section(WD_SECTION.NEW_PAGE)
    appendix_section.header.is_linked_to_previous = False
    appendix_section.header.paragraphs[0].text = ""
    appendix_section.different_first_page_header_footer = True
    doc.add_paragraph("Gene List for MLseq (n=358)")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._restore_reviewed_body_headers(str(docx_path))

    doc = Document(docx_path)
    appendix = doc.sections[-1]
    assert not appendix.different_first_page_header_footer
    assert "科技服务人类健康" in "\n".join(p.text for p in appendix.header.paragraphs)
    p_bdr = appendix.footer.paragraphs[0]._p.get_or_add_pPr().find(qn("w:pBdr"))
    assert p_bdr is not None
    assert p_bdr.find(qn("w:top")).get(qn("w:val")) == "single"


def test_gene_list_table_matches_reviewed_layout(tmp_path):
    docx_path = tmp_path / "gene_list.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Gene List for MLseq (n=358)"
    table.rows[0].cells[0].paragraphs[0].add_run().add_break(WD_BREAK.PAGE)
    rendered_break = OxmlElement("w:lastRenderedPageBreak")
    table.rows[0].cells[0].paragraphs[0].runs[0]._r.append(rendered_break)
    table.rows[1].cells[0].text = "ABL1"
    table.rows[1].cells[1].text = "ABL2"
    table.rows[2].cells[0].text = "ZRSR2"
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._compact_gene_list_tables(str(docx_path))

    doc = Document(docx_path)
    body_run = doc.tables[0].rows[1].cells[0].paragraphs[0].runs[0]
    assert body_run.font.size.pt == 10.5
    assert body_run.font.underline is False
    title_xml = doc.tables[0].rows[0].cells[0].paragraphs[0]._p.xml
    assert 'w:type="page"' not in title_xml
    assert "w:lastRenderedPageBreak" not in title_xml
    borders = doc.tables[0]._tbl.tblPr.find(qn("w:tblBorders"))
    assert borders is not None
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        assert border is not None
        assert border.get(qn("w:val")) == "single"
        assert border.get(qn("w:sz")) == "4"
        assert border.get(qn("w:color")) == "BEBEBE"


def test_quality_control_table_disables_inherited_underlines(tmp_path):
    docx_path = tmp_path / "qc.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "质控项"
    table.rows[0].cells[1].text = "质控结果"
    table.rows[1].cells[0].text = "核酸提取与质量控制"
    table.rows[1].cells[1].text = "合格"
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._normalize_quality_control_tables(str(docx_path))

    doc = Document(docx_path)
    body_run = doc.tables[0].rows[1].cells[0].paragraphs[0].runs[0]
    assert body_run.font.underline is False


def test_macos_field_refresh_does_not_use_word_by_default(monkeypatch, tmp_path):
    renderer = TemplateRenderer(log_level="ERROR")
    docx_path = tmp_path / "refresh.docx"
    docx_path.write_bytes(b"placeholder")
    calls = []

    monkeypatch.setattr(sys, "platform", "darwin")
    monkeypatch.delenv("REPORTGEN_REFRESH_WITH_WORD", raising=False)
    monkeypatch.setattr(renderer, "_document_contains_toc", lambda path: True)
    monkeypatch.setattr(
        renderer, "_refresh_fields_with_word_macos", lambda path: calls.append("word")
    )
    monkeypatch.setattr(
        renderer, "_refresh_fields_with_libreoffice", lambda path: calls.append("libreoffice")
    )

    renderer._refresh_fields_with_native_engine(str(docx_path))

    assert calls == ["libreoffice"]


def test_macos_field_refresh_can_opt_into_word(monkeypatch, tmp_path):
    renderer = TemplateRenderer(log_level="ERROR")
    docx_path = tmp_path / "refresh.docx"
    docx_path.write_bytes(b"placeholder")
    calls = []

    monkeypatch.setattr(sys, "platform", "darwin")
    monkeypatch.setenv("REPORTGEN_REFRESH_WITH_WORD", "1")
    monkeypatch.setattr(renderer, "_document_contains_toc", lambda path: True)
    monkeypatch.setattr(
        renderer, "_refresh_fields_with_word_macos", lambda path: calls.append("word")
    )
    monkeypatch.setattr(
        renderer, "_refresh_fields_with_libreoffice", lambda path: calls.append("libreoffice")
    )

    renderer._refresh_fields_with_native_engine(str(docx_path))

    assert calls == ["word"]


def test_variant_table_layout_optimizer_keeps_reviewed_nine_column_template(tmp_path):
    docx_path = tmp_path / "variant_table.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=9)
    headers = [
        "基因名称",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "靶向药物信息",
        "靶向药物信息",
    ]
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for idx in range(9):
        table.rows[1].cells[idx].text = "测试内容"
    values = [
        "TP53",
        "NM_000546.6",
        "17",
        "8",
        "c.817C>T,\np.R273C",
        "点突变",
        "76.12",
        "AZD1775（C）\nAlisertib（C）",
        "--",
    ]
    for idx, text in enumerate(values):
        table.rows[2].cells[idx].text = text
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._optimize_variant_table_layout(str(docx_path))

    doc = Document(docx_path)
    table = doc.tables[0]
    assert len(table.columns) == 9
    assert table.rows[0].cells[0].text == "基因名称"
    assert table.rows[2].cells[1].text == "NM_000546.6"
    assert table.rows[2].cells[6].text == "76.12"


def test_real_no_variants_sample_regression_if_available():
    sample = (
        ROOT
        / "storage/uploads/2026-03-27/0ae95efb-606b-42f0-80d7-e21283c6415c"
        / "case3_no_variants_MLB0003.result.xlsx"
    )
    if not sample.exists():
        pytest.skip("real no-variants sample is not available in this checkout")

    excel_data = ExcelReader(config_dir=str(ROOT / "config"), log_level="ERROR").read(
        str(sample), include_tables=True
    )
    report_data = enhance_report_data(
        ReportData(),
        excel_data,
        base_path=str(ROOT),
    )

    assert report_data.get_field("total_variants_count") == 0
    assert report_data.get_table("variants") == []
    assert not any(
        (row.get("gene") or "").upper() == "ERBB2"
        for row in report_data.get_table("targeted_drug_tips")
    )


def test_qa_report_passes_basic_clean_docx(tmp_path):
    docx_path = tmp_path / "clean.docx"
    doc = Document()
    doc.add_paragraph("已生成报告")
    doc.save(docx_path)

    qa = build_docx_qa_report(output_file=str(docx_path))

    assert qa["status"] == "PASS"
    assert qa["checks"]["docx_openable"]["status"] == "PASS"
    assert qa["checks"]["unrendered_placeholders"]["count"] == 0
    assert qa["checks"]["visual_render"]["status"] == "SKIP"


def test_qa_report_visual_render_passes_with_nonblank_png(tmp_path, monkeypatch):
    docx_path = tmp_path / "visual_pass.docx"
    doc = Document()
    doc.add_paragraph("已生成报告")
    doc.save(docx_path)

    def fake_render(docx_path, *, output_dir, **_kwargs):
        from PIL import Image, ImageDraw

        output_dir.mkdir(parents=True, exist_ok=True)
        page = output_dir / f"{Path(docx_path).stem}-1.png"
        image = Image.new("RGB", (200, 200), "white")
        ImageDraw.Draw(image).rectangle((20, 20, 180, 180), fill="black")
        image.save(page)
        return [page]

    monkeypatch.setattr("reportgen.core.qa_report.render_docx_to_pngs", fake_render)

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        visual_render="first",
        visual_render_output_dir=str(tmp_path / "rendered"),
    )

    assert qa["status"] == "PASS"
    assert qa["checks"]["visual_render"]["status"] == "PASS"
    assert qa["checks"]["blank_page_detection"]["status"] == "PASS"
    assert qa["metrics"]["visual_render_page_count"] == 1


def test_qa_report_detects_body_blank_page_hidden_by_header_footer_and_watermark(
    tmp_path, monkeypatch
):
    docx_path = tmp_path / "body_blank.docx"
    doc = Document()
    doc.add_paragraph("已生成报告")
    doc.save(docx_path)

    def fake_render(docx_path, *, output_dir, **_kwargs):
        from PIL import Image, ImageDraw

        output_dir.mkdir(parents=True, exist_ok=True)
        pages = []
        for page_no in (1, 2, 3):
            page = output_dir / f"{Path(docx_path).stem}-{page_no}.png"
            image = Image.new("RGB", (200, 300), "white")
            draw = ImageDraw.Draw(image)
            if page_no == 2:
                # Simulate the production page header/footer plus a light cyan
                # watermark, but deliberately leave the report body empty.
                draw.rectangle((20, 10, 180, 15), fill="black")
                draw.rectangle((95, 285, 105, 290), fill="black")
                draw.polygon(
                    [(100, 85), (145, 150), (100, 215), (55, 150)],
                    outline=(210, 250, 250),
                    width=8,
                )
            else:
                draw.rectangle((25, 65, 175, 235), fill="black")
            image.save(page)
            pages.append(page)
        return pages

    monkeypatch.setattr("reportgen.core.qa_report.render_docx_to_pngs", fake_render)

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        visual_render="all",
        visual_render_required=True,
        visual_render_output_dir=str(tmp_path / "rendered"),
    )

    assert qa["status"] == "FAIL"
    check = qa["checks"]["blank_page_detection"]
    assert check["status"] == "FAIL"
    assert any(path.endswith("-2.png") for path in check["low_content_pages"])
    page2 = next(
        row
        for row in qa["checks"]["visual_render"]["pixel_check"]["pages"]
        if row["path"].endswith("-2.png")
    )
    assert page2["body_low_content"] is True
    assert page2["body_dark_ratio"] == 0.0


def test_qa_report_accepts_only_reviewed_semantic_sparse_page(tmp_path, monkeypatch):
    docx_path = tmp_path / "semantic_sparse.docx"
    doc = Document()
    doc.add_paragraph("已生成报告")
    doc.save(docx_path)

    def fake_render(docx_path, *, output_dir, **_kwargs):
        from PIL import Image, ImageDraw

        output_dir.mkdir(parents=True, exist_ok=True)
        pages = []
        for page_no in (1, 2, 3):
            page = output_dir / f"{Path(docx_path).stem}-{page_no}.png"
            image = Image.new("RGB", (200, 300), "white")
            draw = ImageDraw.Draw(image)
            if page_no == 2:
                draw.rectangle((30, 80, 110, 81), fill="black")
                draw.rectangle((30, 100, 80, 101), fill="black")
            else:
                draw.rectangle((25, 65, 175, 235), fill="black")
            image.save(page)
            pages.append(page)
        return pages

    monkeypatch.setattr("reportgen.core.qa_report.render_docx_to_pngs", fake_render)
    monkeypatch.setattr(
        "reportgen.core.qa_report._extract_pdf_page_texts",
        lambda *_args, **_kwargs: (
            {
                1: "正文",
                2: "检测结果仅对本次送检样本负责\n咨询电话：00000000。",
                3: "正文",
            },
            None,
        ),
    )
    report_data = ReportData(
        context={
            "panel_style": {
                "visual_qa": {
                    "expected_sparse_pages": [
                        {
                            "id": "legal_notice",
                            "required_text": [
                                "检测结果仅对本次送检样本负责",
                                "咨询电话：",
                            ],
                            "max_matches": 1,
                        }
                    ]
                }
            }
        }
    )

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        report_data=report_data,
        visual_render="all",
        visual_render_required=True,
        visual_render_output_dir=str(tmp_path / "rendered"),
    )

    assert qa["status"] == "PASS"
    check = qa["checks"]["blank_page_detection"]
    assert check["status"] == "PASS"
    accepted = check["expected_sparse_pages"]
    assert len(accepted) == 1
    assert accepted[0]["id"] == "legal_notice"
    assert accepted[0]["page_number"] == 2


def test_qa_report_rejects_sparse_page_when_semantic_marker_is_incomplete(tmp_path, monkeypatch):
    docx_path = tmp_path / "unexpected_sparse.docx"
    Document().save(docx_path)

    def fake_render(docx_path, *, output_dir, **_kwargs):
        from PIL import Image, ImageDraw

        output_dir.mkdir(parents=True, exist_ok=True)
        pages = []
        for page_no in (1, 2, 3):
            page = output_dir / f"{Path(docx_path).stem}-{page_no}.png"
            image = Image.new("RGB", (200, 300), "white")
            draw = ImageDraw.Draw(image)
            if page_no == 2:
                draw.rectangle((30, 80, 110, 81), fill="black")
            else:
                draw.rectangle((25, 65, 175, 235), fill="black")
            image.save(page)
            pages.append(page)
        return pages

    monkeypatch.setattr("reportgen.core.qa_report.render_docx_to_pngs", fake_render)
    monkeypatch.setattr(
        "reportgen.core.qa_report._extract_pdf_page_texts",
        lambda *_args, **_kwargs: (
            {2: "检测结果仅对本次送检样本负责"},
            None,
        ),
    )
    report_data = ReportData(
        context={
            "panel_style": {
                "visual_qa": {
                    "expected_sparse_pages": [
                        {
                            "id": "legal_notice",
                            "required_text": [
                                "检测结果仅对本次送检样本负责",
                                "咨询电话：",
                            ],
                        }
                    ]
                }
            }
        }
    )

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        report_data=report_data,
        visual_render="all",
        visual_render_required=True,
        visual_render_output_dir=str(tmp_path / "rendered"),
    )

    assert qa["status"] == "FAIL"
    check = qa["checks"]["blank_page_detection"]
    assert check["status"] == "FAIL"
    assert any(path.endswith("-2.png") for path in check["unexpected_low_content_pages"])
    assert check["expected_sparse_pages"] == []


def test_qa_report_never_accepts_blank_page_as_sparse_exception(tmp_path, monkeypatch):
    docx_path = tmp_path / "blank_exception.docx"
    Document().save(docx_path)

    def fake_render(docx_path, *, output_dir, **_kwargs):
        from PIL import Image, ImageDraw

        output_dir.mkdir(parents=True, exist_ok=True)
        pages = []
        for page_no in (1, 2, 3):
            page = output_dir / f"{Path(docx_path).stem}-{page_no}.png"
            image = Image.new("RGB", (200, 300), "white")
            if page_no != 2:
                ImageDraw.Draw(image).rectangle((25, 65, 175, 235), fill="black")
            image.save(page)
            pages.append(page)
        return pages

    monkeypatch.setattr("reportgen.core.qa_report.render_docx_to_pngs", fake_render)
    monkeypatch.setattr(
        "reportgen.core.qa_report._extract_pdf_page_texts",
        lambda *_args, **_kwargs: (
            {2: "检测结果仅对本次送检样本负责 咨询电话：00000000。"},
            None,
        ),
    )
    report_data = ReportData(
        context={
            "panel_style": {
                "visual_qa": {
                    "expected_sparse_pages": [
                        {
                            "id": "legal_notice",
                            "required_text": [
                                "检测结果仅对本次送检样本负责",
                                "咨询电话：",
                            ],
                        }
                    ]
                }
            }
        }
    )

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        report_data=report_data,
        visual_render="all",
        visual_render_required=True,
        visual_render_output_dir=str(tmp_path / "rendered"),
    )

    assert qa["status"] == "FAIL"
    check = qa["checks"]["blank_page_detection"]
    assert any(path.endswith("-2.png") for path in check["blank_pages"])
    assert check["expected_sparse_pages"] == []


def test_qa_report_visual_render_optional_failure_warns(tmp_path, monkeypatch):
    docx_path = tmp_path / "visual_warn.docx"
    Document().save(docx_path)

    def fail_render(*_args, **_kwargs):
        raise RuntimeError("LibreOffice failed")

    monkeypatch.setattr("reportgen.core.qa_report.render_docx_to_pngs", fail_render)

    qa = build_docx_qa_report(output_file=str(docx_path), visual_render="first")

    assert qa["status"] == "WARN"
    assert qa["checks"]["visual_render"]["status"] == "WARN"
    assert any(issue["code"] == "VISUAL_RENDER_FAILED" for issue in qa["issues"])


def test_qa_report_visual_render_required_failure_fails(tmp_path, monkeypatch):
    docx_path = tmp_path / "visual_fail.docx"
    Document().save(docx_path)

    monkeypatch.setattr(
        "reportgen.core.qa_report.render_docx_to_pngs",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(RuntimeError("render failed")),
    )

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        visual_render="all",
        visual_render_required=True,
    )

    assert qa["status"] == "FAIL"
    assert qa["checks"]["visual_render"]["status"] == "FAIL"
    assert any(issue["code"] == "VISUAL_RENDER_FAILED" for issue in qa["issues"])


def test_qa_report_records_pipeline_summary(tmp_path):
    docx_path = tmp_path / "pipeline.docx"
    doc = Document()
    doc.add_paragraph("已生成报告")
    doc.save(docx_path)

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        stage_results=[
            {"name": "PanelResolutionStage", "status": "PASS"},
            {"name": "TemplateContractStage", "status": "WARN"},
        ],
        stage_results_file=str(docx_path.with_suffix(".stage_results.json")),
    )

    assert qa["status"] == "WARN"
    assert qa["pipeline"]["status"] == "WARN"
    assert qa["checks"]["pipeline"]["warning_stages"] == ["TemplateContractStage"]
    assert any(issue["code"] == "PIPELINE_WARN" for issue in qa["issues"])


def test_qa_report_records_rule_provenance(tmp_path):
    docx_path = tmp_path / "rules.docx"
    doc = Document()
    doc.add_paragraph("已生成报告")
    doc.save(docx_path)

    rule_provenance = {
        "schema_version": "1.0",
        "panel_id": "crc_358_msi",
        "status": "PASS",
        "ok": True,
        "file_count": 1,
        "files": [
            {
                "rule_name": "report_text",
                "rule_id": "report_text",
                "schema_version": "1.0",
                "version": "0.1.0",
                "status": "draft",
                "updated": "2026-05-21",
                "path": "panels/crc_358_msi/rules/report_text.yaml",
                "sha256": "abc",
            }
        ],
        "issues": [],
    }

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        rule_provenance=rule_provenance,
    )

    assert qa["status"] == "PASS"
    assert qa["rules"]["panel_id"] == "crc_358_msi"
    assert qa["checks"]["rules"]["status"] == "PASS"
    assert qa["checks"]["rules"]["files"][0]["version"] == "0.1.0"


def test_field_provenance_masks_sensitive_values_and_records_sources(tmp_path):
    docx_path = tmp_path / "report.docx"
    doc = Document()
    doc.add_paragraph("报告")
    doc.save(docx_path)
    excel_data = _excel(
        tmp_path,
        single_values={
            "患者姓名": "表单姓名",
            "TMB": "6.5",
            "MSI状态": "MSS",
        },
    )
    excel_data.metadata["sample_id_from_filename"] = "ZZ999999"
    excel_data.metadata["field_source_overrides"] = {
        "patient_name": {
            "source": "form",
            "source_key": "患者姓名",
            "source_detail": "web_clinical_form",
        }
    }
    report_data = ReportData()
    report_data.set_field("patient_name", "表单姓名")
    report_data.set_field("sample_id", "ZZ999999")
    report_data.set_field("tmb_value", "6.5")
    report_data.set_field("tmb_status", "L")
    report_data.set_field("msi_status", "MSS")

    provenance = build_field_provenance_report(
        output_file=str(docx_path),
        report_data=report_data,
        excel_data=excel_data,
        config_loader=ConfigLoader(config_dir=str(ROOT / "config"), log_level="ERROR"),
        project_type="crc_358_msi",
        project_name="结直肠癌358基因+MSI",
    )

    fields = provenance["fields"]
    assert fields["patient_name"]["source"] == "form"
    assert fields["patient_name"]["value"] != "表单姓名"
    assert fields["sample_id"]["source"] == "filename"
    assert fields["sample_id"]["value"] != "ZZ999999"
    assert fields["tmb_value"]["source"] == "excel"
    assert fields["tmb_status"]["source"] == "rule"
    assert fields["msi_status"]["source"] == "excel"


def test_write_field_provenance_report_creates_sidecar_json(tmp_path):
    docx_path = tmp_path / "report.docx"
    doc = Document()
    doc.add_paragraph("报告")
    doc.save(docx_path)
    payload = {"schema_version": "1.0", "fields": {}}

    out_path = Path(write_field_provenance_report(payload, str(docx_path)))

    assert out_path.name == "report.field_provenance.json"
    assert out_path.exists()


def test_qa_report_references_field_provenance(tmp_path):
    docx_path = tmp_path / "clean.docx"
    doc = Document()
    doc.add_paragraph("已生成报告")
    doc.save(docx_path)
    provenance = {
        "fields": {
            "sample_id": {"source": "filename"},
            "patient_name": {"source": "form"},
        }
    }

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        field_provenance=provenance,
        field_provenance_file=str(docx_path.with_suffix(".field_provenance.json")),
    )

    assert qa["status"] == "PASS"
    assert qa["field_provenance_file"].endswith(".field_provenance.json")
    assert qa["checks"]["field_provenance"]["key_field_sources"] == {
        "sample_id": "filename",
        "patient_name": "form",
    }


def test_qa_report_records_template_contract_failure(tmp_path):
    docx_path = tmp_path / "clean.docx"
    doc = Document()
    doc.add_paragraph("已生成报告")
    doc.save(docx_path)

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        template_contract={
            "ok": False,
            "missing_paths": [],
            "missing_lists": [],
            "missing_row_fields": {},
            "declared_contract": {
                "missing_required_variables": ["sample_id"],
                "missing_required_lists": [],
                "missing_required_tables": ["variant_detail"],
                "table_errors": {},
            },
        },
    )

    assert qa["status"] == "FAIL"
    assert qa["checks"]["template_contract"]["status"] == "FAIL"
    assert qa["checks"]["template_contract"]["missing_required_variables"] == ["sample_id"]
    assert any(i["code"] == "TEMPLATE_CONTRACT_FAILED" for i in qa["issues"])


def test_processor_runner_records_skip_success_and_error():
    class _Logger:
        def __init__(self):
            self.warnings = []

        def warning(self, message, **kwargs):
            self.warnings.append((message, kwargs))

    class _Processor:
        def __init__(self, name, status):
            self.name = name
            self.status = status
            self.warning_message = f"{name} failed"

        def enabled(self, _ctx):
            return self.status != "skip"

        def run(self, _ctx):
            if self.status == "error":
                raise RuntimeError("boom")

    logger = _Logger()
    ctx = ProcessorContext(
        renderer=object(),
        output_path="out.docx",
        template_path="template.docx",
        template_context={},
        logger=logger,
    )

    results = run_processors(
        [
            _Processor("ok_processor", "ok"),
            _Processor("skip_processor", "skip"),
            _Processor("bad_processor", "error"),
        ],
        ctx,
    )

    assert [r.status for r in results] == ["OK", "SKIPPED", "ERROR"]
    assert results[2].error == "boom"
    assert logger.warnings[0][0] == "bad_processor failed"


def test_template_renderer_default_processors_include_key_m1_processors():
    renderer = TemplateRenderer(log_level="ERROR")
    names = [processor.name for processor in renderer.build_post_render_processors()]

    assert "bullet_lists" in names
    assert "variant_tables" in names
    assert "toc_refresh" in names
    assert "front_matter_spacing" in names
    assert "blank_page_cleanup" in names
    assert "underlines_and_styles" in names


def test_front_matter_spacing_restores_report_guide_golden_offset(tmp_path):
    docx_path = tmp_path / "front_matter.docx"
    doc = Document()
    doc.add_paragraph("检测报告")
    page_break = doc.add_paragraph("")
    page_break.add_run().add_break(WD_BREAK.PAGE)
    for _ in range(6):
        doc.add_paragraph("")
    doc.add_paragraph("报告导读")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._normalize_front_matter_spacing(str(docx_path))

    import xml.etree.ElementTree as ET
    from zipfile import ZipFile

    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    w_p = f"{{{ns_w}}}p"
    w_t = f"{{{ns_w}}}t"
    w_br = f"{{{ns_w}}}br"
    w_type = f"{{{ns_w}}}type"

    with ZipFile(docx_path) as zin:
        root = ET.fromstring(zin.read("word/document.xml"))

    paragraphs = [elem for elem in root.iter(w_p)]

    def text(elem):
        return "".join((node.text or "") for node in elem.iter(w_t)).strip()

    def has_page_break(elem):
        return any(node.attrib.get(w_type) == "page" for node in elem.iter(w_br))

    guide_idx = next(idx for idx, elem in enumerate(paragraphs) if text(elem) == "报告导读")

    assert has_page_break(paragraphs[guide_idx - 31])
    assert text(paragraphs[guide_idx - 32]) == "检测报告"
    assert all(text(elem) == "" for elem in paragraphs[guide_idx - 30 : guide_idx])
    assert not any(has_page_break(elem) for elem in paragraphs[guide_idx - 30 : guide_idx])


def test_front_matter_spacing_honors_panel_page_flow_config(tmp_path):
    docx_path = tmp_path / "front_matter_panel_flow.docx"
    doc = Document()
    doc.add_paragraph("检测报告")
    page_break = doc.add_paragraph("")
    page_break.add_run().add_break(WD_BREAK.PAGE)
    for _ in range(6):
        doc.add_paragraph("")
    doc.add_paragraph("报告导读")
    notice = doc.add_paragraph("本表仅展示当前面板已输出的精确事件结果；评审用。")
    notice.add_run().add_break(WD_BREAK.PAGE)
    standalone_break = doc.add_paragraph("")
    standalone_break.add_run().add_break(WD_BREAK.PAGE)
    unrelated = doc.add_paragraph("其它段落")
    unrelated.add_run().add_break(WD_BREAK.PAGE)
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    _run_front_matter_spacing(
        ProcessorContext(
            renderer=renderer,
            output_path=str(docx_path),
            template_path=str(docx_path),
            template_context={
                "panel_style": {
                    "front_matter": {
                        "guide_spacer_count": 30,
                        "insert_page_break": False,
                        "remove_page_break_after_text_prefixes": [
                            "本表仅展示当前面板已输出的精确事件结果"
                        ],
                    }
                }
            },
            logger=SimpleNamespace(),
        )
    )

    import xml.etree.ElementTree as ET

    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    w_p = f"{{{ns_w}}}p"
    w_t = f"{{{ns_w}}}t"
    w_br = f"{{{ns_w}}}br"
    w_type = f"{{{ns_w}}}type"
    with ZipFile(docx_path) as zin:
        root = ET.fromstring(zin.read("word/document.xml"))
    paragraphs = [elem for elem in root.iter(w_p)]

    def text(elem):
        return "".join((node.text or "") for node in elem.iter(w_t)).strip()

    def has_page_break(elem):
        return any(node.attrib.get(w_type) == "page" for node in elem.iter(w_br))

    guide_idx = next(idx for idx, elem in enumerate(paragraphs) if text(elem) == "报告导读")
    assert text(paragraphs[guide_idx - 31]) == "检测报告"
    assert not any(has_page_break(elem) for elem in paragraphs[: guide_idx + 1])
    assert all(text(elem) == "" for elem in paragraphs[guide_idx - 30 : guide_idx])
    rendered_notice = next(
        elem for elem in paragraphs if text(elem).startswith("本表仅展示当前面板")
    )
    rendered_standalone_break = paragraphs[paragraphs.index(rendered_notice) + 1]
    rendered_unrelated = next(elem for elem in paragraphs if text(elem) == "其它段落")
    assert has_page_break(rendered_notice) is False
    assert has_page_break(rendered_standalone_break) is False
    assert has_page_break(rendered_unrelated) is True


def test_template_renderer_can_build_panel_declared_processors_only():
    renderer = TemplateRenderer(log_level="ERROR")
    processors = renderer.build_post_render_processors(["bullet_lists"])

    assert [processor.name for processor in processors] == ["bullet_lists"]


def test_template_renderer_can_disable_panel_processors(tmp_path):
    docx_path = tmp_path / "no_processors.docx"
    doc = Document()
    p = doc.add_paragraph("")
    ppr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(num_id)
    ppr.append(num_pr)
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._run_post_render_processors(
        str(docx_path),
        {},
        str(docx_path),
        processor_names=[],
    )

    rendered = Document(docx_path)
    assert rendered.paragraphs[0]._p.pPr is not None
    assert renderer.last_processor_report == []


def test_template_renderer_blocks_critical_part3_processor_error(tmp_path, monkeypatch):
    docx_path = tmp_path / "critical_part3.docx"
    doc = Document()
    doc.add_paragraph("__PART3_MARKER__")
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    monkeypatch.setattr(
        renderer,
        "_render_part3_formatted",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(RuntimeError("part3 failed")),
    )

    with pytest.raises(RuntimeError, match="Critical DOCX post-processing failed"):
        renderer._run_post_render_processors(
            str(docx_path),
            {},
            str(docx_path),
            processor_names=["part3_formatted_sections"],
        )

    assert renderer.last_processor_report[0]["name"] == "part3_formatted_sections"
    assert renderer.last_processor_report[0]["status"] == "ERROR"


def test_template_renderer_does_not_leave_output_after_critical_failure(
    tmp_path,
    monkeypatch,
):
    template_path = tmp_path / "atomic_template.docx"
    output_path = tmp_path / "atomic_output.docx"
    doc = Document()
    doc.add_paragraph("__PART3_MARKER__")
    doc.save(template_path)

    renderer = TemplateRenderer(log_level="ERROR")
    monkeypatch.setattr(
        renderer,
        "_render_part3_formatted",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(RuntimeError("synthetic part3 failure")),
    )

    with pytest.raises(ValueError, match="Critical DOCX post-processing failed"):
        renderer.render(
            str(template_path),
            ReportData(),
            str(output_path),
            post_processor_names=["part3_formatted_sections"],
        )

    assert not output_path.exists()
    assert list(tmp_path.glob(".atomic_output.*.docx")) == []


def _docx_xml_signature(docx_path: Path) -> dict[str, str]:
    """Return a deterministic XML signature for idempotency checks."""
    import re
    from zipfile import ZipFile

    signature = {}
    with ZipFile(docx_path, "r") as zf:
        for name in sorted(zf.namelist()):
            if name == "docProps/core.xml":
                continue
            if not (name.endswith(".xml") or name.endswith(".rels")):
                continue
            text = zf.read(name).decode("utf-8", "ignore")
            text = re.sub(r'\s+w:rsid\w+="[^"]*"', "", text)
            signature[name] = text
    return signature


def _assert_processors_idempotent(
    tmp_path: Path,
    processor_names: list[str],
    build_docx,
    *,
    context: dict | None = None,
    monkeypatch=None,
) -> Path:
    docx_path = tmp_path / f"idempotent_{'_'.join(processor_names)}.docx"
    build_docx(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    if monkeypatch is not None:
        monkeypatch.setattr(
            renderer,
            "_refresh_fields_with_native_engine",
            lambda file_path: None,
        )
        monkeypatch.setattr(
            renderer,
            "_populate_static_toc_page_numbers",
            lambda file_path, context=None: None,
        )

    renderer._run_post_render_processors(
        str(docx_path),
        context or {},
        str(docx_path),
        processor_names=processor_names,
    )
    assert all(row["status"] in {"OK", "SKIPPED"} for row in renderer.last_processor_report)
    first_signature = _docx_xml_signature(docx_path)

    renderer._run_post_render_processors(
        str(docx_path),
        context or {},
        str(docx_path),
        processor_names=processor_names,
    )
    assert all(row["status"] in {"OK", "SKIPPED"} for row in renderer.last_processor_report)
    assert _docx_xml_signature(docx_path) == first_signature
    return docx_path


def _write_empty_numbered_docx(docx_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph("")
    ppr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(num_id)
    ppr.append(num_pr)
    doc.save(docx_path)


def _set_table_widths(table, widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))


def _write_overwide_variant_table_docx(docx_path: Path) -> None:
    doc = Document()
    table = doc.add_table(rows=3, cols=9)
    rows = [
        [
            "基因名称",
            "基因突变信息",
            "基因突变信息",
            "基因突变信息",
            "基因突变信息",
            "基因突变信息",
            "基因突变信息",
            "靶向药物信息",
            "靶向药物信息",
        ],
        [
            "基因名称",
            "转录本号",
            "染色体",
            "外显子",
            "位点",
            "突变类型",
            "频率",
            "潜在获益靶向药物",
            "可能耐药或慎重药物",
        ],
        [
            "KRAS",
            "NM_004985.5",
            "12",
            "2",
            "c.34G>A",
            "点突变",
            "46.29",
            "司美替尼（C）",
            "西妥昔单抗（A）",
        ],
    ]
    for row, values in zip(table.rows, rows):
        for idx, value in enumerate(values):
            row.cells[idx].text = value
    _set_table_widths(table, [2400] * 9)
    doc.save(docx_path)


def _write_trailing_blank_page_docx(docx_path: Path) -> None:
    doc = Document()
    paragraph = doc.add_paragraph("正文结尾")
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.save(docx_path)


def _write_toc_update_docx(docx_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("第一部分\t1")
    doc.add_heading("第一部分", level=1)
    doc.save(docx_path)


def _write_crc_style_docx(docx_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("本次共检出体细胞变异：1 个")
    _add_crc_style_qa_tables(doc)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.underline = True
    doc.save(docx_path)


def test_bullet_list_processor_is_idempotent(tmp_path):
    docx_path = _assert_processors_idempotent(
        tmp_path,
        ["bullet_lists"],
        _write_empty_numbered_docx,
    )

    rendered = Document(docx_path)
    assert not any(
        paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
        for paragraph in rendered.paragraphs
    )


def test_variant_tables_processor_is_idempotent(tmp_path):
    _assert_processors_idempotent(
        tmp_path,
        ["variant_tables"],
        _write_overwide_variant_table_docx,
    )


def test_blank_page_cleanup_processor_is_idempotent(tmp_path):
    _assert_processors_idempotent(
        tmp_path,
        ["blank_page_cleanup"],
        _write_trailing_blank_page_docx,
    )


def test_blank_heading_cleanup_does_not_cross_tables(tmp_path):
    docx_path = tmp_path / "heading_cleanup.docx"
    doc = Document()
    doc.add_paragraph("正文")
    doc.add_table(rows=1, cols=2)
    doc.add_paragraph("")
    doc.add_table(rows=1, cols=2)
    breaker = doc.add_paragraph()
    breaker.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("2. 结直肠癌诊疗知识")
    doc.add_paragraph("章节正文")
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._remove_blank_page_breaks_before_headings(
        str(docx_path),
        ("2. 结直肠癌诊疗知识",),
    )

    cleaned = Document(docx_path)
    assert len(cleaned.tables) == 2
    paragraphs = [p for p in cleaned.paragraphs if p.text.strip()]
    assert [p.text.strip() for p in paragraphs][-2:] == [
        "2. 结直肠癌诊疗知识",
        "章节正文",
    ]
    assert "w:keepNext" in paragraphs[-2]._p.xml


def test_blank_heading_cleanup_removes_legacy_break_before_targeted_results(
    tmp_path,
):
    docx_path = tmp_path / "targeted_results_heading_cleanup.docx"
    doc = Document()
    doc.add_paragraph("免疫治疗说明末行")
    breaker = doc.add_paragraph()
    breaker.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("2. 靶向药物相关检测结果")
    doc.add_paragraph("检测结果正文")
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._remove_blank_page_breaks_before_headings(
        str(docx_path),
        ("2. 靶向药物相关检测结果",),
    )

    cleaned = Document(docx_path)
    assert not any(
        'w:type="page"' in paragraph._p.xml
        for paragraph in cleaned.paragraphs
        if not paragraph.text.strip()
    )
    headings = [
        paragraph
        for paragraph in cleaned.paragraphs
        if paragraph.text.strip() == "2. 靶向药物相关检测结果"
    ]
    assert len(headings) == 1
    assert "w:keepNext" in headings[0]._p.xml


def test_exact_section_headings_use_idempotent_page_break_before(tmp_path):
    docx_path = tmp_path / "forced_heading_breaks.docx"
    headings = (
        "2.2 其它潜在获益上市药物提示*",
        "2.3 NCCN 推荐临床常规靶向药物相关基因检测结果（不限于本癌种）",
        "结肠癌NCCN指南（2022 V1）和直肠癌NCCN指南（2022 V1）",
        "3. 阅读说明",
        "4. 基因检测列表",
        "5. 参考文献",
        "本次检测质控结果",
    )
    doc = Document()
    doc.add_paragraph("上一节正文")
    breaker = doc.add_paragraph()
    breaker.add_run().add_break(WD_BREAK.PAGE)
    # LibreOffice can leave a body-level bookmark boundary between the blank
    # template break and its heading. It must not shield the break from cleanup.
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), "42")
    doc._body._element.insert(len(doc._body._element) - 1, bookmark_end)
    for index, heading in enumerate(headings):
        heading_paragraph = doc.add_paragraph()
        if index == 0:
            # Legacy templates may also embed a second break in the heading.
            heading_paragraph.add_run().add_break(WD_BREAK.PAGE)
        heading_paragraph.add_run(heading)
        doc.add_paragraph("章节正文")
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._remove_blank_page_breaks_before_headings(str(docx_path), headings)
    renderer._enforce_page_break_before_headings(str(docx_path), headings)
    renderer._enforce_page_break_before_headings(str(docx_path), headings)

    rendered = Document(docx_path)
    heading_paragraphs = {
        paragraph.text.strip(): paragraph
        for paragraph in rendered.paragraphs
        if paragraph.text.strip() in headings
    }
    assert set(heading_paragraphs) == set(headings)
    assert all(
        paragraph.paragraph_format.page_break_before is True
        and paragraph.paragraph_format.keep_with_next is True
        for paragraph in heading_paragraphs.values()
    )
    assert all(
        'w:type="page"' not in paragraph._p.xml
        and "w:lastRenderedPageBreak" not in paragraph._p.xml
        for paragraph in heading_paragraphs.values()
    )
    assert not any(
        not paragraph.text.strip() and 'w:type="page"' in paragraph._p.xml
        for paragraph in rendered.paragraphs
    )


def test_collapsed_reference_notice_and_qc_share_the_page(tmp_path):
    docx_path = tmp_path / "collapsed-reference-qc.docx"
    doc = Document()
    notice = doc.add_paragraph("本报告未生成患者级动态参考文献；历史病例固定参考文献已移除。")
    notice.paragraph_format.page_break_before = True
    notice.add_run().add_break(WD_BREAK.PAGE)
    qc = doc.add_paragraph("本次检测质控结果")
    qc.paragraph_format.page_break_before = True
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._enforce_page_break_before_headings(
        str(docx_path),
        ("本次检测质控结果",),
    )
    renderer._enforce_page_break_before_headings(
        str(docx_path),
        ("本次检测质控结果",),
    )

    rendered = Document(docx_path)
    rendered_notice, rendered_qc = rendered.paragraphs[:2]
    assert rendered_notice.paragraph_format.page_break_before is False
    assert rendered_qc.paragraph_format.page_break_before is False
    assert rendered_qc.paragraph_format.keep_with_next is True
    assert 'w:type="page"' not in rendered_notice._p.xml
    assert "w:lastRenderedPageBreak" not in rendered_notice._p.xml


def test_disabled_lung_part3_notice_and_reading_guide_share_the_page(tmp_path):
    docx_path = tmp_path / "disabled-part3-reading-guide.docx"
    doc = Document()
    doc.add_paragraph("第三部分：基因变异及相应靶向/免疫药物解析")
    notice = doc.add_paragraph(
        "本项目肺癌专属知识当前未启用，暂不提供第三部分的患者级基因解释和药物解析。"
    )
    notice.add_run().add_break(WD_BREAK.PAGE)
    reading = doc.add_paragraph("3. 阅读说明")
    reading.paragraph_format.page_break_before = True
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._enforce_page_break_before_headings(
        str(docx_path),
        ("3. 阅读说明",),
    )
    renderer._enforce_page_break_before_headings(
        str(docx_path),
        ("3. 阅读说明",),
    )

    rendered = Document(docx_path)
    rendered_heading, rendered_notice, rendered_reading = rendered.paragraphs[:3]
    assert rendered_heading.paragraph_format.keep_with_next is True
    assert rendered_notice.paragraph_format.page_break_before is False
    assert rendered_reading.paragraph_format.page_break_before is False
    assert rendered_reading.paragraph_format.keep_with_next is True
    assert 'w:type="page"' not in rendered_notice._p.xml


def test_drug_brand_summary_bolds_only_square_bracket_fragments(tmp_path):
    docx_path = tmp_path / "brand_brackets.docx"
    marker = "上表涉及的已上市的药物名称及对应的商品名称"
    doc = Document()
    targeted = doc.add_paragraph(f"2. {marker}：奥拉帕利[利普卓]、芦卡帕利[Rubraca]。")
    targeted.runs[0].font.bold = False
    immune = doc.add_paragraph(f"3. {marker}：帕博利珠单抗[可瑞达]。")
    immune.runs[0].font.bold = False
    doc.add_paragraph("其他正文[不应处理]")
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._bold_drug_brand_brackets(str(docx_path))
    renderer._bold_drug_brand_brackets(str(docx_path))

    rendered = Document(docx_path)
    summaries = [p for p in rendered.paragraphs if marker in p.text]
    assert len(summaries) == 2
    for paragraph in summaries:
        bracket_runs = [run for run in paragraph.runs if run.text.startswith("[")]
        non_bracket_runs = [
            run for run in paragraph.runs if run.text and not run.text.startswith("[")
        ]
        assert bracket_runs and all(run.bold is True for run in bracket_runs)
        assert non_bracket_runs and all(run.bold is not True for run in non_bracket_runs)


def test_part3_variant_heading_has_ten_point_spacing_before_and_after(tmp_path):
    docx_path = tmp_path / "part3_variant_spacing.docx"
    doc = Document()
    for text in [
        "第三部分：基因变异及相应靶向/免疫药物解析",
        "u GNAS：c.1030G>A，p.E344K；19.04%",
        "基因简介：",
        "GNAS基因简介正文。",
        "3. 阅读说明",
    ]:
        doc.add_paragraph(text)
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._restore_part3_dynamic_styles(str(docx_path), {})

    rendered = Document(docx_path)
    heading = next(
        paragraph
        for paragraph in rendered.paragraphs
        if paragraph.text.strip().startswith("u GNAS：")
    )
    assert heading.paragraph_format.space_before.pt == pytest.approx(10.0)
    assert heading.paragraph_format.space_after.pt == pytest.approx(10.0)


def test_part3_drug_analysis_labels_do_not_emit_keepnext_marker(tmp_path):
    docx_path = tmp_path / "part3_labels.docx"
    doc = Document()
    for text in [
        "第三部分：基因变异及相应靶向/免疫药物解析",
        "u GNAS：c.1030G>A，p.E344K；19.04%",
        "基因简介：",
        "GNAS基因简介正文。",
        "靶向药物/免疫用药提示解析",
        "潜在获益靶向/免疫药物解析",
        "KRAS：c.34G>A，p.G12S突变相应靶向药物",
        "西妥昔单抗",
        "基因变异与药物关联分析：",
        "关联分析正文",
        "药物疗效临床解析：",
        "临床解析正文",
        "3. 阅读说明",
    ]:
        doc.add_paragraph(text)
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._restore_part3_dynamic_styles(str(docx_path), {})

    rendered = Document(docx_path)
    labels = [
        paragraph
        for paragraph in rendered.paragraphs
        if paragraph.text.strip() in {"基因变异与药物关联分析：", "药物疗效临床解析："}
    ]
    assert [paragraph.text.strip() for paragraph in labels] == [
        "基因变异与药物关联分析：",
        "药物疗效临床解析：",
    ]
    assert all("w:keepNext" not in paragraph._p.xml for paragraph in labels)

    subheading = next(
        paragraph
        for paragraph in rendered.paragraphs
        if paragraph.text.strip() == "潜在获益靶向/免疫药物解析"
    )
    assert "w:keepNext" in subheading._p.xml
    main_heading = next(
        paragraph
        for paragraph in rendered.paragraphs
        if paragraph.text.strip() == "靶向药物/免疫用药提示解析"
    )
    assert "w:keepNext" in main_heading._p.xml
    variant_heading = next(
        paragraph
        for paragraph in rendered.paragraphs
        if paragraph.text.strip().startswith("u GNAS：")
    )
    assert "w:keepNext" in variant_heading._p.xml


def test_pdf_footer_page_number_scans_bottom_lines():
    text = "表格内容\n--\n9\n--\n"

    assert TemplateRenderer._extract_pdf_footer_page_number(text) == 9


def test_toc_refresh_processor_is_idempotent(tmp_path, monkeypatch):
    _assert_processors_idempotent(
        tmp_path,
        ["toc_refresh"],
        _write_toc_update_docx,
        monkeypatch=monkeypatch,
    )


def test_final_refresh_cleanup_processor_is_idempotent(tmp_path, monkeypatch):
    _assert_processors_idempotent(
        tmp_path,
        ["final_refresh_cleanup"],
        _write_trailing_blank_page_docx,
        monkeypatch=monkeypatch,
    )


def test_underlines_and_styles_processor_is_idempotent(tmp_path):
    _assert_processors_idempotent(
        tmp_path,
        ["underlines_and_styles"],
        _write_crc_style_docx,
        context={"panel_style": _crc_panel_style()},
    )


def test_underlines_and_styles_finalizes_toc_after_layout_normalizers(monkeypatch):
    renderer = TemplateRenderer(log_level="ERROR")
    sequence = []
    no_op_methods = (
        "_remove_template_underlines",
        "_restore_variant_summary_table_style",
        "_restore_variant_detail_table_style",
        "_restore_biomarker_table_style",
        "_restore_clinical_result_table_style",
        "_restore_detection_content_underlines",
        "_render_patient_salutation",
        "_normalize_repeated_terminal_punctuation",
        "_restore_patient_letter_fill_underlines",
        "_restore_msi_result_emphasis",
        "_restore_part3_dynamic_styles",
        "_bold_drug_brand_brackets",
        "_render_inline_signatures",
        "_remove_empty_numbered_paragraphs",
    )
    for name in no_op_methods:
        monkeypatch.setattr(renderer, name, lambda *_args, **_kwargs: None)
    monkeypatch.setattr(
        renderer,
        "_recolor_part3_intro_marker",
        lambda *_args: sequence.append("recolor"),
    )
    monkeypatch.setattr(
        renderer,
        "_enforce_page_break_before_headings",
        lambda *_args: sequence.append("page_breaks"),
    )
    monkeypatch.setattr(
        renderer,
        "_normalize_reference_section_style",
        lambda *_args: sequence.append("references"),
    )
    monkeypatch.setattr(
        renderer,
        "_normalize_legal_notice_style",
        lambda *_args: sequence.append("legal_notice"),
    )
    monkeypatch.setattr(
        renderer,
        "_normalize_back_cover_artwork",
        lambda *_args: sequence.append("back_cover"),
    )
    monkeypatch.setattr(
        renderer,
        "_remove_page_breaks_after_text_prefixes",
        lambda *_args: sequence.append("page_flow"),
    )
    monkeypatch.setattr(
        renderer,
        "_populate_static_toc_page_numbers",
        lambda *_args: sequence.append("toc"),
    )

    _run_underlines_and_styles(
        SimpleNamespace(
            renderer=renderer,
            output_path="report.docx",
            template_context={
                "report_content": {"force_page_break_before_headings": ["5. 参考文献"]},
                "panel_style": {
                    "front_matter": {
                        "remove_page_break_after_text_prefixes": [
                            "本表仅展示当前面板已输出的精确事件结果"
                        ]
                    }
                },
            },
            logger=SimpleNamespace(warning=lambda *_args, **_kwargs: None),
        )
    )

    assert sequence == [
        "recolor",
        "page_breaks",
        "legal_notice",
        "references",
        "back_cover",
        "page_flow",
        "toc",
    ]


def test_underlines_and_styles_does_not_swallow_toc_convergence_failure(
    monkeypatch,
):
    renderer = TemplateRenderer(log_level="ERROR")
    no_op_methods = (
        "_remove_template_underlines",
        "_restore_variant_summary_table_style",
        "_restore_variant_detail_table_style",
        "_restore_biomarker_table_style",
        "_restore_clinical_result_table_style",
        "_restore_detection_content_underlines",
        "_render_patient_salutation",
        "_normalize_repeated_terminal_punctuation",
        "_restore_patient_letter_fill_underlines",
        "_restore_msi_result_emphasis",
        "_restore_part3_dynamic_styles",
        "_bold_drug_brand_brackets",
        "_render_inline_signatures",
        "_remove_empty_numbered_paragraphs",
        "_recolor_part3_intro_marker",
        "_normalize_legal_notice_style",
        "_normalize_reference_section_style",
        "_normalize_back_cover_artwork",
    )
    for name in no_op_methods:
        monkeypatch.setattr(renderer, name, lambda *_args, **_kwargs: None)
    monkeypatch.setattr(
        renderer,
        "_populate_static_toc_page_numbers",
        lambda *_args: (_ for _ in ()).throw(RuntimeError("目录页码未收敛")),
    )

    with pytest.raises(RuntimeError, match="目录页码未收敛"):
        _run_underlines_and_styles(
            SimpleNamespace(
                renderer=renderer,
                output_path="report.docx",
                template_context={},
                logger=SimpleNamespace(warning=lambda *_args, **_kwargs: None),
            )
        )

    assert "underlines_and_styles" not in CRITICAL_DOCX_PROCESSOR_NAMES
    assert "underlines_and_styles" in critical_docx_processor_names("crc_358_msi")
    assert "final_refresh_cleanup" in critical_docx_processor_names("crc_358_msi")
    assert "underlines_and_styles" not in critical_docx_processor_names("endometrial_29")


def test_fast_toc_skips_final_libreoffice_refresh(monkeypatch):
    calls = {"set_update_fields": 0, "refresh": 0}
    sequence = []

    class FakeRenderer:
        def _normalize_final_section_layout(self, *_args):
            pass

        def _cleanup_section_spacing(self, *_args):
            pass

        def _remove_standalone_page_breaks_before_pathway_tables(self, *_args):
            pass

        def _compact_gene_list_tables(self, *_args):
            pass

        def _normalize_quality_control_tables(self, *_args):
            pass

        def _optimize_variant_table_layout(self, *_args):
            pass

        def _cleanup_trailing_blank_page(self, *_args):
            pass

        def _remove_blank_page_breaks_before_headings(self, _path, headings):
            sequence.append(("heading_cleanup", tuple(headings)))

        def _refresh_fields_with_native_engine(self, *_args):
            calls["refresh"] += 1
            sequence.append("refresh")

        def _set_update_fields(self, *_args):
            calls["set_update_fields"] += 1

        def _normalize_toc_decoration_layout(self, *_args):
            pass

        def _restore_reviewed_body_headers(self, *_args):
            pass

    class FakeLogger:
        def info(self, *_args, **_kwargs):
            pass

        def warning(self, *_args, **_kwargs):
            pass

    monkeypatch.setenv("REPORTGEN_FAST_TOC", "1")
    ctx = SimpleNamespace(
        renderer=FakeRenderer(),
        output_path="report.docx",
        template_context={},
        logger=FakeLogger(),
    )

    _run_final_refresh_cleanup(ctx)

    assert calls["refresh"] == 0
    assert calls["set_update_fields"] == 1
    assert [name for name, _headings in sequence] == [
        "heading_cleanup",
        "heading_cleanup",
    ]
    assert all("2. 靶向药物相关检测结果" in headings for _name, headings in sequence)


def test_fast_toc_is_blocked_for_deterministic_production_panel(monkeypatch):
    monkeypatch.setenv("REPORTGEN_FAST_TOC", "1")
    renderer = TemplateRenderer(log_level="ERROR")
    for name in (
        "_normalize_final_section_layout",
        "_cleanup_section_spacing",
        "_remove_standalone_page_breaks_before_pathway_tables",
        "_compact_gene_list_tables",
        "_normalize_quality_control_tables",
        "_optimize_variant_table_layout",
        "_cleanup_trailing_blank_page",
        "_remove_blank_page_breaks_before_headings",
    ):
        monkeypatch.setattr(renderer, name, lambda *_args, **_kwargs: None)

    with pytest.raises(RuntimeError, match="禁止跳过最终 LibreOffice"):
        _run_final_refresh_cleanup(
            SimpleNamespace(
                renderer=renderer,
                output_path="report.docx",
                template_context={"_require_deterministic_layout": True},
                logger=SimpleNamespace(
                    info=lambda *_args, **_kwargs: None,
                    warning=lambda *_args, **_kwargs: None,
                ),
            )
        )


def test_final_refresh_recleans_heading_breaks_after_native_refresh(monkeypatch):
    sequence = []

    class FakeRenderer:
        def _normalize_final_section_layout(self, *_args):
            pass

        def _cleanup_section_spacing(self, *_args):
            pass

        def _remove_standalone_page_breaks_before_pathway_tables(self, *_args):
            pass

        def _compact_gene_list_tables(self, *_args):
            sequence.append("gene_list")

        def _restore_gene_list_table_borders(self, *_args):
            sequence.append("gene_list_borders")

        def _normalize_quality_control_tables(self, *_args):
            pass

        def _optimize_variant_table_layout(self, *_args):
            pass

        def _cleanup_trailing_blank_page(self, *_args):
            pass

        def _remove_blank_page_breaks_before_headings(self, *_args):
            sequence.append("heading_cleanup")

        def _refresh_fields_with_native_engine(self, *_args):
            sequence.append("refresh")

        def _set_update_fields(self, *_args):
            pass

        def _normalize_toc_decoration_layout(self, *_args):
            pass

        def _restore_reviewed_body_headers(self, *_args):
            pass

    class FakeLogger:
        def info(self, *_args, **_kwargs):
            pass

        def warning(self, *_args, **_kwargs):
            pass

    monkeypatch.delenv("REPORTGEN_FAST_TOC", raising=False)
    monkeypatch.delenv("REPORTGEN_SKIP_FINAL_LO_REFRESH", raising=False)
    _run_final_refresh_cleanup(
        SimpleNamespace(
            renderer=FakeRenderer(),
            output_path="report.docx",
            template_context={},
            logger=FakeLogger(),
        )
    )

    assert sequence == [
        "gene_list",
        "heading_cleanup",
        "refresh",
        "heading_cleanup",
        "gene_list_borders",
    ]


def test_final_refresh_restores_gene_list_borders_dropped_by_libreoffice(tmp_path, monkeypatch):
    """The native field refresh must not erase the reviewed gene-list frame."""
    docx_path = tmp_path / "gene_list_after_refresh.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Gene List for MLseq (n=358)"
    table.rows[1].cells[0].text = "ABL1"
    table.rows[1].cells[1].text = "ABL2"
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")

    def simulate_libreoffice_refresh(path):
        refreshed = Document(path)
        tbl_pr = refreshed.tables[0]._tbl.tblPr
        borders = tbl_pr.find(qn("w:tblBorders"))
        assert borders is not None  # the pre-refresh normalizer added it
        tbl_pr.remove(borders)
        refreshed.save(path)

    monkeypatch.setattr(
        renderer,
        "_refresh_fields_with_native_engine",
        simulate_libreoffice_refresh,
    )
    monkeypatch.setattr(renderer, "_set_update_fields", lambda *_args: None)
    monkeypatch.setattr(renderer, "_normalize_toc_decoration_layout", lambda *_args: None)
    monkeypatch.setattr(renderer, "_restore_reviewed_body_headers", lambda *_args: None)
    monkeypatch.delenv("REPORTGEN_FAST_TOC", raising=False)
    monkeypatch.delenv("REPORTGEN_SKIP_FINAL_LO_REFRESH", raising=False)

    _run_final_refresh_cleanup(
        SimpleNamespace(
            renderer=renderer,
            output_path=str(docx_path),
            template_context={},
            logger=SimpleNamespace(
                info=lambda *_args, **_kwargs: None,
                warning=lambda *_args, **_kwargs: None,
            ),
        )
    )

    rendered = Document(docx_path)
    borders = rendered.tables[0]._tbl.tblPr.find(qn("w:tblBorders"))
    assert borders is not None
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        assert border is not None
        assert border.get(qn("w:val")) == "single"


def test_fast_toc_skips_static_toc_pdf_detection(tmp_path, monkeypatch):
    docx_path = tmp_path / "report.docx"
    doc = Document()
    doc.add_paragraph("目    录")
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    calls = {"set_update_fields": 0}

    monkeypatch.setenv("REPORTGEN_FAST_TOC", "1")
    monkeypatch.setattr(
        renderer,
        "_set_update_fields",
        lambda *_args: calls.__setitem__("set_update_fields", calls["set_update_fields"] + 1),
    )
    monkeypatch.setattr(
        renderer,
        "_document_contains_toc_or_static_toc",
        lambda *_args: pytest.fail("fast TOC mode should skip PDF TOC detection"),
    )

    renderer._populate_static_toc_page_numbers(str(docx_path), {})

    assert calls["set_update_fields"] == 1


def test_static_toc_missing_tools_blocks_deterministic_panel(tmp_path, monkeypatch):
    docx_path = tmp_path / "report.docx"
    Document().save(docx_path)
    renderer = TemplateRenderer(log_level="ERROR")
    monkeypatch.delenv("REPORTGEN_FAST_TOC", raising=False)
    monkeypatch.delenv("REPORTGEN_SKIP_STATIC_TOC_PAGE_NUMBERS", raising=False)
    monkeypatch.setattr(
        renderer,
        "_document_contains_toc_or_static_toc",
        lambda *_args: True,
    )
    monkeypatch.setattr(shutil, "which", lambda _name: None)

    with pytest.raises(RuntimeError, match="缺少目录分页工具"):
        renderer._populate_static_toc_page_numbers(
            str(docx_path),
            {"_require_deterministic_layout": True},
        )


def test_pdf_layout_probe_freezes_only_the_disposable_docx_fields(tmp_path):
    source_path = tmp_path / "source.docx"
    probe_path = tmp_path / "probe.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)
    instruction_run = paragraph.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.text = " PAGEREF _ReportGenToc_1 \\h "
    instruction_run._r.append(instruction)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end.set(qn("w:dirty"), "true")
    end_run._r.append(end)
    doc.save(source_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._set_update_fields(str(source_path))
    renderer._set_word_compat_pagination(str(source_path))
    shutil.copy2(source_path, probe_path)
    renderer._freeze_fields_for_pdf_layout_probe(str(probe_path))

    source_document = _read_docx_part(source_path, "word/document.xml")
    source_settings = _read_docx_part(source_path, "word/settings.xml")
    assert 'w:dirty="true"' in source_document
    assert "<w:updateFields" in source_settings
    assert "<w:usePrinterMetrics" in source_settings

    probe_document = _read_docx_part(probe_path, "word/document.xml")
    probe_settings = _read_docx_part(probe_path, "word/settings.xml")
    assert "w:dirty=" not in probe_document
    assert 'w:fldLock="true"' in probe_document
    assert "<w:updateFields" not in probe_settings
    assert "<w:usePrinterMetrics" not in probe_settings


def test_static_toc_page_numbers_converge_after_toc_rebuild(tmp_path, monkeypatch):
    docx_path = tmp_path / "report.docx"
    Document().save(docx_path)
    renderer = TemplateRenderer(log_level="ERROR")
    first = {"检测结果小结": 2, "参考文献": 67}
    shifted = {"检测结果小结": 3, "参考文献": 81}
    detections = iter([first, shifted, shifted])
    visible = {}
    writes = []

    monkeypatch.delenv("REPORTGEN_FAST_TOC", raising=False)
    monkeypatch.delenv("REPORTGEN_SKIP_STATIC_TOC_PAGE_NUMBERS", raising=False)
    monkeypatch.setattr(shutil, "which", lambda name: f"/fake/{name}")
    monkeypatch.setattr(renderer, "_document_contains_toc_or_static_toc", lambda *_: True)
    monkeypatch.setattr(renderer, "_set_word_compat_pagination", lambda *_: None)
    monkeypatch.setattr(
        renderer,
        "_detect_toc_page_numbers_from_pdf_layout",
        lambda **_kwargs: next(detections),
    )

    def fake_write(_path, page_numbers, _context):
        writes.append(dict(page_numbers))
        visible.clear()
        visible.update(page_numbers)
        return True

    monkeypatch.setattr(renderer, "_write_static_toc_page_numbers", fake_write)
    monkeypatch.setattr(renderer, "_read_static_toc_page_numbers", lambda *_: dict(visible))

    renderer._populate_static_toc_page_numbers(str(docx_path), {})

    assert writes == [first, shifted]
    assert visible == shifted


def test_static_toc_page_numbers_fail_when_final_layout_never_converges(tmp_path, monkeypatch):
    docx_path = tmp_path / "report.docx"
    Document().save(docx_path)
    renderer = TemplateRenderer(log_level="ERROR")
    detections = iter(
        [
            {"参考文献": 67},
            {"参考文献": 68},
            {"参考文献": 69},
            {"参考文献": 70},
            {"参考文献": 71},
            {"参考文献": 72},
            {"参考文献": 73},
        ]
    )
    visible = {}

    monkeypatch.delenv("REPORTGEN_FAST_TOC", raising=False)
    monkeypatch.delenv("REPORTGEN_SKIP_STATIC_TOC_PAGE_NUMBERS", raising=False)
    monkeypatch.setattr(shutil, "which", lambda name: f"/fake/{name}")
    monkeypatch.setattr(renderer, "_document_contains_toc_or_static_toc", lambda *_: True)
    monkeypatch.setattr(renderer, "_set_word_compat_pagination", lambda *_: None)
    monkeypatch.setattr(
        renderer,
        "_detect_toc_page_numbers_from_pdf_layout",
        lambda **_kwargs: next(detections),
    )

    def fake_write(_path, page_numbers, _context):
        visible.clear()
        visible.update(page_numbers)
        return True

    monkeypatch.setattr(renderer, "_write_static_toc_page_numbers", fake_write)
    monkeypatch.setattr(renderer, "_read_static_toc_page_numbers", lambda *_: dict(visible))

    with pytest.raises(RuntimeError, match="目录页码.*未收敛"):
        renderer._populate_static_toc_page_numbers(str(docx_path), {})


def test_qa_report_records_post_processor_errors(tmp_path):
    docx_path = tmp_path / "clean.docx"
    doc = Document()
    doc.add_paragraph("已生成报告")
    doc.save(docx_path)

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        processor_report=[
            {"name": "bullet_lists", "status": "OK"},
            {"name": "toc_refresh", "status": "ERROR", "error": "refresh failed"},
        ],
    )

    assert qa["status"] == "WARN"
    assert qa["checks"]["post_processors"]["error_count"] == 1
    assert any(i["code"] == "POST_PROCESSOR_ERRORS" for i in qa["issues"])


def test_qa_report_fails_on_critical_part3_processor_error(tmp_path):
    docx_path = tmp_path / "clean.docx"
    doc = Document()
    doc.add_paragraph("已生成报告")
    doc.save(docx_path)

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        processor_report=[
            {
                "name": "part3_formatted_sections",
                "status": "ERROR",
                "error": "marker failed",
            }
        ],
    )

    assert qa["status"] == "FAIL"
    assert qa["checks"]["post_processors"]["status"] == "FAIL"
    assert qa["checks"]["post_processors"]["critical_error_count"] == 1
    assert any(issue["code"] == "CRITICAL_POST_PROCESSOR_ERROR" for issue in qa["issues"])


def test_qa_report_scopes_final_pagination_failure_to_crc_panels(tmp_path):
    docx_path = tmp_path / "clean.docx"
    doc = Document()
    doc.add_paragraph("已生成报告")
    doc.save(docx_path)
    processor_report = [
        {
            "name": "underlines_and_styles",
            "status": "ERROR",
            "error": "目录页码未收敛",
        }
    ]

    crc = build_docx_qa_report(
        output_file=str(docx_path),
        project_type="crc_358_msi",
        processor_report=processor_report,
    )
    draft = build_docx_qa_report(
        output_file=str(docx_path),
        project_type="endometrial_29",
        processor_report=processor_report,
    )

    assert crc["checks"]["post_processors"]["status"] == "FAIL"
    assert crc["checks"]["post_processors"]["critical_error_count"] == 1
    assert draft["checks"]["post_processors"]["status"] == "WARN"
    assert draft["checks"]["post_processors"]["critical_error_count"] == 0


def test_qa_report_detects_placeholder_residue(tmp_path):
    docx_path = tmp_path / "placeholder.docx"
    doc = Document()
    doc.add_paragraph("患者：{{ patient_name }}")
    doc.save(docx_path)

    qa = build_docx_qa_report(output_file=str(docx_path))

    assert qa["status"] == "FAIL"
    assert any(i["code"] == "UNRENDERED_PLACEHOLDER" for i in qa["issues"])


def test_qa_report_detects_empty_numbered_paragraph(tmp_path):
    docx_path = tmp_path / "empty_numbered.docx"
    doc = Document()
    p = doc.add_paragraph("")
    ppr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    ppr.append(num_pr)
    doc.save(docx_path)

    qa = build_docx_qa_report(output_file=str(docx_path))

    assert qa["status"] == "FAIL"
    assert qa["checks"]["empty_numbered_paragraphs"]["count"] == 1
    assert any(i["code"] == "EMPTY_NUMBERED_PARAGRAPH" for i in qa["issues"])


def test_qa_report_ignores_numbered_image_paragraph(tmp_path):
    from PIL import Image

    docx_path = tmp_path / "numbered_image.docx"
    image_path = tmp_path / "figure.png"
    Image.new("RGB", (80, 40), color=(0, 180, 180)).save(image_path)

    doc = Document()
    p = doc.add_paragraph("")
    ppr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(num_id)
    ppr.append(num_pr)
    p.add_run().add_picture(str(image_path), width=Inches(2.0))
    doc.save(docx_path)

    qa = build_docx_qa_report(output_file=str(docx_path))

    assert qa["checks"]["empty_numbered_paragraphs"]["status"] == "PASS"
    assert qa["checks"]["empty_numbered_paragraphs"]["count"] == 0
    assert not any(i["code"] == "EMPTY_NUMBERED_PARAGRAPH" for i in qa["issues"])


def test_renderer_clears_empty_numbered_paragraphs_inside_tables(tmp_path):
    docx_path = tmp_path / "table_empty_numbered.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    p = table.rows[0].cells[0].paragraphs[0]
    ppr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(num_id)
    ppr.append(num_pr)
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._remove_empty_numbered_paragraphs(str(docx_path))

    rendered = Document(docx_path)
    cleaned = rendered.tables[0].rows[0].cells[0].paragraphs[0]
    assert cleaned._p.pPr is None or cleaned._p.pPr.numPr is None
    qa = build_docx_qa_report(output_file=str(docx_path))
    assert qa["checks"]["empty_numbered_paragraphs"]["count"] == 0


def test_qa_report_checks_crc_tables_and_counts(tmp_path):
    docx_path = tmp_path / "crc.docx"
    doc = Document()
    doc.add_paragraph("本次共检出体细胞变异：8 个")
    doc.add_paragraph("与靶向药物用药相关的变异有：4 个")
    doc.add_paragraph("与靶向/免疫药物相关的变异：4 个")
    doc.add_paragraph("6.5 mutations/Mb，TMB-L；微卫星稳定型，MSS")

    tips = doc.add_table(rows=2, cols=4)
    for idx, text in enumerate(["基因", "突变位点", "潜在获益靶向药物", "可能耐药"]):
        tips.rows[0].cells[idx].text = text

    summary = doc.add_table(rows=2, cols=4)
    for idx, text in enumerate(["基因", "基因突变信息", "潜在获益靶向药物", "可能耐药"]):
        summary.rows[0].cells[idx].text = text

    detail = doc.add_table(rows=2, cols=9)
    headers = [
        "基因名称",
        "转录本号",
        "染色体",
        "外显子",
        "核苷酸变化",
        "氨基酸变化",
        "突变频率",
        "潜在获益",
        "可能耐药",
    ]
    for idx, text in enumerate(headers):
        detail.rows[0].cells[idx].text = text

    biomarker = doc.add_table(rows=2, cols=3)
    biomarker.rows[0].cells[0].text = "TMB/MSI/其它生物标志物检测结果"
    biomarker.rows[0].cells[2].text = "用药提示"
    biomarker.rows[1].cells[0].text = "MSI"
    doc.save(docx_path)

    report_data = ReportData()
    report_data.set_field("total_variants_count", 8)
    report_data.set_field("drug_related_count", 4)
    report_data.set_field("targeted_or_immune_related_count", 4)
    report_data.set_field("tmb_status", "TMB-L")
    report_data.set_field("msi_status", "MSS")
    report_data.set_table("variants", [{"gene": "TP53"}])

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        report_data=report_data,
        project_type="crc_358_msi",
    )

    assert qa["status"] == "PASS"
    assert qa["checks"]["variant_detail_table_shape"]["status"] == "PASS"
    assert qa["checks"]["total_variant_count_text"]["status"] == "PASS"
    assert qa["checks"]["drug_related_count_text"]["status"] == "PASS"
    assert qa["checks"]["targeted_or_immune_related_count_text"]["status"] == "PASS"


def _add_crc_style_qa_tables(doc: Document) -> None:
    summary = doc.add_table(rows=2, cols=4)
    for idx, text in enumerate(["基因", "突变位点", "潜在获益靶向药物", "可能耐药或慎重药物"]):
        summary.rows[0].cells[idx].text = text
    for idx, text in enumerate(["KRAS", "c.34G>A", "司美替尼（C）", "西妥昔单抗（A）"]):
        summary.rows[1].cells[idx].text = text

    detail = doc.add_table(rows=3, cols=9)
    row0 = [
        "基因名称",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "基因突变信息",
        "靶向药物信息",
        "靶向药物信息",
    ]
    row1 = [
        "基因名称",
        "转录本号",
        "染色体",
        "外显子",
        "位点",
        "突变类型",
        "频率",
        "潜在获益靶向药物",
        "可能耐药或慎重药物",
    ]
    row2 = [
        "KRAS",
        "NM_004985.5",
        "12",
        "2",
        "c.34G>A",
        "点突变",
        "46.29",
        "司美替尼（C）",
        "西妥昔单抗（A）",
    ]
    for row, values in zip(detail.rows, [row0, row1, row2]):
        for idx, value in enumerate(values):
            row.cells[idx].text = value

    biomarker = doc.add_table(rows=2, cols=3)
    biomarker.rows[0].cells[0].text = "TMB/MSI/其它生物标志物检测结果"
    biomarker.rows[0].cells[1].text = "TMB/MSI/其它生物标志物检测结果"
    biomarker.rows[0].cells[2].text = "用药提示"
    biomarker.rows[1].cells[0].text = "MSI"
    biomarker.rows[1].cells[1].text = "MSS"
    biomarker.rows[1].cells[2].text = "研究表明，MSI-H的实体瘤通常具有免疫原性"


def _crc_panel_style() -> dict:
    return ReportGenerator._load_panel_style_config(
        load_panel_package("crc_358_msi", project_root=ROOT)
    )


def test_qa_report_checks_crc_style_rules_pass_after_postprocessing(tmp_path):
    docx_path = tmp_path / "crc_style_pass.docx"
    doc = Document()
    doc.add_paragraph("本次共检出体细胞变异：1 个")
    doc.add_paragraph("与靶向药物用药相关的变异有：1 个")
    doc.add_paragraph("TMB-L；MSS")
    _add_crc_style_qa_tables(doc)
    doc.save(docx_path)

    context = {"panel_style": _crc_panel_style()}
    renderer = TemplateRenderer(log_level="ERROR")
    renderer._restore_variant_summary_table_style(str(docx_path), context)
    renderer._restore_variant_detail_table_style(str(docx_path), context)
    renderer._restore_biomarker_table_style(str(docx_path), context)

    report_data = ReportData(context=context)
    report_data.set_field("total_variants_count", 1)
    report_data.set_field("drug_related_count", 1)
    report_data.set_field("tmb_status", "TMB-L")
    report_data.set_field("msi_status", "MSS")
    report_data.set_table("variants", [{"gene": "KRAS"}])

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        report_data=report_data,
        project_type="crc_358_msi",
    )

    assert qa["status"] == "PASS"
    assert qa["checks"]["docx_style_rules"]["status"] == "PASS"
    assert qa["checks"]["docx_style_rules"]["checked_table_count"] == 3


def test_qa_report_treats_variant_detail_plain_marker_substrings_as_plain(tmp_path):
    docx_path = tmp_path / "crc_style_plain_marker.docx"
    doc = Document()
    doc.add_paragraph("本次共检出体细胞变异：1 个")
    doc.add_paragraph("与靶向药物用药相关的变异有：1 个")
    doc.add_paragraph("TMB-L；MSS")
    _add_crc_style_qa_tables(doc)
    detail = doc.tables[1]
    detail.rows[2].cells[4].text = "未见突变（质控通过）"
    doc.save(docx_path)

    context = {"panel_style": _crc_panel_style()}
    renderer = TemplateRenderer(log_level="ERROR")
    renderer._restore_variant_summary_table_style(str(docx_path), context)
    renderer._restore_variant_detail_table_style(str(docx_path), context)
    renderer._restore_biomarker_table_style(str(docx_path), context)

    report_data = ReportData(context=context)
    report_data.set_field("total_variants_count", 1)
    report_data.set_field("drug_related_count", 1)
    report_data.set_field("tmb_status", "TMB-L")
    report_data.set_field("msi_status", "MSS")
    report_data.set_table("variants", [{"gene": "KRAS"}])

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        report_data=report_data,
        project_type="crc_358_msi",
    )

    assert qa["status"] == "PASS"
    assert qa["checks"]["docx_style_rules"]["status"] == "PASS"


def test_qa_report_flags_crc_style_rule_violations(tmp_path):
    docx_path = tmp_path / "crc_style_fail.docx"
    doc = Document()
    doc.add_paragraph("本次共检出体细胞变异：1 个")
    doc.add_paragraph("与靶向药物用药相关的变异有：1 个")
    doc.add_paragraph("TMB-L；MSS")
    _add_crc_style_qa_tables(doc)
    run = doc.tables[0].rows[1].cells[0].paragraphs[0].runs[0]
    run.font.underline = True
    doc.save(docx_path)

    report_data = ReportData(context={"panel_style": _crc_panel_style()})
    report_data.set_field("total_variants_count", 1)
    report_data.set_field("drug_related_count", 1)
    report_data.set_field("tmb_status", "TMB-L")
    report_data.set_field("msi_status", "MSS")
    report_data.set_table("variants", [{"gene": "KRAS"}])

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        report_data=report_data,
        project_type="crc_358_msi",
    )

    assert qa["status"] == "FAIL"
    assert qa["checks"]["docx_style_rules"]["status"] == "FAIL"
    assert any(i["code"] == "DOCX_STYLE_RULES" for i in qa["issues"])


def test_qa_report_flags_crc_missing_required_tables(tmp_path):
    docx_path = tmp_path / "crc_missing_tables.docx"
    doc = Document()
    doc.add_paragraph("本次共检出体细胞变异：1 个")
    doc.add_paragraph("与靶向药物用药相关的变异有：1 个")
    doc.add_paragraph("TMB-L；MSS")
    doc.save(docx_path)

    report_data = ReportData()
    report_data.set_field("total_variants_count", 1)
    report_data.set_field("drug_related_count", 1)
    report_data.set_field("tmb_status", "TMB-L")
    report_data.set_field("msi_status", "MSS")
    report_data.set_table("variants", [{"gene": "TP53"}])

    qa = build_docx_qa_report(
        output_file=str(docx_path),
        report_data=report_data,
        project_type="crc_358_msi",
    )

    assert qa["status"] == "FAIL"
    issue_codes = {i["code"] for i in qa["issues"]}
    assert "VARIANT_DETAIL_TABLE_SHAPE" in issue_codes
    assert "BIOMARKER_TABLE_PRESENT" in issue_codes


def test_write_qa_report_creates_sidecar_json(tmp_path):
    docx_path = tmp_path / "clean.docx"
    doc = Document()
    doc.add_paragraph("已生成报告")
    doc.save(docx_path)

    qa = build_docx_qa_report(output_file=str(docx_path))
    qa_path = Path(write_docx_qa_report(qa, str(docx_path)))

    assert qa_path.name == "clean.qa.json"
    assert qa_path.exists()
    assert "PASS" in qa_path.read_text(encoding="utf-8")


def test_crc358_golden_excel_fixture_is_synthetic_and_parseable(tmp_path):
    xlsx_path = build_crc_358_msi_golden_excel(tmp_path / "LZ999001_crc_358_msi_golden.xlsx")

    excel_data = ExcelReader(config_dir=str(ROOT / "config"), log_level="ERROR").read(
        str(xlsx_path), include_tables=True
    )

    assert excel_data.metadata["sample_id_from_filename"] == "LZ999001"
    assert excel_data.single_values["患者姓名"] == "黄金测试患者"
    assert excel_data.single_values["TMB"] == 6.5
    assert excel_data.single_values["MSI状态"] == "MSS"
    assert len(excel_data.get_table_data("Variations")) == 2
    assert {r["Gene_Symbol"] for r in excel_data.get_table_data("Variations")} == {
        "ERBB2",
        "FBXW7",
    }


def test_crc358_missing_required_msi_blocks_generation_even_without_strict_mode(tmp_path):
    from openpyxl import load_workbook

    xlsx_path = build_crc_358_msi_golden_excel(tmp_path / "LZ999002_missing_msi.xlsx")
    workbook = load_workbook(xlsx_path)
    del workbook["Msisensor"]
    workbook.save(xlsx_path)
    package = load_panel_package("crc_358_msi", project_root=ROOT)

    result = ReportGenerator(config_dir=str(ROOT / "config"), log_level="ERROR").generate(
        excel_file=str(xlsx_path),
        template_file=str(package.resolve_template_file()),
        output_dir=str(tmp_path / "out"),
        strict_mode=False,
        project_type="crc_358_msi",
    )

    assert result["success"] is False
    assert any("必检生物标志物" in error for error in result["errors"])
    assert not list((tmp_path / "out").glob("*.docx"))


def test_crc301_panel_package_basic_generation_passes(tmp_path):
    xlsx_path = build_crc_301_msi_golden_excel(tmp_path / "LZ999301_crc_301_msi_basic.xlsx")

    result = ReportGenerator(
        config_dir=str(ROOT / "config"),
        log_level="ERROR",
    ).generate(
        excel_file=str(xlsx_path),
        template_file=str(ROOT / "templates" / "aligned_template_with_cnv_fusion_hla_FIXED.docx"),
        output_dir=str(tmp_path / "out"),
        output_filename="crc301_basic.docx",
        strict_mode=True,
        return_context=True,
        template_contract_mode="fail",
        project_type="crc_301",
        project_name="结直肠癌301基因+MSI",
    )

    assert result["success"], result.get("errors")
    assert result["qa_status"] == "PASS"
    assert result["panel_package_validation"]["status"] == "PASS"
    assert result["rule_provenance"]["status"] == "PASS"
    assert result["rule_provenance"]["file_count"] >= 6
    assert result["generation_id"] == Path(result["output_file"]).stem
    stage_results_file = Path(result["stage_results_file"])
    assert stage_results_file.exists()
    stage_payload = json.loads(stage_results_file.read_text(encoding="utf-8"))
    assert stage_payload["generation_id"] == result["generation_id"]
    assert stage_payload["pipeline"]["status"] == "PASS"
    assert stage_payload["stage_results"] == result["stage_results"]
    assert result["qa_report"]["pipeline"]["status"] == "PASS"
    assert result["qa_report"]["rules"]["status"] == "PASS"
    assert result["qa_report"]["checks"]["docx_style_rules"]["status"] == "PASS"
    assert result["qa_report"]["checks"]["rules"]["file_count"] >= 6
    assert result["qa_report"]["checks"]["pipeline"]["stage_results_file"] == str(
        stage_results_file
    )
    stage_names = [stage["name"] for stage in result["stage_results"]]
    assert stage_names[:3] == [
        "PanelResolutionStage",
        "PanelPackageValidationStage",
        "ExcelReadStage",
    ]
    assert "TemplateRenderStage" in stage_names
    assert "QAStage" in stage_names
    assert all(stage["duration_ms"] is not None for stage in result["stage_results"])
    assert Path(result["output_file"]).exists()
    assert result["context"]["project_name"] == "结直肠癌301基因+MSI"
    assert result["context"]["panel_style"]["variant_summary_table"]["link_underline"] is False
    assert result["context"]["panel_style"]["biomarker_table"]["header_fill"] == ("00B7C7")
    chemotherapy = result["context"]["chemotherapy"]
    assert len(chemotherapy) == 7
    assert "瑞戈非尼" in chemotherapy[0]["Drug"]
    assert "结直肠癌" in chemotherapy[0]["药物适应情况"]
    assert result["context"]["report_text_rule_keys"]["immuno_tips"] == ("tmb_table_immuno_tips")
    assert "TMB-H的肿瘤" in result["context"]["immuno_tips"]
    assert "2020年6月，FDA批准帕博利珠单抗" in result["context"]["tmb_detail_interpretation"]
    assert "MSI-H的实体瘤通常具有免疫原性" in result["context"]["msi_tips"]
    assert result["template_contract"]["declared_contract"]["ok"] is True


def test_crc301_golden_case_passes(tmp_path):
    result = run_golden_case(
        GoldenCaseOptions(
            panel="crc_301_msi",
            config_dir=str(ROOT / "config"),
            output_root=str(tmp_path / "golden"),
            template_contract_mode="fail",
        )
    )

    assert result["ok"], result.get("errors")
    assert result["qa_status"] == "PASS"
    assert result["panel"] == "crc_301_msi"
    assert Path(result["input_excel"]).name == "LZ999301_crc_301_msi_golden.xlsx"
    assert Path(result["output_file"]).name == "golden_crc_301_msi.docx"
    assert not [row for row in result["checks"] if not row["passed"]]


def test_lung_methylation_golden_excel_fixture_is_parseable(tmp_path):
    xlsx_path = build_lung_methylation_golden_excel(
        tmp_path / "LUNG999001_lung_methylation_golden.xlsx"
    )

    excel_data = ExcelReader(config_dir=str(ROOT / "config"), log_level="ERROR").read(
        str(xlsx_path), include_tables=True
    )

    assert excel_data.metadata["sample_id_from_filename"] == "LUNG999001"
    assert excel_data.single_values["患者姓名"] == "黄金甲基化患者"
    assert excel_data.single_values["甲基化结果"] == "阳性"
    rows = excel_data.get_table_data("甲基化位点")
    assert len(rows) == 2
    assert {row["基因"] for row in rows} == {"SHOX2", "RASSF1A"}


def test_lung_methylation_golden_case_passes(tmp_path):
    result = run_golden_case(
        panel="lung_methylation",
        config_dir=str(ROOT / "config"),
        output_root=str(tmp_path / "lung_methylation_golden"),
        log_level="ERROR",
        template_contract_mode="fail",
    )

    assert result["ok"], result.get("errors")
    assert result["qa_status"] == "PASS"
    assert Path(result["output_file"]).exists()
    assert result["panel"] == LUNG_METHYLATION_EXPECTATIONS["project_type"]


@pytest.mark.parametrize(
    ("builder", "filename", "panel_id", "project_name"),
    [
        (
            build_lung_329_pdl1_golden_input,
            "SYN-L329-GOLDEN.xlsx",
            "lung_329_pdl1",
            "肺癌329基因+PD-L1",
        ),
        (
            build_lung_588_pdl1_golden_input,
            "SYN-L588-GOLDEN.xlsx",
            "lung_588_pdl1",
            "肺癌588基因+PD-L1",
        ),
    ],
)
def test_lung_pdl1_golden_input_has_explicit_synthetic_provenance(
    tmp_path, builder, filename, panel_id, project_name
):
    fixture = builder(tmp_path / filename)

    assert fixture.excel_file.is_file()
    assert ZipFile(fixture.excel_file).testzip() is None
    assert fixture.excel_data is not None
    assert fixture.excel_data.metadata["synthetic_fixture"] is True
    assert fixture.excel_data.metadata["panel_id"] == panel_id
    assert fixture.excel_data.single_values["检测项目"] == project_name
    assert fixture.excel_data.single_values["PD-L1 TPS"] == 50
    assert fixture.excel_data.single_values["PD-L1 CPS"] == 52
    assert fixture.excel_data.single_values["PD-L1原始记录编号"].startswith(
        "SYNTHETIC-IHC-"
    )
    assert Path(fixture.excel_data.single_values["PD-L1病例图片"]).is_file()
    assert (
        fixture.excel_data.metadata["field_source_overrides"]["pdl1_tps"]["source"]
        == "form"
    )


def test_current_output_snapshot_extracts_pdl1_integer_display(tmp_path):
    output = tmp_path / "synthetic_lung_pdl1.docx"
    document = Document()
    document.add_paragraph(
        "本次PD-L1免疫组化检测结果：TPS 50%，CPS 52，阳性（高表达）。"
    )
    document.save(output)

    snapshot = snapshot_docx_report(output, panel="lung_588_pdl1")

    assert snapshot["features"]["pdl1_tps"] == 50.0
    assert snapshot["features"]["pdl1_cps"] == 52.0
    assert snapshot["features"]["pdl1_result"] == "阳性（高表达）"


def test_golden_case_assertions_pass_on_expected_report_shape(tmp_path):
    docx_path = tmp_path / "golden_like.docx"
    doc = Document()
    doc.add_paragraph("本次共检出体细胞变异：2 个")
    doc.add_paragraph("与靶向药物用药相关的变异有：1 个")
    doc.add_paragraph("与靶向/免疫药物相关的变异：1 个")
    doc.add_paragraph("6.5 mutations/Mb，TMB-L；微卫星稳定型，MSS")
    doc.add_paragraph("多项临床研究表明，TMB-H的肿瘤对免疫检查点抑制剂有更强的免疫应答效果")
    doc.add_paragraph("研究表明，MSI-H的实体瘤通常具有免疫原性和广泛的T细胞浸润性")
    doc.add_paragraph("ERBB2：c.1979G>A，p.G660D")

    tips = doc.add_table(rows=2, cols=4)
    for idx, text in enumerate(["基因", "突变位点", "潜在获益靶向药物", "可能耐药"]):
        tips.rows[0].cells[idx].text = text
    tips.rows[1].cells[0].text = "ERBB2"
    tips.rows[1].cells[1].text = "c.1979G>A，p.G660D"

    summary = doc.add_table(rows=2, cols=4)
    for idx, text in enumerate(["基因", "基因突变信息", "潜在获益靶向药物", "可能耐药"]):
        summary.rows[0].cells[idx].text = text
    summary.rows[1].cells[0].text = "ERBB2"

    detail = doc.add_table(rows=2, cols=9)
    headers = [
        "基因名称",
        "转录本号",
        "染色体",
        "外显子",
        "核苷酸变化",
        "氨基酸变化",
        "突变频率",
        "潜在获益",
        "可能耐药",
    ]
    for idx, text in enumerate(headers):
        detail.rows[0].cells[idx].text = text
    detail.rows[1].cells[0].text = "ERBB2"
    detail.rows[1].cells[4].text = "c.1979G>A"
    detail.rows[1].cells[5].text = "p.G660D"

    biomarker = doc.add_table(rows=2, cols=3)
    biomarker.rows[0].cells[0].text = "TMB/MSI/其它生物标志物检测结果"
    biomarker.rows[0].cells[2].text = "用药提示"
    biomarker.rows[1].cells[0].text = "MSI"
    doc.save(docx_path)

    report_data = ReportData()
    report_data.set_field("total_variants_count", 2)
    report_data.set_field("drug_related_count", 1)
    report_data.set_field("targeted_or_immune_related_count", 1)
    report_data.set_field("tmb_status", "L")
    report_data.set_field("msi_status", "MSS")
    report_data.set_table("variants", [{"gene": "ERBB2"}])
    processor_report = [{"name": "underlines_and_styles", "status": "OK"}]
    qa = build_docx_qa_report(
        output_file=str(docx_path),
        report_data=report_data,
        project_type="crc_358_msi",
        field_provenance={"fields": {"patient_name": {"source": "excel"}}},
        field_provenance_file=str(docx_path.with_suffix(".field_provenance.json")),
        processor_report=processor_report,
    )

    result = {
        "success": True,
        "output_file": str(docx_path),
        "context": report_data.context,
        "qa_report": qa,
        "field_provenance_file": str(docx_path.with_suffix(".field_provenance.json")),
        "post_processors": processor_report,
        "errors": [],
    }
    assertion = assert_golden_case_output(result, expectations=CRC_358_MSI_EXPECTATIONS)

    assert assertion["ok"]


def test_docx_render_uses_configured_tmp_dir_and_timeout(tmp_path, monkeypatch):
    docx_path = tmp_path / "source.docx"
    Document().save(docx_path)
    render_tmp = tmp_path / "render_tmp"
    output_dir = tmp_path / "pages"
    calls = []

    monkeypatch.setattr(
        docx_render,
        "_which_or_raise",
        lambda name, *, hint: f"/bin/{name}",
    )

    def fake_run(cmd, *, timeout_seconds, stage):
        calls.append((stage, list(cmd), timeout_seconds))
        if stage == "docx_to_pdf":
            outdir = Path(cmd[cmd.index("--outdir") + 1])
            assert str(outdir).startswith(str(render_tmp))
            assert Path(cmd[-1]).name == "input.docx"
            (outdir / "input.pdf").write_bytes(b"%PDF-1.4\n")
        elif stage == "pdf_to_png":
            prefix = Path(cmd[-1])
            prefix.parent.mkdir(parents=True, exist_ok=True)
            Path(f"{prefix}-1.png").write_bytes(b"png")
        return subprocess.CompletedProcess(cmd, 0, "", "")

    monkeypatch.setattr(docx_render, "_run_checked", fake_run)
    monkeypatch.setenv("REPORTGEN_RENDER_TMPDIR", str(render_tmp))
    monkeypatch.setenv("REPORTGEN_LIBREOFFICE_PROFILE_MODE", "isolated")

    pngs = docx_render.render_docx_to_pngs(
        docx_path,
        output_dir=output_dir,
        timeout_seconds=37,
    )

    assert [call[0] for call in calls] == ["docx_to_pdf", "pdf_to_png"]
    assert {call[2] for call in calls} == {37}
    assert pngs == [output_dir / "source-1.png"]


def test_docx_render_does_not_hide_isolated_profile_failure(tmp_path, monkeypatch):
    docx_path = tmp_path / "source.docx"
    Document().save(docx_path)
    output_dir = tmp_path / "pages"
    calls = []

    monkeypatch.setattr(
        docx_render,
        "_which_or_raise",
        lambda name, *, hint: f"/bin/{name}",
    )

    def fake_run(cmd, *, timeout_seconds, stage):
        calls.append((stage, list(cmd)))
        if stage == "docx_to_pdf":
            raise docx_render.DocxRenderError(
                "isolated profile failed",
                stage=stage,
                command=cmd,
                stderr="profile crash",
            )
        return subprocess.CompletedProcess(cmd, 0, "", "")

    monkeypatch.setattr(docx_render, "_run_checked", fake_run)
    monkeypatch.setenv("REPORTGEN_LIBREOFFICE_PROFILE_MODE", "isolated")

    with pytest.raises(docx_render.DocxRenderError, match="isolated profile failed"):
        docx_render.render_docx_to_pngs(docx_path, output_dir=output_dir)

    assert [call[0] for call in calls] == ["docx_to_pdf"]


def test_docx_render_uses_isolated_profile_by_default_on_macos(tmp_path, monkeypatch):
    docx_path = tmp_path / "source.docx"
    Document().save(docx_path)
    output_dir = tmp_path / "pages"
    calls = []
    profile_initializations = []

    monkeypatch.delenv("REPORTGEN_LIBREOFFICE_PROFILE_MODE", raising=False)
    monkeypatch.delenv("REPORTGEN_RENDER_PROFILE_MODE", raising=False)
    monkeypatch.setattr(docx_render.platform, "system", lambda: "Darwin")
    monkeypatch.setattr(
        docx_render,
        "_which_or_raise",
        lambda name, *, hint: f"/bin/{name}",
    )
    monkeypatch.setattr(
        docx_render,
        "initialize_libreoffice_profile",
        lambda profile_dir, *, require_available: profile_initializations.append(
            (Path(profile_dir), require_available)
        ),
    )

    def fake_run(cmd, *, timeout_seconds, stage):
        calls.append((stage, list(cmd)))
        if stage == "docx_to_pdf":
            assert any("UserInstallation" in part for part in cmd)
            outdir = Path(cmd[cmd.index("--outdir") + 1])
            (outdir / "input.pdf").write_bytes(b"%PDF-1.4\n")
        elif stage == "pdf_to_png":
            prefix = Path(cmd[-1])
            prefix.parent.mkdir(parents=True, exist_ok=True)
            Path(f"{prefix}-1.png").write_bytes(b"png")
        return subprocess.CompletedProcess(cmd, 0, "", "")

    monkeypatch.setattr(docx_render, "_run_checked", fake_run)

    pngs = docx_render.render_docx_to_pngs(docx_path, output_dir=output_dir)

    assert [call[0] for call in calls] == ["docx_to_pdf", "pdf_to_png"]
    assert len(profile_initializations) == 1
    assert profile_initializations[0][1] is True
    assert pngs == [output_dir / "source-1.png"]


def test_docx_render_timeout_reports_stage():
    with pytest.raises(docx_render.DocxRenderError) as exc_info:
        docx_render._run_checked(
            [sys.executable, "-c", "import time; time.sleep(2)"],
            timeout_seconds=1,
            stage="docx_to_pdf",
        )

    assert exc_info.value.stage == "docx_to_pdf"
    assert "timed out" in str(exc_info.value)


def test_visual_render_optional_failure_is_reported_not_blocking(tmp_path, monkeypatch):
    docx_path = tmp_path / "report.docx"
    Document().save(docx_path)

    def fail_render(*args, **kwargs):
        raise RuntimeError("LibreOffice failed")

    monkeypatch.setattr("reportgen.core.golden_case.render_docx_to_pngs", fail_render)

    result = run_visual_render(
        str(docx_path),
        output_root=tmp_path,
        mode="first",
        required=False,
    )

    assert result["status"] == "WARN"
    assert "LibreOffice failed" in result["error"]


def test_visual_render_required_failure_blocks_golden_check(tmp_path, monkeypatch):
    docx_path = tmp_path / "report.docx"
    Document().save(docx_path)

    monkeypatch.setattr(
        "reportgen.core.golden_case.render_docx_to_pngs",
        lambda *args, **kwargs: (_ for _ in ()).throw(RuntimeError("render failed")),
    )

    result = run_visual_render(
        str(docx_path),
        output_root=tmp_path,
        mode="all",
        required=True,
    )

    assert result["status"] == "FAIL"
    assert result["requested"] == "all"


def test_golden_case_options_accept_visual_render_controls():
    opts = GoldenCaseOptions(
        panel="crc_358_msi",
        render="first",
        render_dpi=140,
        render_timeout_seconds=45,
        render_required=True,
        render_tmp_dir="/tmp/reportgen-render",
    )

    assert opts.render == "first"
    assert opts.render_dpi == 140
    assert opts.render_timeout_seconds == 45
    assert opts.render_required is True


def test_report_diff_passes_for_identical_docx(tmp_path):
    reference = tmp_path / "reference.docx"
    candidate = tmp_path / "candidate.docx"
    doc = Document()
    doc.add_paragraph("本次共检出体细胞变异：2 个")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "基因"
    table.rows[0].cells[1].text = "位点"
    table.rows[1].cells[0].text = "KRAS"
    table.rows[1].cells[1].text = "c.34G>A"
    doc.save(reference)
    doc.save(candidate)

    result = compare_reports(
        ReportDiffOptions(
            reference_docx=str(reference),
            candidate_docx=str(candidate),
        )
    )

    assert result["status"] == "PASS"
    assert result["summary"]["failures"] == 0
    assert result["sections"]["text"]["similarity"] == 1.0


def test_report_diff_fails_on_table_shape_change_and_writes_outputs(tmp_path):
    reference = tmp_path / "reference.docx"
    candidate = tmp_path / "candidate.docx"
    ref_doc = Document()
    ref_doc.add_paragraph("报告摘要")
    ref_table = ref_doc.add_table(rows=2, cols=2)
    ref_table.rows[0].cells[0].text = "基因"
    ref_table.rows[0].cells[1].text = "位点"
    ref_table.rows[1].cells[0].text = "KRAS"
    ref_table.rows[1].cells[1].text = "c.34G>A"
    ref_doc.save(reference)

    cand_doc = Document()
    cand_doc.add_paragraph("报告摘要")
    cand_table = cand_doc.add_table(rows=2, cols=3)
    cand_table.rows[0].cells[0].text = "基因"
    cand_table.rows[0].cells[1].text = "位点"
    cand_table.rows[0].cells[2].text = "药物"
    cand_table.rows[1].cells[0].text = "KRAS"
    cand_table.rows[1].cells[1].text = "c.34G>A"
    cand_table.rows[1].cells[2].text = "西妥昔单抗"
    cand_doc.save(candidate)

    result = compare_reports(
        ReportDiffOptions(
            reference_docx=str(reference),
            candidate_docx=str(candidate),
            output_dir=str(tmp_path / "diff"),
        )
    )

    assert result["status"] == "FAIL"
    assert any(i["code"] == "TABLE_SHAPE_DIFF" for i in result["issues"])
    assert (tmp_path / "diff" / "report_diff.json").exists()
    assert "TABLE_SHAPE_DIFF" in (tmp_path / "diff" / "report_diff.md").read_text(encoding="utf-8")


def test_report_diff_warns_on_style_and_qa_changes(tmp_path):
    reference = tmp_path / "reference.docx"
    candidate = tmp_path / "candidate.docx"
    ref_doc = Document()
    ref_doc.add_paragraph("药物提示")
    ref_doc.save(reference)

    cand_doc = Document()
    para = cand_doc.add_paragraph()
    run = para.add_run("药物提示")
    run.font.underline = True
    cand_doc.save(candidate)

    reference.with_suffix(".qa.json").write_text(
        json.dumps({"status": "PASS", "issues": []}, ensure_ascii=False),
        encoding="utf-8",
    )
    candidate.with_suffix(".qa.json").write_text(
        json.dumps(
            {
                "status": "WARN",
                "issues": [
                    {
                        "level": "warning",
                        "code": "TOC_PAGE_NUMBERS_MISSING",
                        "message": "目录页码缺失",
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    result = compare_reports(
        ReportDiffOptions(
            reference_docx=str(reference),
            candidate_docx=str(candidate),
        )
    )

    assert result["status"] == "WARN"
    assert any(i["code"] == "STYLE_DIFF" for i in result["issues"])
    assert any(i["code"] == "QA_STATUS_DIFF" for i in result["issues"])
    assert result["sections"]["qa"]["candidate_status"] == "WARN"


def test_report_diff_can_suppress_golden_template_noise(tmp_path):
    reference = tmp_path / "reference.docx"
    candidate = tmp_path / "candidate.docx"
    ref_doc = Document()
    ref_doc.add_paragraph("33333333333333333333333333")
    ref_doc.add_paragraph("报告摘要 ")
    ref_table = ref_doc.add_table(rows=1, cols=1)
    ref_table.rows[0].cells[0].text = "司美替尼（C） \n曲美替尼（C） "
    ref_doc.save(reference)

    cand_doc = Document()
    para = cand_doc.add_paragraph()
    run = para.add_run("报告摘要")
    run.font.underline = True
    cand_table = cand_doc.add_table(rows=1, cols=1)
    cand_table.rows[0].cells[0].text = "司美替尼（C）\n曲美替尼（C）"
    cand_doc.save(candidate)

    result = compare_reports(
        ReportDiffOptions(
            reference_docx=str(reference),
            candidate_docx=str(candidate),
            normalize_whitespace=True,
            ignore_reference_artifacts=True,
            style_metric_policy="summary",
        )
    )

    assert result["status"] == "PASS"
    assert result["sections"]["text"]["similarity"] == 1.0
    assert result["sections"]["tables"]["status"] == "PASS"
    assert result["sections"]["styles"]["policy"] == "summary"
    assert result["sections"]["styles"]["samples"]


def test_gene_knowledge_provider_applies_reviewed_part3_yaml_overlay(tmp_path):
    overlay = tmp_path / "reviewed_part3.yaml"
    overlay.write_text(
        yaml.safe_dump(
            {
                "gene_sections": [
                    {
                        "gene": "SETD2",
                        "c_hgvs": "c.4930G>T",
                        "p_hgvs": "p.G1644*",
                        "intro": "SETD2 reviewed intro",
                        "mutation_analysis": "SETD2 reviewed analysis",
                    }
                ],
                "drug_sections": [
                    {
                        "gene": "SETD2",
                        "c_hgvs": "c.4930G>T",
                        "p_hgvs": "p.G1644*",
                        "type": "benefit",
                        "header": "SETD2：c.4930G>T，p.G1644*突变相应靶向药物",
                        "drug_name": "AZD1775",
                        "relation": "reviewed relation",
                        "clinical": "reviewed clinical",
                    }
                ],
            },
            allow_unicode=True,
            sort_keys=False,
        ),
        encoding="utf-8",
    )
    provider = GeneKnowledgeProvider(
        {
            "enabled": True,
            "gene_knowledge_db": {
                "enabled": True,
                "path": "missing.xlsx",
                "reviewed_part3_overlay_path": str(overlay),
            },
        }
    )
    assert provider.load(base_path=str(tmp_path))

    section = provider.build_gene_knowledge_section(
        gene="SETD2",
        c_hgvs="c.4930G>T",
        p_hgvs="p.G1644*",
        frequency=22.15,
        mutation_type="Nonsense",
        has_drug=True,
    )
    assert section["intro"] == "SETD2 reviewed intro"
    assert "22.15%" in section["mutation_desc"]
    assert section["mutation_analysis"] == "SETD2 reviewed analysis"

    drug_sections = provider.build_drug_analysis_sections(
        [
            {
                "gene": "SETD2",
                "cHGVS": "c.4930G>T",
                "pHGVS": "p.G1644*",
                "benefit_drugs": "AZD1775（C）",
                "caution_drugs": "--",
            }
        ]
    )
    assert len(drug_sections) == 1
    assert drug_sections[0]["drug_name"] == "AZD1775"
    assert drug_sections[0]["relation"] == "reviewed relation"
    assert drug_sections[0]["clinical"] == "reviewed clinical"


def test_report_diff_fails_on_part3_gene_and_drug_text_mismatch(tmp_path):
    reference = tmp_path / "reference.docx"
    candidate = tmp_path / "candidate.docx"

    ref_doc = Document()
    for text in [
        "第三部分：基因变异及相应靶向/免疫药物解析",
        "基因变异解析",
        "u SETD2：c.4930G>T，p.G1644*；22.15%",
        "基因简介：",
        "SETD2 reviewed intro",
        "基因变异说明：",
        "SETD2 reviewed desc",
        "基因变异解析：",
        "SETD2 reviewed analysis",
        "靶向药物/免疫用药提示解析",
        "潜在获益靶向/免疫药物解析",
        "SETD2：c.4930G>T，p.G1644*突变相应靶向药物",
        "AZD1775",
        "基因变异与药物关联分析：",
        "reviewed relation",
        "药物疗效临床解析：",
        "reviewed clinical",
        "3. 阅读说明",
    ]:
        ref_doc.add_paragraph(text)
    ref_doc.save(reference)

    cand_doc = Document()
    for text in [
        "第三部分：基因变异及相应靶向/免疫药物解析",
        "基因变异解析",
        "u SETD2：c.4930G>T，p.G1644*；22.15%",
        "基因简介：",
        "SETD2 generic intro",
        "基因变异说明：",
        "SETD2 reviewed desc",
        "基因变异解析：",
        "SETD2 reviewed analysis",
        "靶向药物/免疫用药提示解析",
        "潜在获益靶向/免疫药物解析",
        "SETD2：c.4930G>T，p.G1644*突变相应靶向药物",
        "AZD1775",
        "基因变异与药物关联分析：",
        "generic relation",
        "药物疗效临床解析：",
        "reviewed clinical",
        "3. 阅读说明",
    ]:
        cand_doc.add_paragraph(text)
    cand_doc.save(candidate)

    result = compare_reports(
        ReportDiffOptions(
            reference_docx=str(reference),
            candidate_docx=str(candidate),
        )
    )

    assert result["status"] == "FAIL"
    codes = {issue["code"] for issue in result["issues"]}
    assert "PART3_GENE_SECTION_DIFF" in codes
    assert "PART3_DRUG_SECTION_DIFF" in codes


def test_quality_gate_can_run_panel_validation_only(tmp_path):
    from reportgen.core.qa_gate import QualityGateOptions, run_quality_gate

    result = run_quality_gate(
        QualityGateOptions(
            project_root=str(ROOT),
            output_root=str(tmp_path / "gate"),
            run_lint=False,
            run_pytest=False,
            run_golden=False,
        )
    )

    assert result["status"] == "PASS"
    assert result["summary"]["failed"] == 0
    assert result["steps"][0]["name"] == "panel_validate"
    assert result["steps"][1]["name"] == "knowledge_release_gate"
    assert result["steps"][1]["status"] == "PASS"
    assert Path(result["report_file"]).exists()


def test_quality_gate_records_the_repository_pinned_ruff_toolchain():
    from reportgen.core.qa_gate import _run_qa_toolchain_check

    result = _run_qa_toolchain_check(ROOT)

    assert result["status"] == "PASS"
    assert result["required_ruff_version"] == "0.15.8"
    assert result["installed_ruff_version"] == "0.15.8"
    assert result["python_executable"] == sys.executable


def test_quality_gate_reports_missing_ruff_as_an_environment_contract_error(tmp_path, monkeypatch):
    from reportgen.core import qa_gate

    (tmp_path / "requirements-qa.txt").write_text("ruff==0.15.8\n", encoding="utf-8")

    def missing_distribution(_name):
        raise qa_gate.importlib_metadata.PackageNotFoundError("ruff")

    monkeypatch.setattr(qa_gate.importlib_metadata, "version", missing_distribution)

    result = qa_gate._run_qa_toolchain_check(tmp_path)

    assert result["status"] == "FAIL"
    assert result["installed_ruff_version"] is None
    assert result["issues"][0]["code"] == "QA_ENV_MISSING_DEPENDENCY"


def test_quality_gate_fails_when_panel_validation_has_warnings(tmp_path, monkeypatch):
    from reportgen.core import qa_gate

    class FakeReport:
        errors = []
        warnings = [{"code": "WARN"}]

        def to_dict(self):
            return {
                "summary": {"errors": 0, "warnings": 1},
                "panels_checked": ["demo_panel"],
                "issues": [{"level": "warning", "code": "WARN", "message": "demo"}],
            }

    monkeypatch.setattr(qa_gate, "validate_panel_registry", lambda project_root: FakeReport())

    result = qa_gate.run_quality_gate(
        qa_gate.QualityGateOptions(
            project_root=str(ROOT),
            output_root=str(tmp_path / "gate_warn"),
            run_lint=False,
            run_pytest=False,
            run_golden=False,
            fail_on_warn=True,
        )
    )

    assert result["status"] == "FAIL"
    assert result["steps"][0]["status"] == "FAIL"


def _add_red_bullet_numbering(doc, abstract_id: str, num_id: str) -> None:
    """Append an abstractNum/num pair whose level-0 ❖ bullet is red (FF0000)."""
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), abstract_id)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "")  # Wingdings ❖
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "FF0000")
    rpr.append(color)
    lvl.append(num_fmt)
    lvl.append(lvl_text)
    lvl.append(rpr)
    abstract.append(lvl)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    numbering.append(abstract)
    numbering.append(num)


def _bullet_color(doc, abstract_id: str):
    numbering = doc.part.numbering_part.element
    for abstract in numbering.findall(qn("w:abstractNum")):
        if abstract.get(qn("w:abstractNumId")) != abstract_id:
            continue
        for lvl in abstract.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl")) == "0":
                rpr = lvl.find(qn("w:rPr"))
                color = rpr.find(qn("w:color")) if rpr is not None else None
                return color.get(qn("w:val")) if color is not None else None
    return None


def _add_numbered_paragraph(doc, text: str, num_id: str):
    paragraph = doc.add_paragraph(text)
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), num_id)
    num_pr.append(ilvl)
    num_pr.append(nid)
    ppr.append(num_pr)
    return paragraph


def test_recolor_part3_intro_marker_turns_red_bullet_black(tmp_path):
    """第三部分 1.基因变异解析 引导段的 ❖ 红色装饰符应回黑（仅该段，作用域受限）。"""
    docx_path = tmp_path / "part3_marker.docx"
    doc = Document()
    _add_red_bullet_numbering(doc, abstract_id="901", num_id="91")
    _add_red_bullet_numbering(doc, abstract_id="902", num_id="92")
    _add_numbered_paragraph(
        doc,
        "在本次检测范围内，检出体细胞变异：11个，其中与靶向/免疫药物相关的变异：7个。",
        num_id="91",
    )
    _add_numbered_paragraph(doc, "其它无关的项目符号段落", num_id="92")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._recolor_part3_intro_marker(str(docx_path))

    rendered = Document(docx_path)
    # 引导段对应的 ❖ 变黑
    assert _bullet_color(rendered, "901") == "000000"
    # 无关项目符号保持原红色（作用域不外溢）
    assert _bullet_color(rendered, "902") == "FF0000"


def test_recolor_part3_intro_marker_is_idempotent_and_noop_without_intro(tmp_path):
    docx_path = tmp_path / "no_intro.docx"
    doc = Document()
    _add_red_bullet_numbering(doc, abstract_id="903", num_id="93")
    _add_numbered_paragraph(doc, "完全无关的段落", num_id="93")
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._recolor_part3_intro_marker(str(docx_path))
    renderer._recolor_part3_intro_marker(str(docx_path))

    rendered = Document(docx_path)
    # 没有引导段 → 不应改动任何编号颜色
    assert _bullet_color(rendered, "903") == "FF0000"


def test_reviewed_part3_knowledge_ships_dnmt3a_and_flt3_overrides():
    """历史终版中的 DNMT3A/FLT3 策展解析必须随包发布，避免回退到兜底文案。"""
    overlay_path = ROOT / "panels" / "crc_358_msi" / "rules" / "reviewed_part3_knowledge.yaml"
    data = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    sections = {
        (s["gene"], s.get("c_hgvs"), s.get("p_hgvs")): s for s in data.get("gene_sections", [])
    }

    dnmt3a = sections.get(("DNMT3A", "c.1367delA", "p.K456Sfs*195"))
    assert dnmt3a is not None, "DNMT3A 策展段落缺失"
    assert "表观遗传学" in dnmt3a["intro"]
    assert "p.K456Sfs*195" in dnmt3a["mutation_analysis"]

    flt3 = sections.get(("FLT3", "c.2537G>A", "p.G846D"))
    assert flt3 is not None, "FLT3 策展段落缺失"
    assert "酪氨酸激酶" in flt3["intro"]
    assert "p.G846D" in flt3["mutation_analysis"]


def test_reviewed_part3_legacy_crc_gene_level_candidates_are_promoted():
    """旧肠癌知识库通过审核后，应作为 CRC gene-level reviewed 内容随包发布。"""
    overlay_path = ROOT / "panels" / "crc_358_msi" / "rules" / "reviewed_part3_knowledge.yaml"
    data = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    gene_level_sections = {
        row.get("gene"): row
        for row in data.get("gene_sections", [])
        if row.get("gene") and not row.get("c_hgvs") and not row.get("p_hgvs")
    }

    legacy_genes = {
        "APC",
        "ARID1A",
        "ATM",
        "ATR",
        "BRAF",
        "BRCA1",
        "BRCA2",
        "CTNNB1",
        "EPCAM",
        "FBXW7",
        "KMT2C",
        "KMT2D",
        "KRAS",
        "MLH1",
        "MSH2",
        "MSH6",
        "NF1",
        "NRAS",
        "NTRK1",
        "NTRK2",
        "NTRK3",
        "PIK3CA",
        "PMS2",
        "PTEN",
        "SMAD4",
        "SMARCA4",
        "SMARCB1",
        "TCF7L2",
        "TP53",
    }
    assert not (legacy_genes - set(gene_level_sections))

    unsafe_phrases = (
        "{XX癌",
        "运营系统调取",
        "该样本检出的突变可能导致",
    )
    for gene in legacy_genes:
        section = gene_level_sections[gene]
        combined = f"{section.get('intro', '')}\n{section.get('mutation_analysis', '')}"
        assert section.get("intro"), f"{gene} 缺 gene-level intro"
        assert section.get("mutation_analysis"), f"{gene} 缺 gene-level mutation_analysis"
        for phrase in unsafe_phrases:
            assert phrase not in combined, f"{gene} gene-level 内容含不应泛化的旧库话术"

    provider = GeneKnowledgeProvider(
        {
            "enabled": True,
            "gene_knowledge_db": {
                "enabled": True,
                "path": "missing.xlsx",
                "reviewed_part3_overlay_path": str(overlay_path),
            },
        }
    )
    assert provider.load(base_path=str(ROOT))

    samples = {
        "BRAF": "BRAF V600E",
        "MLH1": "错配修复",
        "APC": "WNT信号通路",
        "TP53": "DNA结合结构域",
    }
    for gene, expected in samples.items():
        section = provider.build_gene_knowledge_section(
            gene=gene,
            c_hgvs="c.1A>G",
            p_hgvs="p.M1V",
            frequency=1.23,
            mutation_type="Missense",
        )
        combined = f"{section['intro']}\n{section['mutation_analysis']}"
        assert expected in combined


def test_reviewed_part3_crc_pressure_sites_are_variant_level_only():
    """报告组初审通过的 CRC358 压测缺口应按位点级入库，不能误作 gene-level 套话。"""
    overlay_path = ROOT / "panels" / "crc_358_msi" / "rules" / "reviewed_part3_knowledge.yaml"
    data = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    sections = {
        (row.get("gene"), row.get("c_hgvs"), row.get("p_hgvs", "")): row
        for row in data.get("gene_sections", [])
        if row.get("gene")
    }

    approved_sites = {
        ("ERBB2", "c.1133C>T", "p.P378L"): "现有证据不足以直接等同于ERBB2扩增",
        ("ERBB2", "c.2521C>A", "p.L841I"): "p.L841I错义变异",
        ("KMT2A", "c.3950_3954del", "p.K1317Sfs*7"): "p.K1317Sfs*7为移码变异",
        ("KMT2A", "c.8848A>G", "p.S2950G"): "p.S2950G为错义变异",
        ("ALK", "c.3388G>T", "p.V1130L"): "p.V1130L为错义变异",
        ("CARD11", "c.2059G>A", "p.A687T"): "p.A687T为错义变异",
        ("CCND1", "c.127T>G", "p.S43A"): "p.S43A为错义变异",
        ("FAT3", "c.11941G>A", "p.G3981S"): "p.G3981S为错义变异",
        ("GNAS", "c.1030G>A", "p.E344K"): "p.E344K为错义变异",
        ("KMT2B", "c.1681C>T", "p.P561S"): "p.P561S为错义变异",
        ("LRP1B", "c.1987G>A", "p.D663N"): "p.D663N为错义变异",
        ("PDGFB", "c.671G>A", "p.R224Q"): "p.R224Q为错义变异",
        ("PTPRS", "c.3073G>A", "p.V1025I"): "p.V1025I为错义变异",
    }
    unsafe_fragments = {
        "突变丰度为",
        "拷贝数为",
        "c.6445C>T",
        "p.R2149",
        "c.729G>T",
        "p.W243C",
        "c.3304G>A",
        "p.V1102M",
        "c.641G>A",
        "p.G214D",
        "c.4622A>T",
        "p.Q1541L",
        "c.659+1G>A",
        "c.3268G>A",
        "p.E1090K",
        "c.7677C>A",
        "p.C2559",
        "c.274G>A",
        "p.A92T",
        "c.604C>T",
        "p.Q202",
    }
    for key, expected in approved_sites.items():
        section = sections.get(key)
        assert section is not None, f"{key} 位点级 reviewed 内容缺失"
        combined = f"{section.get('intro', '')}\n{section.get('mutation_analysis', '')}"
        assert expected in combined
        for fragment in unsafe_fragments:
            assert fragment not in combined, f"{key} 含旧报告其他位点/丰度片段: {fragment}"

    rejected_sites = {
        ("BCL2L11", "c.561dup", "p.R188Tfs*77"),
        ("PCLO", "c.9368C>T", "p.T3123M"),
        ("PDE4DIP", "c.525C>G", "p.I175M"),
        ("SMAD3", "c.860G>A", "p.R287Q"),
        ("SOS1", "c.2536G>A", "p.E846K"),
    }
    for key in rejected_sites:
        assert key not in sections, f"{key} 初审暂不通过，不应进入 reviewed YAML"


def test_mutation_description_build_variant_lead_by_type():
    from reportgen.knowledge.mutation_description import MutationDescriptionGenerator

    gen = MutationDescriptionGenerator()
    assert (
        gen.build_variant_lead("PIK3CA", "c.1624G>A", "p.E542K")
        == "该样本检出PIK3CA基因c.1624G>A，p.E542K错义突变。"
    )
    assert (
        gen.build_variant_lead("TP53", "c.499C>T", "p.Q167*")
        == "该样本检出TP53基因c.499C>T，p.Q167*无义突变。"
    )
    assert (
        gen.build_variant_lead("PMS2", "c.1273delT", "p.S425Lfs*23")
        == "该样本检出PMS2基因c.1273delT，p.S425Lfs*23移码突变。"
    )
    # 剪接：无 p.HGVS → 省略蛋白部分，不留逗号
    assert (
        gen.build_variant_lead("MET", "c.153+2T>C", "") == "该样本检出MET基因c.153+2T>C剪接突变。"
    )
    # 缺基因或 c.HGVS → 空串
    assert gen.build_variant_lead("", "c.1A>T", "p.M1L") == ""
    assert gen.build_variant_lead("BRAF", "", "p.V600E") == ""
    assert "第4666位核苷酸重复" in gen.generate(
        "APC", "c.4666dup", "p.T1556Nfs*3", 17.5, "重复突变"
    )


def test_drug_analysis_relation_prepends_variant_lead():
    """2.1 用药提示解析中每条“基因变异与药物关联分析”应以变异描述开头。"""
    provider = GeneKnowledgeProvider({"enabled": False})
    provider._loaded = True
    provider.get_drug_full_info = lambda gene: [
        {
            "type": "benefit",
            "drug": "依维莫司",
            "c_point": "",
            "p_point": "",
            "relation": "PIK3CA基因激活突变，会导致信号通路异常活化。",
            "clinical": "",
        }
    ]
    sections = provider.build_drug_analysis_sections(
        [
            {
                "gene": "PIK3CA",
                "cHGVS": "c.1624G>A",
                "pHGVS": "p.E542K",
                "benefit_drugs": "依维莫司",
                "caution_drugs": "--",
            }
        ]
    )
    assert len(sections) == 1
    relation = sections[0]["relation"]
    assert relation.startswith("该样本检出PIK3CA基因c.1624G>A，p.E542K错义突变。")
    # 原 KB 正文仍保留在开头句之后
    assert "PIK3CA基因激活突变，会导致信号通路异常活化。" in relation


def test_drug_analysis_relation_lead_not_duplicated_when_already_present():
    provider = GeneKnowledgeProvider({"enabled": False})
    provider._loaded = True
    provider.get_drug_full_info = lambda gene: [
        {
            "type": "benefit",
            "drug": "依维莫司",
            "c_point": "",
            "p_point": "",
            "relation": "该样本检出PIK3CA基因c.1624G>A，p.E542K错义突变。已有开头。",
            "clinical": "",
        }
    ]
    sections = provider.build_drug_analysis_sections(
        [
            {
                "gene": "PIK3CA",
                "cHGVS": "c.1624G>A",
                "pHGVS": "p.E542K",
                "benefit_drugs": "依维莫司",
                "caution_drugs": "--",
            }
        ]
    )
    assert sections[0]["relation"].count("该样本检出") == 1


def _sig_label_inline_drawings(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text or ""
        if "检测者" in text and "审核者" in text:
            return paragraph._p.findall(".//" + qn("w:drawing") + "/" + qn("wp:inline"))
    return None


def test_run_processors_records_changed_and_deltas(tmp_path):
    """后处理器可观测性：每个处理器记录是否真改动文档 + 粗粒度增减，
    OK 但 changed=False 即静默失效信号。"""
    from reportgen.core.processors.base import ProcessorContext, run_processors

    docx_path = tmp_path / "obs.docx"
    doc = Document()
    doc.add_paragraph("起始段")
    doc.save(docx_path)

    class _P:
        def __init__(self, name, fn):
            self.name = name
            self.warning_message = f"{name} 失败"
            self._fn = fn

        def enabled(self, ctx):
            return True

        def run(self, ctx):
            self._fn(ctx.output_path)

    def add_para(path):
        d = Document(path)
        d.add_paragraph("新增")
        d.save(path)

    def silent_noop(path):
        return  # 模拟锚点没匹配上、静默跳过

    class _Logger:
        def warning(self, *a, **k):
            pass

    ctx = ProcessorContext(
        renderer=None,
        output_path=str(docx_path),
        template_path="",
        template_context={},
        logger=_Logger(),
    )
    results = run_processors([_P("changer", add_para), _P("noop", silent_noop)], ctx)

    by_name = {r.name: r for r in results}
    assert by_name["changer"].status == "OK"
    assert by_name["changer"].changed is True
    assert by_name["changer"].deltas.get("paragraphs") == 1
    # 静默 no-op：OK 但无改动 + 空 deltas（= 静默失效信号）
    assert by_name["noop"].status == "OK"
    assert by_name["noop"].changed is False
    assert by_name["noop"].deltas == {}
    # changed/deltas 进入 to_dict（落侧车，随时可查）
    assert "changed" in by_name["changer"].to_dict()
    assert "deltas" in by_name["changer"].to_dict()


def test_gene_list_table_first_column_not_bold(tmp_path):
    """基因检测列表表体统一不加粗（模板把首列基因加粗了，应与其它列一致）。"""
    docx_path = tmp_path / "genelist.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.rows[0].cells[0].text = "Gene List for MLseq (n=358)"
    genes = [["ABL1", "ABL2", "ACVR1B"], ["AMER1", "APC", "AR"]]
    for ri in (1, 2):
        for ci in range(3):
            cell = table.rows[ri].cells[ci]
            cell.text = genes[ri - 1][ci]
            if ci == 0:  # 模拟模板里首列加粗
                cell.paragraphs[0].runs[0].font.bold = True
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._compact_gene_list_tables(str(docx_path))

    rendered = Document(docx_path)
    rendered_table = rendered.tables[0]
    for ri in (1, 2):
        for run in rendered_table.rows[ri].cells[0].paragraphs[0].runs:
            assert run.font.bold is False


def test_rebuild_reference_section_covers_cited_pmids(tmp_path):
    """末尾 5.参考文献 应按正文实际引用重建：被引 PMID/NCT 全覆盖、示例排除、静态替换。"""
    docx_path = tmp_path / "refs.docx"
    doc = Document()
    doc.add_paragraph("BRCA1 突变对PARP抑制剂有预测作用[24579064]。又见[20664172，NCT04305496]。")
    doc.add_paragraph("说明：如编号[99999999]为参考文献PMID号示例。")  # 应被排除
    doc.add_paragraph("5. 参考文献")
    doc.add_paragraph("PMID:0000001 stale ref one")
    doc.add_paragraph("PMID:0000002 stale ref two")
    doc.save(docx_path)

    ctx = {
        "reference_lookup": {
            "pmid": {
                "24579064": "PMID: 24579064 BRCA paper",
                "20664172": "PMID: 20664172 PI3K paper",
            },
            "trial": {"NCT04305496": "NCT04305496 https://clinicaltrials.gov."},
            "other": [],
        }
    }
    TemplateRenderer(log_level="ERROR")._rebuild_reference_section(str(docx_path), ctx)

    rendered = Document(docx_path)
    start = next(
        i for i, p in enumerate(rendered.paragraphs) if (p.text or "").strip() == "5. 参考文献"
    )
    refs = [
        (p.text or "").strip() for p in rendered.paragraphs[start + 1 :] if (p.text or "").strip()
    ]
    blob = "\n".join(refs)
    assert "PMID: 24579064 BRCA paper" in blob
    assert "PMID: 20664172 PI3K paper" in blob
    assert "NCT04305496 https://clinicaltrials.gov." in blob
    assert "stale ref one" not in blob  # 静态条目被替换
    assert "99999999" not in blob  # “如编号”示例被排除
    assert blob.index("20664172") < blob.index("24579064")  # PMID 升序


def test_rebuild_reference_section_removes_uncited_historical_case_list(tmp_path):
    docx_path = tmp_path / "uncited_refs.docx"
    doc = Document()
    doc.add_paragraph("正文没有患者级文献引用。")
    doc.add_paragraph("5. 参考文献")
    stale = doc.add_paragraph("PMID:0000001 historical case ref")
    stale.paragraph_format.page_break_before = True
    stale.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("PMID:0000002 another historical case ref")
    qc_heading = doc.add_paragraph("本次检测质控结果")
    qc_heading.paragraph_format.page_break_before = True
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._rebuild_reference_section(
        str(docx_path),
        {
            "reference_lookup": {"pmid": {}, "trial": {}, "other": []},
            "report_content": {
                "reference_style": {
                    "font_name": "Calibri",
                    "east_asia_font_name": "微软雅黑",
                    "font_size_pt": 10.5,
                }
            },
        },
    )

    rendered = Document(docx_path)
    blob = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "historical case ref" not in blob
    assert "本报告未生成患者级动态参考文献" in blob
    assert "本次检测质控结果" in blob
    notice = next(
        paragraph
        for paragraph in rendered.paragraphs
        if "本报告未生成患者级动态参考文献" in paragraph.text
    )
    assert notice.paragraph_format.page_break_before is False
    assert not [node for node in notice._p.iter(qn("w:br")) if node.get(qn("w:type")) == "page"]
    assert notice.runs[0]._r.rPr.rFonts.get(qn("w:eastAsia")) == "微软雅黑"
    rendered_qc = next(
        paragraph for paragraph in rendered.paragraphs if paragraph.text == "本次检测质控结果"
    )
    assert rendered_qc.paragraph_format.page_break_before is False


def test_rebuild_reference_section_applies_configured_entry_font_only(tmp_path):
    """参考文献条目显式使用 Calibri 10.5，标题和后续章节不受影响。"""
    docx_path = tmp_path / "refs_style.docx"
    doc = Document()
    doc.add_paragraph("正文引用[24579064]。")
    heading = doc.add_paragraph()
    heading_run = heading.add_run("5. 参考文献")
    heading_run.font.name = "Microsoft YaHei"
    heading_run.font.size = Pt(12)
    stale = doc.add_paragraph("PMID:0000001 stale ref")
    stale.runs[0].font.name = "Tahoma"
    stale.runs[0].font.size = Pt(11)
    tail = doc.add_paragraph("本次检测质控结果")
    tail.runs[0].font.name = "Tahoma"
    tail.runs[0].font.size = Pt(11)
    doc.save(docx_path)

    ctx = {
        "reference_lookup": {
            "pmid": {"24579064": "PMID: 24579064 BRCA paper"},
            "trial": {},
            "other": [],
        },
        "report_content": {
            "reference_style": {
                "font_name": "Calibri",
                "font_size_pt": 10.5,
                "line_spacing": 1.0,
                "space_before_pt": 0,
                "space_after_pt": 0,
                "keep_with_next": False,
                "keep_together": False,
                "widow_control": False,
                "page_break_before": False,
            }
        },
    }
    TemplateRenderer(log_level="ERROR")._rebuild_reference_section(str(docx_path), ctx)

    rendered = Document(docx_path)
    start, end = find_reference_section_bounds(rendered.paragraphs)
    assert start is not None
    reference_runs = [
        run
        for paragraph in rendered.paragraphs[start + 1 : end]
        if (paragraph.text or "").strip()
        for run in paragraph.runs
    ]
    assert reference_runs
    for run in reference_runs:
        assert run.font.name == "Calibri"
        assert run.font.size.pt == pytest.approx(10.5)
        r_fonts = run._r.get_or_add_rPr().rFonts
        assert r_fonts is not None
        for attr in ("ascii", "hAnsi", "eastAsia"):
            assert r_fonts.get(qn(f"w:{attr}")) == "Calibri"
    for paragraph in rendered.paragraphs[start + 1 : end]:
        if not paragraph.text.strip():
            continue
        fmt = paragraph.paragraph_format
        assert fmt.line_spacing == 1.0
        assert fmt.space_before.pt == pytest.approx(0.0)
        assert fmt.space_after.pt == pytest.approx(0.0)
        assert fmt.keep_with_next is False
        assert fmt.keep_together is False
        assert fmt.widow_control is False
        assert fmt.page_break_before is False

    rendered_heading = rendered.paragraphs[start].runs[0]
    assert rendered_heading.font.name == "Microsoft YaHei"
    assert rendered_heading.font.size.pt == pytest.approx(12.0)
    rendered_tail = next(
        paragraph
        for paragraph in rendered.paragraphs
        if paragraph.text.strip() == "本次检测质控结果"
    ).runs[0]
    assert rendered_tail.font.name == "Tahoma"
    assert rendered_tail.font.size.pt == pytest.approx(11.0)


def test_back_cover_artwork_uses_idempotent_configured_anchor_position(tmp_path):
    """只移动替代说明为“报告末页”的浮动图，重复处理不累计偏移。"""
    docx_path = tmp_path / "back_cover.docx"
    doc = Document()

    def add_floating_anchor(paragraph, description: str, offset: int) -> None:
        drawing = OxmlElement("w:drawing")
        anchor = OxmlElement("wp:anchor")
        position_h = OxmlElement("wp:positionH")
        position_h.set("relativeFrom", "column")
        pos_offset = OxmlElement("wp:posOffset")
        pos_offset.text = str(offset)
        position_h.append(pos_offset)
        doc_pr = OxmlElement("wp:docPr")
        doc_pr.set("id", str(abs(offset) + 1))
        doc_pr.set("name", description)
        doc_pr.set("descr", description)
        anchor.append(position_h)
        anchor.append(doc_pr)
        drawing.append(anchor)
        paragraph.add_run()._r.append(drawing)

    add_floating_anchor(doc.add_paragraph(), "封面", 123456)
    add_floating_anchor(doc.add_paragraph(), "报告末页", -448945)
    doc.save(docx_path)

    ctx = {
        "report_content": {
            "back_cover_artwork": {
                "description_marker": "报告末页",
                "horizontal_relative_from": "column",
                "horizontal_alignment": "page_left",
                "page_left_inset_cm": 0.0,
            }
        }
    }
    renderer = TemplateRenderer(log_level="ERROR")
    renderer._normalize_back_cover_artwork(str(docx_path), ctx)
    renderer._normalize_back_cover_artwork(str(docx_path), ctx)

    rendered = Document(docx_path)
    offsets = {}
    for anchor in rendered.element.body.iter(qn("wp:anchor")):
        doc_pr = anchor.find(qn("wp:docPr"))
        position_h = anchor.find(qn("wp:positionH"))
        assert doc_pr is not None and position_h is not None
        pos_offset = position_h.find(qn("wp:posOffset"))
        assert pos_offset is not None
        offsets[doc_pr.get("descr")] = (
            int(pos_offset.text),
            position_h.get("relativeFrom"),
        )
    assert offsets["封面"] == (123456, "column")
    assert offsets["报告末页"] == (-1143000, "column")


def test_rebuild_reference_section_preserves_following_report_sections(tmp_path):
    """重建文献只能改文献区，不能吞掉其后的 QC/方法/局限/公司正文。"""
    docx_path = tmp_path / "refs_with_tail.docx"
    doc = Document()
    doc.add_paragraph("正文引用[24579064]。")
    doc.add_paragraph("5. 参考文献")
    doc.add_paragraph("PMID:0000001 stale ref one")
    doc.add_paragraph("PMID:0000002 stale ref two")
    doc.add_paragraph("")
    doc.add_paragraph("本次检测质控结果")
    qc_table = doc.add_table(rows=2, cols=2)
    qc_table.cell(0, 0).text = "质控指标"
    qc_table.cell(0, 1).text = "结果"
    qc_table.cell(1, 0).text = "Q30"
    qc_table.cell(1, 1).text = "合格"
    doc.add_paragraph("*质控结果分为“合格”和“风险”两个等级。")
    doc.add_paragraph("高通量测序检测方法说明")
    doc.add_paragraph("检测方法正文必须保留。")
    doc.add_paragraph("高通量测序局限性")
    doc.add_paragraph("检测局限正文必须保留。")
    doc.add_paragraph("脉络医学检验简介")
    doc.add_paragraph("公司简介正文必须保留。")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._rebuild_reference_section(
        str(docx_path),
        {
            "reference_lookup": {
                "pmid": {"24579064": "PMID: 24579064 BRCA paper"},
                "trial": {},
                "other": [],
            }
        },
    )

    rendered = Document(docx_path)
    paragraphs = [(p.text or "").strip() for p in rendered.paragraphs]
    blob = "\n".join(paragraphs)
    assert "PMID: 24579064 BRCA paper" in blob
    assert "stale ref one" not in blob
    assert "stale ref two" not in blob
    tail = [
        "本次检测质控结果",
        "*质控结果分为“合格”和“风险”两个等级。",
        "高通量测序检测方法说明",
        "检测方法正文必须保留。",
        "高通量测序局限性",
        "检测局限正文必须保留。",
        "脉络医学检验简介",
        "公司简介正文必须保留。",
    ]
    assert all(text in paragraphs for text in tail)
    assert [paragraphs.index(text) for text in tail] == sorted(
        paragraphs.index(text) for text in tail
    )
    assert any(
        "质控指标" in " ".join(cell.text for row in table.rows for cell in row.cells)
        for table in rendered.tables
    )


def test_reference_bounds_ignore_local_colon_heading_in_endometrial_appendix():
    package = load_panel_package("endometrial_29", project_root=ROOT)
    document = Document(package.resolve_template_file())
    texts = [(paragraph.text or "").strip() for paragraph in document.paragraphs]

    start, end = find_reference_section_bounds(document.paragraphs)

    assert start is not None
    assert texts[start].startswith("5") and "参考文献" in texts[start]
    assert texts[end].startswith("第四部分")
    assert texts.index("参考文献：") > end


def test_rebuild_references_blocks_cited_document_without_main_heading(tmp_path):
    docx_path = tmp_path / "missing_reference_heading.docx"
    doc = Document()
    doc.add_paragraph("正文引用[24579064]。")
    doc.save(docx_path)

    with pytest.raises(ValueError, match="reference heading is missing"):
        TemplateRenderer(log_level="ERROR")._rebuild_reference_section(
            str(docx_path),
            {"reference_lookup": {"pmid": {}, "trial": {}, "other": []}},
        )


def test_build_reference_lookup_includes_extra_references(tmp_path):
    overlay = tmp_path / "overlay.yaml"
    overlay.write_text(
        yaml.safe_dump(
            {"extra_references": ["PMID: 28351930 Hyperprogressors after Immunotherapy"]},
            allow_unicode=True,
        ),
        encoding="utf-8",
    )
    provider = GeneKnowledgeProvider(
        {
            "enabled": True,
            "gene_knowledge_db": {
                "enabled": True,
                "path": "missing.xlsx",
                "reviewed_part3_overlay_path": str(overlay),
            },
        }
    )
    provider.load(base_path=str(tmp_path))
    lookup = provider.build_reference_lookup()
    assert lookup["pmid"].get("28351930", "").startswith("PMID: 28351930")


def test_render_inline_signatures_places_images_on_label_line(tmp_path):
    from PIL import Image

    img = tmp_path / "sig.png"
    Image.new("RGB", (80, 30), "white").save(img)
    docx_path = tmp_path / "sig.docx"
    doc = Document()
    doc.add_paragraph("检测者：                审核者：")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._render_inline_signatures(
        str(docx_path),
        {
            "detector_signature_image_path": str(img),
            "reviewer_signature_image_path": str(img),
        },
    )

    rendered = Document(docx_path)
    inline = _sig_label_inline_drawings(rendered)
    assert inline is not None and len(inline) == 2
    # 不应残留浮动锚点
    for paragraph in rendered.paragraphs:
        assert not paragraph._p.findall(".//" + qn("wp:anchor"))


def test_render_inline_signatures_no_path_leaves_label_clean(tmp_path):
    docx_path = tmp_path / "sig_empty.docx"
    doc = Document()
    doc.add_paragraph("检测者：                审核者：")
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._render_inline_signatures(str(docx_path), {})

    rendered = Document(docx_path)
    inline = _sig_label_inline_drawings(rendered)
    assert inline is not None and len(inline) == 0
    # 标签文字保留
    assert any(
        "检测者" in (p.text or "") and "审核者" in (p.text or "") for p in rendered.paragraphs
    )


def test_signature_resolution_semantics_for_warning(tmp_path):
    """无对应签名图 → 解析返回空串（驱动“应有提示”分支）；有则返回路径。"""
    from reportgen.core.signature_library import resolve_signature_path

    (tmp_path / "signatures.yaml").write_text(
        "detector:\n  张三: sigs/zhangsan.png\nreviewer:\n  李四: sigs/lisi.png\n",
        encoding="utf-8",
    )
    assert resolve_signature_path(str(tmp_path), "detector", "张三")
    # 签名库中没有的人 → 空串（report_generator 据此产生警告而非静默空白）
    assert resolve_signature_path(str(tmp_path), "detector", "王医生") == ""
    assert resolve_signature_path(str(tmp_path), "reviewer", "赵主任") == ""
