"""Native TOC and historical appendix boundaries; synthetic documents only."""

import copy
from types import SimpleNamespace

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
from reportgen.core.golden_case import inspect_lung_draft_table_boundaries
from reportgen.core.legacy_reference import (
    _extract_features,
    _section_presence,
    snapshot_docx_report,
)
from reportgen.core.qa_gate import (
    _current_output_profile_from_package,
    _run_single_current_output_contract,
)
from reportgen.core.qa_report import (
    _build_business_checks,
    _build_style_checks,
    _inspect_toc,
    _is_empty_numbered_paragraph,
    _paragraph_has_zero_first_line_indent,
    _read_part3_text,
)
from reportgen.core.report_generator import ReportGenerator, _GenerationState
from reportgen.core.template_bridge_358 import load_panel_config
from reportgen.core.template_renderer import TemplateRenderer
from reportgen.docx_sections import find_reference_section_bounds
from reportgen.models.report_data import ReportData
from reportgen.panels.loader import load_panel_package
from scripts.build_lung_draft_packages import (
    EMPTY_APPROVED_DRUG_NOTICE,
    ensure_table_separator,
    install_refreshable_toc,
    install_shared_modules,
    normalize_b_family_faq_flow,
    normalize_draft_case_fields,
    normalize_draft_variant_flow,
    normalize_optional_drug_block,
)
from scripts.validate_lung_small_panel_drafts import canonical_parity_value, inspect_word_scope

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def synthetic_variant_table(document):
    table = document.add_table(rows=3, cols=9)
    table.cell(0, 0).text = "基因突变信息"
    table.cell(1, 1).text = "转录本号"
    for cell in table.rows[2].cells:
        cell.text = "合成字段"
    return table


def test_small_draft_qa_rejects_empty_drug_header_and_unprotected_variant_rows():
    document = Document()
    synthetic_variant_table(document)
    empty = document.add_table(rows=1, cols=2)
    empty.cell(0, 0).text = "药物名称"
    empty.cell(0, 1).text = "相关基因"
    check = _build_style_checks(document, "lung_13", {})["docx_style_rules"]
    assert check["status"] == "FAIL"
    codes = {item["code"] for item in check["failures"]}
    assert {
        "HEADER_ONLY_DRUG_TABLE", "VARIANT_ROW_CAN_SPLIT", "VARIANT_HEADER_NOT_REPEATED",
    } <= codes


@pytest.mark.parametrize("panel", ["lung_13", "lung_62", "lung_62_pdl1", "lung_588"])
def test_installed_draft_variants_have_explicit_zero_indent_and_pagination_guards(panel):
    package = load_panel_package(panel)
    document = Document(package.resolve_template_file())
    result = _build_style_checks(document, panel, {})["docx_style_rules"]
    assert result["status"] == "PASS", result


@pytest.mark.parametrize("has_drug", [False, True])
def test_optional_drug_empty_state_is_explicit_without_removing_real_rows(tmp_path, has_drug):
    document = Document()
    document.add_paragraph("前述结果")
    document.add_page_break()
    title = document.add_paragraph("3.2 其它潜在获益上市药物提示*")
    title.paragraph_format.page_break_before = True
    document.add_paragraph("历史固定栏目说明")
    table = document.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "药物名称"
    table.cell(0, 1).text = "相关基因"
    table.cell(1, 0).text = "{%tr for row in chemotherapy %}"
    table.cell(2, 0).text = "{{ row.drug }}"
    table.cell(2, 1).text = "{{ row.gene }}"
    table.cell(3, 0).text = "{%tr endfor %}"
    document.add_page_break()
    next_title = document.add_paragraph("4. 检测结果说明")
    next_title.paragraph_format.page_break_before = True
    document.add_paragraph("固定说明完整保留")
    normalize_optional_drug_block(document, table)
    path = tmp_path / "optional-drugs.docx"
    document.save(path)
    template = DocxTemplate(path)
    rows = [{"drug": "合成药物", "gene": "合成基因"}] if has_drug else []
    template.render({"chemotherapy": rows})
    template.save(path)
    actual = Document(path)
    visible = "".join(p.text for p in actual.paragraphs)
    assert (EMPTY_APPROVED_DRUG_NOTICE in visible) is (not has_drug)
    assert "固定说明完整保留" in visible and "历史固定栏目说明" in visible
    assert not actual.element.body.findall(".//" + qn("w:br"))
    assert all(not p.paragraph_format.page_break_before for p in actual.paragraphs)
    if has_drug:
        assert [[c.text for c in row.cells] for row in actual.tables[0].rows] == [
            ["药物名称", "相关基因"], ["合成药物", "合成基因"],
        ]
    else:
        assert not actual.tables


def test_variant_flow_preserves_cells_widths_and_runs_while_overriding_indent():
    document = Document()
    table = synthetic_variant_table(document)
    before = [[cell.text for cell in row.cells] for row in table.rows]
    widths = [cell.width for cell in table.rows[2].cells]
    normalize_draft_variant_flow(table)
    assert before == [[cell.text for cell in row.cells] for row in table.rows]
    assert widths == [cell.width for cell in table.rows[2].cells]
    assert _build_style_checks(document, "lung_13", {})["docx_style_rules"]["status"] == "PASS"
    assert _build_style_checks(document, "lung_588_pdl1", {}) == {}
    table.rows[2]._tr.trPr.find(qn("w:cantSplit")).set(qn("w:val"), "false")
    check = _build_style_checks(document, "lung_13", {})["docx_style_rules"]
    assert check["status"] == "FAIL"
    assert any(item["code"] == "VARIANT_ROW_CAN_SPLIT" for item in check["failures"])
    normalize_draft_variant_flow(table)
    table.cell(2, 1).paragraphs[0]._p.pPr.ind.set(qn("w:firstLine"), "420")
    check = _build_style_checks(document, "lung_13", {})["docx_style_rules"]
    assert "VARIANT_CELL_INHERITED_INDENT" in check["failure_codes"]


@pytest.mark.parametrize("attribute", ["firstLine", "hanging", "firstLineChars", "hangingChars"])
@pytest.mark.parametrize("value,expected", [("0", True), ("420", False), ("-20", False)])
def test_native_zero_indent_forms_are_equivalent_but_nonzero_still_blocks(
    attribute, value, expected,
):
    document = Document()
    paragraph = document.add_paragraph("合成数字字段")
    paragraph._p.get_or_add_pPr().get_or_add_ind().set(qn("w:" + attribute), value)
    assert _paragraph_has_zero_first_line_indent(paragraph, document) is expected


def test_indent_checker_resolves_style_defaults_and_direct_zero_override():
    document = Document()
    paragraph = document.add_paragraph("合成数字字段")
    style_indent = paragraph.style.element.get_or_add_pPr().get_or_add_ind()
    style_indent.set(qn("w:firstLine"), "420")
    assert not _paragraph_has_zero_first_line_indent(paragraph, document)
    direct = paragraph._p.get_or_add_pPr().get_or_add_ind()
    direct.set(qn("w:start"), "0")  # Margin reset alone must not erase indentation.
    assert not _paragraph_has_zero_first_line_indent(paragraph, document)
    direct.set(qn("w:hanging"), "0")  # LibreOffice's exported equivalent of firstLine=0.
    assert _paragraph_has_zero_first_line_indent(paragraph, document)
    direct.set(qn("w:hanging"), "invalid")
    assert not _paragraph_has_zero_first_line_indent(paragraph, document)


@pytest.mark.parametrize("panel", ["lung_13", "lung_62", "lung_62_pdl1", "lung_588", "crc_358_msi"])
@pytest.mark.parametrize("explicit", [False, True])
def test_draft_web_default_filename_never_claims_final_and_crc_naming_is_unchanged(
    tmp_path, panel, explicit,
):
    package = load_panel_package(panel)
    generator = ReportGenerator(log_level="ERROR")
    report = ReportData()
    for key, value in {
        "patient_name": "合成测试", "sample_id": "SYNTHETIC-ONLY",
        "project_name": package.display_name, "report_date": "2000-01-01",
    }.items():
        report.set_field(key, value)
    state = _GenerationState(
        excel_file="synthetic.xlsx", template_file="unused.docx", output_dir=str(tmp_path),
        excel_data=SimpleNamespace(file_path="synthetic.xlsx", metadata={}), report_data=report,
        panel_package=package, output_filename="explicit-name.docx" if explicit else None,
    )
    generator._stage_output_path(SimpleNamespace(artifacts={}, metrics={}), state)
    if explicit:
        assert state.output_filename == "explicit-name.docx"
    elif panel.startswith("lung_"):
        assert state.output_filename.endswith("_评审草稿.docx")
        assert "终版" not in state.output_filename
    else:
        assert state.output_filename.endswith("_终版.docx")


@pytest.mark.parametrize("panel", ["lung_13", "lung_62", "lung_62_pdl1", "lung_588"])
@pytest.mark.parametrize("omit_biomarker_heading", [False, True])
def test_current_output_gate_retains_package_declared_section_aliases(
    tmp_path, panel, omit_biomarker_heading
):
    document = Document()
    for value in (
        "本次检出体细胞变异：1个",
        "" if omit_biomarker_heading else "补充检测结果：TMB、MSI 与化疗药物基因组学",
        "肿瘤突变负荷为10.0 mutations/Mb，TMB-H；MSS",
        "本次 PD-L1 免疫组化检测结果：TPS 50%，CPS 52，阳性（高表达）。",
        "基因检测列表",
    ):
        document.add_paragraph(value)
    document.add_table(rows=1, cols=1).cell(0, 0).text = "基因突变信息"
    path = tmp_path / "synthetic-current-output.docx"
    document.save(path)
    profile = _current_output_profile_from_package(load_panel_package(panel))
    result = _run_single_current_output_contract(
        panel, output_file=path, output_root=tmp_path / "snapshot",
        profile=profile, source_step="synthetic-test", fail_on_warn=True,
    )
    assert result["status"] == ("FAIL" if omit_biomarker_heading else "PASS"), result
    if omit_biomarker_heading:
        assert any(row["code"] == "LEGACY_SECTION_MISSING" for row in result["issues"])


@pytest.mark.parametrize("phrase", ["本次共检出体细胞变异：4个", "本次检出体细胞变异：4 个"])
def test_snapshot_counts_both_reviewed_summary_wordings(phrase):
    assert _extract_features(phrase, tables=[])["total_variants_count"] == 4
    assert _extract_features("本次检出其他基因：13个", tables=[])["total_variants_count"] is None


def test_panel_section_aliases_do_not_weaken_other_panel_or_missing_section(tmp_path):
    document = Document()
    document.add_paragraph("补充检测结果：TMB、MSI 与化疗药物基因组学")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "基因突变信息"
    path = tmp_path / "synthetic-sections.docx"
    document.save(path)
    aliases = {
        "variant_summary": ["基因突变信息"],
        "biomarkers": ["补充检测结果：TMB、MSI 与化疗药物基因组学"],
    }
    plain = snapshot_docx_report(path, panel="crc_358_msi")["section_presence"]
    scoped = snapshot_docx_report(path, panel="lung_62", section_aliases=aliases)[
        "section_presence"
    ]
    assert not plain["variant_summary"] and not plain["biomarkers"]
    assert scoped["variant_summary"] and scoped["biomarkers"]
    assert not scoped["gene_list"]
    assert not _section_presence("实际章节缺失", aliases)["biomarkers"]


@pytest.mark.parametrize("values", [[], [""], "任意字符串", [False]])
def test_empty_or_malformed_section_alias_cannot_match_every_document(values):
    with pytest.raises(ValueError):
        _section_presence("无相应章节", {"biomarkers": values})


@pytest.mark.parametrize(
    "boundary",
    [
        "第三部分：附录", "5. 附录", "补充检测结果：TMB、MSI 与化疗药物基因组学",
        "肺癌相关重要基因变异及药物提示",
    ],
)
def test_reference_cleanup_never_deletes_following_lung_modules(tmp_path, boundary):
    document = Document()
    for value in (
        "检测结果解析",
        "参考文献",
        "合成旧文献",
        boundary,
        "CNV待复核：合成来源待审",
        "固定附录保留",
    ):
        document.add_paragraph(value)
    document.add_table(rows=1, cols=1).cell(0, 0).text = "合成 PGx 行"
    assert find_reference_section_bounds(document.paragraphs) == (1, 3)
    path = tmp_path / "reference-boundary.docx"
    document.save(path)
    TemplateRenderer(log_level="ERROR")._rebuild_reference_section(str(path), {})
    actual = Document(path)
    text = "\n".join(p.text for p in actual.paragraphs)
    assert boundary in text and "CNV待复核：合成来源待审" in text and "固定附录保留" in text
    assert "合成旧文献" not in text
    assert actual.tables[0].cell(0, 0).text == "合成 PGx 行"


def test_empty_reference_list_does_not_consume_following_guideline_caption(tmp_path):
    document = Document()
    document.add_paragraph("参考文献")
    document.add_paragraph("肺癌相关重要基因变异及药物提示")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "来源指南行"
    path = tmp_path / "reference-immediate-guideline.docx"
    document.save(path)
    renderer = TemplateRenderer(log_level="ERROR")
    renderer._rebuild_reference_section(str(path), {})
    actual = Document(path)
    assert [p.text for p in actual.paragraphs] == [
        "参考文献",
        "本报告未生成患者级动态参考文献；历史病例固定参考文献已移除，待报告组根据最终启用的解释内容复核。",
        "肺癌相关重要基因变异及药物提示",
    ]
    assert actual.tables[0].cell(0, 0).text == "来源指南行"
    renderer._rebuild_reference_section(str(path), {})
    assert [p.text for p in Document(path).paragraphs] == [p.text for p in actual.paragraphs]


def test_flat_historical_toc_becomes_refreshable_without_old_page_numbers(tmp_path):
    document = Document()
    document.styles.add_style("toc 1", 1)
    document.add_paragraph("目录")
    labels = ["第一部分：检测结果", "第二部分：结果解析", "第三部分：附录"]
    for label, page in zip(labels, [42, 56, 99]):
        document.add_paragraph(label + "\t" + str(page), style="toc 1")
    body = [document.add_paragraph(label) for label in labels]
    install_refreshable_toc(document)
    assert all("\t" not in p.text for p in document.paragraphs)
    assert all('w:outlineLvl w:val="0"' in p._p.xml for p in body)
    assert " TOC " in document.element.xml
    path = tmp_path / "converted-toc.docx"
    document.save(path)
    # Empty cache must still fail the page-number check until real LO refresh.
    result = _inspect_toc(document.paragraphs, output_path=path, require_fields=True)
    assert result["status"] == "FAIL"


def test_flat_c_family_toc_excludes_inherited_citations_and_local_reference_captions():
    document = Document()
    document.styles.add_style("toc 1", 1)
    document.add_paragraph("目录")
    labels = ["第一部分：检测结果", "第二部分：结果解析", "4. 参考文献", "第三部分：附录"]
    for name in labels:
        document.add_paragraph(name + "\t99", style="toc 1")
    for name in labels:
        document.add_paragraph(name)
    caption = document.add_paragraph("参考文献：", style="Heading 2")
    citation = document.add_paragraph(
        "[3] Synthetic long reference, 2015: 100–200", style="Heading 2"
    )
    cell = document.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0]
    cell.style = "Heading 2"
    install_refreshable_toc(document)
    for paragraph in (caption, citation, cell):
        assert paragraph._p.pPr.find(qn("w:outlineLvl")).get(qn("w:val")) == "9"
    global_reference = next(p for p in document.paragraphs if p.text == "4. 参考文献")
    assert global_reference._p.pPr.find(qn("w:outlineLvl")).get(qn("w:val")) == "0"
    assert caption.text == "参考文献：" and citation.text.startswith("[3] Synthetic")


def test_empty_b_family_toc_uses_actual_major_headings():
    document = Document()
    document.add_paragraph("目    录")
    labels = ["第一部分：基本信息", "第二部分：检测结果", "第三部分：结果解析", "第四部分：附录"]
    for label in labels:
        document.add_paragraph(label)
    install_refreshable_toc(document)
    assert " TOC " in document.element.xml
    assert [p.text for p in document.paragraphs if p.text in labels] == labels
    assert "目录待排版引擎刷新" in [p.text for p in document.paragraphs]


def test_b_family_floating_toc_is_replaced_and_section_boundary_follows_field():
    document = Document()
    title = document.add_paragraph("目录")
    labels = [
        "第一部分：基本信息", "第二部分：检测结果", "第三部分：附录", "肺癌诊疗知识", "参考文献",
    ]
    cache = "".join(
        f"<w:p><w:r><w:tab/><w:t>{name}</w:t><w:tab/><w:t>99</w:t></w:r></w:p>"
        for name in labels
    )
    title._p.append(parse_xml(
        f'<w:r xmlns:w="{W}" xmlns:v="urn:schemas-microsoft-com:vml">'
        f'<w:pict><v:shape><v:textbox><w:txbxContent>{cache}'
        '</w:txbxContent></v:textbox></v:shape></w:pict></w:r>'
    ))
    decoration = parse_xml(f'<w:r xmlns:w="{W}"><w:pict/></w:r>')
    title._p.append(decoration)
    title._p.get_or_add_pPr().append(parse_xml(
        f'<w:sectPr xmlns:w="{W}"><w:pgNumType w:start="1"/></w:sectPr>'
    ))
    document.add_paragraph(labels[0])
    cell_p = document.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0]
    cell_p._p.get_or_add_pPr().append(parse_xml(f'<w:outlineLvl xmlns:w="{W}" w:val="1"/>'))
    for name in labels[1:3]:
        document.add_paragraph(name)
    document.add_table(rows=1, cols=1).cell(0, 0).text = "固定 FAQ 内容保留"
    document.add_paragraph()
    document.add_paragraph()
    document.add_page_break()
    appendix_heading = document.add_paragraph("2. 肺癌诊疗知识")
    prose = document.add_paragraph("不进入目录的历史正文", style="Heading 2")
    local_references = document.add_paragraph("参考文献：")
    global_references = document.add_paragraph("5. 参考文献")
    install_refreshable_toc(document)
    assert "99" not in document.element.xml
    assert not list(document.element.iter(qn("w:txbxContent")))
    assert not list(title._p.iter(qn("w:pict")))
    assert title._p.find(".//" + qn("w:sectPr")) is None
    field = title._p.getnext()
    assert " TOC " in field.xml
    assert field.getnext().find(".//" + qn("w:sectPr")) is not None
    assert cell_p._p.pPr.find(qn("w:outlineLvl")).get(qn("w:val")) == "9"
    assert prose._p.pPr.find(qn("w:outlineLvl")).get(qn("w:val")) == "9"
    assert appendix_heading._p.pPr.find(qn("w:outlineLvl")).get(qn("w:val")) == "1"
    assert appendix_heading._p.getprevious().tag == qn("w:tbl")
    assert appendix_heading.paragraph_format.page_break_before
    assert appendix_heading.paragraph_format.keep_with_next
    assert document.tables[-1].cell(0, 0).text == "固定 FAQ 内容保留"
    assert local_references._p.pPr.find(qn("w:outlineLvl")).get(qn("w:val")) == "9"
    assert global_references._p.pPr.find(qn("w:outlineLvl")).get(qn("w:val")) == "1"


def test_b_family_faq_flow_binds_short_blocks_without_changing_answers():
    document = Document()
    title = document.add_paragraph("常见问题解答")
    faq = document.add_table(rows=2, cols=2)
    faq.cell(0, 0).text = "问题7"
    faq.cell(0, 1).text = "合成问题"
    faq.cell(1, 0).merge(faq.cell(1, 1)).text = "合成答案第一行\n合成答案第二行"
    long_faq = document.add_table(rows=2, cols=1)
    long_faq.cell(0, 0).text = "问题8"
    long_faq.cell(1, 0).text = "长答案" * 300
    next_heading = document.add_paragraph("2. 肺癌诊疗知识")
    next_heading.paragraph_format.page_break_before = True
    before = [table._tbl.xml for table in (faq, long_faq)]
    normalize_b_family_faq_flow(document)
    assert title.text == "1. 常见问题解答" and title._p.pPr.numPr.numId.val == 0
    assert not next_heading.paragraph_format.page_break_before
    assert next_heading.paragraph_format.keep_with_next
    assert faq.cell(1, 0).text == "合成答案第一行\n合成答案第二行"
    assert all(row._tr.trPr.find(qn("w:cantSplit")) is not None for row in faq.rows)
    assert faq.cell(0, 0).paragraphs[0].paragraph_format.keep_with_next
    assert not faq.cell(1, 0).paragraphs[-1].paragraph_format.keep_with_next
    assert faq._tbl.xml != before[0] and long_faq._tbl.xml == before[1]


@pytest.mark.parametrize("number_id, visible", [(0, False), (1, True), (9, True), (None, True)])
def test_empty_numbering_distinguishes_explicit_suppression(number_id, visible):
    document = Document()
    paragraph = document.add_paragraph()
    num = "" if number_id is None else f'<w:numId w:val="{number_id}"/>'
    paragraph._p.get_or_add_pPr().append(parse_xml(
        f'<w:numPr xmlns:w="{W}"><w:ilvl w:val="0"/>{num}</w:numPr>'
    ))
    assert _is_empty_numbered_paragraph(paragraph) is visible


@pytest.mark.parametrize("old_count", [2, 9, 123])
def test_case_summary_and_basic_values_never_inherit_historical_case_data(old_count):
    document = Document()
    basic = document.add_table(rows=1, cols=2)
    basic.cell(0, 0).text = "姓名"
    basic.cell(0, 1).text = "{{ patient_name }}"
    document.add_paragraph(f"在本次检测范围内，检出体细胞变异：{old_count}个，其中与靶向药物相关的变异：0个。暂未筛选到合适的靶向药物。")
    document.add_page_break()
    title = document.add_paragraph("肺癌相关重要基因变异及药物提示")
    normalize_draft_case_fields(document, basic)
    paragraphs = [p.text for p in document.paragraphs]
    assert "暂未筛选" not in "".join(paragraphs)
    assert f"变异：{old_count}个" not in "".join(paragraphs)
    assert "{{ total_variants_count }}" in "".join(paragraphs)
    assert "{{ drug_related_count }}" in "".join(paragraphs)
    assert str(basic.cell(0, 1).paragraphs[0].runs[0].font.color.rgb) == "000000"
    assert title.paragraph_format.keep_with_next
    assert 'w:type="page"' not in document.element.xml


def test_required_native_toc_cannot_pass_from_numbered_body_prose(tmp_path):
    document = Document()
    document.add_paragraph("目录")
    document.add_paragraph("2.1 检测结果")
    document.add_paragraph("3.1 合成正文数值 52")
    path = tmp_path / "missing-native.docx"
    document.save(path)
    result = _inspect_toc(document.paragraphs, output_path=path, require_fields=True)
    assert result["status"] == "FAIL"


@pytest.mark.parametrize("stale", [False, True])
def test_every_case_summary_count_must_agree_even_when_one_correct_summary_exists(stale):
    text = "本次检出体细胞变异：1个；与靶向药物用药相关的变异有：1个。"
    if stale:
        text += "检出体细胞变异：2个，其中与靶向药物相关的变异：0个。"
    result = _build_business_checks(
        text, {"total_variants_count": 1, "drug_related_count": 1}, "lung_62_pdl1"
    )
    for name in ("case_total_count_consistency", "case_drug_count_consistency"):
        assert result[name]["status"] == ("FAIL" if stale else "PASS")


def test_native_pagination_never_invokes_crc_toc_rewrite(monkeypatch):
    renderer = TemplateRenderer(log_level="ERROR")
    events = []
    monkeypatch.setattr(renderer, "_document_contains_toc", lambda _path: True)
    monkeypatch.setattr(
        renderer, "_refresh_fields_with_native_engine", lambda _path: events.append("refresh")
    )
    monkeypatch.setattr(renderer, "_set_update_fields", lambda _path: events.append("fields"))
    context = {"panel_style": {"toc": {"mode": "native"}}}
    renderer._populate_static_toc_page_numbers("synthetic.docx", context)
    assert events == ["refresh", "fields"]
    monkeypatch.setattr(renderer, "_document_contains_toc", lambda _path: False)
    with pytest.raises(RuntimeError, match="Native TOC field was lost"):
        renderer._populate_static_toc_page_numbers("synthetic.docx", context)


def native_toc(tmp_path, pages=(1, 2), target="section", closed=True):
    document = Document()
    entries = "".join(
        '<w:p><w:hyperlink w:anchor="section"><w:r><w:t>检测项目 62</w:t></w:r>'
        + (f"<w:r><w:tab/><w:t>{page}</w:t></w:r>" if page is not None else "")
        + "</w:hyperlink></w:p>"
        for page in pages
    )
    block = parse_xml(
        f'<w:sdt xmlns:w="{W}"><w:sdtContent><w:p><w:r><w:instrText>'
        ' TOC \\o "1-3" \\h </w:instrText></w:r></w:p>' + entries + "</w:sdtContent></w:sdt>"
    )
    document._element.body.insert(0, block)
    anchor = document.add_paragraph("正文")._p
    anchor.append(parse_xml(f'<w:bookmarkStart xmlns:w="{W}" w:id="7" w:name="{target}"/>'))
    if closed:
        anchor.append(parse_xml(f'<w:bookmarkEnd xmlns:w="{W}" w:id="7"/>'))
    path = tmp_path / "native.docx"
    document.save(path)
    return path


def test_native_sdt_toc_reads_all_cached_pages(tmp_path):
    path = native_toc(tmp_path)
    assert all("检测项目" not in p.text for p in Document(path).paragraphs)
    result = _inspect_toc([], output_path=path)
    assert result["mode"] == "native_TOC"
    assert result["status"] == "PASS"
    assert result["page_numbered_line_count"] == result["line_count"] == 2


@pytest.mark.parametrize("label", ["[3] Synthetic reference, 2015: 100–200", "https://example.org/reference"])
def test_native_toc_rejects_bibliography_entries_even_with_valid_pages(tmp_path, label):
    path = native_toc(tmp_path)
    document = Document(path)
    for node in document.element.iter(qn("w:t")):
        if node.text == "检测项目 62":
            node.text = label
    document.save(path)
    result = _inspect_toc([], output_path=path)
    assert result["status"] == "FAIL"
    assert result["citation_entry_count"] == 2
    assert result["page_numbered_line_count"] == 2


@pytest.mark.parametrize("pages", [(None,), (1, None), (0,), (-1,)])
def test_native_toc_missing_pages_never_pass_from_heading_numbers(tmp_path, pages):
    result = _inspect_toc([], output_path=native_toc(tmp_path, pages=pages))
    assert result["status"] == "WARN"


@pytest.mark.parametrize("target,closed", [("wrong", True), ("section", False)])
def test_native_toc_broken_or_unclosed_bookmarks_fail(tmp_path, target, closed):
    result = _inspect_toc([], output_path=native_toc(tmp_path, target=target, closed=closed))
    assert result["status"] == "FAIL"
    assert result["missing_bookmarks"] or result["unclosed_bookmarks"]


def test_native_toc_cannot_hide_a_stale_floating_directory(tmp_path):
    path = native_toc(tmp_path)
    document = Document(path)
    paragraph = document.add_paragraph()
    value = "第一部分：基本信息 99 第二部分：检测结果 99 第三部分：附录 99"
    paragraph._p.append(parse_xml(
        f'<w:r xmlns:w="{W}" xmlns:v="urn:schemas-microsoft-com:vml">'
        '<w:pict><v:shape><v:textbox><w:txbxContent>'
        f'<w:p><w:r><w:t>{value}</w:t></w:r></w:p>'
        '</w:txbxContent></v:textbox></v:shape></w:pict></w:r>'
    ))
    document.save(path)
    result = _inspect_toc([], output_path=path)
    assert result["status"] == "FAIL"
    assert result["stale_floating_cache_count"] == 1


@pytest.mark.parametrize("unsafe_inside", [False, True])
def test_family_specific_part3_bounds_exclude_only_fixed_appendix(unsafe_inside):
    document = Document()
    document.add_paragraph("结直肠癌：章节外的固定文献")
    document.add_paragraph("第二部分：检测结果解析")
    document.add_paragraph("结直肠癌错误解析" if unsafe_inside else "肺癌安全解析")
    document.add_paragraph("附录：基因介绍")
    document.add_paragraph("结直肠癌：附录固定文献")
    policy = {
        "enabled": True,
        "scan_scope": "part3",
        "terms": ["结直肠癌"],
        "start_heading": "第二部分：检测结果解析",
        "end_heading": "附录：基因介绍",
    }
    section = _read_part3_text(document, policy)
    assert "固定文献" not in section
    check = _build_business_checks(
        "结直肠癌", {"part3_cross_cancer_residual_scan": policy}, "lung_13", section
    )["part3_cross_cancer_residuals"]
    assert check["status"] == ("WARN" if unsafe_inside else "PASS")


def test_missing_configured_part3_heading_keeps_full_document_fallback():
    document = Document()
    document.add_paragraph("结直肠癌未分节的错误解析")
    assert _read_part3_text(document, {"start_heading": "不存在的标题"}) is None


@pytest.mark.parametrize(
    "tamper",
    [None, "gene", "guideline", "pgx", "merged_guideline", "merged_pgx", "spanned", "extra_pgx"],
)
def test_final_word_scope_detects_missing_genes_guidelines_and_changed_pgx(tamper):
    package = load_panel_package("lung_13")
    config = load_panel_config(panel_package=package)
    context = {
        "lung_guideline_drug_results": [
            {
                "gene": row.get("display") or "/".join(row["genes"]),
                "drugs": row["drugs"],
                "clinical_note": row["clinical_note"],
                "result": "未检出",
            }
            for row in config.lung_guideline_drug_rows
        ],
        "drug_shunbo": [
            {
                "DrugDisplay": "合成药物",
                "Gene": "合成基因",
                "Locus": "rsTEST",
                "Level": "1B",
                "Genotype": "AA",
                "Result": "合成来源结果",
            }
        ],
    }
    document = Document()
    genes = document.add_table(rows=1, cols=1)
    genes.cell(0, 0).text = "肺癌13基因检测列表"
    for gene in sorted(config.crc_important_genes):
        genes.add_row().cells[0].text = gene
    guide = document.add_table(rows=1, cols=4)
    for cell, value in zip(
        guide.rows[0].cells, ["检测基因", "本癌种相关治疗药物", "临床提示", "检测结果"]
    ):
        cell.text = value
    for row in context["lung_guideline_drug_results"]:
        for cell, key in zip(guide.add_row().cells, ["gene", "drugs", "clinical_note", "result"]):
            cell.text = row[key]
    pgx = document.add_table(rows=1, cols=6)
    for cell, value in zip(
        pgx.rows[0].cells, ["合成药物", "基因", "检测位点", "等级", "检测结果", "用药提示"]
    ):
        cell.text = value
    for cell, key in zip(
        pgx.add_row().cells, ["DrugDisplay", "Gene", "Locus", "Level", "Genotype", "Result"]
    ):
        cell.text = context["drug_shunbo"][0][key]
    second = copy.deepcopy(pgx._tbl)
    pgx._tbl.addnext(second)
    context["drug_kabo"] = copy.deepcopy(context["drug_shunbo"])
    if tamper == "extra_pgx":
        second.addnext(copy.deepcopy(pgx._tbl))
        context["drug_boleihuahewu"] = copy.deepcopy(context["drug_shunbo"])
    if tamper in {"spanned", "merged_pgx"}:
        # Simulate LO's 16-column grid with only six logical cells per row.
        for table in (pgx._tbl, second):
            for _ in range(10):
                table.tblGrid.append(copy.deepcopy(table.tblGrid[0]))
            for row in table.tr_lst:
                for cell, span in zip(row.tc_lst, [1, 3, 3, 2, 3, 4]):
                    cell.get_or_add_tcPr().get_or_add_gridSpan().val = span
    if tamper == "gene":
        genes.cell(1, 0).text = "UNASSAYED"
    elif tamper == "guideline":
        guide._tbl.remove(guide.rows[-1]._tr)
    elif tamper == "pgx":
        pgx.cell(1, 5).text = "与来源不符"
    elif tamper == "merged_guideline":
        prefix = copy.deepcopy(guide.rows[0]._tr)
        for node in prefix.iter(qn("w:t")):
            node.text = "上一个表的内容"
        guide.rows[0]._tr.addprevious(prefix)
    elif tamper == "merged_pgx":
        for row in list(second.tr_lst):
            pgx._tbl.append(row)
        second.getparent().remove(second)
    result = inspect_word_scope(document, context, package)
    assert bool(result["failures"]) is (tamper not in {None, "spanned", "extra_pgx"})
    boundaries = inspect_lung_draft_table_boundaries(document)
    assert boundaries["passed"] is (tamper not in {"merged_guideline", "merged_pgx"})
    if tamper in {"merged_guideline", "merged_pgx", "spanned"}:
        assert result["pgx_detail_row_count"] == result["expected_pgx_detail_row_count"] == 2
        assert result["guideline_row_count"] == len(context["lung_guideline_drug_results"])
        assert "word_guideline_rows_mismatch" not in result["failures"]
        assert "word_pgx_detail_rows_mismatch" not in result["failures"]
    if tamper == "merged_guideline":
        assert "word_guideline_table_merged" in result["failures"]
    if tamper == "merged_pgx":
        assert "word_pgx_tables_merged" in result["failures"]


def test_parity_normalizes_empty_cells_without_masking_clinical_evidence():
    source = {"drug": [{"Reference": None, "\r": None, "Result": "来源结论", "Level": "1B"}]}
    rendered = {"drug": [{"Reference": "", "\r": "", "Result": "来源结论", "Level": "1B"}]}
    assert canonical_parity_value(source) == canonical_parity_value(rendered)
    for key in ("Result", "Level"):
        changed = copy.deepcopy(rendered)
        changed["drug"][0][key] = "篡改"
        assert canonical_parity_value(source) != canonical_parity_value(changed)
    assert canonical_parity_value(source) != canonical_parity_value({"drug": []})
    assert canonical_parity_value([0, None, "a"]) != canonical_parity_value(["", 0, "a"])


@pytest.mark.parametrize("existing_caption", [False, True])
def test_guide_boundary_uses_a_surviving_caption(existing_caption):
    document = Document()
    document.add_table(rows=1, cols=3)
    if existing_caption:
        document.add_paragraph("已有指南标题")
    document.add_paragraph("{%p endif %}")
    guide = document.add_table(rows=1, cols=4)
    ensure_table_separator(guide, "新增指南标题")
    assert [p.text for p in document.paragraphs].count("新增指南标题") == (not existing_caption)


def test_pgx_captions_survive_render_only_for_nonempty_drug_groups(tmp_path):
    document = Document()
    variants = document.add_table(rows=1, cols=1)
    document.add_paragraph("合成附录边界")
    install_shared_modules(document, {"rich_end": "合成附录边界"}, {"variants": variants})
    source = tmp_path / "synthetic-modules.docx"
    output = tmp_path / "synthetic-modules-rendered.docx"
    document.save(source)
    template = DocxTemplate(source)
    template.render({
        "drug_shunbo": [{"Result": "来源甲"}],
        "drug_kabo": [{"Result": "来源乙"}],
    })
    template.save(output)
    actual = Document(output)
    pgx = [
        table for table in actual.tables
        if len(table.rows[0].cells) == 6 and table.rows[0].cells[1].text.strip() == "基因"
    ]
    assert len(pgx) == 2
    for table in pgx:
        previous = table._tbl.getprevious()
        assert previous.tag == qn("w:p")
        caption = "".join(node.text or "" for node in previous.iter(qn("w:t")))
        assert caption == table.cell(0, 0).text.strip() + "药物基因组学明细"
    assert len([p for p in actual.paragraphs if p.text.endswith("药物基因组学明细")]) == 2
