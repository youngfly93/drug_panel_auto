# ruff: noqa: E402, I001
"""Regression contracts for the 2026-07-20 report-group feedback."""

from __future__ import annotations

import copy
import sys
from pathlib import Path
from types import SimpleNamespace

import pytest
import yaml
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from reportgen.core.report_generator import ReportGenerator
from reportgen.core.report_summary import build_report_summary
from reportgen.core.qa_report import (
    _build_business_checks,
    _build_style_checks,
    _business_issues,
)
from reportgen.core.field_mapper import FieldMapper
from reportgen.core.template_bridge_358 import (
    PanelConfig,
    _build_crc_approved_drugs,
    _build_nccn_and_immune_fields,
    build_immune_variants,
    enhance_report_data,
    load_panel_config,
)
from reportgen.core.template_renderer import TemplateRenderer
from reportgen.knowledge.gene_knowledge import GeneKnowledgeProvider
from reportgen.knowledge.governance import load_and_validate_overlay
from reportgen.knowledge.quality import (
    is_generic_mutation_analysis,
    profile_panel_targeted_drug_contracts,
)
from reportgen.models.excel_data import ExcelDataSource
from reportgen.models.report_data import ReportData
from reportgen.panels.loader import load_panel_package
from reportgen.rules.targeted_drugs import load_targeted_drug_rule_context


def _panel_provider(panel_id: str) -> GeneKnowledgeProvider:
    settings = yaml.safe_load((ROOT / "config/settings.yaml").read_text(encoding="utf-8"))
    gene_config = copy.deepcopy(settings["knowledge_bases"]["gene_knowledge_db"])
    package = load_panel_package(panel_id, project_root=ROOT)
    gene_config["panel_id"] = panel_id
    gene_config["reviewed_part3_overlay_paths"] = (
        ReportGenerator._resolve_panel_reviewed_part3_overlays(package)
    )
    provider = GeneKnowledgeProvider(
        {
            "enabled": True,
            "panel_id": panel_id,
            "gene_symbol_aliases": package.raw.get("gene_symbol_aliases") or {},
            "gene_knowledge_db": gene_config,
        }
    )
    assert provider.load(base_path=str(ROOT))
    return provider


def _crc358_provider() -> GeneKnowledgeProvider:
    return _panel_provider("crc_358_msi")


def _section(provider: GeneKnowledgeProvider, gene: str, c_hgvs: str, p_hgvs: str):
    return provider.build_gene_knowledge_section(
        gene=gene,
        c_hgvs=c_hgvs,
        p_hgvs=p_hgvs,
        frequency=12.34,
        mutation_type="Missense",
        has_drug=False,
    )


def test_fixed_domain_survives_erbb2_variant_overlay():
    section = _section(_crc358_provider(), "ERBB2", "c.2521C>A", "p.L841I")
    combined = f"{section['intro']}\n{section['mutation_analysis']}"

    assert "蛋白全长为1255个氨基酸" in combined
    assert "Protein kinase" in combined
    assert "ERBB2基因编码的蛋白全长" in combined
    assert combined.count("蛋白全长为1255个氨基酸") == 1


def test_complementary_kras_domains_are_preserved_once():
    section = _section(_crc358_provider(), "KRAS", "c.35G>A", "p.G12D")
    combined = f"{section['intro']}\n{section['mutation_analysis']}"

    assert combined.count("RAS结构域") == 1
    assert combined.count("Hypervariable region") == 1
    assert combined.count("蛋白全长为189个氨基酸") == 1


def test_base_domain_column_wins_when_legacy_intro_tail_conflicts():
    settings = yaml.safe_load(
        (ROOT / "config/settings.yaml").read_text(encoding="utf-8")
    )
    gene_config = copy.deepcopy(settings["knowledge_bases"]["gene_knowledge_db"])
    gene_config["reviewed_part3_overlay_paths"] = []
    gene_config["reviewed_part3_overlay_path"] = ""
    provider = GeneKnowledgeProvider(
        {"enabled": True, "gene_knowledge_db": gene_config}
    )
    assert provider.load(base_path=str(ROOT))

    smad4 = _section(provider, "SMAD4", "c.1A>G", "p.M1V")
    pik3ca = _section(provider, "PIK3CA", "c.2A>G", "p.M2V")

    assert smad4["fixed_domain_text"].count("编码的蛋白全长") == 1
    assert "552个氨基酸" in smad4["fixed_domain_text"]
    assert "3056个氨基酸" not in smad4["fixed_domain_text"]
    assert pik3ca["fixed_domain_text"].count("编码的蛋白全长") == 1
    assert "1068个氨基酸" in pik3ca["fixed_domain_text"]


def test_missing_gene_domains_ship_as_governed_first_review_overlay():
    overlay_path = ROOT / "panels/crc_358_msi/rules/reviewed_part3_domain_overlay_20260720.yaml"
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    assert overlay_path in {
        Path(path) for path in ReportGenerator._resolve_panel_reviewed_part3_overlays(package)
    }

    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {"approved_for_runtime": 6}
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]

    provider = _crc358_provider()
    expected = {
        "FANCA": ("1455个氨基酸", "ArcN结构域"),
        "PALB2": ("1186个氨基酸", "WD40重复结构域"),
        "DNMT3A": ("912个氨基酸", "ADD结构域"),
        "RAD51D": ("328个氨基酸", "RecA样ATP酶家族"),
        "PIK3CA": ("1068个氨基酸", "PI3K/PI4K catalytic结构域"),
        "SMAD4": ("552个氨基酸", "MH2结构域"),
    }
    for gene, markers in expected.items():
        section = _section(provider, gene, "c.999A>G", "p.X333Y")
        combined = f"{section['intro']}\n{section['mutation_analysis']}"
        assert all(marker in combined for marker in markers), (gene, combined)
        assert combined.count("编码的蛋白全长") == 1, (gene, combined)


def test_high_risk_gene_fallbacks_are_governed_specific_and_drug_free():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/reviewed_part3_gene_fallbacks_20260721.yaml"
    )
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    assert overlay_path in {
        Path(path) for path in ReportGenerator._resolve_panel_reviewed_part3_overlays(package)
    }

    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {"approved_for_runtime": 9}
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]

    raw = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    assert raw["drug_sections"] == []
    assert all(row["panels"] == ["crc_301_msi", "crc_358_msi"] for row in raw["gene_sections"])

    provider = _crc358_provider()
    expected_boundaries = {
        "AKT1": "未收录位点不得仅因基因名称推导靶向用药",
        "ALK": "普通单核苷酸变异不能等同于ALK融合",
        "EGFR": "不能由任意EGFR单核苷酸变异单独推出",
        "ERBB2": "不得等同于HER2扩增",
        "FLT3": "不能套用血液肿瘤中的FLT3证据",
        "MET": "序列变异、扩增和过表达属于不同的生物学事件",
        "RET": "重排或融合与点突变的适用癌种及证据不可互换",
        "ROS1": "不能自动视为致癌激活事件",
        "SETD2": "区分有依据支持的功能缺失事件与普通错义变异",
    }
    for gene, marker in expected_boundaries.items():
        section = _section(provider, gene, "c.999999A>G", "p.X999999Y")
        assert not is_generic_mutation_analysis(gene, section["mutation_analysis"])
        assert marker in section["mutation_analysis"]


def test_crc_important_gene_fallbacks_are_panel_scoped_and_drug_free():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/reviewed_part3_crc_important_fallbacks_20260721.yaml"
    )
    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {"approved_for_runtime": 10}
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]

    raw = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    assert raw["drug_sections"] == []
    crc301_expected = {"ERBB3", "FAT1", "GNAS", "LRP1B", "ZFHX3"}
    crc358_only = {"ACVR2A", "AMER1", "FAT4", "MUTYH", "PTPRT"}
    assert {
        row["gene"]
        for row in raw["gene_sections"]
        if "crc_301_msi" in row["panels"]
    } == crc301_expected
    assert {
        row["gene"]
        for row in raw["gene_sections"]
        if row["panels"] == ["crc_358_msi"]
    } == crc358_only

    crc301 = _panel_provider("crc_301_msi")
    crc358 = _crc358_provider()
    for gene in crc301_expected:
        section = _section(crc301, gene, "c.999999A>G", "p.X999999Y")
        assert not is_generic_mutation_analysis(gene, section["mutation_analysis"])
    for gene in crc301_expected | crc358_only:
        section = _section(crc358, gene, "c.999999A>G", "p.X999999Y")
        assert not is_generic_mutation_analysis(gene, section["mutation_analysis"])


def test_p2a_drug_candidate_gene_fallbacks_are_scoped_specific_and_drug_free():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_drug_candidate_fallbacks_p2a_20260721.yaml"
    )
    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {"approved_for_runtime": 18}
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]

    raw = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    assert raw["drug_sections"] == []
    crc358_only = {"AKT3", "CSF3R"}
    shared = {
        "ABL1",
        "ARAF",
        "BTK",
        "CSF1R",
        "DDR2",
        "EPHA2",
        "ERBB4",
        "FGFR1",
        "FGFR3",
        "FLT1",
        "IGF1R",
        "KIT",
        "MTOR",
        "PDGFRA",
        "RAF1",
        "SMO",
    }
    assert {
        row["gene"]
        for row in raw["gene_sections"]
        if "crc_301_msi" in row["panels"]
    } == shared
    assert {
        row["gene"]
        for row in raw["gene_sections"]
        if row["panels"] == ["crc_358_msi"]
    } == crc358_only

    expected_boundaries = {
        "ABL1": "不得套用BCR::ABL1相关血液肿瘤证据",
        "AKT3": "不得自动推导AKT通路药物敏感性",
        "ARAF": "不得按BRAF或其他RAF基因",
        "BTK": "不得直接生成BTK抑制剂",
        "CSF1R": "不得据此推导CSF1R靶向获益",
        "CSF3R": "不得套用血液肿瘤治疗证据",
        "DDR2": "不得由其他癌种热点或抑制剂研究外推",
        "EPHA2": "不得由表达研究或其他癌种资料直接生成",
        "ERBB4": "普通单核苷酸变异不得视为受体激活",
        "FGFR1": "不得仅凭基因名称推导FGFR抑制剂获益",
        "FGFR3": "不得套用其他癌种的热点或融合药物证据",
        "FLT1": "不得作为抗血管生成药物敏感或耐药标志",
        "IGF1R": "不得据此直接生成IGF1R通路药物结论",
        "KIT": "不得外推到结直肠癌普通KIT变异",
        "MTOR": "不能因检出MTOR变异即推导mTOR抑制剂获益",
        "PDGFRA": "不得直接外推到结直肠癌未收录变异",
        "RAF1": "未收录RAF1位点不得按BRAF热点",
        "SMO": "不得套用其他肿瘤中的Hedgehog通路药物结论",
    }
    crc301 = _panel_provider("crc_301_msi")
    crc358 = _crc358_provider()
    for gene, marker in expected_boundaries.items():
        section = _section(crc358, gene, "c.999999A>G", "p.X999999Y")
        assert not is_generic_mutation_analysis(gene, section["mutation_analysis"])
        assert marker in section["mutation_analysis"]
    for gene in shared:
        section = _section(crc301, gene, "c.999999A>G", "p.X999999Y")
        assert not is_generic_mutation_analysis(gene, section["mutation_analysis"])
    for gene in crc358_only:
        section = _section(crc301, gene, "c.999999A>G", "p.X999999Y")
        assert is_generic_mutation_analysis(gene, section["mutation_analysis"])


def test_p2b_drug_candidate_gene_fallbacks_separate_expression_from_snv():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_drug_candidate_fallbacks_p2b_20260721.yaml"
    )
    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {"approved_for_runtime": 16}
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]

    raw = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    assert raw["drug_sections"] == []
    crc301_only = {"ESR1"}
    shared = {
        "AR",
        "BCL2",
        "CCND1",
        "CCND2",
        "CDK4",
        "CDK6",
        "FGF3",
        "FGF4",
        "HIF1A",
        "MDM2",
        "MYC",
        "NRG1",
        "PDGFB",
        "TOP2A",
        "VEGFA",
    }
    assert {
        row["gene"]
        for row in raw["gene_sections"]
        if "crc_358_msi" in row["panels"]
    } == shared
    assert {
        row["gene"]
        for row in raw["gene_sections"]
        if row["panels"] == ["crc_301_msi"]
    } == crc301_only

    expected_boundaries = {
        "AR": "不得套用前列腺癌中的雄激素依赖",
        "BCL2": "不得直接生成BCL2抑制剂获益或耐药结论",
        "CCND1": "普通错义变异不能替代拷贝数或表达结果",
        "CCND2": "普通错义变异不得单独证明细胞周期通路依赖",
        "CDK4": "不能脱离RB和CDKN2A等通路状态",
        "CDK6": "不得在缺少完整通路证据时生成",
        "FGF3": "配体改变不能直接等同于FGFR受体激活",
        "FGF4": "配体基因位点也不能替代FGFR受体的精确事件",
        "HIF1A": "不得作为抗血管生成或缺氧通路药物的敏感性标志",
        "MDM2": "未收录位点不得仅凭基因名称推导MDM2抑制剂获益",
        "MYC": "未收录位点不能证明MYC依赖",
        "NRG1": "未收录位点不得按NRG1融合处理",
        "PDGFB": "不得直接生成PDGFR抑制剂用药结论",
        "TOP2A": "未收录位点不得单独预测药物敏感或耐药",
        "VEGFA": "不能由单个VEGFA位点直接推导",
    }
    crc301 = _panel_provider("crc_301_msi")
    crc358 = _crc358_provider()
    for gene, marker in expected_boundaries.items():
        for provider in (crc301, crc358):
            section = _section(provider, gene, "c.999999A>G", "p.X999999Y")
            assert not is_generic_mutation_analysis(
                gene, section["mutation_analysis"]
            )
            assert marker in section["mutation_analysis"]
    esr1 = _section(crc301, "ESR1", "c.999999A>G", "p.X999999Y")
    assert not is_generic_mutation_analysis("ESR1", esr1["mutation_analysis"])
    assert "不得套用乳腺癌或子宫内膜癌" in esr1["mutation_analysis"]
    crc358_esr1 = _section(crc358, "ESR1", "c.999999A>G", "p.X999999Y")
    assert is_generic_mutation_analysis("ESR1", crc358_esr1["mutation_analysis"])


def test_p2c_drug_candidate_gene_fallbacks_block_hereditary_and_hrd_inference():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_drug_candidate_fallbacks_p2c_20260721.yaml"
    )
    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {"approved_for_runtime": 21}
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]

    raw = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    assert raw["drug_sections"] == []
    crc358_only = {"BCOR"}
    shared = {
        "CBL",
        "CDH1",
        "CDK12",
        "CDKN2A",
        "CHEK1",
        "CHEK2",
        "EZH2",
        "FANCC",
        "FANCL",
        "FLCN",
        "GNA11",
        "GNAQ",
        "MYD88",
        "NF2",
        "NOTCH2",
        "NPM1",
        "RAD51B",
        "RAD54L",
        "RB1",
        "TERT",
    }
    assert {
        row["gene"]
        for row in raw["gene_sections"]
        if "crc_301_msi" in row["panels"]
    } == shared
    assert {
        row["gene"]
        for row in raw["gene_sections"]
        if row["panels"] == ["crc_358_msi"]
    } == crc358_only

    expected_boundaries = {
        "BCOR": "不得套用血液系统或其他肿瘤中的分型",
        "CBL": "不得套用急性髓系白血病",
        "CDH1": "不能自动证明胚系来源或遗传性肿瘤综合征",
        "CDK12": "不能独立证明同源重组修复缺陷",
        "CDKN2A": "不能自动确定胚系状态或CDK4/6抑制剂敏感性",
        "CHEK1": "不能证明CHEK1依赖",
        "CHEK2": "不能自动证明胚系易感、双等位失活",
        "EZH2": "不得由表达研究或其他肿瘤热点直接推导",
        "FANCC": "不能自动证明胚系来源、双等位失活",
        "FANCL": "不能自动证明胚系状态、双等位功能缺失",
        "FLCN": "不能自动证明胚系来源、双等位失活",
        "GNA11": "不得套用黑色素细胞来源肿瘤",
        "GNAQ": "其他癌种证据不可直接迁移",
        "MYD88": "不得套用淋巴系统肿瘤分型",
        "NF2": "不能自动证明胚系神经纤维瘤病2型",
        "NOTCH2": "普通错义变异不能仅按基因名称分类",
        "NPM1": "不得套用血液肿瘤分型、预后或治疗证据",
        "RAD51B": "不能自动证明双等位失活、同源重组修复缺陷",
        "RAD54L": "不足以证明HRD或药物敏感",
        "RB1": "不能自动证明胚系视网膜母细胞瘤易感",
        "TERT": "启动子热点、拷贝数或表达改变与编码区普通序列变异",
    }
    crc301 = _panel_provider("crc_301_msi")
    crc358 = _crc358_provider()
    for gene, marker in expected_boundaries.items():
        section = _section(crc358, gene, "c.999999A>G", "p.X999999Y")
        assert not is_generic_mutation_analysis(gene, section["mutation_analysis"])
        assert marker in section["mutation_analysis"]
    for gene in shared:
        section = _section(crc301, gene, "c.999999A>G", "p.X999999Y")
        assert not is_generic_mutation_analysis(gene, section["mutation_analysis"])
    bcor = _section(crc301, "BCOR", "c.999999A>G", "p.X999999Y")
    assert is_generic_mutation_analysis("BCOR", bcor["mutation_analysis"])


def test_p3a_dna_hereditary_fallbacks_are_specific_shared_and_drug_free():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3a_dna_hereditary_fallbacks_20260721.yaml"
    )
    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {"approved_for_runtime": 26}
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]

    raw = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    assert raw["drug_sections"] == []
    expected_boundaries = {
        "ATRX": "不能证明胚系ATRX综合征",
        "BLM": "不能自动证明胚系来源、双等位功能缺失",
        "BMPR1A": "不足以诊断遗传性息肉病",
        "FAM175A": "现行官方符号为ABRAXAS1",
        "FANCE": "不能自动证明胚系来源、双等位失活",
        "FANCF": "需结合合子状态、功能和癌种证据",
        "FANCG": "不得直接生成铂类、PARP抑制剂或免疫治疗结论",
        "FANCI": "不能自动证明胚系Fanconi贫血",
        "FH": "不能自动证明胚系来源、双等位FH失活",
        "GALNT12": "不能证明胚系致病性",
        "HNF1A": "不能自动证明胚系疾病",
        "HOXB13": "不能自动证明遗传易感",
        "MRE11A": "现行官方符号为MRE11",
        "MSH3": "不能自动证明错配修复缺陷、MSI-H",
        "NBN": "不得仅凭基因名称推导PARP抑制剂",
        "PMS1": "不能替代MLH1、MSH2、MSH6、PMS2",
        "POLD1": "经验证的校对外切酶结构域致病变异",
        "RAD51": "通路成员身份或单个未收录位点不能证明",
        "RAD52": "不能自动证明RAD52功能缺失",
        "RAD54B": "不足以证明HRD",
        "RECQL4": "不能自动证明胚系来源",
        "SDHC": "不能自动证明胚系副神经节瘤易感",
        "TERC": "不能自动证明胚系来源、端粒功能缺陷",
        "TMEM127": "不能自动证明胚系来源",
        "TP53BP1": "不能证明修复通路重编程",
        "XRCC2": "不能证明双等位功能缺失",
    }
    assert {row["gene"] for row in raw["gene_sections"]} == set(
        expected_boundaries
    )
    assert all(
        row["panels"] == ["crc_301_msi", "crc_358_msi"]
        for row in raw["gene_sections"]
    )

    for panel_id in ("crc_301_msi", "crc_358_msi"):
        provider = _panel_provider(panel_id)
        for gene, marker in expected_boundaries.items():
            section = _section(provider, gene, "c.999999A>G", "p.X999999Y")
            assert not is_generic_mutation_analysis(
                gene, section["mutation_analysis"]
            )
            assert marker in section["mutation_analysis"]


def test_p3a_exact_corrections_remove_cross_event_and_gene_leakage():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3a_legacy_exact_corrections_20260721.yaml"
    )
    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {"approved_for_runtime": 6}
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]
    assert yaml.safe_load(overlay_path.read_text(encoding="utf-8"))["drug_sections"] == []

    expected = {
        ("FANCI", "c.2879G>A", "p.R960Q"): (
            "不得据此解释为HRD",
            ("Abstract IA25", "免疫治疗临床获益相关"),
        ),
        ("ATRX", "c.3277del", "p.R1093Gfs*25"): (
            "仅按潜在功能缺失事件保守解释",
            ("Janat Fazal-Salom",),
        ),
        ("GALNT12", "c.592C>T", "p.R198C"): (
            "不能自动证明遗传性结直肠癌",
            ("可能对蛋白功能有一定的影响",),
        ),
        ("HOXB13", "c.583del", "p.W195Gfs*84"): (
            "甲基化、表达改变或其他癌种的胚系易感研究",
            ("上游CpG岛甲基化在结肠癌中是特有的",),
        ),
        ("RAD52", "c.481G>A", "p.A161T"): (
            "不得直接生成PARP抑制剂",
            ("CHEK2", ""),
        ),
        ("XRCC2", "c.121+2T>C", ""): (
            "仍需RNA或其他功能证据确认",
            ("因此，该样本检出的突变可能与疾病的发生发展相关",),
        ),
    }
    for panel_id in ("crc_301_msi", "crc_358_msi"):
        provider = _panel_provider(panel_id)
        for selector, (marker, forbidden) in expected.items():
            section = _section(provider, *selector)
            assert marker in section["mutation_analysis"], (selector, section)
            assert not any(
                token in section["mutation_analysis"] for token in forbidden
            ), (selector, section)


def test_p3b_immune_fallbacks_are_specific_shared_and_drug_free():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3b_immune_fallbacks_20260721.yaml"
    )
    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {"approved_for_runtime": 12}
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]

    raw = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    assert raw["drug_sections"] == []
    expected_boundaries = {
        "B2M": "单个肿瘤序列变异不能自动证明双等位失活",
        "HLA-A": "单个HLA-A序列变异不足以证明抗原呈递功能缺失",
        "HLA-B": "单个肿瘤序列变异不足以证明抗原呈递缺陷",
        "HLA-DPB1": "胚系等位基因分型与肿瘤体细胞序列变异不是同一证据类型",
        "HLA-DQA1": "单个序列变异不足以证明抗原呈递功能异常",
        "HLA-DQB1": "单个位点不能自动证明HLA II类功能缺失",
        "HLA-DRB1": "不能用于解释结直肠肿瘤中的任意体细胞位点",
        "HLA-DRB5": "不能混为同一事件",
        "HLA-G": "单个序列变异不足以证明免疫抑制表型",
        "IFNGR1": "单个IFNGR1肿瘤序列变异不足以证明IFN-γ通路失活",
        "IFNGR2": "单个IFNGR2序列变异不能自动证明受体复合物",
        "PDCD1LG2": "普通PDCD1LG2序列变异不是同一检测指标",
    }
    assert {row["gene"] for row in raw["gene_sections"]} == set(
        expected_boundaries
    )
    assert all(
        row["panels"] == ["crc_301_msi", "crc_358_msi"]
        for row in raw["gene_sections"]
    )

    for panel_id in ("crc_301_msi", "crc_358_msi"):
        provider = _panel_provider(panel_id)
        for gene, marker in expected_boundaries.items():
            section = _section(provider, gene, "c.999999A>G", "p.X999999Y")
            assert not is_generic_mutation_analysis(
                gene, section["mutation_analysis"]
            )
            assert marker in section["mutation_analysis"]


def test_p3b_immune_rules_do_not_infer_direction_from_sequence_variant(tmp_path):
    variation_rows = []
    for index, gene in enumerate(
        ("PDCD1LG2", "B2M", "IFNGR1", "IFNGR2", "MLH1"), start=1
    ):
        variation_rows.append(
            {
                "Gene_Symbol": gene,
                "cHGVS": f"c.{index}A>G",
                "pHGVS_S": f"p.X{index}Y",
                "ExistInsmall358": 1,
                "ExistIn552": "Ⅱ类",
                "CLNSIG": "Pathogenic",
                "Function": "Missense",
                "Freq(%)": 10.0,
            }
        )
    source = tmp_path / "synthetic-p3b-immune-boundary.xlsx"
    source.touch()
    excel_data = ExcelDataSource(
        file_path=str(source),
        table_data={"Variations": variation_rows},
    )

    for panel_id in ("crc_301_msi", "crc_358_msi"):
        panel_config = load_panel_config(base_path=str(ROOT), panel_id=panel_id)
        positive_modes = {
            row["key"]: row["mode"] for row in panel_config.immune_positive_rows
        }
        negative_modes = {
            row["key"]: row["mode"] for row in panel_config.immune_negative_rows
        }
        assert positive_modes["PDCD1LG2"] == "non_sequence_biomarker"
        assert negative_modes["B2M"] == "confirmed_functional_loss"
        assert negative_modes["IFNGR12"] == "confirmed_functional_loss"

        grouped = build_immune_variants(
            excel_data,
            panel_config=panel_config,
        )
        positive_genes = {row["gene"] for row in grouped["positive"]}
        negative_genes = {row["gene"] for row in grouped["negative"]}
        assert "MLH1" in positive_genes
        assert "PDCD1LG2" not in positive_genes
        assert {"B2M", "IFNGR1", "IFNGR2"}.isdisjoint(negative_genes)

        report_data = ReportData()
        all_variants = [
            {
                "gene": row["Gene_Symbol"],
                "cHGVS": row["cHGVS"],
                "pHGVS": row["pHGVS_S"],
                "gene_class": row["ExistIn552"],
            }
            for row in variation_rows
        ]
        _build_nccn_and_immune_fields(
            report_data,
            all_variants,
            excel_data,
            panel_config=panel_config,
        )
        assert report_data.get_field("imm_pos_PDCD1LG2") == "未检出有害变异"
        assert report_data.get_field("imm_neg_B2M") == "未检出有害变异"
        assert report_data.get_field("imm_neg_IFNGR12") == "未检出有害变异"


def test_p3c_chromatin_transcription_fallbacks_are_shared_and_event_bounded():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3c_chromatin_transcription_fallbacks_20260721.yaml"
    )
    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {"approved_for_runtime": 26}
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]

    raw = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    assert raw["drug_sections"] == []
    expected_boundaries = {
        "ASXL1": "克隆性造血事件与结直肠肿瘤中的普通序列变异不能等同",
        "BRD4": "BRD4融合、扩增或过表达以及药物抑制实验",
        "CIC": "中枢神经系统肿瘤或肉瘤中的分子分类证据",
        "CREBBP": "乙酰转移酶结构域热点、明确功能缺失、融合或拷贝数事件",
        "DEK": "DEK融合、扩增或过表达与普通DEK序列变异不是同一事件",
        "EP300": "乙酰转移酶结构域事件、融合和拷贝数改变",
        "EP400": "普通错义或框内重复事件不能等同",
        "H3F3A": "H3F3A现行官方符号为H3-3A",
        "JUN": "JUN表达升高、上游信号激活、AP-1活性或特定融合",
        "KAT6A": "KAT6A融合、扩增、表达异常和明确功能缺失",
        "KDM6A": "经证实的KDM6A功能缺失、蛋白缺失或双等位事件",
        "KMT2A": "KMT2A重排或融合是特定血液肿瘤中的独立事件",
        "MED12": "特定肿瘤中的MED12复发热点、表达改变或复合物功能异常",
        "RARA": "不能套用急性早幼粒细胞白血病的诊断或治疗语境",
        "RBM10": "其他癌种中的复发事件也不能直接迁移到结直肠癌",
        "RUNX1": "RUNX1融合、血液肿瘤中的驱动变异和胚系易感变异",
        "SF3B1": "SF3B1复发热点具有位点和疾病背景依赖性",
        "SMARCA1": "明确ATP酶功能缺失、复合物异常",
        "SMARCA2": "蛋白缺失、明确失活或完整SWI/SNF复合物缺陷",
        "STAG1": "双等位事件或cohesin复合物缺陷",
        "STAG2": "经证实的STAG2功能缺失或蛋白缺失",
        "TET1": "TET1表达或甲基化改变、融合以及明确催化功能缺失",
        "U2AF1": "U2AF1复发热点及其剪接改变具有位点和疾病背景依赖性",
        "WHSC1": "WHSC1现行官方符号为NSD2",
        "WHSC1L1": "WHSC1L1现行官方符号为NSD3",
        "ZRSR2": "小剪接体功能缺陷、血液肿瘤相关性、胚系易感",
    }
    assert {row["gene"] for row in raw["gene_sections"]} == set(
        expected_boundaries
    )
    assert all(
        row["panels"] == ["crc_301_msi", "crc_358_msi"]
        for row in raw["gene_sections"]
    )

    for panel_id in ("crc_301_msi", "crc_358_msi"):
        provider = _panel_provider(panel_id)
        for gene, marker in expected_boundaries.items():
            section = _section(provider, gene, "c.999999A>G", "p.X999999Y")
            combined = f"{section['intro']}\n{section['mutation_analysis']}"
            assert not is_generic_mutation_analysis(
                gene, section["mutation_analysis"]
            )
            assert marker in combined, (panel_id, gene, combined)


def test_p3c_exact_corrections_remove_cross_event_prose_and_keep_kmt2a_exact():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3c_legacy_exact_corrections_20260721.yaml"
    )
    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {"approved_for_runtime": 4}
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]
    assert yaml.safe_load(overlay_path.read_text(encoding="utf-8"))[
        "drug_sections"
    ] == []

    expected = {
        ("DEK", "c.123G>C", "p.E41D"): (
            "数据库收录次数也不等于致病性或驱动性",
            ("DEK在多种实体瘤中过表达", "乳腺癌患者3年后的复发率"),
        ),
        ("EP400", "c.8175_8177dup", "p.Q2748dup"): (
            "框内单个氨基酸重复",
            ("膀胱癌和儿童淋巴细胞白血病中观测到",),
        ),
        ("JUN", "c.89del", "p.N30Tfs*5"): (
            "仍需结合转录本表达、无义介导降解或蛋白功能证据",
            ("有研究报道，人类致癌基因JUN",),
        ),
        ("ZRSR2", "c.1337G>A", "p.R446Q"): (
            "个案中的胚系变异，均不能证明本例肿瘤体细胞错义位点",
            ("MDS和CMML患者中不超过8%", "再生型障碍性贫血和结直肠癌患者"),
        ),
    }
    kmt2a_exact = {
        ("KMT2A", "c.3950_3954del", "p.K1317Sfs*7"): "p.K1317Sfs*7为移码变异",
        ("KMT2A", "c.8848A>G", "p.S2950G"): "p.S2950G为错义变异",
    }
    for panel_id in ("crc_301_msi", "crc_358_msi"):
        provider = _panel_provider(panel_id)
        for selector, (marker, forbidden) in expected.items():
            section = _section(provider, *selector)
            assert marker in section["mutation_analysis"], (selector, section)
            assert not any(
                token in section["mutation_analysis"] for token in forbidden
            ), (selector, section)
        for selector, marker in kmt2a_exact.items():
            section = _section(provider, *selector)
            assert marker in section["mutation_analysis"], (selector, section)


def test_p3d_histone_legacy_symbol_map_matches_current_ncbi_identities():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3d_histone_fallbacks_20260721.yaml"
    )
    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {"approved_for_runtime": 23}
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]

    raw = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    assert raw["drug_sections"] == []
    expected = {
        "HIST1H1B": ("H1-5", 3009, "H1.5"),
        "HIST1H1C": ("H1-2", 3006, "H1.2"),
        "HIST1H1D": ("H1-3", 3007, "H1.3"),
        "HIST1H1E": ("H1-4", 3008, "H1.4"),
        "HIST1H2AL": ("H2AC16", 8332, "H2A"),
        "HIST1H2AM": ("H2AC17", 8336, "H2A"),
        "HIST1H2BC": ("H2BC4", 8347, "H2B"),
        "HIST1H2BD": ("H2BC5", 3017, "H2B"),
        "HIST1H2BG": ("H2BC8", 8339, "H2B"),
        "HIST1H2BJ": ("H2BC11", 8970, "H2B"),
        "HIST1H2BK": ("H2BC12", 85236, "H2B"),
        "HIST1H2BO": ("H2BC17", 8348, "H2B"),
        "HIST1H3A": ("H3C1", 8350, "H3.1"),
        "HIST1H3C": ("H3C3", 8352, "H3.1"),
        "HIST1H3D": ("H3C4", 8351, "H3.1"),
        "HIST1H3E": ("H3C6", 8353, "H3.1"),
        "HIST1H3F": ("H3C7", 8968, "H3.1"),
        "HIST1H3G": ("H3C8", 8355, "H3.1"),
        "HIST1H3H": ("H3C10", 8357, "H3.1"),
        "HIST1H3I": ("H3C11", 8354, "H3.1"),
        "HIST1H3J": ("H3C12", 8356, "H3.1"),
        "HIST1H4I": ("H4C9", 8294, "H4"),
        "HIST3H3": ("H3-4", 8290, "H3.4"),
    }
    observed = {
        row["legacy_symbol"]: (
            row["current_symbol"],
            row["ncbi_gene_id"],
            row["encoded_histone"],
        )
        for row in raw["legacy_symbol_map"]
    }
    assert observed == expected
    assert len({value[0] for value in observed.values()}) == 23
    assert len({value[1] for value in observed.values()}) == 23

    gene_rows = {row["gene"]: row for row in raw["gene_sections"]}
    assert set(gene_rows) == set(expected)
    assert all(
        row["panels"] == ["crc_301_msi", "crc_358_msi"]
        for row in gene_rows.values()
    )
    for legacy_symbol, (_, gene_id, _) in expected.items():
        ncbi_refs = [
            source
            for source in gene_rows[legacy_symbol]["source_refs"]
            if source["type"] == "ncbi_gene"
        ]
        assert ncbi_refs == [
            {
                "type": "ncbi_gene",
                "id": f"NCBI-Gene-{gene_id}",
                "url": f"https://www.ncbi.nlm.nih.gov/gene/{gene_id}",
            }
        ]


def test_p3d_histone_fallbacks_block_cross_paralog_and_cross_event_inference():
    current_symbols = {
        "HIST1H1B": "H1-5",
        "HIST1H1C": "H1-2",
        "HIST1H1D": "H1-3",
        "HIST1H1E": "H1-4",
        "HIST1H2AL": "H2AC16",
        "HIST1H2AM": "H2AC17",
        "HIST1H2BC": "H2BC4",
        "HIST1H2BD": "H2BC5",
        "HIST1H2BG": "H2BC8",
        "HIST1H2BJ": "H2BC11",
        "HIST1H2BK": "H2BC12",
        "HIST1H2BO": "H2BC17",
        "HIST1H3A": "H3C1",
        "HIST1H3C": "H3C3",
        "HIST1H3D": "H3C4",
        "HIST1H3E": "H3C6",
        "HIST1H3F": "H3C7",
        "HIST1H3G": "H3C8",
        "HIST1H3H": "H3C10",
        "HIST1H3I": "H3C11",
        "HIST1H3J": "H3C12",
        "HIST1H4I": "H4C9",
        "HIST3H3": "H3-4",
    }
    h1_genes = {"HIST1H1B", "HIST1H1C", "HIST1H1D", "HIST1H1E"}
    h2a_genes = {"HIST1H2AL", "HIST1H2AM"}
    h2b_genes = {
        "HIST1H2BC",
        "HIST1H2BD",
        "HIST1H2BG",
        "HIST1H2BJ",
        "HIST1H2BK",
        "HIST1H2BO",
    }
    h3c_genes = {
        "HIST1H3A",
        "HIST1H3C",
        "HIST1H3D",
        "HIST1H3E",
        "HIST1H3F",
        "HIST1H3G",
        "HIST1H3H",
        "HIST1H3I",
        "HIST1H3J",
    }

    for panel_id in ("crc_301_msi", "crc_358_msi"):
        provider = _panel_provider(panel_id)
        for gene, current_symbol in current_symbols.items():
            section = _section(provider, gene, "c.999999A>G", "p.X999999Y")
            combined = f"{section['intro']}\n{section['mutation_analysis']}"
            assert not is_generic_mutation_analysis(
                gene, section["mutation_analysis"]
            )
            assert f"现行官方符号为{current_symbol}" in combined
            if gene in h1_genes:
                assert "滤泡性淋巴瘤中的复发突变证据" in combined
                assert "不能直接迁移到结直肠癌" in combined
            elif gene in h2a_genes:
                assert "H2A翻译后修饰、表达或核小体状态" in combined
                assert "位点证据不得跨基因拷贝转移" in combined
            elif gene in h2b_genes:
                assert "H2B翻译后修饰、表达或核小体状态" in combined
                assert "实际基因座" in combined or "跨基因拷贝" in combined
            elif gene in h3c_genes:
                if gene == "HIST1H3C":
                    assert "p.K27M在特定中枢神经系统弥漫性中线胶质瘤" in combined
                    assert "不能扩展到HIST1H3C其他位点、其他H3拷贝" in combined
                else:
                    assert "不得在不同H3基因拷贝之间转移" in combined
                    assert "不能套用中枢神经系统肿瘤分类" in combined
            elif gene == "HIST1H4I":
                assert "翻译后修饰状态与HIST1H4I序列变异不是同一事件" in combined
            else:
                assert gene == "HIST3H3"
                assert "H3.3热点或H3C家族的H3.1热点证据不能" in combined
                assert "实际基因座、蛋白亚型、残基和肿瘤类型" in combined


def test_p3e_shared_residual_identity_map_matches_current_ncbi_entities():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3e_shared_residual_fallbacks_20260721.yaml"
    )
    for panel_id in ("crc_301_msi", "crc_358_msi"):
        validation = load_and_validate_overlay(overlay_path, panel_id)
        assert validation["status"] == "PASS", validation["issues"]
        assert validation["gene"]["status_counts"] == {
            "approved_for_runtime": 47
        }
        assert validation["gene"]["secondary_review_complete_rows"] == validation[
            "gene"
        ]["total_rows"]

    raw = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    assert raw["drug_sections"] == []
    expected = {
        "AXIN1": ("AXIN1", 8312),
        "BCL2L11": ("BCL2L11", 10018),
        "CARD11": ("CARD11", 84433),
        "CCNE1": ("CCNE1", 898),
        "CD79B": ("CD79B", 974),
        "CDH11": ("CDH11", 1009),
        "CRKL": ("CRKL", 1399),
        "EPAS1": ("EPAS1", 2034),
        "ERRFI1": ("ERRFI1", 54206),
        "FGF19": ("FGF19", 9965),
        "FOXL2": ("FOXL2", 668),
        "GATA3": ("GATA3", 2625),
        "IRS1": ("IRS1", 3667),
        "JAK3": ("JAK3", 3718),
        "KDR": ("KDR", 3791),
        "KEAP1": ("KEAP1", 9817),
        "KEL": ("KEL", 3792),
        "LBP1B": ("UBP1", 7342),
        "MAP2K2": ("MAP2K2", 5605),
        "MAPK1": ("MAPK1", 5594),
        "MAPK3": ("MAPK3", 5595),
        "MDM4": ("MDM4", 4194),
        "MPL": ("MPL", 4352),
        "MYCN": ("MYCN", 4613),
        "PDGFRB": ("PDGFRB", 5159),
        "PIK3R1": ("PIK3R1", 5295),
        "PIM1": ("PIM1", 5292),
        "PPP2R2A": ("PPP2R2A", 5520),
        "PTCH2": ("PTCH2", 8643),
        "PTPN11": ("PTPN11", 5781),
        "RAD21": ("RAD21", 5885),
        "SERPINB3": ("SERPINB3", 6317),
        "SERPINB4": ("SERPINB4", 6318),
        "SHH": ("SHH", 6469),
        "SLIT2": ("SLIT2", 9353),
        "SRC": ("SRC", 6714),
        "SYK": ("SYK", 6850),
        "TEK": ("TEK", 7010),
        "TGFBR1": ("TGFBR1", 7046),
        "TNFAIP3": ("TNFAIP3", 7128),
        "TP63": ("TP63", 8626),
        "TSHR": ("TSHR", 7253),
        "YAP1": ("YAP1", 10413),
        "YES1": ("YES1", 7525),
        "ZNF217": ("ZNF217", 7764),
        "ZNF278": ("PATZ1", 23598),
        "ZNF521": ("ZNF521", 25925),
    }
    observed = {
        row["panel_symbol"]: (row["current_symbol"], row["ncbi_gene_id"])
        for row in raw["symbol_identity_map"]
    }
    assert observed == expected
    assert len({value[1] for value in observed.values()}) == 47

    gene_rows = {row["gene"]: row for row in raw["gene_sections"]}
    assert set(gene_rows) == set(expected)
    for panel_symbol, (_, gene_id) in expected.items():
        row = gene_rows[panel_symbol]
        assert row["panels"] == ["crc_301_msi", "crc_358_msi"]
        assert row["source_refs"] == [
            {
                "type": "ncbi_gene",
                "id": f"NCBI-Gene-{gene_id}",
                "url": f"https://www.ncbi.nlm.nih.gov/gene/{gene_id}",
            }
        ]

    identities = {row["panel_symbol"]: row for row in raw["symbol_identity_map"]}
    assert identities["LBP1B"]["identity_status"] == "legacy_isoform_alias"
    assert identities["ZNF278"]["identity_status"] == "legacy_symbol"


def test_p3e_shared_fallbacks_are_specific_drug_free_and_event_bounded():
    expected_markers = {
        "AXIN1": ("普通错义变异", "APC或CTNNB1"),
        "BCL2L11": ("表达量、缺失、多态性", "治疗敏感性或耐药性"),
        "CARD11": ("B/T细胞谱系", "不能迁移到结直肠癌"),
        "CCNE1": ("扩增或过表达", "编码区序列变异"),
        "CD79B": ("B细胞谱系", "结直肠癌"),
        "CRKL": ("扩增、蛋白磷酸化", "编码区序列变异"),
        "EPAS1": ("蛋白稳定、缺氧表达特征", "胚系红细胞增多症"),
        "FGF19": ("扩增、过表达", "编码区序列变异"),
        "JAK3": ("血液肿瘤中的激活事件", "胚系双等位失活"),
        "KDR": ("肿瘤血管生成", "KDR肿瘤序列变异"),
        "KEL": ("血型相关胚系等位基因", "肿瘤体细胞测序位点"),
        "LBP1B": ("UBP1", "不是独立现行基因符号"),
        "MAPK1": ("ERK2磷酸化", "编码区位点"),
        "MAPK3": ("不得直接套用MAPK1", "未知错义位点"),
        "MDM4": ("扩增、过表达或剪接改变", "错义变异"),
        "MPL": ("骨髓增殖性肿瘤", "结直肠癌"),
        "MYCN": ("扩增", "编码区序列变异"),
        "PDGFRB": ("融合/重排", "普通序列位点"),
        "PIK3R1": ("不同甚至相反", "不能直接套用PIK3CA"),
        "PTCH2": ("不得直接套用PTCH1", "第二击"),
        "RAD21": ("胚系cohesin", "肿瘤体细胞"),
        "SERPINB3": ("蛋白或血清表达", "编码区序列变异"),
        "SHH": ("胚系发育性变异", "肿瘤体细胞"),
        "SLIT2": ("启动子甲基化", "编码区序列变异"),
        "SRC": ("蛋白磷酸化", "编码区序列变异"),
        "SYK": ("血液肿瘤谱系", "结直肠癌"),
        "TEK": ("血管畸形", "结直肠癌"),
        "TGFBR1": ("胚系Loeys-Dietz", "肿瘤中的明确双等位失活"),
        "TNFAIP3": ("淋巴系统肿瘤", "胚系自身炎症性疾病"),
        "TP63": ("亚型表达、扩增、重排", "转录本"),
        "TSHR": ("甲状腺疾病", "结直肠癌"),
        "YAP1": ("核定位、蛋白表达、扩增或融合", "编码区序列变异"),
        "YES1": ("不能直接套用SRC", "普通编码区位点"),
        "ZNF217": ("扩增、过表达", "编码区序列变异"),
        "ZNF278": ("PATZ1", "融合事件"),
        "ZNF521": ("表达、重排", "普通序列位点"),
    }
    overlay = yaml.safe_load(
        (
            ROOT
            / "panels/crc_358_msi/rules/"
            "reviewed_part3_p3e_shared_residual_fallbacks_20260721.yaml"
        ).read_text(encoding="utf-8")
    )
    expected_genes = {row["gene"] for row in overlay["gene_sections"]}

    for panel_id in ("crc_301_msi", "crc_358_msi"):
        provider = _panel_provider(panel_id)
        for gene in expected_genes:
            section = _section(provider, gene, "c.999999A>G", "p.X999999Y")
            assert section["intro"], gene
            assert not is_generic_mutation_analysis(
                gene, section["mutation_analysis"]
            ), (gene, section)
        for gene, markers in expected_markers.items():
            section = _section(provider, gene, "c.999999A>G", "p.X999999Y")
            combined = f"{section['intro']}\n{section['mutation_analysis']}"
            assert all(marker in combined for marker in markers), (gene, combined)


def test_p3e_exact_corrections_remove_static_counts_and_cross_event_inference():
    correction_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3e_legacy_exact_corrections_20260721.yaml"
    )
    for panel_id in ("crc_301_msi", "crc_358_msi"):
        validation = load_and_validate_overlay(correction_path, panel_id)
        assert validation["status"] == "PASS", validation["issues"]
        assert validation["gene"]["status_counts"] == {
            "approved_for_runtime": 7
        }

    expected = {
        ("CARD11", "c.2059G>A", "p.A687T"): "缺少足够的位点级功能实验",
        ("CDH11", "c.1460A>C", "p.K487T"): "静态数据库收录次数",
        ("MDM4", "c.1127C>T", "p.A376V"): "扩增、过表达或剪接异常",
        ("RAD21", "c.395A>G", "p.Q132R"): "不能单独证明功能正常或受损",
        ("SYK", "c.851G>T", "p.W284L"): "数据库未收录均不能证明良性或致病",
        ("TEK", "c.1149G>T", "p.M383I"): "血管畸形中的特定TEK激活位点",
        ("TSHR", "c.1116C>A", "p.N372K"): "甲状腺疾病中的特定激活/失活位点",
    }
    forbidden = ("COSMIC数据库中记载", "与大肠癌相关的记录")
    for panel_id in ("crc_301_msi", "crc_358_msi"):
        provider = _panel_provider(panel_id)
        for selector, marker in expected.items():
            section = _section(provider, *selector)
            assert marker in section["mutation_analysis"], (selector, section)
            assert not any(
                token in section["mutation_analysis"] for token in forbidden
            ), (selector, section)


def test_p3f1_crc358_signaling_identity_map_is_current_scoped_and_drug_free():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3f1_signaling_fallbacks_20260721.yaml"
    )
    correction_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3f1_legacy_exact_corrections_20260721.yaml"
    )
    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {
        "approved_for_runtime": 31
    }
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]

    raw = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    assert raw["drug_sections"] == []
    expected = {
        "ABL2": 27,
        "ACVR1B": 91,
        "AXL": 558,
        "DDR1": 780,
        "EPHA3": 2042,
        "EPHB1": 2047,
        "FLT4": 2324,
        "INPP4A": 3631,
        "INPPL1": 3636,
        "INSR": 3643,
        "LATS1": 9113,
        "LATS2": 26524,
        "MAP2K7": 5609,
        "MAP3K13": 9175,
        "MAP3K4": 4216,
        "MAP3K6": 9064,
        "PASK": 23178,
        "PIK3CD": 5293,
        "PIK3CG": 5294,
        "PREX2": 80243,
        "PTPRB": 5787,
        "PTPRC": 5788,
        "PTPRD": 5789,
        "PTPRK": 5796,
        "PTPRS": 5802,
        "RICTOR": 253260,
        "RPTOR": 57521,
        "SMAD3": 4088,
        "SOS1": 6654,
        "VAV1": 7409,
        "VAV2": 7410,
    }
    observed = {
        row["panel_symbol"]: row["ncbi_gene_id"]
        for row in raw["symbol_identity_map"]
    }
    assert observed == expected
    assert all(
        row["current_symbol"] == row["panel_symbol"]
        for row in raw["symbol_identity_map"]
    )

    gene_rows = {row["gene"]: row for row in raw["gene_sections"]}
    assert set(gene_rows) == set(expected)
    for gene, gene_id in expected.items():
        row = gene_rows[gene]
        assert row["panels"] == ["crc_358_msi"]
        assert {
            "type": "ncbi_gene",
            "id": f"NCBI-Gene-{gene_id}",
            "url": f"https://www.ncbi.nlm.nih.gov/gene/{gene_id}",
        } in row["source_refs"]

    crc358_paths = {
        Path(path)
        for path in ReportGenerator._resolve_panel_reviewed_part3_overlays(
            load_panel_package("crc_358_msi", project_root=ROOT)
        )
    }
    crc301_paths = {
        Path(path)
        for path in ReportGenerator._resolve_panel_reviewed_part3_overlays(
            load_panel_package("crc_301_msi", project_root=ROOT)
        )
    }
    assert overlay_path in crc358_paths
    assert correction_path in crc358_paths
    assert overlay_path not in crc301_paths
    assert correction_path not in crc301_paths


def test_p3f1_crc358_signaling_fallbacks_are_specific_and_event_bounded():
    overlay = yaml.safe_load(
        (
            ROOT
            / "panels/crc_358_msi/rules/"
            "reviewed_part3_p3f1_signaling_fallbacks_20260721.yaml"
        ).read_text(encoding="utf-8")
    )
    expected_genes = {row["gene"] for row in overlay["gene_sections"]}
    expected_markers = {
        "ABL2": ("ETV6-ABL2融合", "ABL1的热点、融合和用药证据"),
        "ACVR1B": ("双等位改变", "ACVR2A、TGFBR1"),
        "AXL": ("GAS6配体依赖", "蛋白磷酸化或扩增"),
        "DDR1": ("胶原诱导磷酸化", "激酶结构域本身不能证明"),
        "EPHA3": ("拷贝数和编码区位点", "不能跨成员转移"),
        "EPHB1": ("甲基化变化", "不能仅凭结构域位置"),
        "FLT4": ("胚系淋巴水肿", "肿瘤体细胞位点"),
        "INPP4A": ("不能直接等同于PTEN或INPPL1", "催化活性"),
        "INPPL1": ("胚系疾病和肿瘤体细胞位点", "乳腺癌表达研究"),
        "INSR": ("胚系严重胰岛素抵抗", "肿瘤体细胞位点"),
        "LATS1": ("不能直接套用LATS2", "YAP激活"),
        "LATS2": ("LATS1证据也不能跨成员替代", "Hippo通路角色"),
        "MAP2K7": ("MAP2K1/2的激活位点", "JNK通路功能资料"),
        "MAP3K13": ("磷酸化状态", "其他MAP3K成员证据"),
        "MAP3K4": ("其他MAP3K成员", "MAPK通路激活"),
        "MAP3K6": ("MAP3K5等同家族证据", "血管生成改变"),
        "PASK": ("胚系代谢相关观察", "肿瘤体细胞序列位点"),
        "PIK3CD": ("胚系PIK3CD", "PIK3CA热点"),
        "PIK3CG": ("免疫微环境活性", "PIK3CA或PIK3CD"),
        "PREX2": ("截短、结构域破坏", "与PTEN相互作用研究"),
        "PTPRB": ("结构域外位置也不能证明", "血管生成"),
        "PTPRC": ("免疫细胞谱系属性", "样本来源、肿瘤纯度"),
        "PTPRD": ("纯合缺失、截短", "表达或甲基化变化"),
        "PTPRK": ("融合、缺失、表达改变", "PTPRD或PTPRS"),
        "PTPRS": ("拷贝数或表达变化", "其他受体型PTP成员"),
        "RICTOR": ("mTORC2", "RPTOR/mTORC1"),
        "RPTOR": ("mTORC1", "RICTOR/mTORC2"),
        "SMAD3": ("胚系动脉瘤-骨关节炎", "SMAD4或其他SMAD成员"),
        "SOS1": ("胚系Noonan", "SOS1位于RAS上游"),
        "VAV1": ("血液肿瘤中的VAV1融合", "VAV2证据"),
        "VAV2": ("VAV1的造血谱系", "GEF功能证据"),
    }
    assert set(expected_markers) == expected_genes

    provider = _crc358_provider()
    for gene in expected_genes:
        section = _section(provider, gene, "c.999999A>G", "p.X999999Y")
        assert section["intro"], gene
        assert not is_generic_mutation_analysis(
            gene, section["mutation_analysis"]
        ), (gene, section)
        combined = f"{section['intro']}\n{section['mutation_analysis']}"
        assert all(marker in combined for marker in expected_markers[gene]), (
            gene,
            combined,
        )


def test_p3f1_exact_corrections_remove_static_counts_and_domain_inference():
    correction_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3f1_legacy_exact_corrections_20260721.yaml"
    )
    validation = load_and_validate_overlay(correction_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {
        "approved_for_runtime": 4
    }

    expected = {
        ("DDR1", "c.1973G>A", "p.R658Q"): "胃癌EMT或肾脏表达研究均不能证明",
        ("EPHB1", "c.2647C>T", "p.R883W"): "静态数据库次数及位于已注释激酶/SAM区域之外",
        ("PTPRB", "c.4700G>T", "p.G1567V"): "数据库未收录或位于主要催化区域之外",
        ("PTPRS", "c.3073G>A", "p.V1025I"): "其他成员事件不能证明本位点",
    }
    forbidden = (
        "COSMIC数据库中记载1次",
        "COSMIC数据库中未记载",
        "胃癌细胞中",
        "肾癌组织",
    )
    provider = _crc358_provider()
    for selector, marker in expected.items():
        section = _section(provider, *selector)
        assert marker in section["mutation_analysis"], (selector, section)
        assert not any(
            token in section["mutation_analysis"] for token in forbidden
        ), (selector, section)


def test_gene_level_replace_fields_requires_explicit_supersedes(tmp_path):
    base = tmp_path / "base.yaml"
    attempted = tmp_path / "attempted.yaml"
    correction = tmp_path / "correction.yaml"
    base.write_text(
        """
gene_sections:
  - gene: DEMO
    intro: first reviewed intro
    mutation_analysis: first reviewed analysis
drug_sections: []
""".strip(),
        encoding="utf-8",
    )
    attempted.write_text(
        """
gene_sections:
  - gene: DEMO
    intro: untracked replacement
    mutation_analysis: untracked replacement analysis
    replace_fields: [intro, mutation_analysis]
drug_sections: []
""".strip(),
        encoding="utf-8",
    )
    correction.write_text(
        """
gene_sections:
  - gene: DEMO
    intro: governed replacement
    mutation_analysis: governed replacement analysis
    replace_fields: [intro, mutation_analysis]
    supersedes: first_reviewed_demo_row
drug_sections: []
""".strip(),
        encoding="utf-8",
    )

    def provider(paths):
        value = GeneKnowledgeProvider(
            {
                "enabled": True,
                "gene_knowledge_db": {
                    "enabled": True,
                    "path": "missing.xlsx",
                    "reviewed_part3_overlay_paths": [str(path) for path in paths],
                },
            }
        )
        assert value.load(base_path=str(ROOT))
        return value

    blocked = _section(
        provider([base, attempted]), "DEMO", "c.1A>G", "p.M1V"
    )
    replaced = _section(
        provider([base, correction]), "DEMO", "c.1A>G", "p.M1V"
    )
    assert blocked["intro"] == "first reviewed intro"
    assert blocked["mutation_analysis"] == "first reviewed analysis"
    assert replaced["intro"] == "governed replacement"
    assert replaced["mutation_analysis"] == "governed replacement analysis"


def test_p3f2_crc358_chromatin_identity_scope_and_collision_contract():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3f2_chromatin_fallbacks_20260721.yaml"
    )
    correction_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3f2_legacy_exact_corrections_20260721.yaml"
    )
    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {
        "approved_for_runtime": 24
    }
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]

    raw = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    assert raw["drug_sections"] == []
    expected = {
        "ARID1B": 57492,
        "ARID2": 196528,
        "ARID4A": 5926,
        "ASXL2": 55252,
        "BACH2": 60468,
        "BCL11A": 53335,
        "BCORL1": 63035,
        "CHD4": 1108,
        "DNMT1": 1786,
        "DNMT3B": 1789,
        "HDAC4": 9759,
        "HDAC7": 51564,
        "HDAC9": 9734,
        "HIRA": 7290,
        "KDM2B": 84678,
        "KDM4C": 23081,
        "KDM5A": 5927,
        "KMT2B": 9757,
        "NCOR1": 9611,
        "NCOR2": 9612,
        "NSD1": 64324,
        "SETD5": 55209,
        "TET2": 54790,
        "TRRAP": 8295,
    }
    observed = {
        row["panel_symbol"]: row["ncbi_gene_id"]
        for row in raw["symbol_identity_map"]
    }
    assert observed == expected
    assert all(
        row["current_symbol"] == row["panel_symbol"]
        for row in raw["symbol_identity_map"]
    )

    gene_rows = {row["gene"]: row for row in raw["gene_sections"]}
    assert set(gene_rows) == set(expected)
    assert all(row["panels"] == ["crc_358_msi"] for row in gene_rows.values())
    assert all(
        set(row["replace_fields"]) == {"intro", "mutation_analysis"}
        and row["supersedes"]
        for row in gene_rows.values()
    )
    for gene, gene_id in expected.items():
        assert {
            "type": "ncbi_gene",
            "id": f"NCBI-Gene-{gene_id}",
            "url": f"https://www.ncbi.nlm.nih.gov/gene/{gene_id}",
        } in gene_rows[gene]["source_refs"]

    collisions = raw["identity_collision_contract"]
    assert collisions == [
        {
            "canonical_symbol": "NSD2",
            "ncbi_gene_id": 7468,
            "historical_aliases": ["WHSC1"],
            "declared_panel_keys": ["NSD2", "WHSC1"],
            "identity_status": "confirmed_same_biological_gene",
            "runtime_normalization_status": (
                "active_panel_scoped_alias_lookup_and_variant_dedup"
            ),
            "runtime_mapping": {"WHSC1": "NSD2"},
            "coverage_accounting": (
                "retain_both_declared_keys_but_resolve_both_to_one_biological_identity"
            ),
            "report_display_policy": (
                "preserve_first_input_symbol_and_deduplicate_identical_variant_identity"
            ),
            "source_refs": [
                {
                    "type": "ncbi_gene",
                    "id": "NCBI-Gene-7468",
                    "url": "https://www.ncbi.nlm.nih.gov/gene/7468",
                }
            ],
        }
    ]
    assert "NSD2" not in gene_rows

    package = load_panel_package("crc_358_msi", project_root=ROOT)
    assert package.raw["gene_symbol_aliases"] == {"WHSC1": "NSD2"}

    crc358_paths = {
        Path(path)
        for path in ReportGenerator._resolve_panel_reviewed_part3_overlays(
            load_panel_package("crc_358_msi", project_root=ROOT)
        )
    }
    crc301_paths = {
        Path(path)
        for path in ReportGenerator._resolve_panel_reviewed_part3_overlays(
            load_panel_package("crc_301_msi", project_root=ROOT)
        )
    }
    assert overlay_path in crc358_paths
    assert correction_path in crc358_paths
    assert overlay_path not in crc301_paths
    assert correction_path not in crc301_paths


def test_p3f2_crc358_chromatin_fallbacks_are_specific_and_event_bounded():
    overlay = yaml.safe_load(
        (
            ROOT
            / "panels/crc_358_msi/rules/"
            "reviewed_part3_p3f2_chromatin_fallbacks_20260721.yaml"
        ).read_text(encoding="utf-8")
    )
    expected_genes = {row["gene"] for row in overlay["gene_sections"]}
    expected_markers = {
        "ARID1B": ("双等位失活", "MSI结直肠癌中ARID家族"),
        "ARID2": ("PBAF复合物缺陷", "肝细胞癌"),
        "ARID4A": ("pRB/HDAC复合物", "MSI结直肠癌"),
        "ASXL2": ("ASXL1的髓系肿瘤", "不能跨成员套用"),
        "BACH2": ("胚系免疫缺陷", "免疫细胞来源"),
        "BCL11A": ("B细胞恶性肿瘤中的易位", "BCL11B"),
        "BCORL1": ("X染色体等位状态", "BCOR事件"),
        "CHD4": ("NuRD复合物缺陷", "浆液性子宫内膜"),
        "DNMT1": ("全局/位点甲基化表型", "胚系神经系统"),
        "DNMT3B": ("胚系ICF综合征", "miR-124"),
        "HDAC4": ("核质转位", "泛HDAC抑制剂"),
        "HDAC7": ("蛋白磷酸化、核质定位", "其他HDAC成员"),
        "HDAC9": ("小鼠Treg/结肠炎", "HDAC抑制剂"),
        "HIRA": ("H3.3装配", "DiGeorge"),
        "KDM2B": ("胚系神经发育", "KDM2A"),
        "KDM4C": ("扩增、表达升高", "KDM4A/B/D"),
        "KDM5A": ("药物耐受细胞状态", "其他KDM5成员"),
        "KMT2B": ("胚系肌张力障碍", "KMT2A融合"),
        "NCOR1": ("核受体配体状态", "NCOR2"),
        "NCOR2": ("共抑制复合物状态", "NCOR1"),
        "NSD1": ("Sotos/Weaver综合征", "NUP98-NSD1融合"),
        "SETD5": ("确切酶学功能尚未完全确定", "仅凭SET区域名称"),
        "TET2": ("克隆性造血", "配对血液"),
        "TRRAP": ("PIKK家族归属", "PI3K或mTOR"),
    }
    assert set(expected_markers) == expected_genes

    provider = _crc358_provider()
    for gene in expected_genes:
        section = _section(provider, gene, "c.999999A>G", "p.X999999Y")
        assert section["intro"], gene
        assert not is_generic_mutation_analysis(
            gene, section["mutation_analysis"]
        ), (gene, section)
        combined = f"{section['intro']}\n{section['mutation_analysis']}"
        assert all(marker in combined for marker in expected_markers[gene]), (
            gene,
            combined,
        )

    corrected_intro_forbidden = {
        "ARID1B": "成骨细胞系中的研究",
        "DNMT3B": "机体对化疗的敏感性",
        "HDAC7": "包括胰腺癌",
        "HDAC9": "儿童急性淋巴细胞白血病预后",
        "KMT2B": "GeneCards",
    }
    for gene, token in corrected_intro_forbidden.items():
        section = _section(provider, gene, "c.999999A>G", "p.X999999Y")
        assert token not in section["intro"], (gene, section["intro"])


def test_p3f2_exact_corrections_remove_database_and_cross_event_inference():
    correction_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3f2_legacy_exact_corrections_20260721.yaml"
    )
    validation = load_and_validate_overlay(correction_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {
        "approved_for_runtime": 6
    }

    expected = {
        ("ARID1B", "c.618_641del", "p.Q207_Q214del"): "不能证明本位点导致双等位失活",
        ("DNMT3B", "c.1526delinsTGGA", "p.G509delinsVD"): (
            "miR-124下调DNMT3B表达的细胞研究也不能证明"
        ),
        ("HDAC7", "c.615G>T", "p.K205N"): "HDAC7表达、磷酸化、核质定位",
        ("HDAC9", "c.829_830delinsGT", "p.S277V"): (
            "小鼠Treg功能和结肠炎的研究不是结直肠肿瘤位点证据"
        ),
        ("KMT2B", "c.3164G>A", "p.R1055Q"): "基因在其他肿瘤中的泛化描述均不能证明",
        ("KMT2B", "c.1681C>T", "p.P561S"): "其他KMT2家族事件或肿瘤测序中的基因级检出",
    }
    forbidden = (
        "该突变在COSMIC数据库中暂无记载",
        "外显子组测序研究揭示",
        "抑制结直肠癌细胞的增殖、迁移和侵袭",
        "有研究报道，抑制HDAC9的表达可增强",
        "KMT2B基因可能参与人类癌症的发生发展",
        "本位点按染色质修饰通路相关长尾变异解释",
    )
    provider = _crc358_provider()
    for selector, marker in expected.items():
        section = _section(provider, *selector)
        assert marker in section["mutation_analysis"], (selector, section)
        assert not any(
            token in section["mutation_analysis"] for token in forbidden
        ), (selector, section)
    p561s = _section(provider, "KMT2B", "c.1681C>T", "p.P561S")
    assert "可能与肿瘤发生发展相关" not in p561s["intro"]


def test_p3f3_lineage_transcription_overlay_identity_governance_and_scope():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3f3_lineage_transcription_fallbacks_20260721.yaml"
    )
    correction_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3f3_legacy_exact_corrections_20260721.yaml"
    )
    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {
        "approved_for_runtime": 18
    }
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]

    raw = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    assert raw["drug_sections"] == []
    expected = {
        "CUX1": 1523,
        "ETV1": 2115,
        "FLI1": 2313,
        "FOXP1": 27086,
        "FUBP1": 8880,
        "MAX": 4149,
        "MGA": 23269,
        "MYB": 4602,
        "PAX5": 5079,
        "RUNX1T1": 862,
        "RUNX2": 860,
        "SOX9": 6662,
        "TCF12": 6938,
        "TCF3": 6929,
        "TCF4": 6925,
        "TLE3": 7090,
        "TLE4": 7091,
        "NCOA4": 8031,
    }
    observed = {
        row["panel_symbol"]: row["ncbi_gene_id"]
        for row in raw["symbol_identity_map"]
    }
    assert observed == expected
    assert all(
        row["current_symbol"] == row["panel_symbol"]
        for row in raw["symbol_identity_map"]
    )

    gene_rows = {row["gene"]: row for row in raw["gene_sections"]}
    assert set(gene_rows) == set(expected)
    assert all(row["panels"] == ["crc_358_msi"] for row in gene_rows.values())
    assert all(
        set(row["replace_fields"]) == {"intro", "mutation_analysis"}
        and row["supersedes"]
        for row in gene_rows.values()
    )
    for gene, gene_id in expected.items():
        assert {
            "type": "ncbi_gene",
            "id": f"NCBI-Gene-{gene_id}",
            "url": f"https://www.ncbi.nlm.nih.gov/gene/{gene_id}",
        } in gene_rows[gene]["source_refs"]

    assert raw["symbol_disambiguation_contract"] == [
        {
            "panel_symbol": "TCF4",
            "canonical_symbol": "TCF4",
            "canonical_ncbi_gene_id": 6925,
            "identity_status": "confirmed_official_symbol",
            "conflicting_historical_alias": {
                "symbol": "TCF4",
                "canonical_symbol": "TCF7L2",
                "ncbi_gene_id": 6934,
            },
            "runtime_policy": (
                "Resolve the CRC358 panel key TCF4 only to NCBI Gene 6925. "
                "Do not import TCF7L2/Wnt evidence merely because TCF7L2 has "
                "historically been called TCF-4 or TCF4."
            ),
            "source_refs": [
                {
                    "type": "ncbi_gene",
                    "id": "NCBI-Gene-6925",
                    "url": "https://www.ncbi.nlm.nih.gov/gene/6925",
                },
                {
                    "type": "ncbi_gene",
                    "id": "NCBI-Gene-6934",
                    "url": "https://www.ncbi.nlm.nih.gov/gene/6934",
                },
            ],
        }
    ]

    crc358_paths = {
        Path(path)
        for path in ReportGenerator._resolve_panel_reviewed_part3_overlays(
            load_panel_package("crc_358_msi", project_root=ROOT)
        )
    }
    crc301_paths = {
        Path(path)
        for path in ReportGenerator._resolve_panel_reviewed_part3_overlays(
            load_panel_package("crc_301_msi", project_root=ROOT)
        )
    }
    assert overlay_path in crc358_paths
    assert correction_path in crc358_paths
    assert overlay_path not in crc301_paths
    assert correction_path not in crc301_paths


def test_p3f3_crc358_lineage_transcription_fallbacks_are_event_bounded():
    expected_markers = {
        "CUX1": ("造血系统", "不能自动判定抑癌功能缺失"),
        "ETV1": ("ETV1融合/重排", "过表达"),
        "FLI1": ("EWSR1-FLI1融合", "造血来源"),
        "FOXP1": ("异常异构体", "FOXP2/FOXP3"),
        "FUBP1": ("少突胶质瘤", "MYC通路关系"),
        "MAX": ("胚系MAX变异", "MYC扩增"),
        "MGA": ("截短/双等位失活", "MAX、MYC或MXD"),
        "MYB": ("白血病/淋巴瘤", "造血克隆"),
        "PAX5": ("B细胞白血病/淋巴瘤", "结直肠上皮肿瘤"),
        "RUNX1T1": ("RUNX1-RUNX1T1融合", "染色体重排"),
        "RUNX2": ("锁骨颅骨发育不良", "其他癌种的表达"),
        "SOX9": ("发育综合征", "过表达"),
        "TCF12": ("CRC中的过表达/转移", "少突胶质瘤"),
        "TCF3": ("TCF3-PBX1", "淋巴发育缺陷"),
        "TCF4": ("NCBI Gene 6925", "TCF7L2（Gene 6934）"),
        "TLE3": ("卵巢癌", "序列变异证据"),
        "TLE4": ("TLE3表达标志物", "Wnt通路关联"),
        "NCOA4": ("RET-NCOA4融合", "染色体重排"),
    }
    provider = _crc358_provider()
    for gene, markers in expected_markers.items():
        section = _section(provider, gene, "c.999999A>G", "p.X999999Y")
        assert section["intro"], gene
        assert not is_generic_mutation_analysis(
            gene, section["mutation_analysis"]
        ), (gene, section)
        combined = f"{section['intro']}\n{section['mutation_analysis']}"
        assert all(marker in combined for marker in markers), (gene, combined)

    # The addition is CRC358-only; the same sentinel remains a base fallback
    # when directly constructing the CRC301 provider.
    crc301_cux1 = _section(
        _panel_provider("crc_301_msi"),
        "CUX1",
        "c.999999A>G",
        "p.X999999Y",
    )
    assert is_generic_mutation_analysis(
        "CUX1", crc301_cux1["mutation_analysis"]
    )


def test_p3f3_exact_corrections_remove_domain_and_cross_context_inference():
    correction_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3f3_legacy_exact_corrections_20260721.yaml"
    )
    validation = load_and_validate_overlay(correction_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {
        "approved_for_runtime": 2
    }

    provider = _crc358_provider()
    runx2 = _section(provider, "RUNX2", "c.211_213del", "p.Q71del")
    assert "框内单个氨基酸缺失" in runx2["mutation_analysis"]
    assert "肺癌中的表达/凋亡研究" in runx2["mutation_analysis"]
    assert "可能对蛋白功能有一定的影响" not in runx2["mutation_analysis"]
    assert "Runt DNA结合区域" in runx2["intro"]

    tcf12 = _section(provider, "TCF12", "c.572C>T", "p.P191L")
    assert "位于主要bHLH区域之外并不能证明无功能影响" in tcf12[
        "mutation_analysis"
    ]
    assert "CRC中的TCF12过表达/转移相关观察" in tcf12[
        "mutation_analysis"
    ]
    assert "具体的临床意义未明" not in tcf12["mutation_analysis"]


def test_p3f4_structural_residual_overlay_identity_governance_and_scope():
    overlay_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3f4_structural_residual_fallbacks_20260721.yaml"
    )
    correction_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3f4_legacy_exact_corrections_20260721.yaml"
    )
    validation = load_and_validate_overlay(overlay_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {
        "approved_for_runtime": 11
    }
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]

    correction_validation = load_and_validate_overlay(
        correction_path, "crc_358_msi"
    )
    assert correction_validation["status"] == "PASS", correction_validation[
        "issues"
    ]
    assert correction_validation["gene"]["status_counts"] == {
        "approved_for_runtime": 4
    }

    raw = yaml.safe_load(overlay_path.read_text(encoding="utf-8"))
    assert raw["drug_sections"] == []
    expected = {
        "ATXN2": 6311,
        "CLTCL1": 8218,
        "DNM2": 1785,
        "FAT3": 120114,
        "MAGI2": 9863,
        "MGAM": 8972,
        "MUC1": 4582,
        "MYH11": 4629,
        "PCLO": 27445,
        "PDE4DIP": 9659,
        "PRSS1": 5644,
    }
    observed = {
        row["panel_symbol"]: row["ncbi_gene_id"]
        for row in raw["symbol_identity_map"]
    }
    assert observed == expected
    assert all(
        row["current_symbol"] == row["panel_symbol"]
        for row in raw["symbol_identity_map"]
    )

    gene_rows = {row["gene"]: row for row in raw["gene_sections"]}
    assert set(gene_rows) == set(expected)
    assert all(row["panels"] == ["crc_358_msi"] for row in gene_rows.values())
    assert all(
        set(row["replace_fields"]) == {"intro", "mutation_analysis"}
        and row["supersedes"]
        for row in gene_rows.values()
    )
    for gene, gene_id in expected.items():
        assert {
            "type": "ncbi_gene",
            "id": f"NCBI-Gene-{gene_id}",
            "url": f"https://www.ncbi.nlm.nih.gov/gene/{gene_id}",
        } in gene_rows[gene]["source_refs"]
    assert "4557个氨基酸" in gene_rows["FAT3"]["fixed_domain_text"]
    assert "Myosin motor结构域" in gene_rows["MYH11"]["fixed_domain_text"]

    crc358_paths = {
        Path(path)
        for path in ReportGenerator._resolve_panel_reviewed_part3_overlays(
            load_panel_package("crc_358_msi", project_root=ROOT)
        )
    }
    crc301_paths = {
        Path(path)
        for path in ReportGenerator._resolve_panel_reviewed_part3_overlays(
            load_panel_package("crc_301_msi", project_root=ROOT)
        )
    }
    assert overlay_path in crc358_paths
    assert correction_path in crc358_paths
    assert overlay_path not in crc301_paths
    assert correction_path not in crc301_paths


def test_p3f4_crc358_structural_residual_fallbacks_are_event_bounded():
    expected_markers = {
        "ATXN2": ("CAG/多聚谷氨酰胺重复扩增", "普通SNV"),
        "CLTCL1": ("染色体异常", "CLTC"),
        "DNM2": ("胚系变异相关神经肌肉疾病", "DNM1、DNM3"),
        "FAT3": ("双等位失活", "FAT1、FAT2、FAT4"),
        "MAGI2": ("缺失、重排、融合", "MAGI1、MAGI3"),
        "MGAM": ("消化酶活性", "蔗糖酶-异麦芽糖酶"),
        "MUC1": ("串联重复拷贝数", "抗原相关治疗"),
        "MYH11": ("CBFB::MYH11", "胚系变异"),
        "PCLO": ("编码区较长", "结构域内外位置"),
        "PDE4DIP": ("髓系/嗜酸粒细胞疾病", "PDE4D酶活"),
        "PRSS1": ("遗传性胰腺炎", "PRSS2"),
    }
    provider = _crc358_provider()
    for gene, markers in expected_markers.items():
        section = _section(provider, gene, "c.999999A>G", "p.X999999Y")
        assert section["intro"], gene
        assert not is_generic_mutation_analysis(
            gene, section["mutation_analysis"]
        ), (gene, section)
        combined = f"{section['intro']}\n{section['mutation_analysis']}"
        assert all(marker in combined for marker in markers), (gene, combined)

    crc301_atxn2 = _section(
        _panel_provider("crc_301_msi"),
        "ATXN2",
        "c.999999A>G",
        "p.X999999Y",
    )
    assert is_generic_mutation_analysis(
        "ATXN2", crc301_atxn2["mutation_analysis"]
    )


def test_p3f4_exact_corrections_remove_domain_cross_context_and_row_leakage():
    correction_path = (
        ROOT
        / "panels/crc_358_msi/rules/"
        "reviewed_part3_p3f4_legacy_exact_corrections_20260721.yaml"
    )
    validation = load_and_validate_overlay(correction_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["status_counts"] == {
        "approved_for_runtime": 4
    }

    provider = _crc358_provider()
    fat3_q1541l = _section(provider, "FAT3", "c.4622A>T", "p.Q1541L")
    assert "Cadherin 14结构域" in fat3_q1541l["mutation_analysis"]
    assert "4557个氨基酸" in fat3_q1541l["fixed_domain_text"]
    assert "4589个氨基酸" not in fat3_q1541l["mutation_analysis"]
    assert "可能对蛋白功能有一定的影响" not in fat3_q1541l[
        "mutation_analysis"
    ]

    fat3_g3981s = _section(provider, "FAT3", "c.11941G>A", "p.G3981S")
    assert "Laminin G-like结构域" in fat3_g3981s["mutation_analysis"]
    assert "按细胞黏附/迁移相关长尾变异解释" not in fat3_g3981s[
        "mutation_analysis"
    ]

    myh11 = _section(provider, "MYH11", "c.3089del", "p.N1030Tfs*6")
    assert "无义介导降解" in myh11["mutation_analysis"]
    assert "Myosin motor结构域" in myh11["fixed_domain_text"]
    assert "PALB2" not in myh11["mutation_analysis"]
    assert "可能与疾病的发生发展相关" not in myh11["mutation_analysis"]
    assert "2个PFAM结构域" not in myh11["fixed_domain_text"]

    pclo = _section(provider, "PCLO", "c.11722C>A", "p.H3908N")
    assert "结构域外定位不能证明无功能影响" in pclo["mutation_analysis"]
    assert "不能形成预后或治疗结论" in pclo["mutation_analysis"]


def test_crc358_nsd2_whsc1_alias_lookup_and_exact_variant_dedup():
    provider = _crc358_provider()
    nsd2 = _section(provider, "NSD2", "c.100A>G", "p.K34E")
    whsc1 = _section(provider, "WHSC1", "c.100A>G", "p.K34E")

    assert nsd2["gene"] == "NSD2"
    assert nsd2["header"].startswith("NSD2：")
    assert "WHSC1现行官方符号为NSD2" in nsd2["intro"]
    assert "特定易位、扩增、过表达或复发热点" in nsd2[
        "mutation_analysis"
    ]
    assert not is_generic_mutation_analysis("NSD2", nsd2["mutation_analysis"])
    assert whsc1["gene"] == "WHSC1"
    assert whsc1["header"].startswith("WHSC1：")

    nsd2_row = {
        "gene": "NSD2",
        "cHGVS": "c.100A>G",
        "pHGVS": "p.K34E",
        "frequency": "12.3",
    }
    whsc1_row = {**nsd2_row, "gene": "WHSC1", "frequency": "9.8"}
    assert provider.variant_identity_key(nsd2_row) == provider.variant_identity_key(
        whsc1_row
    )
    sections = provider.build_all_gene_knowledge_sections(
        [whsc1_row, nsd2_row]
    )
    assert len(sections) == 1
    assert sections[0]["gene"] == "WHSC1"

    distinct = provider.build_all_gene_knowledge_sections(
        [
            whsc1_row,
            {**nsd2_row, "cHGVS": "c.101A>G", "pHGVS": "p.K34R"},
        ]
    )
    assert len(distinct) == 2

    crc301 = _panel_provider("crc_301_msi")
    assert crc301.variant_identity_key(nsd2_row) != crc301.variant_identity_key(
        whsc1_row
    )


def test_gene_symbol_alias_contract_rejects_cycles():
    with pytest.raises(ValueError, match="cyclic gene_symbol_aliases"):
        GeneKnowledgeProvider(
            {
                "enabled": True,
                "gene_symbol_aliases": {"GENE_A": "GENE_B", "GENE_B": "GENE_A"},
            }
        )


def test_gene_fallbacks_do_not_override_exact_variant_analysis():
    provider = _crc358_provider()
    exact_markers = {
        ("ACVR2A", "c.1310dup", "p.R438Efs*19"): "COSMIC数据库中记载12次",
        ("AMER1", "c.1489C>T", "p.R497*"): "COSMIC数据库中记载29次",
        ("FLT3", "c.2537G>A", "p.G846D"): "IL3非依赖性生长",
        ("EGFR", "c.2387G>A", "p.G796D"): "支持获得性耐药方向",
        ("GNAS", "c.1030G>A", "p.E344K"): "不属于本库已整理的经典热点",
        ("LRP1B", "c.1987G>A", "p.D663N"): "长尾抑癌相关变异解释",
        ("SETD2", "c.4930G>T", "p.G1644*"): "AZD1775",
        ("EPHA2", "c.153+2T>C", ""): "可能导致mRNA剪接异常",
        ("DDR2", "c.1466G>A", "p.R489Q"): "COSMIC数据库中记载3次",
        ("FGFR1", "c.1648G>T", "p.A550S"): "p.A550S错义突变位于酪氨酸激酶催化结构域",
        ("BCL2", "c.234del", "p.G79Afs*17"): "COSMIC数据库中记载2次",
        ("CCND2", "c.480T>G", "p.I160M"): "p.I160M突变位于Cyclin_C结构域",
        ("FGF4", "c.278G>A", "p.G93D"): "p.G93D突变位于FGF结构域",
        ("HIF1A", "c.1737G>T", "p.Q579H"): "p.Q579H突变位于HIF-1结构域",
        ("NRG1", "c.1850G>A", "p.R617H"): "p.R617H突变位于Neuregulin结构域",
        ("PDGFB", "c.274G>A", "p.A92T"): "条件性敲除血小板中的PDGFB",
        ("CCND1", "c.127T>G", "p.S43A"): "p.S43A为错义变异",
        ("PDGFB", "c.671G>A", "p.R224Q"): "p.R224Q为错义变异",
        ("FLCN", "c.1285del", "p.H429Tfs*39"): "COSMIC数据库中记载34次",
        ("EZH2", "c.1506-2A>G", ""): "c.1506-2A>G突变位于内含子与外显子的交界处",
        ("MYD88", "c.534del", "p.M178Ifs*12"): "p.M178Ifs*12突变导致蛋白翻译提前终止",
        ("NPM1", "c.676G>C", "p.E226Q"): "p.E226Q突变位于上述结构域之外",
    }
    generic_boundary_markers = (
        "未命中精确规则时不得生成",
        "未收录的ERBB2单核苷酸变异",
        "未证实功能影响的位点",
        "未收录位点需结合变异类型",
        "普通错义变异不能自动按激活性事件处理",
        "单个未收录位点不能独立作为治疗",
    )

    for selector, marker in exact_markers.items():
        section = _section(provider, *selector)
        assert marker in section["mutation_analysis"], (selector, section)
        assert not any(
            boundary in section["mutation_analysis"]
            for boundary in generic_boundary_markers
        ), (selector, section)


def test_variant_overlay_domain_sentence_is_promoted_without_variant_prose(tmp_path):
    overlay = tmp_path / "reviewed.yaml"
    overlay.write_text(
        """
gene_sections:
  - gene: DEMO
    c_hgvs: c.1A>G
    p_hgvs: p.M1V
    mutation_analysis: >-
      DEMO基因编码的蛋白全长为100个氨基酸，主要包含Demo结构域（10-80位氨基酸）。
      该样本检出的p.M1V可能影响蛋白功能。
drug_sections: []
""".strip(),
        encoding="utf-8",
    )
    provider = GeneKnowledgeProvider(
        {
            "enabled": True,
            "gene_knowledge_db": {
                "enabled": True,
                "path": "missing.xlsx",
                "reviewed_part3_overlay_paths": [str(overlay)],
            },
        }
    )
    assert provider.load(base_path=str(ROOT))

    other = provider.build_gene_knowledge_section(
        gene="DEMO",
        c_hgvs="c.2A>G",
        p_hgvs="p.M2V",
        frequency=10.0,
        mutation_type="Missense",
        has_drug=False,
    )

    assert "蛋白全长为100个氨基酸" in other["fixed_domain_text"]
    assert "p.M1V" not in other["mutation_analysis"]


def test_gene_domain_coverage_fails_loudly_for_missing_part3_gene():
    result = GeneKnowledgeProvider.build_gene_domain_coverage(
        [
            {
                "gene": "ERBB2",
                "source_variant_key": "ERBB2|C.1A>G|P.M1V",
                "fixed_domain_text": "ERBB2基因编码的蛋白全长为1255个氨基酸。",
            },
            {
                "gene": "FANCD2",
                "source_variant_key": "FANCD2|C.2A>G|P.M2V",
                "fixed_domain_text": "",
            },
        ]
    )

    assert result == {
        "status": "FAIL",
        "expected_gene_count": 2,
        "covered_gene_count": 1,
        "coverage_percent": 50.0,
        "missing_genes": ["FANCD2"],
        "missing_variant_keys": ["FANCD2|C.2A>G|P.M2V"],
        "duplicate_fixed_domain_genes": [],
        "duplicate_fixed_domain_variant_keys": [],
    }


def test_gene_domain_coverage_fails_loudly_for_duplicate_fixed_statement():
    result = GeneKnowledgeProvider.build_gene_domain_coverage(
        [
            {
                "gene": "DNMT3A",
                "source_variant_key": "DNMT3A|C.1A>G|P.M1V",
                "fixed_domain_text": (
                    "DNMT3A基因编码的蛋白全长为912位氨基酸。\n"
                    "DNMT3A基因编码的蛋白全长为912个氨基酸。"
                ),
            }
        ]
    )

    assert result["status"] == "FAIL"
    assert result["missing_genes"] == []
    assert result["duplicate_fixed_domain_genes"] == ["DNMT3A"]
    assert result["duplicate_fixed_domain_variant_keys"] == [
        "DNMT3A|C.1A>G|P.M1V"
    ]


def test_crc358_domain_catalog_covers_every_reportable_gene():
    catalog_path = (
        ROOT / "panels/crc_358_msi/rules/reviewed_part3_domain_catalog.yaml"
    )
    validation = load_and_validate_overlay(catalog_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["total_rows"] == 273
    assert validation["gene"]["status_counts"] == {"approved_for_runtime": 273}
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]

    reportable_by_panel = {
        panel_id: {
            str(gene).upper()
            for gene in yaml.safe_load(
                (
                    ROOT
                    / f"panels/{panel_id}/rules/knowledge_coverage.yaml"
                ).read_text(encoding="utf-8")
            )["reportable_genes"]
        }
        for panel_id in ("crc_301_msi", "crc_358_msi")
    }
    catalog = yaml.safe_load(catalog_path.read_text(encoding="utf-8"))
    assert catalog["source"]["consumer_panels"] == [
        "crc_301_msi",
        "crc_358_msi",
    ]
    assert all(
        row["panels"]
        == sorted(
            panel_id
            for panel_id, genes in reportable_by_panel.items()
            if row["gene"] in genes
        )
        for row in catalog["gene_sections"]
    )

    coverage = {"reportable_genes": sorted(reportable_by_panel["crc_358_msi"])}
    provider = _crc358_provider()
    sections = [
        _section(provider, gene, "c.999999A>G", "p.X999999Y")
        for gene in coverage["reportable_genes"]
    ]
    result = provider.build_gene_domain_coverage(sections)

    assert result["status"] == "PASS"
    assert result["expected_gene_count"] == 358
    assert result["covered_gene_count"] == 358
    assert result["missing_genes"] == []


def test_crc301_reuses_shared_domains_and_covers_its_panel_specific_genes():
    catalog_path = (
        ROOT / "panels/crc_301_msi/rules/reviewed_part3_domain_catalog.yaml"
    )
    validation = load_and_validate_overlay(catalog_path, "crc_301_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["gene"]["total_rows"] == 45
    assert validation["gene"]["secondary_review_complete_rows"] == validation["gene"]["total_rows"]
    catalog = yaml.safe_load(catalog_path.read_text(encoding="utf-8"))
    assert catalog["source"]["consumer_panels"] == ["crc_301_msi"]
    assert all(
        row["panels"] == ["crc_301_msi"] for row in catalog["gene_sections"]
    )

    package = load_panel_package("crc_301_msi", project_root=ROOT)
    settings = yaml.safe_load((ROOT / "config/settings.yaml").read_text(encoding="utf-8"))
    gene_config = copy.deepcopy(settings["knowledge_bases"]["gene_knowledge_db"])
    gene_config["panel_id"] = "crc_301_msi"
    gene_config["reviewed_part3_overlay_paths"] = (
        ReportGenerator._resolve_panel_reviewed_part3_overlays(package)
    )
    provider = GeneKnowledgeProvider(
        {"enabled": True, "panel_id": "crc_301_msi", "gene_knowledge_db": gene_config}
    )
    assert provider.load(base_path=str(ROOT))
    coverage = yaml.safe_load(
        (ROOT / "panels/crc_301_msi/rules/knowledge_coverage.yaml").read_text(
            encoding="utf-8"
        )
    )
    result = provider.build_gene_domain_coverage(
        [
            provider.build_gene_knowledge_section(
                gene=gene,
                c_hgvs="c.999999A>G",
                p_hgvs="p.X999999Y",
                frequency=10.0,
                mutation_type="Missense",
                has_drug=False,
            )
            for gene in coverage["reportable_genes"]
        ]
    )
    assert result["status"] == "PASS"
    assert result["expected_gene_count"] == 301
    assert result["covered_gene_count"] == 301


def test_domain_catalog_records_ambiguous_and_nonprotein_boundaries():
    catalog = yaml.safe_load(
        (
            ROOT
            / "panels/crc_358_msi/rules/reviewed_part3_domain_catalog.yaml"
        ).read_text(encoding="utf-8")
    )
    rows = {row["gene"]: row for row in catalog["gene_sections"]}

    assert rows["GNAS"]["uniprot_accession"] == "P63092"
    assert rows["RBM10"]["uniprot_accession"] == "P98175"
    assert rows["LBP1B"]["uniprot_accession"] == "Q9NZI7"
    assert rows["TERC"]["annotation_status"] == "not_applicable_noncoding_rna"
    assert "不编码蛋白" in rows["TERC"]["fixed_domain_text"]


def test_domain_coverage_contract_becomes_a_blocking_business_issue():
    checks = _build_business_checks(
        "",
        {
            "gene_domain_coverage": {
                "status": "FAIL",
                "expected_gene_count": 2,
                "covered_gene_count": 1,
                "coverage_percent": 50.0,
                "missing_genes": ["FANCD2"],
                "missing_variant_keys": ["FANCD2|C.2A>G|P.M2V"],
            }
        },
        "crc_358_msi",
    )
    issues = list(_business_issues(checks))

    assert checks["gene_domain_coverage"]["status"] == "FAIL"
    assert issues == [
        {
            "level": "error",
            "code": "GENE_DOMAIN_COVERAGE",
                "message": (
                    "Part-3 gene knowledge has missing or duplicate fixed protein/domain content."
                ),
        }
    ]


def _approved_drug_names(rows):
    return [str(row.get("Drug") or "").splitlines()[0] for row in rows]


def test_crc358_approved_drug_universe_is_seven_when_part2_has_no_match():
    panel_config = load_panel_config(base_path=str(ROOT), panel_id="crc_358_msi")

    assert panel_config.approved_drug_rows_display_mode == (
        "exclude_if_listed_in_part2"
    )
    assert _approved_drug_names(_build_crc_approved_drugs(panel_config, [])) == [
        "瑞戈非尼",
        "贝伐珠单抗",
        "雷莫西尤单抗",
        "呋喹替尼",
        "阿柏西普",
        "阿帕替尼",
        "盐酸安罗替尼",
    ]


def test_crc358_approved_drugs_subtract_part2_across_names_and_combinations():
    panel_config = load_panel_config(base_path=str(ROOT), panel_id="crc_358_msi")
    part2_rows = [
        {
            "benefit_drugs_full": "瑞戈非尼（C）\nFOLFIRI+Bevacizumab（A）",
            "benefit_drugs": "瑞戈非尼（C）\n另1项详见第三部分",
            "caution_drugs": "Aflibercept（C）\n安罗替尼（C）",
        }
    ]

    assert _approved_drug_names(
        _build_crc_approved_drugs(panel_config, part2_rows)
    ) == ["雷莫西尤单抗", "呋喹替尼", "阿帕替尼"]


def test_crc358_approved_drugs_can_be_fully_suppressed_by_all_seven_aliases():
    panel_config = load_panel_config(base_path=str(ROOT), panel_id="crc_358_msi")
    part2_rows = [
        {
            "benefit_drugs": (
                "Regorafenib（C）\n[安维汀]\n希冉择\nFruquintinib\n"
                "ZALTRAP\n艾坦\nAnlotinib Hydrochloride"
            ),
            "caution_drugs": "--",
        }
    ]

    assert _build_crc_approved_drugs(panel_config, part2_rows) == []


def test_approved_drug_matching_does_not_use_genes_or_near_names():
    panel_config = load_panel_config(base_path=str(ROOT), panel_id="crc_358_msi")
    part2_rows = [
        {
            "gene": "VEGFR2",
            "benefit_drugs": "阿帕他胺（C）\n某艾坦试验",
            "caution_drugs": "--",
        }
    ]

    # Gene overlap never suppresses a row, and the deliberately near-but-not-
    # equal drug name stays distinct.
    assert _approved_drug_names(
        _build_crc_approved_drugs(panel_config, part2_rows)
    ) == [
        "瑞戈非尼",
        "贝伐珠单抗",
        "雷莫西尤单抗",
        "呋喹替尼",
        "阿柏西普",
        "阿帕替尼",
        "盐酸安罗替尼",
    ]


def test_crc358_dynamic_policy_overrides_stale_prefilled_approved_table(tmp_path):
    source = tmp_path / "empty.xlsx"
    source.touch()
    excel_data = ExcelDataSource(file_path=str(source), table_data={"Variations": []})
    report_data = ReportData()
    report_data.set_table(
        "variants_2_1",
        [
            {
                "gene": "DEMO",
                "benefit_drugs": "瑞戈非尼（C）",
                "caution_drugs": "--",
            }
        ],
    )
    report_data.set_table(
        "chemotherapy",
        [{"Drug": "过期占位", "Gene": "--", "药物适应情况": "--"}],
    )

    enhanced = enhance_report_data(
        report_data,
        excel_data,
        base_path=str(ROOT),
        panel_id="crc_358_msi",
    )

    assert _approved_drug_names(enhanced.get_table("chemotherapy")) == [
        "贝伐珠单抗",
        "雷莫西尤单抗",
        "呋喹替尼",
        "阿柏西普",
        "阿帕替尼",
        "盐酸安罗替尼",
    ]
    assert enhanced.get_field("approved_drug_rows_total_count") == 7
    assert enhanced.get_field("approved_drug_rows_suppressed_count") == 1
    assert enhanced.get_field("approved_drug_rows_suppressed_names") == [
        "瑞戈非尼"
    ]
    summary = build_report_summary(report_data=enhanced)
    assert summary["drugs"]["approved_display_policy"] == (
        "exclude_if_listed_in_part2"
    )
    assert summary["drugs"]["approved_universe_count"] == 7
    assert summary["drugs"]["approved_suppressed_count"] == 1
    assert summary["drugs"]["approved_suppressed_names"] == ["瑞戈非尼"]


def test_crc301_keeps_fixed_approved_table_when_part2_lists_a_drug():
    panel_config = load_panel_config(base_path=str(ROOT), panel_id="crc_301_msi")

    assert panel_config.approved_drug_rows_display_mode == "fixed"
    assert _approved_drug_names(
        _build_crc_approved_drugs(
            panel_config,
            [{"benefit_drugs": "瑞戈非尼（C）", "caution_drugs": "--"}],
        )
    ) == [
        "瑞戈非尼",
        "贝伐珠单抗",
        "雷莫西尤单抗",
        "呋喹替尼",
        "阿柏西普",
        "阿帕替尼",
        "盐酸安罗替尼",
    ]


def test_drug_consistency_gate_detects_missing_and_duplicate_items():
    provider = GeneKnowledgeProvider({"enabled": False})
    variant = {
        "gene": "DEMO",
        "cHGVS": "c.1A>G",
        "pHGVS": "p.M1V",
        "benefit_drugs": "甲药（C）\n乙药（C）",
        "caution_drugs": "--",
    }
    sections = [
        {
            "gene": "DEMO",
            "c_hgvs": "c.1A>G",
            "p_hgvs": "p.M1V",
            "drug_type": "benefit",
            "drug_name": "甲药（Drug A）",
        },
        {
            "gene": "DEMO",
            "c_hgvs": "c.1A>G",
            "p_hgvs": "p.M1V",
            "drug_type": "benefit",
            "drug_name": "甲药",
        },
    ]

    result = provider.build_drug_analysis_consistency([variant], sections)

    assert result["status"] == "FAIL"
    assert result["missing"][0]["drugs"] == ["乙药（C）"]
    assert result["duplicates"][0]["drugs"] == ["甲药（Drug A）"]


def test_targeted_drug_ingestion_restores_missing_grade_delimiters():
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")

    assert mapper._normalize_drug_evidence_label(
        "NXP800（C）Tuvusertib+Peposertib（C）"
    ) == "NXP800（C）\nTuvusertib+Peposertib（C）"
    assert mapper._normalize_drug_evidence_label(
        "替西罗莫司（C）Tuvusertib+Peposertib（C）"
    ) == "替西罗莫司（C）\nTuvusertib+Peposertib（C）"
    assert mapper._normalize_drug_evidence_label("A+B（C）") == "A+B（C）"


def test_crc358_historical_gene_class_drug_rows_are_fail_closed():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    rules = load_targeted_drug_rule_context(package)
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")

    arid1a, arid1a_caution, arid1a_score = mapper._lookup_targeted_drugs_for_variant(
        "ARID1A",
        c_point="c.5965C>T",
        p_point="p.R1989*",
        variant_level="Ⅱ类",
        cancer_type="结直肠癌",
        targeted_drug_rules=rules,
    )
    fbxw7, fbxw7_caution, fbxw7_score = mapper._lookup_targeted_drugs_for_variant(
        "FBXW7",
        c_point="c.979G>T",
        p_point="p.E327*",
        variant_level="Ⅱ类",
        cancer_type="结直肠癌",
        targeted_drug_rules=rules,
    )

    assert arid1a_score == 100.0
    assert (arid1a, arid1a_caution) == ("--", "--")
    assert fbxw7_score == 100.0
    assert (fbxw7, fbxw7_caution) == ("--", "--")


def test_positionless_internal_drug_rows_are_rejected_for_every_gene():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    rules = load_targeted_drug_rule_context(package)
    assert rules is not None
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")

    assert mapper._targeted_drug_db_row_applicable(
        gene="KMT2D",
        source_db="internal",
        db_c="",
        db_p="",
        db_variant_type="",
        targeted_drug_rules=rules,
    ) is False
    assert mapper._targeted_drug_db_row_applicable(
        gene="KMT2D",
        source_db="internal",
        db_c="",
        db_p="p.M1V",
        db_variant_type="",
        targeted_drug_rules=rules,
    ) is True
    assert mapper._targeted_drug_db_row_applicable(
        gene="KMT2D",
        source_db="CIVIC",
        db_c="",
        db_p="",
        db_variant_type="",
        targeted_drug_rules=rules,
    ) is True


def test_crc358_uat_first_review_closes_all_positionless_candidate_rows():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    rules = load_targeted_drug_rule_context(package)
    assert rules is not None
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")

    candidates = [
        ("APC", "c.22C>T", "p.Q8*"),
        ("APC", "c.2626C>T", "p.R876*"),
        ("APC", "c.3340C>T", "p.R1114*"),
        ("APC", "c.3944C>G", "p.S1315*"),
        ("APC", "c.3982C>T", "p.Q1328*"),
        ("ARID1A", "c.3634C>T", "p.Q1212*"),
        ("ARID1A", "c.5965C>T", "p.R1989*"),
        ("ATR", "c.6729del", "p.F2243Lfs*8"),
        ("BRCA1", "c.4357+1G>A", "--"),
        ("BRCA2", "c.7008-1G>A", "--"),
        ("BRCA2", "c.8403del", "p.P2802Lfs*19"),
        ("FBXW7", "c.1957del", "p.T653Rfs*8"),
        ("FBXW7", "c.40C>T", "p.R14*"),
        ("FBXW7", "c.979G>T", "p.E327*"),
        ("MLH1", "c.676C>T", "p.R226*"),
        ("MSH6", "c.1636G>T", "p.E546*"),
        ("NF1", "c.3721C>T", "p.R1241*"),
        ("NF1", "c.6007-1G>T", "--"),
        ("PIK3CA", "c.1633G>A", "p.E545K"),
        ("PMS2", "c.1273del", "p.S425Lfs*23"),
        ("PTEN", "c.400A>T", "p.M134L"),
        ("SMARCA4", "c.1189C>T", "p.R397*"),
    ]
    for gene, c_hgvs, p_hgvs in candidates:
        benefit, caution, score = mapper._lookup_targeted_drugs_for_variant(
            gene,
            c_point=c_hgvs,
            p_point=p_hgvs,
            variant_level="Ⅱ类",
            cancer_type="结直肠癌",
            targeted_drug_rules=rules,
        )
        assert (benefit, caution, score) == ("--", "--", 100.0), (
            gene,
            c_hgvs,
            p_hgvs,
            benefit,
            caution,
            score,
        )

    benefit, caution, score = mapper._lookup_targeted_drugs_for_variant(
        "KRAS",
        c_point="c.35G>T",
        p_point="p.G12V",
        variant_level="Ⅱ类",
        cancer_type="结直肠癌",
        targeted_drug_rules=rules,
    )
    assert benefit == "--"
    assert caution == "西妥昔单抗（A）\n帕尼单抗（A）"
    assert "依维莫司" not in caution
    assert score == 100.0


def test_crc358_uat_rules_are_provisional_and_panel_scoped():
    crc358 = load_targeted_drug_rule_context(
        load_panel_package("crc_358_msi", project_root=ROOT)
    )
    crc301 = load_targeted_drug_rule_context(
        load_panel_package("crc_301_msi", project_root=ROOT)
    )
    assert crc358 is not None and crc301 is not None

    new_rows = [
        row
        for row in crc358["reviewed_variant_overrides"]
        if row.get("secondary_review_status")
        == "pending_report_group_secondary_review"
    ]
    assert len(new_rows) == 14
    assert all(row["review_status"] == "provisional_runtime" for row in new_rows)
    assert all(row["runtime_eligible"] is True for row in new_rows)
    assert all(row.get("source_refs") for row in new_rows)
    assert not any(
        row.get("secondary_review_status")
        == "pending_report_group_secondary_review"
        for row in crc301["reviewed_variant_overrides"]
    )
    assert crc301["applicability_rules"] == crc358["applicability_rules"]


def test_crc358_uat_research_sections_do_not_recreate_part2_benefit():
    rules = load_targeted_drug_rule_context(
        load_panel_package("crc_358_msi", project_root=ROOT)
    )
    assert rules is not None
    provider = _crc358_provider()

    def rule_for(gene: str, p_hgvs: str = ""):
        return next(
            row
            for row in rules["reviewed_variant_overrides"]
            if row.get("gene") == gene
            and row.get("secondary_review_status")
            == "pending_report_group_secondary_review"
            and (not p_hgvs or row.get("p_hgvs") == p_hgvs)
        )

    variants = []
    for gene, c_hgvs, p_hgvs in (
        ("APC", "c.22C>T", "p.Q8*"),
        ("BRCA2", "c.8403del", "p.P2802Lfs*19"),
        ("NF1", "c.3721C>T", "p.R1241*"),
        ("KRAS", "c.35G>T", "p.G12V"),
    ):
        rule = rule_for(gene, p_hgvs if gene == "KRAS" else "")
        variants.append(
            {
                "gene": gene,
                "cHGVS": c_hgvs,
                "pHGVS": p_hgvs,
                "benefit_drugs": "--",
                "caution_drugs": (
                    "\n".join(rule.get("caution_drugs") or [])
                    if gene == "KRAS"
                    else "--"
                ),
                "research_drugs": "\n".join(rule.get("research_drugs") or []),
            }
        )

    sections = provider.build_drug_analysis_sections(variants)
    assert not any(row["drug_type"] == "benefit" for row in sections)
    assert {
        row["gene"]
        for row in sections
        if row.get("drug_type") == "research"
    } == {"APC", "BRCA2", "NF1", "KRAS"}
    kras_sections = [row for row in sections if row["gene"] == "KRAS"]
    assert {row["drug_type"] for row in kras_sections} == {"caution", "research"}
    assert "依维莫司" not in "\n".join(
        row.get("drug_name", "") for row in kras_sections
    )


def test_feedback_drug_rows_match_part3_without_overlapping_blocks():
    provider = _crc358_provider()
    variants = [
        {
            "gene": "KRAS",
            "cHGVS": "c.35G>A",
            "pHGVS": "p.G12D",
            "benefit_drugs": "\n".join(
                [
                    "Avutometinib+Defactinib（C）",
                    "司美替尼（C）",
                    "曲美替尼+Navitoclax（C）",
                    "帕尼单抗+曲美替尼（C）",
                    "贝美替尼+哌柏西利（C）",
                    "BI 1701963（C）",
                    "BI 1701963+曲美替尼（C）",
                    "PD0325901+哌柏西利（C）",
                    "奈拉替尼+曲美替尼（C）",
                    "福巴替尼+贝美替尼（C）",
                    "依维莫司+Avutometinib（C）",
                    "GH35（C）",
                    "RMC-6236（C）",
                    "HRS-4642（C）",
                    "ASP3082（C）",
                    "ASP4396（C）",
                    "RMC-9805（C）",
                    "RMC-9805+RMC-6236（C）",
                    "PD0325901（D）",
                ]
            ),
            "caution_drugs": "西妥昔单抗（A）\n帕尼单抗（A）\n依维莫司（C）",
        },
        {
            "gene": "PALB2",
            "cHGVS": "c.47del",
            "pHGVS": "p.K16Sfs*2",
            "benefit_drugs": "--",
            "caution_drugs": "--",
            "research_drugs": "芦卡帕利\nLY2606368/Prexasertib",
        },
        {
            "gene": "RAD51D",
            "cHGVS": "c.685C>T",
            "pHGVS": "p.Q229*",
            "benefit_drugs": "--",
            "caution_drugs": "--",
            "research_drugs": "芦卡帕利\nLY2606368/Prexasertib",
        },
    ]

    sections = provider.build_drug_analysis_sections(variants)
    result = provider.build_drug_analysis_consistency(variants, sections)

    assert result["status"] == "PASS", result
    kras_benefit = [
        row for row in sections if row.get("gene") == "KRAS" and row.get("drug_type") == "benefit"
    ]
    assert len(kras_benefit) == 3
    assert {
        row["gene"]
        for row in sections
        if row.get("drug_type") == "research"
    } == {"PALB2", "RAD51D"}
    assert not any(
        row["gene"] in {"PALB2", "RAD51D"}
        and row.get("drug_type") == "benefit"
        for row in sections
    )


def test_all_configured_crc_drug_rules_have_governed_consistent_part3_contracts():
    provider = _crc358_provider()

    assert provider.has_reviewed_drug_analysis_contract(
        {
            "gene": "KRAS",
            "cHGVS": "c.35G>A",
            "pHGVS": "p.G12D",
        }
    )
    assert provider.has_reviewed_drug_analysis_contract(
        {
            "gene": "PALB2",
            "cHGVS": "c.47del",
            "pHGVS": "p.K16Sfs*2",
        }
    )
    assert provider.has_reviewed_drug_analysis_contract(
        {
            "gene": "ERBB2",
            "cHGVS": "c.1979G>A",
            "pHGVS": "p.G660D",
        }
    )

    expected = {
        "crc_301_msi": (9, 10),
        "crc_358_msi": (31, 32),
    }
    for panel_id, (rules, selector_cases) in expected.items():
        package = load_panel_package(panel_id, project_root=ROOT)
        result = profile_panel_targeted_drug_contracts(ROOT, package)
        assert result["status"] == "PASS", result["issues"]
        assert result["rules_checked"] == rules
        assert result["selector_cases_checked"] == selector_cases
        assert result["expected_item_count"] == result["rendered_item_count"]
        assert result["issues"] == []


def test_setd2_rule_is_limited_to_loss_of_function_and_discloses_phase_ii_result():
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    context = load_targeted_drug_rule_context(package)
    assert context is not None
    rule = next(
        row
        for row in context["reviewed_variant_overrides"]
        if row.get("gene") == "SETD2"
    )
    assert rule["applicability"] == "loss_of_function"
    assert "SETD2" not in context["overrides"]

    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")
    assert (
        mapper._lookup_reviewed_variant_override_drugs(
            "SETD2",
            "c.1A>G",
            "p.M1V",
            variant_level="Ⅱ类",
            targeted_drug_rules=context,
        )
        is None
    )
    assert mapper._lookup_reviewed_variant_override_drugs(
        "SETD2",
        "c.1del",
        "p.M1Rfs*2",
        variant_level="Ⅱ类",
        targeted_drug_rules=context,
    ) == ("--", "--")

    provider = _crc358_provider()
    missense = {
        "gene": "SETD2",
        "cHGVS": "c.1A>G",
        "pHGVS": "p.M1V",
        "benefit_drugs": "--",
        "caution_drugs": "--",
        "research_drugs": "AZD1775/Adavosertib",
    }
    lof = {
        **missense,
        "cHGVS": "c.1del",
        "pHGVS": "p.M1Rfs*2",
    }
    assert provider.build_drug_analysis_sections([missense]) == []
    sections = provider.build_drug_analysis_sections([lof])
    assert len(sections) == 1
    assert sections[0]["drug_type"] == "research"
    assert sections[0]["drug_name"] == "AZD1775/Adavosertib"
    combined = f"{sections[0]['relation']}\n{sections[0]['clinical']}"
    assert "26602815" in combined
    assert "38920407" in combined
    assert "未观察到客观缓解" in combined
    assert "29955114" not in combined


def test_gene_level_atm_and_event_rules_keep_the_current_variant_identity():
    provider = _crc358_provider()
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    context = load_targeted_drug_rule_context(package)
    assert context is not None
    assert context["overrides"] == {}
    atm = next(
        row
        for row in context["reviewed_variant_overrides"]
        if row.get("gene") == "ATM"
    )
    assert atm["applicability"] == "loss_of_function"
    tsc1 = next(
        row
        for row in context["reviewed_variant_overrides"]
        if row.get("gene") == "TSC1"
    )

    def drug_text(value):
        return value if isinstance(value, str) else "\n".join(value or [])

    variants = [
        {
            "gene": "ATM",
            "cHGVS": "c.1del",
            "pHGVS": "p.M1Rfs*2",
            "benefit_drugs": drug_text(atm["benefit_drugs"]),
            "caution_drugs": drug_text(atm["caution_drugs"]),
            "research_drugs": drug_text(atm["research_drugs"]),
        },
        {
            "gene": "TSC1",
            "cHGVS": "c.2del",
            "pHGVS": "p.M2Rfs*2",
            "benefit_drugs": drug_text(tsc1["benefit_drugs"]),
            "caution_drugs": drug_text(tsc1.get("caution_drugs")) or "--",
        },
    ]
    sections = provider.build_drug_analysis_sections(variants)
    identities = {
        (row["gene"], row["c_hgvs"], row["p_hgvs"])
        for row in sections
    }
    assert ("ATM", "c.1del", "p.M1Rfs*2") in identities
    assert ("TSC1", "c.2del", "p.M2Rfs*2") in identities
    assert provider.build_drug_analysis_consistency(variants, sections)[
        "status"
    ] == "PASS"
    atm_sections = [row for row in sections if row["gene"] == "ATM"]
    assert len(atm_sections) == 1
    assert atm_sections[0]["drug_type"] == "research"


def test_kras_everolimus_rules_use_primary_source_and_conservative_scope():
    overlay = yaml.safe_load(
        (ROOT / "panels/crc_358_msi/rules/reviewed_part3_knowledge.yaml").read_text(
            encoding="utf-8"
        )
    )
    rows = [
        row
        for row in overlay["drug_sections"]
        if row.get("gene") == "KRAS"
        and row.get("p_hgvs") in {"p.G12S", "p.G12D"}
        and row.get("type") == "caution"
        and "依维莫司" in str(row.get("drug_name") or "")
    ]

    assert {row["p_hgvs"] for row in rows} == {"p.G12S", "p.G12D"}
    for row in rows:
        combined = f"{row['relation']}\n{row['clinical']}"
        assert "20664172" in combined
        assert "29285035" not in combined
        assert "28544747" not in combined
        assert "并非" in combined and "位点专属" in combined
        assert row["review_status"] == "approved_for_runtime"
        assert row["secondary_review_status"] == "report_group_approved"
        assert row["review_metadata"]["references"] == [
            "https://pmc.ncbi.nlm.nih.gov/articles/PMC2912177/"
        ]


def test_kras_g12c_correction_supersedes_misattributed_historical_entry():
    correction_path = ROOT / "panels/crc_358_msi/rules/reviewed_part3_corrections_20260720.yaml"
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    overlays = [
        Path(path) for path in ReportGenerator._resolve_panel_reviewed_part3_overlays(package)
    ]
    historical_path = ROOT / "panels/crc_358_msi/rules/reviewed_part3_crc358_reviewed_case_a.yaml"
    assert correction_path in overlays
    assert overlays.index(correction_path) > overlays.index(historical_path)

    validation = load_and_validate_overlay(correction_path, "crc_358_msi")
    assert validation["status"] == "PASS", validation["issues"]
    assert validation["drug"]["status_counts"] == {"approved_for_runtime": 2}
    assert validation["drug"]["secondary_review_complete_rows"] == validation["drug"]["total_rows"]

    provider = _crc358_provider()
    variant = {
        "gene": "KRAS",
        "cHGVS": "c.34G>T",
        "pHGVS": "p.G12C",
        "benefit_drugs": "--",
        "caution_drugs": "依维莫司（C）",
    }
    rows = [
        row
        for row in provider.build_drug_analysis_sections([variant])
        if row.get("gene") == "KRAS" and "依维莫司" in row.get("drug_name", "")
    ]
    assert len(rows) == 1
    combined = f"{rows[0]['relation']}\n{rows[0]['clinical']}"
    assert combined.count("20664172") == 2
    assert "29285035" not in combined
    assert "28544747" not in combined
    assert "并非p.G12C位点专属临床证据" in combined


def test_kras_g12c_correction_does_not_leak_to_adjacent_variant():
    provider = _crc358_provider()
    rows = provider.build_drug_analysis_sections(
        [
            {
                "gene": "KRAS",
                "cHGVS": "c.34G>C",
                "pHGVS": "p.G12R",
                "benefit_drugs": "--",
                "caution_drugs": "依维莫司（C）",
            }
        ]
    )
    assert len(rows) == 1
    combined = f"{rows[0]['relation']}\n{rows[0]['clinical']}"
    assert "p.G12C" not in combined
    assert "20664172" in combined
    assert "29285035" not in combined
    assert "28544747" not in combined
    assert "不是任一具体位点的专属临床证据" in combined


def test_nccn_results_annotate_only_class_iii_variants(tmp_path):
    source = tmp_path / "empty.xlsx"
    source.touch()
    excel_data = ExcelDataSource(file_path=str(source))
    report_data = ReportData()
    panel_config = PanelConfig(
        nccn_result_rows=[
            {"key": "ERBB2_MUT", "genes": ["ERBB2"], "match": "突变"},
            {"key": "KRAS_MUT", "genes": ["KRAS"], "match": "突变"},
        ],
        immune_positive_rows=[],
        immune_negative_rows=[],
        immune_hyperprogression_rows=[],
    )

    _build_nccn_and_immune_fields(
        report_data,
        [
            {
                "gene": "ERBB2",
                "cHGVS": "c.2521C>A",
                "pHGVS": "p.L841I",
                "gene_class": "Ⅲ类",
            },
            {
                "gene": "KRAS",
                "cHGVS": "c.35G>A",
                "pHGVS": "p.G12D",
                "gene_class": "Ⅱ类",
            },
        ],
        excel_data,
        panel_config=panel_config,
    )

    assert report_data.get_field("nccn_ERBB2_MUT") == ("c.2521C>A，p.L841I（意义未明变异）")
    assert report_data.get_field("nccn_KRAS_MUT") == "c.35G>A，p.G12D"


def _set_table_widths(table, width_twips: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_twips))

    grid_cols = table._tbl.tblGrid.findall(qn("w:gridCol"))
    per_column = width_twips // len(grid_cols)
    for grid_col in grid_cols:
        grid_col.set(qn("w:w"), str(per_column))
    for tc_w in table._tbl.findall(".//w:tcPr/w:tcW", table._tbl.nsmap):
        tc_w.set(qn("w:type"), "dxa")
        tc_w.set(qn("w:w"), str(per_column))


def test_crc_variant_detail_table_expands_to_content_width(tmp_path):
    docx_path = tmp_path / "variant_detail_width.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=9)
    table.rows[0].cells[0].text = "基因名称"
    table.rows[0].cells[1].text = "基因突变信息"
    table.rows[0].cells[6].text = "靶向药物信息"
    table.rows[1].cells[1].text = "转录本号"
    table.rows[1].cells[6].text = "潜在获益靶向药物"
    _set_table_widths(table, 4500)
    expected_width = int(
        doc.sections[0].page_width.twips
        - doc.sections[0].left_margin.twips
        - doc.sections[0].right_margin.twips
    )
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._fit_tables_to_page_width(
        str(docx_path),
        {
            "panel_style": {
                "variant_detail_table": {
                    "fit_to_content_width": True,
                    "minimum_content_width_ratio": 0.98,
                }
            }
        },
    )

    rendered = Document(docx_path)
    grid_cols = rendered.tables[0]._tbl.tblGrid.findall(qn("w:gridCol"))
    assert sum(int(col.get(qn("w:w"))) for col in grid_cols) == expected_width
    layout = rendered.tables[0]._tbl.tblPr.find(qn("w:tblLayout"))
    assert layout is not None
    assert layout.get(qn("w:type")) == "fixed"


def test_part3_drug_names_and_narrative_have_ten_point_space_after(tmp_path):
    docx_path = tmp_path / "part3_drug_spacing.docx"
    doc = Document()
    for text in [
        "第三部分：基因变异及相应靶向/免疫药物解析",
        "靶向药物/免疫用药提示解析",
        "潜在获益靶向/免疫药物解析",
        "KRAS：c.35G>A，p.G12D突变相应潜在获益药物",
        "依维莫司",
        "药物疗效临床解析：",
        "这是药物解析正文。",
        "3. 阅读说明",
    ]:
        doc.add_paragraph(text)
    doc.save(docx_path)

    TemplateRenderer(log_level="ERROR")._restore_part3_dynamic_styles(
        str(docx_path),
        {"drug_benefit_sections": [{"drug_name": "依维莫司"}]},
    )

    rendered = Document(docx_path)
    by_text = {
        paragraph.text.strip(): paragraph
        for paragraph in rendered.paragraphs
        if paragraph.text.strip()
    }
    assert by_text["依维莫司"].paragraph_format.space_after.pt == 10.0
    assert by_text["这是药物解析正文。"].paragraph_format.space_after.pt == 10.0


def test_approved_drug_table_bolds_complete_brand_brackets_and_qa_enforces_it(
    tmp_path,
):
    docx_path = tmp_path / "approved_drug_brands.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.rows[0].cells[0].text = "药物名称"
    table.rows[0].cells[1].text = "相关基因"
    table.rows[0].cells[2].text = "药物适应情况"
    first = table.rows[1].cells[0].paragraphs[0]
    first.add_run("贝伐珠单抗\nBevacizumab\n[")
    first.add_run("安维汀")
    first.add_run("]")
    table.rows[2].cells[0].text = "阿柏西普\nZiv-aflibercept\n[ZALTRAP]"
    doc.save(docx_path)

    before = Document(docx_path)
    before_check = _build_style_checks(
        before,
        "crc_358_msi",
        {"panel_style": {}},
    )["docx_style_rules"]
    assert before_check["status"] == "FAIL"
    assert {
        failure["actual"] for failure in before_check["failures"]
    } == {"[安维汀]", "[ZALTRAP]"}

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._bold_drug_brand_brackets(str(docx_path))
    renderer._bold_drug_brand_brackets(str(docx_path))

    rendered = Document(docx_path)
    after_check = _build_style_checks(
        rendered,
        "crc_358_msi",
        {"panel_style": {}},
    )["docx_style_rules"]
    assert after_check["status"] == "PASS"
    approved = rendered.tables[0]
    for row in approved.rows[1:]:
        bracket_runs = [
            run
            for run in row.cells[0].paragraphs[0].runs
            if run.text.startswith("[")
        ]
        assert len(bracket_runs) == 1
        assert bracket_runs[0].bold is True


def test_targeted_results_heading_is_configured_for_idempotent_new_page():
    settings = yaml.safe_load((ROOT / "config/settings.yaml").read_text(encoding="utf-8"))
    configured = settings["report_content"]["force_page_break_before_headings"]
    assert "2. 靶向药物相关检测结果" in configured
    assert "2.2 其它潜在获益上市药物提示*" in configured


def test_immune_brand_note_uses_compact_reviewed_style():
    doc = Document()
    paragraph = doc.add_paragraph(
        "3. 上表涉及的已上市的药物名称及对应的商品名称："
        "帕博利珠单抗[可瑞达]、纳武利尤单抗[欧狄沃]。"
    )

    changed = TemplateRenderer(log_level="ERROR")._apply_immune_table_notes_to_doc(
        doc,
        {
            "panel_style": {
                "biomarker_table": {
                    "note_font_size": 8,
                    "note_line_spacing": 1.0,
                }
            }
        },
    )

    assert changed is True
    assert paragraph.paragraph_format.line_spacing == 1.0
    assert paragraph.paragraph_format.space_before.pt == 0.0
    assert paragraph.paragraph_format.space_after.pt == 0.0
    assert all(run.font.size.pt == 8.0 for run in paragraph.runs if run.text)


def test_report_generator_reuses_the_validated_context_for_rendering(tmp_path, monkeypatch):
    generator = ReportGenerator(config_dir=str(ROOT / "config"), log_level="ERROR")
    validated_context = {
        "project_type": "crc_358_msi",
        "_require_deterministic_layout": True,
    }
    captured = {}

    def fake_render(*_args, **kwargs):
        captured.update(kwargs)
        generator.template_renderer.last_processor_report = []
        return str(tmp_path / "report.docx")

    monkeypatch.setattr(generator.template_renderer, "render", fake_render)
    monkeypatch.setattr(
        generator,
        "_get_template_processor_names",
        lambda *_args, **_kwargs: [],
    )
    state = SimpleNamespace(
        report_data=ReportData(),
        output_path=str(tmp_path / "report.docx"),
        template_file=str(tmp_path / "template.docx"),
        panel_package=None,
        canonical_project_type="crc_358_msi",
        template_context=validated_context,
        final_output=None,
        processor_report=[],
        template_processor_names=None,
    )
    stage = SimpleNamespace(artifacts={}, metrics={})

    generator._stage_template_render(stage, state)

    assert captured["template_context"] is validated_context
    assert captured["template_context"]["_require_deterministic_layout"] is True


def test_legal_notice_normalizer_writes_east_asian_font_and_is_idempotent(
    tmp_path,
):
    docx_path = tmp_path / "legal_notice.docx"
    doc = Document()
    notice = doc.add_paragraph()
    notice_run = notice.add_run("检测结果仅对本次送检样本负责，电子版仅供参考。")
    notice_run.font.name = "微软雅黑"
    consultation = doc.add_paragraph()
    consultation_run = consultation.add_run("咨询电话：00000000。")
    consultation_run.font.name = "微软雅黑"
    doc.save(docx_path)

    renderer = TemplateRenderer(log_level="ERROR")
    context = {
        "consultation_line": "咨询电话：00000000。",
        "report_content": {
            "legal_notice_style": {
                "marker": "检测结果仅对本次送检样本负责",
                "font_name": "宋体",
            }
        },
    }
    renderer._normalize_legal_notice_style(str(docx_path), context)

    rendered = Document(docx_path)
    for paragraph in rendered.paragraphs:
        for run in paragraph.runs:
            r_fonts = run._r.get_or_add_rPr().rFonts
            assert r_fonts is not None
            for attr in ("ascii", "hAnsi", "eastAsia"):
                assert r_fonts.get(qn(f"w:{attr}")) == "宋体"

    once = docx_path.read_bytes()
    renderer._normalize_legal_notice_style(str(docx_path), context)
    assert docx_path.read_bytes() == once
