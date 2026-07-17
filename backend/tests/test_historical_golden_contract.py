from __future__ import annotations

import hashlib
import json
import subprocess
import sys
from pathlib import Path

import pytest
import yaml
from docx import Document
from reportgen.core.historical_golden_contract import (
    load_historical_golden_contract,
    validate_historical_golden_docx,
)
from scripts.check_historical_golden_release import (
    _candidate_renderer_fingerprint,
    _sha256_files,
)

ROOT = Path(__file__).resolve().parents[2]
CONTRACT = (
    ROOT
    / "panels"
    / "crc_358_msi"
    / "golden_cases"
    / "crc358_reviewed_case_a.yaml"
)


def test_crc358_historical_contract_is_deidentified_and_structured() -> None:
    contract = load_historical_golden_contract(CONTRACT)
    assert contract["case_alias"] == "crc358_reviewed_case_a"
    assert contract["privacy"]["contains_phi"] is False
    assert contract["expectations"]["targeted_summary"]["row_count"] == 7
    assert contract["expectations"]["variant_counts"] == {
        "targeted_drug_related": 7,
        "targeted_or_immune_related": 8,
    }
    assert contract["medical_uat"]["status"] == "blocked"
    assert contract["medical_uat"]["blockers"][0]["id"] == (
        "KRAS_G12C_DRUG_RECONFIRMATION"
    )
    assert (
        len(contract["expectations"]["targeted_drug_brand_summary"]["ordered_pairs"])
        == 41
    )
    assert contract["expectations"]["part3"]["gene_section_count"] == 11
    assert contract["expectations"]["part3"]["drug_section_count"] == 18
    assert contract["expectations"]["part3"]["group_by_gene"] is True
    assert (
        contract["expectations"]["part3"]["gene_order_by_max_vaf_descending"]
        is True
    )
    assert contract["expectations"]["part3"]["within_gene_vaf_descending"] is True
    assert len(contract["expectations"]["part3"]["gene_section_order"]) == 11


def test_targeted_brand_config_order_matches_historical_contract() -> None:
    contract = load_historical_golden_contract(CONTRACT)
    config = yaml.safe_load(
        (ROOT / "config" / "drug_brands.yaml").read_text(encoding="utf-8")
    )
    expected_drugs = [
        pair.split("[", 1)[0]
        for pair in contract["expectations"]["targeted_drug_brand_summary"][
            "ordered_pairs"
        ]
    ]

    assert config["targeted_summary_order"] == expected_drugs
    assert all(drug in config["brands"] for drug in expected_drugs)


def test_historical_contract_loader_rejects_phi_contract(tmp_path) -> None:
    payload = yaml.safe_load(CONTRACT.read_text(encoding="utf-8"))
    payload["privacy"]["contains_phi"] = True
    path = tmp_path / "bad.yaml"
    path.write_text(yaml.safe_dump(payload, allow_unicode=True), encoding="utf-8")
    with pytest.raises(ValueError, match="contains_phi=false"):
        load_historical_golden_contract(path)


def _minimal_historical_docx(
    path: Path,
    *,
    include_tp53: bool = True,
    brand_pairs: list[str] | None = None,
    part3_headers: list[str] | None = None,
    targeted_count: int | None = None,
    targeted_or_immune_count: int | None = None,
) -> None:
    doc = Document()
    if targeted_count is not None:
        doc.add_paragraph(
            f"*本次共检出体细胞变异：11个，其中与靶向药物用药相关的变异有：{targeted_count}个。"
        )
    if targeted_or_immune_count is not None:
        doc.add_paragraph(
            "在本次检测范围内，检出体细胞变异：11个，"
            f"其中与靶向/免疫药物相关的变异：{targeted_or_immune_count}个。"
        )
    summary = doc.add_table(rows=2 if include_tp53 else 1, cols=4)
    summary.rows[0].cells[0].text = "基因"
    summary.rows[0].cells[1].text = "突变位点"
    summary.rows[0].cells[2].text = "潜在获益靶向药物"
    summary.rows[0].cells[3].text = "可能耐药"
    if include_tp53:
        summary.rows[1].cells[0].text = "TP53"
        summary.rows[1].cells[1].text = "c.499C>T, p.Q167*"
        summary.rows[1].cells[2].text = "AZD1775（C）"
        summary.rows[1].cells[3].text = "--"

    detail = doc.add_table(rows=3 if include_tp53 else 2, cols=9)
    detail.rows[0].cells[0].text = "基因名称"
    detail.rows[0].cells[7].text = "潜在获益靶向药物"
    detail.rows[1].cells[1].text = "转录本号"
    if include_tp53:
        row = detail.rows[2].cells
        row[0].text = "TP53"
        row[4].text = "c.499C>T, p.Q167*"
        row[6].text = "1.16"
        row[7].text = "AZD1775（C）"
        row[8].text = "--"
    if brand_pairs is not None:
        doc.add_paragraph(
            "2.上表涉及的已上市的药物名称及对应的商品名称："
            + "、".join(brand_pairs)
            + "。"
        )
    doc.add_paragraph("非业务差异标记")
    if part3_headers is not None:
        doc.add_paragraph("第三部分：基因变异及相应靶向/免疫药物解析")
        for header in part3_headers:
            doc.add_paragraph(header)
        doc.add_paragraph("第四部分：附录")
    doc.save(path)


def test_reference_sha_verifies_reference_not_candidate(tmp_path) -> None:
    reference = tmp_path / "reference.docx"
    candidate = tmp_path / "candidate.docx"
    _minimal_historical_docx(reference)
    _minimal_historical_docx(candidate)
    # Change only a non-contract paragraph so the candidate binary hash differs.
    doc = Document(candidate)
    doc.add_paragraph("候选版本标记")
    doc.save(candidate)
    reference_sha = hashlib.sha256(reference.read_bytes()).hexdigest()
    assert hashlib.sha256(candidate.read_bytes()).hexdigest() != reference_sha
    contract = {
        "schema_version": "1.0",
        "case_alias": "synthetic_case",
        "panel_id": "crc_358_msi",
        "privacy": {"contains_phi": False},
        "source": {"reference_docx_sha256": reference_sha},
        "expectations": {
            "table_count": 2,
            "targeted_summary": {"row_count": 1, "gene_order": ["TP53"]},
            "part3": {"gene_section_count": 0, "drug_section_count": 0},
            "reviewed_variant_rows": [
                {
                    "gene": "TP53",
                    "c_hgvs": "c.499C>T",
                    "p_hgvs": "p.Q167*",
                    "vaf": 1.16,
                    "benefit_count": 1,
                    "caution_count": 0,
                }
            ],
        },
    }

    result = validate_historical_golden_docx(
        contract=contract,
        docx_path=candidate,
        require_reference_sha=True,
        reference_docx_path=reference,
    )

    assert result["status"] == "PASS", result["errors"]


def test_historical_contract_blocks_missing_reviewed_variant(tmp_path) -> None:
    reference = tmp_path / "reference.docx"
    candidate = tmp_path / "candidate.docx"
    _minimal_historical_docx(reference)
    _minimal_historical_docx(candidate, include_tp53=False)
    contract = {
        "schema_version": "1.0",
        "case_alias": "synthetic_case",
        "panel_id": "crc_358_msi",
        "privacy": {"contains_phi": False},
        "source": {"reference_docx_sha256": hashlib.sha256(reference.read_bytes()).hexdigest()},
        "expectations": {
            "table_count": 2,
            "targeted_summary": {"row_count": 1, "gene_order": ["TP53"]},
            "part3": {"gene_section_count": 0, "drug_section_count": 0},
            "reviewed_variant_rows": [
                {
                    "gene": "TP53",
                    "c_hgvs": "c.499C>T",
                    "p_hgvs": "p.Q167*",
                    "vaf": 1.16,
                    "benefit_count": 1,
                    "caution_count": 0,
                }
            ],
        },
    }

    result = validate_historical_golden_docx(contract=contract, docx_path=candidate)

    assert result["status"] == "FAIL"
    assert {error["code"] for error in result["errors"]} >= {
        "REVIEWED_VARIANT_ROW_PRESENT",
        "TARGETED_SUMMARY_REVIEWED_ROW_PRESENT",
    }


def test_historical_contract_blocks_incomplete_targeted_brand_summary(tmp_path) -> None:
    candidate = tmp_path / "candidate.docx"
    _minimal_historical_docx(
        candidate,
        brand_pairs=["奥拉帕利[利普卓]"],
    )
    contract = {
        "schema_version": "1.0",
        "case_alias": "synthetic_case",
        "panel_id": "crc_358_msi",
        "privacy": {"contains_phi": False},
        "source": {},
        "expectations": {
            "table_count": 2,
            "targeted_summary": {"row_count": 1, "gene_order": ["TP53"]},
            "targeted_drug_brand_summary": {
                "ordered_pairs": [
                    "奥拉帕利[利普卓]",
                    "纳武利尤单抗[欧狄沃]",
                ]
            },
            "part3": {"gene_section_count": 0, "drug_section_count": 0},
            "reviewed_variant_rows": [
                {
                    "gene": "TP53",
                    "c_hgvs": "c.499C>T",
                    "p_hgvs": "p.Q167*",
                    "vaf": 1.16,
                    "benefit_count": 1,
                    "caution_count": 0,
                }
            ],
        },
    }

    result = validate_historical_golden_docx(contract=contract, docx_path=candidate)

    assert result["status"] == "FAIL"
    assert {error["code"] for error in result["errors"]} >= {
        "TARGETED_DRUG_BRAND_SUMMARY_ORDER"
    }


def test_historical_contract_blocks_targeted_immune_union_count_regression(
    tmp_path,
) -> None:
    candidate = tmp_path / "candidate.docx"
    _minimal_historical_docx(
        candidate,
        targeted_count=7,
        targeted_or_immune_count=7,
    )
    contract = {
        "schema_version": "1.0",
        "case_alias": "synthetic_case",
        "panel_id": "crc_358_msi",
        "privacy": {"contains_phi": False},
        "source": {},
        "expectations": {
            "table_count": 2,
            "targeted_summary": {"row_count": 1, "gene_order": ["TP53"]},
            "variant_counts": {
                "targeted_drug_related": 7,
                "targeted_or_immune_related": 8,
            },
            "part3": {"gene_section_count": 0, "drug_section_count": 0},
            "reviewed_variant_rows": [],
        },
    }

    result = validate_historical_golden_docx(contract=contract, docx_path=candidate)

    assert result["status"] == "FAIL"
    assert {error["code"] for error in result["errors"]} >= {
        "TARGETED_OR_IMMUNE_RELATED_COUNT_TEXT"
    }


def test_historical_contract_blocks_split_gene_groups_in_part3(
    tmp_path,
) -> None:
    candidate = tmp_path / "candidate.docx"
    _minimal_historical_docx(
        candidate,
        part3_headers=[
            "u APC：c.994C>T，p.R332*；22.03%",
            "u SMAD4：c.1577A>G，p.E526G；20%",
            "u APC：c.4666dup，p.T1556Nfs*3；17.50%",
        ],
    )
    contract = {
        "schema_version": "1.0",
        "case_alias": "synthetic_case",
        "panel_id": "crc_358_msi",
        "privacy": {"contains_phi": False},
        "source": {},
        "expectations": {
            "table_count": 2,
            "targeted_summary": {"row_count": 1, "gene_order": ["TP53"]},
            "part3": {
                "gene_section_count": 3,
                "drug_section_count": 0,
                "group_by_gene": True,
                "gene_order_by_max_vaf_descending": True,
                "within_gene_vaf_descending": True,
            },
            "reviewed_variant_rows": [
                {
                    "gene": "TP53",
                    "c_hgvs": "c.499C>T",
                    "p_hgvs": "p.Q167*",
                    "vaf": 1.16,
                    "benefit_count": 1,
                    "caution_count": 0,
                }
            ],
        },
    }

    result = validate_historical_golden_docx(contract=contract, docx_path=candidate)

    assert result["status"] == "FAIL"
    assert {error["code"] for error in result["errors"]} >= {
        "PART3_GENE_GROUP_CONTIGUITY"
    }
    assert result["checks"]["part3_gene_section_vafs"] == [22.03, 20.0, 17.5]


def test_committed_historical_contract_registry_passes() -> None:
    process = subprocess.run(
        [
            sys.executable,
            str(ROOT / "scripts/check_historical_golden_release.py"),
            "--contracts-only",
        ],
        cwd=ROOT,
        check=False,
        capture_output=True,
        text=True,
    )
    assert process.returncode == 0, process.stdout + process.stderr
    result = json.loads(process.stdout)
    assert result["status"] == "PASS"
    assert result["medical_status"] == "BLOCKED"
    assert result["registry"]["contracts"][0]["medical_uat_status"] == "blocked"


def test_candidate_renderer_fingerprint_is_explicit_and_complete() -> None:
    assert _candidate_renderer_fingerprint(
        {
            "candidate_renderer": {
                "platform": "Linux",
                "engine": "LibreOffice",
                "version": "7.3.7.2",
                "evidence": "production-equivalent visual QA",
            }
        }
    ) == {
        "platform": "Linux",
        "engine": "LibreOffice",
        "version": "7.3.7.2",
        "evidence": "production-equivalent visual QA",
    }
    with pytest.raises(ValueError, match="missing required fields"):
        _candidate_renderer_fingerprint(
            {"candidate_renderer": {"platform": "Linux", "engine": "LibreOffice"}}
        )


def test_release_hash_ignores_macos_filesystem_metadata(tmp_path) -> None:
    controlled = tmp_path / "rules.yaml"
    controlled.write_text("rules: controlled\n", encoding="utf-8")
    expected = _sha256_files([controlled])

    appledouble = tmp_path / "._rules.yaml"
    appledouble.write_bytes(b"macOS AppleDouble metadata")
    finder_marker = tmp_path / ".DS_Store"
    finder_marker.write_bytes(b"Finder metadata")

    assert _sha256_files([controlled, appledouble, finder_marker]) == expected
