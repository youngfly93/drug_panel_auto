"""Regression tests for the de-identified CRC358 historical-golden migration."""

import re
from pathlib import Path

import yaml
from docx import Document
from reportgen.core.field_mapper import FieldMapper
from reportgen.core.report_generator import ReportGenerator
from reportgen.core.template_renderer import TemplateRenderer
from reportgen.knowledge.gene_knowledge import GeneKnowledgeProvider
from reportgen.panels.loader import load_panel_package
from reportgen.rules.targeted_drugs import load_targeted_drug_rule_context

ROOT = Path(__file__).resolve().parents[2]


def _targeted_context(panel_id: str) -> dict:
    package = load_panel_package(panel_id, project_root=ROOT)
    context = load_targeted_drug_rule_context(package)
    assert context is not None
    return context


def _provider() -> GeneKnowledgeProvider:
    package = load_panel_package("crc_358_msi", project_root=ROOT)
    provider = GeneKnowledgeProvider(
        {
            "enabled": True,
            "panel_id": "crc_358_msi",
            "gene_knowledge_db": {
                "enabled": True,
                "path": "does-not-exist.xlsx",
                "reviewed_part3_overlay_paths": (
                    ReportGenerator._resolve_panel_reviewed_part3_overlays(package)
                ),
            },
        }
    )
    assert provider.load(base_path=str(ROOT))
    return provider


def test_crc358_exact_summary_rules_match_historical_contract_and_stay_panel_scoped():
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")
    crc358 = _targeted_context("crc_358_msi")
    crc301 = _targeted_context("crc_301_msi")

    cases = [
        ("PIK3CA", "c.1624G>A", "p.E542K", 15, 2, False),
        ("TP53", "c.499C>T", "p.Q167*", 5, 0, False),
        # FLT3 summary rule already existed as a shared CRC-family rule; this
        # migration adds its exact Part-3 narrative without changing CRC301.
        ("FLT3", "c.2537G>A", "p.G846D", 3, 0, True),
        ("ATR", "c.1291delA", "p.R431Gfs*8", 8, 0, False),
        ("KRAS", "c.34G>T", "p.G12C", 38, 3, False),
    ]
    for gene, c_hgvs, p_hgvs, benefit_count, caution_count, shared in cases:
        result = mapper._lookup_reviewed_variant_override_drugs(
            gene,
            c_hgvs,
            p_hgvs,
            variant_level="Ⅱ类",
            targeted_drug_rules=crc358,
        )
        assert result is not None
        benefit, caution = result
        assert len(benefit.splitlines()) == benefit_count
        assert (0 if caution == "--" else len(caution.splitlines())) == caution_count

        inherited_result = mapper._lookup_reviewed_variant_override_drugs(
            gene,
            c_hgvs,
            p_hgvs,
            variant_level="Ⅱ类",
            targeted_drug_rules=crc301,
        )
        assert (inherited_result is not None) is shared

    kras = mapper._lookup_reviewed_variant_override_drugs(
        "KRAS",
        "c.34G>T",
        "p.G12C",
        variant_level="Ⅱ类",
        targeted_drug_rules=crc358,
    )
    assert kras is not None
    assert kras[0].splitlines()[:3] == [
        "索托拉西布/阿达格拉西布+帕尼单抗/西妥昔单抗（A）",
        "索托拉西布（C）",
        "阿达格拉西布（C）",
    ]
    assert kras[0].splitlines()[-1] == "PD0325901（D）"


def test_exact_part3_overlay_replaces_dynamic_candidates_without_gene_level_leakage():
    provider = _provider()
    variants = [
        {
            "gene": "KRAS",
            "cHGVS": "c.34G>T",
            "pHGVS": "p.G12C",
            "benefit_drugs": "\n".join(["索托拉西布（C）", "司美替尼（C）"]),
            "caution_drugs": "西妥昔单抗（A）\n依维莫司（C）",
        },
        {
            "gene": "TP53",
            "cHGVS": "c.499C>T",
            "pHGVS": "p.Q167*",
            "benefit_drugs": "AZD1775（C）\nEprenetapopt（C）",
            "caution_drugs": "--",
        },
        {
            "gene": "FLT3",
            "cHGVS": "c.2537G>A",
            "pHGVS": "p.G846D",
            "benefit_drugs": "瑞戈非尼（C）\n索拉非尼（C）\n舒尼替尼（C）",
            "caution_drugs": "--",
        },
        {
            "gene": "ATR",
            "cHGVS": "c.1291delA",
            "pHGVS": "p.R431Gfs*8",
            "benefit_drugs": "奥拉帕利（C）\n芦卡帕利（D）",
            "caution_drugs": "--",
        },
    ]
    rows = provider.build_drug_analysis_sections(variants)
    counts = {
        gene: sum(1 for row in rows if row["gene"] == gene)
        for gene in ("KRAS", "TP53", "FLT3", "ATR")
    }
    assert counts == {"KRAS": 8, "TP53": 3, "FLT3": 1, "ATR": 1}
    assert any(row["drug_name"] == "安卓健（Antroquinonol）" for row in rows)
    assert any(
        row["gene"] == "FLT3" and "NCT02029001" in row["clinical"]
        for row in rows
    )

    # Exact TP53 rule must not become a TP53 gene-level rule.
    unrelated = provider.build_drug_analysis_sections(
        [
            {
                "gene": "TP53",
                "cHGVS": "c.821T>A",
                "pHGVS": "p.V274D",
                "benefit_drugs": "AZD1775（C）",
                "caution_drugs": "--",
            }
        ]
    )
    assert not any("p.Q167*" in row.get("relation", "") for row in unrelated)


def test_historical_overlay_is_deidentified_and_explicitly_pending_reconfirmation():
    path = (
        ROOT
        / "panels/crc_358_msi/rules/reviewed_part3_crc358_reviewed_case_a.yaml"
    )
    raw_text = path.read_text(encoding="utf-8")
    data = yaml.safe_load(raw_text)
    assert data["source"]["source_id"] == "crc358_reviewed_case_a"
    assert data["source"]["privacy"].startswith("No patient name")
    assert data["governance"]["defaults"]["drug"][
        "secondary_review_status"
    ] == "pending_report_group_reconfirmation"
    assert data["replace_variant_drug_sections"] is True
    assert len(data["gene_sections"]) == 11
    assert len(data["drug_sections"]) == 18
    assert re.search(r"\bLZ\d{6}\b", raw_text, flags=re.IGNORECASE) is None


def test_reviewed_vertical_merges_are_data_driven_and_idempotent(tmp_path):
    path = tmp_path / "merge-contract.docx"
    doc = Document()
    variant = doc.add_table(rows=5, cols=9)
    variant.rows[0].cells[0].text = "基因名称"
    variant.rows[0].cells[1].text = "基因突变信息"
    variant.rows[1].cells[0].text = "基因名称"
    variant.rows[1].cells[1].text = "转录本号"
    variant.rows[1].cells[2].text = "染色体"
    for row_index, gene in enumerate(("APC", "APC", "KRAS"), start=2):
        variant.rows[row_index].cells[0].text = gene
        variant.rows[row_index].cells[1].text = "NM_TEST"
        variant.rows[row_index].cells[2].text = "1"

    nccn = doc.add_table(rows=5, cols=3)
    for index, text in enumerate(("检测基因", "检测内容", "检测结果")):
        nccn.rows[0].cells[index].text = text
    for row_index, gene in enumerate(("EGFR", "EGFR", "EGFR", "KRAS"), start=1):
        nccn.rows[row_index].cells[0].text = gene
    doc.save(path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._restore_reviewed_vertical_cell_merges(str(path))
    renderer._restore_reviewed_vertical_cell_merges(str(path))

    rendered = Document(path)
    assert rendered.tables[0]._tbl.xml.count("w:vMerge") == 6
    assert rendered.tables[1]._tbl.xml.count("w:vMerge") == 3


def test_active_crc358_template_salutation_is_data_driven(tmp_path):
    template = (
        ROOT
        / "panels/crc_358_msi/templates/crc_358_msi_golden_template_v0.docx"
    )
    renderer = TemplateRenderer(log_level="ERROR")

    female_path = tmp_path / "female.docx"
    female_path.write_bytes(template.read_bytes())
    renderer._render_patient_salutation(
        str(female_path), {"gender": "女", "patient_salutation": "女士"}
    )
    female_text = "\n".join(
        "".join(node.text or "" for node in paragraph.xpath(".//w:t"))
        for paragraph in Document(female_path).element.xpath(".//w:p")
    )
    assert "先生" not in female_text
    assert female_text.count("女士") >= 2

    male_path = tmp_path / "male.docx"
    male_path.write_bytes(template.read_bytes())
    renderer._render_patient_salutation(
        str(male_path), {"gender": "男", "patient_salutation": "先生"}
    )
    male_text = "\n".join(
        "".join(node.text or "" for node in paragraph.xpath(".//w:t"))
        for paragraph in Document(male_path).element.xpath(".//w:p")
    )
    assert "女士" not in male_text
    assert male_text.count("先生") >= 2


def test_repeated_terminal_punctuation_is_normalized(tmp_path):
    path = tmp_path / "punctuation.docx"
    doc = Document()
    doc.add_paragraph("段落。。下一句。")
    doc.save(path)

    renderer = TemplateRenderer(log_level="ERROR")
    renderer._normalize_repeated_terminal_punctuation(str(path))

    assert Document(path).paragraphs[0].text == "段落。下一句。"
