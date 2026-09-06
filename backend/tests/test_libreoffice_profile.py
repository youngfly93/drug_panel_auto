from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from reportgen.utils.libreoffice_profile import (
    FONT_SUBSTITUTION_PROFILE,
    deterministic_font_substitutions,
    font_substitution_fingerprint,
    initialize_libreoffice_profile,
    stabilize_docx_fonts_for_libreoffice_render,
)


def test_platform_font_substitution_targets_are_explicit() -> None:
    darwin = deterministic_font_substitutions(system_name="Darwin", require_available=False)
    linux = deterministic_font_substitutions(system_name="Linux", require_available=False)

    assert darwin["微软雅黑"] == "Arial Unicode MS"
    assert darwin["宋体"] == "Noto Serif CJK SC"
    assert darwin["幼圆"] == "Noto Sans CJK SC"
    assert linux["微软雅黑"] == "Noto Sans CJK SC"
    assert linux["宋体"] == "Noto Serif CJK SC"


def test_profile_initializer_pins_print_font_substitutions(tmp_path: Path) -> None:
    profile = tmp_path / "lo_profile"
    fingerprint = initialize_libreoffice_profile(
        profile,
        system_name="Darwin",
        require_available=False,
    )
    registry = profile / "user" / "registrymodifications.xcu"
    content = registry.read_text(encoding="utf-8")

    assert fingerprint["font_substitution_profile"] == FONT_SUBSTITUTION_PROFILE
    assert re.fullmatch(
        r"[0-9a-f]{64}",
        str(fingerprint["font_substitution_profile_sha256"]),
    )
    assert "微软雅黑" in content
    assert "Arial Unicode MS" in content
    assert "宋体" in content
    assert "Noto Serif CJK SC" in content
    assert "幼圆" in content
    assert "Noto Sans CJK SC" in content
    assert '<prop oor:name="Replacement"' in content
    assert '<prop oor:name="Always"' in content
    assert "<value>true</value>" in content


def test_font_profile_hash_changes_when_mapping_changes(monkeypatch) -> None:
    first = font_substitution_fingerprint(system_name="Linux", require_available=False)
    monkeypatch.setenv("REPORTGEN_LO_SANS_CJK_FONT", "Pinned Sans")
    second = font_substitution_fingerprint(system_name="Linux", require_available=False)

    assert first["font_substitution_profile_sha256"] != second["font_substitution_profile_sha256"]


def test_isolated_profile_can_include_implicit_blank_pages_without_changing_default(tmp_path):
    for include in (False, True):
        profile = tmp_path / str(include)
        initialize_libreoffice_profile(
            profile, require_available=False, include_automatic_blank_pages=include,
        )
        content = (profile / "user/registrymodifications.xcu").read_text()
        assert ("IsSkipEmptyPages" in content) is include
        if include:
            assert '<value>false</value></prop></item>' in content


@pytest.mark.parametrize("listener_available", [True, False])
def test_strict_pdf_option_reaches_listener_and_isolated_fallback(
    tmp_path, monkeypatch, listener_available,
):
    from reportgen.utils import docx_render, uno_pdf

    calls = []
    source = tmp_path / "source.docx"
    Document().save(source)
    monkeypatch.setattr(docx_render, "_libreoffice_profile_mode", lambda: "isolated")

    def listener(_source, output, **kwargs):
        calls.append(("listener", kwargs))
        if listener_available:
            output.write_bytes(b"%PDF-synthetic")
        return listener_available

    def profile(_directory, **kwargs):
        calls.append(("profile", kwargs))

    def convert(_command, **_kwargs):
        (tmp_path / "input.pdf").write_bytes(b"%PDF-synthetic")

    monkeypatch.setattr(uno_pdf, "convert_docx_to_pdf_via_listener", listener)
    monkeypatch.setattr(docx_render, "initialize_libreoffice_profile", profile)
    monkeypatch.setattr(docx_render, "_run_checked", convert)
    result = docx_render._docx_to_pdf(
        soffice="synthetic-soffice", tmp_docx=source, workdir=tmp_path,
        profile_dir=tmp_path / "profile", timeout_seconds=10,
        include_automatic_blank_pages=True,
    )
    assert result == tmp_path / "input.pdf"
    assert all(kwargs["include_automatic_blank_pages"] is True for _, kwargs in calls)
    expected_calls = ["listener"] if listener_available else ["listener", "profile"]
    assert [name for name, _ in calls] == expected_calls


def test_disposable_docx_render_copy_pins_explicit_and_theme_fonts(
    tmp_path: Path,
) -> None:
    source = tmp_path / "render-copy.docx"
    document = Document()
    run = document.add_paragraph().add_run("页眉")
    run.font.name = "幼圆"
    run._r.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "幼圆")
    document.save(source)

    result = stabilize_docx_fonts_for_libreoffice_render(
        source,
        system_name="Darwin",
        require_available=False,
    )

    with ZipFile(source) as archive:
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        theme_xml = etree.fromstring(archive.read("word/theme/theme1.xml"))

    namespaces = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    }
    font_nodes = document_xml.xpath(".//w:rFonts", namespaces=namespaces)
    assert font_nodes
    assert {value for node in font_nodes for value in node.attrib.values() if value} == {
        "Noto Sans CJK SC"
    }
    assert all(
        node.get("typeface") == "Noto Serif CJK SC"
        for node in theme_xml.xpath(".//a:fontScheme//a:ea", namespaces=namespaces)
    )
    assert result["explicit_font_replacement_count"] >= 1
    assert result["theme_font_replacement_count"] == 2
