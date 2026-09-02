# ruff: noqa: E402
"""Independent lung588 engineering-contract regression tests."""

import copy
import hashlib
import io
import json
import re
import shutil
import sys
from pathlib import Path
from zipfile import ZipFile

import yaml
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
from PIL import Image

ROOT = Path(__file__).resolve().parents[2]
BACKEND = ROOT / "backend"
for import_path in (str(ROOT), str(BACKEND)):
    if import_path not in sys.path:
        sys.path.insert(0, import_path)

from reportgen.core.project_detector import ProjectDetector
from reportgen.core.report_generator import (
    ReportGenerator,
    apply_pdl1_display_fields,
    validate_panel_biomarker_contracts,
)
from reportgen.core.review_candidate_contract import (
    load_review_candidate_contract,
    validate_review_candidate_template,
)
from reportgen.core.template_bridge_358 import (
    build_all_variants_for_template,
    build_variants_for_template,
    load_panel_config,
)
from reportgen.core.template_renderer import TemplateRenderer
from reportgen.models.excel_data import ExcelDataSource
from reportgen.models.report_data import ReportData
from reportgen.panels.loader import load_panel_package
from reportgen.panels.validation import validate_panel_package
from reportgen.rules.pdl1 import (
    apply_pdl1_product_display_fields,
    load_pdl1_product_contract,
    validate_pdl1_product_contract,
)

from app.api import batch as batch_api
from app.api.batch import _batch_generation_policy_error
from app.api.report import _controlled_pilot_review_required
from app.services.clinical_info_service import get_clinical_form_schema
from app.services.reportgen_bridge import ReportGenBridge

PANEL_DIR = ROOT / "panels" / "lung_588_pdl1"
TEMPLATE = load_panel_package("lung_588_pdl1", project_root=ROOT).resolve_template_file()
PILOT_ACCEPTANCE = PANEL_DIR / "uat" / "lung588_controlled_pilot_acceptance.yaml"
REVIEW_CANDIDATE_CONTRACT = (
    PANEL_DIR / "review_baselines" / "lung588_historical_review_candidate_v1.yaml"
)
EXPECTED_GENE_SHA256 = "f9e6be05c954a4d3df97f031d453fe1f58ea0689290b11de3c173f4a0edf08f1"


def _gene_contract() -> list[str]:
    raw = yaml.safe_load(
        (PANEL_DIR / "rules" / "knowledge_coverage.yaml").read_text(encoding="utf-8")
    )
    return list(raw["reportable_genes"])


def _visible_template_text() -> str:
    document = Document(TEMPLATE)
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def _excel(tmp_path: Path, rows: list[dict]) -> ExcelDataSource:
    path = tmp_path / "synthetic-lung588.xlsx"
    path.write_bytes(b"synthetic")
    return ExcelDataSource(
        file_path=str(path),
        table_data={"Variations": rows},
        sheet_names=["Variations"],
    )


def _synthetic_pdl1_image(tmp_path: Path) -> Path:
    path = tmp_path / "synthetic-pdl1.png"
    Image.new("RGB", (160, 100), color=(220, 205, 185)).save(path)
    return path


def test_lung588_package_is_independent_and_valid():
    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    report = validate_panel_package("lung_588_pdl1", project_root=ROOT)

    assert report.ok, report.to_dict()
    assert package.panel_id == "lung_588_pdl1"
    assert package.display_name == "肺癌588基因+PD-L1"
    assert package.raw["status"] == "pilot"
    assert package.default_template.status == "pilot"
    assert package.raw["part3_knowledge"]["enabled"] is True
    assert package.raw["part3_knowledge"]["release_mode"] == (
        "report_group_pilot_draft"
    )
    residual_scan = package.raw["part3_knowledge"]["cross_cancer_residual_scan"]
    assert residual_scan["enabled"] is True
    assert residual_scan["severity"] == "warn"
    assert package.raw["release_governance"] == {
        "uat_policy": "uat/lung588_risk_based_release_policy.yaml",
        "report_group_uat_decisions": ("uat/lung588_report_group_uat_decisions.yaml"),
    }
    assert package.raw["review_candidate_contract"] == (
        "review_baselines/lung588_historical_review_candidate_v1.yaml"
    )
    assert package.resolve_template_file() == TEMPLATE.resolve()


def test_lung588_review_candidate_contract_freezes_default_template(tmp_path):
    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    contract = load_review_candidate_contract(REVIEW_CANDIDATE_CONTRACT)
    identity = ReportGenerator._template_identity(package, str(TEMPLATE))

    result = validate_review_candidate_template(
        contract,
        TEMPLATE,
        template_id=str(identity["template_id"]),
        template_version=str(identity["version"]),
        template_status=str(identity["status"]),
    )

    assert result["status"] == "PASS", result["errors"]
    assert identity == {
        "panel_id": "lung_588_pdl1",
        "template_id": "lung_588_pdl1_historical_golden_v1",
            "version": "0.5.1-review.5",
        "status": "pilot",
        "filename": "lung_588_pdl1_historical_golden_v1.docx",
            "sha256": "fe53e15483f7d397ac0507a98a0b9116fce71fe86efad6e8ffdafddf0ac88732",
        "is_default": True,
    }
    assert contract["lifecycle"]["clinical_release_status"] == "blocked"
    assert contract["runtime_policy_revision"] == (
        "2026-09-01-report-group-self-service-pilot"
    )
    assert contract["rendered_output"]["qa"]["allowed_statuses"] == [
        "PASS",
        "WARN",
    ]

    changed = tmp_path / "changed-candidate.docx"
    document = Document(TEMPLATE)
    notice = next(
        paragraph
        for paragraph in document.paragraphs
        if "报告组评审候选稿（非临床交付）" in paragraph.text
    )
    notice.text = "候选标识被意外删除"
    document.save(changed)
    failed = validate_review_candidate_template(contract, changed)
    failed_codes = {item["code"] for item in failed["errors"]}
    assert failed["status"] == "FAIL"
    assert {"TEMPLATE_SHA256", "TEMPLATE_REQUIRED_TEXT"} <= failed_codes


def test_lung588_review_candidate_body_header_uses_stable_tabs():
    document = Document(TEMPLATE)
    seen: set[str] = set()
    matches = []
    for section in document.sections:
        for header in (
            section.header,
            section.first_page_header,
            section.even_page_header,
        ):
            part_name = str(header.part.partname)
            if part_name in seen:
                continue
            seen.add(part_name)
            matches.extend(
                paragraph
                for paragraph in header.paragraphs
                if "姓名：" in (paragraph.text or "")
                and "科技服务人类健康" in (paragraph.text or "")
            )

    assert len(matches) == 1
    paragraph = matches[0]
    assert paragraph.text == "\t姓名：{{ patient_name }}\t科技服务人类健康"
    paragraph_properties = paragraph._p.get_or_add_pPr()
    assert paragraph_properties.find(qn("w:ind")) is None
    assert [
        (tab.get(qn("w:val")), tab.get(qn("w:pos")))
        for tab in paragraph_properties.findall("w:tabs/w:tab", paragraph._p.nsmap)
    ] == [("left", "3400"), ("right", "8200")]
    assert len(paragraph._p.findall("w:r/w:tab", paragraph._p.nsmap)) == 2


def test_lung588_web_bridge_and_preview_resolve_review_candidate_by_default():
    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    bridge = ReportGenBridge(
        config_dir=str(ROOT / "config"),
        template_dir=str(ROOT / "templates"),
    )

    resolved = Path(bridge._resolve_template_path(None, "lung_588_pdl1"))
    explicit = Path(
        bridge._resolve_template_path(
            "lung_588_pdl1_historical_golden_v1",
            "lung_588_pdl1",
        )
    )
    identity = bridge.generator._template_identity(package, str(resolved))

    assert resolved == TEMPLATE.resolve()
    assert explicit == TEMPLATE.resolve()
    assert identity["template_id"] == "lung_588_pdl1_historical_golden_v1"
    assert identity["is_default"] is True


def test_lung588_current_uat_policy_has_no_fixed_case_denominator():
    policy = yaml.safe_load(
        (PANEL_DIR / "uat" / "lung588_risk_based_release_policy.yaml").read_text(encoding="utf-8")
    )
    decisions = yaml.safe_load(
        (PANEL_DIR / "uat" / "lung588_report_group_uat_decisions.yaml").read_text(encoding="utf-8")
    )

    assert policy["status"] == "active_policy"
    assert policy["supersedes"] == ["fixed_ten_real_case_threshold"]
    assert policy["real_case_policy"]["fixed_minimum_real_case_count"] is None
    assert policy["real_case_policy"]["selection"] == (
        "all_registered_real_cases_available_at_release_freeze"
    )
    assert policy["real_case_policy"]["required_review_fraction"] == 0.0
    assert policy["real_case_policy"]["required_pass_fraction"] == 0.0
    assert policy["case_requirements"]["report_group_decision_required"] is False
    assert (
        policy["case_requirements"]["report_group_reviewer_and_date_required"]
        is False
    )
    feedback = policy["product_owner_feedback_source"]
    assert feedback["minimum_entry_count"] == 1
    assert len(feedback["entries"]) == 1
    assert decisions["policy_id"] == policy["policy_id"]
    assert decisions["required_for_pilot_generation"] is False
    assert [item["alias"] for item in decisions["cases"]] == [
        "CASE-LUNG-A",
        "CASE-LUNG-B",
        "CASE-LUNG-C",
    ]
    assert {item["decision"] for item in decisions["cases"]} == {"pending"}


def test_lung588_controlled_pilot_does_not_overclaim_real_uat():
    contract = yaml.safe_load(PILOT_ACCEPTANCE.read_text(encoding="utf-8"))

    assert contract["status"] == "passed_controlled_pilot"
    assert contract["release_tier"] == "controlled_internal_pilot"
    eligibility = contract["eligibility"]
    assert eligibility["required_confirmed_real_ngs_cases"] == 3
    assert eligibility["required_synthetic_boundary_cases"] == 7
    assert eligibility["required_real_case_count_for_active_release"] == 10
    assert eligibility["p0_allowed"] == 0
    assert contract["subject_commit"] == ("b97b8afafc0c417514730c19e557c993c2fe5039")
    assert len(contract["real_case_scope"]["aliases"]) == 3
    assert len(contract["synthetic_boundary_cases"]) == 7
    assert {row["id"] for row in contract["synthetic_boundary_cases"]} == {
        "SYN-L588-01-PDL1-NEGATIVE",
        "SYN-L588-02-PDL1-LOW-LOWER",
        "SYN-L588-03-PDL1-LOW-UPPER",
        "SYN-L588-04-PDL1-HIGH-LOWER",
        "SYN-L588-05-PDL1-HIGH-UPPER",
        "SYN-L588-06-MSIH-TMBH",
        "SYN-L588-07-MULTI-VARIANT",
    }
    assert contract["pdl1_pilot_boundary"]["actual_clone_and_platform_verified"] is False
    assert contract["pdl1_pilot_boundary"]["treatment_inference_allowed"] is False
    assert contract["runtime_boundaries"] == {
        "panel_status": "pilot",
        "batch_generation_enabled": False,
        "part3_knowledge_enabled": False,
        "targeted_drug_rules_enabled": False,
        "external_delivery_without_manual_review": False,
        "active_release_promotion_blocked": True,
    }
    results = contract["results"]
    assert results["identity_match"] is True
    assert results["p0_count"] == 0
    assert results["confirmed_real_ngs"]["passed_case_count"] == 3
    assert results["confirmed_real_ngs"]["verified_case_specific_ihc_source_count"] == 0
    assert results["synthetic_boundary"]["passed_case_count"] == 7
    assert results["synthetic_boundary"]["runtime_targeted_drug_count"] == 0
    assert results["synthetic_boundary"]["runtime_part3_section_count"] == 0
    assert contract["release_decision"]["controlled_pilot_status"] == "PASS"
    assert contract["release_decision"]["active_release_status"] == "BLOCKED"
    assert _controlled_pilot_review_required(
        "lung_588_pdl1",
        "draft",
    )
    assert not _controlled_pilot_review_required(
        "lung_588_pdl1",
        "reviewed",
    )
    assert not _controlled_pilot_review_required(
        "crc_358_msi",
        "draft",
    )


def test_lung588_gene_denominator_is_exact_and_ordered():
    genes = _gene_contract()

    assert len(genes) == 588
    assert len(set(genes)) == 588
    assert hashlib.sha256("\n".join(genes).encode("utf-8")).hexdigest() == EXPECTED_GENE_SHA256

    document = Document(TEMPLATE)
    tables = [
        table for table in document.tables if "Gene List for MLseq (n=588)" in table.cell(0, 0).text
    ]
    assert len(tables) == 1
    rendered = [
        cell.text.strip() for row in tables[0].rows[1:] for cell in row.cells if cell.text.strip()
    ]
    assert rendered == genes
    assert all(row.height_rule == WD_ROW_HEIGHT_RULE.EXACTLY for row in tables[0].rows[1:])
    assert max(row.height.cm for row in tables[0].rows[1:]) <= 0.72


def test_lung588_template_is_hardened_and_byte_reproducible(tmp_path):
    visible = _visible_template_text()
    document = Document(TEMPLATE)
    assert "Gene List for MLseq (n=588)" in visible
    assert "报告组评审候选稿（非临床交付）" in visible
    assert "固定附录沿用历史终版版式" in visible
    assert "肺癌329" not in visible
    assert "n=329" not in visible
    assert "__PART3_MARKER__" in visible
    assert "__PDL1_CASE_IMAGE__" in visible
    assert "图1. 免疫组化：PD-L1" in visible
    assert "本表仅展示当前面板已输出的精确事件结果" in visible
    assert "未复核内容不得用于临床决策" in visible
    assert "肺癌专属治疗知识和事件级药物规则当前未启用" not in visible
    assert "本病例未提供可追溯的PD-L1免疫组化图像" not in visible
    assert "原始记录未提供抗体克隆" not in visible
    assert visible.count("肿瘤具有异质性") == 1
    assert len(document.sections) == 5
    assert len(document.tables) == 69
    with ZipFile(TEMPLATE) as archive:
        document_xml = archive.read("word/document.xml")
    xml_visible = re.sub(rb"<[^>]+>", b"", document_xml).decode("utf-8", "ignore")
    assert "肺癌588基因+PD-L1检测项目" in "".join(xml_visible.split())
    assert document_xml.count(b"{%tr for row in drug_") == 30

    template = DocxTemplate(str(TEMPLATE))
    context = {
        "patient_name": "SYNTHETIC_PATIENT",
        "sample_id": "SYNTHETIC_CASE",
            "pdl1_tps": "50",
            "pdl1_cps": "52",
            "pdl1_result": "阳性（高表达）",
            "pdl1_tps_display": "50%",
            "pdl1_cps_display": "52",
            "pdl1_result_display": "阳性（高表达）",
        "variants_2_1": [],
        "targeted_drug_tips": [],
        "nccn_results": [],
        "immune_positive_results": [],
        "immune_negative_results": [],
        "immune_hyperprogression_results": [],
    }
    template.render(context)
    output = io.BytesIO()
    template.save(output)
    with ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8", "ignore")
    rendered = re.sub(r"<[^>]+>", "", xml)
    assert "SYNTHETIC_PATIENT" in rendered
    assert "阳性（高表达）" in rendered
    assert "__PDL1_CASE_IMAGE__" in rendered

    styled = tmp_path / "styled-lung588.docx"
    shutil.copy2(TEMPLATE, styled)
    TemplateRenderer(log_level="ERROR")._compact_gene_list_tables(
        str(styled),
        {
            "panel_style": {
                "gene_list_table": {
                    "row_height_cm": 0.72,
                    "header_row_height_cm": 0.88,
                    "row_height_rule": "exact",
                    "body_font_size": 10,
                }
            }
        },
    )
    styled_document = Document(styled)
    styled_table = next(
        table
        for table in styled_document.tables
        if "Gene List for MLseq (n=588)" in table.cell(0, 0).text
    )
    assert all(row.height_rule == WD_ROW_HEIGHT_RULE.EXACTLY for row in styled_table.rows)
    assert styled_table.rows[0].height.cm <= 0.89
    assert max(row.height.cm for row in styled_table.rows[1:]) <= 0.73


def test_lung588_missing_optional_pdl1_provenance_uses_review_placeholders():
    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    contract = load_pdl1_product_contract(package)
    data = ReportData()

    apply_pdl1_product_display_fields(data, contract)

    assert "待报告组核对后补充" in data.get_field("pdl1_assay_provenance")
    assert "待补充" in data.get_field("pdl1_source_provenance")
    assert data.get_field("pdl1_assay_profile_id") is None


def test_lung588_template_does_not_reuse_scaffold_pdl1_image():
    def compact(value: object) -> str:
        return "".join(str(value or "").split())

    generated = Document(TEMPLATE)
    start = next(
        paragraph
        for paragraph in generated.paragraphs
        if compact(paragraph.text) == compact("3.2 PD-L1表达检测结果")
    )
    end = next(
        paragraph
        for paragraph in generated.paragraphs
        if compact(paragraph.text) == compact("3.3微卫星不稳定性（MSI）检测结果")
    )
    children = list(generated.element.body.iterchildren())
    block = children[children.index(start._p) + 1 : children.index(end._p)]
    assert not [
        child
        for child in block
        if not compact("".join(node.text or "" for node in child.iter(qn("w:t"))))
        and any(True for _ in child.iter(qn("w:drawing")))
    ]

    with ZipFile(TEMPLATE) as archive:
        names = set(archive.namelist())
        media_parts = {name for name in names if name.startswith("word/media/")}
        relationship_targets = {
            "word/media/" + target.decode("utf-8")
            for name in names
            if name.endswith(".rels")
            for target in re.findall(
                rb'Target="(?:\.\./)?media/([^"]+)"',
                archive.read(name),
            )
        }
    assert media_parts == relationship_targets


def test_pdl1_case_image_processor_is_idempotent(tmp_path):
    output = tmp_path / "pdl1-image-idempotent.docx"
    document = Document()
    document.add_paragraph("__PDL1_CASE_IMAGE__")
    document.save(output)
    image_path = _synthetic_pdl1_image(tmp_path)
    context = {
        "project_type": "lung_588_pdl1",
        "pdl1_image_path": str(image_path),
    }
    renderer = TemplateRenderer(log_level="ERROR")

    renderer._run_post_render_processors(
        str(output),
        context,
        str(output),
        processor_names=["pdl1_case_image"],
    )
    assert renderer.last_processor_report[0]["status"] == "OK"
    first_digest = hashlib.sha256(output.read_bytes()).hexdigest()

    renderer._run_post_render_processors(
        str(output),
        context,
        str(output),
        processor_names=["pdl1_case_image"],
    )
    assert renderer.last_processor_report[0]["status"] == "OK"
    assert hashlib.sha256(output.read_bytes()).hexdigest() == first_digest
    rendered = Document(output)
    assert len(rendered.inline_shapes) == 1
    assert not any("__PDL1_CASE_IMAGE__" in p.text for p in rendered.paragraphs)


def test_pdl1_case_image_processor_renders_idempotent_draft_notice(tmp_path):
    output = tmp_path / "pdl1-image-missing.docx"
    document = Document()
    document.add_paragraph("__PDL1_CASE_IMAGE__")
    document.save(output)
    context = {
        "project_type": "lung_588_pdl1",
        "pdl1_image_missing_notice": "未提供合成病例PD-L1图片；待报告组审核。",
    }
    renderer = TemplateRenderer(log_level="ERROR")

    renderer._run_post_render_processors(
        str(output),
        context,
        str(output),
        processor_names=["pdl1_case_image"],
    )
    first_digest = hashlib.sha256(output.read_bytes()).hexdigest()
    renderer._run_post_render_processors(
        str(output),
        context,
        str(output),
        processor_names=["pdl1_case_image"],
    )

    assert renderer.last_processor_report[0]["status"] == "OK"
    assert hashlib.sha256(output.read_bytes()).hexdigest() == first_digest
    rendered = Document(output)
    assert len(rendered.inline_shapes) == 0
    assert [p.text for p in rendered.paragraphs] == ["未提供合成病例PD-L1图片；待报告组审核。"]


def test_part3_disabled_policy_renders_notice_without_shared_knowledge(tmp_path):
    output = tmp_path / "part3-disabled.docx"
    document = Document()
    document.add_paragraph("__PART3_MARKER__")
    document.save(output)

    notice = "肺癌588第三部分尚未完成独立知识二审，当前工程草案不输出患者级解释。"
    TemplateRenderer(log_level="ERROR")._render_part3_formatted(
        str(output),
        {
            "part3_disabled_notice": notice,
            "part3_knowledge_status": "disabled",
            "total_variants_count": 8,
            "gene_knowledge_sections": [
                {"header": "SHOULD_NOT_RENDER", "intro": "SHOULD_NOT_RENDER"}
            ],
        },
    )

    rendered = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    assert notice in rendered
    assert "SHOULD_NOT_RENDER" not in rendered
    assert "在本次检测范围内" not in rendered
    assert "__PART3_MARKER__" not in rendered


def test_lung588_project_detection_requires_trusted_product_text():
    detector = ProjectDetector(config_dir=str(ROOT / "config"), log_level="ERROR")

    detected = detector.detect("CASE-肺癌588基因+PD-L1.xlsx")
    unknown = detector.detect("CASE-LUNG-ONLY.xlsx")

    assert detected["detected"] is True
    assert detected["project_type"] == "lung_588_pdl1"
    assert unknown["detected"] is False
    assert unknown["project_type"] is None


def test_lung588_project_detection_accepts_reviewed_structural_fingerprint(tmp_path):
    path = tmp_path / "CASE-LUNG-B.xlsx"
    path.write_bytes(b"synthetic")
    excel_data = ExcelDataSource(
        file_path=str(path),
        table_data={
            "Variations": [
                {
                    "ExistIn552": "Ⅱ类",
                    "ExistInsmall588": "Ⅱ类",
                    "Gene_Symbol": "TP53",
                    "cHGVS": "c.1A>T",
                }
            ],
            "Hereditary_tumor": [
                {
                    "ExistInsmall588": "否",
                    "ExistIn178": "否",
                    "Gene": "TP53",
                }
            ],
        },
        sheet_names=["Variations", "Hereditary_tumor"],
    )

    result = ProjectDetector(
        config_dir=str(ROOT / "config"),
        log_level="ERROR",
    ).detect(str(path), excel_data=excel_data)

    assert result["detected"] is True
    assert result["project_type"] == "lung_588_pdl1"
    assert result["confidence"] == 1.0
    assert any(
        "lung588_result_workbook_v1" in detail
        for match in result["match_details"]
        if match["type"] == "lung_588_pdl1"
        for detail in match["details"]
    )


def test_lung588_project_detection_rejects_partial_generic_panel_column(tmp_path):
    path = tmp_path / "CASE-UNKNOWN.xlsx"
    path.write_bytes(b"synthetic")
    excel_data = ExcelDataSource(
        file_path=str(path),
        table_data={
            "Variations": [
                {
                    "ExistInsmall588": "Ⅱ类",
                    "Gene_Symbol": "TP53",
                    "cHGVS": "c.1A>T",
                }
            ],
        },
        sheet_names=["Variations"],
    )

    result = ProjectDetector(
        config_dir=str(ROOT / "config"),
        log_level="ERROR",
    ).detect(str(path), excel_data=excel_data)

    assert result["detected"] is False
    assert result["project_type"] is None


def test_lung588_structural_fingerprint_exposes_conflicting_trusted_filename(tmp_path):
    path = tmp_path / "CASE-结直肠癌358基因+MSI.xlsx"
    path.write_bytes(b"synthetic")
    excel_data = ExcelDataSource(
        file_path=str(path),
        table_data={
            "Variations": [
                {
                    "ExistIn552": "Ⅱ类",
                    "ExistInsmall588": "Ⅱ类",
                    "Gene_Symbol": "TP53",
                    "cHGVS": "c.1A>T",
                }
            ],
            "Hereditary_tumor": [
                {
                    "ExistInsmall588": "否",
                    "ExistIn178": "否",
                    "Gene": "TP53",
                }
            ],
        },
        sheet_names=["Variations", "Hereditary_tumor"],
    )

    result = ProjectDetector(
        config_dir=str(ROOT / "config"),
        log_level="ERROR",
    ).detect(str(path), excel_data=excel_data)

    assert set(result["identity_conflicts"]) == {
        "crc_358_msi",
        "lung_588_pdl1",
    }


def test_lung588_explicit_classes_drive_variant_filter(tmp_path):
    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    config = load_panel_config(
        base_path=str(ROOT),
        panel_package=package,
    )
    rows = [
        {
            "ExistIn552": "Ⅰ类",
            "Gene_Symbol": "BRAF",
            "cHGVS": "c.1799T>A",
            "pHGVS_S": "p.V600E",
            "Freq(%)": 40,
        },
        {
            "ExistIn552": "Ⅱ类",
            "Gene_Symbol": "TP53",
            "cHGVS": "c.734G>A",
            "pHGVS_S": "p.G245D",
            "Freq(%)": 30,
        },
        {
            "ExistIn552": "Ⅲ类",
            "Gene_Symbol": "ATM",
            "cHGVS": "c.1236-2A>T",
            "pHGVS_S": "*",
            "Freq(%)": 20,
        },
        {
            "ExistIn552": 1,
            "Gene_Symbol": "EGFR",
            "cHGVS": "c.2573T>G",
            "pHGVS_S": "p.L858R",
            "Freq(%)": 10,
        },
        {
            "ExistIn552": "Ⅱ类",
            "Gene_Symbol": "KRAS",
            "cHGVS": "not-a-coding-hgvs",
            "pHGVS_S": "p.G12D",
            "Freq(%)": 5,
        },
    ]
    excel = _excel(tmp_path, rows)

    main = build_variants_for_template(
        excel,
        filter_column="ExistIn552",
        panel_config=config,
    )
    all_rows = build_all_variants_for_template(
        excel,
        filter_column="ExistIn552",
        panel_config=config,
    )

    assert [row["gene"] for row in main] == ["BRAF", "TP53"]
    assert [row["gene"] for row in all_rows] == ["BRAF", "TP53", "ATM"]
    assert config.variant_filter_columns == ["ExistIn552"]
    assert config.variant_filter_values == ["Ⅰ类", "Ⅱ类", "Ⅲ类"]
    assert len(config.crc_important_genes) == 588
    assert config.nccn_result_rows == []
    assert len(config.immune_positive_rows) == 15
    assert len(config.immune_negative_rows) == 12
    assert len(config.immune_hyperprogression_rows) == 8
    assert config.immune_module_enabled is True


def test_lung588_pdl1_form_is_project_scoped_and_optional_for_draft():
    lung = get_clinical_form_schema("lung_588_pdl1")
    crc = get_clinical_form_schema("crc_358_msi")
    lung_fields = {field.key: field for group in lung.groups for field in group.fields}
    crc_fields = {field.key for group in crc.groups for field in group.fields}

    pdl1_fields = {
        "pdl1_tps",
        "pdl1_cps",
        "pdl1_result",
        "pdl1_image_path",
    }
    source_fields = {
        "pdl1_assay_profile_id",
        "pdl1_source_record_id",
        "pdl1_source_record_date",
        "pdl1_specimen_id",
    }
    hidden_derived = {
        "pdl1_image_disposition",
    }
    assert pdl1_fields | source_fields <= set(lung_fields)
    assert all(not lung_fields[key].required for key in pdl1_fields | source_fields)
    assert lung_fields["pdl1_result"].ui.options == [
        "阳性（高表达）",
        "阳性（低表达）",
        "阴性",
    ]
    assert not hidden_derived & set(lung_fields)
    assert lung_fields["pdl1_assay_profile_id"].ui.options == [
        "legacy_unspecified_ihc_transcription_v1"
    ]
    assert lung_fields["pdl1_image_path"].ui.component == "pdl1-image-upload"
    assert not pdl1_fields & crc_fields

    treatment_context = {
        "lung_histology": [
            "非小细胞肺癌",
            "小细胞肺癌",
            "其他",
            "未明确",
        ],
        "disease_extent": [
            "可切除早期",
            "不可切除局部晚期",
            "转移性",
            "未明确",
        ],
        "prior_systemic_therapy": ["已接受", "未接受", "未明确"],
        "companion_diagnostic_status": [
            "已确认符合",
            "待确认",
            "不符合",
        ],
    }
    assert (
        next(group.label for group in lung.groups if group.id == "treatment_context")
        == "靶向用药适应证上下文（不属于PD-L1检测字段）"
    )
    for key, options in treatment_context.items():
        assert lung_fields[key].required is False
        assert lung_fields[key].ui.component == "select"
        assert lung_fields[key].ui.options == options
    assert not set(treatment_context) & crc_fields


def test_lung588_batch_is_enabled_with_case_isolated_pdl1():
    error = _batch_generation_policy_error("lung_588_pdl1")

    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    assert error is None
    assert package.raw["batch_generation"]["enabled"] is True
    assert package.raw["batch_generation"]["clinical_info_mode"] == (
        "case_isolated_optional"
    )
    assert _batch_generation_policy_error("crc_358_msi") is None


def test_lung588_batch_strips_shared_pdl1_values_and_aliases():
    isolated = batch_api._isolate_batch_case_fields(
        {
            "patient_name": "SYNTHETIC_PATIENT",
            "pdl1_tps": 50,
            "PD-L1 CPS": 52,
            "pdl1_result": "阳性（高表达）",
            "PD-L1病例图片": "/tmp/other-case.png",
            "pdl1_source_record_id": "OTHER-CASE-IHC",
        },
        "lung_588_pdl1",
    )

    assert isolated == {"patient_name": "SYNTHETIC_PATIENT"}


def test_lung588_batch_marks_only_a_missing_patient_name_as_not_provided():
    missing = batch_api._apply_batch_missing_display_defaults(
        {"sample_id": "SYNTHETIC_L588"},
        "lung_588_pdl1",
    )
    existing = batch_api._apply_batch_missing_display_defaults(
        {"patient_name": "SYNTHETIC_PATIENT", "sample_id": "SYNTHETIC_L588"},
        "lung_588_pdl1",
    )
    unrelated = batch_api._apply_batch_missing_display_defaults(
        {"sample_id": "SYNTHETIC_CRC"},
        "crc_358_msi",
    )

    assert missing == {"sample_id": "SYNTHETIC_L588", "patient_name": "未提供"}
    assert existing["patient_name"] == "SYNTHETIC_PATIENT"
    assert unrelated == {"sample_id": "SYNTHETIC_CRC"}


def test_lung588_pdl1_contract_fails_closed_on_missing_range_and_classification():
    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    contracts = package.input_contract["biomarkers"]

    missing = ReportData()
    assert {
        failure["field"] for failure in validate_panel_biomarker_contracts(missing, contracts)
    } == {
        "tmb_value",
        "msi_status",
    }

    apply_pdl1_display_fields(missing)
    assert "先生成NGS报告草稿" in missing.get_field("pdl1_table_interpretation")

    def add_pdl1_provenance(data: ReportData) -> None:
        data.set_field(
            "pdl1_assay_profile_id",
            "nsclc_22c3_pharmdx_tps_v1",
        )
        data.set_field("pdl1_source_record_id", "SYNTHETIC-IHC-001")
        data.set_field("pdl1_source_record_date", "2026-07-23")
        data.set_field("pdl1_specimen_id", "SYNTHETIC-SPECIMEN-001")
        data.set_field(
            "pdl1_image_disposition",
            "病例专属图像（报告展示）",
        )
        data.set_field("pdl1_image_path", "2026-08-09/synthetic.png")

    invalid = ReportData()
    invalid.set_field("tmb_value", 7.5)
    invalid.set_field("msi_status", "MSS")
    invalid.set_field("pdl1_tps", 5)
    invalid.set_field("pdl1_cps", 101)
    invalid.set_field("pdl1_result", "阳性（高表达）")
    add_pdl1_provenance(invalid)
    reasons = {
        (failure["field"], failure["reason"])
        for failure in validate_panel_biomarker_contracts(invalid, contracts)
    }
    assert ("pdl1_cps", "above_maximum") in reasons
    assert not any(field == "pdl1_result" for field, _ in reasons)

    negative = ReportData()
    negative.set_field("tmb_value", 0)
    negative.set_field("msi_status", "MSS")
    negative.set_field("pdl1_tps", 0)
    negative.set_field("pdl1_cps", 0)
    negative.set_field("pdl1_result", "阴性")
    add_pdl1_provenance(negative)
    assert validate_panel_biomarker_contracts(negative, contracts) == []

    negative.set_field("pdl1_cps", float("nan"))
    nan_failures = validate_panel_biomarker_contracts(negative, contracts)
    assert ("pdl1_cps", "not_numeric") in {
        (failure["field"], failure["reason"]) for failure in nan_failures
    }

    valid = ReportData()
    valid.set_field("tmb_value", 6.3)
    valid.set_field("msi_status", "MSS")
    valid.set_field("pdl1_tps", 5)
    valid.set_field("pdl1_cps", 6)
    valid.set_field("pdl1_result", "阳性（低表达）")
    add_pdl1_provenance(valid)
    assert validate_panel_biomarker_contracts(valid, contracts) == []
    apply_pdl1_display_fields(valid)
    interpretation = valid.get_field("pdl1_table_interpretation")
    assert valid.get_field("pdl1_tps") == "5"
    assert valid.get_field("pdl1_cps") == "6"
    assert "TPS 5%" in interpretation
    assert "CPS 6" in interpretation
    assert "推荐" not in interpretation
    assert "帕博利珠单抗" not in interpretation

    valid.set_field("pdl1_tps", 5.25)
    valid.set_field("pdl1_cps", "6.5")
    apply_pdl1_display_fields(valid)
    assert valid.get_field("pdl1_tps") == "5.25"
    assert valid.get_field("pdl1_cps") == "6.5"


def test_lung588_pdl1_product_profiles_are_traceable_and_fail_closed():
    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    contract = load_pdl1_product_contract(package)

    assert contract["status"] == "pilot"
    governance = contract["governance"]
    assert governance["runtime_enabled"] is True
    assert governance["report_text_allowed"] is True
    assert governance["promotion_blocked"] is True
    assert governance["secondary_review_status"] == ("product_owner_authorized_controlled_pilot")
    assert governance["runtime_mode"] == "controlled_pilot_transcription"
    assert governance["treatment_inference_allowed"] is False
    assert contract["runtime_profiles"] == ["legacy_unspecified_ihc_transcription_v1"]
    assert contract["input_provenance"]["ngs_excel_is_pdl1_source"] is False
    assert contract["image_policy"]["static_template_patient_image_allowed"] is False
    assert contract["image_policy"]["case_specific_image_pipeline_implemented"] is True
    assert contract["image_policy"]["allowed_runtime_dispositions"] == ["病例专属图像（报告展示）"]

    profiles = contract["candidate_profiles"]
    assert len(profiles) == 2
    pilot_profile = next(
        row for row in profiles if row["profile_id"] == "legacy_unspecified_ihc_transcription_v1"
    )
    assert pilot_profile["runtime_eligible"] is True
    assert pilot_profile["report_text_allowed"] is True
    assert pilot_profile["validation_mode"] == "verbatim_source_record"
    assert pilot_profile["treatment_inference_allowed"] is False
    assert pilot_profile["antibody_clone"] == "原始记录未提供"
    assert pilot_profile["staining_platform"] == "原始记录未提供"
    assert pilot_profile["report_classification_notice"].startswith("4、")
    assert "不使用通用TPS/CPS阈值" in pilot_profile["report_classification_notice"]

    profile = next(row for row in profiles if row["profile_id"] == "nsclc_22c3_pharmdx_tps_v1")
    assert profile["profile_id"] == "nsclc_22c3_pharmdx_tps_v1"
    assert profile["runtime_eligible"] is False
    assert profile["report_text_allowed"] is False
    assert profile["secondary_review_status"] == ("pending_report_group_review")
    assert profile["antibody_clone"] == "22C3"
    assert profile["staining_platform"] == "Autostainer Link 48"
    assert profile["primary_scoring_method"] == "TPS"
    assert all(
        source["supports"] and source["does_not_support"] for source in profile["source_refs"]
    )

    data = ReportData()
    data.set_field("pdl1_assay_profile_id", profile["profile_id"])
    data.set_field("pdl1_source_record_id", "SYNTHETIC-IHC-001")
    data.set_field("pdl1_source_record_date", "2026-07-23")
    data.set_field("pdl1_specimen_id", "SYNTHETIC-SPECIMEN-001")
    data.set_field("pdl1_tps", 5)
    data.set_field("pdl1_cps", 6)
    data.set_field("pdl1_result", "阳性（低表达）")
    data.set_field(
        "pdl1_image_disposition",
        "病例专属图像（报告展示）",
    )
    data.set_field("pdl1_image_path", "2026-08-09/synthetic.png")
    reasons = {failure["reason"] for failure in validate_pdl1_product_contract(data, contract)}
    assert reasons == {"assay_profile_not_runtime_approved"}

    data.set_field("pdl1_result", "阳性（高表达）")
    reasons = {failure["reason"] for failure in validate_pdl1_product_contract(data, contract)}
    assert "classification_inconsistent_with_assay_profile" in reasons
    data.set_field("pdl1_result", "阳性（低表达）")

    approved = copy.deepcopy(contract)
    approved["governance"].update(
        {
            "runtime_enabled": True,
            "report_text_allowed": True,
            "promotion_blocked": False,
            "secondary_review_status": "approved_by_report_group",
        }
    )
    approved["runtime_profiles"] = [profile["profile_id"]]
    approved_profile = next(
        row for row in approved["candidate_profiles"] if row["profile_id"] == profile["profile_id"]
    )
    approved_profile.update(
        {
            "runtime_eligible": True,
            "report_text_allowed": True,
            "secondary_review_status": "approved_by_report_group",
        }
    )
    assert validate_pdl1_product_contract(data, approved) == []

    data.set_field("pdl1_tps", 0)
    data.set_field("pdl1_cps", 0)
    data.set_field("pdl1_result", "阴性")
    assert validate_pdl1_product_contract(data, approved) == []

    data.set_field("pdl1_tps", 5)
    data.set_field("pdl1_cps", 6)
    data.set_field("pdl1_result", "阳性（低表达）")
    apply_pdl1_product_display_fields(data, approved)
    assert "22C3" in data.get_field("pdl1_assay_provenance")
    assert "Autostainer Link 48" in data.get_field("pdl1_assay_provenance")
    assert "SYNTHETIC-IHC-001" in data.get_field("pdl1_source_provenance")
    assert "SYNTHETIC-SPECIMEN-001" in data.get_field("pdl1_source_provenance")

    pilot_data = ReportData()
    pilot_data.set_field(
        "pdl1_assay_profile_id",
        pilot_profile["profile_id"],
    )
    pilot_data.set_field("pdl1_source_record_id", "SYNTHETIC-IHC-002")
    pilot_data.set_field("pdl1_source_record_date", "2026-07-24")
    pilot_data.set_field("pdl1_specimen_id", "SYNTHETIC-SPECIMEN-002")
    pilot_data.set_field("pdl1_tps", 5)
    pilot_data.set_field("pdl1_cps", 6)
    # The unspecified legacy profile must transcribe the source-record
    # category instead of pretending that a known assay threshold applies.
    pilot_data.set_field("pdl1_result", "阳性（高表达）")
    pilot_data.set_field(
        "pdl1_image_disposition",
        "病例专属图像（报告展示）",
    )
    pilot_data.set_field("pdl1_image_path", "2026-08-09/synthetic.png")
    assert validate_pdl1_product_contract(pilot_data, contract) == []
    apply_pdl1_product_display_fields(pilot_data, contract)
    provenance = pilot_data.get_field("pdl1_assay_provenance")
    assert "原始记录未提供" in provenance
    assert "不据此推导" in provenance
    assert "22C3" not in provenance
    classification_notice = pilot_data.get_field("pdl1_classification_notice")
    assert classification_notice.startswith("4、")
    assert "不使用通用TPS/CPS阈值" in classification_notice
    assert "TPS<1%" not in classification_notice

    pilot_data.set_field("pdl1_result", "未经允许的分层")
    assert {
        failure["reason"]
        for failure in validate_pdl1_product_contract(
            pilot_data,
            contract,
        )
    } == {"source_record_classification_not_allowed"}


def test_lung588_generation_surfaces_pending_pdl1_profile_in_draft(tmp_path):
    xlsx_path = tmp_path / "SYNTHETIC-LUNG588.xlsx"
    xlsx_path.write_bytes(b"synthetic")
    image_path = _synthetic_pdl1_image(tmp_path)
    excel_data = ExcelDataSource(
        file_path=str(xlsx_path),
        single_values={
            "患者姓名": "SYNTHETIC_PATIENT",
            "样本编号": "SYNTHETIC_CASE",
            "TMB": 5,
            "MSI状态": "MSS",
            "PD-L1 TPS": 5,
            "PD-L1 CPS": 6,
            "PD-L1结果": "阳性（低表达）",
            "PD-L1检测方案": "nsclc_22c3_pharmdx_tps_v1",
            "PD-L1原始记录编号": "SYNTHETIC-IHC-001",
            "PD-L1原始记录日期": "2026-07-23",
            "PD-L1检测标本标识": "SYNTHETIC-SPECIMEN-001",
            "PD-L1图像处置": "病例专属图像（报告展示）",
            "PD-L1病例图片": str(image_path),
            "肺癌病理类型": "非小细胞肺癌",
            "疾病范围": "转移性",
            "既往系统治疗": "已接受",
            "伴随诊断适配状态": "已确认符合",
        },
        table_data={
            "Variations": [
                {
                    "ExistIn552": "Ⅱ类",
                    "Gene_Symbol": "TP53",
                    "Transcript": "NM_000546.6",
                    "cHGVS": "c.734G>A",
                    "pHGVS_S": "p.G245D",
                    "Freq(%)": 30,
                }
            ]
        },
        sheet_names=["Variations", "TMB", "Msisensor"],
    )
    output_dir = tmp_path / "out"

    result = ReportGenerator(
        config_dir=str(ROOT / "config"),
        log_level="ERROR",
    ).generate(
        excel_file=str(xlsx_path),
        excel_data=excel_data,
        template_file=str(TEMPLATE),
        output_dir=str(output_dir),
        project_type="lung_588_pdl1",
    )

    assert result["success"] is True, result["errors"]
    assert result["qa_status"] == "WARN"
    qa = json.loads(Path(result["qa_report_file"]).read_text(encoding="utf-8"))
    residual_check = qa["checks"]["part3_cross_cancer_residuals"]
    assert residual_check["status"] == "PASS"
    assert residual_check["matched_terms"] == []
    suppression_check = qa["checks"]["part3_cross_cancer_suppression"]
    assert suppression_check["status"] == "WARN"
    assert suppression_check["suppressed_field_count"] > 0
    assert any(
        issue["code"] == "PART3_CROSS_CANCER_SUPPRESSION"
        for issue in qa["issues"]
    )
    assert any(
        issue["code"] == "PART3_CROSS_CANCER_FIELDS_SUPPRESSED"
        for stage in result["stage_results"]
        for issue in stage.get("issues") or []
    )
    assert any(
        issue["code"] == "PANEL_PDL1_PRODUCT_CONTRACT_WARNING"
        for stage in result["stage_results"]
        for issue in stage.get("issues") or []
    )
    assert Path(result["output_file"]).is_file()


def test_lung588_generation_without_pdl1_or_image_creates_review_draft(tmp_path):
    xlsx_path = tmp_path / "SYNTHETIC-LUNG588-NGS-ONLY.xlsx"
    xlsx_path.write_bytes(b"synthetic")
    excel_data = ExcelDataSource(
        file_path=str(xlsx_path),
        single_values={
            "患者姓名": "SYNTHETIC_PATIENT",
            "样本编号": "SYNTHETIC_NGS_ONLY",
            "检测项目": "肺癌588基因+PD-L1",
            "报告日期": "2026-08-10",
            "TMB": 5,
            "MSI状态": "MSS",
        },
        table_data={
            "Variations": [
                {
                    "ExistIn552": "Ⅱ类",
                    "Gene_Symbol": "TP53",
                    "Transcript": "NM_000546.6",
                    "cHGVS": "c.734G>A",
                    "pHGVS_S": "p.G245D",
                    "Freq(%)": 30,
                }
            ]
        },
        sheet_names=["Variations", "TMB", "Msisensor"],
    )

    result = ReportGenerator(
        config_dir=str(ROOT / "config"),
        log_level="ERROR",
    ).generate(
        excel_file=str(xlsx_path),
        excel_data=excel_data,
        template_file=str(TEMPLATE),
        output_dir=str(tmp_path / "out"),
        project_type="lung_588_pdl1",
    )

    assert result["success"] is True, result["errors"]
    assert result["qa_status"] == "WARN"
    assert result["template_identity"]["template_id"] == ("lung_588_pdl1_historical_golden_v1")
    assert result["template_identity"]["sha256"] == (
        "fe53e15483f7d397ac0507a98a0b9116fce71fe86efad6e8ffdafddf0ac88732"
    )
    assert result["report_summary"]["template"]["id"] == ("lung_588_pdl1_historical_golden_v1")
    assert any("PD-L1逐病例结果" in warning for warning in result["warnings"])
    assert any(
        issue["code"] == "PANEL_PDL1_PRODUCT_CONTRACT_WARNING"
        for stage in result["stage_results"]
        for issue in stage.get("issues") or []
    )
    output = Path(result["output_file"])
    rendered = Document(output)
    visible = "\n".join(
        [
            *(paragraph.text for paragraph in rendered.paragraphs),
            *(cell.text for table in rendered.tables for row in table.rows for cell in row.cells),
        ]
    )
    assert "先生成NGS报告草稿供报告解读组审核" in visible
    assert "未提供本病例PD-L1免疫组化图片" in visible
    assert "4、未提供可核验的病例专属PD-L1来源记录" in visible
    assert "4、定性结果判定标准：TPS<1%" not in visible
    assert "__PDL1_CASE_IMAGE__" not in visible
    assert len(rendered.inline_shapes) == len(Document(TEMPLATE).inline_shapes)
    pdl1_rows = [
        [cell.text.strip() for cell in row.cells]
        for table in rendered.tables
        for row in table.rows
        if row.cells and "PD-L1蛋白表达" in row.cells[0].text
    ]
    assert any(row[2:5] == ["未提供", "未提供", "未提供"] for row in pdl1_rows)
    assert "第三部分：基因变异及相应靶向/免疫药物解析" in visible
    assert "TP53基因" in visible


def test_lung588_generation_allows_source_record_only_pilot_profile(
    tmp_path,
):
    xlsx_path = tmp_path / "SYNTHETIC-LUNG588-PILOT.xlsx"
    xlsx_path.write_bytes(b"synthetic")
    image_path = _synthetic_pdl1_image(tmp_path)
    excel_data = ExcelDataSource(
        file_path=str(xlsx_path),
        single_values={
            "患者姓名": "SYNTHETIC_PATIENT",
            "样本编号": "SYNTHETIC_CASE_PILOT",
            "检测项目": "肺癌588基因+PD-L1",
            "报告日期": "2026-07-24",
            "TMB": 5,
            "MSI状态": "MSS",
            "PD-L1 TPS": 5,
            "PD-L1 CPS": 6,
            # Verbatim transcription is intentional: the unknown assay must
            # not silently inherit the 22C3/TPS display thresholds.
            "PD-L1结果": "阳性（高表达）",
            "PD-L1检测方案": ("legacy_unspecified_ihc_transcription_v1"),
            "PD-L1原始记录编号": "SYNTHETIC-IHC-PILOT-001",
            "PD-L1原始记录日期": "2026-07-24",
            "PD-L1检测标本标识": "SYNTHETIC-SPECIMEN-PILOT-001",
            "PD-L1图像处置": "病例专属图像（报告展示）",
            "PD-L1病例图片": str(image_path),
            "肺癌病理类型": "非小细胞肺癌",
            "疾病范围": "转移性",
            "既往系统治疗": "已接受",
            "伴随诊断适配状态": "已确认符合",
        },
        table_data={
            "Variations": [
                {
                    "ExistIn552": "Ⅱ类",
                    "Gene_Symbol": "TP53",
                    "Transcript": "NM_000546.6",
                    "cHGVS": "c.734G>A",
                    "pHGVS_S": "p.G245D",
                    "Freq(%)": 30,
                }
            ]
        },
        sheet_names=["Variations", "TMB", "Msisensor"],
    )
    output_dir = tmp_path / "out"

    result = ReportGenerator(
        config_dir=str(ROOT / "config"),
        log_level="ERROR",
    ).generate(
        excel_file=str(xlsx_path),
        excel_data=excel_data,
        template_file=str(TEMPLATE),
        output_dir=str(output_dir),
        project_type="lung_588_pdl1",
    )

    assert result["success"] is True, result["errors"]
    output = Path(result["output_file"])
    assert output.is_file()
    visible = "\n".join(
        [
            *(paragraph.text for paragraph in Document(output).paragraphs),
            *(
                cell.text
                for table in Document(output).tables
                for row in table.rows
                for cell in row.cells
            ),
        ]
    )
    assert "__PDL1_CASE_IMAGE__" not in visible
    with ZipFile(output) as archive:
        assert any(name.startswith("word/media/") for name in archive.namelist())
    assert "22C3" not in visible
    assert "结果判定沿用病例专属来源记录" in visible
    assert "定性结果判定标准：TPS<1%" not in visible
