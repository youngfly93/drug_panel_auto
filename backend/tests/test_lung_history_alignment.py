# ruff: noqa: E402
"""Regression coverage for the lung329/588 historical-final alignment."""

import sys
from pathlib import Path

import pandas as pd
import pytest
import yaml
from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
BACKEND = ROOT / "backend"
for import_path in (str(ROOT), str(BACKEND)):
    if import_path not in sys.path:
        sys.path.insert(0, import_path)

from reportgen.core.batch_runner import (
    BatchValidateOptions,
    iter_excel_inputs,
    run_batch_generate_validate,
)
from reportgen.core.excel_reader import ExcelReader
from reportgen.core.field_mapper import FieldMapper
from reportgen.core.report_generator import apply_pdl1_display_fields
from reportgen.core.template_bridge_358 import (
    _build_lung_chemotherapy_tables,
    _build_nccn_and_immune_fields,
    _build_targeted_drug_introductions,
    _compact_drug_display_tables,
    _compact_drug_display_value,
    build_immune_variants,
    load_panel_config,
)
from reportgen.models.excel_data import ExcelDataSource
from reportgen.models.report_data import ReportData
from reportgen.panels.loader import load_panel_package

GUIDELINE_KEYS = [
    "EGFR",
    "ALK",
    "ROS1",
    "BRAF",
    "KRAS",
    "NTRK123",
    "MET",
    "RET",
    "ERBB2",
    "MET_EGFR_JOINT",
]
IMMUNE_DISPLAYS = {
    "positive": [
        "MLH1",
        "MSH2",
        "MSH6",
        "PMS2",
        "POLE",
        "POLD1",
        "CD274（PD-L1）",
        "PDCD1LG2（PD-L2）",
        "PBRM1",
        "KRAS",
        "KRAS/TP53共突变",
        "TET1",
        "SERPINB3",
        "SERPINB4",
        "DDR基因（DNA损伤修复相关基因）",
    ],
    "negative": [
        "PTEN",
        "JAK1",
        "JAK2",
        "B2M",
        "CTNNB1",
        "EGFR（L858R/EX19del）",
        "ALK",
        "MET",
        "STK11",
        "KRAS/STK11共突变",
        "KEAP1",
        "IFNGR1/2",
    ],
    "hyperprogression": [
        "MDM2",
        "MDM4",
        "DNMT3A",
        "EGFR（扩增）",
        "CCND1",
        "FGF3",
        "FGF4",
        "FGF19",
    ],
}


def _panel_config(panel_id: str):
    package = load_panel_package(panel_id, project_root=ROOT)
    return load_panel_config(base_path=str(ROOT), panel_package=package)


@pytest.mark.parametrize("panel_id", ["lung_588_pdl1", "lung_329_pdl1"])
def test_lung_comprehensive_panels_share_historical_fixed_tables(panel_id):
    config = _panel_config(panel_id)

    assert [row["key"] for row in config.lung_guideline_drug_rows] == GUIDELINE_KEYS
    assert [row["display"] for row in config.immune_positive_rows] == (
        IMMUNE_DISPLAYS["positive"]
    )
    assert [row["display"] for row in config.immune_negative_rows] == (
        IMMUNE_DISPLAYS["negative"]
    )
    assert [row["display"] for row in config.immune_hyperprogression_rows] == (
        IMMUNE_DISPLAYS["hyperprogression"]
    )
    positive_by_key = {row["key"]: row for row in config.immune_positive_rows}
    negative_by_key = {row["key"]: row for row in config.immune_negative_rows}
    assert positive_by_key["MLH1"]["mode"] == "exact_variant"
    assert positive_by_key["PMS2"]["mode"] == "exact_variant"
    assert positive_by_key["DDR"]["mode"] == "exact_variant_group"
    assert {
        selector["genes"][0]
        for selector in positive_by_key["DDR"]["selectors"]
    } == {"ATM", "BRIP1", "MSH3", "BRCA2"}
    assert negative_by_key["PTEN"]["mode"] == "exact_variant"
    assert config.chemotherapy_module_enabled is True
    assert config.chemotherapy_rule["detail_allowed_levels"] == ["1B", "2A", "2B"]
    assert len(config.chemotherapy_rule["prediction_rows"]) == 27
    assert len(config.chemotherapy_rule["regimen_rows"]) == 22
    assert len(config.chemotherapy_rule["dosage_rows"]) == 11


@pytest.mark.parametrize("panel_id", ["lung_588_pdl1", "lung_329_pdl1"])
def test_lung_comprehensive_front_matter_uses_existing_section_break(panel_id):
    package = load_panel_package(panel_id, project_root=ROOT)
    payload = yaml.safe_load(package.resolve_rule_file("style").read_text())
    front_matter = payload["style"]["front_matter"]

    assert front_matter["guide_spacer_count"] == 30
    assert front_matter["insert_page_break"] is False


def test_lung_fixed_immune_tables_keep_historical_event_exactness(tmp_path):
    config = _panel_config("lung_588_pdl1")
    source_path = tmp_path / "SYNTHETIC-LUNG-IMMUNE-EXACT.xlsx"
    source_path.write_bytes(b"synthetic")
    raw_rows = [
        {
            "ExistIn552": "Ⅱ类",
            "Gene_Symbol": "ATM",
            "Transcript": "NM_000051.4",
            "cHGVS": "c.1236-2A>T",
            "pHGVS_S": "",
        },
        {
            "ExistIn552": "Ⅱ类",
            "Gene_Symbol": "ATM",
            "Transcript": "NM_000051.3",
            "cHGVS": "c.1236-2A>T",
            "pHGVS_S": "",
        },
        {
            "ExistIn552": "Ⅱ类",
            "Gene_Symbol": "BRIP1",
            "Transcript": "NM_032043.3",
            "cHGVS": "c.2142G>A",
            "pHGVS_S": "p.W714*",
        },
        {
            "ExistIn552": "Ⅱ类",
            "Gene_Symbol": "MSH3",
            "Transcript": "NM_002439.5",
            "cHGVS": "c.2905C>T",
            "pHGVS_S": "p.Q969*",
        },
        {
            "ExistIn552": "Ⅱ类",
            "Gene_Symbol": "BRCA2",
            "Transcript": "NM_000059.4",
            "cHGVS": "c.7007G>A",
            "pHGVS_S": "p.R2336H",
        },
        {
            "ExistIn552": "Ⅱ类",
            "Gene_Symbol": "PTEN",
            "Transcript": "NM_000314.8",
            "cHGVS": "c.802-2A>T",
            "pHGVS_S": "",
        },
        {
            "ExistIn552": "Ⅱ类",
            "Gene_Symbol": "PTEN",
            "Transcript": "NM_000314.8",
            "cHGVS": "c.801A>T",
            "pHGVS_S": "p.K267N",
        },
        {
            "ExistIn552": "Ⅱ类",
            "Gene_Symbol": "TP53",
            "Transcript": "NM_000546.6",
            "cHGVS": "c.734G>A",
            "pHGVS_S": "p.G245D",
        },
    ]
    excel = ExcelDataSource(
        file_path=str(source_path),
        table_data={"Variations": raw_rows},
        sheet_names=["Variations"],
    )

    immune = build_immune_variants(
        excel,
        filter_column="ExistIn552",
        panel_config=config,
    )
    assert [row["gene"] for row in immune["positive"]] == [
        "ATM",
        "BRIP1",
        "MSH3",
        "BRCA2",
    ]
    assert [row["gene"] for row in immune["negative"]] == ["PTEN"]
    assert all(row["gene"] != "TP53" for row in immune["positive"])

    report_data = ReportData()
    _build_nccn_and_immune_fields(
        report_data,
        [],
        excel,
        panel_config=config,
    )
    positive_rows = {
        row["key"]: row for row in report_data.get_table("immune_positive_results")
    }
    negative_rows = {
        row["key"]: row for row in report_data.get_table("immune_negative_results")
    }
    ddr_result = positive_rows["DDR"]["result"]
    assert all(gene in ddr_result for gene in ("ATM", "BRIP1", "MSH3", "BRCA2"))
    assert "TP53" not in ddr_result
    assert "c.802-2A>T" in negative_rows["PTEN"]["result"]
    assert "c.801A>T" not in negative_rows["PTEN"]["result"]


def test_lung_guideline_results_are_derived_while_fixed_copy_is_preserved(tmp_path):
    config = _panel_config("lung_588_pdl1")
    source_path = tmp_path / "SYNTHETIC-LUNG.xlsx"
    source_path.write_bytes(b"synthetic")
    excel = ExcelDataSource(
        file_path=str(source_path),
        table_data={
            "Variations": [
                {
                    "Gene_Symbol": "BRAF",
                    "ExistIn552": "Ⅱ类",
                    "cHGVS": "c.1799T>A",
                    "pHGVS_S": "p.V600E",
                }
            ],
            "Fusion": [{"Gene1": "EML4", "Gene2": "ALK"}],
        },
        sheet_names=["Variations", "Fusion"],
    )
    data = ReportData()

    _build_nccn_and_immune_fields(data, [], excel, panel_config=config)

    rows = data.get_table("lung_guideline_drug_results")
    assert len(rows) == 10
    by_key = {row["key"]: row for row in rows}
    assert "c.1799T>A，p.V600E" in by_key["BRAF"]["检测结果"]
    assert "融合:EML4-ALK" in by_key["ALK"]["检测结果"]
    assert by_key["EGFR"]["检测结果"] == "未见变异"
    assert "奥希替尼" in by_key["EGFR"]["本癌种相关治疗药物"]
    assert "NCCN" in by_key["EGFR"]["临床提示"]


def test_tmb_reference_follows_tcga_measurement_sample_not_overall_sample_type(
    tmp_path,
):
    frame = pd.DataFrame(
        [
            ["TCGA fit", None, None],
            ["sampletp", "TMB", "note"],
            ["tissue", 7.5, "selected"],
            ["blood", 15.0, "not selected"],
        ]
    )
    reader = ExcelReader(config_dir=str(ROOT / "config"), log_level="ERROR")
    value, sample_type = reader._extract_tmb_measurement(frame)
    assert (value, sample_type) == (7.5, "tissue")

    source_path = tmp_path / "SYNTHETIC-LUNG.xlsx"
    source_path.write_bytes(b"synthetic")
    source = ExcelDataSource(
        file_path=str(source_path),
        single_values={
            "样本编号": "SYNTHETIC-LUNG",
            "患者姓名": "SYNTHETIC",
            "样本类型": "组织+血液",
            "TMB": value,
            "TMB样本类型": sample_type,
        },
        table_data={},
        sheet_names=["TMB"],
    )
    mapped = FieldMapper(
        config_dir=str(ROOT / "config"), log_level="ERROR"
    ).map(
        source,
        panel_package=load_panel_package("lung_588_pdl1", project_root=ROOT),
    )

    assert mapped.get_field("sample_type") == "组织+血液"
    assert mapped.get_field("tmb_sample_type") == "tissue"
    assert mapped.get_field("tmb_reference") == 10
    assert mapped.get_field("tmb_summary").startswith("7.5 mutations/Mb")


@pytest.mark.parametrize("panel_id", ["lung_588_pdl1", "lung_329_pdl1"])
def test_lung_absent_gene_baseline_has_no_crc_rows(panel_id):
    genes = [
        row["gene"]
        for row in FieldMapper(
            config_dir=str(ROOT / "config"), log_level="ERROR"
        )._load_variants_2_1_baseline(panel_id)
    ]

    assert genes == [
        "ALK",
        "EGFR",
        "KRAS",
        "MET",
        "RET",
        "ROS1",
        "NTRK1",
        "NTRK2",
        "NTRK3",
        "ERBB2",
        "BRAF",
        "NOTCH1",
        "MSH2",
        "MSH6",
        "PIK3CA",
    ]
    assert not {"FBXW7", "NF1", "NRAS", "SMAD4", "SMARCA4", "TCF7L2"} & set(
        genes
    )


def test_lung_chemotherapy_filters_appendix_and_normalizes_historical_display(
    tmp_path,
):
    config = _panel_config("lung_588_pdl1")
    data = ReportData(
        context={
            "drug_shunbo": [
                {
                    "Drug": "platinum compounds（顺铂）",
                    "Level": "1B",
                    "Result": "Uncovered",
                },
                {
                    "Drug": "platinum compounds（顺铂）",
                    "Level": "3",
                    "Result": "疗效较好",
                },
            ]
        }
    )
    source_path = tmp_path / "SYNTHETIC-LUNG.xlsx"
    source_path.write_bytes(b"synthetic")
    excel = ExcelDataSource(
        file_path=str(source_path),
        table_data={
            "CtDrug": [
                {
                    "药物": "platinum compounds（顺铂）",
                    "检测基因": "ERCC1",
                    "用药提示（仅供参考）": "疗效较好",
                }
            ]
        },
        sheet_names=["CtDrug"],
    )

    _build_lung_chemotherapy_tables(data, excel, config)

    assert data.get_table("drug_shunbo") == [
        {
            "Drug": "platinum compounds（顺铂）",
            "Level": "1B",
            "Result": "未覆盖",
            "DrugDisplay": "顺铂",
            "药物名称": "顺铂",
        }
    ]
    assert len(data.get_table("chemotherapy_predictions")) == 27
    assert len(data.get_table("chemotherapy_regimen_predictions")) == 22
    assert len(data.get_table("chemotherapy_dosage_rows")) == 11
    cisplatin = data.get_table("chemotherapy_predictions")[0]
    assert cisplatin["药物"] == "顺铂"
    assert cisplatin["相关基因"] == "ERCC1"
    assert cisplatin["有效性"] == "可能较高"


def test_ct1000_explicit_summary_is_preferred_and_reproduces_history(tmp_path):
    config = _panel_config("lung_588_pdl1")
    source_path = tmp_path / "SYNTHETIC-LUNG.xlsx"
    source_path.write_bytes(b"synthetic")
    source = ExcelDataSource(
        file_path=str(source_path),
        table_data={
            "Ct1000_summary": [
                {
                    "source_drug": "顺铂（cisplatin）",
                    "drug": "顺铂",
                    "genes": ["XPC", "MTHFR", "GSTP1", "ERCC1", "XRCC1", "GSTM1"],
                    "efficacy": "可能居中",
                    "toxicity": "可能较高",
                },
                {
                    "source_drug": "吉西他滨（gemcitabine）",
                    "drug": "吉西他滨",
                    "genes": ["CDA", "RRM1"],
                    "efficacy": "可能较高",
                    "toxicity": "可能较低",
                },
            ],
            # This conflicting free-text value must not override Ct1000.
            "CtDrug": [
                {
                    "药物": "顺铂（cisplatin）",
                    "检测基因": "ERCC1",
                    "用药提示（仅供参考）": "药物敏感性可能较高",
                }
            ],
        },
        sheet_names=["Ct1000", "CtDrug"],
    )
    data = ReportData()

    _build_lung_chemotherapy_tables(data, source, config)

    cisplatin = data.get_table("chemotherapy_predictions")[0]
    assert cisplatin == {
        "drug": "顺铂",
        "regimen": "顺铂",
        "genes": "XPC、MTHFR、GSTP1、ERCC1、XRCC1、GSTM1",
        "efficacy": "可能居中",
        "toxicity": "可能较高",
        "药物": "顺铂",
        "化疗方案": "顺铂",
        "相关基因": "XPC、MTHFR、GSTP1、ERCC1、XRCC1、GSTM1",
        "有效性": "可能居中",
        "毒副作用": "可能较高",
    }
    assert data.get_field("chemotherapy_summary_text") == (
        "经分析，可考虑优先选择的化疗方案有吉西他滨单药方案。"
    )


def test_ct1000_parser_keeps_source_gene_order_and_explicit_ratings():
    frame = pd.DataFrame(
        [
            [1, "顺铂（cisplatin）", "", "rs1", "XPC"],
            [2, "顺铂（cisplatin）", "", "rs2", "ERCC1"],
            ["sensitive M! Toxicity H!", None, None, None, None],
            [3, "三胺硫磷（thiotepa）", "", "rs3", "ALDH1A1"],
            ["Toxicity L!", None, None, None, None],
        ]
    )

    rows = ExcelReader._extract_ct1000_summaries(frame)

    assert rows == [
        {
            "source_drug": "顺铂（cisplatin）",
            "drug": "顺铂",
            "genes": ["XPC", "ERCC1"],
            "gene_display": "XPC、ERCC1",
            "efficacy_code": "M",
            "toxicity_code": "H",
            "efficacy": "可能居中",
            "toxicity": "可能较高",
        },
        {
            "source_drug": "三胺硫磷（thiotepa）",
            "drug": "三胺硫磷",
            "genes": ["ALDH1A1"],
            "gene_display": "ALDH1A1",
            "efficacy_code": "",
            "toxicity_code": "L",
            "efficacy": "/",
            "toxicity": "可能较低",
        },
    ]


def test_drug_compaction_preserves_review_notice_and_counts_only_drugs():
    value = "\n".join(
        [
            "药物1（C）",
            "药物2（C）",
            "药物3（C）",
            "药物4（C）",
            "药物5（C）",
            "药物6（C）",
            "【待报告组审】",
        ]
    )

    assert _compact_drug_display_value(value, max_items=5).splitlines() == [
        "药物1（C）",
        "药物2（C）",
        "药物3（C）",
        "药物4（C）",
        "药物5（C）",
        "另1项详见第三部分",
        "【待报告组审】",
    ]


def test_targeted_summary_blanks_only_consecutive_repeated_gene_displays():
    data = ReportData(
        context={
            "targeted_drug_tips": [
                {"gene": "TP53", "benefit_drugs": "药物1（C）"},
                {"gene": "TP53", "benefit_drugs": "药物2（C）"},
                {"gene": "ATM", "benefit_drugs": "药物3（C）"},
                {"gene": "TP53", "benefit_drugs": "药物4（C）"},
            ]
        }
    )

    _compact_drug_display_tables(
        data,
        max_items=5,
        blank_repeated_gene=True,
    )

    rows = data.get_table("targeted_drug_tips")
    assert [row["gene"] for row in rows] == ["TP53", "TP53", "ATM", "TP53"]
    assert [row["gene_display"] for row in rows] == ["TP53", "", "ATM", "TP53"]


@pytest.mark.parametrize(
    "template_path",
    [
        ROOT
        / "panels/lung_588_pdl1/templates/lung_588_pdl1_historical_golden_v1.docx",
        ROOT / "panels/lung_329_pdl1/templates/lung_329_pdl1_golden_template_v2.docx",
    ],
)
def test_lung_templates_bind_targeted_summary_to_gene_display(template_path):
    document = Document(template_path)
    marker = "{%tr for row in targeted_drug_tips %}"
    loop_table = next(
        table
        for table in document.tables
        if marker in "\n".join(cell.text for row in table.rows for cell in row.cells)
    )
    marker_index = next(
        index
        for index, row in enumerate(loop_table.rows)
        if marker in " ".join(cell.text for cell in row.cells)
    )

    assert loop_table.rows[marker_index + 1].cells[0].text.strip() == (
        "{{ row.gene_display }}"
    )


def test_targeted_drug_introduction_compacts_long_lists_without_losing_source():
    full_value = "\n".join(
        [*(f"药物{index}（C）" for index in range(1, 8)), "【待报告组审】"]
    )
    data = ReportData(
        context={
            "targeted_drug_tips": [
                {
                    "gene": "BRAF",
                    "variant_site": "c.1A>T",
                    "benefit_drugs": full_value,
                    "caution_drugs": "--",
                }
            ]
        }
    )

    _build_targeted_drug_introductions(data, max_items=5)

    assert data.get_table("targeted_drug_tips")[0]["benefit_drugs"] == full_value
    assert data.get_table("targeted_drug_introductions")[0]["drug_name"] == (
        "药物1（C）、药物2（C）、药物3（C）、药物4（C）、药物5（C）、"
        "另2项详见第三部分"
    )


def test_ctdrug_short_tip_populates_the_historical_result_alias():
    row = {"用药提示（仅供参考）": "Uncovered"}

    FieldMapper(
        config_dir=str(ROOT / "config"), log_level="ERROR"
    )._apply_ctdrug_template_aliases(row)

    assert row["Result"] == "Uncovered"
    assert row["result"] == "Uncovered"
    assert row["检测结果"] == "Uncovered"


def test_batch_missing_pdl1_displays_explicit_not_provided():
    data = ReportData()

    apply_pdl1_display_fields(data)

    assert data.get_field("pdl1_tps_display") == "未提供"
    assert data.get_field("pdl1_cps_display") == "未提供"
    assert data.get_field("pdl1_result_display") == "未提供"


def test_lung329_template_has_panel_specific_gene_count_and_cjk_notice_font():
    document = Document(
        ROOT / "panels/lung_329_pdl1/templates/lung_329_pdl1_golden_template_v2.docx"
    )
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "与肿瘤密切相关的329个基因进行检测" in full_text
    assert "与肿瘤密切相关的588个基因进行检测" not in full_text
    notice = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("依据病例Excel的CtDrug表生成化疗药物小结")
    )
    assert notice.runs
    assert notice.runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")) == "微软雅黑"


def test_batch_failure_before_context_returns_a_serializable_failure(tmp_path):
    invalid_excel = tmp_path / "invalid.xlsx"
    invalid_excel.write_bytes(b"not-an-excel-workbook")

    run = run_batch_generate_validate(
        BatchValidateOptions(
            inputs=[str(invalid_excel)],
            config_dir=str(ROOT / "config"),
            output_root=str(tmp_path / "batch"),
            emit_context=False,
            emit_meta=False,
        )
    )

    assert run.report["failures"] == 1
    assert run.report["results"][0]["patient_snapshot"] == {}
    assert run.report["results"][0]["errors"]


def test_batch_input_discovery_ignores_macos_resource_forks(tmp_path):
    (tmp_path / "._case.xlsx").write_bytes(b"apple-double")
    expected = tmp_path / "case.xlsx"
    expected.write_bytes(b"excel-placeholder")

    assert iter_excel_inputs([str(tmp_path)]) == [expected.resolve()]
