# ruff: noqa: E402,I001
"""Lung588 Phase-C evidence boundary and historical-audit regression tests."""

from __future__ import annotations

import csv
import hashlib
import json
import subprocess
import sys
import zipfile
from pathlib import Path
from types import SimpleNamespace

import pytest
import yaml
from docx import Document
from lxml import etree

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from reportgen.core.field_mapper import FieldMapper
from reportgen.core.template_bridge_358 import _variant_override_matches
from reportgen.knowledge.quality import (
    build_panel_gene_provider,
    is_generic_mutation_analysis,
    profile_panel_runtime_content,
)
from reportgen.knowledge.release_gate import run_knowledge_release_gate
from reportgen.knowledge.redactions import (
    load_panel_knowledge_redactions,
)
from reportgen.panels.loader import load_panel_package
from reportgen.rules.targeted_drugs import (
    evaluate_required_clinical_context,
    load_targeted_drug_rule_context,
    reviewed_variant_selector_specificity,
)
from scripts.repair_docx_relationships import repair_docx
from scripts import (
    validate_lung588_pilot_boundary_suite,
    validate_lung588_real_inputs,
)


PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
OFFICE_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
DRAWING_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
PANEL_DIR = ROOT / "panels" / "lung_588_pdl1"
PRE_UAT_RECORD = PANEL_DIR / "uat" / "lung588_machine_pre_uat_20260723.yaml"
DOMAIN_CANDIDATES = PANEL_DIR / "rules" / "reviewed_part3_domain_candidates.yaml"
EVENT_NARRATIVE_CANDIDATES = (
    PANEL_DIR / "rules" / "reviewed_part3_p0_event_narrative_candidates.yaml"
)
CROSS_CANCER_NARRATIVE_CANDIDATES = (
    PANEL_DIR / "rules" / "reviewed_part3_p0_cross_cancer_narrative_candidates.yaml"
)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _inject_broken_image_relationship(path: Path) -> None:
    with zipfile.ZipFile(path) as archive:
        infos = archive.infolist()
        content = {item.filename: archive.read(item.filename) for item in infos}

    rels_name = "word/_rels/document.xml.rels"
    rels = etree.fromstring(content[rels_name])
    etree.SubElement(
        rels,
        f"{{{PACKAGE_REL_NS}}}Relationship",
        Id="rIdBroken",
        Type=f"{OFFICE_REL_NS}/image",
        Target="../NULL",
    )
    content[rels_name] = etree.tostring(
        rels,
        encoding="UTF-8",
        xml_declaration=True,
    )

    document = etree.fromstring(content["word/document.xml"])
    body = document.find(f"{{{WORD_NS}}}body")
    assert body is not None
    paragraph = etree.SubElement(body, f"{{{WORD_NS}}}p")
    run = etree.SubElement(paragraph, f"{{{WORD_NS}}}r")
    drawing = etree.SubElement(run, f"{{{WORD_NS}}}drawing")
    etree.SubElement(
        drawing,
        f"{{{DRAWING_NS}}}blip",
        {f"{{{OFFICE_REL_NS}}}embed": "rIdBroken"},
    )
    content["word/document.xml"] = etree.tostring(
        document,
        encoding="UTF-8",
        xml_declaration=True,
    )

    replacement = path.with_suffix(".injected.docx")
    with zipfile.ZipFile(
        replacement,
        mode="w",
        compression=zipfile.ZIP_DEFLATED,
    ) as archive:
        for item in infos:
            archive.writestr(item, content[item.filename])
    replacement.replace(path)


def _add_row(table, values: list[str]) -> None:
    cells = table.add_row().cells
    for cell, value in zip(cells, values):
        cell.text = value


def _synthetic_historical_report(path: Path) -> None:
    document = Document()

    cover = document.add_table(rows=1, cols=2)
    cover.cell(0, 0).text = "受控合成封面"
    cover.cell(0, 1).text = "CASE-LUNG-SYNTHETIC"

    targeted = document.add_table(rows=1, cols=4)
    for cell, value in zip(
        targeted.rows[0].cells,
        ["基因", "突变位点", "潜在获益靶向药物（证据等级）", "慎用药物"],
    ):
        cell.text = value
    _add_row(
        targeted,
        ["BRAF", "c.1799T>A，p.V600E", "达拉非尼+曲美替尼（A）", "--"],
    )

    biomarkers = document.add_table(rows=1, cols=2)
    biomarkers.cell(0, 0).text = "项目"
    biomarkers.cell(0, 1).text = "结果"
    _add_row(biomarkers, ["肿瘤突变负荷（TMB）", "6.3 mutations/Mb"])
    _add_row(biomarkers, ["微卫星不稳定性（MSI）", "MSS"])
    _add_row(biomarkers, ["PD-L1表达", "TPS 5%；CPS 6"])
    _add_row(
        biomarkers,
        ["免疫正相关基因", "检出（1个） BRAF：c.1799T>A，p.V600E"],
    )

    pgx = document.add_table(rows=1, cols=6)
    for cell, value in zip(
        pgx.rows[0].cells,
        ["药物", "基因", "检测位点", "等级", "检测结果", "用药提示"],
    ):
        cell.text = value
    _add_row(pgx, ["铂类", "ERCC1", "rs11615", "C", "CT", "历史候选说明"])

    document.add_paragraph("靶向/免疫药物用药提示解析")
    document.add_paragraph("潜在获益靶向药物解析")
    document.add_paragraph("BRAF：c.1799T>A，p.V600E突变")
    document.add_paragraph("达拉非尼+曲美替尼")
    document.add_paragraph("基因变异与药物关联分析：")
    document.add_paragraph("非小细胞肺癌研究见 PMID 28919011。")
    document.add_paragraph("3. 阅读说明")
    document.save(path)


def test_lung588_medical_candidates_are_registered_but_fail_closed():
    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    candidate_path = package.resolve_rule_file("medical_candidates")
    candidates = yaml.safe_load(candidate_path.read_text(encoding="utf-8"))
    runtime = yaml.safe_load(package.resolve_rule_file("drugs").read_text(encoding="utf-8"))

    assert candidates["status"] == "draft"
    governance = candidates["governance"]
    assert governance["runtime_policy"]["enabled"] is False
    assert governance["runtime_policy"]["runtime_rule_source"] is False
    assert governance["runtime_policy"]["report_text_allowed"] is False
    context_contract = governance["promotion_context_contract"]
    assert context_contract["status"] == "exposed_optional_in_engineering_draft"
    assert context_contract["contract_rule"] == "clinical_context"
    assert context_contract["runtime_enforcement"] == "implemented_fail_closed"
    assert context_contract["promotion_blocked"] is True
    assert context_contract["missing_or_uncertain_policy"] == "keep_candidate_hidden"
    runtime_context = yaml.safe_load(
        package.resolve_rule_file("clinical_context").read_text(encoding="utf-8")
    )["clinical_context_contract"]
    assert set(runtime_context["fields"]) == {
        "lung_histology",
        "disease_extent",
        "prior_systemic_therapy",
        "companion_diagnostic_status",
    }
    assert governance["historical_inventory"] == {
        "source": "two de-identified repaired historical lung588 final reports",
        "exact_targeted_event_count": 15,
        "targeted_drug_claim_count": 169,
        "historical_level_counts": {"A": 9, "B": 0, "C": 146, "D": 14},
        "selected_candidate_claim_count": 4,
        "not_migrated_claim_count": 165,
        "disposition": (
            "Historical report text is evidence of an old display contract, "
            "not current medical truth. Only the four exact-event candidates "
            "below were retained for secondary review; all other claims remain "
            "outside runtime.\n"
        ),
    }

    rules = candidates["candidate_rules"]
    assert len(rules) == 4
    assert {rule["gene"] for rule in rules} == {"BRAF", "ERBB2"}
    assert all(rule["selector"]["c_hgvs"].startswith("c.") for rule in rules)
    assert all(rule["selector"]["p_hgvs"].startswith("p.") for rule in rules)
    assert all(rule["selector"]["transcript"].startswith("NM_") for rule in rules)
    assert all(rule["runtime_eligible"] is False for rule in rules)
    assert all(rule["report_text_allowed"] is False for rule in rules)
    assert all(rule["review_status"] == "needs_review" for rule in rules)
    assert all(rule["secondary_review_status"] == "pending_report_group_review" for rule in rules)
    assert all(rule["source_refs"] for rule in rules)
    for rule in rules:
        required_context = set(rule["required_context_fields"])
        assert set(rule["context_requirements"]) == required_context
        assert {
            "lung_histology",
            "disease_extent",
            "companion_diagnostic_status",
        } <= required_context
        if rule["gene"] == "ERBB2":
            assert "prior_systemic_therapy" in required_context

    candidate_events = {
        (
            rule["gene"],
            rule["selector"]["c_hgvs"],
            rule["selector"]["p_hgvs"],
        )
        for rule in rules
    }
    assert ("BRAF", "c.1781A>G", "p.D594G") not in candidate_events
    candidate_drugs = {rule["therapy"]["generic_name_zh"] for rule in rules}
    assert not {"宗艾替尼", "恩美曲妥珠单抗", "吡咯替尼"} & candidate_drugs

    assert runtime["targeted_drug_rules"]["enabled"] is False
    assert runtime["targeted_drug_rules"]["base_db_enabled"] is False
    assert runtime["targeted_drug_rules"]["allowed_source_dbs"] == []
    assert runtime["targeted_drug_rules"]["clinical_context_rule"] == "clinical_context"
    assert runtime["approved_drug_rows"] == []
    assert runtime["governance"]["schema_version"] == "1.0"
    assert runtime["governance"]["defaults"]["targeted_drug"]["runtime_eligible"] is False
    assert candidates["non_target_domains"]["immune_gene_associations"]["enabled"] is False
    assert candidates["non_target_domains"]["chemotherapy_pharmacogenomics"]["enabled"] is False

    if runtime["targeted_drug_rules"]["enabled"]:
        assert context_contract["runtime_enforcement"] == "implemented_fail_closed"
        assert context_contract["promotion_blocked"] is False


def test_lung588_candidate_evidence_review_matches_exact_nonruntime_queue():
    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    candidates = yaml.safe_load(
        package.resolve_rule_file("medical_candidates").read_text(encoding="utf-8")
    )
    contract = yaml.safe_load(
        package.resolve_rule_file("candidate_evidence_review").read_text(encoding="utf-8")
    )

    governance = contract["governance"]
    assert contract["status"] == "draft"
    assert governance["primary_review_status"] == ("completed_non_authoritative_scope_review")
    assert governance["secondary_review_status"] == ("pending_report_group_review")
    assert governance["runtime_rule_source"] is False
    assert governance["runtime_eligible"] is False
    assert governance["report_text_allowed"] is False
    assert governance["promotion_blocked"] is True

    candidate_by_id = {row["candidate_id"]: row for row in candidates["candidate_rules"]}
    review_by_id = {row["candidate_id"]: row for row in contract["reviews"]}
    assert len(review_by_id) == 4
    assert set(review_by_id) == set(candidate_by_id)
    for candidate_id, review in review_by_id.items():
        candidate = candidate_by_id[candidate_id]
        assert review["gene"] == candidate["gene"]
        assert review["selector"] == candidate["selector"]
        assert review["therapy"] == {
            key: candidate["therapy"][key] for key in ("generic_name_zh", "generic_name_en")
        }
        assert review["source_claim_reviews"]
        assert all(
            source["supports"] and source["does_not_support"]
            for source in review["source_claim_reviews"]
        )
        assert review["primary_review"]["status"] == ("completed_non_authoritative_scope_review")
        assert review["secondary_review"]["status"] == ("pending_report_group_review")
        assert review["runtime_eligible"] is False
        assert review["report_text_allowed"] is False
        assert review["scope_assessment"]["patient_report_display"].startswith("prohibited_")

    for candidate_id in (
        "lung588_braf_v600e_dabrafenib_trametinib",
        "lung588_braf_v600e_encorafenib_binimetinib",
    ):
        review = review_by_id[candidate_id]
        assert (
            review["scope_assessment"]["direct_exact_drug_event_clinical_outcome"] == "identified"
        )
        assert review["primary_review"]["decision"] == (
            "eligible_to_remain_candidate_pending_secondary_review"
        )

    deruxtecan = review_by_id["lung588_erbb2_g660d_trastuzumab_deruxtecan"]
    assert deruxtecan["scope_assessment"]["event_scope"] == (
        "exact_g660d_functional_and_assay_support"
    )
    assert deruxtecan["scope_assessment"]["assay_scope"] == (
        "exact_g660d_reportable_in_companion_diagnostic_documentation"
    )
    assert (
        deruxtecan["scope_assessment"]["direct_exact_drug_event_clinical_outcome"]
        == "not_identified"
    )
    assert deruxtecan["primary_review"]["decision"] == (
        "eligible_to_remain_candidate_with_indirect_exact_event_chain_pending_secondary_review"
    )

    rezetecan = review_by_id["lung588_erbb2_g660d_trastuzumab_rezetecan"]
    assert (
        rezetecan["scope_assessment"]["direct_exact_drug_event_clinical_outcome"]
        == "not_identified"
    )
    assert rezetecan["scope_assessment"]["assay_scope"] == (
        "official_full_label_and_exact_variant_eligibility_not_captured"
    )
    assert rezetecan["primary_review"]["decision"] == (
        "hold_pending_official_china_label_and_secondary_review"
    )


def test_lung588_candidate_context_evaluator_rejects_missing_uncertain_and_out_of_scope():
    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    candidates = yaml.safe_load(
        package.resolve_rule_file("medical_candidates").read_text(encoding="utf-8")
    )
    contract = yaml.safe_load(
        package.resolve_rule_file("clinical_context").read_text(encoding="utf-8")
    )["clinical_context_contract"]
    rules = {rule["candidate_id"]: rule for rule in candidates["candidate_rules"]}
    braf = rules["lung588_braf_v600e_dabrafenib_trametinib"]
    erbb2 = rules["lung588_erbb2_g660d_trastuzumab_deruxtecan"]

    braf_context = {
        "lung_histology": "非小细胞肺癌",
        "disease_extent": "转移性",
        "companion_diagnostic_status": "已确认符合",
    }
    assert evaluate_required_clinical_context(
        braf,
        clinical_context=braf_context,
        contract=contract,
    ).eligible

    missing = dict(braf_context)
    missing.pop("companion_diagnostic_status")
    assert evaluate_required_clinical_context(
        braf,
        clinical_context=missing,
        contract=contract,
    ).reasons == ("CONTEXT_VALUE_MISSING:companion_diagnostic_status",)

    uncertain = {**braf_context, "lung_histology": "未明确"}
    assert evaluate_required_clinical_context(
        braf,
        clinical_context=uncertain,
        contract=contract,
    ).reasons == ("CONTEXT_VALUE_UNCERTAIN:lung_histology",)

    wrong_stage = {**braf_context, "disease_extent": "可切除早期"}
    assert evaluate_required_clinical_context(
        braf,
        clinical_context=wrong_stage,
        contract=contract,
    ).reasons == ("CONTEXT_OUT_OF_SCOPE:disease_extent",)

    erbb2_context = {
        **braf_context,
        "disease_extent": "不可切除局部晚期",
        "prior_systemic_therapy": "已接受",
    }
    assert evaluate_required_clinical_context(
        erbb2,
        clinical_context=erbb2_context,
        contract=contract,
    ).eligible
    not_previously_treated = {
        **erbb2_context,
        "prior_systemic_therapy": "未接受",
    }
    assert evaluate_required_clinical_context(
        erbb2,
        clinical_context=not_previously_treated,
        contract=contract,
    ).reasons == ("CONTEXT_OUT_OF_SCOPE:prior_systemic_therapy",)


def test_runtime_loader_moves_context_ineligible_exact_rule_to_blocked_set(tmp_path):
    panel_root = tmp_path / "panels" / "synthetic_lung"
    rules_root = panel_root / "rules"
    rules_root.mkdir(parents=True)
    (rules_root / "clinical_context.yaml").write_text(
        yaml.safe_dump(
            {
                "schema_version": "1.0",
                "panel_id": "synthetic_lung",
                "rule_id": "clinical_context",
                "clinical_context_contract": {
                    "fields": {
                        "lung_histology": {
                            "allowed_values": ["非小细胞肺癌", "未明确"],
                            "uncertain_values": ["未明确"],
                        }
                    }
                },
            },
            allow_unicode=True,
        ),
        encoding="utf-8",
    )
    (rules_root / "drugs.yaml").write_text(
        yaml.safe_dump(
            {
                "schema_version": "1.0",
                "panel_id": "synthetic_lung",
                "rule_id": "drugs",
                "targeted_drug_rules": {
                    "enabled": True,
                    "clinical_context_rule": "clinical_context",
                    "reviewed_variant_overrides": [
                        {
                            "gene": "BRAF",
                            "c_hgvs": "c.1799T>A",
                            "p_hgvs": "p.V600E",
                            "required_context_fields": ["lung_histology"],
                            "context_requirements": {"lung_histology": ["非小细胞肺癌"]},
                            "benefit_drugs": ["受控合成药物（A）"],
                        }
                    ],
                },
            },
            allow_unicode=True,
        ),
        encoding="utf-8",
    )

    class SyntheticPackage:
        panel_id = "synthetic_lung"
        root_dir = panel_root

        @staticmethod
        def resolve_rule_file(name: str) -> Path:
            return rules_root / f"{name}.yaml"

    active = load_targeted_drug_rule_context(
        SyntheticPackage(),
        clinical_context={"lung_histology": "非小细胞肺癌"},
    )
    assert active is not None
    assert active["clinical_context_enforced"] is True
    assert len(active["reviewed_variant_overrides"]) == 1
    assert active["blocked_reviewed_variant_overrides"] == []
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")
    assert mapper._lookup_reviewed_variant_override_drugs(
        "BRAF",
        "c.1799T>A",
        "p.V600E",
        targeted_drug_rules=active,
    ) == ("受控合成药物（A）", "--")

    blocked = load_targeted_drug_rule_context(
        SyntheticPackage(),
        clinical_context={"lung_histology": "未明确"},
    )
    assert blocked is not None
    assert blocked["reviewed_variant_overrides"] == []
    assert len(blocked["blocked_reviewed_variant_overrides"]) == 1
    assert blocked["blocked_reviewed_variant_overrides"][0]["_clinical_context_block_reasons"] == [
        "CONTEXT_VALUE_UNCERTAIN:lung_histology"
    ]
    assert mapper._lookup_reviewed_variant_override_drugs(
        "BRAF",
        "c.1799T>A",
        "p.V600E",
        targeted_drug_rules=blocked,
    ) == ("--", "--")


def test_reviewed_variant_transcript_selector_is_exact_and_fail_closed():
    mapper = FieldMapper(config_dir=str(ROOT / "config"), log_level="ERROR")
    generic = {
        "gene": "BRAF",
        "c_hgvs": "c.1799T>A",
        "p_hgvs": "p.V600E",
        "benefit_drugs": ["通用合成药物（A）"],
    }
    transcript_bound = {
        **generic,
        "transcript": "NM_004333.6",
        "benefit_drugs": ["转录本限定合成药物（A）"],
    }
    rules = {
        "enabled": True,
        "reviewed_variant_overrides": [generic],
        "blocked_reviewed_variant_overrides": [transcript_bound],
    }

    # A transcript-bound pending selector outranks the otherwise identical
    # generic selector and keeps the reviewed event closed.
    assert mapper._lookup_reviewed_variant_override_drugs(
        "BRAF",
        "c.1799T>A",
        "p.V600E",
        transcript="NM_004333.6",
        targeted_drug_rules=rules,
    ) == ("--", "--")
    # Missing or version-mismatched transcript must not match that selector;
    # the older rule without a transcript declaration remains compatible.
    assert mapper._lookup_reviewed_variant_override_drugs(
        "BRAF",
        "c.1799T>A",
        "p.V600E",
        targeted_drug_rules=rules,
    ) == ("通用合成药物（A）", "--")
    assert mapper._lookup_reviewed_variant_override_drugs(
        "BRAF",
        "c.1799T>A",
        "p.V600E",
        transcript="NM_004333.5",
        targeted_drug_rules=rules,
    ) == ("通用合成药物（A）", "--")

    assert _variant_override_matches(
        transcript_bound,
        "BRAF",
        "c.1799T>A",
        "p.V600E",
        transcript="NM_004333.6",
    )
    assert not _variant_override_matches(
        transcript_bound,
        "BRAF",
        "c.1799T>A",
        "p.V600E",
    )
    assert not _variant_override_matches(
        transcript_bound,
        "BRAF",
        "c.1799T>A",
        "p.V600E",
        transcript="NM_004333.5",
    )
    assert reviewed_variant_selector_specificity(
        transcript_bound
    ) > reviewed_variant_selector_specificity(generic)
    assert reviewed_variant_selector_specificity(
        {"gene": "BRAF", "transcript": "NM_004333.6"}
    ) < reviewed_variant_selector_specificity(generic)


def test_docx_relationship_repair_preserves_source_and_removes_orphan(tmp_path):
    source = tmp_path / "CASE-LUNG-SYNTHETIC.reference.docx"
    output = tmp_path / "CASE-LUNG-SYNTHETIC.repaired.docx"
    document = Document()
    document.add_paragraph("synthetic content")
    document.save(source)
    _inject_broken_image_relationship(source)
    before = _sha256(source)

    result = repair_docx(source, output)

    assert _sha256(source) == before
    assert result["source_sha256"] == before
    assert result["removed_reference_nodes"] == 1
    assert result["removed_relationships"] == [
        {
            "relation_id": "rIdBroken",
            "relationship_type": "image",
            "reason": "null_target",
        }
    ]
    assert Document(output).paragraphs[0].text == "synthetic content"
    with zipfile.ZipFile(output) as archive:
        assert b"rIdBroken" not in archive.read("word/document.xml")
        assert b"rIdBroken" not in archive.read("word/_rels/document.xml.rels")


def test_historical_semantic_audit_emits_deidentified_nonruntime_inventory(tmp_path):
    reference = tmp_path / "CASE-LUNG-SYNTHETIC.reference.docx"
    contract = tmp_path / "case_lung_synthetic.yaml"
    output_dir = tmp_path / "audit"
    _synthetic_historical_report(reference)
    contract.write_text(
        yaml.safe_dump(
            {
                "tables": {
                    "all_variants": {
                        "rows": [
                            {
                                "match": {
                                    "gene": {"equals": "BRAF"},
                                    "cHGVS": {"equals": "c.1799T>A"},
                                },
                                "expect": {"pHGVS": {"equals": "p.V600E"}},
                            }
                        ]
                    }
                }
            },
            allow_unicode=True,
            sort_keys=False,
        ),
        encoding="utf-8",
    )

    completed = subprocess.run(
        [
            sys.executable,
            str(ROOT / "scripts" / "audit_lung588_historical_semantics.py"),
            "--reference",
            f"CASE-LUNG-SYNTHETIC={reference}",
            "--contract",
            f"CASE-LUNG-SYNTHETIC={contract}",
            "--output-dir",
            str(output_dir),
        ],
        check=False,
        capture_output=True,
        text=True,
        cwd=ROOT,
    )

    assert completed.returncode == 0, completed.stderr
    summary = json.loads(completed.stdout)
    assert summary == {
        "status": "PASS",
        "cases": 1,
        "targeted_candidates": 1,
        "immune_observations": 1,
        "pgx_observations": 1,
        "output": "semantic_inventory.json",
    }
    inventory = json.loads((output_dir / "semantic_inventory.json").read_text(encoding="utf-8"))
    case = inventory["cases"]["CASE-LUNG-SYNTHETIC"]
    assert case["contract_check"]["status"] == "PASS"
    assert case["pgx_observation_count"] == 1
    assert inventory["cross_case"]["targeted_candidate_count"] == 1

    emitted = "\n".join(
        path.read_text(encoding="utf-8")
        for path in output_dir.iterdir()
        if path.suffix in {".json", ".tsv"}
    )
    assert "runtime_eligible\tFalse" not in emitted
    assert "\tFalse\tpending_report_group_review" in emitted
    assert "姓名" not in emitted
    assert "报告编号" not in emitted
    assert "LZ258" not in emitted


def test_real_input_validator_reads_quiet_release_revision(tmp_path, monkeypatch):
    revision = "1" * 40
    (tmp_path / "REVISION").write_text(f"{revision}\n", encoding="utf-8")
    monkeypatch.setattr(validate_lung588_real_inputs, "ROOT", tmp_path)

    assert validate_lung588_real_inputs._source_revision() == revision


def test_synthetic_pilot_validator_reads_explicit_release_revision(
    monkeypatch,
):
    revision = "2" * 40
    monkeypatch.setenv("REPORTGEN_SOURCE_REVISION", revision)

    assert validate_lung588_pilot_boundary_suite._source_revision() == revision


def test_real_input_validator_handles_empty_release_revision(tmp_path, monkeypatch):
    (tmp_path / "REVISION").write_text("", encoding="utf-8")
    monkeypatch.setattr(validate_lung588_real_inputs, "ROOT", tmp_path)

    assert validate_lung588_real_inputs._source_revision() == ""


def test_real_input_validator_separates_ngs_pre_uat_from_formal_uat():
    rows = [
        {
            "alias": f"CASE-LUNG-{alias}",
            "auto_detection": {
                "detected": False,
                "project_type": None,
            },
            "targeted_drug_count": 0,
            "biomarker_contract_status": "PASS",
            "pdl1_product_contract_status": "FAIL",
            "pdl1_input_provenance": "synthetic_visual_qa_only",
            "context_contract": {"status": "PASS"},
        }
        for alias in ("A", "B", "C")
    ]

    readiness = validate_lung588_real_inputs._build_uat_readiness(rows)

    assert readiness["scope"] == "risk_based_all_available_real_cases"
    assert readiness["policy_id"] == "lung588_risk_based_all_available_cases_v1"
    assert readiness["fixed_minimum_real_case_count"] is None
    assert readiness["observed_real_input_count"] == 3
    assert readiness["required_report_group_review_case_count"] == 3
    assert readiness["ngs_structure_pass_count"] == 3
    assert readiness["ngs_structure_status"] == "PASS"
    assert readiness["pdl1_product_pass_count"] == 0
    assert readiness["pdl1_product_status"] == "BLOCKED"
    assert readiness["verified_case_pdl1_source_count"] == 0
    assert readiness["report_group_reviewed_case_count"] == 0
    assert readiness["formal_uat_status"] == "BLOCKED"
    assert readiness["formal_uat_requirement_met"] is False
    assert {blocker["code"] for blocker in readiness["blockers"]} == {
        "PDL1_PRODUCT_CONTRACT_BLOCKED",
        "PDL1_CASE_SOURCE_NOT_VERIFIED",
        "REPORT_GROUP_UAT_INCOMPLETE",
    }


def test_real_input_validator_can_pass_three_cases_without_fixed_denominator():
    rows = [
        {
            "alias": f"CASE-LUNG-{alias}",
            "auto_detection": {
                "detected": False,
                "project_type": None,
            },
            "targeted_drug_count": 0,
            "biomarker_contract_status": "PASS",
            "pdl1_product_contract_status": "PASS",
            "pdl1_input_provenance": "case_specific_verified_ihc_source",
            "context_contract": {"status": "PASS"},
        }
        for alias in ("A", "B", "C")
    ]
    decisions = {
        f"CASE-LUNG-{alias}": {
            "decision": "pass",
            "reviewer": "报告组审核人",
            "reviewed_at": "2026-07-25",
            "p0_count": 0,
        }
        for alias in ("A", "B", "C")
    }

    readiness = validate_lung588_real_inputs._build_uat_readiness(
        rows,
        report_group_decisions=decisions,
    )

    assert readiness["observed_real_input_count"] == 3
    assert readiness["required_report_group_review_case_count"] == 3
    assert readiness["report_group_reviewed_case_count"] == 3
    assert readiness["report_group_passed_case_count"] == 3
    assert readiness["formal_uat_status"] == "PASS"
    assert readiness["formal_uat_requirement_met"] is True
    assert readiness["blockers"] == []


def test_real_input_validator_blocks_empty_real_case_set():
    readiness = validate_lung588_real_inputs._build_uat_readiness(
        [],
        report_group_decisions={},
    )

    assert readiness["fixed_minimum_real_case_count"] is None
    assert readiness["formal_uat_status"] == "BLOCKED"
    assert {blocker["code"] for blocker in readiness["blockers"]} == {
        "NO_REGISTERED_REAL_CASES"
    }


def test_real_input_validator_keeps_failure_and_p0_as_release_blockers():
    rows = [
        {
            "alias": "CASE-LUNG-A",
            "auto_detection": {
                "detected": False,
                "project_type": None,
            },
            "targeted_drug_count": 0,
            "biomarker_contract_status": "PASS",
            "pdl1_product_contract_status": "PASS",
            "pdl1_input_provenance": "case_specific_verified_ihc_source",
            "context_contract": {"status": "PASS"},
        }
    ]
    decisions = {
        "CASE-LUNG-A": {
            "decision": "fail",
            "reviewer": "报告组审核人",
            "reviewed_at": "2026-07-25",
            "p0_count": 1,
        }
    }

    readiness = validate_lung588_real_inputs._build_uat_readiness(
        rows,
        report_group_decisions=decisions,
    )

    assert readiness["formal_uat_status"] == "BLOCKED"
    assert {blocker["code"] for blocker in readiness["blockers"]} == {
        "REPORT_GROUP_UAT_FAILED",
        "P0_DEFECTS_PRESENT",
    }


def test_lung588_machine_pre_uat_is_traceable_and_does_not_overclaim():
    record = yaml.safe_load(PRE_UAT_RECORD.read_text(encoding="utf-8"))

    assert record["status"] == "partial_not_release_eligible"
    assert record["source_commit"] == ("c23e9f044c55102c0c0d58217fe3efd2f3ccdc88")
    assert record["environment"]["host_alias"] == "iyun129"
    assert record["environment"]["production_switched"] is False
    assert record["environment"]["production_health_after_test"] == "PASS"
    assert record["release_requirements"] == {
        "required_real_case_count": 10,
        "machine_executed_real_case_count": 3,
        "machine_pass_count": 3,
        "machine_observed_pass_rate": 1.0,
        "report_group_reviewed_case_count": 0,
        "formal_uat_requirement_met": False,
        "p0_count_in_observed_cases": 0,
    }

    cases = record["cases"]
    assert [case["alias"] for case in cases] == [
        "CASE-LUNG-A",
        "CASE-LUNG-B",
        "CASE-LUNG-C",
    ]
    assert [case["report_variant_count"] for case in cases] == [7, 8, 9]
    assert all(case["biomarker_contract_status"] == "PASS" for case in cases)
    assert all(case["targeted_drug_runtime_row_count"] == 0 for case in cases)
    assert all(case["report_qa_status"] == "PASS" for case in cases)
    assert all(case["page_count"] == 27 for case in cases)
    assert all(case["blank_page_count"] == 0 for case in cases)
    assert all(case["unexpected_low_content_page_count"] == 0 for case in cases)
    assert all(case["content_failure_count"] == 0 for case in cases)
    assert len(record["remaining_blockers"]) >= 4


def test_lung588_case_a_contract_is_registered_and_deidentified():
    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    contract_path = package.resolve_context_contract_file("case_lung_a")
    contract = yaml.safe_load(contract_path.read_text(encoding="utf-8"))

    assert contract["contract_id"] == "case_lung_a"
    assert contract["panel_id"] == "lung_588_pdl1"
    assert contract["privacy"]["contains_phi"] is False
    assert len(contract["privacy"]["source_excel_sha256"]) == 64
    assert contract["fields"]["total_variants_count"]["equals"] == 7
    assert contract["tables"]["all_variants"]["row_count"] == 7
    assert {
        row["match"]["gene"]["equals"] for row in contract["tables"]["all_variants"]["rows"]
    } == {"TP53", "ESR1", "FLT3", "GNAS", "APC", "IFNGR1", "TSC1"}
    assert all(
        row["expect"]["transcript"]["equals"].startswith("NM_")
        and row["expect"]["chromosome"]["equals"]
        and row["expect"]["exon"]["equals"]
        and row["expect"]["gene_class"]["equals"] in {"Ⅱ类", "Ⅲ类"}
        and row["expect"]["frequency"]["equals"]
        for row in contract["tables"]["all_variants"]["rows"]
    )
    emitted = contract_path.read_text(encoding="utf-8")
    assert "LZ258" not in emitted
    assert "patient_name" not in emitted


def test_lung588_stk11_unsafe_claims_are_exactly_retracted():
    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    genes = yaml.safe_load(
        package.resolve_rule_file("knowledge_coverage").read_text(encoding="utf-8")
    )["reportable_genes"]
    profile = profile_panel_runtime_content(ROOT, package, genes)

    integrity = profile["citation_integrity"]
    assert integrity["unresolved_pmids"] == []
    assert integrity["source_mismatches"] == []

    redactions = profile["knowledge_redactions"]
    assert len(redactions["rows"]) == 4
    assert redactions["unmatched"] == []
    assert all(row["hit_count"] == 1 for row in redactions["rows"])

    provider = build_panel_gene_provider(ROOT, package)
    section = provider.build_gene_knowledge_section(
        gene="STK11",
        c_hgvs="c.999999A>G",
        p_hgvs="p.X999999Y",
        frequency=10,
        mutation_type="Missense",
        has_drug=False,
    )
    analysis = section["mutation_analysis"]
    narrative = section["mutation_narrative"]
    assert "PMID:25980754" not in analysis
    assert "结直肠癌" not in analysis
    assert "可能预示对免疫检查点抑制剂的耐药" not in analysis
    assert "作为肿瘤抑制基因发挥功能" in analysis
    assert narrative
    assert section["fixed_domain_text"] not in narrative
    assert (
        provider._mutation_narrative_from_composed(
            "固定结构域。",
            "固定结构域。",
        )
        == ""
    )
    assert (
        provider._mutation_narrative_from_composed(
            "固定结构域。",
            "固定结构域。\n独立变异叙述。",
        )
        == "独立变异叙述。"
    )

    gate = run_knowledge_release_gate(
        ROOT,
        panel_ids=["lung_588_pdl1"],
    )
    assert gate["status"] == "FAIL"
    codes = {issue["code"] for issue in gate["panels"][0]["issues"]}
    assert "UNRESOLVED_RUNTIME_PMID" not in codes
    assert "RUNTIME_CITATION_SOURCE_MISMATCH" not in codes
    assert "RUNTIME_KNOWLEDGE_REDACTION_UNMATCHED" not in codes
    analysis_gap = next(
        issue
        for issue in gate["panels"][0]["issues"]
        if issue["code"] == "RUNTIME_MUTATION_ANALYSIS_GAP"
    )
    assert analysis_gap["requires_separate_mutation_narrative"] is True
    assert "non-domain mutation narrative" in analysis_gap["message"]
    assert len(analysis_gap["genes"]) == 247


def test_lung588_knowledge_redaction_contract_rejects_replacement_text(
    tmp_path,
):
    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    redactions = load_panel_knowledge_redactions(package)

    assert len(redactions) == 4
    assert all(row["adds_medical_claim"] is False for row in redactions)
    assert all(row["action"] == "remove_exact_literal" for row in redactions)

    raw = yaml.safe_load(
        package.resolve_rule_file("knowledge_redactions").read_text(encoding="utf-8")
    )
    raw["governance"]["replacement_text_allowed"] = True
    invalid = tmp_path / "knowledge_redactions.yaml"
    invalid.write_text(
        yaml.safe_dump(raw, allow_unicode=True, sort_keys=False),
        encoding="utf-8",
    )
    fake_package = SimpleNamespace(
        panel_id="lung_588_pdl1",
        resolve_rule_file=lambda _name: invalid,
    )

    with pytest.raises(ValueError, match="replacement-free"):
        load_panel_knowledge_redactions(fake_package)


def test_lung588_domain_catalog_is_complete_candidate_only_and_sourced():
    catalog = yaml.safe_load(DOMAIN_CANDIDATES.read_text(encoding="utf-8"))
    defaults = catalog["governance"]["defaults"]["gene"]
    rows = catalog["gene_sections"]

    assert catalog["source"]["panel"] == "lung_588_pdl1"
    assert catalog["source"]["activation_mode"] == ("candidate_only_pending_secondary_review")
    assert defaults["review_status"] == "needs_review"
    assert defaults["runtime_eligible"] is False
    assert defaults["secondary_review_status"] == ("pending_report_group_review")
    assert len(rows) == 551
    assert len({row["gene"] for row in rows}) == 551
    assert all(
        row["panels"] == ["lung_588_pdl1"]
        and row["fixed_domain_text"]
        and row["source_refs"]
        and row.get("runtime_eligible") is not True
        for row in rows
    )
    ambiguous = [row for row in rows if row.get("accession_selection")]
    assert len(ambiguous) == 15
    assert {
        row["gene"]
        for row in ambiguous
        if row["accession_selection"]["requires_transcript_product_review"]
    } == {"CDKN2A", "CUX1", "GNAS", "RBM10"}
    assert catalog["drug_sections"] == []

    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    genes = yaml.safe_load(
        package.resolve_rule_file("knowledge_coverage").read_text(encoding="utf-8")
    )["reportable_genes"]
    profile = profile_panel_runtime_content(ROOT, package, genes)
    assert profile["fixed_domain_covered_genes"] == 37
    assert len(profile["missing_fixed_domain_genes"]) == 551
    assert set(profile["missing_fixed_domain_genes"]) == {row["gene"] for row in rows}


def test_lung588_medical_knowledge_queue_is_complete_and_deidentified(tmp_path):
    completed = subprocess.run(
        [
            sys.executable,
            str(ROOT / "scripts" / "analysis" / "23_profile_lung588_medical_knowledge.py"),
            "--project-root",
            str(ROOT),
            "--as-of",
            "2026-07-23",
            "--output-dir",
            str(tmp_path),
        ],
        check=False,
        capture_output=True,
        text=True,
        cwd=ROOT,
    )

    assert completed.returncode == 0, completed.stderr
    inventory = json.loads(
        (tmp_path / "knowledge_depth_inventory.json").read_text(encoding="utf-8")
    )
    assert inventory["denominator"]["total_genes"] == 588
    assert sum(inventory["summary"]["priority_counts"].values()) == 588
    assert inventory["summary"]["citation_source_mismatch_count"] == 0
    assert inventory["summary"]["complete_mutation_analysis_count"] == 341
    assert inventory["summary"]["complete_mutation_narrative_count"] == 341
    assert inventory["summary"]["composed_analysis_without_narrative_count"] == 0
    assert inventory["summary"]["specific_mutation_narrative_count"] == 32
    assert inventory["summary"]["generic_mutation_narrative_count"] == 309
    assert inventory["summary"]["fixed_domain_count"] == 37
    assert inventory["summary"]["fixed_domain_candidate_count"] == 551
    assert inventory["summary"]["fixed_domain_candidate_runtime_eligible_count"] == 0
    assert inventory["summary"]["fixed_domain_candidate_ambiguous_mapping_count"] == 15
    assert inventory["summary"]["fixed_domain_candidate_transcript_product_review_count"] == 4
    assert inventory["summary"]["fixed_domain_runtime_or_candidate_count"] == 588
    assert inventory["summary"]["depth_strata"] == {
        "missing_narrative_and_domain": 247,
        "narrative_present_domain_missing": 304,
        "domain_present_narrative_missing": 0,
        "narrative_and_domain_present": 37,
    }
    manifest = json.loads(
        (tmp_path / "knowledge_review_batch_manifest.json").read_text(encoding="utf-8")
    )
    assert manifest["total_rows"] == 588
    assert manifest["batch_size"] == 25
    assert manifest["batch_count"] == 25
    assert all(
        batch["secondary_review_completed_count"] == 0
        and batch["patient_visible_allowed_count"] == 0
        for batch in manifest["batches"]
    )
    with (tmp_path / "knowledge_review_batches.tsv").open(
        encoding="utf-8-sig",
        newline="",
    ) as handle:
        batch_rows = list(csv.DictReader(handle, delimiter="\t"))
    assert len(batch_rows) == 588
    assert len(list((tmp_path / "review_batches").glob("*.tsv"))) == 25
    stk11 = next(row for row in inventory["rows"] if row["gene"] == "STK11")
    assert stk11["priority"] == "P4"
    assert stk11["citation_source_mismatches"] == []
    assert stk11["fixed_domain_candidate_available"] is False
    gnas = next(row for row in inventory["rows"] if row["gene"] == "GNAS")
    assert gnas["fixed_domain_candidate_available"] is True
    assert (
        gnas["fixed_domain_candidate"]["accession_selection"]["requires_transcript_product_review"]
        is True
    )
    emitted = "\n".join(
        path.read_text(encoding="utf-8-sig")
        for path in tmp_path.iterdir()
        if path.suffix in {".json", ".tsv"}
    )
    assert "LZ258" not in emitted
    assert "patient_name" not in emitted


def test_lung588_p0_event_review_packet_is_event_scoped_and_fail_closed(
    tmp_path,
):
    completed = subprocess.run(
        [
            sys.executable,
            str(ROOT / "scripts" / "analysis" / "24_build_lung588_p0_event_review.py"),
            "--project-root",
            str(ROOT),
            "--reviewed-at",
            "2026-07-23",
            "--output-dir",
            str(tmp_path),
        ],
        check=False,
        capture_output=True,
        text=True,
        cwd=ROOT,
    )

    assert completed.returncode == 0, completed.stderr
    packet = json.loads((tmp_path / "p0_event_review.json").read_text(encoding="utf-8"))
    assert packet["status"] == ("primary_review_complete_secondary_review_pending")
    assert packet["summary"]["review_unit_count"] == 28
    assert packet["summary"]["unit_type_counts"] == {
        "citation_source_mismatch": 1,
        "targeted_drug_candidate": 4,
        "variant_narrative": 23,
    }
    assert packet["summary"]["secondary_review_completed_count"] == 0
    assert packet["summary"]["patient_visible_part3_allowed_count"] == 0
    assert packet["summary"]["patient_visible_drug_allowed_count"] == 0
    assert packet["summary"]["event_narrative_candidate_count"] == 23
    assert packet["summary"]["event_narrative_candidate_gene_count"] == 19
    assert packet["summary"]["event_narrative_candidate_runtime_eligible_count"] == 0
    assert packet["scope"]["event_narrative_candidates_are_runtime_content"] is False
    assert packet["scope"]["case_aliases"] == [
        "CASE-LUNG-A",
        "CASE-LUNG-B",
        "CASE-LUNG-C",
    ]
    assert all(unit["runtime_eligible"] is False for unit in packet["units"])
    assert all(
        unit["primary_review"]["status"] == "completed_ai_assisted_triage"
        for unit in packet["units"]
    )
    assert all(
        unit["secondary_review"]["status"] == "pending_report_group_review"
        for unit in packet["units"]
    )
    candidates = [
        unit for unit in packet["units"] if unit["unit_type"] == "targeted_drug_candidate"
    ]
    assert len(candidates) == 4
    assert all(unit["source_scope_review"] for unit in candidates)
    assert all(
        unit["source_scope_review"]["candidate_id"] == unit["review_unit_id"] for unit in candidates
    )
    assert {unit["primary_review"]["decision"] for unit in candidates} == {
        "eligible_to_remain_candidate_pending_secondary_review",
        ("eligible_to_remain_candidate_with_indirect_exact_event_chain_pending_secondary_review"),
        "hold_pending_official_china_label_and_secondary_review",
    }
    assert all(
        unit["source_scope_review"]["runtime_eligible"] is False
        and unit["source_scope_review"]["report_text_allowed"] is False
        for unit in candidates
    )

    variants = [unit for unit in packet["units"] if unit["unit_type"] == "variant_narrative"]
    narrative_candidates = [
        unit["candidate_narrative_review"]
        for unit in variants
        if unit["candidate_narrative_review"]
    ]
    assert len(narrative_candidates) == 23
    assert len(variants) == 23
    assert all(unit["candidate_narrative_review"] for unit in variants)
    assert all(
        candidate["review_status"] == "needs_review"
        and candidate["runtime_eligible"] is False
        and candidate["report_text_allowed"] is False
        and candidate["patient_visible"] is False
        and candidate["secondary_review_status"] == "pending_report_group_review"
        and candidate["runtime_selector_contract"]["current_overlay_matches_transcript"] is False
        for candidate in narrative_candidates
    )
    atm = next(
        unit for unit in variants if unit["gene"] == "ATM" and unit["c_hgvs"] == "c.1236-2A>T"
    )
    assert [row["case_alias"] for row in atm["case_observations"]] == [
        "CASE-LUNG-B",
        "CASE-LUNG-C",
    ]
    assert [row["frequency"] for row in atm["case_observations"]] == [
        "1.67",
        "0.92",
    ]
    assert atm["transcript"] == "NM_000051.4"
    assert atm["chromosome"] == "11"
    assert atm["exon"] == "10"
    braf_d594g = next(
        unit for unit in variants if unit["gene"] == "BRAF" and unit["p_hgvs"] == "p.D594G"
    )
    assert braf_d594g["explicit_non_promotion"] is True
    assert braf_d594g["patient_visible_drug_conclusion_allowed"] is False

    mismatch = next(
        unit for unit in packet["units"] if unit["unit_type"] == "citation_source_mismatch"
    )
    assert mismatch["identifier"] == "PMID:25980754"
    assert mismatch["suggested_replacement_identifier"] == "PMID:29773717"
    assert mismatch["runtime_claim_retracted"] is True
    assert mismatch["runtime_retraction_ids"] == [
        "lung588_stk11_remove_unsupported_pmid_25980754_claim"
    ]

    emitted = "\n".join(
        path.read_text(encoding="utf-8-sig")
        for path in tmp_path.iterdir()
        if path.suffix in {".json", ".tsv"}
    )
    assert "LZ258" not in emitted
    assert "patient_name" not in emitted


def test_lung588_p0_event_narrative_candidates_are_exact_sourced_and_hidden():
    contract = yaml.safe_load(EVENT_NARRATIVE_CANDIDATES.read_text(encoding="utf-8"))
    source = contract["source"]
    governance = contract["governance"]
    defaults = governance["defaults"]["gene"]
    rows = contract["gene_sections"]

    assert source["panel"] == "lung_588_pdl1"
    assert source["activation_mode"] == ("candidate_only_pending_secondary_review")
    assert defaults["review_status"] == "needs_review"
    assert defaults["runtime_eligible"] is False
    assert defaults["report_text_allowed"] is False
    assert defaults["patient_visible"] is False
    assert defaults["secondary_review_status"] == ("pending_report_group_review")
    assert governance["runtime_selector_contract"] == {
        "current_overlay_matches_transcript": False,
        "disposition": "promotion_blocked_until_transcript_is_enforced",
        "reason": (
            "The current reviewed-overlay runtime key uses gene, c_hgvs and "
            "p_hgvs. Transcript is preserved below for review identity but "
            "is not yet enforced by the runtime loader.\n"
        ),
    }
    assert contract["drug_sections"] == []
    assert len(rows) == 10
    assert len({row["candidate_id"] for row in rows}) == 10
    assert len({row["gene"] for row in rows}) == 9

    expected_events = {
        ("BRIP1", "NM_032043.3", "c.2142G>A", "p.W714*"),
        ("ERBB2", "NM_004448.4", "c.1979G>A", "p.G660D"),
        ("ESR1", "NM_000125.4", "c.1242G>T", "p.Q414H"),
        ("FLT3", "NM_004119.3", "c.1159C>T", "p.R387*"),
        ("GNAS", "NM_000516.7", "c.695G>T", "p.R232L"),
        (
            "IFNGR1",
            "NM_000416.3",
            "c.1132_1133del",
            "p.S378Ffs*6",
        ),
        ("MSH3", "NM_002439.5", "c.2905C>T", "p.Q969*"),
        ("TSC1", "NM_000368.5", "c.474del", "p.F158Lfs*9"),
        ("TSC1", "NM_000368.5", "c.2074C>T", "p.R692*"),
        ("TSC2", "NM_000548.5", "c.2647C>T", "p.Q883*"),
    }
    assert {
        (
            row["gene"],
            row["transcript"],
            row["c_hgvs"],
            row["p_hgvs"],
        )
        for row in rows
    } == expected_events

    ncbi_gene_ids = {
        "BRIP1": "GeneID:83990",
        "ERBB2": "GeneID:2064",
        "ESR1": "GeneID:2099",
        "FLT3": "GeneID:2322",
        "GNAS": "GeneID:2778",
        "IFNGR1": "GeneID:3459",
        "MSH3": "GeneID:4437",
        "TSC1": "GeneID:7248",
        "TSC2": "GeneID:7249",
    }
    for row in rows:
        assert row["panels"] == ["lung_588_pdl1"]
        assert row["review_status"] == "needs_review"
        assert row["runtime_eligible"] is False
        assert row["report_text_allowed"] is False
        assert row["patient_visible"] is False
        assert row["secondary_review_status"] == ("pending_report_group_review")
        assert row["intro"]
        assert row["mutation_analysis"]
        assert row["source_refs"] == [
            {
                "type": "ncbi_gene",
                "authority": "NCBI",
                "id": ncbi_gene_ids[row["gene"]],
                "url": (
                    "https://www.ncbi.nlm.nih.gov/gene/"
                    f"{ncbi_gene_ids[row['gene']].split(':', 1)[1]}"
                ),
                "supports": "gene_identity_and_function_only",
            }
        ]
        assert all(
            row["evidence_boundaries"][field] is False
            for field in (
                "treatment_inference_allowed",
                "immune_inference_allowed",
                "prognostic_inference_allowed",
                "hereditary_inference_allowed",
            )
        )

    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    provider = build_panel_gene_provider(ROOT, package)
    mutation_types = {
        "frameshift": "Frameshift",
        "stop_gained": "Nonsense",
        "missense": "Missense",
    }
    for row in rows:
        section = provider.build_gene_knowledge_section(
            gene=row["gene"],
            c_hgvs=row["c_hgvs"],
            p_hgvs=row["p_hgvs"],
            frequency=1.0,
            mutation_type=mutation_types[row["variant_kind"]],
            has_drug=False,
        )
        assert is_generic_mutation_analysis(
            row["gene"],
            section["mutation_narrative"],
        )
        assert section["mutation_narrative"] != row["mutation_analysis"]


def test_lung588_cross_cancer_narrative_candidates_are_exact_and_hidden():
    contract = yaml.safe_load(CROSS_CANCER_NARRATIVE_CANDIDATES.read_text(encoding="utf-8"))
    source = contract["source"]
    governance = contract["governance"]
    defaults = governance["defaults"]["gene"]
    rows = contract["gene_sections"]

    assert source["panel"] == "lung_588_pdl1"
    assert source["activation_mode"] == ("candidate_only_pending_secondary_review")
    assert defaults["review_status"] == "needs_review"
    assert defaults["runtime_eligible"] is False
    assert defaults["report_text_allowed"] is False
    assert defaults["patient_visible"] is False
    assert defaults["secondary_review_status"] == ("pending_report_group_review")
    assert governance["runtime_selector_contract"] == {
        "current_overlay_matches_transcript": False,
        "disposition": "promotion_blocked_until_transcript_is_enforced",
        "reason": (
            "The current reviewed-overlay runtime key uses gene, c_hgvs and "
            "p_hgvs. Transcript is preserved below for review identity but "
            "is not yet enforced by the runtime loader.\n"
        ),
    }
    assert contract["drug_sections"] == []
    assert len(rows) == 13
    assert len({row["candidate_id"] for row in rows}) == 13
    assert len({row["gene"] for row in rows}) == 10

    expected_events = {
        ("APC", "NM_000038.6", "c.1269G>A", "p.W423*"),
        ("ATM", "NM_000051.4", "c.1236-2A>T", ""),
        ("BRAF", "NM_004333.6", "c.1781A>G", "p.D594G"),
        ("BRAF", "NM_004333.6", "c.1799T>A", "p.V600E"),
        ("BRCA2", "NM_000059.4", "c.7007G>A", "p.R2336H"),
        ("MLH1", "NM_000249.4", "c.790+1G>A", ""),
        ("PIK3CA", "NM_006218.4", "c.3197C>T", "p.A1066V"),
        ("PMS2", "NM_000535.7", "c.59G>A", "p.R20Q"),
        ("PTEN", "NM_000314.8", "c.802-2A>T", ""),
        (
            "SMAD4",
            "NM_005359.6",
            "c.1389_1396del",
            "p.A464Rfs*27",
        ),
        ("TP53", "NM_000546.6", "c.578A>T", "p.H193L"),
        ("TP53", "NM_000546.6", "c.707A>G", "p.Y236C"),
        ("TP53", "NM_000546.6", "c.734G>A", "p.G245D"),
    }
    assert {
        (
            row["gene"],
            row["transcript"],
            row["c_hgvs"],
            row["p_hgvs"],
        )
        for row in rows
    } == expected_events
    assert {row["gene"]: row["source_refs"][0]["id"] for row in rows} == {
        "APC": "GeneID:324",
        "ATM": "GeneID:472",
        "BRAF": "GeneID:673",
        "BRCA2": "GeneID:675",
        "MLH1": "GeneID:4292",
        "PIK3CA": "GeneID:5290",
        "PMS2": "GeneID:5395",
        "PTEN": "GeneID:5728",
        "SMAD4": "GeneID:4089",
        "TP53": "GeneID:7157",
    }
    for row in rows:
        assert row["panels"] == ["lung_588_pdl1"]
        assert row["review_status"] == "needs_review"
        assert row["runtime_eligible"] is False
        assert row["report_text_allowed"] is False
        assert row["patient_visible"] is False
        assert row["secondary_review_status"] == ("pending_report_group_review")
        assert row["intro"]
        assert row["mutation_analysis"]
        assert row["superseded_risk"]
        assert row["source_refs"] == [
            {
                "type": "ncbi_gene",
                "authority": "NCBI",
                "id": row["source_refs"][0]["id"],
                "url": (
                    "https://www.ncbi.nlm.nih.gov/gene/"
                    f"{row['source_refs'][0]['id'].split(':', 1)[1]}"
                ),
                "supports": "gene_identity_and_function_only",
            }
        ]
        assert all(
            row["evidence_boundaries"][field] is False
            for field in (
                "treatment_inference_allowed",
                "immune_inference_allowed",
                "prognostic_inference_allowed",
                "hereditary_inference_allowed",
            )
        )

    splice_rows = [row for row in rows if row["variant_kind"] == "splice_region_or_site"]
    assert {(row["gene"], row["c_hgvs"], row["p_hgvs"]) for row in splice_rows} == {
        ("ATM", "c.1236-2A>T", ""),
        ("MLH1", "c.790+1G>A", ""),
        ("PTEN", "c.802-2A>T", ""),
    }

    package = load_panel_package("lung_588_pdl1", project_root=ROOT)
    provider = build_panel_gene_provider(ROOT, package)
    mutation_types = {
        "frameshift": "Frameshift",
        "stop_gained": "Nonsense",
        "missense": "Missense",
        "splice_region_or_site": "Splice",
    }
    current_runtime_narratives = {}
    for row in rows:
        section = provider.build_gene_knowledge_section(
            gene=row["gene"],
            c_hgvs=row["c_hgvs"],
            p_hgvs=row["p_hgvs"],
            frequency=1.0,
            mutation_type=mutation_types[row["variant_kind"]],
            has_drug=False,
        )
        current_runtime_narratives[(row["gene"], row["c_hgvs"], row["p_hgvs"])] = section[
            "mutation_narrative"
        ]
        assert section["mutation_narrative"] != row["mutation_analysis"]

    assert "结直肠癌" in current_runtime_narratives[("APC", "c.1269G>A", "p.W423*")]
    assert "肠癌预后" in current_runtime_narratives[("BRAF", "c.1799T>A", "p.V600E")]
    assert "结直肠癌" in current_runtime_narratives[("TP53", "c.734G>A", "p.G245D")]
